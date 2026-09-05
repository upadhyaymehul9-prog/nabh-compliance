# -*- coding: utf-8 -*-
"""
generate_hco_aac_rewrites.py
Generates HCO AAC.1-8 v2 rewrite reference DOCX files.

All-8 fixes: "Policy on [X]" title; no subtitle; HN only guillemet;
generic Rule-2.5 exclusion line; References last; 5-col Traceability.
AAC.6/7/8 extra: no md5 hash; no "Spell out:" in Scope; version "2.0".
"""
import os
from docx import Document

HN = "«Hospital Name»"   # «Hospital Name»
OUT = "policies/build/rewrite_reference"
os.makedirs(OUT, exist_ok=True)


def h(doc, lv, txt):
    return doc.add_paragraph(txt, style={0:"Title",1:"Heading 1",2:"Heading 2"}[lv])

def p(doc, txt=""):
    return doc.add_paragraph(txt, style="Normal")

def ln(doc, txt):
    return doc.add_paragraph(txt, style="List Number")

def lb(doc, txt):
    return doc.add_paragraph(txt, style="List Bullet")

def tbl(doc, rows, cols):
    t = doc.add_table(rows=rows, cols=cols)
    try: t.style = "Table Grid"
    except KeyError: pass
    return t

def doc_ctrl(doc, no, prep, appr="Medical Superintendent"):
    dc = tbl(doc, 6, 4)
    for ri,(a,b,c,d) in enumerate([
        ("Document No.", no, "Version", "2.0"),
        ("Issue No.", "01", "Review due", "One year from implementation"),
        ("Date created", "________", "Date of implementation", "________"),
    ]):
        dc.cell(ri,0).text=a; dc.cell(ri,1).text=b
        dc.cell(ri,2).text=c; dc.cell(ri,3).text=d
    for ri,(lbl,txt) in enumerate([
        ("Prepared by", f"{prep}  Name: ________  Signature: ________"),
        ("Reviewed by", "Quality Coordinator  Name: ________  Signature: ________"),
        ("Approved by", f"{appr}  Name: ________  Signature: ________"),
    ], start=3):
        dc.cell(ri,0).text=lbl
        c1=dc.cell(ri,1); c1.text=txt; c1.merge(dc.cell(ri,3))

def gov_tbl(doc, rows):
    t=tbl(doc,len(rows)+1,2)
    t.cell(0,0).text="Role"; t.cell(0,1).text="Responsibility"
    for ri,(role,resp) in enumerate(rows,1):
        t.cell(ri,0).text=role; t.cell(ri,1).text=resp

def abbrev_tbl(doc, rows):
    t=tbl(doc,len(rows)+1,2)
    t.cell(0,0).text="Abbreviation"; t.cell(0,1).text="Meaning"
    for ri,(a,m) in enumerate(rows,1):
        t.cell(ri,0).text=a; t.cell(ri,1).text=m

def trace_tbl(doc, rows):
    t=tbl(doc,len(rows)+1,5)
    for ci,h_ in enumerate(("OE","Level","Requirement","Where addressed","Responsible")):
        t.cell(0,ci).text=h_
    for ri,row in enumerate(rows,1):
        for ci,v in enumerate(row): t.cell(ri,ci).text=v

def disclaimer(doc, p2=None):
    p(doc, f"This document is a template prepared for the guidance of {HN} and must be reviewed, "
           f"adapted and formally approved by {HN} before use. Every entry marked ________ "
           f"must be completed before the document is issued.")
    p(doc, p2 or (
        f"The requirements in this document are accreditation requirements of the NABH Accreditation "
        f"Standards for Hospitals, 6th Edition, not duties under a named Act of Parliament. This "
        f"policy does not import the Consumer Protection Act, 2019, the Clinical Establishments Act, "
        f"2010, or the Mental Healthcare Act, 2017 as a checklist. Statutory duties that arise under "
        f"other documents of {HN} remain those documents. {HN} is responsible for verifying any "
        f"statutory duty that applies to it; this document does not constitute legal advice."))
    p(doc, f"{HN} remains responsible for verifying that it is current and consistent with the "
           f"edition of the accreditation standard against which it is being assessed. The clinical "
           f"and technical content reflects recognised national and international guidance current "
           f"at the date of preparation.")
    p(doc, "This document is not issued by, endorsed by, or affiliated with NABH, the World Health "
           "Organization, the National Centre for Disease Control, the Food Safety and Standards "
           "Authority of India, any Pollution Control Board, or any other body named in it. Wording "
           "is original; no text has been reproduced from the standards, rules or guidelines referenced.")

def save_and_verify(doc, fname):
    import sys
    out = sys.stdout
    def pr(s):
        try: out.write(s+"\n")
        except UnicodeEncodeError: out.write(s.encode("ascii","replace").decode()+"\n")
    pr(f"\n=== {fname} ===")
    for i,para in enumerate(doc.paragraphs[:50]):
        sn = para.style.name if para.style else "(None)"
        pr(f"{i:3d}  {sn!r:30s}  {para.text[:60]!r}")
    pr(f"  Total paras: {len(doc.paragraphs)}")
    path=os.path.join(OUT,fname)
    doc.save(path)
    print(f"  Saved: {path}")


# ══════════════════════════════════════════════════════════════════════════════
# AAC.1 — Defined and Displayed Healthcare Services
# ══════════════════════════════════════════════════════════════════════════════
def gen_aac1():
    doc=Document()
    h(doc,0,"Policy on Defined and Displayed Healthcare Services")
    p(doc,HN)
    h(doc,1,"Document control")
    doc_ctrl(doc,"HCO/AAC/POL/01","Quality Coordinator")
    p(doc,"A blank marked ________ must be completed before issue.")
    h(doc,1,"Statement of intent")
    p(doc,"Patients, families and referrers must know what this hospital provides — services defined "
          "with the community in mind, staffed and equipped, scoped by department, and displayed where "
          "people can read them.")
    h(doc,1,"1. Purpose")
    p(doc,f"This policy says how {HN} defines the healthcare services it provides in consonance with "
          f"community needs; ensures each defined clinical service has diagnostic and treatment "
          f"capability with suitably qualified personnel for out-patient, in-patient, day-care and "
          f"emergency cover; defines the scope of clinical services of each department; and prominently "
          f"displays those defined clinical services.")
    p(doc,"The chapter intent is that patients are informed of the services provided, that the scope of "
          "each healthcare service including diagnostic and therapeutic services is well defined and "
          "available to patients and families, and that only patients the organisation can care for are admitted.")
    p(doc,"This policy does not cover patient registration, admission, transfer, referral, laboratory "
          "services, or imaging services — those are covered in other hospital policies. The other AAC "
          "standards have their own policies too.")
    h(doc,1,"2. Scope")
    p(doc,f"This policy applies to every clinical and diagnostic department at {HN}. It binds the "
          f"Medical Superintendent, Heads of departments (broad speciality, super speciality and "
          f"diagnostic), the Quality Coordinator, registration and front-office staff, and clinical "
          f"staff who deliver or describe services.")
    p(doc,"It covers AAC.1.a–d: defining services with community needs; diagnostic and treatment "
          "capability with qualified personnel including day-care and emergency cover; department scope "
          "(including outsourced clinical and diagnostic services); and permanent bi-lingual display.")
    p(doc,"Boundaries:")
    lb(doc,"AAC.2 owns registration, admission and acceptance against the services defined here.")
    lb(doc,"AAC.6 and AAC.8 own laboratory and imaging service detail. This policy owns that those "
           "services appear in the department scope and display.")
    lb(doc,"PRE.1 owns the patient-rights display. This policy owns the clinical-services display.")
    lb(doc,"Outsourced services are named and known under this policy; the contract and quality "
           "oversight of those vendors sit with the owning clinical or diagnostic policy.")
    h(doc,1,"3. Policy standards")
    p(doc,f"{HN} defines the healthcare services it provides, in consonance with the needs of the "
          f"community it serves. Each defined clinical service has diagnostic and treatment services "
          f"with suitably qualified medical, nursing and paramedical personnel who provide out-patient, "
          f"in-patient, day-care and emergency cover. The scope of clinical services of each department "
          f"is defined. The defined clinical services are permanently and prominently displayed in at "
          f"least two languages.")
    p(doc,f"{HN} does not claim a clinical service it cannot staff, diagnose for or treat, and does "
          f"not leave the display out of date or unreadable.")
    h(doc,1,"4. Non-negotiable rules")
    ln(doc,"Do not claim a clinical service the hospital cannot provide with qualified personnel and "
           "diagnostic and treatment capability for the care settings that service covers.")
    ln(doc,"Do not operate a department whose clinical scope has not been defined and approved by the "
           "Medical Superintendent.")
    ln(doc,"Do not omit documentation of clinical or diagnostic services that are outsourced; staff "
           "must know what is outsourced.")
    ln(doc,"Do not remove the clinical-services display from a location visible to patients and "
           "visitors, allow it to become temporary when a permanent display is required, or leave it "
           "in only one language.")
    ln(doc,f"Staff who find the displayed services differ from what is actually provided report it the "
            f"same shift to the Quality Coordinator or the Medical Superintendent.")
    h(doc,1,"5. What we do")
    h(doc,2,"5.1 Define healthcare services in consonance with community needs")
    p(doc,f"The Medical Superintendent, with Heads of departments, defines the healthcare services "
          f"{HN} provides. Senior management owns the definition. Community needs are considered when "
          f"planning new services — captured through patient and family feedback, referral patterns, "
          f"district disease-burden data, and changing disease patterns. Starting a new service on the "
          f"hospital's own judgment is allowed; community need still informs the decision.")
    p(doc,f"The defined services are recorded in the service directory of {HN}. A service not in the "
          f"directory is not claimed. The directory is reviewed annually and whenever a service is "
          f"added, suspended or withdrawn.")
    h(doc,2,"5.2 Diagnostic and treatment capability with qualified personnel")
    p(doc,"Each defined clinical service (broad speciality and super speciality) has:")
    lb(doc,"diagnostic and treatment services appropriate to that clinical service;")
    lb(doc,"suitably qualified medical, nursing and paramedical staff for the patient's clinical needs;")
    lb(doc,"out-patient services, in-patient services, day-care where the service uses day-care, and "
           "emergency cover by the consultant(s).")
    p(doc,f"Infrastructure for diagnostics and treatment follows regulatory requirements and professional-"
          f"body guidance where available. The Medical Superintendent ensures personnel qualifications "
          f"are verified at appointment and kept current. Where a service limitation exists, it is "
          f"written in the department scope (section 5.3) and reflected in the display (section 5.4).")
    h(doc,2,"5.3 Define scope of clinical services of each department")
    p(doc,"Each department — super speciality, broad speciality and diagnostic — has a written scope. "
          "Scope may be by inclusion or by exclusion relative to services practised in the department. "
          "Example content for a nephrology department could include biopsy, shunts, fistulas, "
          "haemodialysis and CAPD where those are practised.")
    p(doc,f"All clinical and diagnostic outsourced services are documented. Staff know what is "
          f"outsourced. That information is also available to patients through the website, display "
          f"boards and department brochures as the hospital chooses.")
    p(doc,f"The Medical Superintendent approves each scope. The Quality Coordinator holds the current "
          f"set. Scopes are reviewed annually and on any material change in capability.")
    h(doc,2,"5.4 Prominent permanent bi-lingual display of defined clinical services")
    p(doc,f"The display states the names of clinical and diagnostic departments of {HN}. It is "
          f"permanent (board, citizen's charter or equivalent; electronic display is allowed as a "
          f"supplement or alternative where it remains permanently visible). It is placed where "
          f"patients and visitors can see it — at least at the main entrance and the registration area.")
    p(doc,"The display is at least bi-lingual: the State language or the language spoken by the "
          "majority of people in the catchment, and English. Brochures, standees and the website may "
          "supplement the permanent display; they do not replace it.")
    p(doc,f"When a service is added or withdrawn, the display is updated within seven working days. "
          f"The Quality Coordinator checks the display quarterly against the service directory and "
          f"records any mismatch for correction.")
    h(doc,1,"6. Governance and responsibility")
    gov_tbl(doc,[
        ("Medical Superintendent","Accountable that healthcare services are defined, resourced and "
         "displayed. Approves the service directory and each department clinical scope."),
        ("Heads of departments","Define and maintain the clinical scope of their department, including "
         "outsourced elements. Ensure diagnostic and treatment capability and qualified personnel for "
         "out-patient, in-patient, day-care and emergency cover as applicable."),
        ("Quality Coordinator","Holds the service directory and department scope statements. Audits "
         "the permanent display quarterly."),
        ("Registration / front-office staff","Direct patients only to services listed in the current "
         "directory and display."),
        ("All clinical staff","Report any mismatch between displayed and actual services the same shift."),
    ])
    h(doc,1,"7. Quality monitoring (RCA → CAPA)")
    p(doc,"The Quality Coordinator audits this policy quarterly. What is monitored each quarter:")
    lb(doc,"Service directory is current and matches what is actually provided.")
    lb(doc,"Each listed clinical service has diagnostic and treatment capability and qualified personnel "
           "recorded, including day-care and emergency cover where applicable.")
    lb(doc,"Department scope statements exist, are approved, and document outsourced clinical and "
           "diagnostic services.")
    lb(doc,"Permanent bi-lingual display is legible, current and matches the service directory at "
           "every display point.")
    p(doc,"Root-cause analysis is required when the same service-display mismatch recurs within six months.")
    p(doc,"This policy is reviewed annually, and sooner when services are added, suspended or withdrawn.")
    h(doc,1,"8. Training and staff acknowledgement")
    p(doc,"All staff are informed of this policy at induction and once a year after that. Training "
          "covers the service directory, department scopes, outsourced services, how to check the "
          "display, and how to report a mismatch.")
    p(doc,"Staff acknowledgement")
    p(doc,f"I have read the Policy on Defined and Displayed Healthcare Services of {HN}. I will "
          f"follow the processes described.")
    p(doc,"Name: ___________________________    Designation: ___________________________")
    p(doc,"Department / floor: ____________________    Date: ____________")
    p(doc,"Signature: ___________________________")
    p(doc,"(One row per staff member. The Quality Coordinator holds signed acknowledgements with "
          "the induction record.)")
    h(doc,1,"9. Distribution")
    p(doc,f"Official master copy: office of the Medical Superintendent, {HN}, with the Quality Coordinator.")
    p(doc,"Copies issued to: registration; every clinical and diagnostic department; out-patient; "
          "emergency; day-care; nursing administration.")
    p(doc,"The current version is available to all staff at the front-office policy file and, if the "
          "hospital keeps an intranet, at staff intranet / policies.")
    p(doc,"When a new version is issued, take old copies out of use.")
    h(doc,1,"10. Abbreviations")
    abbrev_tbl(doc,[
        ("AAC","Access, Assessment and Continuity of Care (NABH Hospitals chapter)"),
        ("CAPA","corrective and preventive action"),
        ("CAPD","continuous ambulatory peritoneal dialysis"),
        ("HCO","Hospital (Full Accreditation programme under NABH Hospitals 6th Edition)"),
        ("NABH","National Accreditation Board for Hospitals and Healthcare Providers"),
        ("OE","objective element"),
        ("OPD","out-patient department"),
    ])
    h(doc,1,"11. Traceability to NABH HCO Full Accreditation 6th Edition AAC.1")
    p(doc,"This table is an index. It is not how the policy is organised. An asterisk in the Level "
          "column means documentation of the process is required.")
    trace_tbl(doc,[
        ("AAC.1.a","Commitment",
         "The healthcare services being provided are defined and are in consonance with the needs of the community.",
         "Section 3; 5.1","Medical Superintendent (approve); Heads of departments (define); Quality Coordinator (hold directory)"),
        ("AAC.1.b","Commitment",
         "Each defined clinical service shall have diagnostic and treatment services with suitably qualified personnel who provide out-patient, in-patient, daycare and emergency cover.",
         "Section 3; 5.2","Heads of departments (ensure capability); Medical Superintendent (verify qualifications)"),
        ("AAC.1.c","Commitment*",
         "Scope of the clinical services of each department is defined.",
         "Section 3; 5.3","Heads of departments (write scope); Medical Superintendent (approve); Quality Coordinator (hold)"),
        ("AAC.1.d","Commitment",
         "The organisation's defined clinical services are prominently displayed.",
         "Section 3; 5.4","Quality Coordinator (maintain display); registration staff (direct patients per directory)"),
    ])
    h(doc,1,"12. Required Records / Evidence Checklist")
    p(doc,"Records the hospital holds under this policy, listed by objective element.")
    h(doc,2,"AAC.1.a — Healthcare services defined in consonance with community needs")
    lb(doc,"Service directory listing every healthcare service provided.")
    lb(doc,"Evidence that community needs informed service planning (feedback, referral, disease-burden or similar).")
    lb(doc,"Minutes or record of annual review of the service directory.")
    h(doc,2,"AAC.1.b — Diagnostic and treatment capability with qualified personnel")
    lb(doc,"Service directory entry for each clinical service showing diagnostic capability, treatment capability and personnel.")
    lb(doc,"Personnel qualification and verification records for medical, nursing and paramedical staff per service.")
    lb(doc,"Record of out-patient, in-patient, day-care and emergency cover arrangements for each defined clinical service.")
    lb(doc,"Record of any service limitation and the defined referral pathway.")
    h(doc,2,"AAC.1.c — Scope of clinical services of each department defined")
    lb(doc,"Written clinical scope statement for each department (super speciality, broad speciality and diagnostic), approved and dated.")
    lb(doc,"Documented list of clinical and diagnostic outsourced services known to staff.")
    lb(doc,"Record of annual or change-triggered scope review.")
    lb(doc,"Brochure or equivalent detailing department scope where the hospital uses one.")
    h(doc,2,"AAC.1.d — Clinical services prominently displayed")
    lb(doc,"Photographs or records of the permanent display at each mandated location.")
    lb(doc,"Evidence the display is at least bi-lingual (State/majority language and English).")
    lb(doc,"Quarterly audit log comparing display against the service directory.")
    lb(doc,"Record of display update when services changed, with date of update.")
    h(doc,1,"13. References")
    lb(doc,"National Accreditation Board for Hospitals and Healthcare Providers (NABH), Guidebook to "
           "Accreditation Standards for Hospitals, 6th Edition — Access, Assessment and Continuity of "
           "Care (AAC), standard AAC.1.")
    lb(doc,f"Internal documents of {HN}: service directory; department clinical scope statements; "
            f"outsourced-services list; personnel qualification records; display maintenance log.")
    h(doc,1,"Disclaimer")
    disclaimer(doc)
    save_and_verify(doc,"HCO_AAC_1_v2_REWRITE_DRAFT.docx")


# ══════════════════════════════════════════════════════════════════════════════
# AAC.2 — Registration and Admission
# ══════════════════════════════════════════════════════════════════════════════
def gen_aac2():
    doc=Document()
    h(doc,0,"Policy on Registration and Admission")
    p(doc,HN)
    h(doc,1,"Document control")
    doc_ctrl(doc,"HCO/AAC/POL/02","Registration In-Charge")
    p(doc,"A blank marked ________ must be completed before issue.")
    h(doc,1,"Statement of intent")
    p(doc,"Every patient who is assessed is registered; every admission is authorised by a doctor; "
          "a unique number follows the patient; beds and clinical priority are managed so that the "
          "hospital accepts only what it can care for.")
    h(doc,1,"1. Purpose")
    p(doc,f"This policy says how {HN} uses written guidance to register and admit patients (including "
          f"unidentified patients); generates a unique identification number at the end of registration; "
          f"accepts patients only when it can provide the required service (with life-saving treatment "
          f"first in emergency); manages non-availability of beds; and prioritises access according "
          f"to clinical need.")
    p(doc,"The chapter intent is that only those patients who can be cared for by the organisation "
          "are admitted, and that emergency patients receive life-stabilising treatment and are then "
          "either admitted or transferred appropriately.")
    p(doc,"This policy does not cover which services are offered, transfer-out when services are "
          "unavailable, or specific informed-consent procedures for invasive procedures — those are "
          "covered in other hospital policies. The other AAC standards have their own policies too.")
    h(doc,1,"2. Scope")
    p(doc,f"This policy applies at every registration and admission point at {HN}: front office, "
          f"out-patient registration, emergency, day-care and in-patient admission desks. It binds "
          f"registration and front-office staff, treating doctors who authorise admission, nurses who "
          f"receive admitted patients, the Medical Superintendent, department heads, and the Quality Coordinator.")
    p(doc,"It covers AAC.2.a–e for out-patients, day-care, in-patients and emergency patients.")
    p(doc,"Boundaries:")
    lb(doc,"AAC.1 owns the service directory used to decide acceptance.")
    lb(doc,"AAC.3 owns transfer-out when beds or services are unavailable beyond temporary holding.")
    lb(doc,"PRE.3 (informed consent) owns the method for specific and invasive consent. This policy "
           "owns that general consent for treatment is obtained at entry and that its scope is defined "
           "and explained; general consent must not include invasive procedures that need specific consent.")
    lb(doc,"Billing owns tariffs; this policy owns explaining financial implications when the desired "
           "bed category is unavailable.")
    lb(doc,"Medical Records owns the unique-number master index in coordination with registration.")
    h(doc,1,"3. Policy standards")
    p(doc,f"{HN} uses written guidance for registering and admitting patients, including unidentified "
          f"patients. All patients assessed in the hospital are registered. Every admission is authorised "
          f"by a doctor. A unique identification number is generated at the end of registration and "
          f"used across the organisation. Patients are accepted only if the organisation can provide "
          f"the required service; in emergency, life-saving treatment is started before any acceptance "
          f"decision. Written guidance manages non-availability of beds. Access is prioritised "
          f"according to clinical need.")
    p(doc,f"{HN} does not admit without doctor authorisation, does not leave an assessed patient "
          f"unregistered, and does not refuse life-saving emergency treatment while deciding acceptance.")
    h(doc,1,"4. Non-negotiable rules")
    ln(doc,"Do not leave an assessed patient unregistered — including unidentified patients, who are "
           "registered under the unidentified-patient process.")
    ln(doc,"Do not admit a patient without authorisation by a doctor.")
    ln(doc,"Do not complete registration without generating the unique identification number at the "
           "end of registration, and do not give a second unique number to the same patient.")
    ln(doc,"Do not accept a non-emergency patient for a service the hospital cannot provide; in "
           "emergency, start life-saving treatment before deciding acceptance.")
    ln(doc,"Do not hold a patient on a temporary bed beyond the defined time without a decision to "
           "transfer or place; explain financial implications when the desired bed category is unavailable.")
    ln(doc,"Do not ignore clinical priority in out-patient or diagnostic queues when a patient needs "
           "earlier response.")
    ln(doc,f"Staff who find registration, unique-number or bed-management rules not followed report "
            f"it the same shift to the Medical Superintendent or the Quality Coordinator.")
    h(doc,1,"5. What we do")
    h(doc,2,"5.1 Written guidance for registering and admitting patients")
    p(doc,f"{HN} keeps written guidance for registration and admission covering out-patients, day-care, "
          f"in-patients and emergency patients. The guidance includes unidentified patients. All "
          f"patients who are assessed in the hospital are registered. Government regulations that "
          f"apply to registration and admission are followed. Documentation is designed to avoid "
          f"duplication — information once generated is available to departments that need it.")
    p(doc,"Identity verification at registration is done by government photo ID where available, or "
          "two identifiers stated by the patient or accompanying person. All admissions are authorised "
          "by a doctor. Additional documentation for foreign nationals is collected as required by "
          "applicable rules.")
    p(doc,"General consent for treatment is obtained when the patient enters the organisation. The "
          "organisation defines the scope of that general consent and explains it to the patient "
          "and/or family. General consent does not include invasive procedures or other procedures "
          "that need specific consent under the PRE consent policy.")
    h(doc,2,"5.2 Unique identification number at end of registration")
    p(doc,f"At the end of the patient's first registration interaction with {HN}, a unique "
          f"identification number is generated. That number identifies the patient across the "
          f"organisation and supports continuity of care. All hospital records of the patient "
          f"carry this number.")
    p(doc,"Unique means a one-time assignment: a patient has only one unique number. For later "
          "out-patient or in-patient visits a visit or encounter number may be generated in addition; "
          "those numbers are linked to the unique number. Registration staff must not create a second "
          "unique number for the same person.")
    h(doc,2,"5.3 Accept patients only when the required service can be provided")
    p(doc,f"Registration and admission staff know the services {HN} can provide (from the AAC.1 "
          f"service directory) and whom to contact for clarification — the duty doctor, the Head of "
          f"the relevant department, or the Medical Superintendent.")
    p(doc,"Patients are accepted only if the organisation can provide the required service. In "
          "emergency, life-saving treatment is initiated before any decision about acceptance. If "
          "after stabilisation the hospital still cannot provide the needed service, AAC.3 "
          "transfer-out applies.")
    h(doc,2,"5.4 Managing non-availability of beds")
    p(doc,f"{HN} maintains a current list of alternate organisations where patients may be directed "
          f"when beds are unavailable.")
    p(doc,"If patients are admitted to a temporary holding area, that area has adequate infrastructure "
          "to care for them. The guidance defines how long patients may remain on temporary beds "
          "before a transfer-out decision is taken — not more than 12 hours unless the Medical "
          "Superintendent extends in writing for a named patient.")
    p(doc,"When a bed is not available in the desired bed category or unit, staff manage placement "
          "per the guidance and explain the financial implications to the patient and/or family before "
          "confirming the alternative category.")
    h(doc,2,"5.5 Prioritise access according to clinical needs")
    p(doc,"Access to healthcare services is prioritised according to the patient's clinical needs in "
          "out-patient and diagnostic settings. Patients whose clinical problem warrants an earlier "
          "response are identified and seen sooner.")
    p(doc,f"All staff who handle queues and diagnostic flow are oriented to these guidelines. {HN} "
          f"uses a visual identification mechanism — a coloured sticker on the file or on the patient "
          f"clothing — so that all concerned staff can recognise priority patients. The mechanism is "
          f"defined in writing and taught at induction.")
    h(doc,1,"6. Governance and responsibility")
    gov_tbl(doc,[
        ("Medical Superintendent","Accountable for registration and admission written guidance, bed "
         "non-availability rules and clinical-priority rules. Approves alternate-organisation list "
         "and temporary-holding time limits."),
        ("Registration / front-office staff","Register all assessed patients, including unidentified "
         "patients. Generate the unique identification number at end of registration and never "
         "duplicate it. Obtain general consent at entry and explain its scope. Apply clinical-priority "
         "visual identification in OPD and diagnostics."),
        ("Treating doctors","Authorise every admission. Decide acceptance against available services; "
         "start life-saving treatment in emergency before acceptance decisions."),
        ("Nursing administration / ward nurses","Receive admitted patients; escalate when temporary "
         "holding exceeds the defined time."),
        ("Medical Records","Maintain the unique-number master index linked to visit numbers."),
        ("Quality Coordinator","Audits registration, unique-number integrity, bed-holding and "
         "priority identification quarterly."),
        ("Department heads","Clarify service availability for registration staff when contacted."),
    ])
    h(doc,1,"7. Quality monitoring (RCA → CAPA)")
    p(doc,"The Quality Coordinator audits this policy quarterly. What is monitored each quarter:")
    lb(doc,"Written registration and admission guidance is current and covers unidentified patients, "
           "OP/IP/day-care/emergency.")
    lb(doc,"Sample of admissions shows doctor authorisation and unique identification number on all records.")
    lb(doc,"No duplicate unique numbers in the sample; visit numbers are linked where used.")
    lb(doc,"Temporary-holding episodes stay within the defined time; financial explanation for "
           "category mismatch is documented where applicable.")
    lb(doc,"Clinical-priority visual identification is in use in OPD and diagnostics; staff "
           "orientation records exist.")
    p(doc,"Root-cause analysis is required when a duplicate unique number, an unauthorised admission, "
          "or a temporary-holding over-run recurs within six months.")
    p(doc,"This policy is reviewed annually, and sooner when registration systems or bed capacity change.")
    h(doc,1,"8. Training and staff acknowledgement")
    p(doc,"Registration, front-office, medical records, treating doctors and nurses who admit patients "
          "are informed of this policy at induction and once a year after that. Training covers "
          "unidentified-patient registration, unique-number rules, general consent scope at entry, "
          "bed non-availability, and clinical-priority visual identification.")
    p(doc,"Staff acknowledgement")
    p(doc,f"I have read the Policy on Registration and Admission of {HN}. I will follow the "
          f"processes described.")
    p(doc,"Name: ___________________________    Designation: ___________________________")
    p(doc,"Department / floor: ____________________    Date: ____________")
    p(doc,"Signature: ___________________________")
    p(doc,"(One row per staff member. Registration in-charge and the Quality Coordinator hold "
          "signed acknowledgements.)")
    h(doc,1,"9. Distribution")
    p(doc,f"Official master copy: office of the Medical Superintendent, {HN}, with the Quality Coordinator.")
    p(doc,"Copies issued to: registration; front office; medical records; emergency; day-care; "
          "every ward; nursing administration; billing (bed-category financial implications).")
    p(doc,"The current version is available to all staff at the front-office policy file and, if "
          "the hospital keeps an intranet, at staff intranet / policies.")
    p(doc,"When a new version is issued, take old copies out of use.")
    h(doc,1,"10. Abbreviations")
    abbrev_tbl(doc,[
        ("AAC","Access, Assessment and Continuity of Care (NABH Hospitals chapter)"),
        ("CAPA","corrective and preventive action"),
        ("HCO","Hospital (Full Accreditation programme under NABH Hospitals 6th Edition)"),
        ("NABH","National Accreditation Board for Hospitals and Healthcare Providers"),
        ("OE","objective element"),
        ("OP / OPD","out-patient / out-patient department"),
        ("PRE","Patient Rights and Education (NABH chapter)"),
    ])
    h(doc,1,"11. Traceability to NABH HCO Full Accreditation 6th Edition AAC.2")
    p(doc,"This table is an index. It is not how the policy is organised. An asterisk in the Level "
          "column means documentation of the process is required.")
    trace_tbl(doc,[
        ("AAC.2.a","Commitment*",
         "The organisation uses written guidance for registering and admitting patients.",
         "Section 3; 5.1","Medical Superintendent (approve guidance); Registration / front-office (execute); Treating doctors (authorise admission)"),
        ("AAC.2.b","CORE",
         "A unique identification number is generated at the end of the registration.",
         "Section 3; 5.2","Registration / front-office (generate); Medical Records (master index)"),
        ("AAC.2.c","Commitment",
         "Patients are accepted only if the organisation can provide the required service.",
         "Section 3; 5.3","Registration / front-office (screen); Treating doctors (acceptance and emergency life-saving); Department heads (clarify services)"),
        ("AAC.2.d","Commitment*",
         "The written guidance also addresses managing patients during non-availability of beds.",
         "Section 3; 5.4","Medical Superintendent (approve guidance and time limits); Nursing administration (temporary holding); Registration (alternate organisations and financial explanation)"),
        ("AAC.2.e","Commitment*",
         "Access to the healthcare services in the organisation is prioritised according to the clinical needs of the patient.",
         "Section 3; 5.5","Registration / front-office and diagnostic reception (apply priority); Quality Coordinator (audit orientation)"),
    ])
    h(doc,1,"12. Required Records / Evidence Checklist")
    p(doc,"Records the hospital holds under this policy, listed by objective element.")
    h(doc,2,"AAC.2.a — Written guidance for registering and admitting patients")
    lb(doc,"Written registration and admission guidance covering OP, day-care, IP, emergency and unidentified patients.")
    lb(doc,"Sample registration records showing assessed patients registered and general consent obtained at entry with scope explained.")
    lb(doc,"Admission orders or notes showing doctor authorisation for each admission in the audit sample.")
    h(doc,2,"AAC.2.b — Unique identification number generated at end of registration")
    lb(doc,"Unique-number master index showing one unique number per patient.")
    lb(doc,"Sample clinical records carrying the unique number across departments.")
    lb(doc,"Audit log of duplicate-number checks with corrections.")
    h(doc,2,"AAC.2.c — Patients accepted only when the required service can be provided")
    lb(doc,"Current AAC.1 service directory available at registration points.")
    lb(doc,"Emergency records showing life-saving treatment started before acceptance decisions where applicable.")
    lb(doc,"Log of non-accepted non-emergency presentations with reason and advice given.")
    h(doc,2,"AAC.2.d — Written guidance for non-availability of beds")
    lb(doc,"Written bed non-availability guidance including temporary holding, time limits and category mismatch.")
    lb(doc,"Current list of alternate organisations for redirection.")
    lb(doc,"Temporary-holding log with arrival time, decision time and outcome.")
    lb(doc,"Documented explanation of financial implications when desired bed category is unavailable.")
    h(doc,2,"AAC.2.e — Access prioritised according to clinical needs")
    lb(doc,"Written clinical-priority guidelines for OPD and diagnostic services.")
    lb(doc,"Description of the visual identification mechanism (sticker or equivalent) and where it is placed.")
    lb(doc,"Staff orientation records for queue and diagnostic staff.")
    lb(doc,"Sample of priority patients identified and seen earlier, with visual marker noted.")
    h(doc,1,"13. References")
    lb(doc,"National Accreditation Board for Hospitals and Healthcare Providers (NABH), Guidebook to "
           "Accreditation Standards for Hospitals, 6th Edition — Access, Assessment and Continuity of "
           "Care (AAC), standard AAC.2.")
    lb(doc,f"Internal documents of {HN}: registration and admission written guidance; unique-number "
            f"procedure; alternate-organisation list; temporary-holding and bed-category guidance; "
            f"clinical-priority and visual-identification guidance; PRE informed-consent policy.")
    h(doc,1,"Disclaimer")
    disclaimer(doc)
    save_and_verify(doc,"HCO_AAC_2_v2_REWRITE_DRAFT.docx")


# ══════════════════════════════════════════════════════════════════════════════
# AAC.3 — Transfer In, Transfer Out and Referral  [HAS STOP-WORK]
# ══════════════════════════════════════════════════════════════════════════════
def gen_aac3():
    doc=Document()
    h(doc,0,"Policy on Transfer In, Transfer Out and Referral")
    p(doc,HN)
    h(doc,1,"Document control")
    doc_ctrl(doc,"HCO/AAC/POL/03","Emergency In-Charge")
    p(doc,"A blank marked ________ must be completed before issue.")
    h(doc,1,"Statement of intent")
    p(doc,"Patients who arrive by transfer, leave by transfer or referral, or move for diagnostics "
          "do so safely — stabilised where needed, accompanied by the right staff, with a written "
          "summary of condition and treatment.")
    h(doc,1,"1. Purpose")
    p(doc,f"This policy says how {HN} receives transfer-in patients appropriately (planned and "
          f"unplanned); transfers out or refers patients to another facility appropriately; assigns "
          f"accompanying staff appropriate to the clinical condition; and gives a summary of the "
          f"patient's condition and treatment given.")
    p(doc,"The chapter intent is that emergency patients receive life-stabilising treatment and are "
          "then either admitted or transferred appropriately, and that transfer and discharge protocols "
          "are well defined.")
    p(doc,"This policy does not cover in-hospital registration, admission, routine discharge-summary "
          "content, or continuity of care within the organisation — those are covered in other hospital "
          "policies. The other AAC standards have their own policies too.")
    h(doc,1,"2. Scope")
    p(doc,f"This policy applies to emergency, wards, intensive care, day-care and any unit at {HN} "
          f"that receives a transfer-in or initiates transfer-out or referral. It binds treating "
          f"doctors, accompanying nurses and technologists, ambulance coordinators, registration "
          f"staff who receive transfer-in patients, Medical Records, the Medical Superintendent, "
          f"and the Quality Coordinator.")
    p(doc,"It covers AAC.3.a–d.")
    p(doc,"Boundaries:")
    lb(doc,"AAC.2 owns bed non-availability holding; once the decision is to transfer out, this "
           "policy owns the transfer.")
    lb(doc,"AAC.12 / AAC.13 own discharge-summary content for patients discharged and transferred "
           "(including LAMA); this policy requires that a discharge or transfer summary is given and "
           "a copy retained.")
    lb(doc,"COP standards own clinical stabilisation techniques; this policy owns that stabilisation, "
           "mode, equipment, monitoring and accompanying staff are decided and documented before transfer.")
    lb(doc,"Lab and imaging own the receiving diagnostic service when a patient is shifted for tests; "
           "this policy owns safe transfer to that service when it is outside the immediate unit.")
    h(doc,1,"3. Policy standards")
    p(doc,f"{HN} receives transfer-in patients appropriately for both planned and unplanned transfers. "
          f"Transfer-out and referral to another facility are done appropriately, including pre-transfer "
          f"stabilisation where appropriate, choice of mode and vehicle, equipment and monitoring, with "
          f"gaps documented when requirements cannot be met. Accompanying staff match the patient's "
          f"clinical condition. A transfer summary — or a discharge summary when the patient is "
          f"discharged and transferred, including leaving against medical advice — is given, and a "
          f"copy is retained.")
    p(doc,f"{HN} does not transfer an unstable patient without meeting transfer criteria, completing "
          f"appropriate stabilisation, and assigning appropriate accompanying staff.")
    h(doc,1,"4. Non-negotiable rules")
    ln(doc,"Do not accept or dispatch a transfer-in without recording planned versus unplanned status "
           "and the clinical information received.")
    ln(doc,"Do not transfer out or refer a patient without consulting the patient and/or family where "
           "they can be consulted, and without addressing stabilisation, mode, equipment and monitoring "
           "— or documenting why a stated requirement could not be met.")
    ln(doc,"Do not send an unstable admitted patient for transfer-out or for diagnostic shift without "
           "a doctor accompanying.")
    ln(doc,"Do not send any transfer or referral without accompanying staff who are at least a trained "
           "trauma technologist, emergency technologist or nurse with basic or advanced cardiopulmonary "
           "resuscitation training as appropriate, and who know the transfer procedure.")
    ln(doc,"Do not complete a transfer without giving a transfer summary (or discharge summary if "
           "discharged and transferred, including LAMA) and retaining a copy.")
    ln(doc,f"Staff who find transfer rules not followed report it the same shift to the treating "
            f"doctor or the Medical Superintendent.")
    h(doc,1,"5. What we do")
    h(doc,2,"5.1 Transfer-in done appropriately")
    p(doc,f"Transfer-in covers planned and unplanned arrivals. For unplanned transfers and suspected "
          f"unstable patients, {HN} may send a suitably trained person with the ambulance, guided by "
          f"the information received. Feedback on the patient's clinical status is provided to the "
          f"referring organisation or doctor as good practice.")
    p(doc,"On arrival, registration under AAC.2 applies. The receiving doctor documents the "
          "transfer-in, clinical status on arrival, and whether the transfer was planned or unplanned. "
          "The Emergency In-Charge owns the unplanned transfer-in pathway; department heads own "
          "planned specialty transfer-ins.")
    h(doc,2,"5.2 Transfer-out and referral done appropriately")
    p(doc,"Patients needing transfer-out or referral include those who present to emergency but need "
          "another organisation, those already admitted who now need care elsewhere, and patients being "
          "shifted for diagnostic tests outside the unit's immediate capability.")
    p(doc,"Transfer is done in consultation with the patient and/or family, in a safe manner that includes:")
    lb(doc,"pre-transfer stabilisation where appropriate;")
    lb(doc,"choice of mode and vehicle for transport;")
    lb(doc,"equipment required during transfer;")
    lb(doc,"monitoring required during transfer.")
    p(doc,f"If {HN} cannot meet some of these stated requirements, the reasons are documented in the "
          f"transfer record before the patient leaves. The treating doctor authorises the transfer-out "
          f"or referral.")
    h(doc,2,"5.3 Accompanying staff appropriate to clinical condition")
    p(doc,"Staff accompanying a transfer or referral are at least a trained trauma technologist, "
          "emergency technologist or nurse. That person has undergone training in basic or advanced "
          "cardiopulmonary resuscitation as appropriate to the role, and knows the transfer procedure.")
    p(doc,"A doctor accompanies an unstable admitted patient who is being transferred out or being "
          "shifted for diagnostic purposes. Stability is judged by the treating doctor against airway, "
          "breathing, circulation, conscious level and any specialty-specific instability criteria "
          "written in the transfer guidance.")
    p(doc,"The names and roles of accompanying staff are recorded on the transfer summary.")
    h(doc,2,"5.4 Summary of condition and treatment given")
    p(doc,f"{HN} gives a transfer summary stating significant findings and treatment given to every "
          f"patient transferred from the emergency ward or transferred for diagnostic or therapeutic purposes.")
    p(doc,"When a patient is discharged from the organisation and transferred out, a discharge summary "
          "is given — including patients leaving against medical advice (LAMA). A copy of the transfer "
          "or discharge summary is retained by the hospital in the medical record.")
    p(doc,"Medical Records verifies that the retained copy is filed under the unique identification "
          "number before the record is closed for that episode.")
    h(doc,1,"6. Stop-work authority")
    p(doc,f"Do not transfer an unstable patient out of {HN}, or shift an unstable admitted patient "
          f"for diagnostics, unless transfer criteria are met, appropriate pre-transfer stabilisation "
          f"has been done, and accompanying staff appropriate to the clinical condition (including a "
          f"doctor for an unstable admitted patient) are assigned and ready.")
    p(doc,f"If those conditions are not met, stop the transfer, continue stabilisation and escalate "
          f"to the treating doctor and the Medical Superintendent the same shift. Document why the "
          f"transfer was held.")
    p(doc,"Stop-work does not block emergency life-saving measures. It blocks unsafe transfer movement "
          "until criteria, stabilisation and accompanying staff are in place.")
    p(doc,"Refusing an unsafe transfer is not a disciplinary matter.")
    h(doc,1,"7. Governance and responsibility")
    gov_tbl(doc,[
        ("Medical Superintendent","Accountable for transfer-in/out and referral written guidance and "
         "stop-work enforcement. Approves ambulance and accompanying-staff arrangements."),
        ("Treating doctors","Authorise transfer-out and referral; judge stability; accompany unstable "
         "admitted patients; write or approve transfer/discharge summaries."),
        ("Emergency In-Charge / department heads","Own unplanned and planned transfer-in pathways "
         "respectively. Ensure feedback to referring organisations where practicable."),
        ("Accompanying nurses / technologists","Hold current CPR training appropriate to role; know "
         "the transfer procedure; monitor during transfer."),
        ("Registration / front office","Register transfer-in patients under AAC.2; assist family "
         "communication for transfer-out."),
        ("Medical Records","Retain copies of transfer and discharge summaries under the unique number."),
        ("Quality Coordinator","Audits transfer documentation, accompanying-staff fitness and "
         "stop-work events quarterly."),
    ])
    h(doc,1,"8. Quality monitoring (RCA → CAPA)")
    p(doc,"The Quality Coordinator audits this policy quarterly. What is monitored each quarter:")
    lb(doc,"Sample of transfer-in records shows planned/unplanned status and clinical information on arrival.")
    lb(doc,"Sample of transfer-out/referral records shows stabilisation, mode, equipment, monitoring, "
           "and documented gaps where requirements were not met.")
    lb(doc,"Accompanying staff match clinical condition; doctor accompanies unstable admitted "
           "transfers in the sample.")
    lb(doc,"Transfer or discharge summary (including LAMA where applicable) given and copy retained.")
    lb(doc,"Stop-work events (held unsafe transfers) are logged with outcome.")
    p(doc,"Root-cause analysis is required when an unstable transfer proceeds without criteria, "
          "stabilisation or appropriate accompanying staff.")
    p(doc,"This policy is reviewed annually, and sooner after any adverse event during transfer.")
    h(doc,1,"9. Training and staff acknowledgement")
    p(doc,"Doctors, nurses and technologists who accompany transfers, emergency staff and medical "
          "records staff are informed of this policy at induction and once a year after that. Training "
          "covers planned/unplanned transfer-in, stabilisation and documentation of gaps, "
          "accompanying-staff rules including doctor for unstable patients, CPR expectations, "
          "transfer/discharge summaries including LAMA, and stop-work.")
    p(doc,"Staff acknowledgement")
    p(doc,f"I have read the Policy on Transfer In, Transfer Out and Referral of {HN}. I will "
          f"follow the processes described.")
    p(doc,"Name: ___________________________    Designation: ___________________________")
    p(doc,"Department / floor: ____________________    Date: ____________")
    p(doc,"Signature: ___________________________")
    p(doc,"(One row per staff member. The Quality Coordinator holds signed acknowledgements with "
          "clinical induction records.)")
    h(doc,1,"10. Distribution")
    p(doc,f"Official master copy: office of the Medical Superintendent, {HN}, with the Quality Coordinator.")
    p(doc,"Copies issued to: emergency; every ward and ICU; day-care; ambulance coordination; "
          "medical records; registration; nursing administration.")
    p(doc,"The current version is available to all staff at the clinical policy file and, if the "
          "hospital keeps an intranet, at staff intranet / policies.")
    p(doc,"When a new version is issued, take old copies out of use.")
    h(doc,1,"11. Abbreviations")
    abbrev_tbl(doc,[
        ("AAC","Access, Assessment and Continuity of Care (NABH Hospitals chapter)"),
        ("CAPA","corrective and preventive action"),
        ("CPR","cardiopulmonary resuscitation"),
        ("HCO","Hospital (Full Accreditation programme under NABH Hospitals 6th Edition)"),
        ("ICU","intensive care unit"),
        ("LAMA","left against medical advice"),
        ("NABH","National Accreditation Board for Hospitals and Healthcare Providers"),
        ("OE","objective element"),
    ])
    h(doc,1,"12. Traceability to NABH HCO Full Accreditation 6th Edition AAC.3")
    p(doc,"This table is an index. It is not how the policy is organised. An asterisk in the Level "
          "column means documentation of the process is required.")
    trace_tbl(doc,[
        ("AAC.3.a","Commitment*",
         "Transfer-in of patients to the organisation is done appropriately.",
         "Section 3; 5.1","Emergency In-Charge (unplanned); department heads (planned); treating doctors (receive); registration (AAC.2)"),
        ("AAC.3.b","Commitment*",
         "Transfer-out / referral of patients to another facility is done appropriately.",
         "Section 3; 5.2; Section 6 Stop-work","Treating doctors (authorise); accompanying clinical staff (execute); Medical Superintendent (accountable)"),
        ("AAC.3.c","Commitment",
         "During transfer or referral, accompanying staff are appropriate to the clinical condition of the patient.",
         "Section 3; 5.3; Section 6 Stop-work","Treating doctors (judge stability and assign); accompanying nurses/technologists (execute)"),
        ("AAC.3.d","Commitment",
         "The organisation gives a summary of the patient's condition and the treatment given.",
         "Section 3; 5.4","Treating doctors (write/approve summary); Medical Records (retain copy)"),
    ])
    h(doc,1,"13. Required Records / Evidence Checklist")
    p(doc,"Records the hospital holds under this policy, listed by objective element.")
    h(doc,2,"AAC.3.a — Transfer-in done appropriately")
    lb(doc,"Transfer-in records showing planned or unplanned status and clinical information received.")
    lb(doc,"Ambulance crew or escort assignment notes for unplanned/suspected unstable transfer-ins where used.")
    lb(doc,"Feedback notes to referring organisation/doctor where provided.")
    lb(doc,"Registration of transfer-in patients under the unique identification number.")
    h(doc,2,"AAC.3.b — Transfer-out / referral done appropriately")
    lb(doc,"Transfer-out/referral orders with patient/family consultation noted where consulted.")
    lb(doc,"Documentation of pre-transfer stabilisation, mode/vehicle, equipment and monitoring.")
    lb(doc,"Documented reasons when stated transfer requirements could not be met.")
    lb(doc,"Log of transfers for emergency, admitted patients and diagnostic shifts.")
    h(doc,2,"AAC.3.c — Accompanying staff appropriate to clinical condition")
    lb(doc,"Transfer records naming accompanying staff and role.")
    lb(doc,"Evidence of CPR training (basic or advanced as appropriate) for accompanying staff.")
    lb(doc,"Records showing a doctor accompanied unstable admitted patients transferred out or "
           "shifted for diagnostics.")
    lb(doc,"Written stability/transfer criteria used to decide accompanying level.")
    h(doc,2,"AAC.3.d — Summary of condition and treatment given")
    lb(doc,"Transfer summaries for emergency and diagnostic/therapeutic transfers.")
    lb(doc,"Discharge summaries for patients discharged and transferred out, including LAMA.")
    lb(doc,"Retained copies filed under the unique identification number.")
    lb(doc,"Audit sample confirming summary given before departure.")
    h(doc,1,"14. References")
    lb(doc,"National Accreditation Board for Hospitals and Healthcare Providers (NABH), Guidebook to "
           "Accreditation Standards for Hospitals, 6th Edition — Access, Assessment and Continuity of "
           "Care (AAC), standard AAC.3.")
    lb(doc,f"Internal documents of {HN}: transfer-in/out and referral written guidance; ambulance and "
            f"accompanying-staff roster; CPR training records; transfer-summary and discharge-summary "
            f"templates; AAC.12/AAC.13 discharge policies.")
    h(doc,1,"Disclaimer")
    disclaimer(doc)
    save_and_verify(doc,"HCO_AAC_3_v2_REWRITE_DRAFT.docx")


# ══════════════════════════════════════════════════════════════════════════════
# AAC.4 — Initial Assessment
# ══════════════════════════════════════════════════════════════════════════════
def gen_aac4():
    doc=Document()
    h(doc,0,"Policy on Initial Assessment")
    p(doc,HN)
    h(doc,1,"Document control")
    doc_ctrl(doc,"HCO/AAC/POL/04","Quality Coordinator")
    p(doc,"A blank marked ________ must be completed before issue.")
    h(doc,1,"Statement of intent")
    p(doc,"Every patient undergoes a standardised initial assessment by qualified personnel within a "
          "defined time — producing a documented care plan for in-patients that the clinician in-charge "
          "authorises and that notes special needs after discharge.")
    h(doc,1,"1. Purpose")
    p(doc,f"This policy says how {HN} performs standardised initial assessment for out-patients, "
          f"day-care, in-patients and emergency patients; ensures assessment is done by qualified "
          f"personnel within their scope of practice; completes assessment within a time frame based "
          f"on patient need; includes nursing assessment at admission for day-care and in-patients; "
          f"documents a care plan from the in-patient initial assessment; obtains countersignature by "
          f"the clinician in-charge within 24 hours; and identifies special needs regarding care "
          f"following discharge.")
    p(doc,"The chapter intent is that patients undergo initial assessment and periodic reassessments "
          "resulting in a care plan.")
    p(doc,"This policy does not cover re-assessment, registration, or clinical pathways for specific "
          "procedures — those are covered in other hospital policies. The other AAC and COP standards "
          "have their own policies too.")
    h(doc,1,"2. Scope")
    p(doc,f"This policy applies to treating doctors, nurses and other caregivers who perform initial "
          f"assessment at {HN} in out-patient, day-care (including dialysis day-care), in-patient "
          f"wards and emergency. It binds department heads who define who may assess, the Medical "
          f"Superintendent, Medical Records, and the Quality Coordinator.")
    p(doc,"It covers AAC.4.a–g.")
    p(doc,"Boundaries:")
    lb(doc,"AAC.5 owns re-assessment intervals, care-plan modification after initial assessment, and "
           "early-warning escalation.")
    lb(doc,"AAC.2 owns that the patient is registered and admitted before or alongside assessment.")
    lb(doc,"COP.10 / COP.11 own specific procedural pathways referenced by the guidebook; initial "
           "assessment for those pathways still uses this standard's formats.")
    lb(doc,"Lab and imaging own result reporting; this policy owns that available diagnostics inform "
           "the care plan.")
    h(doc,1,"3. Policy standards")
    p(doc,f"{HN} performs initial assessment of out-patients, day-care, in-patients and emergency "
          f"patients in a standardised manner. Assessment is performed by qualified personnel within "
          f"their scope of practice, registration and applicable laws. Assessment is completed within "
          f"a documented time frame based on patient needs — for in-patients within a maximum of 24 "
          f"hours from arrival at the ward; for emergency from the time of arrival. Day-care and "
          f"in-patient initial assessment includes nursing assessment at admission. The in-patient "
          f"initial assessment results in a documented care plan. The care plan is countersigned by "
          f"the clinician in-charge within 24 hours. The care plan identifies special needs regarding "
          f"care following discharge.")
    p(doc,f"{HN} does not leave an in-patient without a documented care plan, and does not treat "
          f"the 24-hour maximum as a target when clinical need requires earlier assessment.")
    h(doc,1,"4. Non-negotiable rules")
    ln(doc,"Do not use a non-standard assessment format in an area that has an approved standardised "
           "format for that area.")
    ln(doc,"Do not omit vital parameters in emergency initial assessment; do not omit history, "
           "examination including vitals, drug allergies and provisional diagnosis for in-patients; "
           "do not omit medication reconciliation for in-patients.")
    ln(doc,"Do not allow a person to perform an assessment outside their defined scope of practice "
           "and privilege.")
    ln(doc,"Do not exceed the defined time frames — in-patient initial assessment maximum 24 hours "
           "from ward arrival; emergency from arrival to completion as defined.")
    ln(doc,"Do not omit nursing assessment at admission for day-care and in-patients.")
    ln(doc,"Do not continue planned in-patient treatment beyond emergency stabilisation without a "
           "documented care plan from the initial assessment.")
    ln(doc,"Do not leave a care plan initiated by a junior doctor without countersignature by the "
           "clinician in-charge within 24 hours.")
    ln(doc,f"Staff who find assessment or care-plan rules not followed report it the same shift to "
            f"the treating doctor or the Medical Superintendent.")
    h(doc,1,"5. What we do")
    h(doc,2,"5.1 Standardised initial assessment (OP, day-care, IP, emergency)")
    p(doc,f"{HN} uses a standardised format for initial assessment in the OPD, day-care, in-patient "
          f"and emergency areas. Formats may differ by department need but are the same within a "
          f"particular area or unit. Emergency initial assessment includes recording vital parameters.")
    p(doc,"In-patient initial assessment covers history, examination including vital signs, "
          "documentation of any drug allergies, and provisional diagnosis. If a detailed assessment "
          "was done earlier the same day in OPD or emergency, it need not be rewritten in full; a "
          "comment links to the earlier assessment, and findings of all such assessments are reviewed "
          "and/or verified.")
    p(doc,"Initial assessment includes reconciliation of medications for in-patients. Abridged "
          "documentation may be used for day-care as appropriate, including patients coming for dialysis.")
    h(doc,2,"5.2 Performed by qualified personnel within scope of practice")
    p(doc,f"{HN} determines who can do which assessment. Caregivers perform initial assessment "
          f"within their scope of practice, registration and applicable laws and regulations. "
          f"Psychological, spiritual, cultural, social and economic aspects of in-patient initial "
          f"assessment may be done by various healthcare professionals as named on the privilege list.")
    p(doc,"Department heads maintain the assessment privilege list; the Medical Superintendent approves it.")
    h(doc,2,"5.3 Within a time frame based on patient needs")
    p(doc,f"{HN} defines and documents the time frame for completing initial assessment for day-care, "
          f"in-patients and emergency, and implements it.")
    p(doc,"For in-patients, the clock starts when the patient arrives at the ward and stops when "
          "initial assessment is completed. The maximum time is 24 hours. Patients are assessed "
          "earlier when clinical need requires — unstable or high-dependency within 1 hour; routine "
          "ward within 6 hours where staffing allows, and always within 24 hours.")
    p(doc,"For emergency, the clock starts at arrival at emergency — primary survey and vitals "
          "immediately; full emergency initial assessment within 30 minutes unless resuscitation is "
          "ongoing. Day-care time frames are defined per service — before the planned procedure or "
          "dialysis session begins.")
    h(doc,2,"5.4 Nursing assessment at admission for day-care and in-patients")
    p(doc,"Initial assessment of day-care and in-patients includes nursing assessment at the time of "
          "admission, documented in the record. It identifies nursing needs and any special needs of "
          "the patient, completed within a defined time frame — within 1 hour of ward or day-care arrival.")
    p(doc,"A checklist or template may be used. Abridged nursing documentation may be used for "
          "day-care as appropriate. Templates may be specific to the speciality or type of admission.")
    h(doc,2,"5.5 Documented care plan from in-patient initial assessment")
    p(doc,"For in-patients, the initial assessment results in a documented care plan. The care plan "
          "is written by the treating doctor or a doctor member of the treating team in the patient "
          "record and is followed.")
    p(doc,"The care plan is based on the initial assessment and results of diagnostic tests if "
          "available. It includes provisional diagnosis or differential diagnosis, relevant diagnostic "
          "investigations when required, initial treatment suggested, and specific instructions if "
          "any. It reflects the desired results of the treatment, care or service.")
    h(doc,2,"5.6 Care plan countersigned by clinician in-charge within 24 hours")
    p(doc,"Treatment may be initiated by a junior doctor, but the care plan is countersigned and "
          "authorised by the treating doctor — the clinician in-charge — within 24 hours. "
          "Countersignature is dated and timed in the record.")
    p(doc,"If the clinician in-charge will be unavailable beyond 24 hours, the covering consultant "
          "named for that period countersigns and the Medical Superintendent is informed.")
    h(doc,2,"5.7 Special needs regarding care following discharge")
    p(doc,"The care plan includes identification of special needs regarding care following discharge. "
          "Identification is critical for groups such as extremes of age, restricted mobility, "
          "continuing nursing and rehabilitation needs, and assistance with activities of daily living.")
    p(doc,f"{HN} begins identifying special discharge needs early in the assessment process — not "
          f"only on the day of discharge. Nursing and medical assessments both contribute. Detailed "
          f"discharge process remains under AAC.12; this step owns early identification inside the care plan.")
    h(doc,1,"6. Governance and responsibility")
    gov_tbl(doc,[
        ("Medical Superintendent","Accountable for standardised formats, privilege lists and time "
         "frames. Approves assessment formats and privilege lists."),
        ("Department heads","Maintain area-specific formats and who may assess. Ensure medication "
         "reconciliation and care-plan quality in their units."),
        ("Treating doctors / clinician in-charge","Perform or supervise initial assessment; document "
         "care plan; countersign within 24 hours; identify special discharge needs."),
        ("Junior doctors","May initiate assessment and treatment; must obtain countersignature "
         "within 24 hours."),
        ("Nurses","Complete nursing assessment at admission for day-care and in-patients; contribute "
         "special-needs identification."),
        ("Medical Records","Ensure formats are available in records; flag missing care plans or "
         "countersignatures in record review."),
        ("Quality Coordinator","Audits formats, time frames, nursing assessment, care plans, "
         "countersignatures and discharge-needs identification quarterly."),
    ])
    h(doc,1,"7. Quality monitoring (RCA → CAPA)")
    p(doc,"The Quality Coordinator audits this policy quarterly. What is monitored each quarter:")
    lb(doc,"Standardised formats in use in OPD, day-care, IP and emergency; vitals in emergency; "
           "IP content and medication reconciliation complete.")
    lb(doc,"Persons who performed initial assessment match the privilege list.")
    lb(doc,"In-patient assessments completed within 24 hours of ward arrival; emergency within "
           "defined time from arrival.")
    lb(doc,"Nursing assessment documented at admission for day-care and IP sample.")
    lb(doc,"Care plan present for IP sample; countersigned by clinician in-charge within 24 hours.")
    lb(doc,"Special discharge needs identified in care plans for applicable patients.")
    p(doc,"Root-cause analysis is required when missing care plans or late countersignatures recur "
          "within six months.")
    p(doc,"This policy is reviewed annually, and sooner when formats or privilege lists change.")
    h(doc,1,"8. Training and staff acknowledgement")
    p(doc,"Doctors, nurses and other caregivers who assess patients are informed of this policy at "
          "induction and once a year after that. Training covers area formats, privilege limits, time "
          "frames, medication reconciliation, nursing assessment, care-plan content, 24-hour "
          "countersignature, and early identification of special discharge needs.")
    p(doc,"Staff acknowledgement")
    p(doc,f"I have read the Policy on Initial Assessment of {HN}. I will follow the processes described.")
    p(doc,"Name: ___________________________    Designation: ___________________________")
    p(doc,"Department / floor: ____________________    Date: ____________")
    p(doc,"Signature: ___________________________")
    p(doc,"(One row per staff member. The Quality Coordinator holds signed acknowledgements with "
          "clinical induction records.)")
    h(doc,1,"9. Distribution")
    p(doc,f"Official master copy: office of the Medical Superintendent, {HN}, with the Quality Coordinator.")
    p(doc,"Copies issued to: every clinical department; emergency; day-care; nursing administration; "
          "medical records; OPD.")
    p(doc,"The current version is available to all staff at the clinical policy file and, if the "
          "hospital keeps an intranet, at staff intranet / policies.")
    p(doc,"When a new version is issued, take old copies out of use.")
    h(doc,1,"10. Abbreviations")
    abbrev_tbl(doc,[
        ("AAC","Access, Assessment and Continuity of Care (NABH Hospitals chapter)"),
        ("CAPA","corrective and preventive action"),
        ("COP","Care of Patients (NABH chapter)"),
        ("HCO","Hospital (Full Accreditation programme under NABH Hospitals 6th Edition)"),
        ("IP","in-patient"),
        ("NABH","National Accreditation Board for Hospitals and Healthcare Providers"),
        ("OE","objective element"),
        ("OP / OPD","out-patient / out-patient department"),
    ])
    h(doc,1,"11. Traceability to NABH HCO Full Accreditation 6th Edition AAC.4")
    p(doc,"This table is an index. It is not how the policy is organised. An asterisk in the Level "
          "column means documentation of the process is required.")
    trace_tbl(doc,[
        ("AAC.4.a","CORE*",
         "The initial assessment of the out-patients, daycare, in-patients and emergency patients is done in a standardised manner.",
         "Section 3; 5.1","Department heads (formats); treating doctors (complete); nurses (contribute); Quality Coordinator (hold formats)"),
        ("AAC.4.b","Commitment*",
         "The initial assessment is performed by qualified personnel.",
         "Section 3; 5.2","Medical Superintendent (approve privilege list); department heads (maintain); caregivers (assess within scope)"),
        ("AAC.4.c","Commitment*",
         "The initial assessment is performed within a time frame based on the needs of the patient.",
         "Section 3; 5.3","Treating doctors (complete on time); department heads (define specialty defaults); Quality Coordinator (audit)"),
        ("AAC.4.d","Commitment",
         "Initial assessment of daycare and in-patients includes nursing assessment, which is done at the time of admission and documented.",
         "Section 3; 5.4","Nurses (perform and document); nursing administration (templates and time frames)"),
        ("AAC.4.e","CORE",
         "The initial assessment for in-patients results in a documented care plan.",
         "Section 3; 5.5","Treating doctor or doctor member of treating team (document); Medical Records (presence in record)"),
        ("AAC.4.f","Achievement",
         "The care plan is countersigned by the clinician in-charge of the patient within 24 hours.",
         "Section 3; 5.6","Clinician in-charge / treating doctor (countersign); junior doctors (seek countersignature)"),
        ("AAC.4.g","Excellence",
         "The care plan includes the identification of special needs regarding care following discharge.",
         "Section 3; 5.7","Treating doctors and nurses (identify early); AAC.12 owners (execute discharge planning)"),
    ])
    h(doc,1,"12. Required Records / Evidence Checklist")
    p(doc,"Records the hospital holds under this policy, listed by objective element.")
    h(doc,2,"AAC.4.a — Standardised initial assessment")
    lb(doc,"Approved standardised initial-assessment formats for OPD, day-care, IP and emergency.")
    lb(doc,"Emergency records showing vital parameters recorded at initial assessment.")
    lb(doc,"IP records showing history, examination including vitals, drug allergies, provisional "
           "diagnosis and medication reconciliation.")
    lb(doc,"Day-care records using abridged documentation where appropriate (including dialysis).")
    h(doc,2,"AAC.4.b — Performed by qualified personnel")
    lb(doc,"Assessment privilege list naming who may perform which assessment.")
    lb(doc,"Credentials and registration evidence for privileged clinicians who perform assessment.")
    lb(doc,"Sample records showing the person who assessed matches the privilege list.")
    h(doc,2,"AAC.4.c — Within a time frame based on patient needs")
    lb(doc,"Written time-frame document for day-care, IP and emergency initial assessment.")
    lb(doc,"IP sample with ward-arrival time and assessment-completion time within 24 hours.")
    lb(doc,"Emergency sample with arrival time and assessment-completion time within the defined frame.")
    lb(doc,"Day-care sample showing assessment before procedure/session as defined.")
    h(doc,2,"AAC.4.d — Nursing assessment at admission")
    lb(doc,"Nursing-assessment templates/checklists for IP and day-care.")
    lb(doc,"Documented nursing assessments at admission identifying nursing and special needs.")
    lb(doc,"Defined time frame for completing nursing assessment and audit against it.")
    h(doc,2,"AAC.4.e — Documented care plan from in-patient initial assessment")
    lb(doc,"Documented care plans in IP records including provisional/differential diagnosis.")
    lb(doc,"Care-plan entries for relevant investigations, initial treatment and specific instructions.")
    lb(doc,"Audit sample of IP admissions with care plan present after initial assessment.")
    h(doc,2,"AAC.4.f — Care plan countersigned by clinician in-charge within 24 hours")
    lb(doc,"Care plans with dated/timed countersignature by clinician in-charge within 24 hours.")
    lb(doc,"Covering-consultant countersignatures when treating doctor unavailable, with Medical "
           "Superintendent notified.")
    lb(doc,"Audit log of late or missing countersignatures with corrective action.")
    h(doc,2,"AAC.4.g — Special needs regarding care following discharge identified")
    lb(doc,"Care-plan fields or notes identifying special post-discharge needs.")
    lb(doc,"Evidence identification began early in the assessment process, not only on discharge day.")
    lb(doc,"Sample of applicable patients with special needs flagged and handed to discharge planning.")
    h(doc,1,"13. References")
    lb(doc,"National Accreditation Board for Hospitals and Healthcare Providers (NABH), Guidebook to "
           "Accreditation Standards for Hospitals, 6th Edition — Access, Assessment and Continuity of "
           "Care (AAC), standard AAC.4.")
    lb(doc,"Correlated standards named in the guidebook interpretation: COP.10.d; COP.11.e.")
    lb(doc,f"Internal documents of {HN}: initial-assessment formats (OP, day-care, IP, emergency); "
            f"assessment privilege list; nursing-assessment templates; care-plan template; "
            f"medication-reconciliation form; AAC.5 re-assessment policy; AAC.12 discharge policy.")
    h(doc,1,"Disclaimer")
    disclaimer(doc)
    save_and_verify(doc,"HCO_AAC_4_v2_REWRITE_DRAFT.docx")


# ══════════════════════════════════════════════════════════════════════════════
# AAC.5 — Re-assessment  [HAS STOP-WORK]
# ══════════════════════════════════════════════════════════════════════════════
def gen_aac5():
    doc=Document()
    h(doc,0,"Policy on Re-assessment")
    p(doc,HN)
    h(doc,1,"Document control")
    doc_ctrl(doc,"HCO/AAC/POL/05","Quality Coordinator")
    p(doc,"A blank marked ________ must be completed before issue.")
    h(doc,1,"Statement of intent")
    p(doc,"After the first assessment, patients are reassessed often enough to see whether treatment "
          "is working — the care plan moves with the patient, notes are real clinical notes, and early "
          "warning signs trigger prompt action.")
    h(doc,1,"1. Purpose")
    p(doc,f"This policy says how {HN} reassesses patients at appropriate intervals to determine "
          f"response to treatment and to plan further treatment or discharge; informs out-patients of "
          f"the next follow-up where appropriate; monitors and modifies the in-patient care plan "
          f"during re-assessment; requires direct clinical care staff to document re-assessments "
          f"properly; and lays down guidelines to identify early warning signs of change or "
          f"deterioration for prompt intervention.")
    p(doc,"The chapter intent is that patients undergo initial assessment and periodic reassessments "
          "resulting in a care plan.")
    p(doc,"This policy does not cover initial assessment, care plan creation, clinical emergency "
          "protocols, or human resource staffing — those are covered in other hospital policies. "
          "The other AAC and COP standards have their own policies too.")
    h(doc,1,"2. Scope")
    p(doc,f"This policy applies to treating doctors, nurses and other caregivers who reassess "
          f"patients at {HN} in out-patient, day-care, in-patient wards, ICU/high-dependency and "
          f"areas where patients await admission or a bed. It binds the Medical Superintendent, "
          f"department heads, Medical Records, and the Quality Coordinator.")
    p(doc,"It covers AAC.5.a–e.")
    p(doc,"Boundaries:")
    lb(doc,"AAC.4 owns initial assessment and creation of the care plan; this policy owns monitoring "
           "and modifying that plan at re-assessment.")
    lb(doc,"AAC.12 owns the discharge process; this policy owns re-assessment that informs the "
           "decision to discharge.")
    lb(doc,"COP.5 / COP.6 own related clinical emergency and acute-care pathways named in the "
           "guidebook correlation; this policy owns early-warning identification and escalation "
           "under AAC.5.e.")
    lb(doc,"HRM.1 / HRM.5 own staffing competence and related HR elements named in the guidebook "
           "correlation; this policy requires trained staff to use early-warning parameters, not the "
           "full HRM programme.")
    h(doc,1,"3. Policy standards")
    p(doc,f"{HN} reassesses patients at appropriate intervals to determine response to treatment and "
          f"to plan further treatment or discharge. Out-patients are informed of their next follow-up "
          f"where appropriate. For in-patients, the care plan is monitored and modified during "
          f"re-assessment where necessary. Staff involved in direct clinical care document "
          f"re-assessments. The organisation lays down guidelines and implements processes to "
          f"identify early warning signs of change or deterioration for prompt intervention.")
    p(doc,f"{HN} does not leave a deteriorating patient without escalation when early-warning "
          f"triggers fire, and does not accept reassessment notes that say only 'patient well' or "
          f"'condition better'.")
    h(doc,1,"4. Non-negotiable rules")
    ln(doc,"Do not go a calendar day without re-assessment of an in-patient by the treating doctor "
           "or a doctor from the treating team; reassess more often in ICU and when the condition "
           "changes; reassess day-care patients before discharge.")
    ln(doc,"Do not send an out-patient away without documenting the next follow-up where a repeat "
           "visit is appropriate.")
    ln(doc,"Do not leave an ineffective in-patient care plan unchanged when re-assessment shows "
           "it must be modified.")
    ln(doc,"Do not document re-assessment with only phrases such as 'patient well' or 'condition "
           "better' — include vitals, systemic findings and medication orders at minimum (nurses "
           "may document vitals).")
    ln(doc,"Do not ignore early-warning triggers — escalate to appropriate medical personnel and "
           "intervene promptly.")
    ln(doc,f"Staff who find re-assessment or early-warning rules not followed report it the same "
            f"shift to the treating doctor or the Medical Superintendent.")
    h(doc,1,"5. What we do")
    h(doc,2,"5.1 Re-assess at appropriate intervals")
    p(doc,"After initial assessment, every patient is reassessed periodically and the re-assessment "
          "is documented in the case sheet. Re-assessment is done by all applicable caregivers within "
          "their scope of practice, registration and applicable laws.")
    p(doc,"Every patient is reassessed at least once every day by the treating doctor or a doctor "
          "from the treating team. Day-care patients are reassessed before discharge. Patients "
          "awaiting admission or a bed are also reassessed as applicable.")
    p(doc,"The hospital defines and documents re-assessment frequencies for each clinical area — "
          "including ICU/high-dependency and ward — based on patient condition, clinical need and "
          "speciality. Department heads and the ICU in-charge set these frequencies in area-specific guidance.")
    h(doc,2,"5.2 Inform out-patients of next follow-up where appropriate")
    p(doc,"Out-patients are informed of their next follow-up where appropriate. The information is "
          "a specific date or a period (weeks/months) and is documented in the medical record or OP "
          "consultation sheet.")
    p(doc,"Follow-up information may be omitted when the patient came only for an opinion or when "
          "the condition does not warrant a repeat visit — note that reason briefly in the OP record.")
    h(doc,2,"5.3 Monitor and modify the in-patient care plan during re-assessment")
    p(doc,"For in-patients, during re-assessment the care plan is monitored for effectiveness in "
          "achieving the desired results of treatment, care or service. The care plan is dynamic. "
          "The treating doctor or a doctor member of the treating team modifies it where necessary "
          "according to the patient's condition.")
    p(doc,"Changes are documented in the medical record. Progress notes, doctor's orders or "
          "medication charts may show the change — the change must be findable, not only spoken.")
    h(doc,2,"5.4 Direct clinical care staff document re-assessments")
    p(doc,"Actions taken under re-assessment are documented by the treating doctor or any member "
          "of the team per their domain of responsibility. At a minimum, documentation includes "
          "vitals, systemic examination findings and medication orders. Nursing staff may document "
          "the patient's vitals.")
    p(doc,"Phrases alone such as 'patient well' or 'condition better' are not acceptable as a "
          "re-assessment note.")
    h(doc,2,"5.5 Early warning signs — guidelines and prompt intervention")
    p(doc,f"{HN} lays down guidelines and implements processes to identify early warning signs of "
          f"change or deterioration in clinical condition and to initiate prompt intervention.")
    p(doc,"Staff use defined physiological parameters — which may include vital parameters, airway, "
          "circulation, neurological status, and any other concerns felt by staff or the patient/family. "
          "Parameters may be tailored to speciality and age group.")
    p(doc,"There is a mechanism to make this information available to appropriate medical personnel "
          "to initiate prompt and appropriate actions. Effectiveness of the early-warning system is "
          "monitored monthly by the Quality Coordinator with clinical leads.")
    h(doc,1,"6. Stop-work authority")
    p(doc,"Do not leave a deteriorating patient without escalation when early-warning signs or "
          "defined physiological triggers indicate change or deterioration.")
    p(doc,"If a trigger is met and appropriate medical personnel have not yet been informed and "
          "acted, do not defer escalation for non-urgent work; escalate immediately per the "
          "early-warning guideline, start prompt intervention within your scope, and document the "
          "trigger, time of escalation, responder and actions.")
    p(doc,"Stop-work applies to failure to escalate a deteriorating patient. It does not block "
          "emergency life-saving measures — those continue while escalation happens.")
    p(doc,"The person responsible tells the treating doctor and, if there is no timely response, "
          "the Medical Superintendent the same shift. Refusing to leave a deteriorating patient "
          "unescalated is not a disciplinary matter.")
    h(doc,1,"7. Governance and responsibility")
    gov_tbl(doc,[
        ("Medical Superintendent","Accountable for re-assessment intervals, documentation standards "
         "and early-warning guidelines. Receives escalation when treating-doctor response is not timely."),
        ("Treating doctors / treating team","Daily medical re-assessment at minimum; more often as "
         "needed; modify care plans; respond to early-warning alerts."),
        ("Nurses and other direct clinical caregivers","Reassess within scope; document vitals and "
         "nursing findings; trigger early-warning escalation; never rely on 'patient well' alone."),
        ("OPD doctors","Document next follow-up where appropriate."),
        ("Department heads / ICU in-charge","Set area-specific frequencies and early-warning "
         "parameter sets for speciality and age group."),
        ("Quality Coordinator","Audits daily re-assessment, follow-up documentation, care-plan "
         "changes, note quality and early-warning effectiveness quarterly (early-warning "
         "effectiveness also reviewed monthly)."),
        ("Medical Records","Flag missing daily notes and inadequate re-assessment phrases during "
         "record review."),
    ])
    h(doc,1,"8. Quality monitoring (RCA → CAPA)")
    p(doc,"The Quality Coordinator audits this policy quarterly and reviews early-warning "
          "effectiveness monthly. What is monitored:")
    lb(doc,"IP sample shows at least daily medical re-assessment; ICU more frequent; day-care "
           "re-assessed before discharge.")
    lb(doc,"OP sample shows follow-up documented where appropriate, or a brief reason when not.")
    lb(doc,"Care-plan modifications documented when re-assessment required change.")
    lb(doc,"Re-assessment notes include vitals, systemic findings and medication orders — not only "
           "'patient well'.")
    lb(doc,"Early-warning triggers escalated with time, responder and action documented; failed "
           "escalations treated as incidents.")
    p(doc,"Root-cause analysis is required when a deteriorating patient was not escalated, or when "
          "inadequate re-assessment notes recur within six months.")
    p(doc,"This policy is reviewed annually, and sooner after any failure-to-escalate event.")
    h(doc,1,"9. Training and staff acknowledgement")
    p(doc,"Doctors, nurses and other direct clinical caregivers are informed of this policy at "
          "induction and once a year after that. Training covers interval rules, OP follow-up "
          "documentation, care-plan modification, acceptable re-assessment content, early-warning "
          "parameters, escalation routes and stop-work.")
    p(doc,"Staff acknowledgement")
    p(doc,f"I have read the Policy on Re-assessment of {HN}. I will follow the processes described.")
    p(doc,"Name: ___________________________    Designation: ___________________________")
    p(doc,"Department / floor: ____________________    Date: ____________")
    p(doc,"Signature: ___________________________")
    p(doc,"(One row per staff member. The Quality Coordinator holds signed acknowledgements with "
          "clinical induction records.)")
    h(doc,1,"10. Distribution")
    p(doc,f"Official master copy: office of the Medical Superintendent, {HN}, with the Quality Coordinator.")
    p(doc,"Copies issued to: every ward and ICU; day-care; OPD; emergency; nursing administration; "
          "medical records.")
    p(doc,"The current version is available to all staff at the clinical policy file and, if the "
          "hospital keeps an intranet, at staff intranet / policies.")
    p(doc,"When a new version is issued, take old copies out of use.")
    h(doc,1,"11. Abbreviations")
    abbrev_tbl(doc,[
        ("AAC","Access, Assessment and Continuity of Care (NABH Hospitals chapter)"),
        ("CAPA","corrective and preventive action"),
        ("COP","Care of Patients (NABH chapter)"),
        ("HCO","Hospital (Full Accreditation programme under NABH Hospitals 6th Edition)"),
        ("HRM","Human Resource Management (NABH chapter)"),
        ("ICU","intensive care unit"),
        ("NABH","National Accreditation Board for Hospitals and Healthcare Providers"),
        ("OE","objective element"),
        ("OP / OPD","out-patient / out-patient department"),
    ])
    h(doc,1,"12. Traceability to NABH HCO Full Accreditation 6th Edition AAC.5")
    p(doc,"This table is an index. It is not how the policy is organised. An asterisk in the Level "
          "column means documentation of the process is required.")
    trace_tbl(doc,[
        ("AAC.5.a","CORE",
         "Patients are re-assessed at appropriate intervals to determine their response to treatment and to plan further treatment or discharge.",
         "Section 3; 5.1","Treating doctors (daily medical re-assessment); nurses and caregivers (within scope); ICU in-charge (ICU frequency)"),
        ("AAC.5.b","Commitment",
         "Out-patients are informed of their next follow-up, where appropriate.",
         "Section 3; 5.2","OPD treating doctors (inform and document)"),
        ("AAC.5.c","Commitment",
         "For in-patients during re-assessment, the care plan is monitored and modified, where found necessary.",
         "Section 3; 5.3","Treating doctor or doctor member of treating team (monitor and modify)"),
        ("AAC.5.d","Commitment",
         "Staff involved in direct clinical care document re-assessments.",
         "Section 3; 5.4","Treating doctors and team members (document per domain); nurses (vitals)"),
        ("AAC.5.e","Commitment*",
         "The organisation lays down guidelines and implements processes to identify early warning signs of change or deterioration in clinical conditions for initiating prompt intervention.",
         "Section 3; 5.5; Section 6 Stop-work","Medical Superintendent (approve guideline); treating doctors (respond); nurses (trigger); Quality Coordinator (monitor effectiveness)"),
    ])
    h(doc,1,"13. Required Records / Evidence Checklist")
    p(doc,"Records the hospital holds under this policy, listed by objective element.")
    h(doc,2,"AAC.5.a — Re-assessment at appropriate intervals")
    lb(doc,"IP case sheets showing at least daily re-assessment by treating doctor or treating-team doctor.")
    lb(doc,"ICU observation charts showing more frequent re-assessment than ward.")
    lb(doc,"Day-care records showing re-assessment before discharge.")
    h(doc,2,"AAC.5.b — Out-patients informed of next follow-up")
    lb(doc,"OP consultation sheets with next follow-up date or interval documented.")
    lb(doc,"Sample where follow-up was not applicable with brief reason noted.")
    h(doc,2,"AAC.5.c — In-patient care plan monitored and modified")
    lb(doc,"Progress notes, doctor's orders or medication charts showing care-plan changes.")
    lb(doc,"Sample where ineffective plan was modified with clinical reason.")
    h(doc,2,"AAC.5.d — Re-assessments documented by direct clinical care staff")
    lb(doc,"Re-assessment notes including vitals, systemic examination findings and medication orders.")
    lb(doc,"Nursing vital-sign charts contributing to re-assessment documentation.")
    lb(doc,"Record-review log rejecting notes that only say 'patient well' / 'condition better'.")
    h(doc,2,"AAC.5.e — Early warning signs — guidelines and prompt intervention")
    lb(doc,"Written early-warning guideline with physiological parameters (speciality/age tailored where used).")
    lb(doc,"Escalation records showing trigger, time informed, responder and actions.")
    lb(doc,"Monthly early-warning effectiveness review notes.")
    lb(doc,"Incident/RCA records for failure-to-escalate events.")
    h(doc,1,"14. References")
    lb(doc,"National Accreditation Board for Hospitals and Healthcare Providers (NABH), Guidebook to "
           "Accreditation Standards for Hospitals, 6th Edition — Access, Assessment and Continuity of "
           "Care (AAC), standard AAC.5.")
    lb(doc,"Correlated standards named in the guidebook interpretation: COP.5; COP.6; HRM.1; HRM.5.")
    lb(doc,f"Internal documents of {HN}: re-assessment interval guidance by area; early-warning / "
            f"track-and-trigger guideline; care-plan modification method; AAC.4 initial-assessment "
            f"policy; related COP and HRM policies.")
    h(doc,1,"Disclaimer")
    disclaimer(doc)
    save_and_verify(doc,"HCO_AAC_5_v2_REWRITE_DRAFT.docx")


# ══════════════════════════════════════════════════════════════════════════════
# AAC.6 — Laboratory Services
# ══════════════════════════════════════════════════════════════════════════════
def gen_aac6():
    doc=Document()
    h(doc,0,"Policy on Laboratory Services")
    p(doc,HN)
    h(doc,1,"Document control")
    doc_ctrl(doc,"HCO/AAC/POL/06","Laboratory In-Charge")
    p(doc,"A blank marked ________ must be completed before issue.")
    h(doc,1,"Statement of intent")
    p(doc,"Laboratory services are provided as per the scope of services of the organisation — so "
          "that results are available round the clock where clinical care needs them, handled safely "
          "from requisition to disposal, reported on time, and quality-assured.")
    h(doc,1,"1. Purpose")
    p(doc,f"This policy says how {HN} provides laboratory services commensurate with its clinical "
          f"scope, keeps infrastructure and human resources adequate, uses qualified staff under "
          f"pathologist/microbiologist/biochemist supervision, follows written guidance for the "
          f"specimen journey, delivers results within defined turnaround time (TAT), intimates "
          f"critical results within one hour with read-back, reports in a standardised manner, "
          f"recalls or amends reports when needed, and outsources unavailable tests under a "
          f"quality-assured MoU.")
    p(doc,"The chapter intent is that laboratory services are provided by competent staff so that "
          "patient care is continuous and diagnostic results support safe clinical decisions.")
    p(doc,"This policy does not cover laboratory quality assurance, laboratory safety, infection "
          "control, or biomedical-waste compliance — those are covered in other hospital policies. "
          "The other AAC standards have their own policies too.")
    h(doc,1,"2. Scope")
    p(doc,f"This policy applies to the laboratory in-charge, pathologist, microbiologist, "
          f"biochemist, laboratory technologists, treating doctors who requisition tests, nurses "
          f"who collect or transport specimens, and the Quality Coordinator at {HN}.")
    p(doc,"It covers the ten objective elements AAC.6.a-j: scope and round-the-clock availability; "
          "infrastructure and layout; human resources; qualified performance and supervision; written "
          "specimen guidance; TAT; critical results; standardised reporting; recall/amendment; "
          "outsourcing with MoU.")
    p(doc,"Boundaries:")
    lb(doc,"AAC.7 owns the laboratory quality assurance programme and the laboratory safety "
           "programme (IQC/EQA, MSDS, PPE training). This policy owns service scope, specimen "
           "process, TAT, critical intimation, reporting and outsourcing.")
    lb(doc,"HIC and the organisation's biomedical-waste / infection-control documents own waste "
           "segregation categories and statutory waste compliance. This policy requires disposal "
           "per those documents; it does not import that statute into paragraph 2 of the disclaimer.")
    lb(doc,"ROM.6.e owns MoU structure for outsourced diagnostic services; this policy requires "
           "an MoU that incorporates quality assurance for laboratory outsourcing.")
    lb(doc,"AAC.1 / AAC.2 own displayed clinical services and registration identity; the unique "
           "identification number used on specimens comes from registration.")
    h(doc,1,"3. Policy standards")
    p(doc,f"{HN} provides laboratory services commensurate with the healthcare services it offers, "
          f"available round the clock so that patient care is not disrupted, with emergency-management "
          f"test results available on its premises. Infrastructure, layout, human resources and "
          f"qualified supervision are adequate for the defined scope. Specimens are requisitioned, "
          f"collected, identified, handled, transported, processed and disposed of according to "
          f"written guidance. Results are available within a defined TAT, critical results are "
          f"intimated within one hour with documented read-back, and reports follow a standardised "
          f"format. There is a mechanism to recall or amend reports. Tests not available in-house "
          f"are outsourced only to organisations with a quality assurance system under a written MoU.")
    p(doc,f"{HN} does not delay a critical laboratory result, and does not release an altered "
          f"outsourced report as if it were an in-house result.")
    h(doc,1,"4. Non-negotiable rules")
    ln(doc,"Do not offer a laboratory service outside the approved test menu without laboratory "
           "in-charge agreement and menu update or a documented outsourced arrangement.")
    ln(doc,"Do not accept or process a specimen that is unlabelled, mislabelled, or collected "
           "without following written guidance.")
    ln(doc,"Do not report a result without the authorised signatory required by statute and by "
           "this hospital's authorisation list.")
    ln(doc,"Do not leave a critical result uncommunicated beyond one hour after the report is ready.")
    ln(doc,"Do not alter or modify content of an outsourced laboratory report.")
    ln(doc,"Do not outsource a test to a laboratory that has no documented quality assurance "
           "system and no MoU with this hospital.")
    ln(doc,"Do not leave a recalled report in clinical areas, the medical record, LIS or HIS after "
           "recall without replacement by the amended report.")
    ln(doc,f"Staff who see a laboratory service rule broken report it the same shift to the "
            f"Laboratory In-Charge or the Medical Superintendent.")
    h(doc,1,"5. What we do")
    h(doc,2,"5.1 Scope of laboratory services")
    p(doc,f"Scope of the laboratory services at {HN} is commensurate with the healthcare services "
          f"the organisation provides. The laboratory in-charge maintains a test menu listing every "
          f"test available in-house and every test outsourced. The menu is reviewed annually and "
          f"whenever clinical services change.")
    p(doc,"Laboratory services are available round the clock so that patient care is not disrupted. "
          "Test results required for emergency management are available within the premises. Tests "
          "not on the menu are not offered as in-house services.")
    h(doc,2,"5.2 Infrastructure and layout")
    p(doc,"The infrastructure — physical space and equipment — is adequate to provide the defined "
          "scope of services. Reports must not be delayed for lack of adequate equipment.")
    p(doc,"The layout of the laboratory prevents cross-contamination. Clean and dirty workflows are "
          "separated; specimen receipt, processing and reporting areas are arranged so that "
          "contamination routes are interrupted. The laboratory in-charge reviews layout fitness "
          "annually and after any renovation or major equipment change.")
    h(doc,2,"5.3 Human resources")
    p(doc,"Human resource is adequate to provide the defined scope of services. The number of "
          "laboratory personnel is commensurate with workload, with sufficient staff for each shift "
          "and for emergencies. Reports must not be delayed for lack of adequate human resource, "
          "including personnel authorised to report results.")
    h(doc,2,"5.4 Qualified performance and supervision")
    p(doc,"Qualified and trained personnel perform and supervise the investigations and report the "
          "results. Staff employed in the laboratory hold an appropriate degree and are trained to "
          "carry out the tests they perform.")
    p(doc,"A Pathologist, Microbiologist and Biochemist supervise the staff for their respective "
          "disciplines as this hospital's scope requires. Statutory requirements regarding the "
          "authorised signatory are adhered to. The laboratory in-charge keeps the current list of "
          "authorised signatories and does not release reports outside that list.")
    h(doc,2,"5.5 Written guidance for requisition through disposal")
    p(doc,"Requisition for tests, collection, identification, handling, safe transportation, "
          "processing and disposal of a specimen is performed according to written guidance.")
    p(doc,"The unique identification number from registration is used for identification of the "
          "patient. Two-patient identifiers are checked at collection and at receipt. Rejected "
          "specimens (haemolysed, clotted, insufficient, mislabelled) are logged with reason and "
          "a fresh specimen is requested.")
    p(doc,"Disposal of waste follows the statutory biomedical-waste duties owned by this hospital's "
          "infection-control / waste documents. This step requires laboratory staff to segregate and "
          "hand over laboratory waste under those documents.")
    h(doc,2,"5.6 Results within defined turnaround time")
    p(doc,"Laboratory results are available within a defined time frame. The organisation defines "
          "the turnaround time for all tests and ensures adequate staff, materials and equipment so "
          "that results are available within that frame.")
    p(doc,"The turnaround time may differ for different tests and is decided based on the nature of "
          "the test, criticality and urgency desired by the treating doctor. The laboratory in-charge "
          "documents the defined TAT for each test or test category in the laboratory test menu or a "
          "separate TAT schedule.")
    p(doc,"The laboratory in-charge monitors TAT compliance monthly and escalates persistent breaches "
          "to the Quality Coordinator.")
    h(doc,2,"5.7 Critical results intimated within one hour")
    p(doc,"Critical results are intimated to the person concerned at the earliest.")
    p(doc,"Critical limits for tests that require immediate attention for patient management are "
          "documented. Critical results of outsourced investigations are included. Critical test "
          "results are communicated to a person from the treating team at the earliest, and not "
          "later than one hour after completion of the test or the report being ready.")
    p(doc,"The intimation includes: name of the patient; unique ID; date and time of sample "
          "collection; test name, result, measure unit and reference range; identity of who "
          "communicated the value; identity of the recipient; read-back; and date and time of "
          "acknowledgement.")
    h(doc,2,"5.8 Standardised reporting")
    p(doc,"Results are reported in a standardised manner. At a minimum the report includes: the "
          "name of the organisation (or, for an outsourced laboratory, the name of that laboratory); "
          "the patient's name; the unique identification number; the reference range of the test "
          "where applicable; and the name and signature of the person reporting the test result.")
    p(doc,"The organisation does not alter or modify anything in the outsourced report. If on this "
          "organisation's letterhead, the report includes at least the name of the outsourced "
          "laboratory, the date and the reference number of the report given by the outsourced laboratory.")
    h(doc,2,"5.9 Recall and amendment of reports")
    p(doc,"There is a mechanism to address the recall or amendment of reports whenever applicable. "
          "When a particular report is recalled, it is withdrawn from clinical areas, medical "
          "records, the Laboratory Information System (LIS) and the Hospital Information System "
          "(HIS). If already issued to the patient, the amended report is made available with a "
          "caution to ignore the earlier one.")
    h(doc,2,"5.10 Outsourced laboratory tests")
    p(doc,"Laboratory tests not available in the organisation are outsourced to organisation(s) "
          "based on their quality assurance system.")
    p(doc,"Written guidance for outsourcing includes: a list of tests for outsourcing; identity of "
          "personnel in the outsourced facilities; manner of packaging and labelling of specimens; "
          "a methodology to check the performance of the service. The organisation has a Memorandum "
          "of Understanding (MoU) / agreement that incorporates quality assurance and the "
          "requirements of this standard. The panel and MoUs are reviewed annually.")
    h(doc,1,"6. Governance and responsibility")
    gov_tbl(doc,[
        ("Medical Superintendent","Accountable that laboratory services match clinical scope, "
         "remain available round the clock for defined emergency tests, and meet TAT and "
         "critical-result duties."),
        ("Laboratory In-Charge","Maintains the test menu, staffing plan, authorised-signatory list, "
         "written specimen guidance, TAT definitions, critical-limit list, standardised report "
         "format, recall mechanism and outsourced panel with MoUs."),
        ("Pathologist / Microbiologist / Biochemist","Supervise investigations and reporting in "
         "their disciplines; act as or nominate authorised signatories as statute and hospital "
         "authorisation require."),
        ("Laboratory technologists","Perform tests they are trained for; follow specimen guidance; "
         "communicate critical results when designated; do not alter outsourced reports."),
        ("Treating doctors and ward nurses","Requisition appropriately; receive critical-result "
         "intimation with read-back; act on amended reports after recall."),
        ("Quality Coordinator","Audits this policy quarterly; tracks CAPA when TAT, critical-result "
         "or recall defects recur."),
    ])
    h(doc,1,"7. Quality monitoring (RCA → CAPA)")
    p(doc,"The Quality Coordinator audits this policy quarterly. What is monitored each quarter:")
    lb(doc,"Test menu current and commensurate with clinical services; emergency tests available on premises.")
    lb(doc,"Infrastructure, layout and staffing adequate for scope.")
    lb(doc,"Authorised-signatory list current; unqualified reporting absent.")
    lb(doc,"Specimen guidance followed; rejection log reviewed.")
    lb(doc,"TAT compliance by category.")
    lb(doc,"Critical-result communication within one hour with full documentation and read-back.")
    lb(doc,"Standardised report format; outsourced letterhead rules followed; no alteration of outsourced content.")
    lb(doc,"Recall/amendment log complete with withdrawal evidence and CAPA.")
    lb(doc,"Outsourced panel and MoUs current and quality-based.")
    p(doc,"Root-cause analysis is required when the same laboratory service defect recurs within six months.")
    p(doc,"This policy is reviewed annually, and sooner when clinical scope, equipment or outsourcing arrangements change.")
    h(doc,1,"8. Training and staff acknowledgement")
    p(doc,"All laboratory staff, nurses who collect or transport specimens and doctors who requisition "
          "tests are trained on this policy at induction and once a year after that. Training covers "
          "specimen guidance, TAT, critical-result intimation with read-back, standardised reporting, "
          "recall/amendment and outsourcing rules.")
    p(doc,"Staff acknowledgement")
    p(doc,f"I have read the Policy on Laboratory Services of {HN}. I will follow the processes described.")
    p(doc,"Name: ___________________________    Designation: ___________________________")
    p(doc,"Department / floor: ____________________    Date: ____________")
    p(doc,"Signature: ___________________________")
    p(doc,"(One row per staff member. The Quality Coordinator holds signed acknowledgements with the induction record.)")
    h(doc,1,"9. Distribution")
    p(doc,f"Official master copy: office of the Medical Superintendent, {HN}, with the Quality Coordinator.")
    p(doc,"Copies issued to: laboratory; emergency; every in-patient ward; out-patient; nursing "
          "administration; intensive-care areas if present.")
    p(doc,"The current version is available to all staff at the front-office policy file and, if "
          "the hospital keeps an intranet, at staff intranet / policies.")
    p(doc,"When a new version is issued, take old copies out of use.")
    h(doc,1,"10. Abbreviations")
    abbrev_tbl(doc,[
        ("AAC","Access, Assessment and Continuity of Care (NABH HCO chapter)"),
        ("CAPA","corrective and preventive action"),
        ("HCO","Hospital (Full Accreditation programme)"),
        ("HIS","Hospital Information System"),
        ("HIC","Hospital Infection Control"),
        ("LIS","Laboratory Information System"),
        ("MoU","Memorandum of Understanding"),
        ("NABL","National Accreditation Board for Testing and Calibration Laboratories"),
        ("NABH","National Accreditation Board for Hospitals and Healthcare Providers"),
        ("OE","objective element"),
        ("PPE","personal protective equipment"),
        ("PSQ","Patient Safety and Quality"),
        ("ROM","Responsibilities of Management"),
        ("TAT","turnaround time"),
    ])
    h(doc,1,"11. Traceability to NABH HCO Full Accreditation 6th Edition AAC.6")
    p(doc,"This table is an index. It is not how the policy is organised. An asterisk in the Level "
          "column means documentation of the process is required.")
    trace_tbl(doc,[
        ("AAC.6.a","Commitment","Scope of the laboratory services is commensurate to the services provided by the organisation.","Section 3; 5.1","Laboratory In-Charge (maintain menu); Medical Superintendent (approve scope)"),
        ("AAC.6.b","Commitment","The infrastructure (physical and equipment) is adequate to provide the defined scope of services; layout prevents cross-contamination.","Section 3; 5.2","Laboratory In-Charge (manage); Medical Superintendent (resource)"),
        ("AAC.6.c","Commitment","Human resource is adequate to provide the defined scope of services.","Section 3; 5.3","Laboratory In-Charge (roster); Medical Superintendent (resource)"),
        ("AAC.6.d","Commitment","Qualified and trained personnel perform and supervise the investigations and report the results.","Section 3; 5.4","Pathologist/Microbiologist/Biochemist (supervise); Laboratory In-Charge (authorised signatories)"),
        ("AAC.6.e","Commitment*","Requisition for tests, collection, identification, handling, safe transportation, processing and disposal of a specimen is performed according to written guidance.","Section 3; 5.5","Laboratory In-Charge (own SOPs); technologists and collecting nurses (follow)"),
        ("AAC.6.f","Commitment*","Laboratory results are available within a defined time frame.","Section 3; 5.6","Laboratory In-Charge (define and monitor TAT); Quality Coordinator (audit)"),
        ("AAC.6.g","Commitment*","Critical results are intimated to the person concerned at the earliest.","Section 3; 5.7","Designated laboratory personnel (intimate); treating team (read-back); Laboratory In-Charge (limits list)"),
        ("AAC.6.h","Commitment","Results are reported in a standardised manner.","Section 3; 5.8","Laboratory In-Charge (format); authorised signatory (sign); technologists (do not alter outsourced content)"),
        ("AAC.6.i","Commitment*","There is a mechanism to address the recall / amendment of reports whenever applicable.","Section 3; 5.9","Laboratory In-Charge (run mechanism); Quality Coordinator (CAPA)"),
        ("AAC.6.j","Commitment*","Laboratory tests not available in the organisation are outsourced to organisation(s) based on their quality assurance system.","Section 3; 5.10","Laboratory In-Charge (panel and MoU); Medical Superintendent (approve MoU)"),
    ])
    h(doc,1,"12. Required Records / Evidence Checklist")
    p(doc,"Records the hospital holds under this policy, listed by objective element.")
    h(doc,2,"AAC.6.a — Scope commensurate with services")
    lb(doc,"Laboratory test menu listing in-house and outsourced tests.")
    lb(doc,"Annual review of menu against clinical services.")
    lb(doc,"Record that emergency-management tests required by scope are available on premises.")
    h(doc,2,"AAC.6.b — Infrastructure and layout adequate")
    lb(doc,"Equipment inventory with model, serial and maintenance schedule link.")
    lb(doc,"Layout / workflow description showing separation that prevents cross-contamination.")
    lb(doc,"Annual infrastructure adequacy review.")
    h(doc,2,"AAC.6.c — Human resource adequate")
    lb(doc,"Staff list by shift including emergency cover.")
    lb(doc,"Annual staff-to-workload review.")
    lb(doc,"List of personnel authorised to report results.")
    h(doc,2,"AAC.6.d — Qualified performance and supervision")
    lb(doc,"Qualification and training records for laboratory staff.")
    lb(doc,"Supervision arrangement by Pathologist, Microbiologist and Biochemist as applicable.")
    lb(doc,"Current authorised-signatory list meeting statutory requirements.")
    h(doc,2,"AAC.6.e — Written guidance for requisition through disposal")
    lb(doc,"Written guidance covering requisition, collection, identification, handling, transport, processing and disposal.")
    lb(doc,"Specimen rejection log with reasons.")
    lb(doc,"Evidence of unique ID (and lab number if used) on specimens and requisitions.")
    h(doc,2,"AAC.6.f — Results within defined TAT")
    lb(doc,"Defined TAT for all tests or test categories.")
    lb(doc,"Monthly TAT compliance report.")
    lb(doc,"Escalation records for persistent TAT breaches.")
    h(doc,2,"AAC.6.g — Critical results intimated within one hour")
    lb(doc,"Documented critical limits and biological reference intervals (or evaluated published intervals).")
    lb(doc,"Critical-result communication log with patient ID, result, caller, recipient, read-back, date and time.")
    lb(doc,"List of personnel authorised to report critical results.")
    h(doc,2,"AAC.6.h — Standardised reporting")
    lb(doc,"Standardised report template with minimum required fields.")
    lb(doc,"Sample outsourced reports on outsourced letterhead or organisation letterhead with required attribution.")
    lb(doc,"Audit finding that outsourced report content was not altered.")
    h(doc,2,"AAC.6.i — Recall / amendment mechanism")
    lb(doc,"Recall/amendment log with reason and date.")
    lb(doc,"Evidence of withdrawal from clinical areas, medical records, LIS and HIS.")
    lb(doc,"Amended report issued to patient with caution where previously issued; CAPA record.")
    h(doc,2,"AAC.6.j — Outsourced laboratory tests")
    lb(doc,"Written outsourcing guidance with test list, transport contacts, packaging and performance checks.")
    lb(doc,"Current MoU/agreement incorporating quality assurance (ROM.6.e).")
    lb(doc,"Annual review of outsourced laboratory performance.")
    h(doc,1,"13. References")
    lb(doc,"National Accreditation Board for Hospitals and Healthcare Providers (NABH), Guidebook to "
           "Accreditation Standards for Hospitals, 6th Edition — Access, Assessment and Continuity of "
           "Care (AAC), standard AAC.6.")
    lb(doc,"ISO 15189:2022 and NABL 112 — guidance references for laboratory quality systems owned "
           "under AAC.7; cited here only as boundary.")
    lb(doc,f"Internal documents of {HN}: laboratory test menu; specimen SOPs; critical-limit list; "
            f"authorised-signatory list; TAT definitions; recall/amendment log; outsourced laboratory "
            f"MoUs (ROM.6.e); infection-control / biomedical-waste procedures for disposal.")
    h(doc,1,"Disclaimer")
    disclaimer(doc)
    save_and_verify(doc,"HCO_AAC_6_v2_REWRITE_DRAFT.docx")


# ══════════════════════════════════════════════════════════════════════════════
# AAC.7 — Laboratory Quality Assurance and Safety Programme
# ══════════════════════════════════════════════════════════════════════════════
def gen_aac7():
    doc=Document()
    h(doc,0,"Policy on Laboratory Quality Assurance and Safety Programme")
    p(doc,HN)
    h(doc,1,"Document control")
    doc_ctrl(doc,"HCO/AAC/POL/07","Laboratory In-Charge")
    p(doc,"A blank marked ________ must be completed before issue.")
    h(doc,1,"Statement of intent")
    p(doc,"There is an established laboratory quality assurance and safety programme — so that test "
          "results are accurate and precise, and laboratory staff work under documented safety "
          "controls aligned with the organisation's safety programme.")
    h(doc,1,"1. Purpose")
    p(doc,f"This policy says how {HN} implements a laboratory quality assurance programme covering "
          f"point-of-care testing (POCT), pre-analytic, analytic and post-analytic phases, internal "
          f"quality control (IQC), external quality assurance (EQA) / proficiency testing (PT), "
          f"calibration, corrective and preventive action (CAPA), clinico-pathological meetings, "
          f"and a laboratory safety programme with Material Safety Data Sheets (MSDS), training and "
          f"appropriate safety measures including personal protective equipment (PPE) and immunisation.")
    p(doc,"The chapter intent is that laboratory services are quality-assured and safe for staff "
          "and patients.")
    p(doc,"This policy does not cover laboratory service scope, specimen handling, turnaround times, "
          "critical-result intimation, standardised reporting, or laboratory outsourcing — those are "
          "covered in other hospital policies. The other AAC standards have their own policies too.")
    h(doc,1,"2. Scope")
    p(doc,f"This policy applies to the laboratory in-charge, pathologist, microbiologist, biochemist, "
          f"laboratory technologists, POCT operators outside the central laboratory where laboratory "
          f"QA covers them, and the Quality Coordinator at {HN}.")
    p(doc,"It covers the seven objective elements AAC.7.a-g: QA programme implementation; IQC; "
          "PT/EQA participation; clinico-pathological meetings; laboratory safety programme; training "
          "in safe practices; appropriate safety measures.")
    p(doc,"Boundaries:")
    lb(doc,"AAC.6 owns service delivery (menu, specimen process, TAT, critical intimation, reporting, "
           "outsourcing). Quality failures discovered under this policy feed AAC.6 recall/amendment "
           "when a released report is wrong.")
    lb(doc,"Organisation safety / occupational health documents own hospital-wide safety governance; "
           "this laboratory safety programme is aligned with them and does not replace them.")
    lb(doc,"HIC owns hospital infection-control and biomedical-waste streams; laboratory PPE and "
           "standard precautions here are the laboratory application of those duties.")
    lb(doc,"ISO 15189:2022 and NABL 112 are guidance references for the QA programme, not this "
           "policy's disclaimer statutes.")
    h(doc,1,"3. Policy standards")
    p(doc,f"{HN} implements a documented laboratory quality assurance programme that covers POCT "
          f"and the pre-analytic, analytic and post-analytic phases; uses IQC to ensure quality of "
          f"test results; participates in PT/EQA or documented alternate approaches where formal PT "
          f"is unavailable; documents CAPA for deviations; conducts clinico-pathological meetings "
          f"at defined intervals; and implements a laboratory safety programme with MSDS, staff "
          f"training and appropriate safety measures including PPE, immunisation and standard "
          f"precautions, aligned with the organisation's safety programme.")
    p(doc,f"{HN} does not release patient results from a measuring system that has failed IQC "
          f"without documented override and CAPA, and does not leave laboratory staff without "
          f"required PPE and safety training for the work they perform.")
    h(doc,1,"4. Non-negotiable rules")
    ln(doc,"Do not operate a test method without verification or validation as required by the QA programme.")
    ln(doc,"Do not ignore an IQC failure: stop patient reporting for the affected parameter until "
           "IQC is restored or a documented clinical override with CAPA is recorded.")
    ln(doc,"Do not skip scheduled PT/EQA enrolment for analytes where a formal programme exists "
           "and this hospital performs the test.")
    ln(doc,"Do not leave laboratory chemicals or hazards without accessible MSDS for the reagents in use.")
    ln(doc,"Do not allow laboratory personnel to work without induction training in safe practices "
           "and the relevant MSDS for their job.")
    ln(doc,"Do not work without required PPE or without documented immunisation status as required "
           "for laboratory staff.")
    ln(doc,f"Staff who see a laboratory QA or safety rule broken report it the same shift to the "
            f"Laboratory In-Charge or the Medical Superintendent.")
    h(doc,1,"5. What we do")
    h(doc,2,"5.1 Laboratory quality assurance programme")
    p(doc,"The laboratory quality assurance programme is implemented. The organisation has a "
          "documented programme to ensure accuracy and precision. The programme includes point-of-care "
          "testing (POCT) and covers pre-analytic, analytic and post-analytic phases of the testing "
          "cycle, including test standardisation, IQC, EQA or inter-laboratory testing, verification "
          "and/or validation of test methods, periodic calibration and maintenance of equipment, "
          "and documentation of CAPA whenever deviations are observed.")
    p(doc,"There is a mechanism to obtain feedback from stakeholders to evaluate laboratory services "
          "at least once a year. ISO 15189:2022 and NABL 112 are good reference guides for designing "
          "and reviewing this programme.")
    h(doc,2,"5.2 Internal quality control")
    p(doc,"The programme ensures the quality of test results through internal quality control, "
          "including performing IQC to ensure precision and repeatability for all test parameters "
          "(quantitative and qualitative) and peer review where relevant. The programme includes "
          "comparability of results when more than one measuring system is used.")
    p(doc,"Patient results for a parameter with failed IQC are not released until IQC is acceptable "
          "or a documented override authorised by the Laboratory In-Charge or supervising "
          "Pathologist/Biochemist/Microbiologist is recorded with CAPA.")
    h(doc,2,"5.3 Proficiency testing and external quality assurance")
    p(doc,"The laboratory participates in proficiency testing / external quality assurance schemes. "
          "Based on the EQA/PT evaluation report, the laboratory implements and documents corrective "
          "actions for outliers.")
    p(doc,"Where formal EQA/PT is not a practical option — for example non-availability of a formal "
          "national PT programme for the analyte; only few laboratories performing the test; unstable "
          "analyte (blood gases, ammonia, G6PD); or control material of the same matrix not available "
          "— the laboratory adopts alternate approaches to validate performance.")
    h(doc,2,"5.4 Clinico-pathological meetings")
    p(doc,"The organisation conducts clinico-pathological meetings at pre-defined intervals for "
          "correlating histopathology reports with referring clinicians and uses them as a tool for "
          "improving quality. Meetings are scheduled at least quarterly unless the Medical "
          "Superintendent defines a different interval suited to case volume.")
    h(doc,2,"5.5 Laboratory safety programme")
    p(doc,"The laboratory safety programme is implemented. A laboratory safety manual is available "
          "in the laboratory. It addresses safety of the workforce and of equipment, in consonance "
          "with identified risks and hazards. The manual incorporates appropriate MSDS. The programme "
          "may follow an Occupational Health and Safety Management System approach. It is aligned "
          "with the organisation's safety programme.")
    h(doc,2,"5.6 Training in safe practices")
    p(doc,"Laboratory personnel are appropriately trained in safe practices. All laboratory staff "
          "undergo training regarding safe practices in the laboratory and in the relevant MSDS. "
          "Training-need identification is commensurate with the job description of the staff.")
    p(doc,"Training is provided at induction and annually, and when new hazards or equipment are "
          "introduced. Training records are held by the laboratory in-charge.")
    h(doc,2,"5.7 Appropriate safety measures")
    p(doc,"Laboratory personnel are provided with appropriate safety measures. Adequate safety "
          "measures are available in the laboratory — PPE, dressing materials, disinfectants, fire "
          "extinguishers — addressing safety issues at all levels.")
    p(doc,"All laboratory personnel always adhere to standard precautions. All laboratory staff are "
          "appropriately immunised — hepatitis B at minimum; other immunisations as per the "
          "organisation occupational-health schedule. PPE for routine work includes gloves, laboratory "
          "coat, eye protection, and face shield where splash risk exists. Spill kits are available "
          "and staff are trained in their use.")
    h(doc,1,"6. Governance and responsibility")
    gov_tbl(doc,[
        ("Medical Superintendent","Accountable that laboratory QA and safety programmes are "
         "resourced and aligned with organisation safety."),
        ("Laboratory In-Charge","Authors and keeps current the QA programme, IQC/EQA matrix, "
         "safety manual, MSDS set, training records and CAPA log."),
        ("Pathologist / Microbiologist / Biochemist","Supervise technical quality in their "
         "disciplines; authorise IQC overrides; contribute to clinico-pathological meetings."),
        ("Laboratory technologists and POCT operators","Perform IQC as scheduled; follow safety "
         "manual and PPE rules; report deviations the same shift."),
        ("Quality Coordinator","Audits this policy quarterly; tracks CAPA to closure; reviews "
         "annual stakeholder feedback on laboratory services."),
    ])
    h(doc,1,"7. Quality monitoring (RCA → CAPA)")
    p(doc,"The Quality Coordinator audits this policy quarterly. What is monitored each quarter:")
    lb(doc,"QA programme document current; POCT included; method verification/validation records present.")
    lb(doc,"IQC performance and CAPA for failures; multi-instrument comparability where applicable.")
    lb(doc,"PT/EQA enrolment and outlier CAPA; alternate-approach records where PT unavailable.")
    lb(doc,"Clinico-pathological meeting minutes and follow-up actions.")
    lb(doc,"Laboratory safety manual and MSDS current; alignment with organisation safety.")
    lb(doc,"Training and immunisation records complete.")
    lb(doc,"PPE and safety equipment available and used.")
    p(doc,"Root-cause analysis is required when the same IQC failure mode or safety incident recurs "
          "within six months.")
    p(doc,"This policy is reviewed annually, and sooner after a serious laboratory incident or major "
          "method change.")
    h(doc,1,"8. Training and staff acknowledgement")
    p(doc,"All laboratory staff and POCT operators covered by this programme are trained on this "
          "policy at induction and once a year after that. Training covers QA phases, IQC, PT/EQA, "
          "CAPA, the safety manual, MSDS relevant to their job, PPE and standard precautions.")
    p(doc,"Staff acknowledgement")
    p(doc,f"I have read the Policy on Laboratory Quality Assurance and Safety Programme of {HN}. "
          f"I will follow the processes described.")
    p(doc,"Name: ___________________________    Designation: ___________________________")
    p(doc,"Department / floor: ____________________    Date: ____________")
    p(doc,"Signature: ___________________________")
    p(doc,"(One row per staff member. The Quality Coordinator holds signed acknowledgements with the induction record.)")
    h(doc,1,"9. Distribution")
    p(doc,f"Official master copy: office of the Medical Superintendent, {HN}, with the Quality Coordinator.")
    p(doc,"Copies issued to: laboratory; POCT locations covered by this programme; nursing "
          "administration for ward-based POCT operators.")
    p(doc,"The current version is available to all staff at the front-office policy file and, if "
          "the hospital keeps an intranet, at staff intranet / policies.")
    p(doc,"When a new version is issued, take old copies out of use.")
    h(doc,1,"10. Abbreviations")
    abbrev_tbl(doc,[
        ("AAC","Access, Assessment and Continuity of Care (NABH HCO chapter)"),
        ("CAPA","corrective and preventive action"),
        ("EQA","external quality assurance"),
        ("HCO","Hospital (Full Accreditation programme)"),
        ("IQC","internal quality control"),
        ("MSDS","Material Safety Data Sheet"),
        ("NABL","National Accreditation Board for Testing and Calibration Laboratories"),
        ("NABH","National Accreditation Board for Hospitals and Healthcare Providers"),
        ("OE","objective element"),
        ("POCT","point-of-care testing"),
        ("PPE","personal protective equipment"),
        ("PT","proficiency testing"),
        ("QA","quality assurance"),
    ])
    h(doc,1,"11. Traceability to NABH HCO Full Accreditation 6th Edition AAC.7")
    p(doc,"This table is an index. It is not how the policy is organised. An asterisk in the Level "
          "column means documentation of the process is required.")
    trace_tbl(doc,[
        ("AAC.7.a","Commitment*","The laboratory quality assurance programme is implemented.","Section 3; 5.1","Laboratory In-Charge (own programme); Quality Coordinator (annual review)"),
        ("AAC.7.b","Commitment*","The programme ensures the quality of test results through Internal quality control.","Section 3; 5.2","Laboratory technologists (perform IQC); Laboratory In-Charge (CAPA); supervising Pathologist/Biochemist/Microbiologist (overrides)"),
        ("AAC.7.c","Commitment","Laboratory participates in proficiency testing / external quality assurance scheme.","Section 3; 5.3","Laboratory In-Charge (enrolment and alternate approaches); Quality Coordinator (audit)"),
        ("AAC.7.d","Commitment","The programme addresses the clinico-pathological meeting(s).","Section 3; 5.4","Pathologist (convene); referring clinicians (participate); Laboratory In-Charge (minutes)"),
        ("AAC.7.e","Commitment*","The laboratory safety programme is implemented.","Section 3; 5.5","Laboratory In-Charge (safety manual); Medical Superintendent (alignment with organisation safety)"),
        ("AAC.7.f","Commitment","Laboratory personnel are appropriately trained in safe practices.","Section 3; 5.6","Laboratory In-Charge (training); staff (attend)"),
        ("AAC.7.g","Commitment","Laboratory personnel are provided with appropriate safety measures.","Section 3; 5.7","Laboratory In-Charge (provide measures); staff (use PPE and standard precautions)"),
    ])
    h(doc,1,"12. Required Records / Evidence Checklist")
    p(doc,"Records the hospital holds under this policy, listed by objective element.")
    h(doc,2,"AAC.7.a — QA programme implemented")
    lb(doc,"Documented QA programme covering POCT and pre-/analytic/post-analytic phases.")
    lb(doc,"Method verification/validation records; calibration and maintenance schedule.")
    lb(doc,"Annual stakeholder feedback on laboratory services; CAPA for deviations.")
    h(doc,2,"AAC.7.b — Internal quality control")
    lb(doc,"IQC schedules and results for quantitative and qualitative parameters.")
    lb(doc,"Comparability records when more than one measuring system is used.")
    lb(doc,"CAPA and documented overrides for IQC failures.")
    h(doc,2,"AAC.7.c — Proficiency testing and EQA")
    lb(doc,"PT/EQA enrolment matrix by analyte.")
    lb(doc,"EQA/PT evaluation reports with corrective actions for outliers.")
    lb(doc,"Documented alternate approaches where formal PT is unavailable.")
    h(doc,2,"AAC.7.d — Clinico-pathological meetings")
    lb(doc,"Meeting schedule at pre-defined intervals.")
    lb(doc,"Minutes with cases, correlations and discrepancies.")
    lb(doc,"Follow-up actions linked to quality improvement or CAPA.")
    h(doc,2,"AAC.7.e — Laboratory safety programme implemented")
    lb(doc,"Laboratory safety manual available in the laboratory.")
    lb(doc,"MSDS set for reagents in use.")
    lb(doc,"Annual review record and alignment note with organisation safety programme.")
    h(doc,2,"AAC.7.f — Training in safe practices")
    lb(doc,"Training-need identification by job description.")
    lb(doc,"Induction and annual safe-practice training records.")
    lb(doc,"MSDS training records for reagents handled by each role.")
    h(doc,2,"AAC.7.g — Appropriate safety measures provided")
    lb(doc,"PPE and safety-equipment inventory (including spill kit and extinguishers).")
    lb(doc,"Immunisation records for laboratory staff.")
    lb(doc,"Observation or audit notes on adherence to standard precautions.")
    h(doc,1,"13. References")
    lb(doc,"National Accreditation Board for Hospitals and Healthcare Providers (NABH), Guidebook to "
           "Accreditation Standards for Hospitals, 6th Edition — Access, Assessment and Continuity of "
           "Care (AAC), standard AAC.7.")
    lb(doc,"ISO 15189:2022 — Medical laboratories — Requirements for quality and competence (guidance).")
    lb(doc,"NABL 112 — Specific criteria for accreditation of medical laboratories (guidance).")
    lb(doc,f"Internal documents of {HN}: laboratory QA manual; IQC/EQA records; calibration and "
            f"maintenance logs; laboratory safety manual and MSDS file; clinico-pathological meeting "
            f"minutes; organisation safety programme; occupational-health immunisation records.")
    h(doc,1,"Disclaimer")
    disclaimer(doc)
    save_and_verify(doc,"HCO_AAC_7_v2_REWRITE_DRAFT.docx")


# ══════════════════════════════════════════════════════════════════════════════
# AAC.8 — Imaging Services  [HAS STOP-WORK]
# ══════════════════════════════════════════════════════════════════════════════
def gen_aac8():
    doc=Document()
    h(doc,0,"Policy on Imaging Services")
    p(doc,HN)
    h(doc,1,"Document control")
    doc_ctrl(doc,"HCO/AAC/POL/08","Imaging In-Charge")
    p(doc,"A blank marked ________ must be completed before issue.")
    h(doc,1,"Statement of intent")
    p(doc,"Imaging services are provided as per the scope of services of the organisation — "
          "complying with legal requirements, available when clinical care needs them, reported on "
          "time, and quality-assured when outsourced.")
    h(doc,1,"1. Purpose")
    p(doc,f"This policy says how {HN} provides imaging services that comply with legal requirements "
          f"(including Atomic Energy Regulatory Board (AERB) clearance and Pre-Conception and "
          f"Pre-Natal Diagnostic Techniques (PC-PNDT) duties), keeps scope, infrastructure and "
          f"human resources commensurate with clinical services, uses qualified personnel to perform, "
          f"supervise and interpret investigations, delivers results within defined turnaround time "
          f"(TAT) while monitoring waiting, performance and report times, intimates critical results "
          f"within one hour with read-back, reports in a standardised manner (including teleradiology "
          f"naming), recalls or amends reports when needed, and outsources unavailable tests under a "
          f"quality-assured MoU.")
    p(doc,"The chapter intent is that imaging services are reliable, timely, legally compliant and "
          "continuous with patient care.")
    p(doc,"This policy does not cover imaging quality assurance, radiation safety programmes, "
          "patient screening before imaging, or informed consent for imaging procedures — those are "
          "covered in other hospital policies. The other AAC standards have their own policies too.")
    h(doc,1,"2. Scope")
    p(doc,f"This policy applies to the imaging/radiology in-charge, Radiation Safety Officer (RSO), "
          f"radiographers and technologists, reporting doctors (including teleradiology reporters), "
          f"treating doctors who request imaging, nurses who accompany patients, and the Quality "
          f"Coordinator at {HN}.")
    p(doc,"It covers the nine objective elements AAC.8.a-i: legal compliance; scope; infrastructure "
          "and human resources (with round-the-clock availability / safe transfer when outsourced); "
          "qualified performance, supervision and interpretation; TAT with time monitoring; critical "
          "results; standardised reporting; recall/amendment; outsourcing with MoU.")
    p(doc,"Boundaries:")
    lb(doc,"AAC.9 owns the imaging QA programme, radiation-safety programme, patient screening "
           "before imaging, radiation-safety devices, training and signage. This policy owns legal "
           "licences, service scope, TAT, critical intimation, reporting format and outsourcing MoUs.")
    lb(doc,"AAC.3.c owns appropriateness of accompanying staff during transfer; safe transfer of "
           "patients to outsourced imaging is required here as part of availability.")
    lb(doc,"ROM.6.e owns MoU quality content; this policy requires an MoU incorporating quality "
           "assurance for imaging outsourcing.")
    h(doc,1,"3. Policy standards")
    p(doc,f"{HN} provides imaging services that comply with legal and other requirements, including "
          f"AERB clearance, dosimeters, lead shields and aprons, PC-PNDT displays and reports, and "
          f"a Radiation Safety Officer of appropriate level. Scope, infrastructure and human resources "
          f"are commensurate with clinical services; imaging is available round the clock, with safe "
          f"transfer and timely reports when a modality is outsourced. Qualified and trained personnel "
          f"perform, supervise and interpret investigations. Results are available within defined TAT "
          f"with monitoring of waiting, performance and report times. Critical results are intimated "
          f"immediately and not later than one hour with documented read-back. Reports follow a "
          f"standardised format including teleradiology naming. There is a mechanism to recall or "
          f"amend reports. Tests not available in-house are outsourced only under a quality-assured MoU.")
    p(doc,f"{HN} does not operate radiation imaging without current legal compliance, and does not "
          f"delay a critical imaging result beyond one hour after the report is ready.")
    h(doc,1,"4. Non-negotiable rules")
    ln(doc,"Do not operate radiation-emitting imaging equipment without current AERB clearance and "
           "a named Radiation Safety Officer of appropriate level.")
    ln(doc,"Do not operate ultrasound used for prenatal diagnosis without current PC-PNDT "
           "registration, required displays and reports to the competent authority.")
    ln(doc,"Do not report an imaging examination without qualified interpretation as required "
           "for the modality.")
    ln(doc,"Do not leave a critical imaging result uncommunicated beyond one hour after the "
           "report is ready.")
    ln(doc,"Do not alter or modify content of an outsourced imaging report.")
    ln(doc,"Do not outsource imaging to a provider that has no documented quality assurance "
           "system and no MoU with this hospital.")
    ln(doc,"Do not leave a recalled report in clinical areas, the medical record, RIS or HIS "
           "after recall without replacement by the amended report.")
    ln(doc,"Do not proceed with imaging when required legal compliance checks for the modality "
           "are missing — stop-work applies.")
    ln(doc,f"Staff who see an imaging service rule broken report it the same shift to the "
            f"Imaging In-Charge, the Radiation Safety Officer or the Medical Superintendent.")
    h(doc,1,"5. What we do")
    h(doc,2,"5.1 Legal and other requirements")
    p(doc,f"Imaging services comply with legal and other requirements. {HN} is aware of legal and "
          f"other requirements for imaging, documents them for information and compliance, and "
          f"maintains and updates compliance status regularly.")
    p(doc,"Statutory requirements met include: AERB clearance; dosimeters; lead shields; lead aprons; "
          "signage; display as per the Pre-Conception and Pre-Natal Diagnostic Techniques (Prohibition "
          "of Sex Selection) Act, 1994; reports to the competent authority; and a Radiation Safety "
          "Officer (RSO) of appropriate level.")
    p(doc,"The imaging in-charge and RSO maintain a register of licences, registrations, RSO "
          "appointment and renewal dates. Renewal is initiated 60 days before expiry. Equipment is "
          "not used when clearance has lapsed — stop-work applies.")
    h(doc,2,"5.2 Scope of imaging services")
    p(doc,f"Scope of the imaging services is commensurate with the services provided by {HN}. "
          f"Example from the standard: an organisation providing neurosurgery services including "
          f"head injuries shall have facilities for CT scan.")
    p(doc,"Services are available round the clock so that patient care is not disrupted. Imaging "
          "modalities required for emergency management are preferably available within the premises. "
          "The imaging in-charge maintains a service menu of in-house and outsourced modalities, "
          "reviewed annually and whenever clinical services change.")
    h(doc,2,"5.3 Infrastructure and human resources")
    p(doc,"The infrastructure (physical and equipment) and human resources are adequate to provide "
          "for the defined scope of services. Imaging has adequate space and equipment; reports must "
          "not be delayed for lack of equipment or human resources, including personnel authorised "
          "to report results.")
    p(doc,"Round-the-clock availability is organised by roster for in-house modalities and by "
          "documented safe-transfer and turnaround arrangements for outsourced emergency modalities "
          "(see section 5.9). Staffing and authorised reporters are reviewed annually against workload.")
    h(doc,2,"5.4 Qualified performance, supervision and interpretation")
    p(doc,"Qualified and trained personnel perform, supervise and interpret the investigations. AERB "
          "guidelines are used as a reference document for radiation-based imaging.")
    p(doc,"The imaging in-charge keeps competency records for radiographers/technologists and the "
          "list of doctors authorised to interpret and report each modality, including teleradiology "
          "reporters. Unauthorised staff do not release reports.")
    h(doc,2,"5.5 Results within defined turnaround time")
    p(doc,"Imaging results are available within a defined time frame. The organisation documents TAT "
          "of imaging results for all modalities and monitors waiting times, time taken to perform "
          "the tests, and time taken to prepare the reports for all modalities.")
    p(doc,"The defined timeframes may differ for different tests and are decided based on the nature "
          "of the test, modality, criticality and urgency required by the treating doctor. The imaging "
          "in-charge documents the defined TAT for each modality in the imaging service menu or a "
          "separate TAT schedule.")
    p(doc,"The imaging in-charge reviews waiting, performance and report-time data monthly.")
    h(doc,2,"5.6 Critical imaging results")
    p(doc,"Critical results are intimated immediately to the personnel concerned. The organisation "
          "defines and documents critical results that require immediate attention for patient "
          "management for each modality. Critical results of outsourced investigations are also intimated.")
    p(doc,"Critical results are communicated to a person from the treating team at the earliest, "
          "and not later than one hour after completion of the test or the report being ready. The "
          "intimation includes: name of the patient; unique ID; date and time of imaging; "
          "investigation name and result; identity of who communicated; identity of the recipient; "
          "read-back; and date and time of acknowledgement.")
    h(doc,2,"5.7 Standardised reporting")
    p(doc,"Results are reported in a standardised manner. At a minimum the report includes: the name "
          "of the organisation (or outsourced imaging centre); the patient's name; the unique "
          "identification number; and the name and signature of the person reporting the result.")
    p(doc,"For teleradiology, the report includes the name of the reporting doctor and a remark to "
          "that effect, and the name of the reporting organisation if outsourced to an organisation. "
          "The hospital does not alter or modify anything in the outsourced report.")
    h(doc,2,"5.8 Recall and amendment of reports")
    p(doc,"There is a mechanism to address the recall or amendment of reports whenever applicable. "
          "When a report is recalled, withdrawal from clinical areas, medical records, RIS and HIS "
          "is ensured. If already issued to the patient, the amended report is made available with "
          "a caution to ignore the earlier one.")
    h(doc,2,"5.9 Outsourced imaging tests")
    p(doc,"Imaging tests not available in the organisation are outsourced to organisation(s) based "
          "on their quality assurance system. Written guidance includes: list of tests for outsourcing; "
          "identity of personnel in outsourced facilities to ensure safe transportation of patients "
          "and completion of imaging results; manner of identification of patients and test requisition "
          "details; methodology to check selection and performance of the outsourced facility; "
          "reporting of critical results; TAT for emergency and routine requests; and prioritisation "
          "for urgent investigations.")
    p(doc,"The organisation has an MoU/agreement that incorporates quality assurance and the "
          "requirements of this standard. The panel and MoUs are reviewed annually.")
    h(doc,1,"6. Stop-work authority")
    p(doc,"Do not scan or expose a patient when required legal compliance for the modality is "
          "missing — AERB clearance and RSO cover for radiation equipment; PC-PNDT registration "
          "and displays for covered ultrasound.")
    p(doc,"Do not use radiation-emitting equipment when mandatory personnel monitoring "
          "(dosimeter/TLD) or required lead protection for the examination is unavailable.")
    p(doc,"Stop-work applies to the imaging examination, not to emergency clinical stabilisation "
          "outside the imaging room.")
    p(doc,"The person who stops tells the Imaging In-Charge or the Radiation Safety Officer the "
          "same shift, and the Medical Superintendent if legal clearance is the barrier. Refusing "
          "to scan without required checks is not a disciplinary matter.")
    h(doc,1,"7. Governance and responsibility")
    gov_tbl(doc,[
        ("Medical Superintendent","Accountable that imaging services are legally compliant, "
         "available as scoped, and meet TAT and critical-result duties."),
        ("Imaging / Radiology In-Charge","Maintains service menu, licence register (with RSO), "
         "staffing, TAT definitions, critical-result list, report format, recall mechanism and "
         "outsourced panel with MoUs."),
        ("Radiation Safety Officer (RSO)","Holds AERB-facing compliance for radiation equipment, "
         "dosimetry and related statutory duties under this policy's legal section; coordinates "
         "with AAC.9 radiation-safety programme."),
        ("Radiographers / technologists","Perform examinations they are trained for; monitor "
         "times; communicate critical results when designated; do not alter outsourced reports."),
        ("Reporting doctors (including teleradiology)","Interpret and report within authorisation; "
         "support recall/amendment."),
        ("Treating doctors and accompanying nurses","Request appropriately; complete transfer "
         "safely when imaging is off-site; receive critical-result intimation with read-back."),
        ("Quality Coordinator","Audits this policy quarterly; tracks CAPA when TAT, "
         "critical-result or legal-compliance defects recur."),
    ])
    h(doc,1,"8. Quality monitoring (RCA → CAPA)")
    p(doc,"The Quality Coordinator audits this policy quarterly. What is monitored each quarter:")
    lb(doc,"AERB clearance, PC-PNDT registration, RSO appointment and dosimeter/lead-protection status current.")
    lb(doc,"Service menu commensurate with clinical services; round-the-clock arrangements documented.")
    lb(doc,"Infrastructure and staffing adequate; authorised reporters listed.")
    lb(doc,"Waiting, performance and report-time monitoring against defined TAT.")
    lb(doc,"Critical-result communication within one hour with read-back documentation.")
    lb(doc,"Standardised reports including teleradiology naming; no alteration of outsourced content.")
    lb(doc,"Recall/amendment log complete with RIS/HIS withdrawal evidence.")
    lb(doc,"Outsourced MoUs current and quality-based.")
    p(doc,"Root-cause analysis is required when the same imaging service defect recurs within six months.")
    p(doc,"This policy is reviewed annually, and sooner when modalities, legal registrations or "
          "outsourcing arrangements change.")
    h(doc,1,"9. Training and staff acknowledgement")
    p(doc,"All imaging staff, RSO, nurses who accompany patients to imaging and doctors who request "
          "or report imaging are trained on this policy at induction and once a year after that. "
          "Training covers legal compliance triggers for stop-work, TAT monitoring, critical-result "
          "intimation with read-back, standardised reporting, recall/amendment and outsourcing rules.")
    p(doc,"Staff acknowledgement")
    p(doc,f"I have read the Policy on Imaging Services of {HN}. I will follow the processes described.")
    p(doc,"Name: ___________________________    Designation: ___________________________")
    p(doc,"Department / floor: ____________________    Date: ____________")
    p(doc,"Signature: ___________________________")
    p(doc,"(One row per staff member. The Quality Coordinator holds signed acknowledgements with the induction record.)")
    h(doc,1,"10. Distribution")
    p(doc,f"Official master copy: office of the Medical Superintendent, {HN}, with the Quality Coordinator.")
    p(doc,"Copies issued to: imaging/radiology; RSO; emergency; every in-patient ward; out-patient; "
          "nursing administration; intensive-care areas if present.")
    p(doc,"The current version is available to all staff at the front-office policy file and, if "
          "the hospital keeps an intranet, at staff intranet / policies.")
    p(doc,"When a new version is issued, take old copies out of use.")
    h(doc,1,"11. Abbreviations")
    abbrev_tbl(doc,[
        ("AAC","Access, Assessment and Continuity of Care (NABH HCO chapter)"),
        ("AERB","Atomic Energy Regulatory Board"),
        ("CAPA","corrective and preventive action"),
        ("CT","computed tomography"),
        ("HCO","Hospital (Full Accreditation programme)"),
        ("HIS","Hospital Information System"),
        ("MoU","Memorandum of Understanding"),
        ("MRI","magnetic resonance imaging"),
        ("NABH","National Accreditation Board for Hospitals and Healthcare Providers"),
        ("OE","objective element"),
        ("PC-PNDT","Pre-Conception and Pre-Natal Diagnostic Techniques (Prohibition of Sex Selection) Act, 1994"),
        ("PRE","Patient Rights and Education"),
        ("RIS","Radiology Information System"),
        ("ROM","Responsibilities of Management"),
        ("RSO","Radiation Safety Officer"),
        ("TAT","turnaround time"),
        ("TLD","thermo-luminescent dosimeter"),
    ])
    h(doc,1,"12. Traceability to NABH HCO Full Accreditation 6th Edition AAC.8")
    p(doc,"This table is an index. It is not how the policy is organised. An asterisk in the Level "
          "column means documentation of the process is required.")
    trace_tbl(doc,[
        ("AAC.8.a","CORE","Imaging services comply with legal and other requirements.","Section 3; 5.1; Section 6 Stop-work","Imaging In-Charge and RSO (maintain compliance); Medical Superintendent (accountable)"),
        ("AAC.8.b","Commitment","Scope of the imaging services is commensurate to the services provided by the organisation.","Section 3; 5.2","Imaging In-Charge (maintain menu); Medical Superintendent (approve)"),
        ("AAC.8.c","Commitment","The infrastructure (physical and equipment) and human resources are adequate to provide for its defined scope of services.","Section 3; 5.3","Imaging In-Charge (manage); Medical Superintendent (resource)"),
        ("AAC.8.d","Commitment","Qualified and trained personnel perform, supervise and interpret the investigations.","Section 3; 5.4","Imaging In-Charge (competency records); reporting doctors (interpret); radiographers (perform)"),
        ("AAC.8.e","Commitment*","Imaging results are available within a defined time frame.","Section 3; 5.5","Imaging In-Charge (define and monitor); Quality Coordinator (audit)"),
        ("AAC.8.f","Commitment*","Critical results are intimated immediately to the personnel concerned.","Section 3; 5.6","Designated imaging personnel (intimate); treating team (read-back); Imaging In-Charge (critical list)"),
        ("AAC.8.g","Commitment","Results are reported in a standardised manner.","Section 3; 5.7","Imaging In-Charge (format); reporting doctor (sign); staff (no alteration of outsourced content)"),
        ("AAC.8.h","Commitment*","There is a mechanism to address the recall / amendment of reports whenever applicable.","Section 3; 5.8","Imaging In-Charge (run mechanism); Quality Coordinator (CAPA)"),
        ("AAC.8.i","Commitment","Imaging tests not available in the organisation are outsourced to organisation(s) based on their quality assurance system.","Section 3; 5.9","Imaging In-Charge (panel and MoU); Medical Superintendent (approve MoU)"),
    ])
    h(doc,1,"13. Required Records / Evidence Checklist")
    p(doc,"Records the hospital holds under this policy, listed by objective element.")
    h(doc,2,"AAC.8.a — Legal and other requirements")
    lb(doc,"Register of AERB clearances, dosimeters, lead protection and RSO appointment.")
    lb(doc,"PC-PNDT registration, displays and reports to competent authority where applicable.")
    lb(doc,"Periodic update of legal-compliance status.")
    h(doc,2,"AAC.8.b — Scope commensurate with services")
    lb(doc,"Imaging service menu listing in-house and outsourced modalities.")
    lb(doc,"Annual review of menu against clinical services.")
    lb(doc,"Round-the-clock availability / safe-transfer arrangement note.")
    h(doc,2,"AAC.8.c — Infrastructure and human resources adequate")
    lb(doc,"Equipment inventory linked to maintenance/QA under AAC.9.")
    lb(doc,"Staff and authorised-reporter list with shift cover.")
    lb(doc,"Annual adequacy review of space, equipment and human resources.")
    h(doc,2,"AAC.8.d — Qualified performance, supervision and interpretation")
    lb(doc,"Qualification and training records for imaging personnel.")
    lb(doc,"List of doctors authorised to interpret/report by modality (including teleradiology).")
    lb(doc,"Note of AERB guidance used as reference for radiation-based imaging.")
    h(doc,2,"AAC.8.e — Results within defined TAT")
    lb(doc,"Defined TAT for all modalities.")
    lb(doc,"Monitoring records of waiting, performance and report times for IP/OP/emergency.")
    lb(doc,"Monthly TAT exception/escalation records.")
    h(doc,2,"AAC.8.f — Critical results intimated immediately")
    lb(doc,"Documented critical-result definitions per modality.")
    lb(doc,"Critical-result communication log with patient ID, result, caller, recipient, read-back, date and time.")
    lb(doc,"List of personnel authorised to report critical imaging results.")
    h(doc,2,"AAC.8.g — Standardised reporting")
    lb(doc,"Standardised report template with minimum required fields.")
    lb(doc,"Teleradiology reports naming reporting doctor and organisation as required.")
    lb(doc,"Audit finding that outsourced report content was not altered.")
    h(doc,2,"AAC.8.h — Recall / amendment mechanism")
    lb(doc,"Recall/amendment log with reason and date.")
    lb(doc,"Evidence of withdrawal from clinical areas, medical records, RIS and HIS.")
    lb(doc,"Amended report issued to patient with caution where previously issued; CAPA record.")
    h(doc,2,"AAC.8.i — Outsourced imaging tests")
    lb(doc,"Written outsourcing guidance including critical-result and TAT/prioritisation rules.")
    lb(doc,"Current MoU/agreement incorporating quality assurance (AAC.3.c, ROM.6.e).")
    lb(doc,"Annual review of outsourced imaging performance.")
    h(doc,1,"14. References")
    lb(doc,"National Accreditation Board for Hospitals and Healthcare Providers (NABH), Guidebook to "
           "Accreditation Standards for Hospitals, 6th Edition — Access, Assessment and Continuity of "
           "Care (AAC), standard AAC.8.")
    lb(doc,"Atomic Energy (Radiation Protection) Rules, 2004 — licensing, radiation protection and "
           "related duties for radiation-emitting imaging.")
    lb(doc,"Pre-Conception and Pre-Natal Diagnostic Techniques (Prohibition of Sex Selection) Act, "
           "1994 — registration, display and reporting for covered prenatal diagnostic imaging.")
    lb(doc,"AERB guidelines — reference for radiation-based imaging practice (AAC.8.d).")
    lb(doc,f"Internal documents of {HN}: imaging service menu; licence and RSO register; "
            f"critical-result list; TAT and time-monitoring records; recall/amendment log; "
            f"outsourced imaging MoUs (AAC.3.c, ROM.6.e).")
    h(doc,1,"Disclaimer")
    aac8_p2 = (
        f"Several requirements in this document are statutory rather than advisory — in particular "
        f"those arising under the Atomic Energy (Radiation Protection) Rules, 2004, and the "
        f"Pre-Conception and Pre-Natal Diagnostic Techniques (Prohibition of Sex Selection) Act, 1994. "
        f"Statutory requirements change, and State authorities impose additional or stricter "
        f"conditions. {HN} is responsible for verifying the current text of any rule cited here "
        f"and the conditions attached to its own authorisations and licences; this document does "
        f"not constitute legal advice."
    )
    disclaimer(doc, aac8_p2)
    save_and_verify(doc,"HCO_AAC_8_v2_REWRITE_DRAFT.docx")


# ══════════════════════════════════════════════════════════════════════════════
# AAC.9 — Imaging Quality Assurance and Safety Programme  [HAS STOP-WORK]
# ══════════════════════════════════════════════════════════════════════════════
def gen_aac9():
    doc = Document()
    h(doc, 0, "Policy on Imaging Quality Assurance and Safety Programme")
    p(doc, HN)
    h(doc, 1, "Document control")
    doc_ctrl(doc, "HCO/AAC/POL/09", "Imaging In-Charge")
    h(doc, 1, "Statement of intent")
    p(doc, f"{HN} implements a comprehensive quality assurance and safety programme for imaging "
          f"services, meeting statutory AERB requirements, screening patients before radiation or "
          f"MRI exposure, using and testing radiation-safety devices, and training imaging and "
          f"ancillary staff in safety practices.")
    h(doc, 1, "1. Purpose")
    p(doc, f"This policy explains how {HN} maintains an imaging quality assurance and safety "
          f"programme: equipment and protocol quality assurance, appropriateness screening of "
          f"investigation orders, peer review of imaging results, clinico-radiological meetings, "
          f"CAPA on QA deviations, the radiation-safety programme and ALARA principle, pre-imaging "
          f"safety screening, radiation-safety device use and testing, staff training, and imaging signage.")
    p(doc, "Informed consent for contrast, sedation or interventional procedures is covered in "
          "the hospital's other policies. AERB clearance and PC-PNDT registration for imaging "
          "services are covered in the hospital's other policies.")
    h(doc, 1, "2. Scope")
    p(doc, f"This policy applies to the Imaging In-Charge, radiologists, radiographers, MRI/CT "
          f"technologists, ancillary imaging staff, and all patients and attendants entering imaging areas.")
    h(doc, 1, "3. Policy standards")
    p(doc, f"{HN} implements a comprehensive quality assurance programme addressing imaging "
          f"equipment, protocols, surveillance and safety, meeting statutory AERB requirements, "
          f"and including CAPA for deviations.")
    p(doc, "A radiation-safety programme is implemented following AERB guidelines and the ALARA "
          "principle. Patients and attendants are appropriately screened for safety and risk before "
          "imaging. Radiation-safety and monitoring devices are used and periodically tested. "
          "Imaging and ancillary personnel are trained in safety practices, and imaging signage "
          "is prominently displayed.")
    h(doc, 1, "4. Non-negotiable rules")
    lb(doc, "Do not proceed with radiation or MRI exposure when required safety/risk screening "
           "has not been completed and documented for the patient, or for any attendant entering "
           "the imaging area.")
    lb(doc, "Do not proceed with a radiation examination when a required radiation-safety or "
           "monitoring device has failed its periodic test — for example a cracked lead "
           "apron — until it is repaired or replaced.")
    lb(doc, "Do not proceed with a radiation examination when required TLD badges are not "
           "available for staff working with radiation sources.")
    lb(doc, "Do not order routine, unjustified radiation exposure — for example a routine "
           "daily chest radiograph for all ICU patients without clinical indication, especially "
           "in paediatric or neonatal ICUs.")
    lb(doc, "Do not display imaging signage inconsistently with regulatory requirements or omit "
           "it from procedure rooms and operation theatres where imaging equipment is used.")
    h(doc, 1, "5. What we do")
    h(doc, 2, "5.1 Implement the imaging quality assurance programme")
    p(doc, "A comprehensive quality assurance programme for imaging services is implemented, "
          "addressing equipment, protocols, surveillance and safety, and meeting statutory AERB "
          "requirements. Equipment tests are performed as applicable to the modalities present. "
          "Quality assurance for nuclear medicine and radiotherapy, where those services exist, "
          "follows AERB guidelines. Imaging protocols are reviewed to ensure optimum image quality "
          "with minimum possible dosage.")
    h(doc, 2, "5.2 Screen investigation orders for appropriateness")
    p(doc, "A system screens investigation orders before imaging or interventional procedures to "
          "ensure they are appropriate for the clinical indication, in terms of current best-practice "
          "guidelines and patient safety. If a request is not appropriate, alternate investigations "
          "are offered in consultation with the treating doctor. A record is maintained of instances "
          "where the requested test was modified after consultation.")
    h(doc, 2, "5.3 Conduct peer review of imaging results")
    p(doc, "A peer review system reviews imaging results of CT and MRI in a structured manner. "
          "Sample size and periodicity for each modality are defined by the organisation, at a "
          "minimum one percent. Discrepancies are graded on severity and impact on patient "
          "management, and CAPA to minimise recurrence is documented.")
    h(doc, 2, "5.4 Hold clinico-radiological meetings")
    p(doc, "Clinico-radiological meetings are conducted at pre-defined intervals to correlate "
          "imaging results — at a minimum CT and MRI — with referring clinicians, and "
          "are used to improve the quality of imaging results.")
    h(doc, 2, "5.5 Document CAPA for QA deviations")
    p(doc, "When deviations from the laid-down quality assurance programme are noted, the "
          "organisation institutes corrective and preventive action as appropriate, and this is "
          "documented as part of the QA programme.")
    h(doc, 2, "5.6 Implement the radiation-safety programme")
    p(doc, "The radiation-safety programme follows AERB guidelines and implements the As Low As "
          "Reasonably Achievable (ALARA) principle in all investigations involving radiation, "
          "including screening of patients at high risk for radiation. Routine unjustified radiation "
          "is avoided. CT protocols use the lowest exposure parameters that maintain image quality "
          "appropriate for the clinical indication. The programme is aligned with the "
          "organisation's safety programme.")
    h(doc, 2, "5.7 Screen patients for safety and risk before imaging")
    p(doc, "Patients in the child-bearing age group requiring radiation exposure are screened for "
          "pregnancy before the examination. Patients undergoing MRI are screened for any magnetic "
          "substance. Screening is documented and applies to attendants accompanying a patient or "
          "child into the imaging area. Informed consent is taken for contrast injection, "
          "moderate-deep sedation, interventional procedures, and higher-risk findings on risk screening.")
    h(doc, 2, "5.8 Use radiation-safety and monitoring devices")
    p(doc, "Imaging personnel and patients use appropriate radiation-safety and monitoring devices "
          "where applicable. Shielding of body parts of staff, patients and attendants uses "
          "appropriate aprons and shields, in numbers adequate for all workers. Staff directly "
          "working with radiation sources possess and use thermo-luminescent dosimeter (TLD) badges.")
    h(doc, 2, "5.9 Test radiation-safety devices periodically")
    p(doc, "Radiation-safety and monitoring devices are periodically tested and results documented. "
          "Protective devices such as lead aprons are exposed to X-ray, fluoroscopy or CT scout "
          "view to verify cracks and damage; images of the checks are stored. CAPA is taken and "
          "documented where a device test result warrants it. TLD badges are replaced by fresh "
          "badges from accredited laboratories at the frequency recommended by AERB.")
    h(doc, 2, "5.10 Train imaging and ancillary personnel")
    p(doc, "Imaging personnel are trained in imaging safety practices — such as MRI safety, "
          "tube kinking, fall prevention and patient handling — and radiation-safety measures. "
          "Ancillary staff posted in imaging, including nurses, helpers, stretcher-bearers, "
          "housekeeping and security as applicable, are also trained.")
    h(doc, 2, "5.11 Display imaging signage")
    p(doc, "Imaging signage, including safety signage and signage required by regulatory "
          "authorities, is prominently displayed in all appropriate locations — including "
          "procedure rooms and operation theatres where imaging equipment is used.")
    h(doc, 1, "6. Stop-work authority")
    p(doc, "Do not proceed with radiation or MRI exposure when required safety/risk screening has "
          "not been completed and documented for the patient (and for any attendant entering the imaging area).")
    p(doc, "As hospital safety practice (not an NABH objective-element requirement), do not "
          "proceed when pregnancy status is positive or uncertain for a radiation examination "
          "until the hospital radiologist or imaging in-charge has reviewed and "
          "documented the decision. Do not proceed with MRI when magnetic-substance screening "
          "is incomplete or positive for a contraindication.")
    p(doc, "Do not proceed when a required radiation-safety or monitoring device has failed its "
          "periodic test (for example a cracked lead apron) or when required TLD badges are not "
          "available for staff working with radiation sources.")
    p(doc, "Stop-work applies to the imaging exposure or MRI examination, not to emergency "
          "clinical stabilisation outside the imaging room.")
    p(doc, "The person who stops tells the Imaging In-Charge or the "
          "Radiation Safety Officer the same shift. Refusing to expose without "
          "screening or with failed safety devices is not a disciplinary matter.")
    h(doc, 1, "7. Governance and responsibility")
    gov_tbl(doc, [
        ("Medical Superintendent",
         "Accountable for the imaging quality assurance and safety programme; receives stop-work escalations."),
        ("Imaging In-Charge",
         "Owns implementation of this policy; maintains the QA programme, appropriateness screening "
         "system, peer review, and clinico-radiological meeting schedule."),
        ("Radiation Safety Officer",
         "Oversees the radiation-safety programme, ALARA implementation, device testing, and TLD "
         "badge management; receives stop-work escalations."),
        ("Radiologists/Radiographers/Technologists",
         "Perform pre-imaging screening, use safety devices, and follow appropriateness and ALARA protocols."),
        ("Quality Coordinator",
         "Audits this policy; holds training records, CAPA logs, and stop-work event logs."),
    ])
    h(doc, 1, "8. Quality monitoring")
    mon9 = [
        ("QA programme",
         "Equipment test records; protocol review file; nuclear medicine/radiotherapy QA records where applicable."),
        ("Appropriateness screening", "Modification log; consultation records with treating doctors."),
        ("Peer review", "Discrepancy grading records; CAPA for recurring findings."),
        ("Clinico-radiological meetings", "Meeting minutes; quality improvement actions arising."),
        ("CAPA on deviations", "CAPA log with root cause, owner, and closure evidence."),
        ("Radiation-safety programme", "ALARA protocol records; unjustified-exposure avoidance records."),
        ("Pre-imaging screening", "Pregnancy and magnetic-substance screening records; consent records."),
        ("Safety device use and testing",
         "Apron/shield inventory; TLD issue and wear records; periodic test images and results."),
        ("Training", "Induction and periodic training records for imaging and ancillary staff."),
        ("Signage", "Signage location checklist; periodic visibility check log."),
    ]
    mt = tbl(doc, len(mon9)+1, 2)
    mt.cell(0,0).text="Monitoring area"; mt.cell(0,1).text="What is monitored"
    for ri,(a,b) in enumerate(mon9, 1): mt.cell(ri,0).text=a; mt.cell(ri,1).text=b
    h(doc, 1, "9. Training and staff acknowledgement")
    p(doc, "All imaging and ancillary personnel shall be trained in imaging safety practices and "
          "radiation-safety measures, including the stop-work triggers in this policy, at induction and periodically.")
    p(doc, f"I have read the Policy on Imaging Quality Assurance and Safety Programme of {HN} "
          f"and understand my responsibilities under it.")
    h(doc, 1, "10. Distribution")
    p(doc, "This policy shall be available to all imaging and ancillary staff, the Radiation "
          "Safety Officer, the Imaging In-Charge, and the Quality Coordinator.")
    h(doc, 1, "11. Abbreviations")
    abbrev_tbl(doc, [
        ("AERB",  "Atomic Energy Regulatory Board"),
        ("ALARA", "As Low As Reasonably Achievable"),
        ("CAPA",  "Corrective and Preventive Action"),
        ("CT",    "Computed Tomography"),
        ("MRI",   "Magnetic Resonance Imaging"),
        ("NABH",  "National Accreditation Board for Hospitals and Healthcare Providers"),
        ("TLD",   "Thermo-Luminescent Dosimeter"),
    ])
    h(doc, 1, "12. Traceability table")
    p(doc, "This table is an index. It is not how the policy is organised. An asterisk in the "
          "Level column means the objective element is starred and requires documented evidence.")
    tr9 = [
        ("AAC.9.a", "Commitment*", "Sections 3 and 5.1 address the imaging QA programme."),
        ("AAC.9.b", "Achievement",  "Section 5.2 addresses appropriateness screening of investigation orders."),
        ("AAC.9.c", "Achievement",  "Section 5.3 addresses peer review of imaging results."),
        ("AAC.9.d", "Excellence",   "Section 5.4 addresses clinico-radiological meetings."),
        ("AAC.9.e", "Commitment*",  "Section 5.5 addresses CAPA documentation for QA deviations."),
        ("AAC.9.f", "Commitment*",  "Sections 3 and 5.6 address the radiation-safety programme and ALARA principle."),
        ("AAC.9.g", "Commitment",   "Sections 3, 5.7 and 6 address pre-imaging safety screening and the related stop-work trigger."),
        ("AAC.9.h", "Commitment",   "Section 5.8 addresses use of radiation-safety and monitoring devices."),
        ("AAC.9.i", "Commitment*",  "Sections 5.9 and 6 address periodic testing of safety devices and the related stop-work trigger."),
        ("AAC.9.j", "Commitment",   "Section 5.10 addresses training of imaging and ancillary personnel."),
        ("AAC.9.k", "Commitment",   "Section 5.11 addresses imaging signage."),
    ]
    tt = tbl(doc, len(tr9)+1, 3)
    tt.cell(0,0).text="Objective Element"; tt.cell(0,1).text="Level"
    tt.cell(0,2).text="Traceability to this policy"
    for ri,(a,b,c) in enumerate(tr9, 1): tt.cell(ri,0).text=a; tt.cell(ri,1).text=b; tt.cell(ri,2).text=c
    h(doc, 1, "13. Required records/evidence checklist")
    h(doc, 2, "QA programme — AAC.9.a (Commitment*)")
    lb(doc, "QA programme document, equipment test records, protocol review file.")
    lb(doc, "Nuclear medicine/radiotherapy QA records where applicable.")
    h(doc, 2, "Appropriateness screening — AAC.9.b (Achievement)")
    lb(doc, "Appropriateness screening method or checklist.")
    lb(doc, "Modification log and consultation records with treating doctors.")
    h(doc, 2, "Peer review — AAC.9.c (Achievement)")
    lb(doc, "Peer-review records with discrepancy grading and CAPA.")
    lb(doc, "Discrepancy meeting notes where held.")
    h(doc, 2, "Clinico-radiological meetings — AAC.9.d (Excellence)")
    lb(doc, "Meeting schedule and minutes recording correlations with clinicians.")
    lb(doc, "Documented quality improvement actions.")
    h(doc, 2, "CAPA on QA deviations — AAC.9.e (Commitment*)")
    lb(doc, "CAPA log for QA and radiation-safety deviations with root-cause notes, owners, "
           "due dates and closure evidence.")
    h(doc, 2, "Radiation-safety programme — AAC.9.f (Commitment*)")
    lb(doc, "Radiation-safety programme document aligned with AERB and organisation safety.")
    lb(doc, "ALARA protocol examples and unjustified-exposure avoidance records.")
    h(doc, 2, "Pre-imaging screening — AAC.9.g (Commitment)")
    lb(doc, "Documented pregnancy screening and MRI magnetic-substance screening.")
    lb(doc, "Consent records for contrast/sedation/interventional procedures.")
    h(doc, 2, "Safety device use — AAC.9.h (Commitment)")
    lb(doc, "Apron and shield inventory adequate for all workers.")
    lb(doc, "TLD badge issue and wear records.")
    h(doc, 2, "Safety device testing — AAC.9.i (Commitment*)")
    lb(doc, "Periodic lead-apron/shield test images and results.")
    lb(doc, "TLD replacement records per AERB-recommended frequency.")
    lb(doc, "CAPA/removal-from-use records for failed devices.")
    h(doc, 2, "Staff training — AAC.9.j (Commitment)")
    lb(doc, "Induction and periodic training records for imaging and ancillary staff, with "
           "content and attendance lists.")
    h(doc, 2, "Imaging signage — AAC.9.k (Commitment)")
    lb(doc, "Signage location checklist including procedure rooms and OTs with imaging equipment.")
    lb(doc, "Periodic signage visibility check log.")
    h(doc, 1, "14. References")
    ln(doc, "National Accreditation Board for Hospitals and Healthcare Providers. NABH "
           "Accreditation Standards for Hospitals, 6th Edition. AAC.9.")
    ln(doc, "Guidebook interpretation supplied for AAC.9.a through AAC.9.k.")
    ln(doc, f"Atomic Energy (Radiation Protection) Rules, 2004 — applicable to {HN} to "
           f"the extent its imaging services fall within the scope of these rules.")
    ln(doc, f"Internal documents of {HN}: QA programme document; radiation-safety programme "
           f"document; screening and consent records; device test and TLD records; training "
           f"records; signage checklist.")
    h(doc, 1, "Disclaimer")
    p(doc, f"This policy reorganises the supplied AAC.9 objective-element wording and Guidebook "
          f"interpretation into plain-language guidance for {HN}. It does not replace the NABH "
          f"Accreditation Standards for Hospitals, 6th Edition, or the Guidebook, which remain "
          f"the authoritative source in case of any conflict.")
    p(doc, f"This policy references the Atomic Energy (Radiation Protection) Rules, 2004, as "
          f"the statutory framework for radiation-safety duties to the extent {HN}’s imaging "
          f"services fall within its scope. This policy is not a statement of, or substitute for, "
          f"compliance with that or any other statute; {HN} remains separately responsible for "
          f"meeting all applicable statutory and regulatory requirements, including AERB clearance "
          f"and monitoring obligations.")
    p(doc, "This policy shall be reviewed at least once every year, or earlier if NABH standards "
          "or applicable AERB regulations change.")
    p(doc, f"This policy is the property of {HN} and is not to be reproduced or distributed "
          f"outside the organisation without authorisation.")
    save_and_verify(doc, "HCO_AAC_9_v2_REWRITE_DRAFT.docx")


# ══════════════════════════════════════════════════════════════════════════════
# AAC.10 — Continuous and Multi-Disciplinary Patient Care
# ══════════════════════════════════════════════════════════════════════════════
def gen_aac10():
    doc = Document()
    h(doc, 0, "Policy on Continuous and Multi-Disciplinary Patient Care")
    p(doc, HN)
    h(doc, 1, "Document control")
    doc_ctrl(doc, "HCO/AAC/POL/10", "Medical Superintendent")
    h(doc, 1, "Statement of intent")
    p(doc, f"{HN} ensures patient care is continuous and multi-disciplinary: a named responsible "
          f"doctor at all times, coordinated care across settings, shared clinical information, "
          f"standardised hand-over communication, safe internal transfer, written referral guidance, "
          f"predictable service delivery, and monitoring of the response to critical value alerts.")
    h(doc, 1, "1. Purpose")
    p(doc, f"This policy explains how {HN} maintains continuity and coordination of patient care "
          f"across all care settings: naming a responsible doctor, coordinating care, sharing "
          f"clinical information among providers, standardising hand-over, transferring patients "
          f"safely, referring within the organisation, communicating schedule changes, and "
          f"monitoring the response to critical value alerts.")
    p(doc, "Multi-disciplinary care-plan discussion and transfer to another facility are covered "
          "in the hospital's other policies.")
    h(doc, 1, "2. Scope")
    p(doc, f"This policy applies to all doctors and nurses providing direct patient care in "
          f"out-patient, emergency, in-patient, intensive care, and any other care setting at {HN}.")
    h(doc, 1, "3. Policy standards")
    p(doc, f"{HN} identifies a qualified individual responsible for each patient's care during "
          f"all phases of care, coordinates care among providers in each setting, and shares "
          f"information about care and response to treatment among medical, nursing and other care providers.")
    p(doc, "Standardised hand-over communication is implemented during shifts, between shifts, "
          "and during transfers; patient transfer within the organisation is done safely; referral "
          "to other departments follows written guidance; service delivery is predictable and "
          "schedule changes are communicated; and a mechanism monitors the clinical response to "
          "critical value alerts.")
    h(doc, 1, "4. Non-negotiable rules")
    lb(doc, "Do not leave a patient without a named responsible doctor in any phase or setting of care.")
    lb(doc, "Do not allow a hand-over between shifts, units or departments to proceed "
           "undocumented, or without acknowledgement by the receiving clinician.")
    lb(doc, "Do not transfer a patient within the organisation without documented hand-over and "
           "takeover, or without defined accompaniment and monitoring for an unstable patient.")
    lb(doc, "Do not refer a patient to another department or speciality outside the organisation's "
           "written referral guidance on grades, timeframes and who may accept a referral.")
    lb(doc, "Do not leave a critical value alert without a documented, monitored clinical response.")
    h(doc, 1, "5. What we do")
    h(doc, 2, "5.1 Name a responsible doctor for every patient")
    p(doc, "During all phases of care, a qualified doctor is identified as responsible for the "
          "patient's care. The responsible doctor's name is recorded in the patient record at "
          "first contact and updated when responsibility changes. When the named doctor is off "
          "duty, the covering doctor is identified and the change is recorded.")
    h(doc, 2, "5.2 Coordinate care across settings")
    p(doc, "Patient care is coordinated among care providers in each care setting — "
          "out-patient, emergency, in-patient, intensive care, and any other setting the "
          "organisation uses. The patient's requirements are communicated to providers in the "
          "receiving setting so that orders, observations and pending tasks are not lost when "
          "the patient moves.")
    h(doc, 2, "5.3 Share care and response information among providers")
    p(doc, "Information about the patient's care and response to treatment is shared among "
          "medical, nursing and other care providers through entries in the case sheet or "
          "electronic medical record. The patient record remains available to authorised providers.")
    h(doc, 2, "5.4 Implement standardised hand-over communication")
    p(doc, "Standardised hand-over communication is implemented during each staffing shift, "
          "between shifts, and during transfers between units or departments. Hand-over by doctors "
          "and nurses in direct patient care is documented and covers the patient's current "
          "condition, recent changes, ongoing treatment, and possible changes or complications. "
          "The receiving clinician acknowledges the hand-over, and interruptions during "
          "hand-over are minimised.")
    h(doc, 2, "5.5 Refer within the organisation following written guidance")
    p(doc, "Referral to other departments or specialities within the organisation follows written "
          "guidance. Referral may be for opinion, co-management or takeover, and the referral note "
          "states the reason. Urgency is graded, and every referral is seen within a defined "
          "timeframe per grade. Written guidance covers grades, timeframes and who may accept a referral.")
    h(doc, 2, "5.6 Transfer patients within the organisation safely")
    p(doc, "Patients are transported in a safe and timely manner within the organisation. Proper "
          "hand-over and takeover are documented. Transfers to and from imaging follow the same "
          "rule so that imaging is not delayed by transport failure. Unstable patients are "
          "transferred with the accompaniment and monitoring defined for their clinical status.")
    h(doc, 2, "5.7 Ensure predictable service delivery")
    p(doc, "The organisation ensures predictable service delivery by adhering to defined timelines "
          "for services such as laboratory, radiology, and OPD waiting time. Patients are informed "
          "of the applicable timelines. When delivery will deviate, the patient, family and/or "
          "caregiver is informed of the changed schedule, and the service area records that the "
          "change was communicated.")
    h(doc, 2, "5.8 Monitor response to critical value alerts")
    p(doc, "A mechanism monitors whether an adequate clinical intervention has taken place in "
          "response to a critical value alert. The attending clinician's response is documented "
          "in progress notes or medication orders. Interventions are periodically reviewed for "
          "timeliness and appropriateness. For out-patients, efforts are taken to alert the "
          "patient or family about critical values.")
    h(doc, 1, "6. Stop-work authority")
    p(doc, "Not applicable. This standard has no stop-work authority. Gaps in continuity of care "
          "— a missing responsible doctor, an undocumented hand-over, an uncoordinated "
          "transfer, or an unmonitored critical value response — are addressed through the "
          "hospital's quality monitoring and CAPA process (Section 8), not through an immediate "
          "stop-work trigger.")
    h(doc, 1, "7. Governance and responsibility")
    gov_tbl(doc, [
        ("Medical Superintendent",
         "Accountable for continuity and coordination of patient care across all settings."),
        ("Treating/Responsible Doctor",
         "Named as responsible for the patient's care; ensures hand-over and referral guidance "
         "is followed; responds to critical value alerts."),
        ("Nursing Superintendent",
         "Ensures standardised hand-over and safe internal transfer practice among nursing staff."),
        ("Quality Coordinator",
         "Audits this policy; holds hand-over, referral, transfer, and critical-value-response records."),
    ])
    h(doc, 1, "8. Quality monitoring")
    mon10 = [
        ("Responsible doctor",
         "Patient record entries naming the responsible doctor across care settings; covering-doctor records."),
        ("Care coordination",
         "Communication or transfer notes when patients move across care settings."),
        ("Information sharing",
         "Case sheet or EMR entries by multiple provider types on the same patient."),
        ("Hand-over",
         "Documented hand-over records for each staffing shift and transfer, with receiving-clinician acknowledgement."),
        ("Internal referral",
         "Referral notes with reason and grade; response-time log against defined timeframes."),
        ("Internal transfer",
         "Internal transfer records with hand-over/takeover; unstable-patient accompaniment and monitoring records."),
        ("Schedule predictability",
         "Defined service timelines; records of schedule-change communication to patient/family/caregiver."),
        ("Critical value response",
         "Critical value alert log; documented clinical intervention; periodic review of timeliness and appropriateness."),
    ]
    mt = tbl(doc, len(mon10)+1, 2)
    mt.cell(0,0).text="Monitoring area"; mt.cell(0,1).text="What is monitored"
    for ri,(a,b) in enumerate(mon10, 1): mt.cell(ri,0).text=a; mt.cell(ri,1).text=b
    h(doc, 1, "9. Training and staff acknowledgement")
    p(doc, "All doctors and nurses providing direct patient care shall be trained on the "
          "responsible-doctor rule, standardised hand-over, safe internal transfer, and internal "
          "referral guidance in this policy, at induction and periodically.")
    p(doc, f"I have read the Policy on Continuous and Multi-Disciplinary Patient Care of {HN} "
          f"and understand my responsibilities under it.")
    h(doc, 1, "10. Distribution")
    p(doc, "This policy shall be available to all clinical and nursing staff, department heads, "
          "and the Quality Coordinator.")
    h(doc, 1, "11. Abbreviations")
    abbrev_tbl(doc, [
        ("CAPA", "Corrective and Preventive Action"),
        ("EMR",  "Electronic Medical Record"),
        ("ICU",  "Intensive Care Unit"),
        ("NABH", "National Accreditation Board for Hospitals and Healthcare Providers"),
        ("OPD",  "Out-Patient Department"),
    ])
    h(doc, 1, "12. Traceability table")
    p(doc, "This table is an index. It is not how the policy is organised. An asterisk in the "
          "Level column means the objective element is starred and requires documented evidence.")
    tr10 = [
        ("AAC.10.a", "Commitment",  "Sections 3 and 5.1 address naming a responsible doctor for all phases of care."),
        ("AAC.10.b", "Commitment",  "Sections 3 and 5.2 address coordination of care across settings."),
        ("AAC.10.c", "Commitment",  "Sections 3 and 5.3 address sharing of care and response information among providers."),
        ("AAC.10.d", "CORE",        "Sections 3 and 5.4 address standardised hand-over communication."),
        ("AAC.10.e", "Commitment",  "Sections 3 and 5.6 address safe internal patient transfer."),
        ("AAC.10.f", "Commitment*", "Section 5.5 addresses written referral guidance within the organisation."),
        ("AAC.10.g", "Commitment",  "Sections 3 and 5.7 address predictable service delivery and schedule-change communication."),
        ("AAC.10.h", "Commitment",  "Sections 3 and 5.8 address monitoring the response to critical value alerts."),
    ]
    tt = tbl(doc, len(tr10)+1, 3)
    tt.cell(0,0).text="Objective Element"; tt.cell(0,1).text="Level"
    tt.cell(0,2).text="Traceability to this policy"
    for ri,(a,b,c) in enumerate(tr10, 1): tt.cell(ri,0).text=a; tt.cell(ri,1).text=b; tt.cell(ri,2).text=c
    h(doc, 1, "13. Required records/evidence checklist")
    h(doc, 2, "Responsible doctor — AAC.10.a (Commitment)")
    lb(doc, "Patient record entry naming the responsible doctor for each phase of care.")
    lb(doc, "Covering-doctor record when the named doctor is off duty.")
    h(doc, 2, "Care coordination — AAC.10.b (Commitment)")
    lb(doc, "Communication or transfer notes when patients move across care settings.")
    h(doc, 2, "Information sharing — AAC.10.c (Commitment)")
    lb(doc, "Case sheet or EMR entries by multiple provider types on the same patient.")
    lb(doc, "Record of team meetings where used.")
    h(doc, 2, "Hand-over — AAC.10.d (CORE)")
    lb(doc, "Documented hand-over records using the standardised framework for each staffing shift.")
    lb(doc, "Transfer hand-over with acknowledgement by the receiving team.")
    h(doc, 2, "Internal transfer — AAC.10.e (Commitment)")
    lb(doc, "Internal transfer record with reason, clinical status, times, hand-over and takeover.")
    lb(doc, "Imaging transfer records; unstable-patient transfer records with accompaniment and monitoring.")
    h(doc, 2, "Internal referral guidance — AAC.10.f (Commitment*)")
    lb(doc, "Current written guidance for internal referral (opinion/co-management/takeover; "
           "urgency grades; timeframes).")
    lb(doc, "Referral notes with reason and grade; response-time log.")
    h(doc, 2, "Predictable service delivery — AAC.10.g (Commitment)")
    lb(doc, "Defined service timelines; record that patients were informed of timelines.")
    lb(doc, "Record of schedule-change communication to patient/family/caregiver.")
    h(doc, 2, "Critical value response — AAC.10.h (Commitment)")
    lb(doc, "Critical value alert log with time and clinician notified.")
    lb(doc, "Progress notes or orders documenting the clinical intervention.")
    lb(doc, "Periodic review records of response timeliness and appropriateness.")
    h(doc, 1, "14. References")
    ln(doc, "National Accreditation Board for Hospitals and Healthcare Providers. NABH "
           "Accreditation Standards for Hospitals, 6th Edition. AAC.10.")
    ln(doc, "Guidebook interpretation supplied for AAC.10.a through AAC.10.h.")
    ln(doc, f"Internal documents of {HN}: hand-over framework; internal referral guidance; "
           f"transfer records; critical value alert log.")
    h(doc, 1, "Disclaimer")
    p(doc, f"This policy reorganises the supplied AAC.10 objective-element wording and Guidebook "
          f"interpretation into plain-language guidance for {HN}. It does not replace the NABH "
          f"Accreditation Standards for Hospitals, 6th Edition, or the Guidebook, which remain "
          f"the authoritative source in case of any conflict.")
    p(doc, f"This policy is intended for accreditation and internal governance purposes and is "
          f"not a statement of, or substitute for, compliance with any specific statute; {HN} "
          f"remains separately responsible for meeting all applicable statutory and regulatory requirements.")
    p(doc, "This policy shall be reviewed at least once every year, or earlier if NABH standards change.")
    p(doc, f"This policy is the property of {HN} and is not to be reproduced or distributed "
          f"outside the organisation without authorisation.")
    save_and_verify(doc, "HCO_AAC_10_v2_REWRITE_DRAFT.docx")


# ══════════════════════════════════════════════════════════════════════════════
# AAC.11 — Preventive and Promotive Health Services
# ══════════════════════════════════════════════════════════════════════════════
def gen_aac11():
    doc = Document()
    h(doc, 0, "Policy on Preventive and Promotive Health Services")
    p(doc, HN)
    h(doc, 1, "Document control")
    doc_ctrl(doc, "HCO/AAC/POL/11", "Medical Superintendent")
    h(doc, 1, "Statement of intent")
    p(doc, f"{HN} provides preventive and promotive health services within its scope of services: "
          f"written guidance for preventive and promotive care, age-appropriate screening for "
          f"non-communicable diseases, mental health screening, immunisation, and multi-disciplinary "
          f"health education on lifestyle modifications.")
    h(doc, 1, "1. Purpose")
    p(doc, f"This policy explains how {HN} provides preventive and promotive health services "
          f"within its defined scope: written guidance governing implementation, non-communicable "
          f"disease screening, mental health screening, paediatric and adult immunisation, and "
          f"multi-disciplinary lifestyle health education.")
    p(doc, f"A hospital that does not offer a particular preventive or promotive service — such "
          f"as community outreach, a specific immunisation, or a specific screening type — does "
          f"not claim that service; the written guidance under Section 5.1 names what is in scope.")
    h(doc, 1, "2. Scope")
    p(doc, f"This policy applies to treating doctors, nurses, and any other clinical staff "
          f"involved in delivering preventive and promotive health services within the scope of "
          f"services defined by {HN}.")
    h(doc, 1, "3. Policy standards")
    p(doc, f"Written guidance governs the implementation of preventive and promotive care as per "
          f"{HN}'s scope of services. The organisation defines evidence-based, age-appropriate "
          f"screening for non-communicable diseases and advises mental health screening and "
          f"appropriate intervention wherever applicable.")
    p(doc, "Evidence-based paediatric and adult immunisation is advised wherever applicable, "
          "and a multi-disciplinary approach is adopted in imparting health education on lifestyle modifications.")
    h(doc, 1, "4. Non-negotiable rules")
    lb(doc, "Do not deliver preventive or promotive care without written guidance defining what "
           "is in scope, who delivers it, and how it is documented.")
    lb(doc, f"Do not claim or promise a preventive or promotive service that falls outside "
           f"{HN}'s defined scope of services.")
    lb(doc, "Do not advise non-communicable disease screening that is not evidence-based or age-appropriate.")
    lb(doc, "Do not adhere to paediatric immunisation outside the Universal Immunisation "
           "Programme without a documented evidence-based reason.")
    h(doc, 1, "5. What we do")
    h(doc, 2, "5.1 Maintain written guidance for preventive and promotive care")
    p(doc, f"Written guidance governs the implementation of preventive and promotive care as per "
          f"{HN}'s scope of services. Preventive care focuses on measures to prevent the onset of "
          f"diseases or health conditions — for example immunisation, health screening, health "
          f"education and lifestyle counselling, based on age, gender and medical history. "
          f"Promotive care involves efforts to promote overall health and well-being — for "
          f"example healthy diet, regular exercise, stress management, mental health support, "
          f"and patient and family education. The guidance names which preventive and promotive "
          f"services are in scope, who delivers them, and how they are documented. Services "
          f"outside scope are not promised or claimed.")
    h(doc, 2, "5.2 Screen for non-communicable diseases")
    p(doc, "The organisation defines evidence-based and contextual age-appropriate screening for "
          "non-communicable diseases — for example screening for various cancers, diabetes and "
          "osteoporosis. Frequency of screening increases when predisposing factors apply. "
          "Treating doctors document the screening advised, accepted or declined, and its basis "
          "(age, risk factors, guidance reference).")
    h(doc, 2, "5.3 Advise mental health screening")
    p(doc, "Mental health screening and appropriate intervention are advised for patients wherever "
          "applicable, using validated tools where available. A positive screen leads to documented "
          "advice, referral or intervention appropriate to the finding and to the hospital's scope.")
    h(doc, 2, "5.4 Advise immunisation")
    p(doc, "Evidence-based and contextual paediatric and adult immunisation is advised wherever "
          "applicable. The Universal Immunisation Programme is adhered to for all neonates and "
          "paediatric patients, including catch-up immunisation. Adult immunisation, where offered, "
          "follows evidence-based guidelines. Advice, administration and refusals are recorded "
          "in the patient record or immunisation register.")
    h(doc, 2, "5.5 Deliver multi-disciplinary lifestyle health education")
    p(doc, "A multi-disciplinary approach is adopted in imparting health education on lifestyle "
          "modifications, with contributors including the clinician, nurse, dietician, "
          "physiotherapist, occupational therapist, psychologist and medico-social worker where "
          "those roles exist. Family members are also educated towards cessation of tobacco, "
          "alcohol and substance abuse where relevant. Education given and disciplines involved are recorded.")
    h(doc, 1, "6. Stop-work authority")
    p(doc, "Not applicable. This standard has no stop-work authority. Gaps in preventive or "
          "promotive care delivery are addressed through the hospital's quality monitoring and "
          "CAPA process (Section 8), not through an immediate stop-work trigger.")
    h(doc, 1, "7. Governance and responsibility")
    gov_tbl(doc, [
        ("Medical Superintendent",
         "Accountable for defining the scope of preventive and promotive services and ensuring "
         "written guidance exists."),
        ("Treating Doctor",
         "Advises and documents NCD screening, mental health screening, and immunisation as "
         "applicable to each patient."),
        ("Nursing Superintendent",
         "Supports delivery of preventive and promotive care in nursing practice; maintains immunisation records."),
        ("Quality Coordinator",
         "Audits this policy; holds written guidance, screening, and education records."),
    ])
    h(doc, 1, "8. Quality monitoring")
    mon11 = [
        ("Written guidance",
         "Current written guidance for preventive and promotive care aligned to scope; "
         "periodic audit confirming practice matches guidance."),
        ("NCD screening",
         "Defined NCD screening list; patient record entries of screening advised, accepted or declined."),
        ("Mental health screening",
         "Documented use of validated screening tools; records of intervention, advice or referral "
         "after a positive screen."),
        ("Immunisation",
         "Paediatric immunisation records aligned to UIP including catch-up; adult immunisation "
         "records where offered."),
        ("Lifestyle education",
         "Patient education records naming disciplines involved; periodic sample showing "
         "multi-disciplinary input where available."),
    ]
    mt = tbl(doc, len(mon11)+1, 2)
    mt.cell(0,0).text="Monitoring area"; mt.cell(0,1).text="What is monitored"
    for ri,(a,b) in enumerate(mon11, 1): mt.cell(ri,0).text=a; mt.cell(ri,1).text=b
    h(doc, 1, "9. Training and staff acknowledgement")
    p(doc, "All treating doctors, nurses and clinical staff involved in preventive and promotive "
          "care shall be trained on the written guidance, screening protocols, and immunisation "
          "requirements in this policy, at induction and periodically.")
    p(doc, f"I have read the Policy on Preventive and Promotive Health Services of {HN} and "
          f"understand my responsibilities under it.")
    h(doc, 1, "10. Distribution")
    p(doc, "This policy shall be available to all clinical staff, department heads, and the Quality Coordinator.")
    h(doc, 1, "11. Abbreviations")
    abbrev_tbl(doc, [
        ("CAPA", "Corrective and Preventive Action"),
        ("NABH", "National Accreditation Board for Hospitals and Healthcare Providers"),
        ("NCD",  "Non-Communicable Disease"),
        ("UIP",  "Universal Immunisation Programme"),
    ])
    h(doc, 1, "12. Traceability table")
    p(doc, "This table is an index. It is not how the policy is organised. An asterisk in the "
          "Level column means the objective element is starred and requires documented evidence.")
    tr11 = [
        ("AAC.11.a", "Commitment*", "Sections 3 and 5.1 address written guidance for preventive and promotive care as per scope."),
        ("AAC.11.b", "Commitment",  "Sections 3 and 5.2 address age-appropriate non-communicable disease screening."),
        ("AAC.11.c", "Commitment",  "Sections 3 and 5.3 address mental health screening and intervention."),
        ("AAC.11.d", "Commitment",  "Sections 3 and 5.4 address paediatric and adult immunisation."),
        ("AAC.11.e", "Commitment",  "Sections 3 and 5.5 address multi-disciplinary lifestyle health education."),
    ]
    tt = tbl(doc, len(tr11)+1, 3)
    tt.cell(0,0).text="Objective Element"; tt.cell(0,1).text="Level"
    tt.cell(0,2).text="Traceability to this policy"
    for ri,(a,b,c) in enumerate(tr11, 1): tt.cell(ri,0).text=a; tt.cell(ri,1).text=b; tt.cell(ri,2).text=c
    h(doc, 1, "13. Required records/evidence checklist")
    h(doc, 2, "Written guidance — AAC.11.a (Commitment*)")
    lb(doc, "Current written guidance for preventive and promotive care aligned to scope, with "
           "version and issue date.")
    h(doc, 2, "NCD screening — AAC.11.b (Commitment)")
    lb(doc, "Defined NCD screening list with age and risk-context criteria.")
    lb(doc, "Patient record entries of screening advised, accepted or declined.")
    h(doc, 2, "Mental health screening — AAC.11.c (Commitment)")
    lb(doc, "Documented use of validated screening tools where applicable.")
    lb(doc, "Records of intervention, advice or referral after a positive screen.")
    h(doc, 2, "Immunisation — AAC.11.d (Commitment)")
    lb(doc, "Paediatric immunisation records aligned to UIP including catch-up.")
    lb(doc, "Adult immunisation advice or administration records where offered.")
    lb(doc, "Immunisation register or patient-file entries.")
    h(doc, 2, "Lifestyle health education — AAC.11.e (Commitment)")
    lb(doc, "Patient education record naming disciplines involved.")
    lb(doc, "Documentation of diet, exercise, ADL, pain or substance-cessation education where given.")
    h(doc, 1, "14. References")
    ln(doc, "National Accreditation Board for Hospitals and Healthcare Providers. NABH "
           "Accreditation Standards for Hospitals, 6th Edition. AAC.11.")
    ln(doc, "Guidebook interpretation supplied for AAC.11.a through AAC.11.e.")
    ln(doc, "Universal Immunisation Programme, as adopted for paediatric immunisation guidance.")
    ln(doc, f"Internal documents of {HN}: written guidance for preventive and promotive care; "
           f"NCD screening list; immunisation register.")
    h(doc, 1, "Disclaimer")
    p(doc, f"This policy reorganises the supplied AAC.11 objective-element wording and Guidebook "
          f"interpretation into plain-language guidance for {HN}. It does not replace the NABH "
          f"Accreditation Standards for Hospitals, 6th Edition, or the Guidebook, which remain "
          f"the authoritative source in case of any conflict.")
    p(doc, f"This policy references the Universal Immunisation Programme as the programme method "
          f"for paediatric immunisation guidance; it is not a statement of, or substitute for, "
          f"compliance with any statute. {HN} remains separately responsible for meeting all "
          f"applicable statutory and regulatory requirements.")
    p(doc, "This policy shall be reviewed at least once every year, or earlier if NABH standards change.")
    p(doc, f"This policy is the property of {HN} and is not to be reproduced or distributed "
          f"outside the organisation without authorisation.")
    save_and_verify(doc, "HCO_AAC_11_v2_REWRITE_DRAFT.docx")


# ══════════════════════════════════════════════════════════════════════════════
# AAC.12 — Established Discharge Process
# ══════════════════════════════════════════════════════════════════════════════
def gen_aac12():
    doc = Document()
    h(doc, 0, "Policy on the Established Discharge Process")
    p(doc, HN)
    h(doc, 1, "Document control")
    doc_ctrl(doc, "HCO/AAC/POL/12", "Medical Superintendent")
    h(doc, 1, "Statement of intent")
    p(doc, f"{HN} has an established discharge process: planned in consultation with the patient "
          f"and/or family, coordinated among departments and agencies including medico-legal and "
          f"absconded cases, governed by written LAMA guidance, issuing a discharge summary to "
          f"every patient, adhering to planned discharge, expanding access through domiciliary "
          f"visits where applicable, and monitoring discharge time for continual improvement.")
    h(doc, 1, "1. Purpose")
    p(doc, f"This policy explains how {HN} manages the discharge process: planning discharge "
          f"with the patient and family, coordinating discharge across departments and agencies, "
          f"governing discharge against medical advice, issuing a discharge summary to every "
          f"leaving patient, adhering to planned discharge, providing domiciliary visits where "
          f"applicable, and monitoring and improving discharge time.")
    p(doc, "The content of the discharge summary is covered in the hospital's other policies.")
    h(doc, 1, "2. Scope")
    p(doc, f"This policy applies to treating doctors, nursing staff, accounts and other "
          f"departments involved in discharge, and all patients and families at {HN}.")
    h(doc, 1, "3. Policy standards")
    p(doc, f"The discharge process is planned in consultation with the patient and/or family and "
          f"coordinated among all departments and agencies involved, including medico-legal and "
          f"absconded cases. Written guidance governs the discharge of patients leaving against "
          f"medical advice, and a discharge summary is given to all patients leaving the organisation.")
    p(doc, f"The organisation adheres to planned discharge, expands access to health practices "
          f"through domiciliary visits wherever applicable, and monitors discharge time, setting "
          f"appropriate benchmarks for continual improvement.")
    h(doc, 1, "4. Non-negotiable rules")
    lb(doc, "Do not discharge a patient without planning the process in consultation with the "
           "patient and/or family.")
    lb(doc, "Do not discharge a medico-legal case without informing the police, or handle an "
           "absconded case outside the organisation's coordination discipline.")
    lb(doc, "Do not allow a patient to leave against medical advice without following the written "
           "LAMA guidance and recording the declaration in the required format.")
    lb(doc, "Do not let any patient leave the organisation — including one leaving against "
           "medical advice or in case of death — without issuing a discharge or death summary.")
    lb(doc, "Do not fabricate a domiciliary visit service that falls outside the "
           "organisation's defined scope.")
    h(doc, 1, "5. What we do")
    h(doc, 2, "5.1 Plan discharge in consultation with the patient and family")
    p(doc, "The discharge process is planned in consultation with the patient and/or family. "
          "The treating doctor determines readiness for discharge during regular reassessments "
          "and discusses the plan with the patient and family, covering expected timing, "
          "continued care needs, and who will receive the discharge summary.")
    h(doc, 2, "5.2 Coordinate discharge among departments and agencies")
    p(doc, "The discharge process is coordinated among all departments and agencies involved, "
          "including accounts, so that discharge papers are completed within time. For "
          "medico-legal cases, the police are informed about the discharge. Absconded-patient "
          "workflows follow the same coordination discipline when the patient is found or the "
          "case is otherwise closed.")
    h(doc, 2, "5.3 Govern discharge against medical advice")
    p(doc, "Written guidance governs the discharge of patients leaving against medical advice "
          "(LAMA). The treating doctor explains the consequences of leaving to the patient or "
          "attendant. The written guidance addresses reasons for LAMA so the organisation can "
          "take corrective and/or preventive action where patterns appear. The "
          "patient/attendant declaration is recorded in a proper format. LAMA does not cancel "
          "the duty to give a discharge summary and reports.")
    h(doc, 2, "5.4 Issue a discharge summary to every leaving patient")
    p(doc, "A discharge summary is given to all patients leaving the organisation, including "
          "patients leaving against medical advice. In case of death, a death summary is given "
          "to the next of kin or relatives. The organisation hands over the discharge summary "
          "and reports to the patient or attendant in all cases and retains a copy in the "
          "medical record.")
    h(doc, 2, "5.5 Adhere to planned discharge")
    p(doc, "Discharge is planned at least 24 hours in advance. Planning includes preparation "
          "of the draft discharge summary, refund of medications where applicable, patient "
          "education on continued care, and identification of special needs — for example "
          "special equipment, devices, or transportation. Unplanned discharges are minimised.")
    h(doc, 2, "5.6 Provide domiciliary visits where applicable")
    p(doc, f"Where the organisation's scope includes domiciliary visits, trained healthcare "
          f"workers visit patients at their residences to deliver care and support — for "
          f"example for chronic conditions, fall prevention, memory care, nutrition, exercise, "
          f"or palliative catheter care. Visits are scheduled at discharge planning when "
          f"applicable and recorded. A hospital whose scope does not include domiciliary visits "
          f"records this as not applicable and does not fabricate the service.")
    h(doc, 2, "5.7 Monitor discharge time and improve it continually")
    p(doc, "The organisation monitors discharge time, setting appropriate benchmarks and making "
          "continual improvement. Discharge time is defined from when the treating doctor "
          "declares the patient fit for discharge to when the patient vacates the bed. "
          "Benchmarks may differ by payer mix. Delays are monitored, reasons identified, "
          "and improvement activities performed.")
    h(doc, 1, "6. Stop-work authority")
    p(doc, "Not applicable. This standard has no stop-work authority. Gaps in discharge "
          "planning, coordination, LAMA guidance, or discharge-summary issuance are addressed "
          "through the hospital's quality monitoring and CAPA process (Section 8), not through "
          "an immediate stop-work trigger. The patient's right to refuse treatment and leave "
          "against medical advice is a patient right, not a stop-work matter.")
    h(doc, 1, "7. Governance and responsibility")
    gov_tbl(doc, [
        ("Medical Superintendent",
         "Accountable for the discharge process across all departments and agencies."),
        ("Treating Doctor",
         "Determines discharge readiness; explains LAMA consequences; ensures discharge summary "
         "is prepared and handed over."),
        ("Nursing Superintendent",
         "Coordinates discharge planning and domiciliary visit scheduling with nursing staff."),
        ("Accounts/Administration",
         "Completes discharge papers within the defined coordination timeframe."),
        ("Quality Coordinator",
         "Audits this policy; monitors discharge time and holds LAMA and CAPA records."),
    ])
    h(doc, 1, "8. Quality monitoring")
    mon12 = [
        ("Discharge planning",
         "Reassessment notes stating readiness; discussion records with patient/family."),
        ("Departmental coordination",
         "Documented discharge procedure; MLC police intimation records; absconded-case "
         "coordination records."),
        ("LAMA",
         "Written LAMA guidance; patient/attendant declarations; record of consequences explained."),
        ("Discharge summary issuance",
         "Copy of discharge summary retained in the medical record for every leaver sampled; "
         "death summary acknowledgement."),
        ("Planned discharge",
         "Expected discharge noted in advance; draft summary and special-needs preparation; "
         "unplanned discharge log with reasons."),
        ("Domiciliary visits",
         "Scope statement; visit schedule and notes for eligible patients."),
        ("Discharge-time monitoring",
         "Discharge-time log from fit-for-discharge to bed vacated; delay analysis with "
         "improvement actions."),
    ]
    mt = tbl(doc, len(mon12)+1, 2)
    mt.cell(0,0).text="Monitoring area"; mt.cell(0,1).text="What is monitored"
    for ri,(a,b) in enumerate(mon12, 1): mt.cell(ri,0).text=a; mt.cell(ri,1).text=b
    h(doc, 1, "9. Training and staff acknowledgement")
    p(doc, "All treating doctors, nursing staff, and departments involved in discharge shall be "
          "trained on the discharge process, LAMA guidance, and discharge-summary requirements "
          "in this policy, at induction and periodically.")
    p(doc, f"I have read the Policy on the Established Discharge Process of {HN} and "
          f"understand my responsibilities under it.")
    h(doc, 1, "10. Distribution")
    p(doc, "This policy shall be available to all clinical staff, accounts and administration, "
          "department heads, and the Quality Coordinator.")
    h(doc, 1, "11. Abbreviations")
    abbrev_tbl(doc, [
        ("CAPA", "Corrective and Preventive Action"),
        ("LAMA", "Leaving/Left Against Medical Advice"),
        ("MLC",  "Medico-Legal Case"),
        ("NABH", "National Accreditation Board for Hospitals and Healthcare Providers"),
    ])
    h(doc, 1, "12. Traceability table")
    p(doc, "This table is an index. It is not how the policy is organised. An asterisk in the "
          "Level column means the objective element is starred and requires documented evidence.")
    tr12 = [
        ("AAC.12.a", "Commitment",  "Sections 3 and 5.1 address discharge planning in consultation with the patient and family."),
        ("AAC.12.b", "Commitment*", "Sections 3 and 5.2 address coordination of discharge among departments and agencies."),
        ("AAC.12.c", "Commitment*", "Sections 3 and 5.3 address written guidance for discharge against medical advice."),
        ("AAC.12.d", "Commitment",  "Sections 3 and 5.4 address issuing a discharge summary to every leaving patient."),
        ("AAC.12.e", "Commitment",  "Sections 3 and 5.5 address adherence to planned discharge."),
        ("AAC.12.f", "Commitment",  "Sections 3 and 5.6 address domiciliary visits where applicable."),
        ("AAC.12.g", "Commitment",  "Sections 3 and 5.7 address monitoring and improving discharge time."),
    ]
    tt = tbl(doc, len(tr12)+1, 3)
    tt.cell(0,0).text="Objective Element"; tt.cell(0,1).text="Level"
    tt.cell(0,2).text="Traceability to this policy"
    for ri,(a,b,c) in enumerate(tr12, 1): tt.cell(ri,0).text=a; tt.cell(ri,1).text=b; tt.cell(ri,2).text=c
    h(doc, 1, "13. Required records/evidence checklist")
    h(doc, 2, "Discharge planning — AAC.12.a (Commitment)")
    lb(doc, "Reassessment note stating readiness for discharge.")
    lb(doc, "Record of discussion with patient and/or family about the discharge plan.")
    h(doc, 2, "Departmental coordination — AAC.12.b (Commitment*)")
    lb(doc, "Documented discharge procedure covering department coordination including accounts.")
    lb(doc, "MLC discharge file with police intimation record.")
    lb(doc, "Absconded-case coordination record where applicable.")
    h(doc, 2, "LAMA guidance — AAC.12.c (Commitment*)")
    lb(doc, "Current LAMA written guidance.")
    lb(doc, "Patient/attendant LAMA declaration in the required format.")
    lb(doc, "Record that the treating doctor explained consequences.")
    h(doc, 2, "Discharge summary issuance — AAC.12.d (Commitment)")
    lb(doc, "Copy of discharge summary retained in the medical record for every leaver sampled.")
    lb(doc, "LAMA cases with summary handed over; death summary acknowledgement by next of kin.")
    h(doc, 2, "Planned discharge — AAC.12.e (Commitment)")
    lb(doc, "Expected discharge noted at least 24 hours in advance for planned cases.")
    lb(doc, "Draft discharge summary and special-needs preparation before discharge.")
    lb(doc, "Log of unplanned discharges with reasons and minimisation actions.")
    h(doc, 2, "Domiciliary visits — AAC.12.f (Commitment)")
    lb(doc, "Scope statement or written note of when domiciliary visits apply.")
    lb(doc, "Domiciliary visit schedule and visit notes for eligible patients.")
    h(doc, 2, "Discharge-time monitoring — AAC.12.g (Commitment)")
    lb(doc, "Defined discharge-time benchmarks (by payer mix where used).")
    lb(doc, "Discharge-time log; delay analysis with continual-improvement actions.")
    h(doc, 1, "14. References")
    ln(doc, "National Accreditation Board for Hospitals and Healthcare Providers. NABH "
           "Accreditation Standards for Hospitals, 6th Edition. AAC.12.")
    ln(doc, "Guidebook interpretation supplied for AAC.12.a through AAC.12.g.")
    ln(doc, f"Internal documents of {HN}: discharge procedure; LAMA written guidance; "
           f"discharge-time benchmarks; domiciliary visit schedule.")
    h(doc, 1, "Disclaimer")
    p(doc, f"This policy reorganises the supplied AAC.12 objective-element wording and Guidebook "
          f"interpretation into plain-language guidance for {HN}. It does not replace the NABH "
          f"Accreditation Standards for Hospitals, 6th Edition, or the Guidebook, which remain "
          f"the authoritative source in case of any conflict.")
    p(doc, f"This policy is intended for accreditation and internal governance purposes and is "
          f"not a statement of, or substitute for, compliance with any specific statute; {HN} "
          f"remains separately responsible for meeting all applicable statutory and regulatory "
          f"requirements, including medico-legal reporting obligations.")
    p(doc, "This policy shall be reviewed at least once every year, or earlier if NABH standards change.")
    p(doc, f"This policy is the property of {HN} and is not to be reproduced or distributed "
          f"outside the organisation without authorisation.")
    save_and_verify(doc, "HCO_AAC_12_v2_REWRITE_DRAFT.docx")


# ══════════════════════════════════════════════════════════════════════════════
# AAC.13 — Content of the Discharge Summary
# ══════════════════════════════════════════════════════════════════════════════
def gen_aac13():
    doc = Document()
    h(doc, 0, "Policy on the Content of the Discharge Summary")
    p(doc, HN)
    h(doc, 1, "Document control")
    doc_ctrl(doc, "HCO/AAC/POL/13", "Medical Superintendent")
    h(doc, 1, "Statement of intent")
    p(doc, f"{HN} defines the content of the discharge summary: provided at the time of "
          f"discharge, signed and acknowledged, standardised in content, written in an "
          f"understandable manner with follow-up and medication instructions, incorporating "
          f"urgent-care instructions, and including cause of death where applicable.")
    h(doc, 1, "1. Purpose")
    p(doc, f"This policy explains what {HN}'s discharge summary contains and how instructions "
          f"are written: provision and acknowledgement at discharge, standardised minimum "
          f"content, understandable follow-up and medication instructions, urgent-care "
          f"instructions, and cause of death in case of death.")
    p(doc, "The discharge process itself — planning, coordination, LAMA guidance, and that a "
          "summary is given — is covered in the hospital's other policies.")
    h(doc, 1, "2. Scope")
    p(doc, f"This policy applies to treating doctors and doctor members of treating teams who "
          f"prepare and sign discharge summaries, and to all patients and families receiving a "
          f"discharge or death summary at {HN}.")
    h(doc, 1, "3. Policy standards")
    p(doc, f"A discharge summary is provided to the patient at the time of discharge, signed by "
          f"the treating doctor or a doctor member of the treating team, and acknowledged by the "
          f"patient or family. The discharge summary has standardised minimum content.")
    p(doc, "The summary contains follow-up advice, medication and other instructions in an "
          "understandable manner, incorporates instructions about when and how to obtain urgent "
          "care, and, in case of death, includes the cause of death.")
    h(doc, 1, "4. Non-negotiable rules")
    lb(doc, "Do not issue a discharge summary without the signature of the treating doctor or a "
           "doctor member of the treating team, or without a recorded acknowledgement of receipt "
           "by the patient or family.")
    lb(doc, "Do not omit any of the standardised minimum content fields from a discharge summary "
           "without recording the reason it is not applicable.")
    lb(doc, "Do not write medication frequencies using abbreviations such as BD, TID or QID — "
           "write them in plain words the patient can understand.")
    lb(doc, "Do not omit urgent-care instructions specific to the patient's diagnosis and "
           "condition, or the hospital or doctor contact number, from the discharge summary.")
    lb(doc, "Do not issue a death summary without stating the cause of death, or without "
           "documenting a post-mortem when the cause is unclear.")
    h(doc, 1, "5. What we do")
    h(doc, 2, "5.1 Provide and acknowledge the discharge summary")
    p(doc, "A discharge summary is provided to the patient at the time of discharge. The "
          "discharge summary is signed by the treating doctor or a doctor member of the treating "
          "team. The patient or family acknowledges receipt of the summary, and this "
          "acknowledgement is recorded with date and name of the person receiving it.")
    h(doc, 2, "5.2 Standardise discharge summary content")
    p(doc, "Discharge summaries have standardised minimum content: patient's name; unique "
          "identification number; name of the treating doctor; date of admission and date of "
          "discharge; reasons for admission; significant findings, diagnosis and the patient's "
          "condition at discharge; investigation results; any procedure performed; medication "
          "administered; and any other treatment given. The summary may also name other "
          "consultants involved in treatment.")
    h(doc, 2, "5.3 Write understandable follow-up and medication instructions")
    p(doc, "The discharge summary contains follow-up advice, medication and other instructions "
          "in an understandable manner. Instructions are explained to the patient and/or "
          "relatives in a language and manner they understand. Medical abbreviations such as BD, "
          "TID and QID are not used; frequencies are written in plain words — for example "
          "morning and night, three times a day. Instructions include safe and effective use of "
          "medical equipment at home where applicable, and wound-care or pressure-ulcer care "
          "instructions for post-operative patients where relevant.")
    h(doc, 2, "5.4 Incorporate urgent-care instructions")
    p(doc, "The discharge summary incorporates instructions about when and how to obtain urgent "
          "care, specific to the patient's diagnosis and clinical condition at discharge — for "
          "example development of fever, or bleeding or discharge from the operative site. "
          "Instructions may state what medicines to take, when to consult a doctor, how to seek "
          "medical help, and the contact number of the hospital or doctor. These instructions "
          "are explained to the patient and/or relatives in a language and manner they understand.")
    h(doc, 2, "5.5 Include cause of death where applicable")
    p(doc, "In case of death, the summary of the case includes the cause of death. When the "
          "cause of death is not clear, a post-mortem is performed — for example in a "
          "medico-legal case — and documented.")
    h(doc, 1, "6. Stop-work authority")
    p(doc, "Not applicable. This standard has no stop-work authority. Gaps in discharge summary "
          "content, signature, acknowledgement, or urgent-care instructions are addressed through "
          "the hospital's quality monitoring and CAPA process (Section 8), not through an "
          "immediate stop-work trigger.")
    h(doc, 1, "7. Governance and responsibility")
    gov_tbl(doc, [
        ("Medical Superintendent",
         "Accountable for the standardised discharge summary template and content requirements."),
        ("Treating Doctor",
         "Prepares, signs, and explains the discharge summary, including urgent-care instructions, "
         "in an understandable manner."),
        ("Nursing Superintendent",
         "Supports acknowledgement recording and equipment/wound-care instruction delivery."),
        ("Quality Coordinator",
         "Audits this policy; holds discharge summary templates and completed-summary samples."),
    ])
    h(doc, 1, "8. Quality monitoring")
    mon13 = [
        ("Provision and acknowledgement",
         "Signed discharge summary issued at discharge; acknowledgement of receipt with date and name."),
        ("Standardised content",
         "Current standardised discharge-summary template; completed summaries with all minimum content fields."),
        ("Follow-up and medication instructions",
         "Summaries with plain-language medication frequencies; equipment and wound-care instructions where applicable."),
        ("Urgent-care instructions",
         "Urgent-care section specific to diagnosis/condition at discharge, with hospital or doctor contact."),
        ("Cause of death",
         "Death summary stating cause of death; post-mortem documentation when required."),
    ]
    mt = tbl(doc, len(mon13)+1, 2)
    mt.cell(0,0).text="Monitoring area"; mt.cell(0,1).text="What is monitored"
    for ri,(a,b) in enumerate(mon13, 1): mt.cell(ri,0).text=a; mt.cell(ri,1).text=b
    h(doc, 1, "9. Training and staff acknowledgement")
    p(doc, "All treating doctors and doctor members of treating teams shall be trained on the "
          "standardised discharge summary content, plain-language instruction requirements, and "
          "urgent-care instruction content in this policy, at induction and periodically.")
    p(doc, f"I have read the Policy on the Content of the Discharge Summary of {HN} and "
          f"understand my responsibilities under it.")
    h(doc, 1, "10. Distribution")
    p(doc, "This policy shall be available to all treating doctors, department heads, and the Quality Coordinator.")
    h(doc, 1, "11. Abbreviations")
    abbrev_tbl(doc, [
        ("CAPA", "Corrective and Preventive Action"),
        ("MLC",  "Medico-Legal Case"),
        ("NABH", "National Accreditation Board for Hospitals and Healthcare Providers"),
    ])
    h(doc, 1, "12. Traceability table")
    p(doc, "This table is an index. It is not how the policy is organised. An asterisk in the "
          "Level column means the objective element is starred and requires documented evidence.")
    tr13 = [
        ("AAC.13.a", "Commitment", "Sections 3 and 5.1 address provision and acknowledgement of the discharge summary."),
        ("AAC.13.b", "Commitment", "Sections 3 and 5.2 address standardised minimum content."),
        ("AAC.13.c", "Commitment", "Sections 3 and 5.3 address understandable follow-up and medication instructions."),
        ("AAC.13.d", "Commitment", "Sections 3 and 5.4 address urgent-care instructions."),
        ("AAC.13.e", "Commitment", "Sections 3 and 5.5 address cause of death in case of death."),
    ]
    tt = tbl(doc, len(tr13)+1, 3)
    tt.cell(0,0).text="Objective Element"; tt.cell(0,1).text="Level"
    tt.cell(0,2).text="Traceability to this policy"
    for ri,(a,b,c) in enumerate(tr13, 1): tt.cell(ri,0).text=a; tt.cell(ri,1).text=b; tt.cell(ri,2).text=c
    h(doc, 1, "13. Required records/evidence checklist")
    h(doc, 2, "Provision and acknowledgement — AAC.13.a (Commitment)")
    lb(doc, "Signed discharge summary issued at discharge.")
    lb(doc, "Acknowledgement of receipt by patient or family (with date and name).")
    h(doc, 2, "Standardised content — AAC.13.b (Commitment)")
    lb(doc, "Current standardised discharge-summary template.")
    lb(doc, "Completed summaries with all minimum content fields.")
    h(doc, 2, "Follow-up and medication instructions — AAC.13.c (Commitment)")
    lb(doc, "Summaries with plain-language medication frequencies.")
    lb(doc, "Equipment and wound-care instructions where applicable.")
    h(doc, 2, "Urgent-care instructions — AAC.13.d (Commitment)")
    lb(doc, "Urgent-care section specific to diagnosis/condition at discharge, with hospital or doctor contact.")
    lb(doc, "Record of explanation to patient/relatives.")
    h(doc, 2, "Cause of death — AAC.13.e (Commitment)")
    lb(doc, "Death summary stating cause of death.")
    lb(doc, "Post-mortem documentation when cause is unclear or MLC.")
    h(doc, 1, "14. References")
    ln(doc, "National Accreditation Board for Hospitals and Healthcare Providers. NABH "
           "Accreditation Standards for Hospitals, 6th Edition. AAC.13.")
    ln(doc, "Guidebook interpretation supplied for AAC.13.a through AAC.13.e.")
    ln(doc, f"Internal documents of {HN}: standardised discharge-summary template; completed "
           f"summaries; death summary records.")
    h(doc, 1, "Disclaimer")
    p(doc, f"This policy reorganises the supplied AAC.13 objective-element wording and Guidebook "
          f"interpretation into plain-language guidance for {HN}. It does not replace the NABH "
          f"Accreditation Standards for Hospitals, 6th Edition, or the Guidebook, which remain "
          f"the authoritative source in case of any conflict.")
    p(doc, f"This policy is intended for accreditation and internal governance purposes and is "
          f"not a statement of, or substitute for, compliance with any specific statute; {HN} "
          f"remains separately responsible for meeting all applicable statutory and regulatory "
          f"requirements, including medico-legal reporting obligations.")
    p(doc, "This policy shall be reviewed at least once every year, or earlier if NABH standards change.")
    p(doc, f"This policy is the property of {HN} and is not to be reproduced or distributed "
          f"outside the organisation without authorisation.")
    save_and_verify(doc, "HCO_AAC_13_v2_REWRITE_DRAFT.docx")


if __name__ == "__main__":
    gen_aac1()
    gen_aac2()
    gen_aac3()
    gen_aac4()
    gen_aac5()
    gen_aac6()
    gen_aac7()
    gen_aac8()
    print("\nAll 8 AAC rewrite drafts generated.")

