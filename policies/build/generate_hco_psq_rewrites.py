# -*- coding: utf-8 -*-
"""
generate_hco_psq_rewrites.py
Generates HCO PSQ.1-7 v2 rewrite-reference DOCX files.

Pipeline : python-docx, identical to generate_hco_fms_rewrites.py.
Output   : policies/build/rewrite_reference/HCO_PSQ_N_v2_REWRITE_DRAFT.docx
Source   : ChatGPT first-draft content (approved) + policies/drafts_hco/hco_psqN_v2_draft.json
"""
import os
from docx import Document

HN  = "«Hospital Name»"
OUT = "policies/build/rewrite_reference"
os.makedirs(OUT, exist_ok=True)


# ── Helpers (identical to generate_hco_fms_rewrites.py) ──────────────────────

def h(doc, lv, txt):
    return doc.add_paragraph(txt, style={0: "Title", 1: "Heading 1", 2: "Heading 2"}[lv])

def p(doc, txt=""):
    return doc.add_paragraph(txt, style="Normal")

def ln(doc, txt):
    return doc.add_paragraph(txt, style="List Number")

def lb(doc, txt):
    return doc.add_paragraph(txt, style="List Bullet")

def tbl(doc, rows, cols):
    t = doc.add_table(rows=rows, cols=cols)
    try:
        t.style = "Table Grid"
    except KeyError:
        pass
    return t

def doc_ctrl(doc, no, prep, appr="Medical Superintendent"):
    dc = tbl(doc, 6, 4)
    for ri, (a, b, c, d) in enumerate([
        ("Document No.", no,         "Version",              "2.0"),
        ("Issue No.",    "01",        "Review due",           "One year from implementation"),
        ("Date created", "________",  "Date of implementation", "________"),
    ]):
        dc.cell(ri, 0).text = a; dc.cell(ri, 1).text = b
        dc.cell(ri, 2).text = c; dc.cell(ri, 3).text = d
    for ri, (lbl, txt) in enumerate([
        ("Prepared by", f"{prep}  Name: ________  Signature: ________"),
        ("Reviewed by", "Quality Coordinator  Name: ________  Signature: ________"),
        ("Approved by", f"{appr}  Name: ________  Signature: ________"),
    ], start=3):
        dc.cell(ri, 0).text = lbl
        c1 = dc.cell(ri, 1); c1.text = txt; c1.merge(dc.cell(ri, 3))

def abbrev_tbl(doc, rows):
    t = tbl(doc, len(rows) + 1, 2)
    t.cell(0, 0).text = "Abbreviation"; t.cell(0, 1).text = "Meaning"
    for ri, (a, m) in enumerate(rows, 1):
        t.cell(ri, 0).text = a; t.cell(ri, 1).text = m

def save_and_verify(doc, fname):
    import sys
    out = sys.stdout
    def pr(s):
        try:
            out.write(s + "\n")
        except UnicodeEncodeError:
            out.write(s.encode("ascii", "replace").decode() + "\n")
    pr(f"\n=== {fname} ===")
    for i, para in enumerate(doc.paragraphs[:60]):
        sn = para.style.name if para.style else "(None)"
        pr(f"{i:3d}  {sn!r:30s}  {para.text[:60]!r}")
    pr(f"  Total paras: {len(doc.paragraphs)}")
    path = os.path.join(OUT, fname)
    doc.save(path)
    print(f"  Saved: {path}")


# ══════════════════════════════════════════════════════════════════════════════
# PSQ.1 — Patient Safety Programme   (no stop-work)
# Content: ChatGPT first draft (approved). Two mechanical fixes applied:
#   1. Section numbers renumbered: Traceability=11, Records=12, References=13
#   2. Literal <b> / </b> stripped from staff acknowledgement sentence
# ══════════════════════════════════════════════════════════════════════════════
def gen_psq1():
    doc = Document()

    # Title + hospital name
    h(doc, 0, "Policy on Patient Safety Programme")
    p(doc, HN)

    # Document control
    h(doc, 1, "Document control")
    doc_ctrl(doc, "HCO/PSQ/POL/01", "Patient Safety Officer")
    p(doc, "A blank marked ________ must be completed before issue.")

    # Statement of intent
    h(doc, 1, "Statement of intent")
    p(doc, f"{HN} implements a structured patient safety programme.")
    p(doc,
      "The patient safety programme is developed, implemented and maintained by a multi-disciplinary "
      "safety committee. The programme covers the major elements related to patient safety and incidents "
      "ranging from “no harm” to “sentinel events”. Designated patient safety officer(s) "
      "coordinate implementation of the patient safety programme. The organisation performs proactive "
      "analysis of patient safety risks and makes improvements accordingly. The programme is reviewed "
      "and updated at least once a year and the organisation adapts and implements national and "
      "international patient-safety goals, solutions or frameworks.")

    # 1. Purpose
    h(doc, 1, "1. Purpose")
    p(doc,
      "The purpose of this policy is to establish and maintain a structured approach to patient safety "
      "and quality improvement through defined goals, resources, quality indicators, data analysis, "
      "improvement actions and communication of patient safety and quality information.")
    p(doc, "This policy does not cover matters outside the patient safety programme.")

    # 2. Scope
    h(doc, 1, "2. Scope")
    p(doc,
      f"This policy applies across {HN} to the patient safety programme and the activities supporting "
      "patient safety and quality improvement.")
    p(doc,
      "It applies to the multi-disciplinary safety committee, designated patient safety officer(s), "
      "leadership, relevant staff, departments and, where applicable, the governing body.")

    # 3. Policy standards
    h(doc, 1, "3. Policy standards")

    h(doc, 2, "3.1 Patient safety and quality goals")
    p(doc, "The organisation identifies patient safety and quality goals for a minimum period of one year.")
    p(doc, "The goals are based on findings from:")
    lb(doc, "Surveillance.")
    lb(doc, "Audit.")
    lb(doc, "Incident reports.")
    lb(doc, "Patient feedback.")
    lb(doc, "Risk assessments.")
    p(doc, "The goals should be measurable and address high-risk, high-volume or problem-prone areas.")

    h(doc, 2, "3.2 Resources for patient safety and quality goals")
    p(doc,
      "Leadership provides adequate human and financial resources to implement and sustain the "
      "identified quality and patient safety goals.")
    p(doc, "Resources include:")
    lb(doc, "Budget.")
    lb(doc, "Staff time.")
    lb(doc, "Training.")
    p(doc,
      "Resource allocation should be visible in planning documents and should not be only stated "
      "as intent.")

    h(doc, 2, "3.3 Quality indicators")
    p(doc,
      "The organisation identifies and monitors quality indicators relevant to clinical and "
      "managerial processes.")
    p(doc, "The indicators are aligned with the identified patient safety and quality goals.")
    p(doc, "Each indicator shall have:")
    lb(doc, "A defined numerator.")
    lb(doc, "A defined denominator.")
    lb(doc, "A defined data-collection frequency.")

    h(doc, 2, "3.4 Quality-data analysis")
    p(doc,
      "The organisation defines a mechanism for aggregating, analysing and using quality-indicator "
      "data on a periodic basis.")
    p(doc, "The mechanism is used to identify:")
    lb(doc, "Trends.")
    lb(doc, "Variances.")
    lb(doc, "Improvement opportunities.")
    p(doc, "Analysis could use run charts, control charts or comparative benchmarking.")

    h(doc, 2, "3.5 Use of quality data for improvement")
    p(doc,
      "The organisation shall document evidence that quality data is used to drive actual "
      "improvement actions.")
    p(doc,
      "Data collection alone, without a documented improvement response, does not meet this "
      "requirement.")

    h(doc, 2, "3.6 Communication of quality and patient safety information")
    p(doc,
      "Quality and patient safety information shall be regularly communicated to relevant staff "
      "and departments.")
    p(doc,
      "Where applicable, the information shall also be communicated to the governing body.")
    p(doc, "Communication could be through:")
    lb(doc, "Dashboards.")
    lb(doc, "Meetings.")
    lb(doc, "Reports.")

    h(doc, 2, "3.7 Quality-improvement tools")
    p(doc,
      "The organisation could use quality-improvement tools to structure improvement projects.")
    p(doc, "Examples include:")
    lb(doc, "PDSA.")
    lb(doc, "Six Sigma.")
    lb(doc, "Lean methodologies.")
    p(doc, "The choice of tool or methodology is not mandated.")

    # 4. Non-negotiable rules
    h(doc, 1, "4. Non-negotiable rules")
    ln(doc,
       "Do not treat quality indicators as complete without defined numerators, denominators and "
       "data-collection frequency.")
    ln(doc,
       "Do not treat quality-data collection alone as sufficient without documented evidence of "
       "the improvement action driven by the data.")
    ln(doc,
       "Do not withhold regular quality and patient safety information from relevant staff and "
       "departments.")
    ln(doc,
       "Do not omit communication of quality and patient safety information to the governing body "
       "where applicable.")
    ln(doc,
       "Do not replace the required quality-data analysis mechanism with a specific analytical tool "
       "as though that tool were mandatory; the use of run charts, control charts or comparative "
       "benchmarking could be selected.")
    ln(doc,
       "Do not make PDSA, Six Sigma or Lean a mandatory methodology for quality-improvement "
       "projects; the choice of tool or methodology is not mandated.")

    # 5. What we do
    h(doc, 1, "5. What we do")

    h(doc, 2, "5.1 Identify patient safety and quality goals")
    p(doc,
      "The organisation identifies patient safety and quality goals for a minimum period of one year.")
    p(doc,
      "The goals are informed by surveillance, audit, incident reports, patient feedback and "
      "risk assessments.")
    p(doc,
      "Goals should be measurable and should address high-risk, high-volume or problem-prone areas.")

    h(doc, 2, "5.2 Provide resources")
    p(doc,
      "Leadership provides adequate human and financial resources for implementation and "
      "sustainability of the identified goals.")
    p(doc, "The resources include budget, staff time and training.")
    p(doc, "Resource allocation should be visible in planning documents.")

    h(doc, 2, "5.3 Define and monitor indicators")
    p(doc,
      "The organisation identifies quality indicators relevant to clinical and managerial processes.")
    p(doc, "Indicators are aligned with identified goals.")
    p(doc,
      "Each indicator shall have a defined numerator, denominator and data-collection frequency.")

    h(doc, 2, "5.4 Aggregate and analyse data")
    p(doc,
      "Quality-indicator data is aggregated, analysed and used through a defined mechanism.")
    p(doc,
      "The analysis is performed on a periodic basis to identify trends, variances and improvement "
      "opportunities.")
    p(doc,
      "Run charts, control charts or comparative benchmarking could be used for analysis.")

    h(doc, 2, "5.5 Act on quality data")
    p(doc, "The organisation uses quality data to drive actual improvement actions.")
    p(doc,
      "The organisation shall document evidence of the improvement actions arising from quality data.")

    h(doc, 2, "5.6 Communicate information")
    p(doc,
      "Quality and patient safety information shall be regularly communicated to relevant staff "
      "and departments.")
    p(doc, "Where applicable, the governing body receives the information.")
    p(doc, "Dashboards, meetings or reports could be used for communication.")

    h(doc, 2, "5.7 Use improvement tools where appropriate")
    p(doc,
      "The organisation could use PDSA, Six Sigma or Lean methodologies to structure improvement "
      "projects.")
    p(doc, "The choice of quality-improvement tool or methodology is not mandated.")

    # 6. Governance and responsibility
    h(doc, 1, "6. Governance and responsibility")

    h(doc, 2, "6.1 Leadership")
    p(doc,
      "Leadership provides adequate human and financial resources to implement and sustain "
      "identified quality and patient safety goals.")
    p(doc, "Leadership supports resource allocation for budget, staff time and training.")

    h(doc, 2, "6.2 Multi-disciplinary safety committee")
    p(doc,
      "The multi-disciplinary safety committee develops, implements and maintains the patient "
      "safety programme.")
    p(doc,
      "The committee supports the identification of patient safety and quality goals and the "
      "implementation of the programme.")

    h(doc, 2, "6.3 Patient safety officer(s)")
    p(doc,
      "Designated patient safety officer(s) coordinate implementation of the patient safety "
      "programme.")

    h(doc, 2, "6.4 Relevant staff and departments")
    p(doc,
      "Relevant staff and departments participate in the implementation of identified patient "
      "safety and quality goals.")
    p(doc,
      "They receive quality and patient safety information through the defined communication "
      "mechanism.")

    h(doc, 2, "6.5 Governing body")
    p(doc,
      "Where applicable, the governing body receives quality and patient safety information "
      "through the established communication mechanism.")

    # 7. Quality monitoring
    h(doc, 1, "7. Quality monitoring")
    mon_rows = [
        ("Patient safety goals",   "Identified patient safety and quality goals"),
        ("Goal period",            "Goals identified for a minimum period of one year"),
        ("Goal basis",             "Surveillance, audit, incident reports, patient feedback and risk assessments"),
        ("Resources",              "Human and financial resources, budget, staff time and training"),
        ("Quality indicators",     "Clinical and managerial quality indicators"),
        ("Indicator definition",   "Numerator, denominator and data-collection frequency"),
        ("Data management",        "Aggregation, analysis and use of quality-indicator data"),
        ("Data analysis",          "Trends, variances and improvement opportunities"),
        ("Improvement response",   "Documented evidence that quality data drives actual improvement actions"),
        ("Communication",          "Regular communication of quality and patient safety information"),
        ("Recipients",             "Relevant staff, departments and, where applicable, governing body"),
        ("Improvement tools",      "Use of quality-improvement tools where the organisation chooses to use them"),
    ]
    mon = tbl(doc, len(mon_rows) + 1, 2)
    mon.cell(0, 0).text = "Monitoring area"
    mon.cell(0, 1).text = "What is monitored"
    for ri, (area, what) in enumerate(mon_rows, 1):
        mon.cell(ri, 0).text = area
        mon.cell(ri, 1).text = what
    p(doc,
      "Quality monitoring shall distinguish between data collection and documented improvement "
      "response. Data collection without a documented improvement response does not meet the "
      "stated requirement.")

    # 8. Training and staff acknowledgement
    h(doc, 1, "8. Training and staff acknowledgement")
    p(doc,
      "Leadership provides training as part of the resources required to implement and sustain "
      "identified quality and patient safety goals.")
    p(doc,
      "Relevant staff participate in activities related to the patient safety programme according "
      "to their responsibilities.")
    p(doc, "Staff acknowledgement:")
    # Fix 2: <b> / </b> stripped — sentence as plain text
    p(doc,
      f"I have read the Policy on Patient Safety Programme of {HN}. I will follow the processes "
      "described.")
    # Sign-off table: 4 columns, 4 rows (header + 3 data)
    sig = tbl(doc, 4, 4)
    for ci, hdr in enumerate(("Staff name", "Designation", "Signature", "Date")):
        sig.cell(0, ci).text = hdr
    for ri in range(1, 4):
        for ci in range(4):
            sig.cell(ri, ci).text = "________"

    # 9. Distribution
    h(doc, 1, "9. Distribution")
    p(doc,
      "This policy shall be available to relevant staff and departments involved in the patient "
      "safety programme.")
    p(doc,
      "Quality and patient safety information shall be regularly communicated to relevant staff "
      "and departments and, where applicable, the governing body.")

    # 10. Abbreviations
    h(doc, 1, "10. Abbreviations")
    abbrev_tbl(doc, [
        ("NABH", "National Accreditation Board for Hospitals and Healthcare Providers"),
        ("PDSA", "Plan-Do-Study-Act"),
        ("PSQ",  "Patient Safety and Quality Improvement"),
    ])

    # 11. Traceability table  [fix 1: was 12 in source]
    h(doc, 1, "11. Traceability table")
    p(doc,
      "This table is an index. It is not how the policy is organised. An asterisk in the Level "
      "column means documentation of the process is required.")
    tr = tbl(doc, 8, 3)
    for ci, hdr in enumerate(("Objective Element", "Level", "Traceability to this policy")):
        tr.cell(0, ci).text = hdr
    trace_rows = [
        ("PSQ.1.a", "CORE*",
         "Sections 3.1, 5.1 and 6.2 address patient safety and quality goals and the development, "
         "implementation and maintenance of the patient safety programme by the multi-disciplinary "
         "safety committee."),
        ("PSQ.1.b", "Commitment",
         "Sections 3.2 and 6.1 address adequate human and financial resources, including budget, "
         "staff time and training, and the visibility of resource allocation in planning documents."),
        ("PSQ.1.c", "Commitment",
         "Sections 3.3 and 5.3 address identification and monitoring of quality indicators and "
         "defined numerators, denominators and data-collection frequency."),
        ("PSQ.1.d", "Commitment",
         "Sections 3.4 and 5.4 address the defined mechanism for aggregating, analysing and using "
         "quality-indicator data on a periodic basis to identify trends, variances and improvement "
         "opportunities."),
        ("PSQ.1.e", "Commitment",
         "Sections 3.5 and 5.5 address documented evidence that quality data is used to drive "
         "actual improvement actions."),
        ("PSQ.1.f", "Commitment",
         "Sections 3.6 and 5.6 address regular communication of quality and patient safety "
         "information to relevant staff, departments and, where applicable, the governing body."),
        ("PSQ.1.g", "CORE",
         "Sections 3.7 and 5.7 address the possible use of quality-improvement tools while "
         "preserving the source position that the choice of tool or methodology is not mandated."),
    ]
    for ri, (oe, lvl, txt) in enumerate(trace_rows, 1):
        tr.cell(ri, 0).text = oe
        tr.cell(ri, 1).text = lvl
        tr.cell(ri, 2).text = txt

    # 12. Required Records/Evidence Checklist  [fix 1: was 13 in source]
    h(doc, 1, "12. Required Records/Evidence Checklist")

    h(doc, 2, "Patient safety and quality goals")
    lb(doc, "Current patient safety and quality goals covering a minimum period of one year.")
    lb(doc,
       "Documents showing the goals were based on surveillance findings, audit findings, incident "
       "reports, patient feedback and risk assessments.")
    lb(doc,
       "Goals that are measurable and address high-risk, high-volume or problem-prone areas.")

    h(doc, 2, "Resources and planning")
    lb(doc,
       "Planning documents showing the human and financial resources allocated to patient safety "
       "and quality goals.")
    lb(doc, "Budget allocation for the identified goals.")
    lb(doc, "Staff-time allocation for the identified goals.")
    lb(doc, "Training resources allocated for the identified goals.")

    h(doc, 2, "Quality indicators")
    lb(doc, "List of quality indicators for relevant clinical and managerial processes.")
    lb(doc, "Document showing how each indicator is aligned with the identified goals.")
    lb(doc, "Defined numerator and denominator for each indicator.")
    lb(doc, "Defined data-collection frequency for each indicator.")

    h(doc, 2, "Quality data management and analysis")
    lb(doc, "Defined mechanism for aggregating quality-indicator data.")
    lb(doc, "Defined mechanism for analysing quality-indicator data.")
    lb(doc,
       "Defined mechanism for using quality-indicator data to identify trends, variances and "
       "improvement opportunities.")
    lb(doc,
       "Periodic quality-data analysis showing identified trends, variances and improvement "
       "opportunities, where applicable.")
    lb(doc,
       "Analysis may include run charts, control charts or comparative benchmarking where these "
       "methods are used.")

    h(doc, 2, "Improvement actions")
    lb(doc, "Documents showing quality data led to actual improvement actions.")
    lb(doc,
       "Documentation of the improvement response taken after quality-data analysis.")

    h(doc, 2, "Communication")
    lb(doc,
       "Dashboards, meeting records or reports showing regular communication of quality and patient "
       "safety information to relevant staff and departments.")
    lb(doc, "Where applicable, information shared with the governing body.")
    lb(doc,
       "Records showing the communication method used, such as dashboards, meetings or reports.")

    h(doc, 2, "Quality-improvement methods")
    lb(doc,
       "Where used, project documents showing the quality-improvement tool or methodology selected, "
       "such as PDSA, Six Sigma or Lean.")
    lb(doc,
       "There is no requirement to use a particular quality-improvement tool or methodology.")

    # 13. References  [fix 1: was 14 in source]
    h(doc, 1, "13. References")
    lb(doc,
       "National Accreditation Board for Hospitals and Healthcare Providers. NABH Accreditation "
       "Standards for Hospitals, 6th Edition. PSQ.1, “The organisation implements a structured "
       "patient safety programme.”")
    lb(doc,
       "National Accreditation Board for Hospitals and Healthcare Providers. NABH Accreditation "
       "Standards for Hospitals, 6th Edition. Guidance on reading standards and objective elements, "
       "including the interpretation of asterisked objective elements as requiring documentation.")

    # Disclaimer
    h(doc, 1, "Disclaimer")
    p(doc,
      "This policy reorganises the supplied PSQ.1 objective-element wording and Guidebook "
      "interpretation into plain-language policy format. The modal strength of the Guidebook has "
      "been preserved: shall requirements are treated as mandatory, while should, could and examples "
      "have not been converted into mandatory requirements. No additional frequency, deadline, "
      "threshold or prescribed methodology has been added beyond what is stated in the supplied source.")

    save_and_verify(doc, "HCO_PSQ_1_v2_REWRITE_DRAFT.docx")


# ══════════════════════════════════════════════════════════════════════════════
# PSQ.5 — Clinical Audit   (no stop-work)
# Content: ChatGPT final draft (approved, PSQ 05.pdf).
# Structure: Document control table, Governance table, Section 12 bullet list.
# ══════════════════════════════════════════════════════════════════════════════
def gov_tbl_psq(doc, rows):
    t = tbl(doc, len(rows) + 1, 2)
    t.cell(0, 0).text = "Role"; t.cell(0, 1).text = "Responsibility"
    for ri, (role, resp) in enumerate(rows, 1):
        t.cell(ri, 0).text = role; t.cell(ri, 1).text = resp

def gen_psq5():
    doc = Document()

    # Title
    h(doc, 0, "Policy on Clinical Audit")
    p(doc, HN)

    # Document control
    h(doc, 1, "Document control")
    doc_ctrl(doc, "HCO/PSQ/POL/05", "Quality Coordinator")
    p(doc, "A blank marked ________ must be completed before issue.")

    # Statement of intent
    h(doc, 1, "Statement of intent")
    p(doc, f"{HN} uses clinical audits as a quality improvement tool to improve the quality of patient care.")
    p(doc, "Clinical audits use predefined parameters, documented findings and remedial measures to complete the audit cycle.")

    # 1. Purpose
    h(doc, 1, "1. Purpose")
    p(doc, "The purpose of this policy is to establish the process for conducting clinical audits to improve the quality of patient care.")
    p(doc, "This policy does not cover matters outside clinical audit.")

    # 2. Scope
    h(doc, 1, "2. Scope")
    p(doc, f"This policy applies to clinical audits conducted by {HN} as per the scope of services.")
    p(doc, "It covers clinical and nursing care and personnel involved in clinical audits.")

    # 3. Policy standards
    h(doc, 1, "3. Policy standards")

    h(doc, 2, "3.1 Clinical audit as a quality improvement tool")
    p(doc, "The organisation shall use clinical audits as a quality improvement tool to improve the quality of patient care.")
    p(doc, "The clinical audit could be retrospective or prospective in nature.")
    p(doc, "The topic for audit could be disease-based, cost-based, community-based or based on morbidity, including length of stay.")
    p(doc, "The organisation shall conduct one clinical audit per clinical department once in two years as per the scope of services.")
    p(doc, "The organisation needs to take care to differentiate clinical audit from research projects.")

    h(doc, 2, "3.2 Predefined audit parameters")
    p(doc, "Clinical audits are standards-based and must be done using predefined parameters so that there is no bias.")
    p(doc,
      "The organisation shall lay down the objectives, the standards against which the audit will be conducted, "
      "a checklist where required, sampling and data-collection guidelines, and preparation of the report.")
    p(doc, "The audit shall encompass aspects of clinical and nursing care.")

    h(doc, 2, "3.3 Personnel participation")
    p(doc, "The organisation shall identify personnel for participation in clinical audit.")
    p(doc, "It could be a mix of clinicians, administrators and nurses.")
    p(doc, "These could be members of the core committee, quality assurance committee, etc.")

    h(doc, 2, "3.4 Patient and staff anonymity")
    p(doc,
      "Names of patients and hospital staff who may figure in audit documents must not be disclosed or "
      "referenced in public discussions or conferences.")
    p(doc, "This applies at both the stage of report preparation and dissemination.")
    p(doc, "Staff participating in the audit shall maintain patient and staff anonymity and not reveal names.")

    h(doc, 2, "3.5 Audit documentation")
    p(doc,
      "The organisation could use a checklist with the predefined parameters, and audit findings could be "
      "recorded on this sheet.")
    p(doc, "After the audit, a report shall be prepared, highlighting the key findings of the audit.")

    h(doc, 2, "3.6 Remedial measures and audit cycle")
    p(doc,
      "All remedial measures as ascertained shall be documented and implemented, and improvements thereof "
      "recorded to complete the audit cycle.")
    p(doc, "This shall preferably be done based on root-cause analysis.")

    # 4. Non-negotiable rules
    h(doc, 1, "4. Non-negotiable rules")
    ln(doc,
       "Do not conduct clinical audits less frequently than one clinical audit per clinical department once "
       "in two years as per the scope of services.")
    ln(doc,
       "Do not conduct a clinical audit without predefined objectives, standards, a checklist where required, "
       "sampling and data-collection guidelines, and preparation of the report.")
    ln(doc, "Do not omit clinical and nursing care from the audit.")
    ln(doc,
       "Do not disclose patient or hospital staff names from audit documents during report preparation or "
       "dissemination, or reveal those names in public discussions or conferences.")
    ln(doc, "Do not complete a clinical audit without preparing the required post-audit report highlighting key findings.")
    ln(doc,
       "Do not leave ascertained remedial measures undocumented or unimplemented, and do not leave the "
       "resulting improvements unrecorded.")

    # 5. What we do
    h(doc, 1, "5. What we do")

    h(doc, 2, "5.1 Plan and conduct clinical audits")
    p(doc, "The organisation uses clinical audits as a quality improvement tool to improve patient care.")
    p(doc,
      "Audits could be retrospective or prospective and topics could be disease-based, cost-based, "
      "community-based or based on morbidity, including length of stay.")
    p(doc, "One clinical audit is conducted per clinical department once in two years as per the scope of services.")
    p(doc, "Clinical audit is distinguished from research projects.")

    h(doc, 2, "5.2 Define audit parameters")
    p(doc, "Predefined parameters are used for standards-based audits.")
    p(doc,
      "The organisation lays down objectives, standards, a checklist where required, sampling and "
      "data-collection guidelines, and preparation of the report.")
    p(doc, "Audits encompass clinical and nursing care.")

    h(doc, 2, "5.3 Identify audit personnel")
    p(doc, "The organisation identifies personnel participating in clinical audit.")
    p(doc,
      "The personnel mix and committee mechanism are selected by the organisation; clinicians, administrators "
      "and nurses, or committee membership, could be used.")

    h(doc, 2, "5.4 Maintain anonymity")
    p(doc, "Patient and staff names are not disclosed in audit documents or referenced in public discussions or conferences.")
    p(doc, "Anonymity is maintained during report preparation and dissemination.")
    p(doc, "Participating staff shall not reveal names.")

    h(doc, 2, "5.5 Document the audit")
    p(doc, "A checklist with predefined parameters could be used and findings could be recorded on it.")
    p(doc, "After the audit, a report shall be prepared highlighting key findings.")

    h(doc, 2, "5.6 Complete the audit cycle")
    p(doc, "All ascertained remedial measures are documented and implemented.")
    p(doc, "Improvements are recorded to complete the audit cycle.")
    p(doc, "Root-cause analysis is preferably used as the basis.")

    # 6. Governance and responsibility — proper table
    h(doc, 1, "6. Governance and responsibility")
    gov_tbl_psq(doc, [
        ("Organisation",
         "Defines audit parameters, identifies participating personnel, conducts the required audits and "
         "ensures documentation and remedial measures."),
        ("Audit personnel",
         "Identified personnel participate in clinical audits and maintain patient and staff anonymity."),
        ("Clinical departments",
         "Each clinical department participates in the required audit cycle according to the scope of services."),
        ("Responsible personnel",
         "Ensure findings are reported and remedial measures are documented, implemented and improvements recorded."),
    ])

    # 7. Quality monitoring
    h(doc, 1, "7. Quality monitoring")
    mon_rows = [
        ("Clinical audit frequency",
         "One clinical audit per clinical department once in two years as per scope of services"),
        ("Audit parameters",
         "Objectives, standards, checklist where required, sampling/data-collection guidelines and report preparation"),
        ("Care coverage",   "Clinical and nursing care"),
        ("Personnel",       "Identified personnel participating in clinical audit"),
        ("Anonymity",       "Patient and staff anonymity during report preparation and dissemination"),
        ("Documentation",   "Post-audit report highlighting key findings"),
        ("Remedial measures", "Documentation, implementation and recording of improvements"),
    ]
    mon = tbl(doc, len(mon_rows) + 1, 2)
    mon.cell(0, 0).text = "Monitoring area"
    mon.cell(0, 1).text = "What is monitored"
    for ri, (area, what) in enumerate(mon_rows, 1):
        mon.cell(ri, 0).text = area
        mon.cell(ri, 1).text = what

    # 8. Training and staff acknowledgement
    h(doc, 1, "8. Training and staff acknowledgement")
    p(doc, "Staff participating in clinical audits are familiar with the processes described in this policy.")
    p(doc, f"I have read the Policy on Clinical Audit of {HN}. I will follow the processes described.")
    sig = tbl(doc, 4, 4)
    for ci, hdr in enumerate(("Staff name", "Designation", "Signature", "Date")):
        sig.cell(0, ci).text = hdr
    for ri in range(1, 4):
        for ci in range(4):
            sig.cell(ri, ci).text = "________"

    # 9. Distribution
    h(doc, 1, "9. Distribution")
    p(doc, "This policy shall be available to personnel involved in clinical audits.")
    p(doc, "The processes for clinical audit are communicated to relevant staff.")

    # 10. Abbreviations
    h(doc, 1, "10. Abbreviations")
    abbrev_tbl(doc, [
        ("QA", "Quality Assurance"),
    ])

    # 11. Traceability table
    h(doc, 1, "11. Traceability table")
    p(doc,
      "This table is an index. It is not how the policy is organised. An asterisk in the Level "
      "column means documentation of the process is required.")
    tr = tbl(doc, 7, 3)
    for ci, hdr in enumerate(("Objective Element", "Level", "Traceability to this policy")):
        tr.cell(0, ci).text = hdr
    trace_rows = [
        ("PSQ.5.a", "Commitment",
         "Sections 3.1 and 5.1 address clinical audits, the exact frequency, optional audit "
         "types/topics, and distinction from research projects."),
        ("PSQ.5.b", "Commitment",
         "Sections 3.2 and 5.2 address predefined parameters, objectives, standards, checklist "
         "where required, sampling/data-collection guidelines, report preparation, and clinical/nursing care."),
        ("PSQ.5.c", "Achievement",
         "Sections 3.3 and 5.3 address identification of participating personnel while preserving "
         "optional personnel composition and committee mechanism."),
        ("PSQ.5.d", "Commitment",
         "Sections 3.4 and 5.4 address anonymity during report preparation and dissemination and "
         "non-disclosure of names."),
        ("PSQ.5.e", "Commitment",
         "Sections 3.5 and 5.5 address documentation, optional checklist use and mandatory post-audit report."),
        ("PSQ.5.f", "Commitment",
         "Sections 3.6 and 5.6 address documentation, implementation and recording of improvements, "
         "with root-cause analysis preferred rather than mandatory."),
    ]
    for ri, (oe, lvl, txt) in enumerate(trace_rows, 1):
        tr.cell(ri, 0).text = oe
        tr.cell(ri, 1).text = lvl
        tr.cell(ri, 2).text = txt

    # 12. Required Records/Evidence Checklist — bulleted list
    h(doc, 1, "12. Required Records/Evidence Checklist")

    h(doc, 2, "Clinical audit planning and frequency")
    lb(doc, "Clinical audit schedule showing one clinical audit per clinical department once in two years as per the scope of services.")
    lb(doc, "Audit plan showing the selected clinical audit topic and approach.")

    h(doc, 2, "Audit parameters")
    lb(doc, "Documented audit objectives and standards.")
    lb(doc, "Checklist where required.")
    lb(doc, "Sampling and data-collection guidelines.")
    lb(doc, "Documentation showing clinical and nursing care are included.")

    h(doc, 2, "Personnel and anonymity")
    lb(doc, "List of personnel identified to participate in clinical audits.")
    lb(doc,
       "Audit documents and processes showing patient and staff names are not disclosed during "
       "report preparation or dissemination.")

    h(doc, 2, "Audit documentation")
    lb(doc, "Post-audit report highlighting key findings.")
    lb(doc, "Checklist or other audit finding record where the organisation has chosen to use one.")

    h(doc, 2, "Remedial measures")
    lb(doc, "Documented ascertained remedial measures.")
    lb(doc, "Records showing remedial measures were implemented.")
    lb(doc, "Records showing improvements were recorded to complete the audit cycle.")
    lb(doc, "Root-cause analysis where it has been used as the basis for remedial measures.")

    # 13. References
    h(doc, 1, "13. References")
    ln(doc,
       "National Accreditation Board for Hospitals and Healthcare Providers. NABH Accreditation "
       "Standards for Hospitals, 6th Edition. PSQ.5.")
    ln(doc, "Guidebook interpretation supplied for PSQ.5.a through PSQ.5.f.")

    # Disclaimer
    h(doc, 1, "Disclaimer")
    p(doc,
      "This policy reorganises the supplied PSQ.5 objective-element wording and Guidebook "
      "interpretation into plain-language policy format. The modal strength of the source has been "
      "preserved. No additional frequency, deadline, threshold or prescribed methodology has been "
      "added beyond the supplied source.")

    save_and_verify(doc, "HCO_PSQ_5_v2_REWRITE_DRAFT.docx")


# ══════════════════════════════════════════════════════════════════════════════
# PSQ.6 — Safety Culture, Leadership and Resources   (no stop-work)
# Content: ChatGPT final draft (approved, PSQ 06.pdf).
# Structure: Document control table, Governance table, Section 12 bullet list.
# ══════════════════════════════════════════════════════════════════════════════
def gen_psq6():
    doc = Document()

    # Title
    h(doc, 0, "Policy on Safety Culture, Leadership and Resources")
    p(doc, HN)

    # Document control
    h(doc, 1, "Document control")
    doc_ctrl(doc, "HCO/PSQ/POL/06", "Medical Superintendent")
    p(doc, "A blank marked ________ must be completed before issue.")

    # Statement of intent
    h(doc, 1, "Statement of intent")
    p(doc,
      f"{HN} promotes a culture of safety through safety-promoting behaviours, leadership "
      "awareness and involvement, adequate resources, performance targets, and use of staff "
      "feedback to improve patient safety and the quality improvement programme.")

    # 1. Purpose
    h(doc, 1, "1. Purpose")
    p(doc,
      "The purpose of this policy is to establish the approach for safety culture, leadership "
      "awareness and involvement, resource availability, performance improvement targets and "
      "use of staff feedback.")
    p(doc, "This policy does not cover matters outside these areas.")

    # 2. Scope
    h(doc, 1, "2. Scope")
    p(doc,
      f"This policy applies to management, leaders at all levels, departmental leaders, staff "
      f"and personnel involved in patient safety and the quality improvement programme at {HN}.")

    # 3. Policy standards
    h(doc, 1, "3. Policy standards")

    h(doc, 2, "3.1 Safety culture")
    p(doc, "The management needs to ensure the adoption of behaviours that promote patient safety.")
    p(doc,
      "Key features required for a culture of safety include sharing information, reporting "
      "occurrences of incidents, learning from safety incident analysis, blame-free culture and "
      "encouragement of collaboration across disciplines and departments.")
    p(doc,
      "The key components of patient safety culture are informed culture, reporting culture, "
      "learning culture, just culture and flexible culture.")
    p(doc, "The management needs to measure its safety culture regularly, at least once a year.")
    p(doc, "This shall be measured using validated surveys.")
    p(doc,
      "Examples include the Manchester Patient Safety Framework (MaPSaF), Safety Attitudes "
      "Questionnaire and AHRQ Surveys on Patient Safety Culture (SOPS).")
    p(doc, "The management shall act on its patient safety culture assessment results.")

    h(doc, 2, "3.2 Leadership awareness")
    p(doc,
      "The organisation and departmental leaders are aware of the patient safety and quality "
      "improvement programme, its intent and applicability to their respective areas and how it "
      "contributes to the organisation as a whole.")

    h(doc, 2, "3.3 Departmental leadership involvement")
    p(doc, "Departmental leaders are involved in patient safety and quality improvement.")
    p(doc,
      "Each department could have department objectives or key performance indicators including "
      "clinical indicators.")
    p(doc, "The responsibility of achieving them could be that of the leader or designated person.")
    p(doc, "This could be discussed during departmental meetings.")

    h(doc, 2, "3.4 Resources")
    p(doc,
      "Appropriate fund allocation is done by the organisation for the smooth functioning of the "
      "patient safety and quality improvement programme.")
    p(doc, "The budget could be earmarked based on previous year's spending.")
    p(doc,
      "If no data is available, the organisation could make a beginning by earmarking a budget "
      "but reviewing it at the end of six months to make any necessary modifications.")
    p(doc,
      "The management makes available adequate resources required for the patient safety and "
      "quality improvement programme.")
    p(doc,
      "Resources shall include men, material, machine, money, milieu, measurement and method.")
    p(doc,
      "These shall be in steady supply to ensure that the programme functions smoothly.")

    h(doc, 2, "3.5 Organisational performance improvement targets")
    p(doc,
      "The management shall identify the organisation and department level quality objectives, "
      "set targets, monitor them at least once in three months and modify the target at least "
      "annually.")
    p(doc,
      "The targets shall be shared with the faculty and staff and regular feedback taken.")

    h(doc, 2, "3.6 Workforce feedback")
    p(doc,
      "The feedback shall be obtained from the staff on their understanding and use of the "
      "safety and quality systems.")
    p(doc, "Feedback may be obtained once a year through staff surveys.")
    p(doc,
      "It is also important that staff workforce feel able to raise concerns whenever they occur.")
    p(doc,
      "These inputs shall be used to improve patient safety and quality improvement programmes.")

    # 4. Non-negotiable rules
    h(doc, 1, "4. Non-negotiable rules")
    ln(doc, "Do not measure safety culture less frequently than at least once a year.")
    ln(doc, "Do not measure safety culture without using a validated survey.")
    ln(doc,
       "Do not leave patient safety culture assessment results without action by management.")
    ln(doc,
       "Do not omit any of the seven required resource categories: men, material, machine, "
       "money, milieu, measurement and method, when providing resources for the patient safety "
       "and quality improvement programme.")
    ln(doc,
       "Do not allow the required resources to fall outside the steady supply required for the "
       "programme to function smoothly.")
    ln(doc,
       "Do not monitor organisational and departmental performance improvement targets less "
       "frequently than at least once in three months.")
    ln(doc,
       "Do not leave performance improvement targets unmodified beyond the required annual "
       "modification.")
    ln(doc,
       "Do not withhold organisational and departmental performance improvement targets from "
       "faculty and staff.")
    ln(doc,
       "Do not fail to obtain staff feedback on their understanding and use of the safety and "
       "quality systems.")
    ln(doc,
       "Do not fail to use the staff feedback obtained to improve patient safety and quality "
       "improvement programmes.")

    # 5. What we do
    h(doc, 1, "5. What we do")

    h(doc, 2, "5.1 Build and measure safety culture")
    p(doc,
      "Management ensures adoption of behaviours that promote patient safety.")
    p(doc,
      "Information sharing, incident reporting, learning from safety incident analysis, a "
      "blame-free culture and collaboration across disciplines and departments support the "
      "culture of safety.")
    p(doc,
      "The target culture is described through informed culture, reporting culture, learning "
      "culture, just culture and flexible culture.")
    p(doc, "Safety culture is measured at least once a year using a validated survey.")
    p(doc,
      "MaPSaF, Safety Attitudes Questionnaire and AHRQ Surveys on Patient Safety Culture "
      "(SOPS) are examples of validated survey instruments; the organisation selects its "
      "validated survey.")
    p(doc, "Management acts on the patient safety culture assessment results.")

    h(doc, 2, "5.2 Maintain leadership awareness")
    p(doc,
      "Organisation and departmental leaders are aware of the patient safety and quality "
      "improvement programme, its intent, applicability to their areas and contribution to the "
      "organisation as a whole.")

    h(doc, 2, "5.3 Involve departmental leaders")
    p(doc, "Departmental leaders are involved in patient safety and quality improvement.")
    p(doc,
      "Departments could use objectives or key performance indicators including clinical "
      "indicators, with responsibility assigned to a leader or designated person and discussion "
      "during departmental meetings.")

    h(doc, 2, "5.4 Provide adequate resources")
    p(doc,
      "The organisation makes appropriate fund allocation for the smooth functioning of the "
      "programme.")
    p(doc, "Management makes adequate resources available.")
    p(doc,
      "The resources shall include men, material, machine, money, milieu, measurement and "
      "method, in steady supply.")
    p(doc,
      "Budgeting based on previous-year spending and a six-month review where no data is "
      "available could be used.")

    h(doc, 2, "5.5 Set and monitor improvement targets")
    p(doc,
      "Management identifies organisation- and department-level quality objectives and sets "
      "targets.")
    p(doc,
      "Targets are monitored at least once in three months and modified at least annually.")
    p(doc, "Targets are shared with faculty and staff and regular feedback is taken.")

    h(doc, 2, "5.6 Obtain and use staff feedback")
    p(doc,
      "Staff feedback is obtained on understanding and use of safety and quality systems.")
    p(doc, "An annual staff survey may be used to obtain feedback.")
    p(doc, "Staff should feel able to raise concerns whenever they occur.")
    p(doc,
      "Inputs obtained are used to improve patient safety and quality improvement programmes.")

    # 6. Governance and responsibility — proper table
    h(doc, 1, "6. Governance and responsibility")
    gov_tbl_psq(doc, [
        ("Management",
         "Ensures safety culture measurement and action on results, makes adequate resources "
         "available, and identifies, monitors and modifies performance improvement targets."),
        ("Organisation and departmental leaders",
         "Leaders are aware of the patient safety and quality improvement programme, its intent, "
         "applicability and contribution to the organisation. Departmental leaders are involved "
         "in patient safety and quality improvement."),
        ("Staff",
         "Faculty and staff receive organisational and departmental performance improvement "
         "targets and provide regular feedback. Staff feedback on understanding and use of "
         "safety and quality systems is obtained."),
        ("Responsible personnel",
         "Responsible personnel support implementation of the patient safety and quality "
         "improvement programme and use feedback to improve it."),
    ])

    # 7. Quality monitoring
    h(doc, 1, "7. Quality monitoring")
    mon_rows = [
        ("Safety culture",
         "Measurement using a validated survey at least once a year and action on results"),
        ("Leadership awareness",
         "Awareness of programme intent, applicability and contribution to the organisation"),
        ("Departmental leadership",
         "Involvement in patient safety and quality improvement"),
        ("Resources",
         "Men, material, machine, money, milieu, measurement and method in steady supply"),
        ("Performance targets",
         "Organisation and department objectives and targets; monitoring at least once in "
         "three months and modification at least annually"),
        ("Target communication",
         "Targets shared with faculty and staff and regular feedback taken"),
        ("Staff feedback",
         "Feedback on understanding and use of safety and quality systems and its use for "
         "improvement"),
    ]
    mon = tbl(doc, len(mon_rows) + 1, 2)
    mon.cell(0, 0).text = "Monitoring area"
    mon.cell(0, 1).text = "What is monitored"
    for ri, (area, what) in enumerate(mon_rows, 1):
        mon.cell(ri, 0).text = area
        mon.cell(ri, 1).text = what

    # 8. Training and staff acknowledgement
    h(doc, 1, "8. Training and staff acknowledgement")
    p(doc,
      "Staff involved in patient safety and the quality improvement programme are familiar "
      "with the processes described in this policy.")
    p(doc,
      f"I have read the Policy on Safety Culture, Leadership and Resources of {HN}. "
      "I will follow the processes described.")
    sig = tbl(doc, 4, 4)
    for ci, hdr in enumerate(("Staff name", "Designation", "Signature", "Date")):
        sig.cell(0, ci).text = hdr
    for ri in range(1, 4):
        for ci in range(4):
            sig.cell(ri, ci).text = "________"

    # 9. Distribution
    h(doc, 1, "9. Distribution")
    p(doc,
      "This policy shall be available to management, leaders, departmental leaders and staff "
      "involved in patient safety and the quality improvement programme.")

    # 10. Abbreviations
    h(doc, 1, "10. Abbreviations")
    abbrev_tbl(doc, [
        ("MaPSaF", "Manchester Patient Safety Framework"),
        ("AHRQ",   "Agency for Healthcare Research and Quality"),
        ("SOPS",   "Surveys on Patient Safety Culture"),
    ])

    # 11. Traceability table
    h(doc, 1, "11. Traceability table")
    p(doc,
      "This table is an index. It is not how the policy is organised. An asterisk in the Level "
      "column means documentation of the process is required.")
    tr = tbl(doc, 7, 3)
    for ci, hdr in enumerate(("Objective Element", "Level", "Traceability to this policy")):
        tr.cell(0, ci).text = hdr
    trace_rows = [
        ("PSQ.6.a", "Achievement",
         "Sections 3.1 and 5.1 address safety-promoting behaviours, annual validated-survey "
         "measurement and action on assessment results."),
        ("PSQ.6.b", "Commitment",
         "Sections 3.2 and 5.2 address awareness among organisation and departmental leaders "
         "of the programme, its intent, applicability and contribution."),
        ("PSQ.6.c", "Commitment",
         "Sections 3.3 and 5.3 address departmental leader involvement while preserving the "
         "illustrative nature of KPIs, designated responsibility and departmental meetings."),
        ("PSQ.6.d", "Commitment",
         "Sections 3.4 and 5.4 address adequate resources and the seven required resource "
         "categories in steady supply while preserving advisory budgeting mechanisms."),
        ("PSQ.6.e", "Achievement",
         "Sections 3.5 and 5.5 address quality objectives, targets, monitoring at least once "
         "in three months, modification at least annually, sharing and feedback."),
        ("PSQ.6.f", "Excellence",
         "Sections 3.6 and 5.6 address mandatory staff feedback and its use for improvement "
         "while preserving the optional annual survey mechanism."),
    ]
    for ri, (oe, lvl, txt) in enumerate(trace_rows, 1):
        tr.cell(ri, 0).text = oe
        tr.cell(ri, 1).text = lvl
        tr.cell(ri, 2).text = txt

    # 12. Required Records/Evidence Checklist — bulleted list
    h(doc, 1, "12. Required Records/Evidence Checklist")

    h(doc, 2, "Safety culture")
    lb(doc, "Validated safety-culture survey used at least once a year.")
    lb(doc, "Safety-culture assessment results and management actions taken on the results.")
    lb(doc,
       "Materials showing safety-promoting behaviours, information sharing, incident reporting, "
       "learning from safety incident analysis, blame-free culture and collaboration across "
       "disciplines and departments.")

    h(doc, 2, "Leadership")
    lb(doc,
       "Information showing organisation and departmental leaders understand the patient safety "
       "and quality improvement programme, its intent, applicability to their areas and "
       "contribution to the organisation as a whole.")
    lb(doc,
       "Documents showing departmental leader involvement in patient safety and quality "
       "improvement.")
    lb(doc,
       "Department objectives or key performance indicators, leader/designated-person "
       "responsibility or departmental meeting discussion where the organisation has chosen to "
       "use these mechanisms.")

    h(doc, 2, "Resources")
    lb(doc,
       "Budget or fund-allocation information for the patient safety and quality improvement "
       "programme.")
    lb(doc,
       "Resources showing men, material, machine, money, milieu, measurement and method are "
       "available in steady supply.")
    lb(doc,
       "Previous-year spending basis or six-month budget review where the organisation has "
       "chosen to use these mechanisms.")

    h(doc, 2, "Performance improvement targets")
    lb(doc, "Organisation-level and department-level quality objectives and targets.")
    lb(doc,
       "Monitoring records showing targets are monitored at least once in three months.")
    lb(doc, "Records showing targets are modified at least annually.")
    lb(doc, "Targets shared with faculty and staff and regular feedback taken.")

    h(doc, 2, "Staff feedback")
    lb(doc, "Staff feedback on understanding and use of the safety and quality systems.")
    lb(doc,
       "Records showing staff feedback is used to improve patient safety and quality "
       "improvement programmes.")
    lb(doc,
       "Annual staff survey where the organisation has chosen to use this mechanism.")

    # 13. References
    h(doc, 1, "13. References")
    ln(doc,
       "National Accreditation Board for Hospitals and Healthcare Providers. NABH Accreditation "
       "Standards for Hospitals, 6th Edition. PSQ.6.")
    ln(doc, "Guidebook interpretation supplied for PSQ.6.a through PSQ.6.f.")

    # Disclaimer
    h(doc, 1, "Disclaimer")
    p(doc,
      "This policy reorganises the supplied PSQ.6 objective-element wording and Guidebook "
      "interpretation into plain-language policy format. The modal strength of the source has "
      "been preserved. Named survey instruments and departmental mechanisms are treated as "
      "examples where the source uses illustrative language. The exact mandatory frequencies "
      "and the seven resource categories have been retained. No additional frequency, deadline, "
      "threshold or prescribed mechanism has been added beyond the supplied source.")

    save_and_verify(doc, "HCO_PSQ_6_v2_REWRITE_DRAFT.docx")


# ══════════════════════════════════════════════════════════════════════════════
# PSQ.7 — Incident Management and Safety Events   (HAS stop-work: PSQ.7.a CORE)
# Content: ChatGPT final draft (approved, PSQ 07.pdf).
# Structure: Document control table, Stop-work Section 6, Governance table,
#            Section 13 bullet checklist.
# ══════════════════════════════════════════════════════════════════════════════
def gen_psq7():
    doc = Document()

    # Title
    h(doc, 0, "Policy on Incident Management and Safety Events")
    p(doc, HN)

    # Document control
    h(doc, 1, "Document control")
    doc_ctrl(doc, "HCO/PSQ/POL/07", "Patient Safety Officer")
    p(doc, "A blank marked ________ must be completed before issue.")

    # Statement of intent
    h(doc, 1, "Statement of intent")
    p(doc,
      f"{HN} implements an incident management system covering identification, reporting, "
      "review and action on incidents. The system supports factual reporting and learning "
      "and is based on the principle of just culture.")
    p(doc,
      "The organisation identifies relevant sentinel events, analyses incidents, takes "
      "corrective and preventive action, manages unidentified risks revealed through incident "
      "analysis, and informs relevant stakeholders after due analysis and classification.")

    # 1. Purpose
    h(doc, 1, "1. Purpose")
    p(doc,
      "The purpose of this policy is to establish the process for identification, reporting, "
      "review, analysis and management of incidents and relevant safety events.")
    p(doc, "This policy does not cover matters outside incident management.")

    # 2. Scope
    h(doc, 1, "2. Scope")
    p(doc,
      f"This policy applies to incidents occurring within {HN} and to personnel involved "
      "in incident reporting, review, analysis, corrective and preventive action, risk "
      "management and stakeholder communication.")
    p(doc,
      "It includes sentinel events, near misses and adverse events as described in the "
      "source requirements.")

    # 3. Policy standards
    h(doc, 1, "3. Policy standards")

    h(doc, 2, "3.1 Incident management system")
    p(doc,
      "The organisation implements an incident management system covering identification, "
      "reporting, review and action on incidents.")
    p(doc,
      "The system supports factual reporting and learning and is based on the principle of "
      "just culture.")
    p(doc,
      "The organisation shall have a mechanism for reporting the occurrence of incidents on "
      "standardised incident report forms.")
    p(doc,
      "While capturing incidents, the organisation shall capture all incidents without going "
      "into the severity or whether harm was caused.")
    p(doc,
      "It is preferable that the reporting system is simple, clear, confidential and focused "
      "on process improvement.")

    h(doc, 2, "3.2 Identification of sentinel events")
    p(doc,
      "Sentinel events relating to system or process deficiencies that are relevant and "
      "important to the organisation must be clearly defined.")
    p(doc, "The list of identified and relevant sentinel events shall be documented.")
    p(doc, "The definition of sentinel events is available in the Glossary.")

    h(doc, 2, "3.3 Analysis of incidents")
    p(doc, "The safety committee shall be responsible for incident analysis.")
    p(doc, "Root-cause identification could preferably be used.")
    p(doc, "Inputs could be sought from the units, disciplines or departments concerned.")
    p(doc,
      "Where possible, patients and other stakeholders could be included in analysing the "
      "feedback and complaint.")
    p(doc,
      "The immediate response to a safety incident shall be to address the urgent care and "
      "support needs of those involved. This shall not await analysis.")
    p(doc,
      "For sentinel events, correction, if any, shall be initiated within 24 working hours "
      "of occurrence or reporting.")
    p(doc,
      "The analysis of sentinel events shall be completed within seven working days of "
      "occurrence or reporting.")

    h(doc, 2, "3.4 Corrective and preventive action")
    p(doc,
      "All corrective and preventive action based on incident analysis shall be documented.")
    p(doc,
      "The findings and recommendations arrived at after analysis shall be communicated to "
      "all personnel concerned to correct systems and processes and prevent recurrences.")
    p(doc,
      "Any change in policy or procedure shall be reflected as an amendment in the "
      "organisation's documentation.")

    h(doc, 2, "3.5 Risk management")
    p(doc,
      "If incident analysis reveals the presence of an unidentified risk, the same shall be "
      "subjected to risk management.")

    h(doc, 2, "3.6 Stakeholder communication")
    p(doc,
      "After due analysis, an incident could be termed as a near miss, adverse event or "
      "sentinel event.")
    p(doc,
      "Based on the nature of the near miss, adverse event or sentinel event, the organisation "
      "shall inform the stakeholders, including patient and family where applicable, regarding "
      "the relevant concerns in addition to initiating corrective and preventive action.")

    # 4. Non-negotiable rules
    h(doc, 1, "4. Non-negotiable rules")
    lb(doc,
       "Do not operate the incident management system without identification, reporting, "
       "review and action on incidents.")
    lb(doc,
       "Do not report incidents without using the organisation's standardised incident report "
       "forms.")
    lb(doc,
       "Do not exclude an incident from capture based on its severity or whether harm was "
       "caused.")
    lb(doc,
       "Do not leave relevant and important sentinel events undefined or their identified "
       "list undocumented.")
    lb(doc,
       "Do not allow the immediate care and support response to a safety incident to wait "
       "for incident analysis.")
    lb(doc,
       "Do not initiate sentinel-event correction later than within 24 working hours of "
       "occurrence or reporting.")
    lb(doc,
       "Do not complete sentinel-event analysis later than within seven working days of "
       "occurrence or reporting.")
    lb(doc, "Do not leave corrective and preventive action undocumented.")
    lb(doc,
       "Do not communicate incident-analysis findings and recommendations only to the "
       "immediate team; they shall be communicated to all personnel concerned.")
    lb(doc,
       "Do not leave changes to policy or procedure unrecorded as amendments in the "
       "organisation's documentation.")
    lb(doc,
       "Do not leave an unidentified risk revealed by incident analysis outside the risk "
       "management system.")
    lb(doc,
       "Do not inform stakeholders about a near miss, adverse event or sentinel event before "
       "due analysis and classification; the required stakeholder communication occurs after "
       "due analysis.")

    # 5. What we do
    h(doc, 1, "5. What we do")

    h(doc, 2, "5.1 Operate the incident management system")
    p(doc,
      "The organisation operates an incident management system covering identification, "
      "reporting, review and action.")
    p(doc,
      "Standardised incident report forms are used for reporting the occurrence of incidents.")
    p(doc, "All incidents are captured regardless of severity or whether harm was caused.")
    p(doc,
      "The reporting system preferably remains simple, clear, confidential and focused on "
      "process improvement.")

    h(doc, 2, "5.2 Identify sentinel events")
    p(doc,
      "The organisation defines sentinel events relating to system or process deficiencies "
      "that are relevant and important to the organisation.")
    p(doc, "The identified and relevant sentinel-event list is documented.")
    p(doc, "The definition of sentinel events is referred to in the Glossary.")

    h(doc, 2, "5.3 Analyse incidents")
    p(doc, "The safety committee is responsible for incident analysis.")
    p(doc, "Root-cause identification could preferably be used.")
    p(doc, "Inputs could be obtained from the concerned units, disciplines or departments.")
    p(doc,
      "Where possible, patients and other stakeholders could be included in analysing "
      "feedback and complaints.")
    p(doc,
      "The immediate response to a safety incident addresses urgent care and support needs "
      "and does not await analysis.")
    p(doc,
      "For sentinel events, correction, if any, is initiated within 24 working hours of "
      "occurrence or reporting.")
    p(doc,
      "Sentinel-event analysis is completed within seven working days of occurrence or "
      "reporting.")

    h(doc, 2, "5.4 Take corrective and preventive action")
    p(doc, "Corrective and preventive action based on incident findings is documented.")
    p(doc, "Findings and recommendations are communicated to all personnel concerned.")
    p(doc, "Systems and processes are corrected to prevent recurrences.")
    p(doc,
      "Changes to policies or procedures are reflected as amendments in organisational "
      "documentation.")

    h(doc, 2, "5.5 Manage unidentified risks")
    p(doc,
      "When incident analysis reveals an unidentified risk, that risk is subjected to risk "
      "management.")

    h(doc, 2, "5.6 Inform stakeholders")
    p(doc,
      "After due analysis, an incident could be termed a near miss, adverse event or "
      "sentinel event.")
    p(doc,
      "Based on the nature of the event, relevant concerns are communicated to stakeholders, "
      "including patient and family where applicable, in addition to initiating corrective "
      "and preventive action.")

    # 6. Stop-work authority
    h(doc, 1, "6. Stop-work authority")

    h(doc, 2, "6.1 Incident management system trigger")
    p(doc,
      "Absence of an implemented incident management system is CORE non-compliance. The "
      "incident management system must include identification, reporting, review and action "
      "on incidents. Failure to have an implemented system covering these components "
      "constitutes the specified stop-work trigger.")

    # 7. Governance and responsibility — proper table
    h(doc, 1, "7. Governance and responsibility")
    gov_tbl_psq(doc, [
        ("Management",
         "Ensures implementation of the incident management system and the processes "
         "described in this policy."),
        ("Safety committee",
         "Shall be responsible for incident analysis."),
        ("Personnel involved in reporting and review",
         "Capture all incidents and participate in the incident management process."),
        ("Personnel concerned with corrective and preventive action",
         "Receive findings and recommendations arising from incident analysis and participate "
         "in correction of systems and processes to prevent recurrences."),
        ("Risk management responsibility",
         "When incident analysis reveals an unidentified risk, the risk is subjected to risk "
         "management."),
        ("Stakeholder communication responsibility",
         "The organisation informs stakeholders, including patient and family where applicable, "
         "after due analysis and classification of the incident."),
    ])

    # 8. Quality monitoring
    h(doc, 1, "8. Quality monitoring")
    mon_rows = [
        ("Incident management system",
         "Identification, reporting, review and action on incidents"),
        ("Incident reporting",
         "Use of standardised incident report forms and capture of all incidents regardless "
         "of severity or harm"),
        ("Sentinel events",
         "Organisation-defined relevant sentinel events and documented list"),
        ("Immediate safety response",
         "Urgent care and support provided without awaiting analysis"),
        ("Sentinel-event correction",
         "Correction initiated within 24 working hours of occurrence or reporting"),
        ("Sentinel-event analysis",
         "Analysis completed within seven working days of occurrence or reporting"),
        ("Corrective and preventive action",
         "Documentation, communication to all personnel concerned and documented amendments "
         "to policies or procedures"),
        ("Risk management",
         "Unidentified risks revealed through incident analysis subjected to risk management"),
        ("Stakeholder communication",
         "Stakeholders informed after due analysis and classification, including patient and "
         "family where applicable"),
    ]
    mon = tbl(doc, len(mon_rows) + 1, 2)
    mon.cell(0, 0).text = "Monitoring area"
    mon.cell(0, 1).text = "What is monitored"
    for ri, (area, what) in enumerate(mon_rows, 1):
        mon.cell(ri, 0).text = area
        mon.cell(ri, 1).text = what

    # 9. Training and staff acknowledgement
    h(doc, 1, "9. Training and staff acknowledgement")
    p(doc,
      "Staff involved in incident reporting, review, analysis, corrective and preventive "
      "action, risk management and stakeholder communication shall be familiar with the "
      "processes described in this policy.")
    p(doc,
      f"I have read the Policy on Incident Management and Safety Events of {HN}. "
      "I will follow the processes described.")
    sig = tbl(doc, 4, 4)
    for ci, hdr in enumerate(("Staff name", "Designation", "Signature", "Date")):
        sig.cell(0, ci).text = hdr
    for ri in range(1, 4):
        for ci in range(4):
            sig.cell(ri, ci).text = "________"

    # 10. Distribution
    h(doc, 1, "10. Distribution")
    p(doc,
      "This policy shall be available to personnel involved in incident management and "
      "related patient safety activities.")
    p(doc,
      "The processes described in this policy shall be communicated to relevant personnel.")

    # 11. Abbreviations
    h(doc, 1, "11. Abbreviations")
    abbrev_tbl(doc, [
        ("CAPA", "Corrective and Preventive Action"),
    ])

    # 12. Traceability table
    h(doc, 1, "12. Traceability table")
    p(doc,
      "This table is an index. It is not how the policy is organised. An asterisk in the "
      "Level column means documentation of the process is required.")
    tr = tbl(doc, 7, 3)
    for ci, hdr in enumerate(("Objective Element", "Level", "Traceability to this policy")):
        tr.cell(0, ci).text = hdr
    trace_rows = [
        ("PSQ.7.a", "CORE*",
         "Sections 3.1, 4, 5.1 and 6 address the incident management system, its four "
         "components, standardised incident reporting, capture of all incidents and the "
         "specified CORE stop-work trigger."),
        ("PSQ.7.b", "Commitment*",
         "Sections 3.2 and 5.2 address organisation-defined relevant sentinel events and "
         "documentation of the identified list."),
        ("PSQ.7.c", "Commitment",
         "Sections 3.3 and 5.3 address incident analysis, safety committee responsibility, "
         "immediate care and support, and the exact sentinel-event correction and analysis "
         "timeframes."),
        ("PSQ.7.d", "Commitment",
         "Sections 3.4 and 5.4 address documented corrective and preventive action, "
         "communication to all personnel concerned and documented amendments to policies "
         "and procedures."),
        ("PSQ.7.e", "Achievement",
         "Sections 3.5 and 5.5 address unidentified risks revealed through incident analysis "
         "and their submission to risk management."),
        ("PSQ.7.f", "Commitment",
         "Sections 3.6 and 5.6 address stakeholder communication after due analysis and "
         "classification, including patient and family where applicable."),
    ]
    for ri, (oe, lvl, txt) in enumerate(trace_rows, 1):
        tr.cell(ri, 0).text = oe
        tr.cell(ri, 1).text = lvl
        tr.cell(ri, 2).text = txt

    # 13. Required Records/Evidence Checklist — bulleted list
    h(doc, 1, "13. Required Records/Evidence Checklist")

    h(doc, 2, "Incident management and reporting")
    lb(doc,
       "Implemented incident management process covering identification, reporting, review "
       "and action on incidents.")
    lb(doc, "Standardised incident report forms.")
    lb(doc,
       "Incident reports showing that all incidents are captured regardless of severity or "
       "whether harm was caused.")
    lb(doc,
       "Incident reporting process supporting factual reporting, learning and just culture.")
    lb(doc,
       "Reporting process characteristics showing simplicity, clarity, confidentiality and "
       "focus on process improvement where these have been adopted.")

    h(doc, 2, "Sentinel events")
    lb(doc,
       "Organisation-defined list of relevant and important sentinel events relating to "
       "system or process deficiencies.")
    lb(doc,
       "Document showing that the identified sentinel-event list is defined and maintained.")
    lb(doc, "Reference to the Glossary for the definition of sentinel events.")

    h(doc, 2, "Incident analysis and response")
    lb(doc, "Safety committee responsibility for incident analysis.")
    lb(doc, "Incident-analysis records.")
    lb(doc,
       "Records of immediate care and support provided after safety incidents without "
       "waiting for analysis.")
    lb(doc,
       "Sentinel-event correction records showing initiation within 24 working hours of "
       "occurrence or reporting.")
    lb(doc,
       "Sentinel-event analysis records showing completion within seven working days of "
       "occurrence or reporting.")
    lb(doc, "Root-cause analysis records where this approach has been used.")
    lb(doc,
       "Inputs from concerned units, disciplines or departments where these have been "
       "obtained.")
    lb(doc,
       "Patient and other stakeholder participation in analysis where this has been used.")

    h(doc, 2, "Corrective and preventive action")
    lb(doc, "Documented corrective and preventive actions arising from incident analysis.")
    lb(doc, "Communication of findings and recommendations to all personnel concerned.")
    lb(doc, "Policy or procedure amendments made following incident analysis.")

    h(doc, 2, "Risk management")
    lb(doc, "Incident-analysis records identifying any previously unidentified risk.")
    lb(doc,
       "Risk-management records for unidentified risks revealed through incident analysis.")

    h(doc, 2, "Stakeholder communication")
    lb(doc,
       "Incident classification after due analysis as a near miss, adverse event or "
       "sentinel event where applicable.")
    lb(doc,
       "Communication records showing relevant concerns were shared with stakeholders, "
       "including patient and family where applicable.")

    # 14. References
    h(doc, 1, "14. References")
    ln(doc,
       "National Accreditation Board for Hospitals and Healthcare Providers. NABH "
       "Accreditation Standards for Hospitals, 6th Edition. PSQ.7.")
    ln(doc, "Guidebook interpretation supplied for PSQ.7.a through PSQ.7.f.")
    ln(doc, "NABH Glossary for the definition of sentinel events.")

    # Disclaimer
    h(doc, 1, "Disclaimer")
    p(doc,
      "This policy reorganises the supplied PSQ.7 objective-element wording and Guidebook "
      "interpretation into plain-language policy format. The modal strength of the source "
      "has been preserved. Optional and preferable mechanisms have not been converted into "
      "mandatory requirements. No specific sentinel-event categories or external statutory "
      "reporting categories have been added. The exact mandatory timeframes of 24 working "
      "hours and seven working days have been retained.")

    save_and_verify(doc, "HCO_PSQ_7_v2_REWRITE_DRAFT.docx")


# ══════════════════════════════════════════════════════════════════════════════
if __name__ == "__main__":
    gen_psq1()
    print("\nPSQ.1 draft generated.")
    gen_psq5()
    print("\nPSQ.5 draft generated.")
    gen_psq6()
    print("\nPSQ.6 draft generated.")
    gen_psq7()
    print("\nPSQ.7 draft generated.")
