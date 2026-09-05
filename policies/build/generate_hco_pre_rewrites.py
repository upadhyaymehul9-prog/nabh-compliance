# -*- coding: utf-8 -*-
"""
generate_hco_pre_rewrites.py
Generates HCO PRE chapter v2 rewrite-reference DOCX files.

Pipeline : python-docx, identical to generate_hco_ipc_rewrites.py.
Output   : policies/build/rewrite_reference/HCO_PRE_N_v2_REWRITE_DRAFT.docx
Source   : policies/build/pre_raw_dump_1-4.txt
"""
import os
from docx import Document

HN  = "«Hospital Name»"
OUT = os.path.join(os.path.dirname(os.path.abspath(__file__)), "rewrite_reference")
os.makedirs(OUT, exist_ok=True)


# ── Helpers ───────────────────────────────────────────────────────────────────

def h(doc, lv, txt):
    return doc.add_paragraph(txt, style={0: "Title", 1: "Heading 1", 2: "Heading 2"}[lv])

def p(doc, txt=""):
    return doc.add_paragraph(txt, style="Normal")

def ln(doc, txt):
    return doc.add_paragraph(txt, style="List Number")

def lb(doc, txt):
    return doc.add_paragraph(txt, style="List Bullet")

def tbl(doc, rows, cols):
    t = doc.add_table(rows=rows, cols=cols)
    try:
        t.style = "Table Grid"
    except KeyError:
        pass
    return t

def doc_ctrl(doc, no, prep, appr="Medical Superintendent"):
    dc = tbl(doc, 6, 4)
    for ri, (a, b, c, d) in enumerate([
        ("Document No.", no,          "Version",                "2.0"),
        ("Issue No.",    "01",         "Review due",             "One year from implementation"),
        ("Date created", "________",   "Date of implementation", "________"),
    ]):
        dc.cell(ri, 0).text = a; dc.cell(ri, 1).text = b
        dc.cell(ri, 2).text = c; dc.cell(ri, 3).text = d
    for ri, (lbl, txt) in enumerate([
        ("Prepared by", f"{prep}  Name: ________  Signature: ________"),
        ("Reviewed by", "Quality Coordinator  Name: ________  Signature: ________"),
        ("Approved by", f"{appr}  Name: ________  Signature: ________"),
    ], start=3):
        dc.cell(ri, 0).text = lbl
        c1 = dc.cell(ri, 1); c1.text = txt; c1.merge(dc.cell(ri, 3))

def abbrev_tbl(doc, rows):
    t = tbl(doc, len(rows) + 1, 2)
    t.cell(0, 0).text = "Abbreviation"; t.cell(0, 1).text = "Meaning"
    for ri, (a, m) in enumerate(rows, 1):
        t.cell(ri, 0).text = a; t.cell(ri, 1).text = m

def gov_tbl(doc, rows):
    t = tbl(doc, len(rows) + 1, 2)
    t.cell(0, 0).text = "Role"; t.cell(0, 1).text = "Responsibility"
    for ri, (role, resp) in enumerate(rows, 1):
        t.cell(ri, 0).text = role; t.cell(ri, 1).text = resp

def mon_tbl(doc, rows):
    t = tbl(doc, len(rows) + 1, 2)
    t.cell(0, 0).text = "Monitoring area"; t.cell(0, 1).text = "What is monitored"
    for ri, (area, what) in enumerate(rows, 1):
        t.cell(ri, 0).text = area; t.cell(ri, 1).text = what

def sig_tbl(doc):
    sig = tbl(doc, 4, 4)
    for ci, hdr in enumerate(("Staff name", "Designation", "Signature", "Date")):
        sig.cell(0, ci).text = hdr
    for ri in range(1, 4):
        for ci in range(4):
            sig.cell(ri, ci).text = "________"

def save_and_verify(doc, fname):
    import sys
    out = sys.stdout

    def pr(s):
        try:
            out.write(s + "\n")
        except UnicodeEncodeError:
            out.write(s.encode("ascii", "replace").decode() + "\n")

    pr(f"\n=== {fname} ===")
    for i, para in enumerate(doc.paragraphs[:80]):
        sn = para.style.name if para.style else "(None)"
        pr(f"{i:3d}  {sn!r:30s}  {para.text[:60]!r}")
    counts = {}
    for para in doc.paragraphs:
        sn = para.style.name if para.style else "(None)"
        counts[sn] = counts.get(sn, 0) + 1
    pr("  Style inventory:")
    for sn, n in sorted(counts.items()):
        pr(f"    {sn}: {n}")
    pr(f"  Total paras: {len(doc.paragraphs)}")
    path = os.path.join(OUT, fname)
    doc.save(path)
    pr(f"  Saved: {path}")


# ══════════════════════════════════════════════════════════════════════════════
# PRE.1 — Patient and Family Rights and Responsibilities   (NO stop-work)
# COREs: c, d, e | Stars: a*, b* | Achievement: b | Excellence: none
# Prepared by: Patient Rights Officer | Doc: HCO/PRE/POL/01
# ══════════════════════════════════════════════════════════════════════════════
def gen_pre1():
    doc = Document()

    h(doc, 0, "Policy on Patient and Family Rights and Responsibilities")
    p(doc, HN)

    h(doc, 1, "Document control")
    doc_ctrl(doc, "HCO/PRE/POL/01", "Patient Rights Officer")
    p(doc, "A blank marked ________ must be completed before issue.")

    h(doc, 1, "Statement of intent")
    p(doc,
      f"{HN} documents and displays patient and family rights and responsibilities, "
      "makes staff aware of their duty to protect these rights, actively promotes them, "
      "and provides a fair mechanism to report and act on any violation.")

    h(doc, 1, "1. Purpose")
    p(doc,
      f"This policy explains how {HN} protects and promotes patient and family rights "
      "and informs them about their responsibilities during care: documenting and "
      "displaying rights and responsibilities, actively promoting them through counselling "
      "and accessible material, protecting rights through staff conduct, providing a "
      "mechanism to report violations, and ensuring top leadership monitors and acts on "
      "violations.")
    p(doc,
      "Informed consent, privacy during examinations, and other decision-making rights "
      "are covered in the hospital's other policies.")

    h(doc, 1, "2. Scope")
    p(doc,
      f"This policy applies to all staff who interact with patients and families, the "
      "Patient Rights Officer, ward and OPD staff responsible for display and counselling, "
      "and top leadership responsible for monitoring rights violations.")

    h(doc, 1, "3. Policy standards")
    p(doc,
      f"{HN} documents, displays, and actively promotes patient and family rights and "
      "responsibilities in a format and language patients and families can understand, "
      "and protects these rights through staff awareness and conduct.")
    p(doc,
      "A mechanism exists for patients and families to report a violation of their "
      "rights, and top leadership monitors, investigates, and takes corrective and "
      "preventive action on every reported violation.")

    h(doc, 1, "4. Non-negotiable rules")
    lb(doc,
       "Do not display patient rights and responsibilities in a single language only "
       "— information, education and communication material shall be at least bilingual.")
    lb(doc,
       "Do not limit rights promotion to a static display board — in-patients shall "
       "receive counselling and out-patients shall have accessible educational material.")
    lb(doc,
       "Do not leave staff untrained on their duty to protect and respect patient and "
       "family rights.")
    lb(doc,
       "Do not operate without a mechanism for patients and families to report a "
       "rights violation.")
    lb(doc,
       "Do not leave a reported violation uninvestigated, undocumented, or without "
       "corrective/preventive action from top leadership.")

    h(doc, 1, "5. What we do")

    h(doc, 2, "5.1 Document and display patient and family rights and responsibilities")
    p(doc,
      "Patient and family rights and responsibilities are documented in consonance with "
      "the Charter of Patients’ Rights laid down by the statutory body. They are "
      "displayed prominently where visible to patients, families and visitors. Information, "
      "education and communication material is at least bilingual.")

    h(doc, 2, "5.2 Actively promote patient and family rights and responsibilities")
    p(doc,
      f"Beyond static display, {HN} takes active steps to promote rights and "
      "responsibilities. In-patients and/or their family are counselled on their rights "
      "and responsibilities in a format and language they can understand. For out-patients, "
      "educational material is easily accessible and prominently displayed in OPD.")

    h(doc, 2, "5.3 Protect patient and family rights through staff conduct")
    p(doc,
      "Staff are made aware of their responsibility in protecting and respecting patient "
      "and family rights. Staff conduct themselves in a manner that conveys a positive "
      "attitude toward the protection of these rights.")

    h(doc, 2, "5.4 Provide a mechanism to report rights violations")
    p(doc,
      "A mechanism exists for patients and/or family to report a violation of their "
      "rights. Violations are reported through an incident reporting form providing "
      "details of how the right was violated and, where applicable, by whom.")

    h(doc, 2, "5.5 Monitor, investigate and act on rights violations")
    p(doc,
      f"Top leadership documents violations of patient and family rights, investigates "
      "them, and maintains records of each incident and its outcome. Corrective and/or "
      f"preventive action is taken in a time frame defined by {HN}.")

    h(doc, 1, "6. Stop-work authority")
    p(doc,
      "Not applicable. This standard has no stop-work authority — violations of "
      "patient and family rights are addressed through the reporting mechanism "
      "(Section 5.4) and leadership review (Section 5.5), not through an "
      "immediate stop-work trigger.")

    h(doc, 1, "7. Governance and responsibility")
    gov_tbl(doc, [
        ("Medical Superintendent",
         "Accountable for ensuring patient and family rights are documented, displayed, "
         "promoted, and that reported violations receive leadership attention."),
        ("Patient Rights Officer",
         "Owns day-to-day implementation of this policy; maintains the rights charter, "
         "display boards, and counselling records; receives violation reports."),
        ("Top Leadership / Governing Body",
         "Reviews the violation log; investigates and closes CAPA arising from reported "
         "violations."),
        ("Nursing Superintendent",
         "Ensures ward staff are trained on patient rights and complete counselling "
         "for in-patients."),
        ("Quality Coordinator",
         "Audits this policy; holds training records and staff acknowledgements."),
    ])

    h(doc, 1, "8. Quality monitoring")
    mon_tbl(doc, [
        ("Display and documentation",
         "Rights and responsibilities documented per the Charter of Patients’ Rights; "
         "display boards current and bilingual."),
        ("Active promotion",
         "In-patient counselling records; OPD educational material accessibility."),
        ("Staff awareness",
         "Staff training records on rights protection."),
        ("Reporting mechanism",
         "Violation log entries; accessibility of the reporting mechanism to "
         "patients/families."),
        ("Leadership review",
         "Leadership review minutes; CAPA records with owners and closure evidence."),
    ])

    h(doc, 1, "9. Training and staff acknowledgement")
    p(doc,
      "All staff shall be familiar with patient and family rights and responsibilities, "
      "their duty to protect these rights, and the mechanism to report a violation, at "
      "induction and periodically.")
    p(doc,
      f"I have read the Policy on Patient and Family Rights and Responsibilities of "
      f"{HN} and understand my responsibilities under it.")
    sig_tbl(doc)

    h(doc, 1, "10. Distribution")
    p(doc,
      "This policy shall be available to all clinical and non-clinical staff, the "
      "Patient Rights Officer, ward and OPD in-charges, and top leadership.")

    h(doc, 1, "11. Abbreviations")
    abbrev_tbl(doc, [
        ("CAPA",  "Corrective and Preventive Action"),
        ("IEC",   "Information, Education and Communication"),
        ("NABH",  "National Accreditation Board for Hospitals and Healthcare Providers"),
        ("OPD",   "Out-Patient Department"),
    ])

    h(doc, 1, "12. Traceability table")
    p(doc,
      "This table is an index. It is not how the policy is organised. An asterisk in "
      "the Level column means the objective element is starred and requires documented "
      "evidence.")
    tr = tbl(doc, 6, 3)
    for ci, hdr in enumerate(("Objective Element", "Level", "Traceability to this policy")):
        tr.cell(0, ci).text = hdr
    trace_rows = [
        ("PRE.1.a", "Commitment*",
         "Sections 3 and 5.1 address documentation, display, and bilingual IEC material "
         "for patient and family rights and responsibilities."),
        ("PRE.1.b", "Achievement*",
         "Sections 3 and 5.2 address active promotion through in-patient counselling "
         "and accessible OPD material."),
        ("PRE.1.c", "CORE",
         "Sections 3 and 5.3 address staff awareness and conduct protecting patient "
         "and family rights."),
        ("PRE.1.d", "CORE",
         "Sections 3 and 5.4 address the mechanism for patients and families to report "
         "a rights violation."),
        ("PRE.1.e", "CORE",
         "Sections 3 and 5.5 address leadership documentation, investigation, and CAPA "
         "for reported violations."),
    ]
    for ri, (oe, lvl, txt) in enumerate(trace_rows, 1):
        tr.cell(ri, 0).text = oe
        tr.cell(ri, 1).text = lvl
        tr.cell(ri, 2).text = txt

    h(doc, 1, "13. Required Records/Evidence Checklist")

    h(doc, 2, "Documentation and display — PRE.1.a (Commitment*)")
    lb(doc,
       "Documented set of patient and family rights and responsibilities, in consonance "
       "with the Charter of Patients’ Rights.")
    lb(doc,
       "Display boards visible to patients, families and visitors, in at least two "
       "languages.")
    lb(doc, "Awareness records with date, language and named staff.")

    h(doc, 2, "Active promotion — PRE.1.b (Achievement*)")
    lb(doc, "In-patient counselling records (format, language, named staff).")
    lb(doc, "OPD display records confirming accessible educational material.")

    h(doc, 2, "Staff awareness and conduct — PRE.1.c (CORE)")
    lb(doc,
       "Staff training records on protecting and respecting patient and family rights.")
    lb(doc, "Conduct or incident investigation records where relevant.")

    h(doc, 2, "Reporting mechanism — PRE.1.d (CORE)")
    lb(doc,
       "Written mechanism for patients/families to report a violation (complaint box, "
       "helpline, form).")
    lb(doc,
       "Violation log with entries showing how the right was violated and, where "
       "applicable, by whom.")

    h(doc, 2, "Leadership monitoring and CAPA — PRE.1.e (CORE)")
    lb(doc, "Dated leadership review minutes.")
    lb(doc, "CAPA records with owners and closure evidence.")
    lb(doc, "Violation log referenced during leadership review.")

    h(doc, 1, "14. References")
    ln(doc,
       "National Accreditation Board for Hospitals and Healthcare Providers. NABH "
       "Accreditation Standards for Hospitals, 6th Edition. PRE.1.")
    ln(doc, "Guidebook interpretation supplied for PRE.1.a through PRE.1.e.")
    ln(doc, "Charter of Patients’ Rights, as laid down by the statutory body.")

    h(doc, 1, "Disclaimer")
    p(doc,
      f"This policy reorganises the supplied PRE.1 objective-element wording and "
      f"Guidebook interpretation into plain-language guidance for {HN}. It does not "
      "replace the NABH Accreditation Standards for Hospitals, 6th Edition, or the "
      "Guidebook, which remain the authoritative source in case of any conflict.")
    p(doc,
      "This policy is intended for accreditation and internal governance purposes and "
      "is not a statement of, or substitute for, compliance with any specific statute; "
      f"{HN} remains separately responsible for meeting all applicable statutory and "
      "regulatory requirements.")
    p(doc,
      "This policy shall be reviewed at least once every year, or earlier if NABH "
      "standards, the Charter of Patients’ Rights, or applicable law changes.")
    p(doc,
      f"This policy is the property of {HN} and is not to be reproduced or distributed "
      "outside the organisation without authorisation.")

    save_and_verify(doc, "HCO_PRE_1_v2_REWRITE_DRAFT.docx")


# ══════════════════════════════════════════════════════════════════════════════
# PRE.2 — Patient and Family Rights in Care and Decision-Making  (HAS stop-work)
# COREs: d, g | Stars: none | Achievement: l | Excellence: none
# Prepared by: Patient Rights Officer | Doc: HCO/PRE/POL/02
# Stop-work triggers: exam/procedure without privacy; neglect/abuse in progress;
#   transfusion/anaesthesia/surgery/research/invasive care without informed consent
# ══════════════════════════════════════════════════════════════════════════════
def gen_pre2():
    doc = Document()

    h(doc, 0, "Policy on Patient and Family Rights in Care and Decision-Making")
    p(doc, HN)

    h(doc, 1, "Document control")
    doc_ctrl(doc, "HCO/PRE/POL/02", "Patient Rights Officer")
    p(doc, "A blank marked ________ must be completed before issue.")

    h(doc, 1, "Statement of intent")
    p(doc,
      f"{HN} respects patient and family values, beliefs, dignity, privacy and "
      "confidentiality; protects patients from neglect or abuse; obtains informed consent "
      "before high-risk care; and supports the patient’s and family’s right to "
      "refuse treatment, seek a second opinion, complain, know treatment costs, access "
      "records, and be informed of their care.")

    h(doc, 1, "1. Purpose")
    p(doc,
      f"This policy explains how {HN} supports patient and family rights that involve "
      "individual beliefs, values, dignity, privacy, confidentiality, protection from "
      "neglect or abuse, informed consent before high-risk care, refusal of treatment, "
      "second opinion, complaints, cost information, record access, and information-"
      "sharing preferences.")
    p(doc,
      "The detailed informed-consent process, complaint redressal mechanism, and cost-"
      "estimate procedure are covered in the hospital’s other policies.")

    h(doc, 1, "2. Scope")
    p(doc,
      "This policy applies to all clinical and nursing staff involved in examination, "
      f"procedures, and treatment, and to all patients and families receiving care at "
      f"{HN}.")

    h(doc, 1, "3. Policy standards")
    p(doc,
      f"{HN} respects patient and family values, beliefs, cultural needs and spiritual "
      "requests; protects dignity and privacy during examination and procedures; protects "
      "patients from neglect and abuse, especially vulnerable patients; and treats patient "
      "information as confidential.")
    p(doc,
      "Informed consent is obtained before transfusion, anaesthesia, surgery, research, "
      "or other invasive/high-risk procedures; patients may refuse treatment, seek a "
      "second opinion, complain, and know treatment costs; and patients may access their "
      "clinical records and determine what information is shared with family.")

    h(doc, 1, "4. Non-negotiable rules")
    lb(doc,
       "Do not continue an examination, procedure or treatment that exposes the patient "
       "beyond what the procedure requires, or that leaves the patient without privacy.")
    lb(doc,
       "Do not leave a patient in a situation of neglect or abuse — stop the unsafe "
       "situation, protect the patient, and report it the same shift.")
    lb(doc,
       "Do not start transfusion of blood or blood components, anaesthesia, surgery, a "
       "research protocol, or any other invasive/high-risk procedure or treatment without "
       "informed consent, except where the documented emergency life-saving rule applies.")
    lb(doc,
       "Do not reveal confidential patient information, including HIV status, without "
       "the patient’s permission, and do not write or display confidential information "
       "where the public can read it.")
    lb(doc,
       "Do not photograph or record a procedure without explicit informed consent, and "
       "do not reveal the patient’s identity in any such material.")
    lb(doc,
       "Do not block a patient or family from seeking a second opinion from within or "
       "outside the organisation.")

    h(doc, 1, "5. What we do")

    h(doc, 2, "5.1 Respect values, beliefs, cultural and spiritual needs")
    p(doc,
      "Patient and family rights include respecting values and beliefs, special "
      "preferences, cultural needs, and responding to requests for spiritual needs — "
      "including, where relevant, how the patient wishes to be addressed, dietary "
      "preferences, worship requirements, and any specific requirement following death.")

    h(doc, 2, "5.2 Protect dignity and privacy during examination and procedures")
    p(doc,
      f"Staff ensure that the patient’s privacy and dignity are maintained during "
      f"all stages of care. {HN} has written guidelines for maintaining privacy and "
      "dignity. The patient is exposed only just before the actual procedure. Explicit "
      "informed consent is taken for photographs or recordings of procedures, and the "
      "patient’s identity is not revealed.")

    h(doc, 2, "5.3 Protect patients from neglect and abuse")
    p(doc,
      "Patient and family rights include protection from neglect or abuse. Special "
      "precautions are taken for vulnerable patients — including elderly, neonates, "
      "physically and mentally challenged patients, comatose patients, and patients under "
      "anaesthesia.")

    h(doc, 2, "5.4 Treat patient information as confidential")
    p(doc,
      "Effective measures are taken to maintain the confidentiality of all patient-"
      "related information. Staff avoid patient-related discussions in public places. "
      "Statutory requirements regarding privileged communication are followed. "
      "Confidential information, including HIV status, is not revealed without the "
      "patient’s permission, and is not written or displayed where the public can "
      "see it.")

    h(doc, 2, "5.5 Respect the right to refuse treatment")
    p(doc,
      "The treating doctor discusses all available options and allows the patient to "
      "make an informed choice. In case of refusal, the treating doctor explains the "
      "consequences of refusal and documents the discussion. After explanation, if the "
      "patient still refuses treatment, the refusal is respected.")

    h(doc, 2, "5.6 Support the right to a second opinion")
    p(doc,
      f"A mechanism exists for patients and families to seek a second opinion from "
      f"within or outside the organisation. {HN} respects the patient’s and "
      "family’s decision and facilitates access to relevant information or clinical "
      "evaluation when a second opinion is sought.")

    h(doc, 2, "5.7 Obtain informed consent before high-risk care")
    p(doc,
      "Informed consent is obtained before transfusion of blood and blood components, "
      "anaesthesia, surgery, initiation of any research protocol, and any other "
      "invasive/high-risk procedure or treatment. Consent is obtained by the treating "
      "doctor or a doctor member of the treating team. Those requirements are covered "
      "in the hospital’s other policies.")

    h(doc, 2, "5.8 Support the right to complain")
    p(doc,
      "Displayed patient rights include the right to make a complaint and the "
      "methodology to voice it. The complaint mechanism is accessible, and redressal "
      "of complaints is fair and transparent. Those requirements are covered in the "
      "hospital’s other policies.")

    h(doc, 2, "5.9 Provide cost information")
    p(doc,
      "Patients and families are explained about the expected costs of treatment — "
      "including consultations, procedures and investigations — in a transparent "
      "manner. Those requirements are covered in the hospital’s other policies.")

    h(doc, 2, "5.10 Provide access to clinical records")
    p(doc,
      "Every patient has access to their clinical record, in consonance with the Code "
      "of Medical Ethics and applicable statutory requirements. Those requirements are "
      "covered in the hospital’s other policies.")

    h(doc, 2, "5.11 Discuss the treating doctor, care plan and progress")
    p(doc,
      "Patients and families are given information on the name of the treating doctor, "
      "care plan, progress, and healthcare needs, discussed with them directly.")

    h(doc, 2, "5.12 Support patient control over information sharing")
    p(doc,
      "A mechanism exists for the patient to determine what information about their "
      "care is provided to self and to family. For minors, information is provided to "
      "at least one parent or guardian.")

    h(doc, 1, "6. Stop-work authority")
    p(doc,
      "Do not continue an examination, procedure or treatment that exposes the patient "
      "beyond what the procedure requires, or that leaves the patient without privacy "
      "(screens, drapes, closed door or equivalent).")
    p(doc,
      "Do not leave a patient in a situation of neglect or abuse (unattended fall risk, "
      "unwarranted repeated examination, manhandling, or failure to protect a vulnerable "
      "patient). Stop the unsafe situation, protect the patient, and report it the same "
      "shift.")
    p(doc,
      "Do not start transfusion of blood or blood components, anaesthesia, surgery, a "
      "research protocol, or any other invasive/high-risk procedure/treatment without "
      "informed consent obtained under the hospital’s informed consent policy, "
      "except where the documented emergency life-saving rule applies.")
    p(doc,
      "Stop-work applies to the examination or procedure start. Immediate life-saving "
      "measures continue while escalation happens, and are documented.")
    p(doc,
      "The person who stops tells the treating doctor and the Medical Superintendent "
      "the same shift. Refusing an unsafe examination or a procedure without consent "
      "is not a disciplinary matter.")

    h(doc, 1, "7. Governance and responsibility")
    gov_tbl(doc, [
        ("Medical Superintendent",
         "Accountable for ensuring patient rights in care and decision-making are "
         "respected across the organisation; receives stop-work escalations."),
        ("Patient Rights Officer",
         "Owns implementation of this policy; maintains privacy/confidentiality "
         "guidelines, awareness records, and second-opinion and complaint mechanisms."),
        ("Treating Doctor",
         "Obtains informed consent for high-risk care; discusses refusal of treatment "
         "and its consequences; documents all such discussions."),
        ("Nursing Superintendent",
         "Ensures staff maintain patient privacy, dignity, and confidentiality in "
         "daily practice; escalates stop-work triggers."),
        ("Quality Coordinator",
         "Audits this policy; holds training records, staff acknowledgements, and "
         "stop-work event logs."),
    ])

    h(doc, 1, "8. Quality monitoring")
    mon_tbl(doc, [
        ("Privacy and dignity",
         "Privacy guidelines on file; audit/observation records of privacy practice; "
         "consent records for photography."),
        ("Neglect/abuse protection",
         "Incident reports; vulnerable-patient protection records; stop-work event logs."),
        ("Confidentiality",
         "Confidentiality guidelines; staff training records; incident records involving "
         "disclosure."),
        ("Informed consent for high-risk care",
         "Consent forms for each listed act; stop-work event log if triggered."),
        ("Refusal of treatment",
         "Medical record entries documenting refusal discussions and outcomes."),
        ("Second opinion and complaints",
         "Records of facilitated second-opinion access; complaint mechanism records."),
        ("Records access and information sharing",
         "Access records; documented patient preferences on information sharing."),
    ])

    h(doc, 1, "9. Training and staff acknowledgement")
    p(doc,
      "All clinical and nursing staff shall be familiar with patient privacy, dignity, "
      "confidentiality, neglect/abuse protection, informed consent for high-risk care, "
      "and the stop-work triggers in this policy, at induction and periodically.")
    p(doc,
      f"I have read the Policy on Patient and Family Rights in Care and Decision-Making "
      f"of {HN} and understand my responsibilities under it.")
    sig_tbl(doc)

    h(doc, 1, "10. Distribution")
    p(doc,
      "This policy shall be available to all clinical and nursing staff, the Patient "
      "Rights Officer, department heads, and the Quality Coordinator.")

    h(doc, 1, "11. Abbreviations")
    abbrev_tbl(doc, [
        ("CAPA", "Corrective and Preventive Action"),
        ("HIV",  "Human Immunodeficiency Virus"),
        ("NABH", "National Accreditation Board for Hospitals and Healthcare Providers"),
    ])

    h(doc, 1, "12. Traceability table")
    p(doc,
      "This table is an index. It is not how the policy is organised. An asterisk in "
      "the Level column means the objective element is starred and requires documented "
      "evidence.")
    tr = tbl(doc, 13, 3)
    for ci, hdr in enumerate(("Objective Element", "Level", "Traceability to this policy")):
        tr.cell(0, ci).text = hdr
    trace_rows = [
        ("PRE.2.a", "Commitment",
         "Sections 3 and 5.1 address respect for values, beliefs, cultural needs and "
         "spiritual requests."),
        ("PRE.2.b", "Commitment",
         "Sections 3, 5.2 and 6 address privacy, dignity, and the stop-work trigger "
         "for exposure/privacy breach."),
        ("PRE.2.c", "Commitment",
         "Sections 3, 5.3 and 6 address protection from neglect/abuse and the stop-work "
         "trigger for active neglect/abuse."),
        ("PRE.2.d", "CORE",
         "Sections 3 and 5.4 address confidentiality of patient information, including "
         "HIV status."),
        ("PRE.2.e", "Commitment",
         "Section 5.5 addresses the right to refuse treatment."),
        ("PRE.2.f", "Commitment",
         "Section 5.6 addresses the right to a second opinion."),
        ("PRE.2.g", "CORE",
         "Sections 3, 5.7 and 6 address informed consent before high-risk care and the "
         "related stop-work trigger."),
        ("PRE.2.h", "Commitment",
         "Section 5.8 addresses the right to complain."),
        ("PRE.2.i", "Commitment",
         "Section 5.9 addresses cost information."),
        ("PRE.2.j", "Commitment",
         "Section 5.10 addresses access to clinical records."),
        ("PRE.2.k", "Commitment",
         "Section 5.11 addresses information on treating doctor, care plan and progress."),
        ("PRE.2.l", "Achievement",
         "Section 5.12 addresses patient control over information sharing with family."),
    ]
    for ri, (oe, lvl, txt) in enumerate(trace_rows, 1):
        tr.cell(ri, 0).text = oe
        tr.cell(ri, 1).text = lvl
        tr.cell(ri, 2).text = txt

    h(doc, 1, "13. Required Records/Evidence Checklist")

    h(doc, 2, "Values, beliefs and cultural needs — PRE.2.a (Commitment)")
    lb(doc,
       "Awareness records and preferences noted in the admission sheet/nursing kardex.")

    h(doc, 2, "Privacy and dignity — PRE.2.b (Commitment)")
    lb(doc, "Privacy guidelines document.")
    lb(doc, "Consent forms for photography/recording of procedures.")
    lb(doc, "Audit/observation records of privacy practice.")
    lb(doc, "Stop-work event log for exposure/privacy triggers.")

    h(doc, 2, "Neglect and abuse protection — PRE.2.c (Commitment)")
    lb(doc, "Incident reports for neglect/abuse events.")
    lb(doc, "Vulnerable-patient protection records.")
    lb(doc, "Stop-work event log for active neglect/abuse triggers.")

    h(doc, 2, "Confidentiality — PRE.2.d (CORE)")
    lb(doc, "Confidentiality guidelines.")
    lb(doc, "Staff training records on confidentiality.")
    lb(doc, "Incident records involving disclosure, if any.")

    h(doc, 2, "Refusal of treatment — PRE.2.e (Commitment)")
    lb(doc,
       "Medical record entries documenting options offered, consequences explained, "
       "and sustained refusal.")

    h(doc, 2, "Second opinion — PRE.2.f (Commitment)")
    lb(doc, "Written mechanism for seeking a second opinion.")
    lb(doc, "Records of facilitated second-opinion access.")

    h(doc, 2, "Informed consent for high-risk care — PRE.2.g (CORE)")
    lb(doc,
       "Consent forms on file for transfusion, anaesthesia, surgery, research, and "
       "other high-risk acts.")
    lb(doc, "Stop-work event log for missing-consent triggers.")

    h(doc, 2, "Complaints — PRE.2.h (Commitment)")
    lb(doc, "Display boards including complaint information.")
    lb(doc, "Complaint mechanism description.")

    h(doc, 2, "Cost information — PRE.2.i (Commitment)")
    lb(doc, "Records of cost explanation or written estimates provided to patients.")
    lb(doc, "Tariff availability mechanism.")

    h(doc, 2, "Records access — PRE.2.j (Commitment)")
    lb(doc, "Written record-access procedure.")
    lb(doc, "Records of access provided to patients.")

    h(doc, 2, "Treating doctor and care-plan information — PRE.2.k (Commitment)")
    lb(doc,
       "Admission notes and clinical records naming the treating doctor and documenting "
       "care-plan discussion.")

    h(doc, 2, "Information-sharing preferences — PRE.2.l (Achievement)")
    lb(doc,
       "Written mechanism for patients to determine information sharing with family.")
    lb(doc, "Records of patient-expressed preferences, including for minors.")

    h(doc, 1, "14. References")
    ln(doc,
       "National Accreditation Board for Hospitals and Healthcare Providers. NABH "
       "Accreditation Standards for Hospitals, 6th Edition. PRE.2.")
    ln(doc, "Guidebook interpretation supplied for PRE.2.a through PRE.2.l.")
    ln(doc, "Code of Medical Ethics, Medical Council of India.")

    h(doc, 1, "Disclaimer")
    p(doc,
      f"This policy reorganises the supplied PRE.2 objective-element wording and "
      f"Guidebook interpretation into plain-language guidance for {HN}. It does not "
      "replace the NABH Accreditation Standards for Hospitals, 6th Edition, or the "
      "Guidebook, which remain the authoritative source in case of any conflict.")
    p(doc,
      "This policy is intended for accreditation and internal governance purposes and "
      "is not a statement of, or substitute for, compliance with any specific statute; "
      f"{HN} remains separately responsible for meeting all applicable statutory and "
      "regulatory requirements.")
    p(doc,
      "This policy shall be reviewed at least once every year, or earlier if NABH "
      "standards, the Code of Medical Ethics, or applicable law changes.")
    p(doc,
      f"This policy is the property of {HN} and is not to be reproduced or distributed "
      "outside the organisation without authorisation.")

    save_and_verify(doc, "HCO_PRE_2_v2_REWRITE_DRAFT.docx")


# ══════════════════════════════════════════════════════════════════════════════
# PRE.3 — Informed Decision-Making, Care Planning and Patient Education  (NO stop-work)
# COREs: a | Stars: none | Achievement: b, e | Commitment: c, d
# Prepared by: Patient Rights Officer | Doc: HCO/PRE/POL/03
# ══════════════════════════════════════════════════════════════════════════════
def gen_pre3():
    doc = Document()

    h(doc, 0, "Policy on Informed Decision-Making, Care Planning and Patient Education")
    p(doc, HN)

    h(doc, 1, "Document control")
    doc_ctrl(doc, "HCO/PRE/POL/03", "Patient Rights Officer")
    p(doc, "A blank marked ________ must be completed before issue.")

    h(doc, 1, "Statement of intent")
    p(doc,
      f"{HN} explains proposed care, risks, benefits, alternatives, and expected outcomes "
      "to patients and families; involves them in preparing and modifying the care plan; "
      "keeps them informed of diagnostic results and any change in condition; and provides "
      "multi-disciplinary counselling where appropriate.")

    h(doc, 1, "1. Purpose")
    p(doc,
      f"This policy explains how {HN} educates patients and/or family members to make "
      "informed decisions and involves them in the care planning and delivery process: "
      "explaining proposed care and its risks, benefits and alternatives; involving them "
      "in the care plan; informing them of diagnostic results and condition changes; and "
      "providing multi-disciplinary counselling when appropriate.")
    p(doc,
      "The formal consent process for high-risk procedures is covered in the hospital's "
      "other policies.")

    h(doc, 1, "2. Scope")
    p(doc,
      f"This policy applies to attending doctors, treating physicians and their teams, "
      f"and all patients and/or family members receiving care at {HN}.")

    h(doc, 1, "3. Policy standards")
    p(doc,
      f"{HN}'s attending doctors explain the proposed care, its risks, benefits, "
      "alternatives, expected results and possible complications to the patient and/or "
      "family in a language they can understand, and involve them in preparing and "
      "modifying the care plan.")
    p(doc,
      "Patients and/or family members are informed about diagnostic results, the "
      "diagnosis, and any change in the patient's condition in a timely manner, and "
      "receive multi-disciplinary counselling when the situation calls for it.")

    h(doc, 1, "4. Non-negotiable rules")
    lb(doc,
       "Do not proceed with proposed care without explaining it, including risks, benefits "
       "and alternatives, to the patient and/or family in a language they can understand.")
    lb(doc,
       "Do not withhold explanation of diagnostic test results and the diagnosis from "
       "the patient and/or family.")
    lb(doc,
       "Do not delay explaining a change in the patient's condition — improvement, "
       "deterioration or complications — once the change is known.")
    lb(doc,
       "Do not exclude the multi-disciplinary team from counselling in situations that "
       "call for it.")

    h(doc, 1, "5. What we do")

    h(doc, 2, "5.1 Explain the proposed care, risks, benefits and alternatives")
    p(doc,
      "The proposed care, including referral to internal and/or external services, is "
      "discussed by the attending doctor with the patient and/or family members, in a "
      "language they can understand. Expected outcomes of treatment are explained at "
      "periodic intervals by the treating physician or their team. Possible complications "
      "are clearly communicated.")

    h(doc, 2, "5.2 Involve the patient and family in the care plan")
    p(doc,
      "The care plan is prepared and modified in consultation with the patient and/or "
      "family members. During preparation, the patient and/or family are explained the "
      "various treatment options, risks and benefits. Religious, cultural and spiritual "
      "views of the patient and/or family are considered during care delivery, within "
      "the limits of applicable statutory requirements.")

    h(doc, 2, "5.3 Inform patients and family of diagnostic results")
    p(doc,
      "Results of all diagnostic tests are explained, at least in broad terms, to the "
      "patient and family members, along with their implication on progress and treatment.")

    h(doc, 2, "5.4 Explain any change in the patient's condition")
    p(doc,
      "The patient and/or family members are explained about any change in the patient's "
      "condition — improvement, deterioration or complications — in a timely manner.")

    h(doc, 2, "5.5 Provide multi-disciplinary counselling when appropriate")
    p(doc,
      f"{HN} identifies situations requiring multi-disciplinary counselling — such as "
      "counselling for family members of critically ill patients, potential organ donors "
      "and/or their family, and long-stay patients. Members of the multi-disciplinary "
      "team — doctors from different specialties, nurses and physiotherapists — take "
      "part in the counselling as relevant to the situation.")

    h(doc, 1, "6. Stop-work authority")
    p(doc,
      "Not applicable. This standard has no stop-work authority. Gaps in explaining "
      "proposed care, care-plan consultation, or condition updates are addressed through "
      "the hospital's quality monitoring and CAPA process (Section 8), not through an "
      "immediate stop-work trigger. Consent requirements for high-risk procedures are "
      "covered in the hospital's other policies.")

    h(doc, 1, "7. Governance and responsibility")
    gov_tbl(doc, [
        ("Medical Superintendent",
         "Accountable for ensuring patients and families are informed and involved in "
         "decision-making across all departments."),
        ("Patient Rights Officer",
         "Owns implementation of this policy; monitors compliance with explanation and "
         "counselling requirements."),
        ("Attending/Treating Doctor",
         "Explains proposed care, risks, benefits, alternatives and outcomes; involves "
         "the patient/family in care-plan preparation; explains diagnostic results and "
         "condition changes; documents these discussions."),
        ("Nursing Superintendent",
         "Supports patient/family understanding through nursing communication and "
         "escalates unaddressed information gaps."),
        ("Quality Coordinator",
         "Audits this policy; holds training records and staff acknowledgements."),
    ])

    h(doc, 1, "8. Quality monitoring")
    mon_tbl(doc, [
        ("Explanation of proposed care",
         "Clinical notes/consultation records showing explanation of risks, benefits "
         "and alternatives."),
        ("Care-plan involvement",
         "Medical record notes of care-plan consultation with patient/family."),
        ("Diagnostic result communication",
         "Clinical notes showing diagnostic results were communicated."),
        ("Condition-change communication",
         "Clinical notes showing timely communication of condition changes."),
        ("Multi-disciplinary counselling",
         "Written situation list; counselling records showing attendance and participation."),
    ])

    h(doc, 1, "9. Training and staff acknowledgement")
    p(doc,
      "All attending doctors, treating physicians and their teams shall be familiar with "
      "the informed decision-making, care-planning, and patient-education requirements "
      "of this policy, at induction and periodically.")
    p(doc,
      f"I have read the Policy on Informed Decision-Making, Care Planning and Patient "
      f"Education of {HN} and understand my responsibilities under it.")
    sig_tbl(doc)

    h(doc, 1, "10. Distribution")
    p(doc,
      "This policy shall be available to all clinical staff, the Patient Rights Officer, "
      "department heads, and the Quality Coordinator.")

    h(doc, 1, "11. Abbreviations")
    abbrev_tbl(doc, [
        ("CAPA", "Corrective and Preventive Action"),
        ("NABH", "National Accreditation Board for Hospitals and Healthcare Providers"),
        ("OPD",  "Out-Patient Department"),
    ])

    h(doc, 1, "12. Traceability table")
    p(doc,
      "This table is an index. It is not how the policy is organised. An asterisk in "
      "the Level column means the objective element is starred and requires documented "
      "evidence.")
    tr = tbl(doc, 6, 3)
    for ci, hdr in enumerate(("Objective Element", "Level", "Traceability to this policy")):
        tr.cell(0, ci).text = hdr
    trace_rows = [
        ("PRE.3.a", "CORE",
         "Sections 3 and 5.1 address explanation of proposed care, risks, benefits, "
         "alternatives, expected results and complications."),
        ("PRE.3.b", "Achievement",
         "Sections 3 and 5.2 address patient/family involvement in care-plan preparation "
         "and modification."),
        ("PRE.3.c", "Commitment",
         "Sections 3 and 5.3 address explanation of diagnostic results and diagnosis."),
        ("PRE.3.d", "Commitment",
         "Sections 3 and 5.4 address timely explanation of condition changes."),
        ("PRE.3.e", "Achievement",
         "Sections 3 and 5.5 address multi-disciplinary counselling when appropriate."),
    ]
    for ri, (oe, lvl, txt) in enumerate(trace_rows, 1):
        tr.cell(ri, 0).text = oe
        tr.cell(ri, 1).text = lvl
        tr.cell(ri, 2).text = txt

    h(doc, 1, "13. Required Records/Evidence Checklist")

    h(doc, 2, "Explanation of proposed care — PRE.3.a (CORE)")
    lb(doc,
       "Clinical notes or consultation records showing explanation of proposed care, "
       "risks, benefits and alternatives was given.")

    h(doc, 2, "Care-plan involvement — PRE.3.b (Achievement)")
    lb(doc,
       "Medical record notes documenting care-plan consultation with the patient and/or "
       "family.")

    h(doc, 2, "Diagnostic result communication — PRE.3.c (Commitment)")
    lb(doc,
       "Clinical notes showing diagnostic results and their implications were communicated.")

    h(doc, 2, "Condition-change communication — PRE.3.d (Commitment)")
    lb(doc,
       "Clinical notes showing timely communication of condition changes, with dates.")

    h(doc, 2, "Multi-disciplinary counselling — PRE.3.e (Achievement)")
    lb(doc, "Written list of situations requiring multi-disciplinary counselling.")
    lb(doc, "Counselling records (attendees, date, what was discussed).")

    h(doc, 1, "14. References")
    ln(doc,
       "National Accreditation Board for Hospitals and Healthcare Providers. NABH "
       "Accreditation Standards for Hospitals, 6th Edition. PRE.3.")
    ln(doc, "Guidebook interpretation supplied for PRE.3.a through PRE.3.e.")

    h(doc, 1, "Disclaimer")
    p(doc,
      f"This policy reorganises the supplied PRE.3 objective-element wording and "
      f"Guidebook interpretation into plain-language guidance for {HN}. It does not "
      "replace the NABH Accreditation Standards for Hospitals, 6th Edition, or the "
      "Guidebook, which remain the authoritative source in case of any conflict.")
    p(doc,
      "This policy is intended for accreditation and internal governance purposes and "
      "is not a statement of, or substitute for, compliance with any specific statute; "
      f"{HN} remains separately responsible for meeting all applicable statutory and "
      "regulatory requirements.")
    p(doc,
      "This policy shall be reviewed at least once every year, or earlier if NABH "
      "standards change.")
    p(doc,
      f"This policy is the property of {HN} and is not to be reproduced or distributed "
      "outside the organisation without authorisation.")

    save_and_verify(doc, "HCO_PRE_3_v2_REWRITE_DRAFT.docx")


# ══════════════════════════════════════════════════════════════════════════════
# PRE.4 — Informed Consent   (HAS stop-work)
# COREs: a*, c, e | Stars: a*, d* | Commitment: b, d*
# Stop-work: procedure started without valid consent; consent by nurse/clerk without
#   performer explanation
# Prepared by: Patient Rights Officer | Doc: HCO/PRE/POL/04
# ══════════════════════════════════════════════════════════════════════════════
def gen_pre4():
    doc = Document()

    h(doc, 0, "Policy on Informed Consent")
    p(doc, HN)

    h(doc, 1, "Document control")
    doc_ctrl(doc, "HCO/PRE/POL/04", "Patient Rights Officer")
    p(doc, "A blank marked ________ must be completed before issue.")

    h(doc, 1, "Statement of intent")
    p(doc,
      f"{HN} obtains valid informed consent, from the patient or the person who may "
      "consent on their behalf, before any procedure on its written consent list — as a "
      "process of effective communication about the procedure, its risks, benefits, "
      "alternatives, and who will perform it, in a language the patient can understand.")

    h(doc, 1, "1. Purpose")
    p(doc,
      f"This policy explains how {HN} obtains informed consent from the patient or "
      "family about their care: maintaining a written list of procedures requiring "
      "consent, following statutory norms for the consent process, ensuring consent "
      "includes procedure information in an understandable language, describing who may "
      "consent when the patient cannot, and ensuring the person performing the procedure "
      "is responsible for the consent process.")
    p(doc,
      "The definition of specific invasive/high-risk procedures requiring consent under "
      "the hospital's rights policy is covered in the hospital's other policies.")

    h(doc, 1, "2. Scope")
    p(doc,
      "This policy applies to all doctors and clinical teams performing procedures on "
      f"the organisation's informed-consent list, and to all patients and persons who "
      f"may consent on a patient's behalf.")

    h(doc, 1, "3. Policy standards")
    p(doc,
      f"{HN} maintains a written list of procedures requiring informed consent and "
      "written guidance on the consent process, taking into account applicable statutory "
      "requirements such as the MTP Act, PC-PNDT Act, Transplantation of Human Organs "
      "Act, and HIV and AIDS (Prevention and Control) Act 2017/NACO HIV-testing policy "
      "where relevant to the services provided.")
    p(doc,
      "Consent is taken before the procedure, witnessed, includes the procedure's risks, "
      "benefits, alternatives and who will perform it, is obtained from the patient or "
      "the person who may consent when the patient is incapable, and is the "
      "responsibility of the person performing the procedure.")

    h(doc, 1, "4. Non-negotiable rules")
    lb(doc,
       "Do not start a procedure on the organisation's informed-consent list until valid "
       "informed consent has been obtained from the patient, or from the person who may "
       "consent when the patient is incapable of independent decision-making, in a "
       "language they can understand.")
    lb(doc,
       "Do not accept a consent form signed by a nurse or clerk as the only consent "
       "when the person performing the procedure (or a doctor member of that team) has "
       "not explained the procedure, its risks, benefits, alternatives and who will "
       "perform it.")
    lb(doc, "Do not allow anyone to consent on behalf of a competent adult.")
    lb(doc,
       "Do not use a consent validity period longer than six months for a procedure "
       "repeated over a long time without obtaining fresh consent.")
    lb(doc,
       "Do not change or add a treatment modality without obtaining fresh consent.")

    h(doc, 1, "5. What we do")

    h(doc, 2, "5.1 Maintain a written consent list and consent process guidance")
    p(doc,
      "A written list of procedures requiring informed consent is maintained, taking "
      "into account this standard's requirements and applicable statutory requirements. "
      "Written guidance explains the steps of the informed consent process and the "
      "person responsible. Staff are made aware of this guidance.")

    h(doc, 2, "5.2 Follow statutory norms in the consent process")
    p(doc,
      "Consent is taken before the procedure. At least one witness — present for the "
      "entire duration of the doctor-patient communication — signs the consent form. "
      "For procedures repeated over a long time (for example, dialysis), consent is "
      "taken at the first instance with a defined validity period not exceeding six "
      "months, and the patient endorses the consent at each repeat treatment. Fresh "
      "consent is obtained if there is a change in treatment modality or an addition "
      "of another modality.")

    h(doc, 2, "5.3 Ensure consent includes procedure information in an understandable language")
    p(doc,
      "Consent includes information regarding the procedure, its risks, benefits, "
      "alternatives, and who will perform it, in a language the patient can understand. "
      "The consent names the doctor performing the procedure; where multiple specialties "
      "are involved, it names the principal surgeon from each specialty, and each doctor "
      "explains their own role, risks, benefits and alternatives. If a doctor under "
      "training performs the procedure, this is specified along with the name of the "
      "supervising qualified doctor. The consent form is at a minimum bilingual, and "
      "where consent is taken in a language other than what the patient understands, "
      "the language of counselling and any interpreter used is clearly documented.")

    h(doc, 2, "5.4 Describe who may consent when the patient is incapable")
    p(doc,
      "Consent is taken from the patient when the patient is capable and above the "
      "legal age for giving consent; no one may consent on behalf of a competent adult. "
      f"When the patient is incapable of independent decision-making, {HN} follows "
      "applicable statutory norms and takes consent from the next of kin/legal guardian "
      "in the order: spouse, son/daughter, parents, brothers/sister. For life-threatening "
      "situations where the patient is incapable and next of kin is not available, the "
      "treating doctor and another clinician may decide to safeguard the patient's life; "
      "both names and the reason are documented.")

    h(doc, 2, "5.5 Ensure the performer is responsible for consent")
    p(doc,
      "The person performing the procedure is responsible for the entire consent process, "
      "including explanation and taking the signature. It is not acceptable for the "
      "person performing to only explain while a nurse or clerk takes the written consent. "
      "A doctor member of the performing team may take consent on behalf of the person "
      "performing the procedure.")

    h(doc, 1, "6. Stop-work authority")
    p(doc,
      "Do not start a procedure on the organisation's informed-consent list until valid "
      "informed consent has been obtained from the patient, or from the person who may "
      "consent when the patient is incapable of independent decision-making, in a "
      "language they can understand.")
    p(doc,
      "Do not accept a consent form signed by a nurse or clerk as the only consent "
      "when the person performing the procedure (or a doctor member of that team) has "
      "not explained the procedure, its risks, benefits, alternatives and who will "
      "perform it.")
    p(doc,
      "Stop-work applies to the procedure start. Immediate life-saving care when the "
      "patient is incapable and next of kin is not available follows the two-clinician "
      "emergency rule in Section 5.4, and is documented the same shift.")
    p(doc,
      "The person who stops tells the person performing the procedure and the Medical "
      "Superintendent the same shift. Refusing to start without valid consent is not "
      "a disciplinary matter.")

    h(doc, 1, "7. Governance and responsibility")
    gov_tbl(doc, [
        ("Medical Superintendent",
         "Accountable for ensuring the informed-consent list and process are maintained "
         "and followed; receives stop-work escalations."),
        ("Patient Rights Officer",
         "Owns implementation of this policy; maintains the consent list and process "
         "guidance; audits consent forms."),
        ("Person Performing the Procedure",
         "Responsible for the entire consent process — explanation and signature — for "
         "their own procedure; receives stop-work escalations for their procedure."),
        ("Nursing Superintendent",
         "Ensures nursing staff do not accept consent forms without the performer's "
         "explanation; escalates stop-work triggers."),
        ("Quality Coordinator",
         "Audits this policy; holds training records, staff acknowledgements, and "
         "stop-work event logs."),
    ])

    h(doc, 1, "8. Quality monitoring")
    mon_tbl(doc, [
        ("Consent list and guidance",
         "Written consent list current; process guidance available; staff awareness "
         "records."),
        ("Statutory norms",
         "Witness signatures; validity periods for repeat procedures; fresh-consent "
         "records on modality change."),
        ("Consent content and language",
         "Bilingual consent forms; documentation of performer name(s), risks, benefits, "
         "alternatives; interpreter records where used."),
        ("Consent for incapable patients",
         "Next-of-kin consent records showing order of preference followed; two-clinician "
         "emergency decision records."),
        ("Performer responsibility",
         "Consent forms showing the performer or team doctor as the responsible party; "
         "stop-work event log."),
    ])

    h(doc, 1, "9. Training and staff acknowledgement")
    p(doc,
      "All doctors and clinical teams who perform procedures on the informed-consent "
      "list shall be familiar with the consent process, statutory norms, and stop-work "
      "triggers in this policy, at induction and periodically.")
    p(doc,
      f"I have read the Policy on Informed Consent of {HN} and understand my "
      "responsibilities under it.")
    sig_tbl(doc)

    h(doc, 1, "10. Distribution")
    p(doc,
      "This policy shall be available to all doctors and clinical teams, the Patient "
      "Rights Officer, department heads, and the Quality Coordinator.")

    h(doc, 1, "11. Abbreviations")
    abbrev_tbl(doc, [
        ("CAPA",    "Corrective and Preventive Action"),
        ("MTP",     "Medical Termination of Pregnancy"),
        ("NABH",    "National Accreditation Board for Hospitals and Healthcare Providers"),
        ("NACO",    "National AIDS Control Organisation"),
        ("PC-PNDT", "Pre-Conception and Pre-Natal Diagnostic Techniques"),
    ])

    h(doc, 1, "12. Traceability table")
    p(doc,
      "This table is an index. It is not how the policy is organised. An asterisk in "
      "the Level column means the objective element is starred and requires documented "
      "evidence.")
    tr = tbl(doc, 6, 3)
    for ci, hdr in enumerate(("Objective Element", "Level", "Traceability to this policy")):
        tr.cell(0, ci).text = hdr
    trace_rows = [
        ("PRE.4.a", "CORE*",
         "Sections 3, 5.1 and 6 address the written consent list, process guidance, "
         "staff awareness, and the stop-work trigger for procedures started without "
         "consent."),
        ("PRE.4.b", "Commitment",
         "Sections 3 and 5.2 address statutory norms in the consent process, including "
         "witness and repeat-procedure validity requirements."),
        ("PRE.4.c", "CORE",
         "Sections 3, 5.3 and 6 address consent content, language, multi-specialty "
         "consent, and the stop-work trigger for consent taken without performer "
         "explanation."),
        ("PRE.4.d", "Commitment*",
         "Sections 3 and 5.4 address who may consent when the patient is incapable, "
         "including the order of preference and the two-clinician emergency rule."),
        ("PRE.4.e", "CORE",
         "Sections 3, 5.5 and 6 address performer responsibility for the consent "
         "process and the related stop-work trigger."),
    ]
    for ri, (oe, lvl, txt) in enumerate(trace_rows, 1):
        tr.cell(ri, 0).text = oe
        tr.cell(ri, 1).text = lvl
        tr.cell(ri, 2).text = txt

    h(doc, 1, "13. Required Records/Evidence Checklist")

    h(doc, 2, "Consent list and process guidance — PRE.4.a (CORE*)")
    lb(doc, "Written list of procedures requiring informed consent.")
    lb(doc, "Written consent process guidance naming the person responsible.")
    lb(doc, "Staff awareness records.")
    lb(doc, "Completed consent forms for listed procedures.")

    h(doc, 2, "Statutory norms in consent — PRE.4.b (Commitment)")
    lb(doc, "Consent forms showing pre-procedure signature and witness signature.")
    lb(doc,
       "Validity period documentation for repeat procedures (not exceeding six months).")
    lb(doc, "Patient endorsement records at repeat treatments.")
    lb(doc, "Fresh-consent records on treatment modality change.")

    h(doc, 2, "Consent content and language — PRE.4.c (CORE)")
    lb(doc,
       "Completed bilingual consent forms naming the performer(s), risks, benefits "
       "and alternatives.")
    lb(doc, "Interpreter records where used.")
    lb(doc, "Multi-specialty consent forms where applicable.")

    h(doc, 2, "Consent for incapable patients — PRE.4.d (Commitment*)")
    lb(doc,
       "Written description of who may consent when the patient is incapable, including "
       "order of preference.")
    lb(doc, "Consent forms showing next-of-kin signatory and relationship.")
    lb(doc, "Two-clinician emergency decision records, when applicable.")

    h(doc, 2, "Performer responsibility — PRE.4.e (CORE)")
    lb(doc, "Consent forms showing the performer or team doctor as the responsible party.")
    lb(doc, "Audit records of consent form review.")

    h(doc, 1, "14. References")
    ln(doc,
       "National Accreditation Board for Hospitals and Healthcare Providers. NABH "
       "Accreditation Standards for Hospitals, 6th Edition. PRE.4.")
    ln(doc, "Guidebook interpretation supplied for PRE.4.a through PRE.4.e.")
    ln(doc,
       "Medical Termination of Pregnancy Act; Pre-Conception and Pre-Natal Diagnostic "
       "Techniques Act, 1994; Transplantation of Human Organs Act; HIV and AIDS "
       "(Prevention and Control) Act, 2017 and the National AIDS Control Organisation's "
       "national HIV-testing policy — applicable to "
       f"{HN} only to the extent its services fall within the scope of these statutes.")

    h(doc, 1, "Disclaimer")
    p(doc,
      f"This policy reorganises the supplied PRE.4 objective-element wording and "
      f"Guidebook interpretation into plain-language guidance for {HN}. It does not "
      "replace the NABH Accreditation Standards for Hospitals, 6th Edition, or the "
      "Guidebook, which remain the authoritative source in case of any conflict.")
    p(doc,
      "This policy references certain statutes (MTP Act, PC-PNDT Act, Transplantation "
      "of Human Organs Act, HIV and AIDS (Prevention and Control) Act 2017) as examples "
      "of statutory requirements that may shape "
      f"{HN}'s consent list, where its services fall within their scope. This policy is "
      "not a statement of, or substitute for, compliance with any of these statutes; "
      f"{HN} remains separately responsible for meeting all applicable statutory and "
      "regulatory requirements.")
    p(doc,
      "This policy shall be reviewed at least once every year, or earlier if NABH "
      "standards or applicable law changes.")
    p(doc,
      f"This policy is the property of {HN} and is not to be reproduced or distributed "
      "outside the organisation without authorisation.")

    save_and_verify(doc, "HCO_PRE_4_v2_REWRITE_DRAFT.docx")


# ══════════════════════════════════════════════════════════════════════════════
# PRE.5 — Information and Education about Healthcare Needs  (NO stop-work)
# COREs: a | Stars: none | Achievement: i | Excellence: j
# Prepared by: Nursing Superintendent | Doc: HCO/PRE/POL/05
# ══════════════════════════════════════════════════════════════════════════════
def gen_pre5():
    doc = Document()

    h(doc, 0, "Policy on Information and Education about Healthcare Needs")
    p(doc, HN)

    h(doc, 1, "Document control")
    doc_ctrl(doc, "HCO/PRE/POL/05", "Nursing Superintendent")
    p(doc, "A blank marked ________ must be completed before issue.")

    h(doc, 1, "Statement of intent")
    p(doc,
      f"Patients and families at {HN} have a right to information and education about "
      "their healthcare needs, delivered in a way they can understand and act on.")

    h(doc, 1, "1. Purpose")
    p(doc,
      f"This policy sets out how {HN} educates patients and families — in a language "
      "and format they understand, on medication safety, food-drug interaction, diet and "
      "nutrition, immunisation, pain management, their specific disease process, and "
      "preventing healthcare-associated infections — and how special educational needs "
      "and patient engagement are addressed.")
    p(doc,
      "This policy does not cover informed consent, patient rights generally, or "
      "information on costs — those are covered in other hospital policies. The other "
      "PRE standards have their own policies too.")

    h(doc, 1, "2. Scope")
    p(doc,
      f"This policy applies to all patients and families receiving care at {HN}, across "
      "out-patient and in-patient settings.")

    h(doc, 1, "3. Policy standards")
    p(doc,
      f"{HN} educates patients and families in a language and format they can understand; "
      "about the safe and effective use of medication and its potential side effects when "
      "appropriate; about food-drug interaction; about diet and nutrition; about "
      "immunisations; on pain management techniques when appropriate; about their "
      "specific disease process, complications and prevention strategies; and about "
      "preventing healthcare-associated infections.")
    p(doc,
      f"{HN} identifies and addresses patients' and families' special educational needs, "
      "and has a mechanism to promote patient engagement to enhance clinical outcomes, "
      "safety and quality.")

    h(doc, 1, "4. Non-negotiable rules")
    lb(doc,
       "Do not educate a patient or family in a language or format they cannot understand "
       "— screen for understanding and language needs first.")
    lb(doc,
       "Do not skip educating a patient or family on the safe and effective use of "
       "medication and its potential side effects, where appropriate.")
    lb(doc,
       "Do not skip educating a patient or family about food-drug interaction for drugs "
       "on the hospital's identified list.")
    lb(doc,
       "Do not give diet and nutrition education that omits general recommendations for "
       "a healthy diet.")
    lb(doc,
       "Do not skip paediatric immunisation education that follows the universal "
       "immunisation programme.")
    lb(doc,
       "Do not deliver pain-management education outside the patient's personal, cultural "
       "and religious beliefs.")
    lb(doc,
       "Do not give disease-specific education that omits lifestyle modifications, diet "
       "changes or immunisation where appropriate.")
    lb(doc,
       "Do not skip educating a patient or family about preventing healthcare-associated "
       "infections.")
    lb(doc, "Do not identify a special educational need without addressing it.")
    lb(doc,
       "Staff who see a rule under this policy broken report it the same shift to the "
       "Nursing Superintendent or the Medical Superintendent.")

    h(doc, 1, "5. What we do")

    h(doc, 2, "5.1 Education in an understood language and format")
    p(doc,
      f"During treatment, patients and/or families at {HN} are screened — at least "
      "informally — for their understanding abilities and language requirements. Education "
      "is delivered through counselling, printed material or audio-visual aids in the "
      "language and format identified from that screening.")

    h(doc, 2, "5.2 Medication safety education")
    p(doc,
      f"{HN} maintains a list of drugs for which education on safe and effective use and "
      "potential side effects is required — for example digoxin — and educates patients "
      "and/or families against that list when those drugs are prescribed or dispensed, "
      "including the importance of taking a drug at a specific time where relevant.")

    h(doc, 2, "5.3 Food-drug interaction education")
    p(doc,
      f"{HN} maintains a list of drugs, or drug classes, for which food-drug interactions "
      "are clinically relevant — for example no alcohol when taking metronidazole — and "
      "educates patients and families about their diet during medication when a drug on "
      "that list is prescribed.")

    h(doc, 2, "5.4 Diet and nutrition education")
    p(doc,
      f"Education on diet and nutrition at {HN} incorporates general recommendations for "
      "following a healthy diet, and may also cover the relationship between specific "
      "foods or supplements and particular health conditions.")

    h(doc, 2, "5.5 Immunisation education")
    p(doc,
      f"Paediatric immunisation education at {HN} follows the universal immunisation "
      "programme. Adult patients are educated on relevant vaccines — for example "
      "influenza, pneumococcal, typhoid, hepatitis B, or meningococcal vaccines — as "
      "determined by the treating doctor based on the patient's condition and risk profile.")

    h(doc, 2, "5.6 Pain management education")
    p(doc,
      f"Where appropriate — typically for patients likely to have long-term pain because "
      "the underlying condition is not treatable — {HN} educates the patient and/or "
      "family on pain management techniques, within the framework of their personal, "
      "cultural and religious beliefs.")

    h(doc, 2, "5.7 Disease-specific education")
    p(doc,
      f"Education on a patient's specific disease process, complications and prevention "
      f"strategies at {HN} includes lifestyle modifications — stress management, physical "
      "exercise, cessation of smoking and substance abuse — and covers diet changes and "
      "immunisations where appropriate to the condition. This may be delivered through "
      "booklets, videos or leaflets.")

    h(doc, 2, "5.8 Healthcare-associated infection prevention education")
    p(doc,
      f"{HN} educates patients and families about preventing healthcare-associated "
      "infections — for example hand washing and avoiding overcrowding near the patient.")

    h(doc, 2, "5.9 Special educational needs")
    p(doc,
      "Special educational needs of patients and/or family members — for example related "
      "to ADHD, autism support, physical disabilities, speech and communication needs, "
      "or social and emotional health needs — are identified during treatment and "
      "addressed through counselling, printed material, audio-visual aids or other "
      "adapted means.")

    h(doc, 2, "5.10 Patient engagement")
    p(doc,
      f"{HN} works to create an enabling environment for partnership between patients, "
      "families, communities and healthcare providers, and encourages patients to become "
      "actively involved in reporting safety incidents, near misses and safety concerns. "
      "This may include disease-based patient support groups, a patient advisory council, "
      "and designated patient safety champions or advocates, where appropriate.")

    h(doc, 1, "6. Stop-work authority")
    p(doc,
      "Not applicable. This standard has no stop-work authority — gaps in patient and "
      "family education are addressed through the hospital's quality monitoring and CAPA "
      "process (Section 8), not through an immediate stop-work trigger.")

    h(doc, 1, "7. Governance and responsibility")
    gov_tbl(doc, [
        ("Medical Superintendent",
         "Accountable for ensuring patients and families receive education as required by "
         "this policy; receives escalations under Section 4."),
        ("Nursing Superintendent",
         "Owns implementation of this policy; coordinates patient and family education "
         "across departments; holds education records."),
        ("Pharmacy In-Charge",
         "Maintains the medication-safety and food-drug interaction education lists; "
         "keeps these current with the formulary."),
        ("Treating Doctor / Care Team",
         "Delivers or assigns disease-specific, pain-management and immunisation education "
         "relevant to the individual patient."),
        ("Quality Coordinator",
         "Audits this policy; holds training records and staff acknowledgements."),
    ])

    h(doc, 1, "8. Quality monitoring")
    mon_tbl(doc, [
        ("Language and format",
         "Evidence of language/format screening; education given in the identified "
         "language and format."),
        ("Medication safety",
         "Written drug-education list; counselling records for patients receiving "
         "listed drugs."),
        ("Food-drug interaction",
         "Written food-drug interaction list; counselling or diet-note records."),
        ("Diet and nutrition",
         "In-patient diet education records; out-patient diet counselling records."),
        ("Immunisation",
         "Adult immunisation advice records; paediatric immunisation counselling "
         "records referencing the universal immunisation programme."),
        ("Pain management",
         "Long-term pain education records for identified patients."),
        ("Disease-specific education",
         "Education records covering disease process, lifestyle, diet and immunisation."),
        ("HAI prevention",
         "Admission-checklist education records confirming person-to-person education."),
        ("Special educational needs",
         "Identified-need records; corresponding adapted-education or counselling records."),
        ("Patient engagement",
         "Engagement activity records — support groups, advisory council, safety "
         "champion appointments, patient safety incident reports received."),
    ])

    h(doc, 1, "9. Training and staff acknowledgement")
    p(doc,
      f"All staff covered by this policy are trained on the education requirements "
      "described in Section 5 at induction and periodically.")
    p(doc,
      f"I have read the Policy on Information and Education about Healthcare Needs of "
      f"{HN} and understand my responsibilities under it.")
    sig_tbl(doc)

    h(doc, 1, "10. Distribution")
    p(doc,
      "This policy shall be available to the Nursing department, Pharmacy, all clinical "
      "department heads, and all staff through the hospital intranet and department "
      "policy folders.")

    h(doc, 1, "11. Abbreviations")
    abbrev_tbl(doc, [
        ("ADHD", "Attention Deficit Hyperactivity Disorder"),
        ("CAPA", "Corrective and Preventive Action"),
        ("HAI",  "Healthcare-Associated Infection"),
        ("NABH", "National Accreditation Board for Hospitals and Healthcare Providers"),
        ("OPD",  "Out-Patient Department"),
        ("PREM", "Patient-Reported Experience Measure"),
    ])

    h(doc, 1, "12. Traceability table")
    p(doc,
      "This table is an index. It is not how the policy is organised. An asterisk in "
      "the Level column means the objective element is starred and requires documented "
      "evidence.")
    tr = tbl(doc, 11, 3)
    for ci, hdr in enumerate(("Objective Element", "Level", "Traceability to this policy")):
        tr.cell(0, ci).text = hdr
    trace_rows = [
        ("PRE.5.a", "CORE",
         "Sections 3 and 5.1 address education in a language and format the patient "
         "and/or family can understand, including informal screening for language needs."),
        ("PRE.5.b", "Commitment",
         "Sections 3 and 5.2 address education on safe and effective medication use and "
         "potential side effects, using the hospital's drug-education list."),
        ("PRE.5.c", "Commitment",
         "Sections 3 and 5.3 address education on food-drug interactions, using the "
         "hospital's food-drug interaction list."),
        ("PRE.5.d", "Commitment",
         "Sections 3 and 5.4 address diet and nutrition education, including general "
         "healthy-diet recommendations."),
        ("PRE.5.e", "Commitment",
         "Sections 3 and 5.5 address immunisation education, following the universal "
         "immunisation programme for paediatric patients."),
        ("PRE.5.f", "Commitment",
         "Sections 3 and 5.6 address pain-management education for patients likely to "
         "have long-term pain, within personal, cultural and religious beliefs."),
        ("PRE.5.g", "Commitment",
         "Sections 3 and 5.7 address disease-specific education, including lifestyle "
         "modifications, diet changes and immunisation where appropriate."),
        ("PRE.5.h", "Commitment",
         "Sections 3 and 5.8 address education on preventing healthcare-associated "
         "infections, including hand washing and avoiding overcrowding."),
        ("PRE.5.i", "Achievement",
         "Sections 3 and 5.9 address identification and addressing of patients' and "
         "families' special educational needs."),
        ("PRE.5.j", "Excellence",
         "Sections 3 and 5.10 address mechanisms to promote patient engagement, including "
         "support groups, advisory council and safety champions."),
    ]
    for ri, (oe, lvl, txt) in enumerate(trace_rows, 1):
        tr.cell(ri, 0).text = oe
        tr.cell(ri, 1).text = lvl
        tr.cell(ri, 2).text = txt

    h(doc, 1, "13. Required Records/Evidence Checklist")

    h(doc, 2, "Education language and format — PRE.5.a (CORE)")
    lb(doc,
       "Language/format screening record (admission note, nursing kardex or counselling "
       "record), education given, and the format used.")

    h(doc, 2, "Medication safety education — PRE.5.b (Commitment)")
    lb(doc, "Written drug-education list (current version).")
    lb(doc, "Counselling records for sampled patients who received a listed drug.")

    h(doc, 2, "Food-drug interaction education — PRE.5.c (Commitment)")
    lb(doc, "Written food-drug interaction list.")
    lb(doc, "Counselling or diet-note records for sampled patients who received a listed drug.")

    h(doc, 2, "Diet and nutrition education — PRE.5.d (Commitment)")
    lb(doc, "In-patient diet education records (for patients on therapeutic diets).")
    lb(doc, "Out-patient diet counselling records when ordered by the treating doctor.")

    h(doc, 2, "Immunisation education — PRE.5.e (Commitment)")
    lb(doc, "Records of immunisation advice given to adults, noting vaccines discussed or due.")
    lb(doc,
       "Paediatric immunisation counselling records referencing the universal immunisation "
       "programme.")

    h(doc, 2, "Pain management education — PRE.5.f (Commitment)")
    lb(doc,
       "Records of long-term pain-management education for identified patients, noting "
       "the treating doctor's rationale and the cultural/religious framework respected.")

    h(doc, 2, "Disease-specific education — PRE.5.g (Commitment)")
    lb(doc,
       "Education records for sampled patients, covering disease process, complications, "
       "lifestyle modifications, diet and immunisation.")
    lb(doc, "Supporting materials used (booklets, videos, leaflets).")

    h(doc, 2, "HAI prevention education — PRE.5.h (Commitment)")
    lb(doc,
       "HAI-prevention education records on the admission checklist confirming "
       "person-to-person education occurred.")

    h(doc, 2, "Special educational needs — PRE.5.i (Achievement)")
    lb(doc, "Identified-need records in the patient chart.")
    lb(doc, "Corresponding adapted-education or counselling records showing the need was met.")

    h(doc, 2, "Patient engagement — PRE.5.j (Excellence)")
    lb(doc,
       "Engagement activity records — support group sessions, advisory council meetings, "
       "safety champion appointments, and patient safety incident reports received from "
       "patients or families.")

    h(doc, 1, "14. References")
    ln(doc,
       "National Accreditation Board for Hospitals and Healthcare Providers. NABH "
       "Accreditation Standards for Hospitals, 6th Edition. PRE.5.")
    ln(doc, "Guidebook interpretation supplied for PRE.5.a through PRE.5.j.")

    h(doc, 1, "Disclaimer")
    p(doc,
      f"This policy reorganises the supplied PRE.5 objective-element wording and "
      f"Guidebook interpretation into plain-language guidance for {HN}. It does not "
      "replace the NABH Accreditation Standards for Hospitals, 6th Edition, or the "
      "Guidebook, which remain the authoritative source in case of any conflict.")
    p(doc,
      "This policy is intended for accreditation and internal governance purposes and "
      "is not a statement of, or substitute for, compliance with any specific statute; "
      f"{HN} remains separately responsible for meeting all applicable statutory and "
      "regulatory requirements.")
    p(doc,
      "This policy shall be reviewed at least once every year, or earlier if NABH "
      "standards change.")
    p(doc,
      f"This policy is the property of {HN} and is not to be reproduced or distributed "
      "outside the organisation without authorisation.")

    save_and_verify(doc, "HCO_PRE_5_v2_REWRITE_DRAFT.docx")


# ══════════════════════════════════════════════════════════════════════════════
# PRE.6 — Information on Expected Costs   (NO stop-work)
# COREs: a | Stars: none | Achievement: none | Excellence: none
# Prepared by: Patient Accounts In-Charge | Doc: HCO/PRE/POL/06
# ══════════════════════════════════════════════════════════════════════════════
def gen_pre6():
    doc = Document()

    h(doc, 0, "Policy on Information on Expected Costs")
    p(doc, HN)

    h(doc, 1, "Document control")
    doc_ctrl(doc, "HCO/PRE/POL/06", "Patient Accounts In-Charge")
    p(doc, "A blank marked ________ must be completed before issue.")

    h(doc, 1, "Statement of intent")
    p(doc,
      f"Patients and families at {HN} have a right to information on expected costs, "
      "so they can make informed decisions about their care without surprise charges.")

    h(doc, 1, "1. Purpose")
    p(doc,
      f"This policy sets out how {HN} makes its pricing policy and tariff available to "
      "patients and families, explains expected treatment costs, and informs them of "
      "financial implications when the care plan changes.")
    p(doc,
      "This policy does not cover informed consent, patient education about healthcare "
      "needs, or complaint redressal — those are covered in other hospital policies. "
      "The other PRE standards have their own policies too.")

    h(doc, 1, "2. Scope")
    p(doc,
      f"This policy applies to all patients and families across out-patient, emergency, "
      f"ICU and in-patient settings at {HN}.")

    h(doc, 1, "3. Policy standards")
    p(doc,
      f"{HN} makes patients and families aware of the pricing policy in each setting — "
      "out-patient, emergency, ICU and in-patient — and makes the relevant tariff list "
      "available to patients.")
    p(doc,
      f"{HN} explains expected costs to patients and families and informs them about "
      "financial implications when there is a change in the care plan.")

    h(doc, 1, "4. Non-negotiable rules")
    lb(doc,
       "Do not leave the pricing policy undisplayed at the registration or admission "
       "desk for any setting — out-patient, emergency, ICU or in-patient.")
    lb(doc,
       "Do not charge outside the current tariff list, and do not leave an additional "
       "charge un-enumerated or uncommunicated to the patient.")
    lb(doc,
       "Do not skip giving a patient an estimate of expected treatment costs, based on "
       "the treatment plan.")
    lb(doc,
       "Do not leave a patient or family uninformed of the financial implications when "
       "the care plan changes.")
    lb(doc,
       "Staff who see a rule under this policy broken report it the same shift to the "
       "Patient Accounts In-Charge or the Medical Superintendent.")

    h(doc, 1, "5. What we do")

    h(doc, 2, "5.1 Pricing policy display")
    p(doc,
      f"Key components of pricing — at minimum consultation charges, bed charges, nursing "
      f"charges and security deposit — are available to patients near the registration "
      f"and/or admission desk at {HN}, based on the billing policy that defines charges "
      "for various healthcare activities. Pricing information is available for "
      "out-patient, emergency, ICU and in-patient settings.")

    h(doc, 2, "5.2 Tariff list availability")
    p(doc,
      f"{HN} maintains an updated tariff list, available for review to patients when "
      "required. Charges follow the tariff list; any additional charge is enumerated in "
      "the tariff and communicated to patients. Tariff rates are uniform within a given "
      "setting and transparent.")

    h(doc, 2, "5.3 Cost estimates")
    p(doc,
      f"Patients at {HN} are given an estimate of expected treatment expenses, based on "
      "the treatment plan, preferably in written form. The estimate may be prepared by "
      "OPD, registration or admission staff in consultation with the treating doctor. "
      "Any limitations of the estimate — for example for emergency admissions — may "
      "also be discussed with the patient.")

    h(doc, 2, "5.4 Financial implications of care plan changes")
    p(doc,
      f"When there is a change in the care plan at {HN} — for example a shift between "
      "settings such as to or from the ICU, a shift from medical to surgical management, "
      "or the need for further expensive investigations — the patient and/or family are "
      "informed of the financial implications.")

    h(doc, 1, "6. Stop-work authority")
    p(doc,
      "Not applicable. This standard has no stop-work authority — gaps in cost "
      "information are addressed through the hospital's quality monitoring and CAPA "
      "process (Section 8), not through an immediate stop-work trigger.")

    h(doc, 1, "7. Governance and responsibility")
    gov_tbl(doc, [
        ("Medical Superintendent",
         "Accountable for ensuring patients and families receive cost information as "
         "required by this policy; receives escalations under Section 4."),
        ("Patient Accounts In-Charge",
         "Owns implementation of this policy; maintains the billing policy, pricing "
         "display and tariff list; keeps these current."),
        ("OPD / Registration / Admission Staff",
         "Prepare cost estimates in consultation with the treating doctor; direct "
         "patients to pricing and tariff information."),
        ("Treating Doctor",
         "Confirms the treatment plan behind each cost estimate; informs the patient "
         "and family of care-plan changes affecting cost."),
        ("Quality Coordinator",
         "Audits this policy; holds training records and staff acknowledgements."),
    ])

    h(doc, 1, "8. Quality monitoring")
    mon_tbl(doc, [
        ("Pricing policy display",
         "Displayed pricing-policy document (with date); billing policy it is based on; "
         "audit or observation record confirming display is in place across settings."),
        ("Tariff list availability",
         "Dated tariff list (current version); charge-versus-tariff reconciliation record "
         "for a sampled bill; confirmation no undisclosed charge was applied."),
        ("Cost estimates",
         "Cost-estimate records given to patients (written where possible), filed in the "
         "patient account or medical record; record of limitations discussed."),
        ("Care plan change communication",
         "Revised-estimate or cost-change notification records; confirmation no surprise "
         "charge appeared at discharge for a known plan change."),
    ])

    h(doc, 1, "9. Training and staff acknowledgement")
    p(doc,
      f"All staff covered by this policy are trained on how to direct patients to pricing "
      "and cost information and on the estimate and care-plan-change notification "
      "requirements, at induction and periodically.")
    p(doc,
      f"I have read the Policy on Information on Expected Costs of {HN} and understand "
      "my responsibilities under it.")
    sig_tbl(doc)

    h(doc, 1, "10. Distribution")
    p(doc,
      "This policy shall be available to the Patient Accounts department, OPD and "
      "Admission desks, treating doctors, and all staff through the hospital intranet "
      "and department policy folders.")

    h(doc, 1, "11. Abbreviations")
    abbrev_tbl(doc, [
        ("CAPA", "Corrective and Preventive Action"),
        ("ICU",  "Intensive Care Unit"),
        ("NABH", "National Accreditation Board for Hospitals and Healthcare Providers"),
        ("OPD",  "Out-Patient Department"),
    ])

    h(doc, 1, "12. Traceability table")
    p(doc,
      "This table is an index. It is not how the policy is organised. An asterisk in "
      "the Level column means the objective element is starred and requires documented "
      "evidence.")
    tr = tbl(doc, 5, 3)
    for ci, hdr in enumerate(("Objective Element", "Level", "Traceability to this policy")):
        tr.cell(0, ci).text = hdr
    trace_rows = [
        ("PRE.6.a", "CORE",
         "Sections 3 and 5.1 address display of pricing policy components across "
         "out-patient, emergency, ICU and in-patient settings."),
        ("PRE.6.b", "Commitment",
         "Sections 3 and 5.2 address availability of an updated tariff list and "
         "charging as per the tariff."),
        ("PRE.6.c", "Commitment",
         "Sections 3 and 5.3 address explanation of expected treatment costs, "
         "preferably in written form, based on the treatment plan."),
        ("PRE.6.d", "Commitment",
         "Sections 3 and 5.4 address informing the patient and/or family of the "
         "financial implications of a care plan change."),
    ]
    for ri, (oe, lvl, txt) in enumerate(trace_rows, 1):
        tr.cell(ri, 0).text = oe
        tr.cell(ri, 1).text = lvl
        tr.cell(ri, 2).text = txt

    h(doc, 1, "13. Required Records/Evidence Checklist")

    h(doc, 2, "Pricing policy display — PRE.6.a (CORE)")
    lb(doc,
       "Displayed pricing-policy document (with date) and the billing policy it is "
       "based on.")
    lb(doc,
       "Audit or observation record confirming the display is in place at registration "
       "and/or admission desk across all applicable settings.")

    h(doc, 2, "Tariff list — PRE.6.b (Commitment)")
    lb(doc, "Dated tariff list (current version).")
    lb(doc, "Charge-versus-tariff reconciliation record for a sampled bill.")
    lb(doc, "Confirmation that no undisclosed additional charge was applied.")

    h(doc, 2, "Cost estimates — PRE.6.c (Commitment)")
    lb(doc,
       "Cost-estimate records given to patients (written where possible), filed in the "
       "patient account or medical record.")
    lb(doc, "Records of any limitations discussed (for example for emergency admissions).")

    h(doc, 2, "Care plan change financial communication — PRE.6.d (Commitment)")
    lb(doc,
       "Records of revised-estimate or cost-change notification when the care plan "
       "shifted, noting the change and the date the patient or family was informed.")
    lb(doc,
       "Confirmation no surprise discharge charge occurred for a known plan change.")

    h(doc, 1, "14. References")
    ln(doc,
       "National Accreditation Board for Hospitals and Healthcare Providers. NABH "
       "Accreditation Standards for Hospitals, 6th Edition. PRE.6.")
    ln(doc, "Guidebook interpretation supplied for PRE.6.a through PRE.6.d.")

    h(doc, 1, "Disclaimer")
    p(doc,
      f"This policy reorganises the supplied PRE.6 objective-element wording and "
      f"Guidebook interpretation into plain-language guidance for {HN}. It does not "
      "replace the NABH Accreditation Standards for Hospitals, 6th Edition, or the "
      "Guidebook, which remain the authoritative source in case of any conflict.")
    p(doc,
      "This policy is intended for accreditation and internal governance purposes and "
      "is not a statement of, or substitute for, compliance with any specific statute; "
      f"{HN} remains separately responsible for meeting all applicable statutory and "
      "regulatory requirements.")
    p(doc,
      "This policy shall be reviewed at least once every year, or earlier if NABH "
      "standards change.")
    p(doc,
      f"This policy is the property of {HN} and is not to be reproduced or distributed "
      "outside the organisation without authorisation.")

    save_and_verify(doc, "HCO_PRE_6_v2_REWRITE_DRAFT.docx")


# ══════════════════════════════════════════════════════════════════════════════
# PRE.7 — Patient Feedback, Experience and Complaint Redressal  (NO stop-work)
# COREs: c* | Stars: c* | Achievement: b | Excellence: none
# Prepared by: Quality Coordinator | Doc: HCO/PRE/POL/07
# ══════════════════════════════════════════════════════════════════════════════
def gen_pre7():
    doc = Document()

    h(doc, 0, "Policy on Patient Feedback, Experience and Complaint Redressal")
    p(doc, HN)

    h(doc, 1, "Document control")
    doc_ctrl(doc, "HCO/PRE/POL/07", "Quality Coordinator")
    p(doc, "A blank marked ________ must be completed before issue.")

    h(doc, 1, "Statement of intent")
    p(doc,
      f"{HN} has a mechanism to capture patient feedback and experience, and to redress "
      "complaints, so patient voices shape how care is delivered.")

    h(doc, 1, "1. Purpose")
    p(doc,
      f"This policy sets out how {HN} captures patient feedback and patient experience, "
      "redresses complaints, makes patients aware of how to give feedback or complain, "
      "reviews and analyses what it receives, and acts on what it finds.")
    p(doc,
      "This policy does not cover informed consent, patient education about healthcare "
      "needs, or information on costs — those are covered in other hospital policies. "
      "The other PRE standards have their own policies too.")

    h(doc, 1, "2. Scope")
    p(doc,
      f"This policy applies to feedback, experience data and complaints from all patients "
      f"and families at {HN}, across out-patient and in-patient settings, including "
      "complaints against healthcare workers.")

    h(doc, 1, "3. Policy standards")
    p(doc,
      f"{HN} has a mechanism to capture feedback from patients, including patient "
      "satisfaction, and a mechanism to capture the patient experience, beyond "
      "satisfaction, including patient-reported experience measures.")
    p(doc,
      f"{HN} redresses patient complaints as per a defined, written mechanism; makes "
      "patients and families aware of the procedure for giving feedback and/or lodging "
      "complaints; reviews and/or analyses feedback and complaints within a defined time "
      "frame; and takes corrective and/or preventive action based on that analysis "
      "where appropriate.")

    h(doc, 1, "4. Non-negotiable rules")
    lb(doc,
       "Do not operate without a mechanism to capture feedback from patients, including "
       "patient satisfaction.")
    lb(doc,
       "Do not limit patient experience capture to satisfaction alone — it must go "
       "beyond, and must include patient-reported experience measures.")
    lb(doc,
       "Do not redress patient complaints outside a defined, written mechanism — and do "
       "not exclude verbal, telephonic, or healthcare-worker complaints from that "
       "mechanism.")
    lb(doc,
       "Do not leave patients and families unaware of how to give feedback or lodge "
       "a complaint.")
    lb(doc,
       "Do not let feedback or complaints sit unreviewed beyond the defined time frame, "
       "or leave the review process undocumented.")
    lb(doc,
       "Do not leave a pattern or systemic gap identified through analysis without "
       "corrective or preventive action.")
    lb(doc,
       "Staff who see a rule under this policy broken report it the same shift to the "
       "Quality Coordinator or the Medical Superintendent.")

    h(doc, 1, "5. What we do")

    h(doc, 2, "5.1 Capturing patient feedback")
    p(doc,
      f"A mechanism to capture feedback from patients — including patient satisfaction "
      f"— exists at {HN}, captured physically or electronically. Out-patient and "
      "in-patient data are kept separate as best practice.")

    h(doc, 2, "5.2 Capturing patient experience")
    p(doc,
      f"Beyond patient feedback, {HN} also captures patient experience — covering, at "
      "a minimum, communication with doctors and nurses, pain management, hospital "
      "environment (cleanliness and quietness), responsiveness of hospital staff, "
      "discharge information, communication about medications, and overall rating. "
      f"{HN} captures patient-reported experience measures (PREMs) as part of this.")

    h(doc, 2, "5.3 Complaint redressal mechanism")
    p(doc,
      f"{HN} has written guidance covering the mechanism for lodging complaints "
      "(including verbal or telephonic complaints), the method of compiling and "
      "analysing complaints and the time frame for doing so, the person(s) responsible, "
      "and how action taken is documented. Patient complaints include those against "
      f"healthcare workers. Whether to give credence to anonymous complaints is for "
      f"{HN} to decide.")

    h(doc, 2, "5.4 Awareness of the feedback and complaint procedure")
    p(doc,
      f"Patients and families at {HN} are made aware of the procedure for giving "
      "feedback and/or lodging complaints, through display or written information. "
      f"{HN} works to create an environment of trust in which patients are comfortable "
      "airing their views.")

    h(doc, 2, "5.5 Review and analysis")
    p(doc,
      f"Feedback and complaints at {HN} are reviewed and/or analysed within a defined "
      "time frame, and the entire process is documented. Where appropriate, the patient "
      "and/or family may be involved in the discussion and informed of the outcome.")

    h(doc, 2, "5.6 Corrective and preventive action")
    p(doc,
      "Where the analysis in Section 5.5 identifies an opportunity for improvement, "
      f"{HN} carries out corrective and/or preventive action.")

    h(doc, 1, "6. Stop-work authority")
    p(doc,
      "Not applicable. This standard has no stop-work authority — gaps in feedback "
      "capture, complaint redressal or analysis are addressed through the hospital's "
      "quality monitoring and CAPA process (Section 8), not through an immediate "
      "stop-work trigger.")

    h(doc, 1, "7. Governance and responsibility")
    gov_tbl(doc, [
        ("Medical Superintendent",
         "Accountable for ensuring the feedback, experience and complaint mechanisms "
         "are in place and effective; receives escalations under Section 4."),
        ("Quality Coordinator",
         "Owns implementation of this policy; maintains the feedback and complaint "
         "mechanisms, complaint log and analysis records; tracks CAPA to closure."),
        ("Department Heads / Unit In-Charges",
         "Support complaint investigation and implement corrective/preventive action "
         "in their area."),
        ("All Staff",
         "Direct patients to the feedback and complaint channels; report identified "
         "risks and rule breaches."),
    ])

    h(doc, 1, "8. Quality monitoring")
    mon_tbl(doc, [
        ("Patient feedback mechanism",
         "Feedback tool or form; response rate record; tabulated satisfaction data "
         "with out-patient and in-patient data kept separate."),
        ("Patient experience capture",
         "Experience capture tool showing the listed parameters; PREM/"
         "experience reports to the Medical Superintendent; action records."),
        ("Complaint redressal mechanism",
         "Written complaint-mechanism guidance; complaint log (including verbal, "
         "telephonic, anonymous and healthcare-worker complaints); documented action "
         "records for sampled complaints."),
        ("Awareness",
         "Display records (rights board, admission pamphlet, ward notice) matching "
         "the actual complaints desk/phone/form."),
        ("Review and analysis",
         "Documented review/analysis process record; timelines for complaints and "
         "feedback; overdue-item list reported to the Medical Superintendent."),
        ("CAPA",
         "CAPA records from feedback and complaint analysis, with owner, due date "
         "and closure."),
    ])

    h(doc, 1, "9. Training and staff acknowledgement")
    p(doc,
      f"All staff covered by this policy are trained on the feedback and complaint "
      "mechanisms described in Section 5 at induction and periodically.")
    p(doc,
      f"I have read the Policy on Patient Feedback, Experience and Complaint Redressal "
      f"of {HN} and understand my responsibilities under it.")
    sig_tbl(doc)

    h(doc, 1, "10. Distribution")
    p(doc,
      "This policy shall be available to the Quality department, all department "
      "heads/unit in-charges, and all staff through the hospital intranet and department "
      "policy folders.")

    h(doc, 1, "11. Abbreviations")
    abbrev_tbl(doc, [
        ("CAPA", "Corrective and Preventive Action"),
        ("NABH", "National Accreditation Board for Hospitals and Healthcare Providers"),
        ("OPD",  "Out-Patient Department"),
        ("PREM", "Patient-Reported Experience Measure"),
    ])

    h(doc, 1, "12. Traceability table")
    p(doc,
      "This table is an index. It is not how the policy is organised. An asterisk in "
      "the Level column means the objective element is starred and requires documented "
      "evidence.")
    tr = tbl(doc, 7, 3)
    for ci, hdr in enumerate(("Objective Element", "Level", "Traceability to this policy")):
        tr.cell(0, ci).text = hdr
    trace_rows = [
        ("PRE.7.a", "Commitment",
         "Sections 3 and 5.1 address the mechanism to capture patient feedback, "
         "including patient satisfaction, with out-patient and in-patient data separate."),
        ("PRE.7.b", "Achievement",
         "Sections 3 and 5.2 address the mechanism to capture patient experience beyond "
         "satisfaction, including PREMs and the listed experience parameters."),
        ("PRE.7.c", "CORE*",
         "Sections 3 and 5.3 address the written complaint redressal mechanism, "
         "covering verbal/telephonic/healthcare-worker complaints and documented action."),
        ("PRE.7.d", "Commitment",
         "Sections 3 and 5.4 address making patients and families aware of the feedback "
         "and complaint procedure through display and written information."),
        ("PRE.7.e", "Commitment",
         "Sections 3 and 5.5 address review and/or analysis of feedback and complaints "
         "within a defined time frame, with the process fully documented."),
        ("PRE.7.f", "Commitment",
         "Sections 3 and 5.6 address corrective and/or preventive action based on "
         "feedback and complaint analysis where appropriate."),
    ]
    for ri, (oe, lvl, txt) in enumerate(trace_rows, 1):
        tr.cell(ri, 0).text = oe
        tr.cell(ri, 1).text = lvl
        tr.cell(ri, 2).text = txt

    h(doc, 1, "13. Required Records/Evidence Checklist")

    h(doc, 2, "Patient feedback — PRE.7.a (Commitment)")
    lb(doc,
       "Feedback tool or form, response rate record against the target, and tabulated "
       "satisfaction data with out-patient and in-patient data kept separate.")

    h(doc, 2, "Patient experience — PRE.7.b (Achievement)")
    lb(doc,
       "Experience capture tool showing the listed parameters (communication, pain, "
       "environment, responsiveness, discharge information, medications, overall rating).")
    lb(doc, "PREM/experience reports to the Medical Superintendent.")
    lb(doc, "Action records from those reports.")

    h(doc, 2, "Complaint redressal mechanism — PRE.7.c (CORE*)")
    lb(doc,
       "Written complaint-mechanism guidance document covering lodging, compiling, "
       "analysing, time frame, responsible person, and documentation of action taken.")
    lb(doc,
       "Complaint log including verbal, telephonic, anonymous and healthcare-worker "
       "complaints.")
    lb(doc, "Documented action records for sampled complaints.")

    h(doc, 2, "Awareness — PRE.7.d (Commitment)")
    lb(doc,
       "Display records (rights board, admission pamphlet, ward notice) naming the "
       "feedback and complaint procedure.")
    lb(doc, "Verification that the displayed path matches the actual complaints desk/phone/form.")

    h(doc, 2, "Review and analysis — PRE.7.e (Commitment)")
    lb(doc,
       "Documented review/analysis process record and the defined review time frame.")
    lb(doc, "Overdue-complaint list reported to the Medical Superintendent.")

    h(doc, 2, "CAPA — PRE.7.f (Commitment)")
    lb(doc,
       "CAPA records from feedback and complaint analysis, with owner, due date "
       "and closure.")
    lb(doc, "Confirmation that analysis led to action, not analysis alone.")

    h(doc, 1, "14. References")
    ln(doc,
       "National Accreditation Board for Hospitals and Healthcare Providers. NABH "
       "Accreditation Standards for Hospitals, 6th Edition. PRE.7.")
    ln(doc, "Guidebook interpretation supplied for PRE.7.a through PRE.7.f.")

    h(doc, 1, "Disclaimer")
    p(doc,
      f"This policy reorganises the supplied PRE.7 objective-element wording and "
      f"Guidebook interpretation into plain-language guidance for {HN}. It does not "
      "replace the NABH Accreditation Standards for Hospitals, 6th Edition, or the "
      "Guidebook, which remain the authoritative source in case of any conflict.")
    p(doc,
      "This policy is intended for accreditation and internal governance purposes and "
      "is not a statement of, or substitute for, compliance with any specific statute; "
      f"{HN} remains separately responsible for meeting all applicable statutory and "
      "regulatory requirements.")
    p(doc,
      "This policy shall be reviewed at least once every year, or earlier if NABH "
      "standards change.")
    p(doc,
      f"This policy is the property of {HN} and is not to be reproduced or distributed "
      "outside the organisation without authorisation.")

    save_and_verify(doc, "HCO_PRE_7_v2_REWRITE_DRAFT.docx")


# ══════════════════════════════════════════════════════════════════════════════
# PRE.8 — Effective Communication with Patients and Families  (NO stop-work)
# COREs: none | Stars: a*, b*, c* | Achievement: e | Commitment: a*, b*, c*, d
# PRE.8.d is a conduct prohibition, not a procedure-start stop-work gate
# Prepared by: Quality Coordinator | Doc: HCO/PRE/POL/08
# ══════════════════════════════════════════════════════════════════════════════
def gen_pre8():
    doc = Document()

    h(doc, 0, "Policy on Effective Communication with Patients and Families")
    p(doc, HN)

    h(doc, 1, "Document control")
    doc_ctrl(doc, "HCO/PRE/POL/08", "Quality Coordinator")
    p(doc, "A blank marked ________ must be completed before issue.")

    h(doc, 1, "Statement of intent")
    p(doc,
      f"{HN} has a system for effective communication with patients and/or families, "
      "so that communication serves its purpose, meets patients where they are, and "
      "never crosses into disrespect.")

    h(doc, 1, "1. Purpose")
    p(doc,
      f"This policy sets out how {HN} communicates effectively with patients and "
      "families, identifies special situations needing enhanced communication, ensures "
      "that enhanced communication is done well, prevents unacceptable communication, "
      "and monitors how well all of this works.")
    p(doc,
      "This policy does not cover informed consent, patient education about healthcare "
      "needs, or complaint redressal — those are covered in other hospital policies. "
      "The other PRE standards have their own policies too.")

    h(doc, 1, "2. Scope")
    p(doc,
      f"This policy applies to all communication between staff and patients and/or "
      f"families at {HN}.")

    h(doc, 1, "3. Policy standards")
    p(doc,
      f"{HN} communicates with patients and/or families effectively; identifies special "
      "situations where enhanced communication is required; carries out enhanced "
      "communication effectively for each identified special situation; and ensures "
      "there is no unacceptable communication.")
    p(doc,
      f"{HN} monitors and reviews the implementation of effective communication.")

    h(doc, 1, "4. Non-negotiable rules")
    lb(doc,
       "Do not communicate with a patient or family in a way that fails to serve its "
       "purpose — identify and work to overcome communication barriers.")
    lb(doc,
       "Do not leave special situations requiring enhanced communication unidentified.")
    lb(doc,
       "Do not leave an identified special situation without a detailed, trained approach "
       "to the enhanced communication it requires.")
    lb(doc,
       "Do not allow unacceptable communication — abusing patients, hurting religious "
       "or cultural sentiments, or communicating with disrespect.")
    lb(doc,
       "Do not leave the effectiveness of communication unmonitored and unreviewed.")
    lb(doc,
       "Staff who see a rule under this policy broken report it the same shift to the "
       "Quality Coordinator or the Medical Superintendent.")

    h(doc, 1, "5. What we do")

    h(doc, 2, "5.1 Effective communication")
    p(doc,
      f"Communication with patients and/or families at {HN} is done effectively — "
      "serving its purpose. {HN} identifies potential communication barriers and has "
      "plans to overcome them — for example, a language barrier may be overcome through "
      "an interpreter. Any recognised model of effective communication may be adopted.")

    h(doc, 2, "5.2 Identifying special situations")
    p(doc,
      f"{HN} identifies special situations where enhanced communication with patients "
      "and/or families is required — for example breaking bad news, handling adverse "
      "events, handling an aggressive patient or family, talking to the family of a "
      "patient who has died, or counselling for a complicated intervention.")

    h(doc, 2, "5.3 Delivering enhanced communication")
    p(doc,
      "For each special situation identified under Section 5.2, "
      f"{HN} details the nature of the enhanced communication required — for example, "
      "one recognised model for delivering bad news is SPIKES (Setting, Perception, "
      "Invitation or information, Knowledge, Empathy, Summarize or strategize). Any "
      "named or documented model may be adopted for each situation.")

    h(doc, 2, "5.4 No unacceptable communication")
    p(doc,
      f"{HN} does not allow unacceptable communication — for example abusing patients, "
      "hurting religious or cultural sentiments, or communicating with disrespect. "
      "Anyone who witnesses such conduct stops the interaction, protects the patient, "
      "and reports it under Section 4.")

    h(doc, 2, "5.5 Monitoring communication")
    p(doc,
      f"{HN} has a system to monitor and review how effective communication is "
      "implemented — for example through feedback or complaint analysis from patients "
      "and other stakeholders.")

    h(doc, 1, "6. Stop-work authority")
    p(doc,
      "Not applicable in the procedure-start sense. PRE.8.d prohibits unacceptable "
      "communication as a conduct obligation — anyone who witnesses such conduct stops "
      "the interaction, protects the patient, and reports it the same shift under "
      "Section 4. The Medical Superintendent acts on a confirmed breach as a conduct "
      "matter. This is not a procedure-start stop-work gate.")

    h(doc, 1, "7. Governance and responsibility")
    gov_tbl(doc, [
        ("Medical Superintendent",
         "Accountable for ensuring effective communication across the organisation; "
         "acts on confirmed unacceptable-communication breaches; receives escalations "
         "under Section 4."),
        ("Quality Coordinator",
         "Owns implementation of this policy; maintains the list of special situations; "
         "monitors communication effectiveness and reports findings."),
        ("Department Heads / Unit In-Charges",
         "Train staff on effective and enhanced communication for their area; escalate "
         "unacceptable-communication incidents."),
        ("All Staff",
         "Communicate effectively; recognise and report unacceptable communication "
         "the same shift."),
    ])

    h(doc, 1, "8. Quality monitoring")
    mon_tbl(doc, [
        ("Effective communication",
         "Training records at induction and periodically; barrier-identification and "
         "interpreter-arrangement records; observed-practice records for sampled cases."),
        ("Special situations identified",
         "Written list of special situations; training pack covering that list; "
         "staff-awareness records for emergency, ICU, OT and ward staff."),
        ("Enhanced communication delivered",
         "Detailed description of enhanced communication for each listed situation; "
         "records confirming enhanced communication was carried out for sampled cases; "
         "training records showing staff know the models in use."),
        ("No unacceptable communication",
         "Conduct-incident records for confirmed breaches; action records by the "
         "Medical Superintendent; training records naming examples."),
        ("Monitoring and review",
         "Observation or record-sample review; findings reported to the "
         "Medical Superintendent; action records when communication failures found."),
    ])

    h(doc, 1, "9. Training and staff acknowledgement")
    p(doc,
      f"All staff covered by this policy are trained on effective and enhanced "
      "communication principles, including recognising and reporting unacceptable "
      "communication, at induction and periodically.")
    p(doc,
      f"I have read the Policy on Effective Communication with Patients and Families "
      f"of {HN} and understand my responsibilities under it.")
    sig_tbl(doc)

    h(doc, 1, "10. Distribution")
    p(doc,
      "This policy shall be available to the Quality department, all department "
      "heads/unit in-charges, and all staff through the hospital intranet and department "
      "policy folders.")

    h(doc, 1, "11. Abbreviations")
    abbrev_tbl(doc, [
        ("CAPA",   "Corrective and Preventive Action"),
        ("ICU",    "Intensive Care Unit"),
        ("NABH",   "National Accreditation Board for Hospitals and Healthcare Providers"),
        ("OT",     "Operation Theatre"),
        ("SPIKES", "Setting, Perception, Invitation or information, Knowledge, Empathy, "
                   "Summarize or strategize (breaking-bad-news model)"),
    ])

    h(doc, 1, "12. Traceability table")
    p(doc,
      "This table is an index. It is not how the policy is organised. An asterisk in "
      "the Level column means the objective element is starred and requires documented "
      "evidence.")
    tr = tbl(doc, 6, 3)
    for ci, hdr in enumerate(("Objective Element", "Level", "Traceability to this policy")):
        tr.cell(0, ci).text = hdr
    trace_rows = [
        ("PRE.8.a", "Commitment*",
         "Sections 3 and 5.1 address effective communication serving its purpose, "
         "identifying and overcoming communication barriers."),
        ("PRE.8.b", "Commitment*",
         "Sections 3 and 5.2 address identification of special situations where "
         "enhanced communication is required."),
        ("PRE.8.c", "Commitment*",
         "Sections 3 and 5.3 address detailing and delivering enhanced communication "
         "for each identified special situation, including the SPIKES model for bad news."),
        ("PRE.8.d", "Commitment",
         "Sections 3 and 5.4 address the prohibition on unacceptable communication "
         "and the conduct-reporting obligation when it occurs."),
        ("PRE.8.e", "Achievement",
         "Sections 3 and 5.5 address the system to monitor and review implementation "
         "of effective communication, with findings reported and acted on."),
    ]
    for ri, (oe, lvl, txt) in enumerate(trace_rows, 1):
        tr.cell(ri, 0).text = oe
        tr.cell(ri, 1).text = lvl
        tr.cell(ri, 2).text = txt

    h(doc, 1, "13. Required Records/Evidence Checklist")

    h(doc, 2, "Effective communication — PRE.8.a (Commitment*)")
    lb(doc,
       "Training records on effective communication principles (induction and periodic).")
    lb(doc,
       "Barrier-identification and interpreter-arrangement records.")
    lb(doc,
       "Observed-practice records confirming effective communication in sampled cases.")

    h(doc, 2, "Special situations identified — PRE.8.b (Commitment*)")
    lb(doc,
       "Written list of special situations requiring enhanced communication.")
    lb(doc,
       "Training pack covering that list, with staff-awareness records for emergency, "
       "ICU, OT and ward staff.")

    h(doc, 2, "Enhanced communication delivered — PRE.8.c (Commitment*)")
    lb(doc,
       "Detailed description of enhanced communication for each listed special situation "
       "(including the adopted model for breaking bad news, for example SPIKES).")
    lb(doc,
       "Records confirming enhanced communication was carried out for sampled cases.")
    lb(doc,
       "Training records showing staff know the models in use.")

    h(doc, 2, "No unacceptable communication — PRE.8.d (Commitment)")
    lb(doc,
       "Conduct-incident records for confirmed unacceptable-communication breaches.")
    lb(doc, "Action records by the Medical Superintendent for confirmed breaches.")
    lb(doc,
       "Training records naming examples so staff can recognise the line.")

    h(doc, 2, "Monitoring and review — PRE.8.e (Achievement)")
    lb(doc,
       "Communication-monitoring review records (observations, record samples, "
       "or feedback/complaint analysis focused on communication).")
    lb(doc, "Findings reported to the Medical Superintendent.")
    lb(doc, "Action records when communication failures were identified.")

    h(doc, 1, "14. References")
    ln(doc,
       "National Accreditation Board for Hospitals and Healthcare Providers. NABH "
       "Accreditation Standards for Hospitals, 6th Edition. PRE.8.")
    ln(doc, "Guidebook interpretation supplied for PRE.8.a through PRE.8.e.")

    h(doc, 1, "Disclaimer")
    p(doc,
      f"This policy reorganises the supplied PRE.8 objective-element wording and "
      f"Guidebook interpretation into plain-language guidance for {HN}. It does not "
      "replace the NABH Accreditation Standards for Hospitals, 6th Edition, or the "
      "Guidebook, which remain the authoritative source in case of any conflict.")
    p(doc,
      "This policy is intended for accreditation and internal governance purposes and "
      "is not a statement of, or substitute for, compliance with any specific statute; "
      f"{HN} remains separately responsible for meeting all applicable statutory and "
      "regulatory requirements.")
    p(doc,
      "This policy shall be reviewed at least once every year, or earlier if NABH "
      "standards change.")
    p(doc,
      f"This policy is the property of {HN} and is not to be reproduced or distributed "
      "outside the organisation without authorisation.")

    save_and_verify(doc, "HCO_PRE_8_v2_REWRITE_DRAFT.docx")


if __name__ == "__main__":
    gen_pre1()
    gen_pre2()
    gen_pre3()
    gen_pre4()
    gen_pre5()
    gen_pre6()
    gen_pre7()
    gen_pre8()
