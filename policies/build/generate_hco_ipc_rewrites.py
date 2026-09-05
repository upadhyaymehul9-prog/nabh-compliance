# -*- coding: utf-8 -*-
"""
generate_hco_ipc_rewrites.py
Generates HCO IPC chapter v2 rewrite-reference DOCX files.

Pipeline : python-docx, identical to generate_hco_mom_rewrites.py.
Output   : policies/build/rewrite_reference/HCO_IPC_N_v2_REWRITE_DRAFT.docx
Source   : policies/build/ipc_raw_dump_1-6.txt
"""
import os
from docx import Document

HN  = "«Hospital Name»"
OUT = os.path.join(os.path.dirname(os.path.abspath(__file__)), "rewrite_reference")
os.makedirs(OUT, exist_ok=True)


# ── Helpers ───────────────────────────────────────────────────────────────────

def h(doc, lv, txt):
    return doc.add_paragraph(txt, style={0: "Title", 1: "Heading 1", 2: "Heading 2"}[lv])

def p(doc, txt=""):
    return doc.add_paragraph(txt, style="Normal")

def ln(doc, txt):
    return doc.add_paragraph(txt, style="List Number")

def lb(doc, txt):
    return doc.add_paragraph(txt, style="List Bullet")

def tbl(doc, rows, cols):
    t = doc.add_table(rows=rows, cols=cols)
    try:
        t.style = "Table Grid"
    except KeyError:
        pass
    return t

def doc_ctrl(doc, no, prep, appr="Medical Superintendent"):
    dc = tbl(doc, 6, 4)
    for ri, (a, b, c, d) in enumerate([
        ("Document No.", no,          "Version",                "2.0"),
        ("Issue No.",    "01",         "Review due",             "One year from implementation"),
        ("Date created", "________",   "Date of implementation", "________"),
    ]):
        dc.cell(ri, 0).text = a; dc.cell(ri, 1).text = b
        dc.cell(ri, 2).text = c; dc.cell(ri, 3).text = d
    for ri, (lbl, txt) in enumerate([
        ("Prepared by", f"{prep}  Name: ________  Signature: ________"),
        ("Reviewed by", "Quality Coordinator  Name: ________  Signature: ________"),
        ("Approved by", f"{appr}  Name: ________  Signature: ________"),
    ], start=3):
        dc.cell(ri, 0).text = lbl
        c1 = dc.cell(ri, 1); c1.text = txt; c1.merge(dc.cell(ri, 3))

def abbrev_tbl(doc, rows):
    t = tbl(doc, len(rows) + 1, 2)
    t.cell(0, 0).text = "Abbreviation"; t.cell(0, 1).text = "Meaning"
    for ri, (a, m) in enumerate(rows, 1):
        t.cell(ri, 0).text = a; t.cell(ri, 1).text = m

def gov_tbl(doc, rows):
    t = tbl(doc, len(rows) + 1, 2)
    t.cell(0, 0).text = "Role"; t.cell(0, 1).text = "Responsibility"
    for ri, (role, resp) in enumerate(rows, 1):
        t.cell(ri, 0).text = role; t.cell(ri, 1).text = resp

def mon_tbl(doc, rows):
    t = tbl(doc, len(rows) + 1, 2)
    t.cell(0, 0).text = "Monitoring area"; t.cell(0, 1).text = "What is monitored"
    for ri, (area, what) in enumerate(rows, 1):
        t.cell(ri, 0).text = area; t.cell(ri, 1).text = what

def sig_tbl(doc):
    sig = tbl(doc, 4, 4)
    for ci, hdr in enumerate(("Staff name", "Designation", "Signature", "Date")):
        sig.cell(0, ci).text = hdr
    for ri in range(1, 4):
        for ci in range(4):
            sig.cell(ri, ci).text = "________"

def save_and_verify(doc, fname):
    import sys
    out = sys.stdout

    def pr(s):
        try:
            out.write(s + "\n")
        except UnicodeEncodeError:
            out.write(s.encode("ascii", "replace").decode() + "\n")

    pr(f"\n=== {fname} ===")
    for i, para in enumerate(doc.paragraphs[:80]):
        sn = para.style.name if para.style else "(None)"
        pr(f"{i:3d}  {sn!r:30s}  {para.text[:60]!r}")
    counts = {}
    for para in doc.paragraphs:
        sn = para.style.name if para.style else "(None)"
        counts[sn] = counts.get(sn, 0) + 1
    pr("  Style inventory:")
    for sn, n in sorted(counts.items()):
        pr(f"    {sn}: {n}")
    pr(f"  Total paras: {len(doc.paragraphs)}")
    path = os.path.join(OUT, fname)
    doc.save(path)
    pr(f"  Saved: {path}")


# ══════════════════════════════════════════════════════════════════════════════
# IPC.1 — Infection Prevention and Control Programme   (NO stop-work)
# COREs: a | Stars: a*, b*, e*, f*, g*, h* | Achievement: d | Excellence: none
# Prepared by: Infection Prevention and Control Officer | Doc: HCO/IPC/POL/01
# ══════════════════════════════════════════════════════════════════════════════
def gen_ipc1():
    doc = Document()

    h(doc, 0, "Policy on Infection Prevention and Control Programme")
    p(doc, HN)

    h(doc, 1, "Document control")
    doc_ctrl(doc, "HCO/IPC/POL/01", "Infection Prevention and Control Officer")
    p(doc, "A blank marked ________ must be completed before issue.")

    h(doc, 1, "Statement of intent")
    p(doc,
      f"The IPC programme at {HN} is comprehensively documented as an IPC Manual, "
      "directed at identified high-risk activities and areas, maintained by a designated "
      "officer and nurse(s) under a multi-disciplinary committee, assessed with a "
      "validated tool at least annually, and extended to the community.")

    h(doc, 1, "1. Purpose")
    p(doc,
      f"This policy explains how {HN} documents its infection prevention and control "
      "programme as a written IPC Manual, identifies high-risk activities and areas, "
      "reviews and updates the programme at least annually using a validated assessment "
      "tool, guides it through a multi-disciplinary IPC committee that meets at least "
      "monthly, and runs it day to day through a designated IPC Officer and IPC Nurse(s) "
      "supported by link nurses. It also explains how the organisation conducts IEC for "
      "the community and participates in managing community outbreaks.")
    p(doc,
      "The clinical IPC processes, support-service IPC, HAI-prevention bundles, and "
      "surveillance are covered in separate hospital policies.")

    h(doc, 1, "2. Scope")
    p(doc,
      f"This policy applies to the IPC Officer, IPC Nurse(s), link nurses, the "
      f"multi-disciplinary IPC committee, and all staff involved in infection prevention "
      f"and control activities at {HN}.")

    h(doc, 1, "3. Policy standards")
    p(doc,
      f"{HN} maintains a documented IPC Manual covering the programme's structure, "
      "overall aims and objectives, all processes, activities and surveillance procedures, "
      "based on scientific knowledge, national/international guidelines, and statutory "
      "requirements. High-risk activities and areas are identified and have written "
      "guidance. The programme is reviewed and updated at least once a year using a "
      "validated assessment tool. A multi-disciplinary IPC committee is constituted with "
      "documented composition, a meeting frequency of at least once a month, a defined "
      "quorum, and documented minutes; it sets risk-reduction goals and measurable "
      "objectives at least annually and reviews them monthly. The IPC team comprises at "
      "minimum a designated IPC Officer (a doctor) and IPC Nurse(s) (registered nurses "
      "with structured IPC training, at least one per 250 beds), supported by link nurses. "
      "An IEC programme reaches the community, and the organisation participates in "
      "managing community outbreaks.")
    p(doc, "Staff follow the written guidance below and keep the records it requires.")

    h(doc, 1, "4. Non-negotiable rules")
    lb(doc,
       "Do not manage infection prevention and control without a documented IPC Manual "
       "covering the programme's structure, aims and objectives, all processes, activities "
       "and surveillance procedures, based on scientific knowledge, national/international "
       "guidelines, and statutory requirements.")
    lb(doc,
       "Do not run the IPC programme without identifying high-risk activities and areas "
       "and having written guidance to prevent and manage infections for each.")
    lb(doc,
       "Do not let the IPC programme go more than a year without review and update; if "
       "the annual review finds no opportunities for improvement, document that finding "
       "in the IPC committee minutes.")
    lb(doc,
       "Do not constitute the IPC committee without documenting its composition, meeting "
       "frequency, quorum and minutes; do not let it meet less often than once a month.")
    lb(doc,
       "Do not operate without risk-reduction goals and measurable objectives established "
       "at least annually and reviewed monthly by the IPC committee.")
    lb(doc,
       "Do not designate a non-doctor as the IPC Officer.")
    lb(doc,
       "Do not designate as IPC Nurse anyone who is not a registered nurse with "
       "additional structured IPC training, and do not staff fewer than one dedicated "
       "IPC Nurse per 250 beds.")

    h(doc, 1, "5. What we do")

    h(doc, 2, "5.1 Document the IPC programme as a manual")
    p(doc,
      "Written infection prevention and control guidance is consolidated into an IPC "
      "Manual that covers the structure of the programme, its overall aims and objectives, "
      "and all processes, activities and surveillance procedures. The Manual is based on "
      f"organisational priorities, current scientific knowledge, and national/international "
      "guidelines and professional body guidance. Statutory requirements are incorporated "
      f"where applicable. {HN} implements strategies, education, and evidence-based "
      "activities based on the Manual to reduce infection risks.")

    h(doc, 2, "5.2 Identify high-risk activities and areas, and write guidance for them")
    p(doc,
      "High-risk activities and high-risk areas across the hospital are identified based "
      "on scientific literature and the potential risk of transmission of infections to "
      "patients and healthcare providers. Examples of high-risk activities include "
      "aerosol-generating procedures, handling blood and body fluids, spills, specimens "
      "and sharps, and exposure to contaminated devices, equipment and BMW. Examples of "
      "high-risk areas include the OT, ICU, CSSD, cath lab, labour room, laboratory and "
      "blood centre. The organisation identifies its own list based on its scope of "
      "services. Written guidance to prevent and manage infections is developed for each "
      "identified high-risk activity and area.")

    h(doc, 2, "5.3 Review and update the programme at least annually")
    p(doc,
      "The IPC programme is reviewed and updated at least once a year, based on newer "
      "scientific literature, infection trends, outbreak prevention developments, and the "
      "outcomes of IPC audit processes. If the annual review does not identify any "
      "opportunities for improvement, that finding is documented in the IPC committee "
      "meeting minutes.")

    h(doc, 2, "5.4 Assess the programme using a validated tool")
    p(doc,
      f"{HN} uses a validated tool to perform a structured assessment of the IPC "
      "programme at least annually. Examples of validated tools include the WHO Infection "
      "Prevention and Control Assessment Framework at the Facility Level (WHO IPCAF) and "
      "the CDC Infection Prevention and Control Assessment Tool for Acute Care Hospitals. "
      "The assessment output is documented, gap analysis is completed, and follow-up "
      "actions are tracked to closure.")

    h(doc, 2, "5.5 Constitute a multi-disciplinary IPC committee")
    p(doc,
      "A multi-disciplinary IPC committee co-ordinates all IPC activities. Its "
      "composition is documented and it includes, where available, representatives from "
      "major clinical departments, administration, nursing, the IPC Officer, IPC Nurse(s), "
      "CSSD, OT, ICU, and support services. The committee's terms of reference — "
      "composition, meeting frequency, and minimum quorum — are defined. The committee "
      "meets at least once a month and its minutes are documented. Risk-reduction goals "
      "and measurable objectives are established at least annually and reviewed monthly "
      "by the committee.")

    h(doc, 2, "5.6 Maintain an IPC team")
    p(doc,
      "An IPC team is responsible for the day-to-day functioning of the IPC programme. "
      "The team supports the surveillance process, detects outbreaks, participates in "
      "audit activity, and is represented in the IPC committee. The team is staffed "
      "according to the organisation's size, risk level, and programme complexity. At "
      "minimum, the team comprises the IPC Officer and IPC Nurse(s). IPC Nurses are "
      "supported by link nurses — trained clinical staff who actively participate in IPC "
      "activities in their work areas.")

    h(doc, 2, "5.7 Designate an IPC officer")
    p(doc,
      "The IPC Officer is a doctor who is knowledgeable in infection prevention and "
      "control practices. It is preferable for the IPCO to be a clinical microbiologist "
      "or infectious disease specialist; where this is not available, a surgeon or "
      "physician may be designated. The IPCO's responsibilities are defined in the IPC "
      "Manual. Those requirements related to HR designation are covered in the hospital's "
      "other policies.")

    h(doc, 2, "5.8 Designate IPC nurse(s)")
    p(doc,
      "IPC Nurses are designated on the basis of qualification as a registered nurse "
      "and completion of additional structured IPC training. Their responsibilities are "
      "defined in the IPC Manual and may include: surveillance of HAIs and healthcare-"
      "associated organisms; compliance monitoring of hand hygiene, transmission-based "
      "precautions, isolation, infection-specific bundles, disinfection and sterilisation; "
      "education; and documentation. At minimum, one dedicated IPC Nurse is maintained "
      "per 250 beds. Those requirements related to HR designation are covered in the "
      "hospital's other policies.")

    h(doc, 2, "5.9 Run community IEC on infection prevention")
    p(doc,
      f"{HN} implements an information, education and communication programme on "
      "infection prevention and control for the community. The organisation may work "
      "with stakeholders to develop IEC messages covering topics such as hand hygiene, "
      "appropriate use of antimicrobials, and use of personal protective equipment.")

    h(doc, 2, "5.10 Participate in managing community outbreaks")
    p(doc,
      f"{HN} coordinates with external agencies, including statutory agencies, to respond "
      "effectively to community outbreaks. Internal and external communication channels "
      "are defined. Roles and responsibilities for staff during community outbreak "
      "response are documented. Staff are trained on their roles.")

    h(doc, 1, "6. Governance and responsibility")
    gov_tbl(doc, [
        ("Medical Superintendent",
         "Accountable for ensuring the IPC programme is resourced and implemented "
         "as required by this policy."),
        ("Infection Prevention and Control Officer (IPCO)",
         "Owns day-to-day implementation of this policy; leads the IPC team; brings "
         "surveillance findings, incidents and audit outcomes to the IPC committee."),
        ("Infection Prevention and Control Nurse(s) (IPCN)",
         "Conducts day-to-day IPC activities including surveillance, compliance "
         "monitoring and education; supported by link nurses in each clinical area."),
        ("IPC Committee",
         "Guides the IPC programme; sets and reviews risk-reduction goals and objectives; "
         "meets at least monthly and documents minutes."),
        ("Quality Coordinator",
         "Audits this policy; holds training records and staff acknowledgements."),
    ])

    h(doc, 1, "7. Quality monitoring")
    mon_tbl(doc, [
        ("IPC Manual",
         "Complete, current, and based on scientific knowledge and national/international "
         "guidelines; updated at least annually."),
        ("High-risk register",
         "High-risk activities and areas identified; written guidance exists for each."),
        ("Annual programme assessment",
         "Validated assessment tool completed at least annually; gap analysis documented; "
         "follow-up actions tracked."),
        ("IPC committee",
         "Composition, quorum and terms of reference documented; meets at least monthly; "
         "minutes documented; goals and objectives set annually and reviewed monthly."),
        ("IPC team staffing",
         "IPCO is a doctor; IPC Nurse ratio at least one per 250 beds; link nurses "
         "in place in clinical areas; roles defined in manual."),
        ("Community IEC",
         "IEC programme implemented; materials or session records available."),
        ("Community outbreak participation",
         "External agency contacts defined; staff roles documented; staff trained."),
    ])

    h(doc, 1, "8. Training and staff acknowledgement")
    p(doc,
      "The IPC Officer, IPC Nurse(s), link nurses, and IPC committee members shall be "
      "familiar with the IPC Manual, the assessment tool process, committee roles, "
      "community IEC responsibilities, and the community outbreak response process.")
    p(doc,
      f"I have read the Policy on Infection Prevention and Control Programme of {HN}. "
      "I will follow the processes described.")
    sig_tbl(doc)

    h(doc, 1, "9. Distribution")
    p(doc,
      "This policy shall be available to the IPC Officer, IPC Nurse(s), link nurses, "
      "the multi-disciplinary IPC committee, department heads, and the Quality Coordinator.")

    h(doc, 1, "10. Abbreviations")
    abbrev_tbl(doc, [
        ("AMR",   "Antimicrobial resistance"),
        ("BMW",   "Biomedical waste"),
        ("CDC",   "Centers for Disease Control and Prevention"),
        ("CSSD",  "Central Sterile Services Department"),
        ("ICU",   "Intensive Care Unit"),
        ("IEC",   "Information, Education and Communication"),
        ("IPC",   "Infection Prevention and Control"),
        ("IPCC",  "Infection Prevention and Control Committee"),
        ("IPCO",  "Infection Prevention and Control Officer"),
        ("IPCN",  "Infection Prevention and Control Nurse"),
        ("IPCAF", "IPC Assessment Framework (WHO tool)"),
        ("NABH",  "National Accreditation Board for Hospitals and Healthcare Providers"),
        ("OT",    "Operation Theatre"),
        ("WHO",   "World Health Organization"),
    ])

    h(doc, 1, "11. Traceability table")
    p(doc,
      "This table is an index. It is not how the policy is organised. An asterisk in "
      "the Level column means documentation of the process is required.")
    tr = tbl(doc, 11, 3)
    for ci, hdr in enumerate(("Objective Element", "Level", "Traceability to this policy")):
        tr.cell(0, ci).text = hdr
    trace_rows = [
        ("IPC.1.a", "CORE*",
         "Sections 3 and 5.1 address the documented IPC Manual covering programme "
         "structure, aims and objectives, all processes, activities and surveillance "
         "procedures, based on scientific knowledge and guidelines."),
        ("IPC.1.b", "Commitment*",
         "Sections 3 and 5.2 address the identification of high-risk activities and "
         "areas and the requirement for written guidance for each."),
        ("IPC.1.c", "Commitment",
         "Sections 3 and 5.3 address the annual programme review and update, including "
         "the documented nil-finding requirement."),
        ("IPC.1.d", "Achievement",
         "Section 5.4 addresses use of a validated assessment tool with documented "
         "gap analysis and follow-up action tracking."),
        ("IPC.1.e", "Commitment*",
         "Sections 3 and 5.5 address the multi-disciplinary IPC committee: composition, "
         "meeting frequency of at least monthly, quorum, documented minutes, and monthly "
         "review of annually-set objectives."),
        ("IPC.1.f", "Commitment*",
         "Sections 3 and 5.6 address the IPC team's day-to-day functions, minimum "
         "staffing (IPCO and IPCN), and link nurse support structure."),
        ("IPC.1.g", "Commitment*",
         "Sections 3 and 5.7 address the IPCO designation (must be a doctor) and "
         "definition of responsibilities in the IPC Manual."),
        ("IPC.1.h", "Commitment*",
         "Sections 3 and 5.8 address IPCN designation criteria (registered nurse + "
         "structured training), responsibilities defined in the manual, and the minimum "
         "staffing ratio of one per 250 beds."),
        ("IPC.1.i", "Commitment",
         "Section 5.9 addresses the community IEC programme on infection prevention "
         "and control."),
        ("IPC.1.j", "Commitment",
         "Section 5.10 addresses community outbreak participation: external coordination, "
         "communication channels, staff roles and staff training."),
    ]
    for ri, (oe, lvl, txt) in enumerate(trace_rows, 1):
        tr.cell(ri, 0).text = oe
        tr.cell(ri, 1).text = lvl
        tr.cell(ri, 2).text = txt

    h(doc, 1, "12. Required Records/Evidence Checklist")

    h(doc, 2, "IPC Manual — IPC.1.a (CORE*)")
    lb(doc,
       "A single documented IPC Manual covering the programme's structure, aims and "
       "objectives, all processes, activities and surveillance procedures.")
    lb(doc,
       "Evidence the Manual is based on current scientific knowledge and "
       "national/international guidelines.")
    lb(doc,
       "Version control record showing the Manual has been reviewed at least annually.")

    h(doc, 2, "High-risk register and written guidance — IPC.1.b (Commitment*)")
    lb(doc,
       "Documented list of high-risk activities and high-risk areas identified for "
       f"{HN}'s scope of services.")
    lb(doc,
       "Written guidance for each identified high-risk activity and area — one document "
       "or section per activity/area.")
    lb(doc,
       "Evidence the list and guidance were last reviewed in line with the annual update.")

    h(doc, 2, "Annual review record — IPC.1.c")
    lb(doc,
       "Dated review record, or IPC committee minute, confirming the annual update was "
       "completed (or, if no gaps were found, documenting that nil-finding explicitly).")

    h(doc, 2, "Assessment tool output — IPC.1.d (Achievement)")
    lb(doc,
       "Completed validated assessment tool (e.g., WHO IPCAF) with scoring and gap "
       "analysis, dated.")
    lb(doc,
       "Action plan derived from the assessment, with owners and due dates, and "
       "records of closure.")

    h(doc, 2, "IPC committee records — IPC.1.e (Commitment*)")
    lb(doc,
       "Committee constitution and terms of reference defining composition, meeting "
       "frequency (at least monthly), and minimum quorum.")
    lb(doc,
       "Documented minutes of monthly meetings, naming decisions, owners, and due dates.")
    lb(doc,
       "Annual risk-reduction goals and measurable objectives document, with evidence "
       "of monthly review.")

    h(doc, 2, "IPC team records — IPC.1.f (Commitment*)")
    lb(doc,
       "IPC team composition list showing IPCO and IPCN(s) as the minimum, with a "
       "link-nurse register by clinical area.")
    lb(doc,
       "Evidence of team's day-to-day functioning: surveillance support, audit "
       "participation, and IPC committee representation records.")

    h(doc, 2, "IPCO designation — IPC.1.g (Commitment*)")
    lb(doc,
       "Designation letter for the IPCO naming a doctor with IPC competency.")
    lb(doc,
       "IPCO responsibilities defined in the IPC Manual.")
    lb(doc,
       "Time-allocation or role-fraction record if the IPCO holds concurrent duties.")

    h(doc, 2, "IPCN designation and staffing ratio — IPC.1.h (Commitment*)")
    lb(doc,
       "Designation letters for all IPC Nurses with qualification (registered nurse) and "
       "structured IPC training record.")
    lb(doc,
       "IPCN responsibilities defined in the IPC Manual.")
    lb(doc,
       "Staffing ratio calculation record showing at least one dedicated IPCN per 250 "
       "approved beds.")

    h(doc, 2, "Community IEC records — IPC.1.i")
    lb(doc,
       "Materials, session records, or other evidence of an IEC programme on infection "
       "prevention and control delivered to the community.")

    h(doc, 2, "Community outbreak records — IPC.1.j")
    lb(doc,
       "Written external-agency contact list (including statutory agencies).")
    lb(doc,
       "Documented internal and external communication path for outbreak response.")
    lb(doc,
       "Staff training records on outbreak participation roles.")
    lb(doc,
       "Record of at least one drill or real community outbreak response event.")

    h(doc, 1, "13. References")
    ln(doc,
       "National Accreditation Board for Hospitals and Healthcare Providers. NABH "
       "Accreditation Standards for Hospitals, 6th Edition. IPC.1.")
    ln(doc, "Guidebook interpretation supplied for IPC.1.a through IPC.1.j.")
    ln(doc,
       "WHO Infection Prevention and Control Assessment Framework at the Facility "
       "Level (WHO IPCAF). World Health Organization.")
    ln(doc,
       "CDC Infection Prevention and Control Assessment Tool for Acute Care Hospitals. "
       "Centers for Disease Control and Prevention.")
    ln(doc,
       f"Internal documents of {HN}: IPC Manual; IPC committee terms of reference; "
       "IPCO and IPCN designation letters; high-risk register; assessment tool outputs.")

    h(doc, 1, "Disclaimer")
    p(doc,
      "This policy reorganises the supplied IPC.1 objective-element wording and "
      "Guidebook interpretation into plain-language policy format. The modal strength "
      "of the source has been preserved. Optional examples and mechanisms have not been "
      "converted into mandatory requirements. The IPC Manual as a named documented "
      "programme record, the minimum monthly committee meeting frequency, the annual "
      "objective-setting and monthly review cycle, the doctor requirement for the IPCO, "
      "the registered-nurse-plus-structured-training requirement for IPCNs, and the "
      "one-IPCN-per-250-beds staffing floor have been retained verbatim. IPC.1 has no "
      "stop-work section.")

    save_and_verify(doc, "HCO_IPC_1_v2_REWRITE_DRAFT.docx")


# ══════════════════════════════════════════════════════════════════════════════
# IPC.2 — Resources for Infection Prevention and Control   (NO stop-work)
# COREs: a, c | Stars: none | Achievement: none | Excellence: none
# Prepared by: Medical Superintendent | Doc: HCO/IPC/POL/02
# ══════════════════════════════════════════════════════════════════════════════
def gen_ipc2():
    doc = Document()

    h(doc, 0, "Policy on Resources for Infection Prevention and Control")
    p(doc, HN)

    h(doc, 1, "Document control")
    doc_ctrl(doc, "HCO/IPC/POL/02", "Medical Superintendent")
    p(doc, "A blank marked ________ must be completed before issue.")

    h(doc, 1, "Statement of intent")
    p(doc,
      f"{HN} allocates a dedicated budget for infection prevention and control, "
      "provides PPE, soaps, and disinfectants at the point of use in every patient-care "
      "area, ensures at least one accessible washbasin with running water and hand rub "
      "in every patient-care area, and maintains defined isolation and barrier-nursing "
      "facilities.")

    h(doc, 1, "1. Purpose")
    p(doc,
      f"This policy explains how {HN} makes available the resources required for "
      "infection prevention and control: a dedicated budget, adequate PPE and consumables "
      "at the point of use, hand-hygiene facilities in every patient-care area, and "
      "isolation and barrier-nursing facilities.")
    p(doc,
      "Staffing and competency requirements for the IPC programme are covered in the "
      "hospital's other policies.")

    h(doc, 1, "2. Scope")
    p(doc,
      f"This policy applies to the Medical Superintendent, the IPC Officer, nursing "
      f"staff, and all staff responsible for providing or using IPC resources at {HN}.")

    h(doc, 1, "3. Policy standards")
    p(doc,
      f"{HN} allocates a budget line demarcated for IPC activities and ensures required "
      "resources — human and material — are available on a continual and sustained basis. "
      "PPE, soaps and disinfectants are available at the point of use with adequate "
      "inventory maintained; staff use PPE appropriate to the risk and remove it when the "
      "purpose is served. At least one easily accessible washbasin with running water and "
      "hand rub are available in every patient-care area. Conditions requiring isolation "
      "or barrier nursing are defined in writing; necessary resources and appropriate "
      "signage are in place to carry this out.")
    p(doc, "Staff follow the written guidance below and keep the records it requires.")

    h(doc, 1, "4. Non-negotiable rules")
    lb(doc,
       "Do not run the IPC programme without a budget line demarcated for IPC activities, "
       "and do not let required resources go unavailable.")
    lb(doc,
       "Do not leave any patient-care area without PPE, soaps and disinfectants at the "
       "point of use, and do not allow staff to use PPE inappropriate to the risk or "
       "to retain PPE after the purpose is served.")
    lb(doc,
       "Do not leave any patient-care area without at least one easily accessible "
       "washbasin with running water and hand rub.")
    lb(doc,
       "Do not carry out isolation or barrier nursing without the conditions being "
       "defined in writing, and do not carry it out without the necessary resources "
       "and appropriate signage.")

    h(doc, 1, "5. What we do")

    h(doc, 2, "5.1 Make IPC resources continually available")
    p(doc,
      "Management ensures that the resources required for the IPC programme — both human "
      "resources and materials — are made available on a continual and sustained basis. "
      "The annual budget includes a line demarcated for IPC activities, prepared taking "
      "into consideration the scope of the programme and previous years' experience.")

    h(doc, 2, "5.2 Provide PPE, soaps and disinfectants at the point of use")
    p(doc,
      "Personal protective equipment — including gloves, protective eyewear, mask, gown, "
      "boots/shoe covers, and cap/hair cover — as well as soaps and disinfectants, are "
      "available at the point of use. Adequate inventory is maintained. Staff use PPE "
      "appropriate to the risks involved and the guidance available, and remove it as "
      "soon as the purpose is served.")

    h(doc, 2, "5.3 Provide hand-hygiene facilities in every patient-care area")
    p(doc,
      "At least one easily accessible washbasin with running water is provided in every "
      "patient-care area for healthcare providers. Hand rub is also available in every "
      "patient-care area — dispensers may be installed at convenient points or carried "
      "by staff as they move between patients.")

    h(doc, 2, "5.4 Provide isolation and barrier-nursing facilities")
    p(doc,
      f"{HN} defines the conditions requiring isolation and the conditions requiring "
      "barrier nursing (or both). These are carried out when the defined conditions apply. "
      "Necessary resources — such as appropriate clothing, masks, gloves, and other PPE "
      "— are provided. Appropriate signage is displayed at isolation areas. Patients "
      "requiring isolation (contact, droplet or airborne) are placed in appropriate "
      "facilities; airborne isolation facilities are ideally negative-pressure rooms with "
      "closed doors.")

    h(doc, 1, "6. Governance and responsibility")
    gov_tbl(doc, [
        ("Medical Superintendent",
         "Accountable for ensuring IPC resources are budgeted, procured, and available "
         "on a continual and sustained basis."),
        ("Infection Prevention and Control Officer",
         "Monitors resource availability; escalates shortfalls in PPE, hand-hygiene "
         "facilities, or isolation capacity to the Medical Superintendent."),
        ("Nursing Superintendent",
         "Ensures PPE, soaps, disinfectants and hand rub are available at the point "
         "of use in all patient-care areas and that isolation signage is in place."),
        ("Stores / Materials Management",
         "Maintains adequate inventory of PPE, soaps, disinfectants and hand rub; "
         "alerts IPC Officer and Nursing Superintendent of low-stock situations."),
        ("Quality Coordinator",
         "Audits this policy; holds training records and staff acknowledgements."),
    ])

    h(doc, 1, "7. Quality monitoring")
    mon_tbl(doc, [
        ("IPC budget",
         "Annual budget includes a demarcated IPC line; resources available on a "
         "continual and sustained basis."),
        ("PPE availability",
         "PPE, soaps and disinfectants available at the point of use in all clinical "
         "areas; adequate inventory confirmed."),
        ("Hand-hygiene facilities",
         "At least one washbasin with running water and hand rub in every patient-care "
         "area — verified by IPC round."),
        ("Isolation facilities",
         "Conditions for isolation and barrier nursing defined in writing; required "
         "resources and appropriate signage in place."),
    ])

    h(doc, 1, "8. Training and staff acknowledgement")
    p(doc,
      "All clinical and nursing staff shall be familiar with the PPE types appropriate "
      "to different risk situations, the correct sequence for donning and doffing PPE, "
      "and the location and use of hand-hygiene facilities and isolation resources.")
    p(doc,
      f"I have read the Policy on Resources for Infection Prevention and Control of "
      f"{HN}. I will follow the processes described.")
    sig_tbl(doc)

    h(doc, 1, "9. Distribution")
    p(doc,
      "This policy shall be available to the IPC Officer, IPC Nurse(s), nursing staff, "
      "stores/materials management, and the Quality Coordinator.")

    h(doc, 1, "10. Abbreviations")
    abbrev_tbl(doc, [
        ("IPC",  "Infection Prevention and Control"),
        ("IPCO", "Infection Prevention and Control Officer"),
        ("NABH", "National Accreditation Board for Hospitals and Healthcare Providers"),
        ("PPE",  "Personal Protective Equipment"),
    ])

    h(doc, 1, "11. Traceability table")
    p(doc,
      "This table is an index. It is not how the policy is organised. An asterisk in "
      "the Level column means documentation of the process is required.")
    tr = tbl(doc, 5, 3)
    for ci, hdr in enumerate(("Objective Element", "Level", "Traceability to this policy")):
        tr.cell(0, ci).text = hdr
    trace_rows = [
        ("IPC.2.a", "CORE",
         "Sections 3 and 5.1 address the dedicated IPC budget line and the continual "
         "and sustained availability of human resources and materials."),
        ("IPC.2.b", "Commitment",
         "Sections 3 and 5.2 address PPE, soaps and disinfectants at the point of use, "
         "adequate inventory, appropriate use by risk, and removal when purpose is served."),
        ("IPC.2.c", "CORE",
         "Sections 3 and 5.3 address at least one washbasin with running water and "
         "hand rub in every patient-care area."),
        ("IPC.2.d", "Commitment",
         "Sections 3 and 5.4 address written definitions of isolation and barrier-nursing "
         "conditions, provision of necessary resources, and appropriate signage."),
    ]
    for ri, (oe, lvl, txt) in enumerate(trace_rows, 1):
        tr.cell(ri, 0).text = oe
        tr.cell(ri, 1).text = lvl
        tr.cell(ri, 2).text = txt

    h(doc, 1, "12. Required Records/Evidence Checklist")

    h(doc, 2, "IPC budget — IPC.2.a (CORE)")
    lb(doc,
       "Annual budget document showing a demarcated IPC line.")
    lb(doc,
       "Resource availability records (stock register, PPE log) showing continual "
       "and sustained availability.")

    h(doc, 2, "PPE availability — IPC.2.b")
    lb(doc,
       "PPE, soap and disinfectant stock records by area, showing point-of-use "
       "availability and adequate inventory.")
    lb(doc,
       "IPC round or audit record confirming PPE is available at the point of use.")

    h(doc, 2, "Hand-hygiene facilities — IPC.2.c (CORE)")
    lb(doc,
       "Facilities inventory or IPC round record confirming at least one washbasin "
       "with running water and hand rub in every patient-care area.")
    lb(doc,
       "Records of any hand-hygiene facility deficiencies and their resolution.")

    h(doc, 2, "Isolation and barrier-nursing facilities — IPC.2.d")
    lb(doc,
       "Written policy or guideline defining conditions for isolation and for "
       "barrier nursing.")
    lb(doc,
       "Resource inventory showing necessary items (PPE, gowns, signage) are in "
       "place for isolation use.")
    lb(doc,
       "Audit or round record confirming appropriate signage displayed in isolation "
       "areas in use.")

    h(doc, 1, "13. References")
    ln(doc,
       "National Accreditation Board for Hospitals and Healthcare Providers. NABH "
       "Accreditation Standards for Hospitals, 6th Edition. IPC.2.")
    ln(doc, "Guidebook interpretation supplied for IPC.2.a through IPC.2.d.")
    ln(doc,
       f"Internal documents of {HN}: IPC budget record; PPE stock records; "
       "hand-hygiene facility inventory; isolation policy.")

    h(doc, 1, "Disclaimer")
    p(doc,
      "This policy reorganises the supplied IPC.2 objective-element wording and "
      "Guidebook interpretation into plain-language policy format. The modal strength "
      "of the source has been preserved. Optional examples and mechanisms have not been "
      "converted into mandatory requirements. The requirement for at least one washbasin "
      "with running water and hand rub in every patient-care area, and the dedicated "
      "IPC budget line, have been retained as mandatory requirements. IPC.2 has no "
      "stop-work section.")

    save_and_verify(doc, "HCO_IPC_2_v2_REWRITE_DRAFT.docx")


# ══════════════════════════════════════════════════════════════════════════════
# IPC.3 — Infection Prevention and Control in Clinical Areas   (HAS stop-work)
# COREs: a, b, d, f | Stars: a*, b*, c*, d*, e* | Achievement: none
# Prepared by: Infection Prevention and Control Officer | Doc: HCO/IPC/POL/03
# Stop-work triggers: start procedure without standard precautions/hand hygiene/
#   safe injection; restricted antimicrobial without stewardship path
# ══════════════════════════════════════════════════════════════════════════════
def gen_ipc3():
    doc = Document()

    h(doc, 0, "Policy on Infection Prevention and Control in Clinical Areas")
    p(doc, HN)

    h(doc, 1, "Document control")
    doc_ctrl(doc, "HCO/IPC/POL/03", "Infection Prevention and Control Officer")
    p(doc, "A blank marked ________ must be completed before issue.")

    h(doc, 1, "Statement of intent")
    p(doc,
      f"{HN} applies standard precautions in every clinical area at all times, adheres "
      "to hand-hygiene and transmission-based precaution guidelines, practices safe "
      "injection and infusion, and operates a documented antimicrobial usage policy with "
      "a stewardship programme that governs the ordering and monitoring of restricted "
      "antimicrobials.")

    h(doc, 1, "1. Purpose")
    p(doc,
      f"This policy explains how {HN} implements infection prevention and control "
      "processes in clinical areas: universal application of standard precautions, "
      "adherence to hand-hygiene and transmission-based precaution guidelines, safe "
      "injection and infusion practices, and an antimicrobial usage policy with a "
      "stewardship programme including a restricted antimicrobial ordering mechanism.")
    p(doc,
      "IPC processes for support services, HAI-prevention bundles for specific devices "
      "and procedures, and IPC surveillance are covered in separate hospital policies.")

    h(doc, 1, "2. Scope")
    p(doc,
      f"This policy applies to all clinical staff — doctors, nurses, allied health "
      f"professionals, and students — working in any clinical area of {HN}.")

    h(doc, 1, "3. Policy standards")
    p(doc,
      f"{HN} adheres to standard precautions in every clinical area at all times. "
      "International or national hand-hygiene guidelines are followed. Transmission-based "
      "precautions for contact, droplet and airborne modes are defined and used with "
      "appropriate PPE. Safe injection practice follows the one-needle-one-syringe-one-"
      "time rule. A documented antimicrobial usage policy — covering the type, "
      "monotherapy versus combination therapy, escalation and de-escalation, dose and "
      "duration, and a restricted list aligned with WHO AWaRe — is reviewed at least "
      "annually based on the organisation's site-specific antibiogram. Antimicrobials "
      "are prescribed only as per the policy; a defined mechanism governs ordering of "
      "restricted agents; deviations are brought to clinicians and CAPA documented; and "
      "the antimicrobial stewardship programme is monitored by a stewardship forum of "
      "clinicians.")
    p(doc, "Staff follow the written guidance below and keep the records it requires.")

    h(doc, 1, "4. Non-negotiable rules")
    lb(doc,
       "Do not start any clinical procedure unless standard precautions — including "
       "appropriate PPE, sharps safety, and respiratory etiquette — are in place.")
    lb(doc,
       "Do not start an invasive step without having performed the required hand-hygiene "
       "moment for that step.")
    lb(doc,
       "Do not reuse a needle or syringe — one needle, one syringe, one time only.")
    lb(doc,
       "Do not start a restricted antimicrobial without following the organisation's "
       "antimicrobial stewardship path (documented indication and, where required, "
       "prior approval).")
    lb(doc,
       "Do not prescribe antimicrobials outside the organisation's policy.")
    lb(doc,
       "Do not operate without a current antimicrobial usage policy aligned with WHO "
       "AWaRe for the restricted list, reviewed at least annually.")

    h(doc, 1, "5. What we do")

    h(doc, 2, "5.1 Apply standard precautions at all times")
    p(doc,
      "Standard precautions are adhered to in every area of the organisation, at all "
      "times. Components include hand hygiene, appropriate use of PPE, respiratory "
      "etiquette, safe injection and infusion practices, use of sterile instruments and "
      "devices, cleaning and disinfection of environmental surfaces and equipment, and "
      "needle-stick and sharps injury prevention. Appropriate preparation of body parts "
      "before any procedure, and the use of disinfected or sterilised instruments, is "
      "ensured.")

    h(doc, 2, "5.2 Adhere to hand-hygiene guidelines")
    p(doc,
      f"{HN} adheres to international or national guidelines on hand hygiene. The WHO "
      "guidelines on hand hygiene in health care (2009) or the national equivalent serve "
      "as the primary reference. Instructions are displayed near hand-washing areas. "
      "Compliance with hand-hygiene guidelines is monitored as part of the hospital's "
      "IPC surveillance programme.")

    h(doc, 2, "5.3 Apply transmission-based precautions")
    p(doc,
      "Transmission-based precautions cover airborne, droplet and contact modes of "
      "transmission. PPE appropriate to each mode and situation is identified and used "
      "accordingly. These precautions are applicable across all areas of the organisation. "
      "CDC or equivalent guidelines are the reference for implementation.")

    h(doc, 2, "5.4 Practice safe injection and infusion")
    p(doc,
      "Safe injection practice at all injection and infusion points follows the "
      "one-needle-one-syringe-one-time rule as recommended by CDC. Sharps are disposed "
      "of in puncture-proof containers at the point of use immediately after use. The "
      "WHO best-practices toolkit for injections and related procedures is a useful "
      "reference for the organisation's safe-injection guidance.")

    h(doc, 2, "5.5 Establish and document the antimicrobial usage policy")
    p(doc,
      f"{HN}'s antimicrobial usage policy identifies the clinical conditions in which "
      "antimicrobial agents — including anti-bacterial, anti-tubercular, anti-fungal, "
      "anti-viral and anti-parasitic agents — are used, in terms of: the type of "
      "antimicrobial agent; monotherapy versus combination therapy; escalation and "
      "de-escalation of therapy; and the dose and duration of therapy. The policy is "
      "based on the organisation's site-specific antibiogram and antimicrobial "
      "susceptibility data (culture-sensitivity based), updated at least annually. "
      "The policy identifies a list of restricted antimicrobial agents. The list of "
      "restricted antibiotics adheres to the WHO AWaRe classification.")

    h(doc, 2, "5.6 Implement the antimicrobial stewardship programme")
    p(doc,
      "Antimicrobials are prescribed as per the organisation's antimicrobial usage "
      "policy. A defined mechanism exists for ordering restricted antimicrobial agents. "
      "Deviations from the policy are brought to the notice of the concerned clinician; "
      "corrective and preventive actions are taken and documented. The appropriate use "
      "of restricted agents is monitored. The antimicrobial stewardship programme is "
      "monitored by an antimicrobial stewardship forum of clinicians, and includes "
      "elements of leadership, accountability, drug expertise, action, tracking, "
      "reporting, and education. The organisation generates awareness among healthcare "
      "providers and the community on antimicrobial resistance and rational use of "
      "antimicrobials.")

    h(doc, 1, "6. Stop-work authority")
    p(doc,
      "Do not start a clinical procedure when standard precautions are not in place "
      "(PPE, sharps safety, cough etiquette as the situation requires), when hand "
      "hygiene has not been done at the required moment, or when injection/infusion "
      "practice would reuse a needle or syringe.")
    p(doc,
      "Do not start a restricted antimicrobial without the organisation's antimicrobial-"
      "stewardship path (documented indication and, where required, approval).")
    p(doc,
      "Stop-work applies to the procedure or first dose start. Immediate life-saving "
      "care continues with the best available precautions and is documented.")
    p(doc,
      f"The person who stops tells the treating doctor and the Infection Prevention and "
      "Control Officer the same shift. Refusing an unsafe start is not a disciplinary "
      "matter.")

    h(doc, 1, "7. Governance and responsibility")
    gov_tbl(doc, [
        ("Medical Superintendent",
         "Accountable for ensuring clinical IPC processes are resourced and implemented "
         "across the organisation."),
        ("Infection Prevention and Control Officer",
         "Owns implementation of this policy; leads the antimicrobial stewardship forum; "
         "receives stop-work escalations the same shift; brings CAPA and audit findings "
         "to the IPC committee."),
        ("Infection Prevention and Control Nurse(s)",
         "Monitors compliance with standard precautions, hand hygiene, transmission-"
         "based precautions, and safe injection; raises stop-work when a trigger is "
         "observed."),
        ("Antimicrobial Stewardship Forum",
         "Monitors the antimicrobial stewardship programme; reviews restricted-agent "
         "orders and deviations; oversees the site-specific antibiogram update."),
        ("Nursing Superintendent",
         "Ensures nursing staff in all clinical areas follow standard and transmission-"
         "based precautions, hand hygiene, and safe-injection practice."),
        ("Quality Coordinator",
         "Audits this policy; holds training records and staff acknowledgements; "
         "logs stop-work events."),
    ])

    h(doc, 1, "8. Quality monitoring")
    mon_tbl(doc, [
        ("Standard precautions compliance",
         "IPC round records confirming standard precautions in use in all clinical "
         "areas; audit results tabled at IPC committee."),
        ("Hand-hygiene compliance",
         "Monthly compliance monitoring as per IPC.6.d; results shared with relevant "
         "staff and tabled at IPC committee."),
        ("Transmission-based precaution compliance",
         "Audit records confirming PPE appropriate to each mode is used and available."),
        ("Safe injection compliance",
         "Audit records confirming one-needle-one-syringe-one-time practice; "
         "needle-stick injury reports reviewed."),
        ("Antimicrobial usage policy",
         "Current policy reviewed at least annually; restricted list aligned with "
         "WHO AWaRe; site-specific antibiogram updated."),
        ("Antimicrobial stewardship",
         "Restricted-agent orders tracked; deviation records with CAPA documented; "
         "stewardship forum meeting minutes."),
        ("Stop-work events",
         "Stop-work events logged with trigger, action taken, and outcome."),
    ])

    h(doc, 1, "9. Training and staff acknowledgement")
    p(doc,
      "All clinical staff shall be familiar with standard precautions, hand-hygiene "
      "guidelines, transmission-based precautions, safe injection practice, the "
      "antimicrobial usage policy, the restricted-agent ordering mechanism, and the "
      "stop-work authority in this policy.")
    p(doc,
      f"I have read the Policy on Infection Prevention and Control in Clinical Areas "
      f"of {HN}. I will follow the processes described, including the stop-work "
      "authority in Section 6.")
    sig_tbl(doc)

    h(doc, 1, "10. Distribution")
    p(doc,
      "This policy shall be available to all clinical staff, the IPC Officer, IPC "
      "Nurse(s), the antimicrobial stewardship forum, department heads, and the "
      "Quality Coordinator.")

    h(doc, 1, "11. Abbreviations")
    abbrev_tbl(doc, [
        ("AMS",   "Antimicrobial Stewardship"),
        ("AMR",   "Antimicrobial Resistance"),
        ("AWaRe", "Access, Watch, Reserve (WHO antibiotic classification)"),
        ("CAPA",  "Corrective and Preventive Action"),
        ("CDC",   "Centers for Disease Control and Prevention"),
        ("IPC",   "Infection Prevention and Control"),
        ("IPCO",  "Infection Prevention and Control Officer"),
        ("IPCN",  "Infection Prevention and Control Nurse"),
        ("NABH",  "National Accreditation Board for Hospitals and Healthcare Providers"),
        ("PPE",   "Personal Protective Equipment"),
        ("WHO",   "World Health Organization"),
    ])

    h(doc, 1, "12. Traceability table")
    p(doc,
      "This table is an index. It is not how the policy is organised. An asterisk in "
      "the Level column means documentation of the process is required.")
    tr = tbl(doc, 7, 3)
    for ci, hdr in enumerate(("Objective Element", "Level", "Traceability to this policy")):
        tr.cell(0, ci).text = hdr
    trace_rows = [
        ("IPC.3.a", "CORE*",
         "Sections 3, 5.1 and 6 address universal application of standard precautions "
         "including PPE, sharps safety, respiratory etiquette, sterile instruments, and "
         "surface cleaning; stop-work triggers when precautions are not in place."),
        ("IPC.3.b", "CORE*",
         "Sections 3, 5.2 and 6 address adherence to hand-hygiene guidelines; "
         "stop-work trigger when required hand-hygiene moment is missed before an "
         "invasive step."),
        ("IPC.3.c", "Commitment*",
         "Sections 3 and 5.3 address transmission-based precautions (airborne, droplet, "
         "contact), with PPE identified for each mode, applicable across the organisation."),
        ("IPC.3.d", "CORE*",
         "Sections 3, 5.4 and 6 address the one-needle-one-syringe-one-time safe "
         "injection rule; stop-work trigger when this rule would be violated."),
        ("IPC.3.e", "Commitment*",
         "Sections 3 and 5.5 address the documented antimicrobial usage policy covering "
         "type, mono/combination therapy, escalation/de-escalation, dose, duration, and "
         "a restricted list aligned with WHO AWaRe, reviewed at least annually."),
        ("IPC.3.f", "CORE",
         "Sections 3, 5.6 and 6 address the AMS programme: prescribing per policy, "
         "restricted-agent ordering mechanism, deviation CAPA, stewardship forum, and "
         "the seven AMS programme elements; stop-work trigger for restricted antimicrobial "
         "without the stewardship path."),
    ]
    for ri, (oe, lvl, txt) in enumerate(trace_rows, 1):
        tr.cell(ri, 0).text = oe
        tr.cell(ri, 1).text = lvl
        tr.cell(ri, 2).text = txt

    h(doc, 1, "13. Required Records/Evidence Checklist")

    h(doc, 2, "Standard precautions — IPC.3.a (CORE*)")
    lb(doc,
       "Written standard-precautions guidance (or section within IPC Manual) covering "
       "all components: hand hygiene, PPE, respiratory etiquette, safe injection, sterile "
       "instruments, environmental cleaning, and sharps injury prevention.")
    lb(doc,
       "IPC round or audit records confirming standard precautions are in use in all "
       "clinical areas.")
    lb(doc,
       "Staff training records on standard precautions.")

    h(doc, 2, "Hand hygiene — IPC.3.b (CORE*)")
    lb(doc,
       "Named hand-hygiene guideline adopted (e.g., WHO 2009 or national equivalent).")
    lb(doc,
       "Monthly hand-hygiene compliance monitoring records (see also IPC.6.d).")
    lb(doc,
       "Evidence of display of hand-hygiene instructions near washbasins.")

    h(doc, 2, "Transmission-based precautions — IPC.3.c (Commitment*)")
    lb(doc,
       "Written transmission-based precaution guidance covering contact, droplet, and "
       "airborne modes with PPE specification for each.")
    lb(doc,
       "Audit records confirming appropriate PPE use by mode in clinical areas.")

    h(doc, 2, "Safe injection — IPC.3.d (CORE*)")
    lb(doc,
       "Written safe-injection and infusion guidance naming the one-needle-one-syringe-"
       "one-time rule.")
    lb(doc,
       "Audit records and needle-stick injury log reviewed for reuse events.")
    lb(doc,
       "Records confirming puncture-proof sharps containers are at the point of use.")

    h(doc, 2, "Antimicrobial usage policy — IPC.3.e (Commitment*)")
    lb(doc,
       "Dated antimicrobial usage policy covering all five required content elements: "
       "type, mono/combination, escalation/de-escalation, dose, and duration.")
    lb(doc,
       "Restricted antimicrobial list aligned with WHO AWaRe.")
    lb(doc,
       "Site-specific antibiogram/culture-sensitivity data used for the policy update.")
    lb(doc,
       "Annual review record confirming the policy remains current.")

    h(doc, 2, "Antimicrobial stewardship programme — IPC.3.f (CORE)")
    lb(doc,
       "Restricted-agent ordering mechanism (forms, electronic approvals, or equivalent) "
       "and records of its use.")
    lb(doc,
       "Deviation records with CAPA documented and tracked to closure.")
    lb(doc,
       "Antimicrobial stewardship forum meeting minutes.")
    lb(doc,
       "Stop-work event log with trigger, action, and outcome for each event.")

    h(doc, 1, "14. References")
    ln(doc,
       "National Accreditation Board for Hospitals and Healthcare Providers. NABH "
       "Accreditation Standards for Hospitals, 6th Edition. IPC.3.")
    ln(doc, "Guidebook interpretation supplied for IPC.3.a through IPC.3.f.")
    ln(doc,
       "WHO Guidelines on Hand Hygiene in Health Care, 2009. World Health Organization.")
    ln(doc,
       "WHO AWaRe (Access, Watch, Reserve) Antibiotic Book. World Health Organization.")
    ln(doc,
       "CDC Safe Injection Practices: One Needle, One Syringe, Only One Time.")
    ln(doc,
       f"Internal documents of {HN}: IPC Manual — standard precautions, hand hygiene, "
       "transmission precautions, safe injection sections; antimicrobial usage policy; "
       "site-specific antibiogram; stewardship forum minutes.")

    h(doc, 1, "Disclaimer")
    p(doc,
      "This policy reorganises the supplied IPC.3 objective-element wording and "
      "Guidebook interpretation into plain-language policy format. The modal strength "
      "of the source has been preserved. Optional examples and mechanisms have not been "
      "converted into mandatory requirements. The one-needle-one-syringe-one-time rule, "
      "the five required elements of the antimicrobial usage policy, the WHO AWaRe "
      "alignment for the restricted list, and the annual review frequency have been "
      "retained verbatim. IPC.3 carries stop-work authority as stated in Section 6: "
      "procedure start without standard precautions or hand hygiene, needle or syringe "
      "reuse, and restricted antimicrobial without the stewardship path.")

    save_and_verify(doc, "HCO_IPC_3_v2_REWRITE_DRAFT.docx")


# ══════════════════════════════════════════════════════════════════════════════
# IPC.4 — Infection Prevention and Control in Support Services   (HAS stop-work)
# COREs: c, d | Stars: a*, b*, c*, e*, f* | Achievement: none
# Prepared by: Infection Prevention and Control Officer | Doc: HCO/IPC/POL/04
# Stop-work triggers: BMW without required segregation/PPE; construction/renovation
#   without the infection-risk plan approved by IPCO
# Statute P2: IPC.4.d (Biomedical Waste Management Rules)
# ══════════════════════════════════════════════════════════════════════════════
def gen_ipc4():
    doc = Document()

    h(doc, 0, "Policy on Infection Prevention and Control in Support Services")
    p(doc, HN)

    h(doc, 1, "Document control")
    doc_ctrl(doc, "HCO/IPC/POL/04", "Infection Prevention and Control Officer")
    p(doc, "A blank marked ________ must be completed before issue.")
    p(doc,
      "STATUTORY NOTE: Section 5.4 of this policy (biomedical waste) is designated "
      "statute P2. Those provisions follow the Biomedical Waste Management Rules "
      "applicable at the time of implementation. The accreditation-specific guidance "
      "in all other sections does not carry statute status.")

    h(doc, 1, "Statement of intent")
    p(doc,
      f"{HN} maintains engineering controls to prevent infection, uses a validated "
      "infection-risk assessment tool before any construction or renovation, follows "
      "documented housekeeping procedures across all areas, handles biomedical waste as "
      "per statutory provisions, and adheres to defined laundry/linen and kitchen "
      "sanitation processes — regardless of whether these services are in-house or "
      "outsourced.")

    h(doc, 1, "1. Purpose")
    p(doc,
      f"This policy explains how {HN} implements infection prevention and control "
      "processes in support services: engineering controls, construction/renovation "
      "infection risk management, housekeeping, biomedical waste, laundry and linen, "
      "and kitchen sanitation.")
    p(doc,
      "Clinical IPC processes, HAI-prevention bundles, and IPC surveillance are "
      "covered in separate hospital policies.")

    h(doc, 1, "2. Scope")
    p(doc,
      f"This policy applies to housekeeping staff, BMW handlers, laundry staff, "
      f"kitchen and dietary staff, engineering/maintenance teams, and the IPC Officer "
      f"at {HN}, including those performing outsourced services.")

    h(doc, 1, "3. Policy standards")
    p(doc,
      f"{HN} maintains engineering controls including patient-care area design, OT "
      "air quality, water supply and HVAC maintenance. A validated infection-risk "
      "assessment tool is used before any construction or renovation; no works begin "
      "in a patient-care area without an approved infection-risk plan. Housekeeping "
      "is addressed at all levels of the organisation, with regular mandatory cleaning "
      "and documented procedures for terminal cleaning, blood/body-fluid cleanup and "
      "isolation rooms; dusting inside clinical areas is prohibited. Biomedical waste "
      "is segregated in colour-coded containers as per statutory provisions, handled "
      "with appropriate PPE, stored per statutory requirements, and handed to an "
      "authorised vendor. Laundry and linen management follows a defined process with "
      "clean linen always separated from dirty. Kitchen and dietary services adhere to "
      "statutory requirements, implement sanitation measures, and ensure food is safely "
      "prepared, handled, stored and distributed.")
    p(doc, "Staff follow the written guidance below and keep the records it requires.")

    h(doc, 1, "4. Non-negotiable rules")
    lb(doc,
       "Do not start construction or renovation in a patient-care area until the "
       "infection-risk plan for that work has been completed and the IPC Officer has "
       "agreed the controls.")
    lb(doc,
       "Do not handle biomedical waste without the required colour-coded segregation "
       "and appropriate PPE.")
    lb(doc,
       "Do not perform any dusting inside clinical areas.")
    lb(doc,
       "Do not use common disinfectants without an established dilution protocol based "
       "on national/international guidelines.")
    lb(doc,
       "Do not store biomedical waste outside statutory requirements, and do not hand "
       "it to anyone other than an authorised vendor.")
    lb(doc,
       "Do not perform routine environmental microbiological sampling of OT and high-risk "
       "areas, do not conduct periodic fumigation, and do not fog unless commissioning a "
       "new OT, following modification/renovation of an OT, or when there is a "
       "significant increase in SSI with environmental evidence.")
    lb(doc,
       "Do not mix clean linen with dirty linen at any point.")

    h(doc, 1, "5. What we do")

    h(doc, 2, "5.1 Maintain engineering controls to prevent infection")
    p(doc,
      "Engineering controls address patient-care area design — including optimum spacing "
      "between beds — OT design and zoning, air quality, and water supply. The "
      "organisation follows NABH guidelines on OT air-conditioning. HVAC plant and "
      "equipment, air-conditioning ducts and filters, air-handling units, and cooling "
      "towers are maintained regularly; filters are replaced on schedule. Water-supply "
      "sources, systems, and water-quality testing are included in the engineering "
      "controls. Fungal colonisation prevention is addressed. Those requirements related "
      "to facilities management are covered in the hospital's other policies.")

    h(doc, 2, "5.2 Plan and implement infection controls for construction and renovation")
    p(doc,
      "Before any construction or renovation in or adjacent to a patient-care area, a "
      "validated infection-risk assessment tool — such as the Infection Control Risk "
      "Assessment (ICRA) — is used to identify the risk of infection. An infection "
      "prevention plan covering architectural segregation, traffic-flow controls, and "
      "materials management is developed and reviewed by the IPC team. Construction or "
      "renovation does not begin until the plan is complete and the IPC Officer has "
      f"agreed the controls. {HN} considers IPC from the design stage when building "
      "new facilities. Surveillance of demolition, construction and repair areas is "
      "conducted as part of the IPC surveillance programme.")

    h(doc, 2, "5.3 Follow housekeeping procedures")
    p(doc,
      "Housekeeping is addressed at all levels of the organisation — wards, OT, public "
      "areas including toilets, and corridors. Regular cleaning to remove visible dirt "
      "and dust is mandatory. The environment, fixtures, fomites, furniture, furnishings, "
      "equipment, and surfaces are cleaned as applicable. Common disinfectants are "
      "identified per national/international guidelines; dilution protocols are "
      "established and followed. Written procedures exist and are followed for terminal "
      "cleaning, blood and body-fluid cleanup, and isolation rooms. Dusting of any sort "
      "inside clinical areas is avoided. The effectiveness of housekeeping is monitored "
      "regularly through outcome parameters rather than only frequency verification — "
      "see also the IPC surveillance policy.")

    h(doc, 2, "5.4 Handle biomedical waste appropriately and safely")
    p(doc,
      "Biomedical waste from patient-care areas is properly segregated and collected in "
      "colour-coded bags and containers as per the statutory provisions of the applicable "
      "Biomedical Waste Management Rules. BMW is handled with appropriate PPE. BMW is "
      "stored in accordance with statutory provisions. BMW is handed over to an authorised "
      "vendor for transport to the site of treatment and disposal. Monitoring of the BMW "
      "management programme is done by members of the IPC committee/team as per statutory "
      "requirements. These obligations follow the Biomedical Waste Management Rules "
      f"applicable at the time of implementation; {HN} ensures its practices remain "
      "current with any amendments to those Rules.")

    h(doc, 2, "5.5 Follow laundry and linen management processes")
    p(doc,
      "Linen management follows a defined process whether the service is in-house or "
      f"outsourced. {HN} documents the change-of-linen procedure. A defined process "
      "covers handling linen in patient-care units, during transport to the laundry, "
      "and inside the laundry. The process for storage and distribution of clean linen "
      "is defined and followed. Clean linen is separated from dirty linen at all times. "
      "Where laundry is outsourced, adequate IPC controls are established to ensure "
      "compliance.")

    h(doc, 2, "5.6 Follow kitchen sanitation and food-handling requirements")
    p(doc,
      "Kitchen sanitation and food-handling requirements apply whether the dietary "
      f"service is in-house or outsourced. {HN} adheres to all statutory requirements, "
      "including screening of kitchen workers and food handlers. Kitchen sanitation "
      "measures are implemented to prevent cross-contamination. Food is prepared, "
      "handled, stored and distributed safely. Dietary services are designed with no "
      "criss-cross of traffic and with activities flowing in sequence. Hygienic "
      "conditions are maintained throughout.")

    h(doc, 1, "6. Stop-work authority")
    p(doc,
      "Do not handle biomedical waste without the required colour-coded segregation "
      "and PPE.")
    p(doc,
      "Do not start construction or renovation in a patient-care area until the "
      "infection-risk plan for that work has been completed and the Infection Prevention "
      "and Control Officer has agreed the controls.")
    p(doc,
      "Stop-work applies to the waste-handling act and to the construction/renovation "
      "start. Emergency repair that cannot wait follows the documented emergency-works "
      "IPC controls and is recorded.")
    p(doc,
      f"The person who stops tells the Infection Prevention and Control Officer and "
      "the Medical Superintendent the same shift. Refusing unsafe waste handling or "
      "unplanned works is not a disciplinary matter.")

    h(doc, 1, "7. Governance and responsibility")
    gov_tbl(doc, [
        ("Medical Superintendent",
         "Accountable for ensuring support-service IPC processes are resourced and "
         "implemented; receives stop-work escalations the same shift for construction "
         "and renovation triggers."),
        ("Infection Prevention and Control Officer",
         "Owns implementation of this policy; approves infection-risk plans before "
         "construction or renovation; receives stop-work escalations the same shift; "
         "monitors BMW management programme compliance."),
        ("Engineering / Maintenance In-Charge",
         "Ensures engineering controls, HVAC and water systems are maintained; "
         "submits planned construction and renovation works to the IPC team for "
         "ICRA before commencement."),
        ("Housekeeping In-Charge",
         "Ensures housekeeping procedures are followed in all areas; reports deviations "
         "to the IPC Officer."),
        ("BMW / Waste Management Officer",
         "Ensures colour-coded segregation, PPE use, statutory storage, and authorised "
         "vendor handover; maintains BMW monitoring records."),
        ("Dietary / Kitchen In-Charge",
         "Ensures statutory requirements, sanitation measures, and safe food-handling "
         "practices are followed, whether in-house or outsourced."),
        ("Quality Coordinator",
         "Audits this policy; holds training records, staff acknowledgements, and "
         "stop-work event logs."),
    ])

    h(doc, 1, "8. Quality monitoring")
    mon_tbl(doc, [
        ("Engineering controls",
         "HVAC and water-system maintenance records; OT air-quality records; "
         "engineering-round records."),
        ("Construction/renovation ICRA",
         "Completed ICRA tool and approved infection-risk plan on file before works "
         "start; IPC Officer sign-off documented."),
        ("Housekeeping effectiveness",
         "Outcome parameters monitored regularly; terminal-clean records for isolation "
         "rooms; no-dusting policy confirmed in clinical areas."),
        ("BMW management",
         "Segregation audit records; PPE compliance; statutory storage records; "
         "authorised vendor handover log; statutory monitoring records."),
        ("Laundry and linen",
         "Linen management process records; clean/dirty separation confirmed in IPC "
         "round."),
        ("Kitchen sanitation",
         "Kitchen worker screening records; statutory compliance records; sanitation "
         "round records."),
        ("Stop-work events",
         "Stop-work events logged with trigger, action taken, and outcome."),
    ])

    h(doc, 1, "9. Training and staff acknowledgement")
    p(doc,
      "Housekeeping staff, BMW handlers, laundry staff, kitchen staff, and "
      "engineering/maintenance staff shall be familiar with their respective "
      "IPC obligations under this policy, including the stop-work authority in "
      "Section 6.")
    p(doc,
      f"I have read the Policy on Infection Prevention and Control in Support Services "
      f"of {HN}. I will follow the processes described, including the stop-work "
      "authority in Section 6.")
    sig_tbl(doc)

    h(doc, 1, "10. Distribution")
    p(doc,
      "This policy shall be available to the IPC Officer, housekeeping staff, BMW "
      "handlers, laundry staff, kitchen staff, engineering/maintenance teams, department "
      "heads, and the Quality Coordinator.")

    h(doc, 1, "11. Abbreviations")
    abbrev_tbl(doc, [
        ("BMW",  "Biomedical Waste"),
        ("CAPA", "Corrective and Preventive Action"),
        ("HEPA", "High-Efficiency Particulate Air"),
        ("HVAC", "Heating, Ventilation and Air Conditioning"),
        ("ICRA", "Infection Control Risk Assessment"),
        ("IPC",  "Infection Prevention and Control"),
        ("IPCO", "Infection Prevention and Control Officer"),
        ("NABH", "National Accreditation Board for Hospitals and Healthcare Providers"),
        ("OT",   "Operation Theatre"),
        ("PPE",  "Personal Protective Equipment"),
        ("SSI",  "Surgical Site Infection"),
    ])

    h(doc, 1, "12. Traceability table")
    p(doc,
      "This table is an index. It is not how the policy is organised. An asterisk in "
      "the Level column means documentation of the process is required.")
    tr = tbl(doc, 7, 3)
    for ci, hdr in enumerate(("Objective Element", "Level", "Traceability to this policy")):
        tr.cell(0, ci).text = hdr
    trace_rows = [
        ("IPC.4.a", "Commitment*",
         "Sections 3 and 5.1 address engineering controls for patient-care area design, "
         "OT air quality, water supply, and HVAC/duct/filter/cooling-tower maintenance."),
        ("IPC.4.b", "Commitment*",
         "Sections 3, 5.2 and 6 address the validated infection-risk assessment "
         "(ICRA or equivalent) and approved infection-risk plan before any "
         "construction/renovation; stop-work trigger at construction/renovation start."),
        ("IPC.4.c", "CORE*",
         "Sections 3 and 5.3 address mandatory regular cleaning, disinfectant "
         "identification with dilution protocols, terminal cleaning/blood spill/isolation "
         "room procedures, and the prohibition on dusting in clinical areas."),
        ("IPC.4.d", "CORE",
         "Sections 3, 5.4 and 6 address colour-coded BMW segregation per statutory "
         "provisions, PPE, statutory storage, authorised vendor handover, and monitoring "
         "per statutory requirements; stop-work trigger for BMW without segregation/PPE. "
         "Statute P2 designated."),
        ("IPC.4.e", "Commitment*",
         "Sections 3 and 5.5 address the documented linen management process: in-unit "
         "handling, transport, laundry, storage and distribution, clean/dirty separation "
         "at all times, and outsourcing controls."),
        ("IPC.4.f", "Commitment*",
         "Sections 3 and 5.6 address kitchen sanitation and food handling: statutory "
         "requirements including worker screening, sanitation measures, safe food "
         "preparation/handling/storage/distribution, and traffic/sequence design."),
    ]
    for ri, (oe, lvl, txt) in enumerate(trace_rows, 1):
        tr.cell(ri, 0).text = oe
        tr.cell(ri, 1).text = lvl
        tr.cell(ri, 2).text = txt

    h(doc, 1, "13. Required Records/Evidence Checklist")

    h(doc, 2, "Engineering controls — IPC.4.a (Commitment*)")
    lb(doc,
       "Engineering controls record (bed layout/spacing, OT zoning, water and air "
       "systems) and annual review record.")
    lb(doc,
       "HVAC and duct/filter/cooling-tower maintenance schedule and service logs.")
    lb(doc,
       "Water-quality test results.")

    h(doc, 2, "Construction/renovation ICRA — IPC.4.b (Commitment*)")
    lb(doc,
       "Completed ICRA tool or equivalent for each construction/renovation project, "
       "before works begin.")
    lb(doc,
       "Written infection-risk plan covering architectural segregation, traffic flow, "
       "and material controls, with IPC Officer sign-off.")
    lb(doc,
       "Records confirming no works started before plan approval.")

    h(doc, 2, "Housekeeping — IPC.4.c (CORE*)")
    lb(doc,
       "Written housekeeping procedures by area (ward, OT, public areas) including "
       "cleaning frequency, products, dilution protocols, and special procedures for "
       "terminal cleaning, blood/body-fluid cleanup, and isolation rooms.")
    lb(doc,
       "Cleaning records and outcome-monitoring records (see also IPC.6.f).")
    lb(doc,
       "IPC round records confirming no dusting inside clinical areas.")

    h(doc, 2, "Biomedical waste — IPC.4.d (CORE, Statute P2)")
    lb(doc,
       "BMW segregation audit records confirming colour-coded bags and containers per "
       "statutory provisions.")
    lb(doc,
       "PPE compliance records for BMW handlers.")
    lb(doc,
       "BMW storage records confirming statutory requirements met.")
    lb(doc,
       "Authorised vendor agreement and handover logs.")
    lb(doc,
       "BMW monitoring records as per statutory requirements.")
    lb(doc,
       "Stop-work event log for BMW-trigger events.")

    h(doc, 2, "Laundry and linen — IPC.4.e (Commitment*)")
    lb(doc,
       "Written linen management process covering handling in units, transport, laundry, "
       "storage and distribution of clean linen.")
    lb(doc,
       "Change-of-linen documentation.")
    lb(doc,
       "Outsourcing IPC control records, where applicable.")

    h(doc, 2, "Kitchen sanitation — IPC.4.f (Commitment*)")
    lb(doc,
       "Kitchen worker and food-handler screening records (statutory requirement).")
    lb(doc,
       "Kitchen sanitation procedure and daily sanitation check records.")
    lb(doc,
       "IPC round record that includes the dietary/kitchen area.")

    h(doc, 1, "14. References")
    ln(doc,
       "National Accreditation Board for Hospitals and Healthcare Providers. NABH "
       "Accreditation Standards for Hospitals, 6th Edition. IPC.4.")
    ln(doc, "Guidebook interpretation supplied for IPC.4.a through IPC.4.f.")
    ln(doc,
       "Biomedical Waste Management Rules (as amended), Government of India. "
       "[Statute P2 — IPC.4.d.]")
    ln(doc,
       f"Internal documents of {HN}: engineering-controls register; ICRA records; "
       "housekeeping procedures and logs; BMW segregation and monitoring records; "
       "authorised vendor agreement; linen management procedure; kitchen sanitation records.")

    h(doc, 1, "Disclaimer")
    p(doc,
      "This policy reorganises the supplied IPC.4 objective-element wording and "
      "Guidebook interpretation into plain-language policy format. The modal strength "
      "of the source has been preserved. Optional examples and mechanisms have not been "
      "converted into mandatory requirements. The requirement for a validated ICRA before "
      "construction/renovation, mandatory regular cleaning, the prohibition on dusting "
      "in clinical areas, the prohibition on routine environmental microbiological "
      "sampling and periodic fumigation, the statutory BMW management obligations, "
      "and the clean/dirty linen separation have been retained. IPC.4 carries stop-work "
      "authority as stated in Section 6: BMW handling without required segregation/PPE, "
      "and construction/renovation start without an approved infection-risk plan. "
      "IPC.4.d carries statute P2 status for the BMW provisions.")

    save_and_verify(doc, "HCO_IPC_4_v2_REWRITE_DRAFT.docx")


# ══════════════════════════════════════════════════════════════════════════════
# IPC.5 — Prevention of Healthcare Associated Infections in Patients
#          (HAS stop-work: Section 6)
# All 4 OEs are Commitment level; no stars; no COREs in this standard.
# Prepared by: Infection Prevention and Control Officer | Doc: HCO/IPC/POL/05
# Stop-work triggers: inserting urinary catheter, central line, starting ventilation,
#   or making a surgical incision without the relevant HAI-prevention bundle in place
# ══════════════════════════════════════════════════════════════════════════════
def gen_ipc5():
    doc = Document()

    h(doc, 0, "Policy on Prevention of Healthcare Associated Infections in Patients")
    p(doc, HN)

    h(doc, 1, "Document control")
    doc_ctrl(doc, "HCO/IPC/POL/05", "Infection Prevention and Control Officer")
    p(doc, "A blank marked ________ must be completed before issue.")

    h(doc, 1, "Statement of intent")
    p(doc,
      f"{HN} develops, implements and monitors care bundles to prevent catheter-associated "
      "urinary tract infections (CAUTI), ventilator-associated pneumonia (VAP), central "
      "line-associated bloodstream infections (CLABSI), and surgical site infections (SSI). "
      "No relevant device insertion or surgical incision is made unless the organisation's "
      "prevention bundle for that HAI is in place for the patient.")

    h(doc, 1, "1. Purpose")
    p(doc,
      f"This policy explains how {HN} takes actions to prevent healthcare associated "
      "infections in patients by developing, implementing and monitoring evidence-based "
      "care bundles for CAUTI, VAP, CLABSI, and SSI prevention, alongside adherence to "
      "standard and transmission-based precautions.")
    p(doc,
      "IPC in clinical areas (standard precautions, hand hygiene, antimicrobial "
      "stewardship) and IPC surveillance are covered in separate hospital policies.")

    h(doc, 1, "2. Scope")
    p(doc,
      f"This policy applies to all clinical staff — doctors, nurses, and allied health "
      f"professionals — involved in the care of patients with urinary catheters, invasive "
      f"ventilation, central venous access, or undergoing surgical procedures at {HN}.")

    h(doc, 1, "3. Policy standards")
    p(doc,
      f"{HN} adheres to standard precautions and transmission-based precautions in all "
      "patient care. Four HAI-prevention care bundles are developed, implemented and "
      "monitored: (1) CAUTI prevention bundle for all patients with indwelling urinary "
      "catheters; (2) VAP prevention bundle for all invasively ventilated patients; "
      "(3) CLABSI prevention bundle for all patients with central venous access; and "
      "(4) SSI prevention bundle for all patients undergoing surgical procedures, "
      "covering pre-operative, intra-operative and post-operative measures. No urinary "
      "catheter insertion, central line insertion, start of invasive ventilation, or "
      "surgical incision is made unless the relevant bundle is in place for that patient.")
    p(doc, "Staff follow the written guidance below and keep the records it requires.")

    h(doc, 1, "4. Non-negotiable rules")
    lb(doc,
       "Do not insert an indwelling urinary catheter unless the CAUTI-prevention "
       "bundle is in place for that patient.")
    lb(doc,
       "Do not start invasive ventilation unless the VAP-prevention bundle is in "
       "place for that patient.")
    lb(doc,
       "Do not insert a central venous line unless the CLABSI-prevention bundle "
       "is in place for that patient.")
    lb(doc,
       "Do not make a surgical incision unless the SSI-prevention bundle items "
       "required before knife-to-skin are in place for that patient.")
    lb(doc,
       "Adhere to standard precautions and transmission-based precautions at all "
       "times in patient care.")

    h(doc, 1, "5. What we do")

    h(doc, 2, "5.1 Prevent catheter-associated urinary tract infections (CAUTI)")
    p(doc,
      "Standard precautions and transmission-based precautions are adhered to during "
      "all catheter-related care. A care bundle to prevent CAUTI is developed, implemented "
      "and monitored. The bundle addresses insertion and maintenance practices for "
      "indwelling urinary catheters. CDC, WHO and SHEA guidelines are references for "
      "developing the CAUTI-prevention bundle.")

    h(doc, 2, "5.2 Prevent ventilator-associated pneumonia (VAP)")
    p(doc,
      "Standard precautions and transmission-based precautions are adhered to during "
      "all ventilator-related care. A care bundle to prevent VAP is developed, implemented "
      "and monitored. The bundle addresses practices for all invasively ventilated "
      "patients. CDC, WHO and SHEA guidelines are references for developing the "
      "VAP-prevention bundle.")

    h(doc, 2, "5.3 Prevent central line-associated bloodstream infections (CLABSI)")
    p(doc,
      "Standard precautions and transmission-based precautions are adhered to during "
      "all central-line care. A care bundle to prevent CLABSI is developed, implemented "
      "and monitored, including compliance monitoring. The bundle addresses insertion and "
      "maintenance practices for central venous access devices. CDC, WHO and SHEA "
      "guidelines are references for developing the CLABSI-prevention bundle.")

    h(doc, 2, "5.4 Prevent surgical site infections (SSI)")
    p(doc,
      "Standard precautions and transmission-based precautions are adhered to throughout "
      "the surgical episode. A care bundle to prevent SSI is developed, implemented and "
      "monitored. The bundle covers pre-operative, intra-operative and post-operative "
      "measures. CDC, WHO and SHEA guidelines are references for developing the "
      "SSI-prevention bundle.")

    h(doc, 1, "6. Stop-work authority")
    p(doc,
      "Do not insert an indwelling urinary catheter, insert a central line, start "
      "invasive ventilation, or make a surgical incision unless the organisation's "
      "prevention bundle for that HAI (CAUTI, CLABSI, VAP or SSI) is in place for "
      "this patient.")
    p(doc,
      "Stop-work applies to the device insertion or incision. Life-saving airway or "
      "haemorrhage control continues with the best available precautions and the bundle "
      "is completed as soon as the patient is stable.")
    p(doc,
      f"The person who stops tells the treating doctor and the Infection Prevention "
      "and Control Officer the same shift. Refusing an insertion without the bundle "
      "is not a disciplinary matter.")

    h(doc, 1, "7. Governance and responsibility")
    gov_tbl(doc, [
        ("Medical Superintendent",
         "Accountable for ensuring HAI-prevention bundles are implemented across "
         "the organisation."),
        ("Infection Prevention and Control Officer",
         "Owns implementation of this policy; monitors bundle compliance; brings HAI "
         "rates and bundle compliance data to the IPC committee; receives stop-work "
         "escalations the same shift."),
        ("Infection Prevention and Control Nurse(s)",
         "Conducts day-to-day bundle compliance monitoring; raises stop-work when a "
         "trigger is observed."),
        ("Clinical Department Heads (ICU, Surgery, Gynaecology, etc.)",
         "Ensure clinical staff in their departments implement and adhere to the "
         "relevant HAI-prevention bundles."),
        ("Nursing Superintendent",
         "Ensures nursing staff consistently implement bundle elements in all relevant "
         "patient-care areas."),
        ("Quality Coordinator",
         "Audits this policy; holds training records, staff acknowledgements, and "
         "stop-work event logs."),
    ])

    h(doc, 1, "8. Quality monitoring")
    mon_tbl(doc, [
        ("CAUTI prevention bundle",
         "Bundle documented, implemented, and compliance monitored; CAUTI rates "
         "tracked and reported."),
        ("VAP prevention bundle",
         "Bundle documented, implemented, and compliance monitored; VAP rates "
         "tracked and reported."),
        ("CLABSI prevention bundle",
         "Bundle documented, implemented, and compliance monitored; CLABSI rates "
         "tracked and reported."),
        ("SSI prevention bundle",
         "Bundle documented for pre-op, intra-op and post-op phases; compliance "
         "monitored; SSI rates tracked and reported."),
        ("Stop-work events",
         "Stop-work events logged with trigger, action taken, and outcome."),
    ])

    h(doc, 1, "9. Training and staff acknowledgement")
    p(doc,
      "All clinical staff involved in catheter insertion/care, invasive ventilation, "
      "central-line insertion/care, or surgical procedures shall be familiar with the "
      "relevant HAI-prevention bundle and the stop-work authority in this policy.")
    p(doc,
      f"I have read the Policy on Prevention of Healthcare Associated Infections in "
      f"Patients of {HN}. I will follow the processes described, including the "
      "stop-work authority in Section 6.")
    sig_tbl(doc)

    h(doc, 1, "10. Distribution")
    p(doc,
      "This policy shall be available to the IPC Officer, IPC Nurse(s), ICU staff, "
      "surgical teams, nursing staff in relevant units, department heads, and the "
      "Quality Coordinator.")

    h(doc, 1, "11. Abbreviations")
    abbrev_tbl(doc, [
        ("CAUTI", "Catheter-Associated Urinary Tract Infection"),
        ("CDC",   "Centers for Disease Control and Prevention"),
        ("CLABSI","Central Line-Associated Bloodstream Infection"),
        ("HAI",   "Healthcare Associated Infection"),
        ("ICU",   "Intensive Care Unit"),
        ("IPC",   "Infection Prevention and Control"),
        ("IPCO",  "Infection Prevention and Control Officer"),
        ("IPCN",  "Infection Prevention and Control Nurse"),
        ("NABH",  "National Accreditation Board for Hospitals and Healthcare Providers"),
        ("PPE",   "Personal Protective Equipment"),
        ("SHEA",  "Society for Healthcare Epidemiology of America"),
        ("SSI",   "Surgical Site Infection"),
        ("VAP",   "Ventilator-Associated Pneumonia"),
        ("WHO",   "World Health Organization"),
    ])

    h(doc, 1, "12. Traceability table")
    p(doc,
      "This table is an index. It is not how the policy is organised. An asterisk in "
      "the Level column means documentation of the process is required.")
    tr = tbl(doc, 5, 3)
    for ci, hdr in enumerate(("Objective Element", "Level", "Traceability to this policy")):
        tr.cell(0, ci).text = hdr
    trace_rows = [
        ("IPC.5.a", "Commitment",
         "Sections 3, 5.1 and 6 address adherence to standard and transmission-based "
         "precautions and the development, implementation and monitoring of the CAUTI "
         "prevention bundle; stop-work trigger at urinary catheter insertion."),
        ("IPC.5.b", "Commitment",
         "Sections 3, 5.2 and 6 address adherence to standard and transmission-based "
         "precautions and the development, implementation and monitoring of the VAP "
         "prevention bundle; stop-work trigger at start of invasive ventilation."),
        ("IPC.5.c", "Commitment",
         "Sections 3, 5.3 and 6 address adherence to standard and transmission-based "
         "precautions and the development, implementation and monitoring (including "
         "compliance monitoring) of the CLABSI prevention bundle; stop-work trigger "
         "at central-line insertion."),
        ("IPC.5.d", "Commitment",
         "Sections 3, 5.4 and 6 address adherence to standard and transmission-based "
         "precautions and the development, implementation and monitoring of the SSI "
         "prevention bundle covering pre-op, intra-op and post-op measures; stop-work "
         "trigger at surgical incision."),
    ]
    for ri, (oe, lvl, txt) in enumerate(trace_rows, 1):
        tr.cell(ri, 0).text = oe
        tr.cell(ri, 1).text = lvl
        tr.cell(ri, 2).text = txt

    h(doc, 1, "13. Required Records/Evidence Checklist")

    h(doc, 2, "CAUTI prevention bundle — IPC.5.a")
    lb(doc, "Documented CAUTI-prevention care bundle.")
    lb(doc, "Compliance monitoring records for the CAUTI bundle.")
    lb(doc, "CAUTI rate tracking records as part of the IPC surveillance programme.")

    h(doc, 2, "VAP prevention bundle — IPC.5.b")
    lb(doc, "Documented VAP-prevention care bundle.")
    lb(doc, "Compliance monitoring records for the VAP bundle.")
    lb(doc, "VAP rate tracking records as part of the IPC surveillance programme.")

    h(doc, 2, "CLABSI prevention bundle — IPC.5.c")
    lb(doc, "Documented CLABSI-prevention care bundle.")
    lb(doc,
       "Compliance monitoring records for the CLABSI bundle, including insertion "
       "and maintenance compliance.")
    lb(doc, "CLABSI rate tracking records as part of the IPC surveillance programme.")

    h(doc, 2, "SSI prevention bundle — IPC.5.d")
    lb(doc,
       "Documented SSI-prevention care bundle covering pre-operative, intra-operative "
       "and post-operative phases.")
    lb(doc, "Compliance monitoring records for the SSI bundle.")
    lb(doc, "SSI rate tracking records as part of the IPC surveillance programme.")

    h(doc, 2, "Stop-work events")
    lb(doc,
       "Stop-work event log with trigger (CAUTI/VAP/CLABSI/SSI), action taken, "
       "and outcome.")

    h(doc, 1, "14. References")
    ln(doc,
       "National Accreditation Board for Hospitals and Healthcare Providers. NABH "
       "Accreditation Standards for Hospitals, 6th Edition. IPC.5.")
    ln(doc, "Guidebook interpretation supplied for IPC.5.a through IPC.5.d.")
    ln(doc,
       "CDC/WHO/SHEA guidelines for prevention of healthcare-associated infections "
       "(referenced for the CAUTI, VAP, CLABSI and SSI bundles).")
    ln(doc,
       f"Internal documents of {HN}: CAUTI, VAP, CLABSI and SSI bundle documents; "
       "compliance monitoring records; HAI rate reports.")

    h(doc, 1, "Disclaimer")
    p(doc,
      "This policy reorganises the supplied IPC.5 objective-element wording and "
      "Guidebook interpretation into plain-language policy format. The modal strength "
      "of the source has been preserved. Optional examples and mechanisms have not been "
      "converted into mandatory requirements. The four care bundles — CAUTI, VAP, "
      "CLABSI, and SSI — are each separately mandatory (develop, implement, and monitor). "
      "The SSI bundle is explicitly required to cover pre-operative, intra-operative and "
      "post-operative phases. IPC.5 carries stop-work authority as stated in Section 6: "
      "no device insertion or surgical incision unless the relevant bundle is in place.")

    save_and_verify(doc, "HCO_IPC_5_v2_REWRITE_DRAFT.docx")


# ══════════════════════════════════════════════════════════════════════════════
# IPC.6 — Infection Prevention and Control Surveillance   (NO stop-work)
# COREs: a, d, f | Stars: h* | Achievement: e | Excellence: none
# Prepared by: Infection Prevention and Control Officer | Doc: HCO/IPC/POL/06
# ══════════════════════════════════════════════════════════════════════════════
def gen_ipc6():
    doc = Document()

    h(doc, 0, "Policy on Infection Prevention and Control Surveillance")
    p(doc, HN)

    h(doc, 1, "Document control")
    doc_ctrl(doc, "HCO/IPC/POL/06", "Infection Prevention and Control Officer")
    p(doc, "A blank marked ________ must be completed before issue.")

    h(doc, 1, "Statement of intent")
    p(doc,
      f"{HN}'s IPC surveillance programme tracks infection risks, rates and trends "
      "using a mix of active and passive surveillance; directs surveillance at identified "
      "high-risk activities; monitors hand-hygiene compliance at least monthly; captures "
      "MDROs; monitors housekeeping effectiveness; provides regular feedback to healthcare "
      "providers; identifies and responds to outbreaks; and analyses surveillance data "
      "for corrective and preventive action.")

    h(doc, 1, "1. Purpose")
    p(doc,
      f"This policy explains how {HN} performs IPC surveillance to capture and monitor "
      "infection risks, rates and trends, verify data, direct surveillance at high-risk "
      "activities and areas, monitor hand-hygiene compliance, capture multi-drug-resistant "
      "organisms, monitor housekeeping effectiveness, feed back surveillance data to "
      "healthcare providers, identify and respond to outbreaks, and take corrective and "
      "preventive action based on surveillance findings.")
    p(doc,
      "The clinical IPC processes, HAI-prevention bundles, and IPC programme governance "
      "are covered in separate hospital policies.")

    h(doc, 1, "2. Scope")
    p(doc,
      f"This policy applies to the IPC Officer, IPC Nurse(s), link nurses, the IPC "
      f"committee, and all clinical staff whose practice is assessed through IPC "
      f"surveillance activities at {HN}.")

    h(doc, 1, "3. Policy standards")
    p(doc,
      f"{HN} uses a judicious mix of active and passive surveillance to track infection "
      "risks, rates and trends on an ongoing basis, calculating risk-adjusted rates (not "
      "only numerator counts), and takes suitable steps based on analysis. Surveillance "
      "data is authenticated by the IPC team. Periodic surveillance is conducted in all "
      "identified high-risk activities and areas. Hand-hygiene compliance is monitored at "
      "least monthly across all categories of direct-patient-care staff, with results "
      "shared with relevant staff. The occurrence and spread of MDROs is monitored. The "
      "effectiveness of housekeeping is monitored regularly using outcome parameters — not "
      "only frequency verification; routine microbiological environmental sampling is not "
      "used as a monitoring method; periodic fumigation is not conducted; fogging is used "
      "only in the three specified conditions. Surveillance feedback is provided regularly "
      "to the appropriate healthcare provider, covering adherence rates, HAI rates, trends "
      "and improvement opportunities. Baseline rates are established; written outbreak "
      "guidance is implemented. The IPC committee analyses data and takes CAPA where "
      "necessary.")
    p(doc, "Staff follow the written guidance below and keep the records it requires.")

    h(doc, 1, "4. Non-negotiable rules")
    lb(doc,
       "Do not conduct surveillance that captures only numerator counts — calculate "
       "risk-adjusted rates using both numerator and denominator.")
    lb(doc,
       "Do not monitor hand-hygiene compliance less often than once a month, and do "
       "not exclude any category of direct-patient-care staff from monitoring.")
    lb(doc,
       "Do not use routine microbiological environmental sampling of OT and high-risk "
       "areas as a monitoring method.")
    lb(doc,
       "Do not conduct periodic fumigation.")
    lb(doc,
       "Do not fog except in the three specified conditions: commissioning a new OT; "
       "modification or renovation of an OT; significant increase in SSI with evidence "
       "of environmental involvement.")
    lb(doc,
       "Do not withhold surveillance feedback from the appropriate healthcare provider.")
    lb(doc,
       "Do not fail to establish baseline rates, and do not operate without written "
       "guidance for handling outbreaks.")

    h(doc, 1, "5. What we do")

    h(doc, 2, "5.1 Track infection risks, rates and trends")
    p(doc,
      f"{HN} uses a judicious mix of active and passive surveillance. Surveillance data "
      "is collected on an ongoing basis at regular intervals and consolidated into reports. "
      "The organisation takes suitable steps based on analysis of this data. Surveillance "
      "goes beyond simple numerator counts — risk-adjusted rates are calculated using both "
      "numerator (infected patients) and denominator (all at-risk patients) data.")

    h(doc, 2, "5.2 Verify surveillance data")
    p(doc,
      "Surveillance data is authenticated by the IPC team by reviewing every datum or "
      "by using random sampling. Where the IPC team collects all surveillance data "
      "itself, separate verification by a second party is not required.")

    h(doc, 2, "5.3 Direct surveillance at identified high-risk activities")
    p(doc,
      f"{HN} conducts periodic surveillance in each identified high-risk activity and "
      "high-risk area. The frequency and mode of surveillance are defined for each. "
      "The surveillance system adheres to national/international guidelines. Surveillance "
      "includes areas where demolition, construction or repairs are undertaken.")

    h(doc, 2, "5.4 Monitor hand-hygiene compliance")
    p(doc,
      "Hand-hygiene compliance is monitored at a minimum once every month. An "
      "appropriate sample size is chosen and all categories of staff involved in direct "
      "patient care are monitored — no category is exempt. Compliance levels are shared "
      "with the relevant staff. The WHO Observation Form or an equivalent validated "
      "observation tool may be used.")

    h(doc, 2, "5.5 Capture multi-drug-resistant organisms")
    p(doc,
      f"{HN} monitors the occurrence and spread of multi-drug-resistant organisms "
      "(MDROs). Examples of organisms monitored include methicillin-resistant "
      "Staphylococcus aureus (MRSA), multi-drug-resistant gram-negative bacteria, and "
      "Vancomycin-resistant enterococci (VRE). The organisation's own antibiogram data "
      "may identify additional MDROs to track.")

    h(doc, 2, "5.6 Monitor the effectiveness of housekeeping")
    p(doc,
      "The effectiveness of housekeeping services is monitored regularly. The "
      f"organisation defines the periodicity of this monitoring. This applies even when "
      "housekeeping is outsourced. Effectiveness is captured through outcome parameters "
      "— such as cleanliness of premises, patient feedback, and staff feedback — not "
      "only by verifying that housekeeping was done at the defined frequency. "
      "IMPORTANT: Routine microbiological environmental sampling of OT and other "
      "high-risk areas is not used as a routine monitoring method. Periodic fumigation "
      f"is not conducted. Fogging is used only when: commissioning a new OT; following "
      "modification or renovation of an OT; or when there is a significant increase in "
      "SSI with evidence of environmental involvement.")

    h(doc, 2, "5.7 Feed surveillance data back to healthcare providers")
    p(doc,
      "Surveillance findings are fed back regularly to the appropriate healthcare "
      "provider — reaching unit in-charges, treating doctors, and others who can act "
      "on the data, not only the IPC office. Feedback includes adherence rates, HAI "
      "rates, trends, and opportunities for improvement, including data from other "
      "surveillance activities. Specific inputs to reduce HAI rates may also be included. "
      "Feedback may be in the form of a bulletin or newsletter.")

    h(doc, 2, "5.8 Identify and respond to outbreaks")
    p(doc,
      f"{HN} establishes baseline infection rates to define what constitutes an "
      "outbreak. Written guidance for identifying and handling outbreaks is implemented. "
      "The surveillance system is designed to enable early identification of outbreaks "
      "by detecting a rate above the established baseline. When an outbreak is identified, "
      "appropriate actions are taken to control it; corrective actions are implemented "
      "to prevent recurrence.")

    h(doc, 2, "5.9 Analyse surveillance data and take corrective and preventive action")
    p(doc,
      "The IPC committee analyses surveillance data at each meeting. Corrective and "
      "preventive actions are taken where the analysis identifies the need. This "
      "includes actions following outbreaks. The IPC Officer tracks open CAPA items "
      "to closure.")

    h(doc, 1, "6. Governance and responsibility")
    gov_tbl(doc, [
        ("Medical Superintendent",
         "Accountable for ensuring the IPC surveillance programme is resourced and "
         "its findings acted upon."),
        ("Infection Prevention and Control Officer",
         "Owns the IPC surveillance programme; ensures ongoing data collection, "
         "analysis and reporting; leads outbreak response; tracks CAPA to closure; "
         "brings surveillance data to the IPC committee."),
        ("Infection Prevention and Control Nurse(s)",
         "Conducts active surveillance, hand-hygiene compliance monitoring, and "
         "housekeeping effectiveness monitoring; collates data for the IPC Officer."),
        ("IPC Committee",
         "Analyses surveillance data at each monthly meeting; decides on CAPA; "
         "reviews outbreak response and corrective actions."),
        ("Department Heads",
         "Ensure compliance with corrective actions arising from surveillance findings "
         "in their departments; disseminate feedback to relevant staff."),
        ("Quality Coordinator",
         "Audits this policy; holds training records and staff acknowledgements."),
    ])

    h(doc, 1, "7. Quality monitoring")
    mon_tbl(doc, [
        ("Surveillance scope and rates",
         "Mix of active and passive surveillance in use; risk-adjusted rates "
         "calculated and reported; suitable steps taken on findings."),
        ("Data verification",
         "Authentication by IPC team documented (or nil-verification note where "
         "IPC team collected all data)."),
        ("High-risk area surveillance",
         "Periodic surveillance records for each identified high-risk area and "
         "activity, with defined frequency and mode."),
        ("Hand-hygiene compliance",
         "Monthly monitoring records covering all direct-patient-care staff categories; "
         "compliance levels shared with relevant staff."),
        ("MDRO monitoring",
         "MDRO occurrence and spread tracked; organism log maintained."),
        ("Housekeeping effectiveness",
         "Effectiveness monitored using outcome parameters; no routine environmental "
         "sampling; no periodic fumigation; fogging only in three specified conditions."),
        ("Feedback to providers",
         "Regular feedback including adherence rates, HAI rates, trends and improvement "
         "opportunities reaching the appropriate healthcare provider."),
        ("Outbreak identification and response",
         "Baseline rates established; written outbreak guidance implemented; outbreak "
         "response records and post-outbreak CAPA."),
        ("CAPA",
         "IPC committee CAPA records with owners, due dates, and closure evidence."),
    ])

    h(doc, 1, "8. Training and staff acknowledgement")
    p(doc,
      "The IPC Officer, IPC Nurse(s), and link nurses shall be familiar with the "
      "surveillance methods, monitoring tools, outbreak definition and response process, "
      "and the data feedback and CAPA cycle described in this policy.")
    p(doc,
      f"I have read the Policy on Infection Prevention and Control Surveillance of "
      f"{HN}. I will follow the processes described.")
    sig_tbl(doc)

    h(doc, 1, "9. Distribution")
    p(doc,
      "This policy shall be available to the IPC Officer, IPC Nurse(s), the IPC "
      "committee, department heads, and the Quality Coordinator.")

    h(doc, 1, "10. Abbreviations")
    abbrev_tbl(doc, [
        ("CAPA",  "Corrective and Preventive Action"),
        ("HAI",   "Healthcare Associated Infection"),
        ("IPC",   "Infection Prevention and Control"),
        ("IPCC",  "Infection Prevention and Control Committee"),
        ("IPCO",  "Infection Prevention and Control Officer"),
        ("IPCN",  "Infection Prevention and Control Nurse"),
        ("MDRO",  "Multi-Drug-Resistant Organism"),
        ("MRSA",  "Methicillin-Resistant Staphylococcus aureus"),
        ("NABH",  "National Accreditation Board for Hospitals and Healthcare Providers"),
        ("OT",    "Operation Theatre"),
        ("SSI",   "Surgical Site Infection"),
        ("VRE",   "Vancomycin-Resistant Enterococci"),
        ("WHO",   "World Health Organization"),
    ])

    h(doc, 1, "11. Traceability table")
    p(doc,
      "This table is an index. It is not how the policy is organised. An asterisk in "
      "the Level column means documentation of the process is required.")
    tr = tbl(doc, 10, 3)
    for ci, hdr in enumerate(("Objective Element", "Level", "Traceability to this policy")):
        tr.cell(0, ci).text = hdr
    trace_rows = [
        ("IPC.6.a", "CORE",
         "Sections 3 and 5.1 address the use of active and passive surveillance, "
         "ongoing regular data collection, suitable steps based on analysis, and "
         "calculation of risk-adjusted rates."),
        ("IPC.6.b", "Commitment",
         "Sections 3 and 5.2 address data authentication by the IPC team through "
         "full review or random sampling, with the carve-out for IPC-team-collected "
         "data."),
        ("IPC.6.c", "Commitment",
         "Sections 3 and 5.3 address periodic surveillance directed at identified "
         "high-risk activities and areas, with defined frequency and mode, including "
         "construction/demolition areas."),
        ("IPC.6.d", "CORE",
         "Sections 3 and 5.4 address monthly hand-hygiene compliance monitoring of "
         "all direct-patient-care staff categories, with results shared with relevant "
         "staff."),
        ("IPC.6.e", "Achievement",
         "Section 5.5 addresses a mechanism to monitor the occurrence and spread "
         "of MDROs."),
        ("IPC.6.f", "CORE",
         "Sections 3 and 5.6 address regular monitoring of housekeeping effectiveness "
         "using outcome parameters; the prohibitions on routine environmental sampling, "
         "periodic fumigation, and non-conditional fogging."),
        ("IPC.6.g", "Commitment",
         "Sections 3 and 5.7 address regular feedback to the appropriate healthcare "
         "provider covering adherence rates, HAI rates, trends and improvement "
         "opportunities."),
        ("IPC.6.h", "Commitment*",
         "Sections 3 and 5.8 address establishment of baseline rates, written outbreak "
         "guidance, early identification through surveillance, and post-outbreak "
         "corrective action."),
        ("IPC.6.i", "Commitment",
         "Sections 3 and 5.9 address IPC committee analysis of surveillance data, "
         "CAPA where necessary, and post-outbreak corrective actions tracked to closure."),
    ]
    for ri, (oe, lvl, txt) in enumerate(trace_rows, 1):
        tr.cell(ri, 0).text = oe
        tr.cell(ri, 1).text = lvl
        tr.cell(ri, 2).text = txt

    h(doc, 1, "12. Required Records/Evidence Checklist")

    h(doc, 2, "Infection risk, rate and trend tracking — IPC.6.a (CORE)")
    lb(doc,
       "Surveillance reports showing ongoing data collection at regular intervals.")
    lb(doc,
       "Evidence that risk-adjusted rates (not only numerator counts) are calculated.")
    lb(doc,
       "Records of suitable steps (actions) taken in response to analysis findings.")

    h(doc, 2, "Data verification — IPC.6.b")
    lb(doc,
       "Authentication records (sample-based or full review by IPC team), or a "
       "documented note that the IPC team collected all data (verification not required).")

    h(doc, 2, "High-risk area surveillance — IPC.6.c")
    lb(doc,
       "Surveillance records for each identified high-risk area and activity, including "
       "construction/demolition areas.")
    lb(doc,
       "Defined frequency and mode of surveillance documented per high-risk area.")

    h(doc, 2, "Hand-hygiene compliance monitoring — IPC.6.d (CORE)")
    lb(doc,
       "Monthly compliance monitoring records covering all categories of direct-patient-"
       "care staff.")
    lb(doc,
       "Evidence that compliance levels are shared with relevant staff (e.g., bulletin, "
       "feedback sheet, department communication).")

    h(doc, 2, "MDRO monitoring — IPC.6.e (Achievement)")
    lb(doc,
       "MDRO occurrence and spread log, with records of organisms detected and "
       "containment measures taken.")
    lb(doc,
       "Evidence the MDRO capture mechanism is operational (microbiologist alert to "
       "IPC team, or equivalent).")

    h(doc, 2, "Housekeeping effectiveness monitoring — IPC.6.f (CORE)")
    lb(doc,
       "Defined outcome parameters for housekeeping effectiveness monitoring.")
    lb(doc,
       "Regular monitoring records using those outcome parameters.")
    lb(doc,
       "Record confirming no routine microbiological environmental sampling is used "
       "as a monitoring method.")
    lb(doc,
       "Records confirming fogging was used only under the three permitted conditions "
       "(or a nil record if no fogging has occurred).")

    h(doc, 2, "Feedback to healthcare providers — IPC.6.g")
    lb(doc,
       "Periodic feedback documents (bulletin, report, or equivalent) that include "
       "adherence rates, HAI rates, trends, and improvement opportunities.")
    lb(doc,
       "Evidence that feedback reached the appropriate healthcare provider — unit "
       "in-charges, treating doctors, or department heads.")

    h(doc, 2, "Outbreak identification and response — IPC.6.h (Commitment*)")
    lb(doc, "Baseline infection rates established and documented.")
    lb(doc, "Written outbreak definition using those baseline rates.")
    lb(doc, "Written outbreak-handling guidance.")
    lb(doc,
       "Outbreak response records (or a drill record if no real event has occurred), "
       "including post-outbreak corrective action.")

    h(doc, 2, "Surveillance data CAPA — IPC.6.i")
    lb(doc,
       "IPC committee minutes showing surveillance data analysis and CAPA decisions.")
    lb(doc,
       "Open-action tracking records showing owners, due dates, and evidence of closure.")

    h(doc, 1, "13. References")
    ln(doc,
       "National Accreditation Board for Hospitals and Healthcare Providers. NABH "
       "Accreditation Standards for Hospitals, 6th Edition. IPC.6.")
    ln(doc, "Guidebook interpretation supplied for IPC.6.a through IPC.6.i.")
    ln(doc,
       "WHO Multimodal Hand Hygiene Improvement Strategy; WHO Observation Form for "
       "Hand Hygiene Compliance.")
    ln(doc,
       f"Internal documents of {HN}: IPC surveillance reports; hand-hygiene compliance "
       "records; MDRO log; housekeeping effectiveness records; outbreak baseline data "
       "and outbreak-response guidance; IPC committee minutes with CAPA.")

    h(doc, 1, "Disclaimer")
    p(doc,
      "This policy reorganises the supplied IPC.6 objective-element wording and "
      "Guidebook interpretation into plain-language policy format. The modal strength "
      "of the source has been preserved. Optional examples and mechanisms have not been "
      "converted into mandatory requirements. The minimum monthly hand-hygiene compliance "
      "monitoring frequency, the requirement to cover all direct-patient-care staff "
      "categories, the prohibition on routine microbiological environmental sampling, "
      "the prohibition on periodic fumigation, the three specific conditions permitting "
      "fogging, and the IPC team verification carve-out have all been retained verbatim. "
      "IPC.6 has no stop-work section.")

    save_and_verify(doc, "HCO_IPC_6_v2_REWRITE_DRAFT.docx")


# ══════════════════════════════════════════════════════════════════════════════
# IPC.7 — Sterilisation and Disinfection of Instruments, Equipment and Devices
#          (HAS stop-work: Section 6)
# COREs: b | Stars: b*, c*, d*, e* | Achievement: none
# Prepared by: CSSD In-Charge | Doc: HCO/IPC/POL/07
# Stop-work trigger: issue from CSSD / sterile store when load validation failed
#   or a recall is in effect
# ══════════════════════════════════════════════════════════════════════════════
def gen_ipc7():
    doc = Document()

    h(doc, 0, "Policy on Sterilisation and Disinfection of Instruments, Equipment and Devices")
    p(doc, HN)

    h(doc, 1, "Document control")
    doc_ctrl(doc, "HCO/IPC/POL/07", "CSSD In-Charge")
    p(doc, "A blank marked ________ must be completed before issue.")

    h(doc, 1, "Statement of intent")
    p(doc,
      f"{HN} provides adequate space and appropriate zoning in the CSSD, follows written "
      "guidance aligned with national/international standards for the full cleaning-packing-"
      "sterilisation-storage-issue cycle, conducts daily physical/chemical and at least "
      "weekly biological validation tests, implements a documented recall procedure on any "
      "sterilisation breakdown, and does not issue from a failed or recalled load.")

    h(doc, 1, "1. Purpose")
    p(doc,
      f"This policy explains how {HN} implements infection prevention through sterilisation "
      "and disinfection of instruments, equipment and devices: providing adequate CSSD space "
      "and zoning; following written guidance for every step of the cleaning-packing-"
      "sterilisation-storage-issue cycle; managing reprocessing of identified re-usable "
      "items; carrying out and documenting regular validation tests; and implementing the "
      "recall procedure when a sterilisation breakdown is identified.")
    p(doc,
      "IPC in clinical areas, support services, HAI-prevention bundles, and IPC surveillance "
      "are covered in separate hospital policies.")

    h(doc, 1, "2. Scope")
    p(doc,
      f"This policy applies to CSSD staff, nursing staff who handle sterile instruments "
      f"and equipment in clinical areas, and all staff responsible for sterilisation, "
      f"disinfection, storage and issue of instruments, equipment and devices at {HN}.")

    h(doc, 1, "3. Policy standards")
    p(doc,
      f"{HN}'s CSSD has a suitable location, unidirectional flow, and defined zoning with "
      "separation of clean and dirty areas and sufficient space. The full cycle — cleaning, "
      "packing, disinfection/sterilisation, storage, and issue — is performed as per written "
      "guidance that is in consonance with national/international guidelines. Cleaning removes "
      "visible biological material and dirt. Spaulding's classification guides the level of "
      "disinfection. Shelf life is defined by packing material and mode of sterilisation. "
      "Flash sterilisation is used only in exceptional situations. Sterilised items are stored "
      "appropriately across the organisation. Instruments identified for re-use are listed; "
      "the number of re-uses and reprocessing method are defined, monitored, and patients are "
      "informed. Physical/chemical tests are performed daily and biological tests at least "
      "weekly; Bowie-Dick and leak rate tests are carried out; each load has a unique number. "
      "Written recall guidance exists; the recall procedure is implemented on any breakdown; "
      "a mock recall is run at least annually.")
    p(doc, "Staff follow the written guidance below and keep the records it requires.")

    h(doc, 1, "4. Non-negotiable rules")
    lb(doc,
       "Do not issue an instrument, device or pack from sterile storage when validation "
       "of that sterilisation load has failed, or when a recall of that load or machine "
       "is in effect.")
    lb(doc,
       "Do not perform flash sterilisation routinely — use it only in exceptional situations "
       "when there is insufficient time to sterilise by the preferred method.")
    lb(doc,
       "Do not skip daily physical/chemical tests or weekly biological tests for any "
       "sterilisation cycle.")
    lb(doc,
       "Do not issue sterilised items without confirming the load number, content, and "
       "expiry date.")
    lb(doc,
       "Do not reprocess an item beyond its defined number of re-uses, and do not "
       "reprocess single-use items unless a documented, accepted method and risk "
       "assessment exist.")
    lb(doc,
       "Do not handle used instruments without adequate infection prevention measures "
       "in place, whether cleaning is done in CSSD or in a patient-care area.")

    h(doc, 1, "5. What we do")

    h(doc, 2, "5.1 Provide adequate space and appropriate zoning")
    p(doc,
      f"{HN}'s CSSD (or designated sterilisation area) has a suitable location with a "
      "proper layout providing unidirectional flow and zoning. Clean and dirty areas are "
      "separated. Sufficient space is available to ensure sterilisation activities can "
      "be performed properly. The layout aims to have separate areas for receiving, "
      "washing, cleaning, packing, sterilisation, sterile storage and issue; this is the "
      "optimal arrangement. The CSSD In-Charge and the IPC Officer walk the zoning "
      "regularly.")

    h(doc, 2, "5.2 Follow written guidance for cleaning, packing, disinfection/sterilisation, "
      "storage and issue")
    p(doc,
      "The full cycle is performed as per written guidance that is in consonance with "
      "national/international guidelines. The CDC Guideline for Disinfection and "
      "Sterilization in Healthcare Facilities and HISI guidelines are reference sources. "
      "Cleaning of used instruments/equipment/devices removes visible biological material "
      "and dirt; it is preferably done in the CSSD, but if done in patient-care areas, "
      "adequate IPC measures are in place. After cleaning, sets are prepared and packed "
      "using appropriate material. Spaulding's classification guides the decision on high/"
      "intermediate/low-level disinfection. Disinfection/sterilisation is performed as per "
      "written guidance. Flash sterilisation is used only in exceptional situations when "
      "there is insufficient time to sterilise by the preferred method. Sterilised/"
      "disinfected equipment and sets are stored appropriately across the organisation — "
      "not only in CSSD. Shelf life and expiry date of sterilised items are guided by the "
      "packing material used and the mode of sterilisation.")

    h(doc, 2, "5.3 Manage reprocessing of instruments, equipment and devices")
    p(doc,
      f"{HN} identifies instruments, equipment and devices that are suitable for re-use. "
      "The number of re-uses and the reprocessing method for each identified item are "
      "defined and monitored. Patients are informed about the re-use of items. Written "
      "guidance addresses cleaning, disinfection or sterilisation between patients for "
      "each reprocessed item. The written guidance is in consonance with available good "
      "practices. The CSSD In-Charge maintains the re-use register.")

    h(doc, 2, "5.4 Carry out and document validation tests for sterilisation")
    p(doc,
      "Validation tests for sterilisation are carried out by accepted methods. "
      "Physical/chemical indicator tests are performed daily. Biological indicator tests "
      "are performed at least weekly. Engineering validations — including the Bowie-Dick "
      "tape test and leak rate test — are carried out. Each sterilisation load has a unique "
      "load number and content description. Where applicable, temperature, pressure and "
      "time-record charts are maintained. The CDC Guideline for Disinfection and "
      "Sterilization in Healthcare Facilities 2008 is a reference for validation methods. "
      "A load that fails any validation test is not issued.")

    h(doc, 2, "5.5 Implement the recall procedure on sterilisation breakdown")
    p(doc,
      "The sterilisation procedure is regularly monitored. When a breakdown in the "
      "sterilisation system is identified, the written recall procedure is implemented: "
      "issue from the affected load/machine is stopped, items from that load (traced by "
      "unique load number, machine number, and date) are retrieved from all areas, users "
      "are notified, and items are reprocessed or discarded. The organisation uses a "
      "batch-processing system with date and machine number to enable effective recall. "
      "The recall procedure is verified through a mock drill at least annually.")

    h(doc, 1, "6. Stop-work authority")
    p(doc,
      "Do not issue an instrument, device or pack from sterile storage when validation "
      "of that sterilisation load has failed, or when a recall of that load or machine "
      "is in effect.")
    p(doc,
      "Stop-work applies to issue from CSSD / sterile store. Immediate life-saving use "
      "of the only available item is documented and the CSSD In-Charge is told the "
      "same shift.")
    p(doc,
      f"The person who stops tells the CSSD In-Charge and the Infection Prevention and "
      "Control Officer the same shift. Refusing to issue a failed or recalled load is "
      "not a disciplinary matter.")

    h(doc, 1, "7. Governance and responsibility")
    gov_tbl(doc, [
        ("Medical Superintendent",
         "Accountable for ensuring the CSSD is adequately resourced, staffed, and "
         "operating as required by this policy."),
        ("CSSD In-Charge",
         "Owns day-to-day implementation of this policy; maintains the processing "
         "guidance, re-use register, and load records; approves issue; leads recalls; "
         "receives stop-work escalations and notifies the IPC Officer the same shift."),
        ("Infection Prevention and Control Officer",
         "Monitors CSSD processes in IPC rounds; receives stop-work escalations; "
         "brings sterilisation audit findings to the IPC committee."),
        ("Nursing Superintendent",
         "Ensures nursing staff in clinical areas check expiry dates and load numbers "
         "of sterile items before use, and escalate any doubt about sterility."),
        ("Quality Coordinator",
         "Audits this policy; holds training records, staff acknowledgements, and "
         "stop-work event logs."),
    ])

    h(doc, 1, "8. Quality monitoring")
    mon_tbl(doc, [
        ("CSSD space and zoning",
         "Regular IPC/CSSD round records confirming adequate space, unidirectional "
         "flow, and clean/dirty separation."),
        ("Processing guidance compliance",
         "Load records confirming cleaning, packing, sterilisation, storage and issue "
         "per written guidance; Spaulding classification applied correctly."),
        ("Re-use register",
         "Re-use register updated; patient information records for re-used items on "
         "file; no item issued beyond defined re-use count."),
        ("Validation tests",
         "Daily physical/chemical test records; weekly biological test records; "
         "Bowie-Dick and leak rate test records; all loads with unique number and "
         "content description."),
        ("Recall procedure",
         "Written recall procedure on file; annual mock recall record; records of any "
         "actual recalls with tracing and retrieval evidence."),
        ("Flash sterilisation",
         "Flash sterilisation log confirming use only in exceptional documented "
         "situations."),
        ("Stop-work events",
         "Stop-work events logged with trigger, action taken, and outcome."),
    ])

    h(doc, 1, "9. Training and staff acknowledgement")
    p(doc,
      "CSSD staff and nursing staff who handle sterile instruments shall be familiar "
      "with the cleaning-packing-sterilisation-storage-issue cycle, the Spaulding "
      "classification, the validation test schedule, the re-use register and patient "
      "information requirement, the recall procedure, and the stop-work authority in "
      "this policy.")
    p(doc,
      f"I have read the Policy on Sterilisation and Disinfection of Instruments, "
      f"Equipment and Devices of {HN}. I will follow the processes described, including "
      "the stop-work authority in Section 6.")
    sig_tbl(doc)

    h(doc, 1, "10. Distribution")
    p(doc,
      "This policy shall be available to CSSD staff, nursing staff, the IPC Officer, "
      "department heads, and the Quality Coordinator.")

    h(doc, 1, "11. Abbreviations")
    abbrev_tbl(doc, [
        ("CAPA",  "Corrective and Preventive Action"),
        ("CDC",   "Centers for Disease Control and Prevention"),
        ("CSSD",  "Central Sterile Services Department"),
        ("HISI",  "Hospital Infection Society India"),
        ("IPC",   "Infection Prevention and Control"),
        ("IPCO",  "Infection Prevention and Control Officer"),
        ("NABH",  "National Accreditation Board for Hospitals and Healthcare Providers"),
        ("OT",    "Operation Theatre"),
        ("PEP",   "Post-Exposure Prophylaxis"),
        ("PPE",   "Personal Protective Equipment"),
        ("SSI",   "Surgical Site Infection"),
    ])

    h(doc, 1, "12. Traceability table")
    p(doc,
      "This table is an index. It is not how the policy is organised. An asterisk in "
      "the Level column means documentation of the process is required.")
    tr = tbl(doc, 6, 3)
    for ci, hdr in enumerate(("Objective Element", "Level", "Traceability to this policy")):
        tr.cell(0, ci).text = hdr
    trace_rows = [
        ("IPC.7.a", "Commitment",
         "Sections 3 and 5.1 address adequate CSSD space, suitable location, "
         "unidirectional flow and zoning, and clean/dirty area separation."),
        ("IPC.7.b", "CORE*",
         "Sections 3, 5.2 and 6 address written guidance for the full cleaning-packing-"
         "sterilisation-storage-issue cycle, in consonance with national/international "
         "guidelines; Spaulding classification; flash sterilisation restriction; shelf "
         "life and expiry; appropriate storage; stop-work trigger at issue of a failed "
         "or recalled load."),
        ("IPC.7.c", "Commitment*",
         "Sections 3 and 5.3 address identification of re-usable items, defined re-use "
         "count and reprocessing method, patient information, and written guidance "
         "for cleaning/disinfection/sterilisation between patients."),
        ("IPC.7.d", "Commitment*",
         "Sections 3 and 5.4 address daily physical/chemical tests, at least weekly "
         "biological tests, Bowie-Dick and leak rate tests, unique load numbers, and "
         "temperature/pressure/time charts."),
        ("IPC.7.e", "Commitment*",
         "Sections 3, 5.5 and 6 address written recall procedure, implementation on "
         "any breakdown, batch tracing by load/machine/date, and at least annual mock "
         "recall; stop-work trigger when a recall is in effect."),
    ]
    for ri, (oe, lvl, txt) in enumerate(trace_rows, 1):
        tr.cell(ri, 0).text = oe
        tr.cell(ri, 1).text = lvl
        tr.cell(ri, 2).text = txt

    h(doc, 1, "13. Required Records/Evidence Checklist")

    h(doc, 2, "CSSD space and zoning — IPC.7.a")
    lb(doc,
       "CSSD layout diagram or floor plan showing unidirectional flow and clean/dirty "
       "area zoning.")
    lb(doc,
       "Regular CSSD/IPC round records confirming adequate space and clean/dirty "
       "separation.")

    h(doc, 2, "Cleaning-packing-sterilisation-storage-issue guidance — IPC.7.b (CORE*)")
    lb(doc,
       "Written processing guidance covering all steps of the cycle, in consonance "
       "with named national/international guidelines.")
    lb(doc,
       "Spaulding classification decision table in the guidance.")
    lb(doc,
       "Flash sterilisation log confirming use only in documented exceptional situations.")
    lb(doc,
       "Shelf-life and expiry-date record per packing material and sterilisation mode.")
    lb(doc,
       "Load records showing sterilised items stored appropriately (not only in CSSD).")
    lb(doc,
       "Stop-work event log for any issue-from-failed-load events.")

    h(doc, 2, "Re-use register and patient information — IPC.7.c (Commitment*)")
    lb(doc,
       "Re-use register listing identified re-usable items with defined re-use count "
       "and reprocessing method for each.")
    lb(doc,
       "Patient information records confirming patients were informed about re-use.")
    lb(doc,
       "Compliance monitoring records confirming no item is issued beyond its defined "
       "re-use count.")

    h(doc, 2, "Validation test records — IPC.7.d (Commitment*)")
    lb(doc,
       "Daily physical/chemical indicator records for each sterilisation load.")
    lb(doc,
       "Weekly biological indicator records.")
    lb(doc,
       "Bowie-Dick tape test and leak rate test records.")
    lb(doc,
       "Load log with unique load number, content description, and (where applicable) "
       "temperature/pressure/time chart.")

    h(doc, 2, "Recall procedure — IPC.7.e (Commitment*)")
    lb(doc,
       "Written recall procedure naming: stop-issue step, retrieval method by load/"
       "machine/date, user notification path, and reprocess-or-discard decision.")
    lb(doc,
       "Annual mock recall exercise record with findings and any CAPA.")
    lb(doc,
       "Records of any actual recalls implemented, with tracing and retrieval evidence.")

    h(doc, 1, "14. References")
    ln(doc,
       "National Accreditation Board for Hospitals and Healthcare Providers. NABH "
       "Accreditation Standards for Hospitals, 6th Edition. IPC.7.")
    ln(doc, "Guidebook interpretation supplied for IPC.7.a through IPC.7.e.")
    ln(doc,
       "CDC Guideline for Disinfection and Sterilization in Healthcare Facilities, 2008. "
       "Centers for Disease Control and Prevention.")
    ln(doc,
       "Hospital Infection Society India (HISI). Guidelines for Central Sterile "
       "Services Department.")
    ln(doc,
       f"Internal documents of {HN}: CSSD processing guidance; re-use register; "
       "load records; validation test records; recall procedure; mock recall records.")

    h(doc, 1, "Disclaimer")
    p(doc,
      "This policy reorganises the supplied IPC.7 objective-element wording and "
      "Guidebook interpretation into plain-language policy format. The modal strength "
      "of the source has been preserved. Optional examples and mechanisms have not been "
      "converted into mandatory requirements, with one exception: the annual mock recall "
      "drill is treated as mandatory in this policy (following the standard method_bodies "
      "implementation) although the Guidebook uses 'could be verified through a mock drill' "
      "(aspirational). This discrepancy is documented in the IPC.7-8 raw source dump. "
      "The prohibition on routine flash sterilisation, the daily physical/chemical and "
      "minimum weekly biological validation test frequencies, the load-unique-number "
      "requirement, the patient-information requirement for re-use, and the recall-on-"
      "breakdown obligation have been retained verbatim. IPC.7 carries stop-work authority "
      "as stated in Section 6: no issue from a failed or recalled load.")

    save_and_verify(doc, "HCO_IPC_7_v2_REWRITE_DRAFT.docx")


# ══════════════════════════════════════════════════════════════════════════════
# IPC.8 — Prevention of Healthcare Associated Infections in Staff
#          (HAS stop-work: Section 6)
# COREs: none | Stars: a*, b*, e* | Achievement: c | Excellence: none
# Prepared by: Infection Prevention and Control Officer | Doc: HCO/IPC/POL/08
# Stop-work triggers: continue duty against a work restriction;
#   skip PEP after blood/body-fluid exposure
# ══════════════════════════════════════════════════════════════════════════════
def gen_ipc8():
    doc = Document()

    h(doc, 0, "Policy on Prevention of Healthcare Associated Infections in Staff")
    p(doc, HN)

    h(doc, 1, "Document control")
    doc_ctrl(doc, "HCO/IPC/POL/08", "Infection Prevention and Control Officer")
    p(doc, "A blank marked ________ must be completed before issue.")

    h(doc, 1, "Statement of intent")
    p(doc,
      f"{HN} implements occupational health and safety practices as per written guidance, "
      "provides Hepatitis B vaccination to all direct-patient-care staff and other "
      "relevant immunisation as per risk, implements work restrictions for staff with "
      "transmissible infections, implements blood and body fluid exposure prevention "
      "measures, and provides post-exposure prophylaxis for Hepatitis B and HIV exposure "
      "the same shift.")

    h(doc, 1, "1. Purpose")
    p(doc,
      f"This policy explains how {HN} takes action to prevent or reduce healthcare "
      "associated infections in staff: implementing occupational health and safety "
      "practices as per written guidance; maintaining a staff immunisation programme "
      "with a Hepatitis B vaccination floor; implementing work restrictions for "
      "healthcare providers with transmissible infections; preventing blood and body "
      "fluid exposures; and providing post-exposure prophylaxis aligned with national/"
      "international guidelines.")
    p(doc,
      "Standard precautions, hand-hygiene guidelines, and IPC resources are covered "
      "in the hospital's other policies.")

    h(doc, 1, "2. Scope")
    p(doc,
      f"This policy applies to all staff — clinical, nursing, allied health, laboratory, "
      f"housekeeping, CSSD, and support staff — working at {HN}, and to the Occupational "
      f"Health Physician and the IPC Officer in their roles under this policy.")

    h(doc, 1, "3. Policy standards")
    p(doc,
      f"{HN} maintains written guidance on occupational health and safety practices and "
      "provides adequate resources to staff for IPC. Pre-employment health screening and "
      "capture of immunisation status are conducted where appropriate. The immunisation "
      "policy is in consonance with available evidence; at minimum, Hepatitis B vaccination "
      "is provided to all staff in direct patient care; other relevant immunisation is "
      "provided as per risk and statutory requirements. Work restrictions for healthcare "
      "providers with transmissible infections are implemented; the organisation encourages "
      "reporting of illness and does not penalise staff for reporting. Healthcare providers "
      "use appropriate PPE to prevent blood and body fluid exposures; the organisation "
      "strives to eliminate needle devices when safe alternatives exist; sharps containers "
      "are at the point of use and replaced before full. PEP for Hepatitis B and HIV "
      "exposure is provided the same shift, aligned with national/international guidelines, "
      "and documented by the IPC Nurse.")
    p(doc, "Staff follow the written guidance below and keep the records it requires.")

    h(doc, 1, "4. Non-negotiable rules")
    lb(doc,
       "Do not continue clinical duty when a work restriction for a transmissible "
       "infection applies to you.")
    lb(doc,
       "Do not leave a blood or body-fluid exposure — including a needle-stick — "
       "without starting the organisation's post-exposure path the same shift.")
    lb(doc,
       "Do not operate without Hepatitis B vaccination being available to and provided "
       "for all staff in direct patient care.")
    lb(doc,
       "Do not leave a staff member's exposure event undocumented in the IPC Nurse's "
       "PEP record file.")
    lb(doc,
       "Do not penalise any healthcare provider for reporting illness, an exposure, "
       "or observing a colleague's work restriction.")
    lb(doc,
       "Do not recap used needles — place sharps in the puncture-proof container at "
       "the point of use immediately after use.")

    h(doc, 1, "5. What we do")

    h(doc, 2, "5.1 Implement occupational health and safety practices as per written guidance")
    p(doc,
      "Written guidance on occupational health and safety practices is available and "
      f"implemented at {HN}. The guidance covers PPE use, access to vaccination, "
      "exposure reporting, and work restriction for transmissible infections. The "
      "organisation provides adequate resources to staff for IPC. Staff adhere to "
      "standard precautions as required by the hospital's IPC policies. Where "
      "appropriate, pre-employment health screening is conducted and immunisation "
      "status is captured. The Occupational Health Physician and the IPC Officer own "
      "this guidance jointly. Staff are trained at induction and at least annually.")

    h(doc, 2, "5.2 Implement a staff immunisation programme")
    p(doc,
      f"{HN}'s immunisation policy for staff is in consonance with available evidence. "
      "At minimum, Hepatitis B vaccination is provided to all staff involved in direct "
      "patient care. Other relevant immunisation is provided as per the risk from time "
      "to time and in accordance with applicable statutory requirements. The Occupational "
      "Health Physician maintains an immunisation register recording staff name, vaccine, "
      "dose, date, and booster schedule. An unvaccinated direct-care worker without a "
      "documented contraindication or documented refusal is a gap at this element.")

    h(doc, 2, "5.3 Implement work restrictions for healthcare providers with transmissible infections")
    p(doc,
      f"{HN} implements work restrictions for healthcare providers with transmissible "
      "infections. Restrictions limit the role and responsibilities of the affected "
      "staff member until they are cleared. Examples of transmissible infections that "
      "may trigger restrictions include acute conjunctivitis, chickenpox, acute "
      "respiratory infections, and smear-positive tuberculosis; the list is not "
      "exhaustive. The organisation actively encourages healthcare providers to report "
      "their illness or exposures and does not penalise them for doing so. Evidence "
      "of restrictions that were actually applied — not only a policy document — is "
      "maintained.")

    h(doc, 2, "5.4 Implement blood and body fluid exposure prevention measures")
    p(doc,
      "Healthcare providers use appropriate PPE to prevent blood and body fluid "
      f"exposures. {HN} strives to eliminate the use of needle devices whenever safe "
      "and effective alternatives are available. Sharps containers are at the point "
      "of use in all clinical areas and are replaced before they are full. A recapping "
      "culture on a unit is treated as an incident and investigated. The IPC Nurse "
      "includes sharps safety in regular IPC rounds.")

    h(doc, 2, "5.5 Provide post-exposure prophylaxis")
    p(doc,
      f"{HN} provides post-exposure prophylaxis (PEP) to all concerned staff members "
      "following blood or body-fluid exposures, including needle-stick injuries. "
      "PEP for Hepatitis B and HIV exposure is provided and aligned with national/"
      "international guidelines. Immediate first aid at the exposure site is the first "
      "step. The Occupational Health Physician or the emergency doctor starts the PEP "
      "path the same shift. The IPC Nurse maintains the PEP record file in confidence, "
      "with the staff member's name, date, type of exposure, PEP given, and follow-up "
      "schedule. Confidentiality follows the hospital's staff confidentiality policy.")

    h(doc, 1, "6. Stop-work authority")
    p(doc,
      "Do not continue clinical duty when a work restriction for a transmissible "
      "infection applies to you.")
    p(doc,
      "Do not leave a blood or body-fluid exposure (including a needle-stick) without "
      "starting the organisation's post-exposure path the same shift.")
    p(doc,
      "Stop-work applies to the restricted duty and to delay of PEP. Immediate first "
      "aid at the exposure site starts at once.")
    p(doc,
      f"The person who stops (or the colleague who sees the restriction being ignored) "
      "tells the Occupational Health Physician and the Infection Prevention and Control "
      "Officer the same shift. Reporting an exposure or observing a restriction is not "
      "a disciplinary matter.")

    h(doc, 1, "7. Governance and responsibility")
    gov_tbl(doc, [
        ("Medical Superintendent",
         "Accountable for ensuring staff occupational health and IPC protections are "
         "resourced and implemented."),
        ("Infection Prevention and Control Officer",
         "Owns implementation of this policy; monitors compliance; receives stop-work "
         "escalations the same shift; brings staff infection data to the IPC committee."),
        ("Occupational Health Physician",
         "Owns the staff immunisation programme and immunisation register; manages "
         "work restrictions; starts or coordinates the PEP path on the same shift as "
         "an exposure; holds exposure event records."),
        ("Infection Prevention and Control Nurse(s)",
         "Maintains PEP documentation file; includes sharps safety and PPE compliance "
         "in regular IPC rounds; raises stop-work when a trigger is observed."),
        ("Nursing Superintendent",
         "Ensures nursing staff comply with blood/body-fluid exposure prevention "
         "measures and report exposures without delay."),
        ("Department Heads",
         "Ensure staff in their departments comply with work restrictions when applied "
         "and support reporting without fear of penalty."),
        ("Quality Coordinator",
         "Audits this policy; holds training records, staff acknowledgements, and "
         "stop-work event logs."),
    ])

    h(doc, 1, "8. Quality monitoring")
    mon_tbl(doc, [
        ("Occupational health guidance",
         "Written guidance available and current; pre-employment screening conducted "
         "where appropriate; staff trained at induction and annually."),
        ("Immunisation programme",
         "Immunisation register current; Hepatitis B coverage confirmed for all direct-"
         "patient-care staff; statutory requirements met."),
        ("Work restrictions",
         "Records of work restrictions applied; reporting culture evidenced (staff "
         "report illness without penalty)."),
        ("Blood/body-fluid exposure prevention",
         "Sharps containers at the point of use in all clinical areas; regular IPC "
         "rounds confirm no recapping; PPE availability confirmed."),
        ("PEP provision",
         "PEP record file maintained by IPC Nurse; PEP aligned with national/"
         "international guidelines; no missed same-shift starts."),
        ("Stop-work events",
         "Stop-work events logged with trigger, action taken, and outcome."),
    ])

    h(doc, 1, "9. Training and staff acknowledgement")
    p(doc,
      "All clinical and nursing staff shall be familiar with the occupational health "
      "and safety guidance, the immunisation programme, the work restriction process, "
      "blood and body fluid exposure prevention (including safe-sharps practice), and "
      "the PEP path and stop-work authority in this policy.")
    p(doc,
      f"I have read the Policy on Prevention of Healthcare Associated Infections in "
      f"Staff of {HN}. I will follow the processes described, including the stop-work "
      "authority in Section 6.")
    sig_tbl(doc)

    h(doc, 1, "10. Distribution")
    p(doc,
      "This policy shall be available to the IPC Officer, the Occupational Health "
      "Physician, IPC Nurse(s), all clinical and nursing staff, department heads, "
      "and the Quality Coordinator.")

    h(doc, 1, "11. Abbreviations")
    abbrev_tbl(doc, [
        ("CAPA",  "Corrective and Preventive Action"),
        ("CSSD",  "Central Sterile Services Department"),
        ("HIV",   "Human Immunodeficiency Virus"),
        ("IPC",   "Infection Prevention and Control"),
        ("IPCO",  "Infection Prevention and Control Officer"),
        ("IPCN",  "Infection Prevention and Control Nurse"),
        ("NABH",  "National Accreditation Board for Hospitals and Healthcare Providers"),
        ("PEP",   "Post-Exposure Prophylaxis"),
        ("PPE",   "Personal Protective Equipment"),
        ("TB",    "Tuberculosis"),
        ("WHO",   "World Health Organization"),
    ])

    h(doc, 1, "12. Traceability table")
    p(doc,
      "This table is an index. It is not how the policy is organised. An asterisk in "
      "the Level column means documentation of the process is required.")
    tr = tbl(doc, 6, 3)
    for ci, hdr in enumerate(("Objective Element", "Level", "Traceability to this policy")):
        tr.cell(0, ci).text = hdr
    trace_rows = [
        ("IPC.8.a", "Commitment*",
         "Sections 3 and 5.1 address written occupational health and safety guidance, "
         "adequate IPC resources for staff, standard-precaution adherence, and where-"
         "appropriate pre-employment screening and immunisation-status capture."),
        ("IPC.8.b", "Commitment*",
         "Sections 3 and 5.2 address the staff immunisation policy in consonance with "
         "evidence, the Hepatitis B vaccination floor for direct-patient-care staff, "
         "and other relevant immunisation as per risk and statutory requirements."),
        ("IPC.8.c", "Achievement",
         "Sections 3, 5.3 and 6 address work restrictions for healthcare providers "
         "with transmissible infections, the culture of reporting without penalty, and "
         "the stop-work trigger when a restriction is disregarded."),
        ("IPC.8.d", "Commitment",
         "Sections 3 and 5.4 address PPE for blood/body-fluid exposure prevention, "
         "striving to eliminate needle devices, sharps containers at the point of use, "
         "and a no-recapping rule."),
        ("IPC.8.e", "Commitment*",
         "Sections 3, 5.5 and 6 address PEP for Hepatitis B and HIV exposure, "
         "alignment with national/international guidelines, same-shift start, IPC Nurse "
         "documentation, and the stop-work trigger for delayed PEP."),
    ]
    for ri, (oe, lvl, txt) in enumerate(trace_rows, 1):
        tr.cell(ri, 0).text = oe
        tr.cell(ri, 1).text = lvl
        tr.cell(ri, 2).text = txt

    h(doc, 1, "13. Required Records/Evidence Checklist")

    h(doc, 2, "Occupational health and safety guidance — IPC.8.a (Commitment*)")
    lb(doc,
       "Written occupational health and safety guidance covering PPE, vaccination "
       "access, exposure reporting, and work restriction.")
    lb(doc,
       "Pre-employment health screening records (where conducted) and immunisation "
       "status capture at induction.")
    lb(doc,
       "Staff training records for occupational health and safety at induction and "
       "annually.")

    h(doc, 2, "Staff immunisation programme — IPC.8.b (Commitment*)")
    lb(doc,
       "Written staff immunisation policy in consonance with available evidence.")
    lb(doc,
       "Immunisation register: staff name, vaccine, dose(s), date(s), booster schedule "
       "— current for all staff.")
    lb(doc,
       "Hepatitis B vaccination coverage record confirming all direct-patient-care "
       "staff are vaccinated (or have a documented contraindication or refusal).")

    h(doc, 2, "Work restrictions — IPC.8.c (Achievement)")
    lb(doc,
       "Written work restriction policy/guideline listing conditions and restriction "
       "types.")
    lb(doc,
       "Records of actual work restrictions applied (case records) — policy alone "
       "does not satisfy this Achievement OE.")
    lb(doc,
       "Evidence of reporting culture: documentation that staff are encouraged to "
       "report illness/exposure and are not penalised.")

    h(doc, 2, "Blood and body fluid exposure prevention — IPC.8.d")
    lb(doc,
       "Sharps container availability records confirming containers are at the point "
       "of use in all clinical areas and replaced before full.")
    lb(doc,
       "Regular IPC round records confirming no recapping practice observed.")
    lb(doc,
       "Needle-stick / exposure incident records and investigation outcomes.")

    h(doc, 2, "PEP provision — IPC.8.e (Commitment*)")
    lb(doc,
       "Written PEP protocol for Hepatitis B and HIV exposure, aligned with named "
       "national/international guidelines.")
    lb(doc,
       "PEP record file maintained by the IPC Nurse: staff name, date, exposure type, "
       "PEP administered, follow-up schedule.")
    lb(doc,
       "Evidence of same-shift start for each PEP event.")
    lb(doc,
       "Stop-work event log for any delayed-PEP events.")

    h(doc, 1, "14. References")
    ln(doc,
       "National Accreditation Board for Hospitals and Healthcare Providers. NABH "
       "Accreditation Standards for Hospitals, 6th Edition. IPC.8.")
    ln(doc, "Guidebook interpretation supplied for IPC.8.a through IPC.8.e.")
    ln(doc,
       f"Internal documents of {HN}: occupational health guidance; staff immunisation "
       "register; work restriction records; PEP protocol and PEP record file.")

    h(doc, 1, "Disclaimer")
    p(doc,
      "This policy reorganises the supplied IPC.8 objective-element wording and "
      "Guidebook interpretation into plain-language policy format. The modal strength "
      "of the source has been preserved. Optional examples and mechanisms have not been "
      "converted into mandatory requirements. The Hepatitis B vaccination floor for "
      "direct-patient-care staff, the same-shift PEP start requirement, the IPC Nurse "
      "documentation obligation, the non-penalisation requirement for staff who report "
      "illness or exposures, and the no-recapping practice requirement have been retained "
      "as mandatory requirements. The examples of transmissible infections triggering work "
      "restrictions (conjunctivitis, chickenpox, ARI, smear-positive TB) are treated as "
      "illustrative, not as a closed list. IPC.8 carries stop-work authority as stated "
      "in Section 6: continued duty against a work restriction, and delay of PEP after "
      "a qualifying exposure.")

    save_and_verify(doc, "HCO_IPC_8_v2_REWRITE_DRAFT.docx")


# ══════════════════════════════════════════════════════════════════════════════
if __name__ == "__main__":
    gen_ipc1()
    gen_ipc2()
    gen_ipc3()
    gen_ipc4()
    gen_ipc5()
    gen_ipc6()
    print("\nAll 6 IPC rewrite-reference DOCX files generated.")
