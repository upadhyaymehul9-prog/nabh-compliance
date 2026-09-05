# -*- coding: utf-8 -*-
"""
generate_hco_mom_rewrites.py
Generates HCO MOM chapter v2 rewrite-reference DOCX files.

Pipeline : python-docx, identical to generate_hco_cop_rewrites.py.
Output   : policies/build/rewrite_reference/HCO_MOM_N_v2_REWRITE_DRAFT.docx
Source   : Approved plain-language content (mom1_content.txt, mom2_content.txt)
           + policies/build/mom_raw_dump_1-6.txt
"""
import os
from docx import Document

HN  = "«Hospital Name»"
# Resolve relative to this file so the script can be run from any CWD
OUT = os.path.join(os.path.dirname(os.path.abspath(__file__)), "rewrite_reference")
os.makedirs(OUT, exist_ok=True)


# ── Helpers ───────────────────────────────────────────────────────────────────

def h(doc, lv, txt):
    return doc.add_paragraph(txt, style={0: "Title", 1: "Heading 1", 2: "Heading 2"}[lv])

def p(doc, txt=""):
    return doc.add_paragraph(txt, style="Normal")

def ln(doc, txt):
    return doc.add_paragraph(txt, style="List Number")

def lb(doc, txt):
    return doc.add_paragraph(txt, style="List Bullet")

def tbl(doc, rows, cols):
    t = doc.add_table(rows=rows, cols=cols)
    try:
        t.style = "Table Grid"
    except KeyError:
        pass
    return t

def doc_ctrl(doc, no, prep, appr="Medical Superintendent"):
    dc = tbl(doc, 6, 4)
    for ri, (a, b, c, d) in enumerate([
        ("Document No.", no,          "Version",                "2.0"),
        ("Issue No.",    "01",         "Review due",             "One year from implementation"),
        ("Date created", "________",   "Date of implementation", "________"),
    ]):
        dc.cell(ri, 0).text = a; dc.cell(ri, 1).text = b
        dc.cell(ri, 2).text = c; dc.cell(ri, 3).text = d
    for ri, (lbl, txt) in enumerate([
        ("Prepared by", f"{prep}  Name: ________  Signature: ________"),
        ("Reviewed by", "Quality Coordinator  Name: ________  Signature: ________"),
        ("Approved by", f"{appr}  Name: ________  Signature: ________"),
    ], start=3):
        dc.cell(ri, 0).text = lbl
        c1 = dc.cell(ri, 1); c1.text = txt; c1.merge(dc.cell(ri, 3))

def abbrev_tbl(doc, rows):
    t = tbl(doc, len(rows) + 1, 2)
    t.cell(0, 0).text = "Abbreviation"; t.cell(0, 1).text = "Meaning"
    for ri, (a, m) in enumerate(rows, 1):
        t.cell(ri, 0).text = a; t.cell(ri, 1).text = m

def gov_tbl(doc, rows):
    t = tbl(doc, len(rows) + 1, 2)
    t.cell(0, 0).text = "Role"; t.cell(0, 1).text = "Responsibility"
    for ri, (role, resp) in enumerate(rows, 1):
        t.cell(ri, 0).text = role; t.cell(ri, 1).text = resp

def mon_tbl(doc, rows):
    t = tbl(doc, len(rows) + 1, 2)
    t.cell(0, 0).text = "Monitoring area"; t.cell(0, 1).text = "What is monitored"
    for ri, (area, what) in enumerate(rows, 1):
        t.cell(ri, 0).text = area; t.cell(ri, 1).text = what

def sig_tbl(doc):
    sig = tbl(doc, 4, 4)
    for ci, hdr in enumerate(("Staff name", "Designation", "Signature", "Date")):
        sig.cell(0, ci).text = hdr
    for ri in range(1, 4):
        for ci in range(4):
            sig.cell(ri, ci).text = "________"

def save_and_verify(doc, fname):
    import sys
    out = sys.stdout

    def pr(s):
        try:
            out.write(s + "\n")
        except UnicodeEncodeError:
            out.write(s.encode("ascii", "replace").decode() + "\n")

    pr(f"\n=== {fname} ===")
    for i, para in enumerate(doc.paragraphs[:80]):
        sn = para.style.name if para.style else "(None)"
        pr(f"{i:3d}  {sn!r:30s}  {para.text[:60]!r}")
    counts = {}
    for para in doc.paragraphs:
        sn = para.style.name if para.style else "(None)"
        counts[sn] = counts.get(sn, 0) + 1
    pr("  Style inventory:")
    for sn, n in sorted(counts.items()):
        pr(f"    {sn}: {n}")
    pr(f"  Total paras: {len(doc.paragraphs)}")
    path = os.path.join(OUT, fname)
    doc.save(path)
    pr(f"  Saved: {path}")


# ══════════════════════════════════════════════════════════════════════════════
# MOM.1 — Safe Pharmacy Services and Medication Management   (NO stop-work)
# Content: mom1_content.txt (approved).
# Structure: Document control table, Governance table, Section 12 bullet list.
# COREs: none | Stars: a*, d* | Achievement: c | Excellence: none
# Exact items verified:
#   5.1: "Medication Management Manual" named as the single document
#   5.2: "at least once every three months" — exact frequency
#   5.3: "at least once a year" — annual system review exact
#   5.5: "within 24 hours" — drug recall notification timeframe exact
# ══════════════════════════════════════════════════════════════════════════════
def gen_mom1():
    doc = Document()

    # Title
    h(doc, 0, "Policy on Safe Pharmacy Services and Medication Management")
    p(doc, HN)

    # Document control
    h(doc, 1, "Document control")
    doc_ctrl(doc, "HCO/MOM/POL/01", "Medication Safety Officer")
    p(doc, "A blank marked ________ must be completed before issue.")

    # Statement of intent
    h(doc, 1, "Statement of intent")
    p(doc,
      f"Pharmacy services and medication management at {HN} are governed by written "
      "guidance, overseen by a multi-disciplinary committee, and stay reliable even "
      "when the pharmacy is closed or stock runs out.")

    # 1. Purpose
    h(doc, 1, "1. Purpose")
    p(doc,
      f"This policy explains how {HN} runs its pharmacy services and medication "
      "management under written guidance, uses a multi-disciplinary committee to guide "
      "and update that guidance, keeps medications available when the pharmacy is closed "
      "or out of stock, and keeps staff informed of key changes.")
    p(doc,
      "This policy does not cover the hospital formulary, medication storage, "
      "prescription, dispensing, or reconciliation in detail — those are covered in "
      "other hospital policies.")

    # 2. Scope
    h(doc, 1, "2. Scope")
    p(doc,
      f"This policy applies to all pharmacy staff and clinical staff involved in "
      f"medication management at {HN}.")

    # 3. Policy standards
    h(doc, 1, "3. Policy standards")
    p(doc,
      f"{HN} runs pharmacy services and medication management under a Medication "
      "Management Manual, guided by a multi-disciplinary committee that meets at least "
      "once every three months and reviews the whole system at least once a year. "
      "There is a standard procedure to obtain medications when the pharmacy is closed "
      "or out of stock, and a process to keep relevant staff informed of shortages, "
      "recalls, and safety incidents.")
    p(doc, "Staff follow the written guidance below and keep the records it requires.")

    # 4. Non-negotiable rules
    h(doc, 1, "4. Non-negotiable rules")
    lb(doc,
       "Do not manage medications without written guidance covering the formulary, "
       "procurement, storage, prescription, dispensing, administration, and monitoring "
       "— documented as a single Medication Management Manual, and supervised by a "
       "qualified individual across every patient care area.")
    lb(doc,
       "Do not run the medication management committee without its defined roles, its "
       "required composition, or its terms of reference, and do not let it meet less "
       "often than once every three months, without documenting the minutes.")
    lb(doc,
       "Do not skip the committee's review of the whole medication management system "
       "at least once a year.")
    lb(doc,
       "Do not leave the hospital without a standard procedure to obtain medications "
       "when the pharmacy is closed or stock runs out.")
    lb(doc,
       "Do not delay informing relevant staff of medication shortages, and do not take "
       "more than 24 hours to inform them of a drug recall.")

    # 5. What we do
    h(doc, 1, "5. What we do")

    h(doc, 2, "5.1 Manage medications under written guidance")
    p(doc,
      "Written guidance — documented as a single Medication Management Manual — covers "
      "the formulary, procurement, storage, prescription, dispensing, administration, "
      "and monitoring of medications. A qualified individual supervises all pharmacy "
      "service activities, and this guidance applies across every patient care area in "
      "the organisation.")

    h(doc, 2, "5.2 Guide medication management through a multi-disciplinary committee")
    p(doc,
      "A multi-disciplinary committee has defined roles and responsibilities for managing "
      "medications, in line with applicable legislation and regulations where relevant. "
      "Its responsibilities include, among other things, developing medication management "
      "processes, developing and revising the formulary, and evaluating medication use "
      "and safety incidents. The committee includes representatives from major clinical "
      "departments, administration, a pharmacist or clinical pharmacologist, the "
      "medication safety officer, nursing, and the Quality department. Its terms of "
      "reference — composition, meeting frequency, and quorum — are defined, and it "
      "meets at least once every three months, with minutes documented.")

    h(doc, 2, "5.3 Review and update the medication management system")
    p(doc,
      "The committee reviews the whole medication management system at least once a year, "
      "covering rational use, medication errors, medication management processes, adverse "
      "drug reactions, patient safety, and high-risk medications.")
    p(doc,
      f"{HN} designates a medication safety officer — this responsibility could be held "
      "by the patient safety officer. Related requirements are covered in the hospital's "
      "other policies.")

    h(doc, 2, "5.4 Keep medications available when the pharmacy is closed or out of stock")
    p(doc,
      "A standard operating procedure sets out how to procure medications when the "
      "pharmacy is closed, or during a stock-out. It is preferable for "
      f"{HN} to run a 24-hour pharmacy.")

    h(doc, 2, "5.5 Keep staff informed of key changes")
    p(doc,
      "A process is in place to communicate medication shortages and stock-outs to "
      "clinicians and nurses. Staff are informed of a drug recall within 24 hours, and "
      "of any serious adverse events or patient safety incidents connected to medication "
      "use. The Pharmacy In-Charge, or a designated authority, is responsible for this "
      "communication, and it is documented. Related requirements are covered in the "
      "hospital's other policies.")

    # 6. Governance and responsibility
    h(doc, 1, "6. Governance and responsibility")
    gov_tbl(doc, [
        ("Medical Superintendent",
         "Accountable for ensuring pharmacy services and medication management are "
         "resourced and implemented as required by this policy."),
        ("Medication Safety Officer",
         "Owns day-to-day implementation of this policy; coordinates medication-safety "
         "processes; brings incidents and audit findings to the multi-disciplinary "
         "committee."),
        ("Multi-disciplinary committee",
         "Guides the formulation and implementation of pharmacy services and medication "
         "management; reviews the whole system at least once a year; meets at least "
         "once every three months and documents minutes."),
        ("Pharmacy In-Charge / designated authority",
         "Responsible for communicating medication shortages, recalls, and safety "
         "incidents to relevant staff; ensures the after-hours and stock-out procedure "
         "is operational."),
        ("Quality Coordinator",
         "Audits this policy; holds training records and staff acknowledgements."),
    ])

    # 7. Quality monitoring
    h(doc, 1, "7. Quality monitoring")
    mon_tbl(doc, [
        ("Medication Management Manual",
         "Current, complete, and supervised by a qualified individual across all patient "
         "care areas."),
        ("Multi-disciplinary committee",
         "Constitution, terms of reference, required composition, meeting frequency "
         "of at least once every three months, and documented minutes."),
        ("Annual system review",
         "Conducted by the committee at least once a year and covering all required "
         "topics (rational use, medication errors, ADRs, patient safety, high-risk "
         "medications)."),
        ("After-hours / stock-out procedure",
         "Standard operating procedure in place, tested, and used when required."),
        ("Staff communication",
         "Drug recalls communicated within 24 hours; shortages and serious adverse "
         "events communicated and documented."),
    ])

    # 8. Training and staff acknowledgement
    h(doc, 1, "8. Training and staff acknowledgement")
    p(doc,
      "Pharmacy staff and clinical staff involved in medication management shall be "
      "familiar with the Medication Management Manual, the committee's role, the "
      "after-hours and stock-out procedure, and the staff communication process.")
    p(doc,
      f"I have read the Policy on Safe Pharmacy Services and Medication Management of "
      f"{HN}. I will follow the processes described.")
    sig_tbl(doc)

    # 9. Distribution
    h(doc, 1, "9. Distribution")
    p(doc,
      "This policy shall be available to pharmacy staff, the multi-disciplinary "
      "committee, clinical staff involved in medication management, the Medication "
      "Safety Officer, the Pharmacy In-Charge, and the Quality Coordinator.")

    # 10. Abbreviations
    h(doc, 1, "10. Abbreviations")
    abbrev_tbl(doc, [
        ("ADR",  "Adverse drug reaction"),
        ("DTC",  "Drug and Therapeutics Committee (the organisation's multi-disciplinary "
                  "medication management committee)"),
        ("MSO",  "Medication Safety Officer"),
        ("NABH", "National Accreditation Board for Hospitals and Healthcare Providers"),
        ("MOM",  "Management of Medication (NABH Hospitals chapter)"),
        ("SOP",  "Standard operating procedure"),
    ])

    # 11. Traceability table
    h(doc, 1, "11. Traceability table")
    p(doc,
      "This table is an index. It is not how the policy is organised. An asterisk in "
      "the Level column means documentation of the process is required.")
    tr = tbl(doc, 6, 3)
    for ci, hdr in enumerate(("Objective Element", "Level", "Traceability to this policy")):
        tr.cell(0, ci).text = hdr
    trace_rows = [
        ("MOM.1.a", "Commitment*",
         "Sections 3 and 5.1 address written guidance documented as a single Medication "
         "Management Manual, qualified supervision, and coverage of all patient care areas."),
        ("MOM.1.b", "Commitment",
         "Sections 3 and 5.2 address the committee's defined roles, required composition, "
         "terms of reference, meeting frequency of at least once every three months, and "
         "documented minutes."),
        ("MOM.1.c", "Achievement",
         "Sections 3 and 5.3 address the committee's annual review of the whole medication "
         "management system covering all required topics."),
        ("MOM.1.d", "Commitment*",
         "Sections 3 and 5.4 address the standard operating procedure to obtain medications "
         "when the pharmacy is closed or during a stock-out."),
        ("MOM.1.e", "Commitment",
         "Sections 3 and 5.5 address the staff communication process for shortages, "
         "drug recalls (within 24 hours), and serious adverse events, with documented "
         "communication by the Pharmacy In-Charge or designated authority."),
    ]
    for ri, (oe, lvl, txt) in enumerate(trace_rows, 1):
        tr.cell(ri, 0).text = oe
        tr.cell(ri, 1).text = lvl
        tr.cell(ri, 2).text = txt

    # 12. Required Records/Evidence Checklist
    h(doc, 1, "12. Required Records/Evidence Checklist")

    h(doc, 2, "Medication Management Manual (MOM.1.a)")
    lb(doc,
       "A single documented Medication Management Manual covering the formulary, "
       "procurement, storage, prescription, dispensing, administration, and monitoring.")
    lb(doc,
       "Evidence that a qualified individual supervises all pharmacy service activities.")
    lb(doc,
       "Evidence that the Manual applies across every patient care area in the "
       "organisation.")

    h(doc, 2, "Multi-disciplinary committee (MOM.1.b)")
    lb(doc,
       "Committee constitution record naming the required representatives: major "
       "clinical departments, administration, pharmacist or clinical pharmacologist, "
       "medication safety officer, nursing, and Quality department.")
    lb(doc,
       "Written terms of reference defining composition, meeting frequency, and quorum.")
    lb(doc,
       "Documented minutes of meetings held at least once every three months, naming "
       "decisions, owners, and due dates.")

    h(doc, 2, "Annual system review (MOM.1.c)")
    lb(doc,
       "Record of the committee's annual review of the medication management system.")
    lb(doc,
       "Evidence the review covered rational use, medication errors, medication "
       "management processes, adverse drug reactions, patient safety, and high-risk "
       "medications.")
    lb(doc,
       "Open-action tracking from the last review until closure.")

    h(doc, 2, "After-hours and stock-out procedure (MOM.1.d)")
    lb(doc,
       "Documented standard operating procedure to procure medications when the pharmacy "
       "is closed or during a stock-out.")
    lb(doc,
       "Evidence the procedure has been tested and is operational.")

    h(doc, 2, "Staff communication (MOM.1.e)")
    lb(doc,
       "Records of communication of medication shortages and stock-outs to clinicians "
       "and nurses.")
    lb(doc,
       "Records showing drug recalls communicated within 24 hours of notification.")
    lb(doc,
       "Records of communication of serious adverse events and patient safety incidents "
       "connected to medication use.")
    lb(doc,
       "Documentation identifying the Pharmacy In-Charge or designated authority as "
       "responsible for communication.")

    # 13. References
    h(doc, 1, "13. References")
    ln(doc,
       "National Accreditation Board for Hospitals and Healthcare Providers. NABH "
       "Accreditation Standards for Hospitals, 6th Edition. MOM.1.")
    ln(doc, "Guidebook interpretation supplied for MOM.1.a through MOM.1.e.")
    ln(doc,
       f"Internal documents of {HN}: Medication Management Manual; Drug and "
       "Therapeutics Committee terms of reference; committee minutes; after-hours and "
       "stock-out SOP; staff communication records.")

    # Disclaimer
    h(doc, 1, "Disclaimer")
    p(doc,
      "This policy reorganises the supplied MOM.1 objective-element wording and "
      "Guidebook interpretation into plain-language policy format. The modal strength "
      "of the source has been preserved. Optional examples and mechanisms have not been "
      "converted into mandatory requirements. The exact requirements of the Medication "
      "Management Manual as a single named document, the committee meeting frequency "
      "of at least once every three months, the annual system review, and the 24-hour "
      "drug recall notification timeframe have been retained verbatim. MOM.1 has no "
      "stop-work section.")

    save_and_verify(doc, "HCO_MOM_1_v2_REWRITE_DRAFT.docx")


# ══════════════════════════════════════════════════════════════════════════════
# MOM.2 — Hospital Formulary   (NO stop-work)
# Content: mom2_content.txt (approved).
# Structure: Document control table, Governance table, Section 12 bullet list.
# COREs: a | Stars: e*, f* | Achievement: d | Excellence: none
# Exact items verified:
#   5.1: molecule / formulation / strength — minimum three fields preserved
#   5.1: "Implants and devices are treated as drugs and included in the formulary"
#   5.2: "at least once a year" — annual review exact
#   5.5: all six acquisition-procedure elements preserved
#   5.6: three-step non-formulary process — evaluation / authorisation / ratification
#   MOM.2.a CORE designation does NOT create a stop-work section (none in this policy)
# ══════════════════════════════════════════════════════════════════════════════
def gen_mom2():
    doc = Document()

    # Title
    h(doc, 0, "Policy on Hospital Formulary")
    p(doc, HN)

    # Document control
    h(doc, 1, "Document control")
    doc_ctrl(doc, "HCO/MOM/POL/02", "Drug and Therapeutics Committee Chair")
    p(doc, "A blank marked ________ must be completed before issue.")

    # Statement of intent
    h(doc, 1, "Statement of intent")
    p(doc,
      f"The hospital formulary at {HN} is developed collaboratively, kept current, "
      "made available to every clinician, and followed — with clear procedures for "
      "acquiring both formulary and non-formulary medications.")

    # 1. Purpose
    h(doc, 1, "1. Purpose")
    p(doc,
      f"This policy explains how {HN} develops and updates its formulary, makes it "
      "available to clinicians, monitors adherence to it, and follows defined "
      "procedures for acquiring formulary and non-formulary medications.")
    p(doc,
      "This policy does not cover medication storage, prescription, or dispensing in "
      "detail — those are covered in other hospital policies.")

    # 2. Scope
    h(doc, 1, "2. Scope")
    p(doc,
      f"This policy applies to the multi-disciplinary medication management committee, "
      f"pharmacy staff, and all treating doctors at {HN}.")

    # 3. Policy standards
    h(doc, 1, "3. Policy standards")
    p(doc,
      f"{HN}'s multi-disciplinary committee develops the formulary to match the "
      "organisation's mission, patient needs, and scope of services, reviews and "
      "updates it at least once a year, and makes the current version available to "
      "every treating doctor. Clinicians' adherence to the formulary is monitored, "
      "and defined procedures govern how formulary and non-formulary medications are "
      "acquired.")
    p(doc, "Staff follow the written guidance below and keep the records it requires.")

    # 4. Non-negotiable rules
    h(doc, 1, "4. Non-negotiable rules")
    lb(doc,
       "Do not use a formulary that was not collaboratively prepared by the "
       "multi-disciplinary committee, or that omits the molecule name, formulation, "
       "and strength for each entry, or that leaves out implants and devices.")
    lb(doc,
       "Do not let the formulary go more than a year without a collaborative review "
       "and update.")
    lb(doc,
       "Do not leave the current formulary unavailable to any treating doctor.")
    lb(doc,
       "Do not acquire formulary medications without a documented procedure covering "
       "vendor selection, vendor evaluation, reorder levels, the indenting process, "
       "purchase order generation, receipt of goods, and managing stock-outs.")
    lb(doc,
       "Do not use a non-formulary medication through a local purchase without "
       "following an evaluation, authorisation, and ratification process, including "
       "a decision on whether it should be added to the formulary.")

    # 5. What we do
    h(doc, 1, "5. What we do")

    h(doc, 2, "5.1 Develop the formulary collaboratively")
    p(doc,
      "The multi-disciplinary committee prepares the formulary to include medications "
      "necessary for the organisation's mission, patient needs, and scope of services. "
      "The National List of Essential Medicines or the WHO Model List of Essential "
      "Medicines could inform this, along with factors like harm potential, drug "
      "interactions, and likelihood of patient safety incidents. The committee could "
      "also consider a system-wise or speciality-wise formulary.")
    p(doc,
      "At minimum, every formulary entry names the molecule, its formulation, and its "
      f"strength. {HN} works to limit the number of concentrations of any one drug in "
      "the formulary. Implants and devices are treated as drugs and included in the "
      "formulary.")

    h(doc, 2, "5.2 Review and update the formulary")
    p(doc,
      "The committee reviews and updates the formulary collaboratively at least once a "
      "year — this could cover all medications or focus on certain categories, and could "
      "be done speciality by speciality. Non-formulary drugs regularly procured the "
      "previous year could be added during this review. Patient safety factors — adverse "
      "drug reactions, changing disease or resistance patterns, and cost — could inform "
      "the review.")

    h(doc, 2, "5.3 Make the current formulary available")
    p(doc,
      f"The current formulary is available to every treating doctor at {HN}, in "
      "physical or electronic form.")

    h(doc, 2, "5.4 Monitor adherence to the formulary")
    p(doc,
      f"{HN} makes sure prescriptions follow the formulary, and monitors how often "
      "prescriptions are rejected or a local purchase is needed because a non-formulary "
      "drug was prescribed.")

    h(doc, 2, "5.5 Follow the procedure for acquiring formulary medications")
    p(doc,
      "A documented procedure governs vendor selection, vendor evaluation, reorder "
      "levels, the indenting process, generating purchase orders, and receiving goods, "
      "and also addresses how stock-outs are managed.")

    h(doc, 2, "5.6 Follow the procedure for non-formulary medications")
    p(doc,
      "Where a medication is not in the formulary — for example through a local "
      "purchase or a hotline arrangement for urgent need — "
      f"{HN} follows a process of evaluation, authorisation, and ratification, and "
      "makes a decision on whether that medication should be added to the formulary.")

    # 6. Governance and responsibility
    h(doc, 1, "6. Governance and responsibility")
    gov_tbl(doc, [
        ("Medical Superintendent",
         "Accountable for ensuring the formulary is in place and that acquisition "
         "and non-formulary procedures are followed."),
        ("Drug and Therapeutics Committee Chair",
         "Owns formulary development and annual review; chairs the committee that "
         "governs all formulary decisions."),
        ("Multi-disciplinary committee (DTC)",
         "Develops, reviews and updates the formulary collaboratively; approves "
         "non-formulary requests; monitors adherence."),
        ("Pharmacy In-Charge",
         "Implements the acquisition procedure; manages non-formulary procurement "
         "records; ensures the current formulary is available at all prescribing "
         "locations."),
        ("Treating doctors / clinicians",
         "Prescribe from the current formulary; follow the non-formulary request "
         "process where a non-formulary drug is required."),
        ("Quality Coordinator",
         "Monitors formulary adherence; audits this policy; holds training records "
         "and staff acknowledgements."),
    ])

    # 7. Quality monitoring
    h(doc, 1, "7. Quality monitoring")
    mon_tbl(doc, [
        ("Formulary development",
         "Collaboratively prepared by the committee; every entry names molecule, "
         "formulation, and strength; implants and devices included as drugs."),
        ("Annual review",
         "Conducted and documented by the committee at least once a year; additions, "
         "deletions, and restrictions minuted with clinical reason."),
        ("Formulary availability",
         "Current version available to every treating doctor in physical or "
         "electronic form; outdated versions removed."),
        ("Adherence monitoring",
         "Frequency of non-formulary prescriptions, rejections, and local purchases "
         "tracked and reported to the committee."),
        ("Acquisition procedure",
         "All six procedure elements present and in use: vendor selection, vendor "
         "evaluation, reorder levels, indenting, purchase order generation, and "
         "receipt of goods; stock-out management addressed."),
        ("Non-formulary process",
         "Evaluation, authorisation, and ratification documented for every local "
         "purchase of a non-formulary medication; formulary-addition decision "
         "recorded."),
    ])

    # 8. Training and staff acknowledgement
    h(doc, 1, "8. Training and staff acknowledgement")
    p(doc,
      "The multi-disciplinary committee, pharmacy staff, and all treating doctors "
      "shall be familiar with the formulary, the formulary acquisition procedure, "
      "and the non-formulary medication process.")
    p(doc,
      f"I have read the Policy on Hospital Formulary of {HN}. "
      "I will follow the processes described.")
    sig_tbl(doc)

    # 9. Distribution
    h(doc, 1, "9. Distribution")
    p(doc,
      "This policy shall be available to the multi-disciplinary committee, pharmacy "
      "staff, all treating doctors, the Medication Safety Officer, and the Quality "
      "Coordinator.")

    # 10. Abbreviations
    h(doc, 1, "10. Abbreviations")
    abbrev_tbl(doc, [
        ("DTC",  "Drug and Therapeutics Committee (the organisation's multi-disciplinary "
                  "medication management committee)"),
        ("MSO",  "Medication Safety Officer"),
        ("NABH", "National Accreditation Board for Hospitals and Healthcare Providers"),
        ("NLEM", "National List of Essential Medicines"),
        ("MOM",  "Management of Medication (NABH Hospitals chapter)"),
        ("WHO",  "World Health Organization"),
    ])

    # 11. Traceability table
    h(doc, 1, "11. Traceability table")
    p(doc,
      "This table is an index. It is not how the policy is organised. An asterisk in "
      "the Level column means documentation of the process is required.")
    tr = tbl(doc, 7, 3)
    for ci, hdr in enumerate(("Objective Element", "Level", "Traceability to this policy")):
        tr.cell(0, ci).text = hdr
    trace_rows = [
        ("MOM.2.a", "CORE",
         "Sections 3 and 5.1 address collaborative formulary preparation by the "
         "committee, minimum content (molecule, formulation, strength for every entry), "
         "and the inclusion of implants and devices as drugs."),
        ("MOM.2.b", "Commitment",
         "Sections 3 and 5.2 address the collaborative annual review and update of the "
         "formulary, including the review of all or certain medication categories."),
        ("MOM.2.c", "Commitment",
         "Sections 3 and 5.3 address making the current formulary available to every "
         "treating doctor in physical or electronic form."),
        ("MOM.2.d", "Achievement",
         "Section 5.4 addresses monitoring of prescription adherence to the formulary "
         "and tracking of rejections and local purchases."),
        ("MOM.2.e", "Commitment*",
         "Sections 3 and 5.5 address the documented acquisition procedure covering all "
         "six elements: vendor selection, vendor evaluation, reorder levels, indenting "
         "process, purchase order generation, receipt of goods, and stock-out "
         "management."),
        ("MOM.2.f", "Commitment*",
         "Sections 3 and 5.6 address the non-formulary medication process — evaluation, "
         "authorisation, and ratification — and the decision on formulary inclusion."),
    ]
    for ri, (oe, lvl, txt) in enumerate(trace_rows, 1):
        tr.cell(ri, 0).text = oe
        tr.cell(ri, 1).text = lvl
        tr.cell(ri, 2).text = txt

    # 12. Required Records/Evidence Checklist
    h(doc, 1, "12. Required Records/Evidence Checklist")

    h(doc, 2, "Formulary development — MOM.2.a (CORE)")
    lb(doc,
       "Formulary prepared collaboratively by the multi-disciplinary committee.")
    lb(doc,
       "Every formulary entry naming the molecule, formulation, and strength.")
    lb(doc,
       "Formulary including implants and devices as drugs.")
    lb(doc,
       "Record showing the committee limited the number of concentrations of any one "
       "drug where possible.")

    h(doc, 2, "Annual review — MOM.2.b")
    lb(doc,
       "Record of the committee's collaborative annual review and update of the "
       "formulary.")
    lb(doc,
       "Additions, deletions, and restrictions minuted with clinical reasons.")
    lb(doc,
       "Dated current-formulary cover showing the version in use.")

    h(doc, 2, "Formulary availability — MOM.2.c")
    lb(doc,
       "Current formulary available at prescribing locations — OPD consulting rooms, "
       "wards, ICU, emergency, OT, and the hospital information system or intranet.")
    lb(doc,
       "Record showing outdated copies removed when a new version is issued.")

    h(doc, 2, "Adherence monitoring — MOM.2.d")
    lb(doc,
       "Prescription-versus-formulary audit records.")
    lb(doc,
       "Tracking records for prescriptions rejected or resulting in local purchase "
       "because a non-formulary drug was prescribed.")
    lb(doc,
       "DTC records of any department's repeated non-adherence tabled for review.")

    h(doc, 2, "Acquisition procedure — MOM.2.e")
    lb(doc,
       "Documented acquisition procedure covering vendor selection, vendor evaluation, "
       "reorder levels, indenting process, purchase order generation, and receipt of "
       "goods.")
    lb(doc,
       "Record showing how stock-outs are managed under the procedure.")
    lb(doc,
       "Procurement records showing approved supplier, receipt quality checks, and "
       "batch and expiry logged.")

    h(doc, 2, "Non-formulary process — MOM.2.f")
    lb(doc,
       "Written non-formulary request record for every local purchase — including "
       "clinical justification (evaluation), approval (authorisation), and committee "
       "ratification.")
    lb(doc,
       "Record of the decision on whether the non-formulary item should be added to "
       "the formulary.")
    lb(doc,
       "Same-shift retrospective documentation record for any emergency non-formulary "
       "use.")

    # 13. References
    h(doc, 1, "13. References")
    ln(doc,
       "National Accreditation Board for Hospitals and Healthcare Providers. NABH "
       "Accreditation Standards for Hospitals, 6th Edition. MOM.2.")
    ln(doc, "Guidebook interpretation supplied for MOM.2.a through MOM.2.f.")
    ln(doc,
       "National List of Essential Medicines, Ministry of Health and Family Welfare, "
       "Government of India.")
    ln(doc,
       "WHO Model List of Essential Medicines, World Health Organization.")
    ln(doc,
       f"Internal documents of {HN}: hospital formulary; DTC meeting minutes; "
       "acquisition procedure; non-formulary request records.")

    # Disclaimer
    h(doc, 1, "Disclaimer")
    p(doc,
      "This policy reorganises the supplied MOM.2 objective-element wording and "
      "Guidebook interpretation into plain-language policy format. The modal strength "
      "of the source has been preserved. Optional examples and mechanisms have not been "
      "converted into mandatory requirements. The exact requirements of minimum "
      "formulary entry fields (molecule, formulation, strength), inclusion of implants "
      "and devices as drugs, the annual review frequency, all six acquisition-procedure "
      "elements, and the three-step non-formulary process (evaluation, authorisation, "
      "ratification) have been retained verbatim. MOM.2.a carries CORE status; this "
      "policy does not contain a stop-work section, which is correct — MOM.2 is not "
      "in the MOM stop-work proposals.")

    save_and_verify(doc, "HCO_MOM_2_v2_REWRITE_DRAFT.docx")


# ══════════════════════════════════════════════════════════════════════════════
# MOM.3 — Storage and Availability of Medications   (HAS stop-work: Section 6)
# Content: mom3_content.txt (approved).
# Structure: Document control, Sec 3 standards, Sec 4 non-negotiables,
#            Sec 5 (7 subsections), Sec 6 Stop-work, Sec 7 Governance,
#            Sec 8 Monitoring, Sec 9 Training, Sec 10 Distribution,
#            Sec 11 Abbreviations, Sec 12 Traceability, Sec 13 Records,
#            Sec 14 References, Disclaimer.
# COREs: a, c, e, g | Stars: c*, e*, f* | Achievement: d | Excellence: none
# Exact items verified:
#   5.1: manufacturer requirements apply to every area INCLUDING clinical areas
#   5.1: cold-storage temperature "at least once a day — or on every working day,
#        for areas not open daily" — both frequencies distinct
#   5.3: high-risk list CORE* — posted in pharmacy AND every clinical area storing it
#   5.5 + Sec 6: LASA/different-concentration separation is an absolute stop-work trigger
#   5.7: daily check distinct from sealed-cart "after each use or once a month"
# ══════════════════════════════════════════════════════════════════════════════
def gen_mom3():
    doc = Document()

    # Title
    h(doc, 0, "Policy on Storage and Availability of Medications")
    p(doc, HN)

    # Document control
    h(doc, 1, "Document control")
    doc_ctrl(doc, "HCO/MOM/POL/03", "Pharmacy In-Charge")
    p(doc, "A blank marked ________ must be completed before issue.")

    # Statement of intent
    h(doc, 1, "Statement of intent")
    p(doc,
      f"Medications at {HN} are stored safely and to manufacturer requirements "
      "everywhere in the hospital, including clinical areas. High-risk and "
      "look-alike/sound-alike drugs are kept physically apart, and emergency "
      "medications are always available and replenished immediately after use.")

    # 1. Purpose
    h(doc, 1, "1. Purpose")
    p(doc,
      f"This policy explains how {HN} stores medications safely, controls inventory, "
      "identifies and separates high-risk and look-alike/sound-alike medications, "
      "and keeps emergency medications available and replenished at all times.")
    p(doc,
      "This policy does not cover prescription, dispensing, or administration of "
      "medications in detail — those are covered in other hospital policies.")

    # 2. Scope
    h(doc, 1, "2. Scope")
    p(doc,
      f"This policy applies to pharmacy staff and clinical staff involved in the "
      f"storage of medications at {HN}.")

    # 3. Policy standards
    h(doc, 1, "3. Policy standards")
    p(doc,
      f"{HN} stores medications in a clean, safe, and secure environment that follows "
      "manufacturer recommendations everywhere in the hospital, including clinical "
      "areas, and applies sound inventory control. A high-risk medication list is "
      "defined, kept current, and posted wherever high-risk medications are stored, "
      "and look-alike, sound-alike medications and different concentrations of the "
      "same drug are kept physically apart. A defined, uniformly stored emergency "
      "medication list is kept available at all times and replenished immediately "
      "after use.")
    p(doc, "Staff follow the written guidance below and keep the records it requires.")

    # 4. Non-negotiable rules
    h(doc, 1, "4. Non-negotiable rules")
    lb(doc,
       "Do not store any medication outside its manufacturer's storage requirements, "
       "in any area of the hospital including clinical areas, and do not leave "
       "medications unprotected from loss or theft.")
    lb(doc,
       "Do not check cold-storage temperature less often than once a day — or on "
       "every working day, for areas not open daily — and do not store expired "
       "medications alongside those intended for patient use.")
    lb(doc,
       "Do not operate without a current, defined high-risk medication list posted "
       "in the pharmacy and every clinical area where high-risk medications are "
       "stored.")
    lb(doc,
       "Do not store high-risk medications outside their predetermined, clinically "
       "justified locations, or without safeguards against inadvertent administration.")
    lb(doc,
       "Do not store look-alike, sound-alike medications, or different concentrations "
       "of the same medication, physically together.")
    lb(doc,
       "Do not leave an emergency medication location without the defined list, or "
       "leave a used emergency medication unreplenished.")
    lb(doc,
       "Do not issue or use a medication that has been stored outside the "
       "manufacturer's temperature, light, or security recommendations until pharmacy "
       "has assessed it.")
    lb(doc,
       "Do not store any other drug alongside emergency medications, or check "
       "emergency medication stock less often than daily — or, for sealed carts, "
       "less often than after each use or once a month.")

    # 5. What we do
    h(doc, 1, "5. What we do")

    h(doc, 2, "5.1 Store medications safely")
    p(doc,
      "Medication storage spaces are clean, safe, and secure, following the "
      "manufacturer's storage requirements — where none exist, "
      f"{HN} develops and applies its own. This applies to every area where "
      "medications are stored, including clinical areas. Medications are protected "
      "from loss or theft — for example by limiting access to authorised staff, "
      "locking medication carts and never leaving them unattended, or keeping them "
      "in a continuously staffed area. It is preferable for the storage area to be "
      "well organised, and overall cleanliness is maintained. Vaccines are kept at "
      "the manufacturer's required temperature, with temperature monitoring of the "
      "room or refrigerator done at least once a day — or on every working day, for "
      "areas not open daily. Medications past their expiry date are stored separately "
      "from those intended for patient use, pending disposal.")

    h(doc, 2, "5.2 Apply sound inventory control")
    p(doc,
      f"{HN} follows recognised inventory control practices — for example ABC, VED, "
      "FSN, or First-Expiry-First-Out analysis, alone or combined. Medicines could "
      "be stored alphabetically by generic name, and stock verification audits could "
      f"run at intervals {HN} defines, to detect loss or theft. There is also a "
      "mechanism for handling medications outside the regular inventory — for example, "
      "physicians' samples not for sale.")

    h(doc, 2, "5.3 Define and maintain the high-risk medication list")
    p(doc,
      f"{HN} defines and periodically updates its list of high-risk medications — "
      "those carrying a heightened risk of serious harm if an error occurs, such as "
      "drugs with a low therapeutic window, controlled substances, psychotherapeutic "
      "medications, look-alike/sound-alike medications, and concentrated electrolytes. "
      "The list is available in the pharmacy and every clinical area where high-risk "
      "medications are stored.")

    h(doc, 2, "5.4 Store high-risk medications where clinically necessary")
    p(doc,
      "High-risk medications are kept in predetermined areas — for example certain "
      "wards, the OT, or the ICU — based on clinical need. Where regulations apply, "
      "such as for narcotics, storage follows them. Safeguards are in place in every "
      "such area to prevent inadvertent administration.")

    h(doc, 2, "5.5 Keep look-alike, sound-alike, and different-concentration medications apart")
    p(doc,
      "Look-alike and sound-alike medications, and different concentrations of the "
      "same drug, are identified periodically, drawn from the hospital formulary, "
      "and listed. This list is made available in every unit where drugs are stored, "
      "not just the pharmacy, and revised whenever the formulary or drug packaging "
      "changes. It is good practice to store these medications as far apart as "
      "possible — but at minimum, they are kept physically apart, in the pharmacy "
      "and in patient care areas alike.")

    h(doc, 2, "5.6 Define and uniformly store emergency medications")
    p(doc,
      "The list of emergency medications is prepared in line with sound clinical "
      "practice and documented — it could vary by department, for example ICU, "
      "physiotherapy, emergency, or the cath lab. A crash cart with defined rows "
      "and drawers is a useful way to store these consistently. No other drug is "
      "kept with emergency medications.")

    h(doc, 2, "5.7 Keep emergency medications available and replenished")
    p(doc,
      "Adequate quantities of emergency medications are stocked at all times, "
      "checked at least daily. Where "
      f"{HN} uses a sealed emergency cart, the check happens after each use of "
      "the cart, or once a month, whichever comes first.")

    # 6. Stop-work authority
    h(doc, 1, "6. Stop-work authority")
    p(doc,
      "Do not store look-alike, sound-alike medications, or different concentrations "
      "of the same medication, physically together.")
    p(doc,
      "Do not leave an emergency medication location without the defined list, or "
      "leave a used emergency medication unreplenished.")
    p(doc,
      "Do not issue or use a medication that has been stored outside the "
      "manufacturer's temperature, light, or security recommendations until pharmacy "
      "has assessed it.")
    p(doc,
      "Stop-work applies to the storage location and to issue from that location. "
      "Immediate life-saving use of the only available dose continues while "
      "escalation happens, and is documented.")
    p(doc,
      "The person who stops tells the Pharmacy In-Charge and the Medication Safety "
      "Officer the same shift. Refusing unsafe storage or issue is not a "
      "disciplinary matter.")

    # 7. Governance and responsibility
    h(doc, 1, "7. Governance and responsibility")
    gov_tbl(doc, [
        ("Medical Superintendent",
         "Accountable for ensuring medication storage requirements are resourced "
         "and implemented across the organisation."),
        ("Pharmacy In-Charge",
         "Owns day-to-day implementation of this policy; maintains the high-risk "
         "and LASA lists; manages emergency medication replenishment; receives "
         "stop-work escalations the same shift."),
        ("Medication Safety Officer",
         "Coordinates stop-work escalations; brings storage incidents and audit "
         "findings to the Drug and Therapeutics Committee."),
        ("Nursing Superintendent",
         "Ensures clinical areas comply with manufacturer storage requirements, "
         "high-risk storage rules, and emergency medication check schedules."),
        ("Quality Coordinator",
         "Audits this policy; holds training records and staff acknowledgements."),
    ])

    # 8. Quality monitoring
    h(doc, 1, "8. Quality monitoring")
    mon_tbl(doc, [
        ("Medication storage conditions",
         "Manufacturer requirements followed in all areas including clinical areas; "
         "cold-storage temperature logged at least once a day or every working day; "
         "expired medications stored separately."),
        ("Inventory control",
         "Recognised inventory control practices in use throughout the organisation."),
        ("High-risk medication list",
         "Current, defined list updated periodically; posted in pharmacy and every "
         "clinical area that stores high-risk medications; staff trained."),
        ("High-risk storage locations",
         "Predetermined locations documented; walk-round confirms high-risk "
         "medications stored only in authorised areas with safeguards."),
        ("LASA and different-concentration separation",
         "List available in every unit where drugs are stored; physical separation "
         "confirmed in pharmacy and all patient care areas."),
        ("Emergency medication list and storage",
         "Defined list documented; uniform storage layout; no other drugs stored "
         "alongside emergency medications."),
        ("Emergency medication availability",
         "Daily inventory check at every location; sealed-cart check after each use "
         "or once a month; replenishment logged immediately after use."),
        ("Stop-work events",
         "Stop-work events logged with trigger, action taken, and outcome."),
    ])

    # 9. Training and staff acknowledgement
    h(doc, 1, "9. Training and staff acknowledgement")
    p(doc,
      "Pharmacy staff and clinical staff involved in the storage of medications "
      "shall be familiar with the storage requirements, high-risk and LASA lists, "
      "emergency medication procedures, and the stop-work authority in this policy.")
    p(doc,
      f"I have read the Policy on Storage and Availability of Medications of {HN}. "
      "I will follow the processes described.")
    sig_tbl(doc)

    # 10. Distribution
    h(doc, 1, "10. Distribution")
    p(doc,
      "This policy shall be available to pharmacy staff, nursing staff, clinical "
      "staff involved in medication storage, the Pharmacy In-Charge, the Medication "
      "Safety Officer, and the Quality Coordinator.")

    # 11. Abbreviations
    h(doc, 1, "11. Abbreviations")
    abbrev_tbl(doc, [
        ("ADR",  "Adverse drug reaction"),
        ("DTC",  "Drug and Therapeutics Committee"),
        ("FEFO", "First expiry, first out"),
        ("ICU",  "Intensive Care Unit"),
        ("LASA", "Look-alike, sound-alike"),
        ("MOM",  "Management of Medication (NABH Hospitals chapter)"),
        ("MSO",  "Medication Safety Officer"),
        ("NABH", "National Accreditation Board for Hospitals and Healthcare Providers"),
        ("OT",   "Operation Theatre"),
    ])

    # 12. Traceability table
    h(doc, 1, "12. Traceability table")
    p(doc,
      "This table is an index. It is not how the policy is organised. An asterisk "
      "in the Level column means documentation of the process is required.")
    tr = tbl(doc, 8, 3)
    for ci, hdr in enumerate(("Objective Element", "Level", "Traceability to this policy")):
        tr.cell(0, ci).text = hdr
    trace_rows = [
        ("MOM.3.a", "CORE",
         "Sections 3 and 5.1 address clean, safe, secure storage following "
         "manufacturer requirements in all areas including clinical areas, protection "
         "from loss or theft, temperature monitoring at least once a day or every "
         "working day, and separation of expired medications."),
        ("MOM.3.b", "Commitment",
         "Section 5.2 addresses recognised inventory control practices throughout "
         "the organisation and a mechanism for non-regular-inventory medications."),
        ("MOM.3.c", "CORE*",
         "Sections 3 and 5.3 address the defined, periodically updated high-risk "
         "medication list posted in the pharmacy and every clinical area where "
         "high-risk medications are stored."),
        ("MOM.3.d", "Achievement",
         "Section 5.4 addresses storage of high-risk medications in predetermined, "
         "clinically justified locations with safeguards against inadvertent "
         "administration."),
        ("MOM.3.e", "CORE*",
         "Sections 3 and 5.5 address the identification, listing, and physical "
         "separation of LASA medications and different concentrations in the pharmacy "
         "and all patient care areas. Physical separation is also a stop-work trigger "
         "in Section 6."),
        ("MOM.3.f", "Commitment*",
         "Section 5.6 addresses the defined emergency-medication list prepared in "
         "line with sound clinical practice, documented, uniformly stored, and kept "
         "free of other drugs."),
        ("MOM.3.g", "CORE",
         "Section 5.7 addresses adequate emergency medication stock checked at least "
         "daily — or, for sealed carts, after each use or once a month — and "
         "replenished immediately after use."),
    ]
    for ri, (oe, lvl, txt) in enumerate(trace_rows, 1):
        tr.cell(ri, 0).text = oe
        tr.cell(ri, 1).text = lvl
        tr.cell(ri, 2).text = txt

    # 13. Required Records/Evidence Checklist
    h(doc, 1, "13. Required Records/Evidence Checklist")

    h(doc, 2, "Medication storage conditions — MOM.3.a (CORE)")
    lb(doc,
       "Storage-condition records showing manufacturer requirements followed in all "
       "areas of the hospital, including clinical areas.")
    lb(doc,
       "Cold-storage and refrigerator temperature log with entries at least once a "
       "day — or on every working day for areas not open daily — and excursion-"
       "reporting records.")
    lb(doc,
       "Access-control records for the main pharmacy and controlled-drug cupboards.")
    lb(doc,
       "Records showing expired medications stored separately from those intended "
       "for patient use, pending disposal.")

    h(doc, 2, "Inventory control — MOM.3.b")
    lb(doc,
       "Documented inventory-control method (for example FEFO, ABC, VED, or FSN) "
       "in use throughout the organisation.")
    lb(doc,
       "Stock-movement traceability records from receipt to issue for sampled items.")

    h(doc, 2, "High-risk medication list — MOM.3.c (CORE*)")
    lb(doc,
       "Current, DTC-approved high-risk medication list, updated periodically and "
       "after any related incident.")
    lb(doc,
       "Posted-list records showing the current list available in the pharmacy and "
       "every clinical area where high-risk medications are stored.")
    lb(doc,
       "Staff training records on the high-risk list.")

    h(doc, 2, "High-risk storage locations — MOM.3.d")
    lb(doc,
       "DTC record naming the predetermined locations where each high-risk "
       "medication is clinically necessary.")
    lb(doc,
       "Walk-round record confirming high-risk medications stored only in "
       "authorised areas.")
    lb(doc,
       "Safeguard documentation in each designated high-risk storage area.")

    h(doc, 2, "LASA and different-concentration separation — MOM.3.e (CORE*)")
    lb(doc,
       "LASA and different-concentration medication list developed from the "
       "hospital formulary, available in every unit where drugs are stored.")
    lb(doc,
       "Physical-separation records showing LASA medications and different "
       "concentrations stored apart — in separate bins or on separate shelves — "
       "in the pharmacy and all patient care areas.")
    lb(doc,
       "List-revision records following formulary changes or drug packaging changes.")

    h(doc, 2, "Emergency medication list and storage — MOM.3.f (Commitment*)")
    lb(doc,
       "Defined, documented emergency-medication list prepared in line with sound "
       "clinical practice.")
    lb(doc,
       "Uniform-storage layout record across crash carts and emergency trolleys.")
    lb(doc,
       "Confirmation that no other drug is stored alongside emergency medications.")

    h(doc, 2, "Emergency medication availability and replenishment — MOM.3.g (CORE)")
    lb(doc,
       "Daily inventory-check records at every emergency-medication location "
       "(unsealed cart).")
    lb(doc,
       "Sealed-cart check records: after each use of the cart, or once a month, "
       "whichever comes first.")
    lb(doc,
       "Immediate-replenishment log after any use of an emergency medication, "
       "before the cart is returned to service.")

    # 14. References
    h(doc, 1, "14. References")
    ln(doc,
       "National Accreditation Board for Hospitals and Healthcare Providers. NABH "
       "Accreditation Standards for Hospitals, 6th Edition. MOM.3.")
    ln(doc, "Guidebook interpretation supplied for MOM.3.a through MOM.3.g.")
    ln(doc,
       f"Internal documents of {HN}: high-risk medication list; LASA list; "
       "emergency-medication list; temperature logs; inventory-control records.")

    # Disclaimer
    h(doc, 1, "Disclaimer")
    p(doc,
      "This policy reorganises the supplied MOM.3 objective-element wording and "
      "Guidebook interpretation into plain-language policy format. The modal strength "
      "of the source has been preserved. Optional examples and mechanisms have not "
      "been converted into mandatory requirements. The exact requirements of "
      "manufacturer storage rules applying universally including clinical areas, "
      "the cold-storage temperature check at least once a day or on every working "
      "day for areas not open daily, the CORE high-risk list posted in all clinical "
      "areas storing high-risk medications, LASA and different-concentration "
      "physical separation as an absolute stop-work trigger, and the distinct "
      "emergency-medication check frequencies (daily for unsealed carts; after each "
      "use or once a month for sealed carts) have been retained verbatim.")

    save_and_verify(doc, "HCO_MOM_3_v2_REWRITE_DRAFT.docx")


# ══════════════════════════════════════════════════════════════════════════════
# MOM.4 — Safe and Rational Prescription of Medications   (HAS stop-work: Sec 6)
# Content: mom4_content.txt (approved).
# Structure: Document control, Sec 3 standards, Sec 4 non-negotiables,
#            Sec 5 (8 subsections), Sec 6 Stop-work, Sec 7 Governance,
#            Sec 8 Monitoring, Sec 9 Training, Sec 10 Distribution,
#            Sec 11 Abbreviations, Sec 12 Traceability, Sec 13 Records,
#            Sec 14 References, Disclaimer.
# COREs: b, e, f | Stars: a*, b*, f* | Achievement: g, h | Excellence: d
# Exact items verified:
#   5.2: all SEVEN minimum prescription elements preserved verbatim
#   5.2: capital letters + no error-prone abbreviations stated explicitly
#   Sec 6: drug-allergy-check stop-work trigger first in the block
#   5.5: all THREE reconciliation transition points (admission, transfer, discharge)
#   5.6: countersignature "within 24 hours" exact
#   5.7: "at least once a month, using a representative sample" exact
# ══════════════════════════════════════════════════════════════════════════════
def gen_mom4():
    doc = Document()

    # Title
    h(doc, 0, "Policy on Safe and Rational Prescription of Medications")
    p(doc, HN)

    # Document control
    h(doc, 1, "Document control")
    doc_ctrl(doc, "HCO/MOM/POL/04", "Medical Superintendent")
    p(doc, "A blank marked ________ must be completed before issue.")

    # Statement of intent
    h(doc, 1, "Statement of intent")
    p(doc,
      f"Medications at {HN} are prescribed rationally and meet defined minimum "
      "requirements every time, drug allergies are checked before every prescription, "
      "verbal orders follow a documented safe process, and medications are reconciled "
      "at every transition in a patient's care.")

    # 1. Purpose
    h(doc, 1, "1. Purpose")
    p(doc,
      f"This policy explains how {HN} ensures rational prescribing, applies minimum "
      "prescription requirements, checks drug allergies and adverse drug reactions "
      "before prescribing, reconciles medications at transitions of care, controls "
      "verbal orders, and audits prescriptions for safety.")
    p(doc,
      "This policy does not cover storage, dispensing, or administration of "
      "medications in detail — those are covered in other hospital policies.")

    # 2. Scope
    h(doc, 1, "2. Scope")
    p(doc,
      f"This policy applies to all doctors who prescribe medications, and to pharmacy "
      f"and nursing staff involved in checking, dispensing, and administering them, "
      f"at {HN}.")

    # 3. Policy standards
    h(doc, 1, "3. Policy standards")
    p(doc,
      f"{HN} prescribes medications in line with rational-prescribing good practice "
      "for both outpatients and inpatients, applies defined minimum requirements to "
      "every prescription, checks drug allergies and previous adverse reactions before "
      "prescribing, gives clinicians a mechanism to help avoid interactions and dosing "
      "errors, reconciles medications at every transition of care, controls verbal "
      "orders through a documented process, and audits prescriptions monthly for "
      "safety and rationality, acting on the findings where appropriate.")
    p(doc, "Staff follow the written guidance below and keep the records it requires.")

    # 4. Non-negotiable rules
    h(doc, 1, "4. Non-negotiable rules")
    lb(doc,
       "Do not prescribe outside of good practice and guidelines for rational "
       "prescribing — this applies to both outpatient and inpatient care — and do "
       "not skip training clinicians on rational prescribing.")
    lb(doc,
       f"Do not accept, dispense, or administer a prescription that fails {HN}'s "
       "determined minimum requirements: patient name, unique hospital number, drug "
       "name (generic composition, except for vitamin or mineral combinations), "
       "strength, dosage instruction, duration and total quantity, and the prescribing "
       "doctor's name, signature, and registration number. Do not use error-prone "
       "abbreviations, and write every prescription in capital letters.")
    lb(doc,
       "Do not prescribe, or transcribe a prescription for action, without first "
       "ascertaining the patient's drug allergies and previous adverse drug reactions.")
    lb(doc,
       "Do not skip medication reconciliation at admission, at transfer between wards "
       "or departments, or at discharge, and do not leave it undocumented.")
    lb(doc,
       f"Do not act on a verbal medication order outside {HN}'s documented verbal-"
       "order process — including read-back, documentation at the time it is received, "
       "and countersignature by the ordering doctor within 24 hours.")
    lb(doc,
       "Do not skip the monthly medication order audit using a representative sample, "
       "and do not leave the corrective or preventive action decision undocumented "
       "where the audit calls for one.")

    # 5. What we do
    h(doc, 1, "5. What we do")

    h(doc, 2, "5.1 Prescribe rationally")
    p(doc,
      "Prescribing follows good practice and guidelines for rational prescription of "
      "medications, across both outpatient and inpatient care. Clinicians are trained "
      "or sensitised on rational prescribing. The Code of Medical Ethics-2002, "
      "published by the Medical Council of India (now the National Medical "
      "Commission), is followed.")

    h(doc, 2, "5.2 Meet minimum prescription requirements")
    p(doc,
      "Every prescription — whether inpatient, outpatient, or emergency — follows "
      "applicable national and international guidelines. At minimum, it includes:")
    lb(doc, "the patient's name")
    lb(doc, "the patient's unique hospital number")
    lb(doc,
       "the drug name in generic composition, except for vitamin or mineral "
       "combinations")
    lb(doc, "strength")
    lb(doc, "dosage instructions")
    lb(doc, "duration and total quantity")
    lb(doc,
       "the prescribing doctor's name, signature, and registration number.")
    p(doc,
      "Error-prone abbreviations are not used, and prescriptions are written in "
      "capital letters. It is preferable to use a digital prescription system to "
      "reduce errors. Where a prescription has an error or is illegible, it is "
      "corrected with a single strikethrough, initialled, and rewritten.")

    h(doc, 2, "5.3 Check for drug allergies and prior reactions")
    p(doc,
      "Drug allergies and any previous adverse drug reactions are ascertained before "
      "prescribing — during the initial consultation, or at any point during the "
      "patient's care. It is good practice to record this prominently in the medical "
      "record for both outpatients and inpatients.")

    h(doc, 2, "5.4 Help clinicians avoid interactions and dosing errors")
    p(doc,
      f"{HN} provides clinicians with a mechanism — electronic or physical — to help "
      "identify drug interactions, food-drug interactions, alcohol-drug interactions, "
      "therapeutic duplication, and dose adjustments.")

    h(doc, 2, "5.5 Reconcile medications at transitions of care")
    p(doc,
      "Prescribed medications are checked for accuracy at three points:")
    lb(doc,
       "at admission — whether direct or after admission from the emergency "
       "department")
    lb(doc, "at transfer between wards or departments")
    lb(doc, "at discharge.")
    p(doc,
      "It is preferable to also reconcile after a cross-consultation. Reconciliation "
      "is documented, and there is a system to communicate it effectively during "
      "handover.")

    h(doc, 2, "5.6 Control verbal orders")
    p(doc,
      "Written guidance sets out who can give a verbal medication order, when it can "
      "be given, and how it is authenticated. Verbal orders are limited to urgent "
      "situations where immediate written or electronic communication is not "
      "practical, and their use is kept to a minimum. An approved list of formulary "
      "drugs that can be ordered verbally is maintained — defined either by what is "
      "included or what is excluded. The process includes read-back or repeat-back, "
      "and every verbal order is countersigned by the ordering doctor within 24 hours.")

    h(doc, 2, "5.7 Audit prescriptions")
    p(doc,
      "Medication order and prescription audits check legibility and use of capitals, "
      "the appropriateness of drug, dose, frequency, and route, therapeutic "
      "duplication, drug and food-drug interactions and how they are avoided, dosage "
      "adjustment for renal or hepatic impairment, IV incompatibility, inappropriate "
      "dilutions or infusion durations, and adherence to this policy's other "
      "requirements. This happens at least once a month, using a representative "
      "sample. It is preferably done by a clinical pharmacologist or clinical "
      "pharmacist — where none is available, a trained multidisciplinary team can "
      "do it instead. It is preferable to audit prescriptions live, before dispensing.")

    h(doc, 2, "5.8 Act on audit findings")
    p(doc,
      "Where appropriate, corrective or preventive action is taken based on the "
      "audit — ideally guided by root-cause analysis — and records of these actions "
      "are kept.")

    # 6. Stop-work authority
    h(doc, 1, "6. Stop-work authority")
    p(doc,
      "Do not prescribe (or transcribe a prescription for action) when drug allergies "
      "and previous adverse drug reactions have not been ascertained.")
    p(doc,
      "Do not act on a verbal medication order except through the organisation's "
      "documented verbal-order process (read-back, documentation, countersignature "
      "within the defined time).")
    p(doc,
      "Do not accept a prescription that fails the organisation's determined minimum "
      "requirements.")
    p(doc,
      "Stop-work applies to writing or acting on the unsafe order. Immediate "
      "life-saving medication in an emergency follows the documented emergency-"
      "prescription rules and is written up as soon as the patient is stable.")
    p(doc,
      "The person who stops tells the treating doctor and the Medication Safety "
      "Officer the same shift. Refusing an unsafe prescription is not a disciplinary "
      "matter.")

    # 7. Governance and responsibility
    h(doc, 1, "7. Governance and responsibility")
    gov_tbl(doc, [
        ("Medical Superintendent",
         "Accountable for rational prescribing implementation; maintains the "
         "authorised-prescriber list; receives stop-work escalations."),
        ("Medication Safety Officer",
         "Owns the prescription audit programme; brings audit findings and incidents "
         "to the Drug and Therapeutics Committee."),
        ("Treating doctors / prescribers",
         "Follow rational-prescribing guidance and the Code of Medical Ethics-2002; "
         "apply minimum prescription requirements; check drug allergies before "
         "prescribing; follow the verbal-order and reconciliation processes."),
        ("Pharmacy In-Charge",
         "Enforces minimum prescription requirements at the point of dispensing; "
         "holds the list of drugs orderable verbally; implements the verbal-order "
         "control process."),
        ("Nursing staff",
         "Check prescriptions before administration; complete medication "
         "reconciliation at transitions of care."),
        ("Quality Coordinator",
         "Audits this policy; holds training records and staff acknowledgements."),
    ])

    # 8. Quality monitoring
    h(doc, 1, "8. Quality monitoring")
    mon_tbl(doc, [
        ("Rational prescribing",
         "Clinicians trained; Code of Medical Ethics-2002 followed; prescribing "
         "in line with rational-prescribing guidelines for outpatients and "
         "inpatients."),
        ("Minimum prescription requirements",
         "All seven elements present (patient name, unique hospital number, generic "
         "drug name, strength, dosage instruction, duration and total quantity, "
         "prescriber name/signature/registration number); capital letters used; "
         "no error-prone abbreviations; non-compliant prescriptions held."),
        ("Drug-allergy check",
         "Drug allergies and previous ADRs ascertained and documented before "
         "prescribing for every patient."),
        ("Prescribing assistance mechanism",
         "Mechanism available and working — electronic or physical — for drug "
         "interactions, food-drug interactions, therapeutic duplication, and dose "
         "adjustments."),
        ("Medication reconciliation",
         "Documented at all three transition points: admission, transfer between "
         "wards or departments, and discharge."),
        ("Verbal orders",
         "Read-back performed and documented; countersignature by ordering doctor "
         "within 24 hours; verbal orders limited to urgent situations."),
        ("Prescription audit",
         "Conducted at least once a month with a representative sample; all scope "
         "parameters covered; findings documented; CAPA recorded where appropriate."),
        ("Stop-work events",
         "Stop-work events logged with trigger, action taken, and outcome."),
    ])

    # 9. Training and staff acknowledgement
    h(doc, 1, "9. Training and staff acknowledgement")
    p(doc,
      "All doctors who prescribe, and pharmacy and nursing staff involved in "
      "checking, dispensing, and administering medications, shall be familiar with "
      "the rational-prescribing guidance, minimum prescription requirements, "
      "allergy-check process, reconciliation requirements, verbal-order process, "
      "and stop-work authority in this policy.")
    p(doc,
      f"I have read the Policy on Safe and Rational Prescription of Medications "
      f"of {HN}. I will follow the processes described.")
    sig_tbl(doc)

    # 10. Distribution
    h(doc, 1, "10. Distribution")
    p(doc,
      "This policy shall be available to all treating doctors, pharmacy staff, "
      "nursing staff, the Medical Superintendent, the Medication Safety Officer, "
      "and the Quality Coordinator.")

    # 11. Abbreviations
    h(doc, 1, "11. Abbreviations")
    abbrev_tbl(doc, [
        ("ADR",  "Adverse drug reaction"),
        ("CAPA", "Corrective and Preventive Action"),
        ("DTC",  "Drug and Therapeutics Committee"),
        ("IV",   "Intravenous"),
        ("MCI",  "Medical Council of India"),
        ("MOM",  "Management of Medication (NABH Hospitals chapter)"),
        ("MSO",  "Medication Safety Officer"),
        ("NABH", "National Accreditation Board for Hospitals and Healthcare Providers"),
        ("NMC",  "National Medical Commission"),
    ])

    # 12. Traceability table
    h(doc, 1, "12. Traceability table")
    p(doc,
      "This table is an index. It is not how the policy is organised. An asterisk "
      "in the Level column means documentation of the process is required.")
    tr = tbl(doc, 9, 3)
    for ci, hdr in enumerate(("Objective Element", "Level", "Traceability to this policy")):
        tr.cell(0, ci).text = hdr
    trace_rows = [
        ("MOM.4.a", "Commitment*",
         "Sections 3 and 5.1 address rational prescribing for outpatients and "
         "inpatients, clinician training, and the Code of Medical Ethics-2002."),
        ("MOM.4.b", "CORE*",
         "Sections 3, 4, and 5.2 address all seven minimum prescription elements, "
         "capital letters, no error-prone abbreviations, and the process for holding "
         "non-compliant prescriptions. Physical-separation of requirements is "
         "presented as a bullet list for assessor legibility."),
        ("MOM.4.c", "Commitment",
         "Sections 3 and 5.3 address ascertainment of drug allergies and previous "
         "adverse drug reactions before prescribing. This is also the first stop-work "
         "trigger in Section 6."),
        ("MOM.4.d", "Excellence",
         "Section 5.4 addresses the clinician assistance mechanism — electronic or "
         "physical — for drug interactions, food-drug interactions, therapeutic "
         "duplication, and dose adjustments."),
        ("MOM.4.e", "CORE",
         "Sections 3 and 5.5 address medication reconciliation at all three mandatory "
         "transition points: admission, transfer between wards or departments, and "
         "discharge, with documentation and handover communication."),
        ("MOM.4.f", "CORE*",
         "Sections 3 and 5.6 address the verbal-order process: written guidance "
         "naming who, when, and how; read-back; countersignature within 24 hours; "
         "and a list of drugs orderable verbally. This is also a stop-work trigger "
         "in Section 6."),
        ("MOM.4.g", "Achievement",
         "Section 5.7 addresses the monthly prescription audit with a representative "
         "sample, covering all required scope parameters."),
        ("MOM.4.h", "Achievement",
         "Section 5.8 addresses corrective or preventive action taken where "
         "appropriate based on audit findings, with records kept."),
    ]
    for ri, (oe, lvl, txt) in enumerate(trace_rows, 1):
        tr.cell(ri, 0).text = oe
        tr.cell(ri, 1).text = lvl
        tr.cell(ri, 2).text = txt

    # 13. Required Records/Evidence Checklist
    h(doc, 1, "13. Required Records/Evidence Checklist")

    h(doc, 2, "Rational prescribing — MOM.4.a (Commitment*)")
    lb(doc,
       "DTC-named rational-prescribing reference (Code of Medical Ethics-2002; "
       "applicable rational-prescribing guidelines).")
    lb(doc,
       "Staff training records showing clinicians trained or sensitised on rational "
       "prescribing.")
    lb(doc,
       "Prescription-audit records sampled against rational-prescribing guidance.")

    h(doc, 2, "Minimum prescription requirements — MOM.4.b (CORE*)")
    lb(doc,
       "Published minimum-prescription-requirement list showing all seven mandatory "
       "elements.")
    lb(doc,
       "Sample prescriptions showing all elements present, written in capital "
       "letters, with no error-prone abbreviations.")
    lb(doc,
       "Hold or reject records for prescriptions failing minimum requirements, "
       "except through the documented emergency path.")

    h(doc, 2, "Drug-allergy check — MOM.4.c")
    lb(doc,
       "Drug-allergy and previous-ADR ascertainment records before prescribing, "
       "including entries stating 'none known'.")
    lb(doc,
       "Prominent allergy documentation in the medical record for outpatients and "
       "inpatients.")

    h(doc, 2, "Prescribing assistance mechanism — MOM.4.d")
    lb(doc,
       "Evidence of a working assistance mechanism — electronic or physical — for "
       "drug interactions, food-drug interactions, alcohol-drug interactions, "
       "therapeutic duplication, and dose adjustments.")
    lb(doc,
       "Records of orders clarified by pharmacy before dispensing where the "
       "mechanism identified a concern.")

    h(doc, 2, "Medication reconciliation — MOM.4.e (CORE)")
    lb(doc,
       "Reconciled-medication-list records at admission (direct or post-emergency), "
       "transfer between wards or departments, and at discharge.")
    lb(doc,
       "Discrepancy-resolution records between the reconciling clinician and the "
       "prescriber.")
    lb(doc,
       "Handover communication records showing reconciliation status transmitted "
       "at the transition point.")

    h(doc, 2, "Verbal orders — MOM.4.f (CORE*)")
    lb(doc,
       "Written verbal-order guidance naming who can give a verbal order, when, "
       "and how it is authenticated.")
    lb(doc,
       "Read-back records for each verbal order — drug, dose, route, frequency, "
       "patient — before administration.")
    lb(doc,
       "Countersignature records by the ordering doctor within 24 hours of the "
       "verbal order.")
    lb(doc,
       "Approved list of formulary drugs that can be ordered verbally.")

    h(doc, 2, "Prescription audit — MOM.4.g")
    lb(doc,
       "Monthly audit records with a representative sample, covering all required "
       "scope parameters.")
    lb(doc,
       "DTC or Medication Safety Officer presentation record of audit findings.")
    lb(doc,
       "Evidence the audit sample size was representative and the frequency was at "
       "least once a month.")

    h(doc, 2, "CAPA from audit — MOM.4.h")
    lb(doc,
       "CAPA records from audit findings, with owner and due date.")
    lb(doc,
       "Closure-tracking records for each open action.")
    lb(doc,
       "Decision records noting where CAPA was considered and judged not required, "
       "where that is the case.")

    # 14. References
    h(doc, 1, "14. References")
    ln(doc,
       "National Accreditation Board for Hospitals and Healthcare Providers. NABH "
       "Accreditation Standards for Hospitals, 6th Edition. MOM.4.")
    ln(doc, "Guidebook interpretation supplied for MOM.4.a through MOM.4.h.")
    ln(doc,
       "Code of Medical Ethics-2002, Medical Council of India (now National Medical "
       "Commission).")
    ln(doc,
       "Institute for Safe Medication Practices (ISMP) guidelines on error-prone "
       "abbreviations.")
    ln(doc,
       f"Internal documents of {HN}: rational-prescribing guidance; minimum-"
       "prescription-requirement list; verbal-order procedure; medication-"
       "reconciliation records; prescription audit reports.")

    # Disclaimer
    h(doc, 1, "Disclaimer")
    p(doc,
      "This policy reorganises the supplied MOM.4 objective-element wording and "
      "Guidebook interpretation into plain-language policy format. The modal strength "
      "of the source has been preserved. Optional examples and mechanisms have not "
      "been converted into mandatory requirements. The exact requirements of all seven "
      "minimum prescription elements, the capital-letters and no-error-prone-"
      "abbreviations rules, the drug-allergy-check stop-work trigger, all three "
      "reconciliation transition points (admission, transfer, discharge), the verbal-"
      "order countersignature within 24 hours, and the monthly audit frequency with a "
      "representative sample have been retained verbatim.")

    save_and_verify(doc, "HCO_MOM_4_v2_REWRITE_DRAFT.docx")


def gen_mom5():
    """MOM.5 — Uniform Medication Orders (no stop-work; sections 1-13 + Disclaimer)"""
    doc = Document()

    # Title
    h(doc, 0, "Policy on Uniform Medication Orders")
    p(doc, HN)

    # Document control
    h(doc, 1, "Document control")
    doc_ctrl(doc, "HCO/MOM/POL/05", "Medical Superintendent")
    p(doc, "A blank marked ________ must be completed before issue.")

    # Statement of intent
    h(doc, 1, "Statement of intent")
    p(doc,
      "Medication orders are written only by authorised personnel, in one uniform "
      "place in the medical record, legibly, and complete with every required detail "
      "every time.")

    # 1. Purpose
    h(doc, 1, "1. Purpose")
    p(doc,
      f"This policy explains how {HN} ensures only authorised personnel write "
      "medication orders, keeps every order in one uniform location, keeps orders "
      "legible and traceable, and requires every order to include the medicine name, "
      "route, strength, and frequency.")
    p(doc,
      "This policy does not cover prescription content requirements or dispensing "
      "in detail — those are covered in other hospital policies.")

    # 2. Scope
    h(doc, 1, "2. Scope")
    p(doc,
      f"This policy applies to all doctors and authorised staff who write medication "
      f"orders at {HN}.")

    # 3. Policy standards
    h(doc, 1, "3. Policy standards")
    p(doc,
      f"{HN} ensures only doctors — or staff specifically authorised by legislation "
      "or government order — write medication orders. Every order is recorded in a "
      "single, uniform location in the medical record, including the patient's name "
      "and unique identification number, and only medications recorded there are "
      "administered. Orders are legible, dated, timed, signed, and traceable to the "
      "person who wrote them, and every order names the medicine, route of "
      "administration, strength, and frequency or time of administration.")
    p(doc, "Staff follow the written guidance below and keep the records it requires.")

    # 4. Non-negotiable rules
    h(doc, 1, "4. Non-negotiable rules")
    lb(doc,
       "Do not let anyone other than a doctor — holding at least an MBBS qualification "
       "— write a medication order, unless another category of staff is specifically "
       "authorised by legislation or a government order.")
    lb(doc,
       "Do not administer a medication that isn't recorded in the uniform medication-"
       "order location of the medical record, and do not accept phrases like "
       "\"continue same treatment,\" \"repeat all,\" or similar shorthand in place of "
       "a written order.")
    lb(doc,
       "Do not modify an existing medication order by striking through or overwriting "
       "it — discontinue the original and write a fresh order instead.")
    lb(doc,
       "Do not use error-prone abbreviations in a medication order, and do not use "
       "any abbreviation outside the hospital's approved, standardised list.")
    lb(doc,
       "Do not write a multi-drug order without stating the strength of every "
       "individual drug, except where the combination is only of vitamins and/or "
       "minerals.")
    lb(doc,
       "Do not leave an incomplete medication order — missing the drug name, route, "
       "strength, or frequency — without a mechanism to catch and correct it.")

    # 5. What we do
    h(doc, 1, "5. What we do")

    h(doc, 2, "5.1 Ensure only authorised personnel write orders")
    p(doc,
      "Medication orders are written by a doctor holding at minimum an MBBS "
      "qualification. Any other staff category authorised to write orders is backed "
      "by legislation or a government order, not an internal hospital decision alone. "
      "Even when transcribing a treating consultant's orders from an OP record or "
      "admission note, a doctor writes the inpatient medication order. Where the "
      "hospital uses an electronic medical record, the doctor enters the prescription "
      "directly using their own login; if an assistant enters it, the doctor verifies "
      "and authorises it.")

    h(doc, 2, "5.2 Keep every order in one uniform location")
    p(doc,
      "Every medication order is recorded in a single, uniform location in the "
      "medical record, which includes the patient's name and unique identification "
      "number. Only medications recorded there are administered — orders written "
      "anywhere else are moved to this location, and electronic orders follow the "
      "same principle. It's preferable for the prescription and administration record "
      "to sit on the same sheet — a drug \"Kardex,\" updated or authorised daily, is "
      "one useful format for this. Phrases like \"continue same treatment,\" "
      "\"repeat all,\" or \"repeat 1,4,5,8\" aren't accepted in place of a written "
      "order. Where a drug's dose changes — for example from four times a day to "
      "twice a day — the original order is discontinued and a fresh order written; a "
      "strike-through or overwrite of the old order isn't acceptable.")

    h(doc, 2, "5.3 Keep orders legible and traceable")
    p(doc,
      "Hand-written medication orders are written in capital letters. Where "
      "abbreviations are used, only the hospital's approved, standardised list "
      "applies, throughout the organisation — the Institute for Safe Medication "
      "Practices guidelines are a useful reference. Error-prone abbreviations aren't "
      "used. The identity of whoever wrote the order is traceable — for example by "
      "name against each order, a master signature list, or an employee code.")

    h(doc, 2, "5.4 Include every required detail")
    p(doc,
      "Where an order includes two or more drugs, the strength of each individual "
      "drug is stated — this may not apply to combinations of vitamins and/or "
      "minerals. Where the strength of a drug differs by time of administration, "
      "separate orders are recorded for each. There's a mechanism to catch and act "
      "on any order that's incomplete on drug name, route, strength, or frequency "
      "and time of administration. (Related requirements are covered in the "
      "hospital's other policies.)")

    # 6. Governance and responsibility
    h(doc, 1, "6. Governance and responsibility")
    gov_tbl(doc, [
        ("Medical Superintendent",
         "Accountable for implementing this policy; maintains the authorised-"
         "prescriber list."),
        ("Treating doctors / prescribers",
         "Write medication orders per this policy — MBBS minimum, uniform location, "
         "legible, complete — and do not delegate order-writing to non-authorised "
         "staff."),
        ("Nursing Superintendent / Nursing In-Charge",
         "Ensures only medications recorded in the uniform location are administered; "
         "checks orders for legibility and completeness before administration."),
        ("Quality Coordinator",
         "Audits adherence to this policy; holds training records."),
    ])

    # 7. Quality monitoring
    h(doc, 1, "7. Quality monitoring")
    mon_tbl(doc, [
        ("Authorised prescribers",
         "Only doctors holding MBBS minimum, or staff specifically authorised by "
         "legislation or government order, write medication orders."),
        ("Uniform order location",
         "All orders recorded in one location with patient name and unique ID; "
         "administration only from that location; no orders written elsewhere "
         "accepted without transfer."),
        ("Order legibility and traceability",
         "Orders written in capital letters; only approved abbreviations used; "
         "prescriber identity traceable on every order."),
        ("Order completeness",
         "Drug name, route, strength, and frequency present on every order; "
         "multi-drug orders state strength of each drug (vitamin/mineral exception "
         "applied correctly); mechanism in place for incomplete orders."),
        ("No shorthand or overwrite",
         "No CST, repeat-all, or similar shorthand; modifications written as fresh "
         "orders; no strike-through or overwrite accepted."),
    ])

    # 8. Training and staff acknowledgement
    h(doc, 1, "8. Training and staff acknowledgement")
    p(doc,
      "All doctors and authorised staff who write medication orders, and nursing "
      "staff who administer them, shall be familiar with the authorised-prescriber "
      "requirements, uniform-location requirement, legibility, traceability, and "
      "order-completeness rules in this policy.")
    p(doc,
      f"I have read the Policy on Uniform Medication Orders of {HN}. "
      "I will follow the processes described.")
    sig_tbl(doc)

    # 9. Distribution
    h(doc, 1, "9. Distribution")
    p(doc,
      "This policy shall be available to all treating doctors, nursing staff, the "
      "Medical Superintendent, the Medication Safety Officer, and the Quality "
      "Coordinator.")

    # 10. Abbreviations
    h(doc, 1, "10. Abbreviations")
    abbrev_tbl(doc, [
        ("EMR",  "Electronic Medical Record"),
        ("ISMP", "Institute for Safe Medication Practices"),
        ("MBBS", "Bachelor of Medicine, Bachelor of Surgery"),
        ("MOM",  "Management of Medication (NABH Hospitals chapter)"),
        ("NABH", "National Accreditation Board for Hospitals and Healthcare Providers"),
        ("OE",   "Objective Element"),
        ("OP",   "Outpatient"),
    ])

    # 11. Traceability table
    h(doc, 1, "11. Traceability table")
    p(doc,
      "This table is an index. It is not how the policy is organised. An asterisk "
      "in the Level column means documentation of the process is required.")
    tr = tbl(doc, 5, 3)
    for ci, hdr in enumerate(("Objective Element", "Level", "Traceability to this policy")):
        tr.cell(0, ci).text = hdr
    trace_rows = [
        ("MOM.5.a", "Commitment*",
         "Sections 3 and 5.1 address the MBBS-minimum qualification requirement, "
         "legislation or government-order basis for any other authorised category, "
         "the requirement for a doctor to write the inpatient order even when "
         "transcribing, and electronic-record direct-entry or doctor-verified "
         "assistant entry."),
        ("MOM.5.b", "Commitment",
         "Sections 3 and 5.2 address the single uniform location with patient name "
         "and unique ID, administration only from that location, transfer of orders "
         "written elsewhere, the ban on shorthand phrases (CST, repeat-all), and "
         "the fresh-order requirement for modifications — strike-through or overwrite "
         "of the old order is not acceptable."),
        ("MOM.5.c", "Commitment",
         "Section 5.3 addresses the capital-letter requirement for hand-written "
         "orders, the approved standardised abbreviation list applied throughout the "
         "organisation, the ban on error-prone abbreviations, and prescriber-identity "
         "traceability on every order."),
        ("MOM.5.d", "Commitment",
         "Section 5.4 addresses the multi-drug strength requirement — with the "
         "explicit exception for vitamin and/or mineral combinations only — separate "
         "orders for different strengths by time, and the mechanism for catching and "
         "acting on incomplete orders."),
    ]
    for ri, (oe, lvl, txt) in enumerate(trace_rows, 1):
        tr.cell(ri, 0).text = oe
        tr.cell(ri, 1).text = lvl
        tr.cell(ri, 2).text = txt

    # 12. Required Records/Evidence Checklist
    h(doc, 1, "12. Required Records/Evidence Checklist")

    h(doc, 2, "Authorised prescribers — MOM.5.a (Commitment*)")
    lb(doc, "Authorised-prescriber list showing MBBS minimum for doctors.")
    lb(doc,
       "Legislation or government order for any non-doctor category authorised "
       "to write medication orders.")
    lb(doc,
       "Electronic-record verification record where an assistant enters the "
       "prescription — showing doctor authorisation.")

    h(doc, 2, "Uniform order location — MOM.5.b")
    lb(doc,
       "Medication-order records showing uniform location containing patient name "
       "and unique identification number.")
    lb(doc,
       "Evidence that only medications from the uniform location are administered.")
    lb(doc,
       "Transfer records for orders originally written outside the uniform location.")
    lb(doc, "Absence of shorthand (CST, repeat-all) in current medication-order records.")
    lb(doc, "Fresh-order records for medication changes, with original discontinued.")

    h(doc, 2, "Legibility and traceability — MOM.5.c")
    lb(doc,
       "Sample medication orders written in capital letters with no error-prone "
       "abbreviations.")
    lb(doc, "Approved abbreviation list in use throughout the organisation.")
    lb(doc,
       "Prescriber-identity traceability on each order — name, master signature "
       "list, or employee code.")

    h(doc, 2, "Order completeness — MOM.5.d")
    lb(doc,
       "Sample orders showing drug name, route, strength, and frequency or time "
       "of administration present.")
    lb(doc,
       "Multi-drug orders with strength of each individual drug stated "
       "(vitamin/mineral exception documented where applied).")
    lb(doc, "Mechanism documentation for catching and acting on incomplete orders.")

    # 13. References
    h(doc, 1, "13. References")
    ln(doc,
       "National Accreditation Board for Hospitals and Healthcare Providers. NABH "
       "Accreditation Standards for Hospitals, 6th Edition. MOM.5.")
    ln(doc, "Guidebook interpretation supplied for MOM.5.a through MOM.5.d.")
    ln(doc,
       "Institute for Safe Medication Practices (ISMP) guidelines on error-prone "
       "abbreviations.")
    ln(doc,
       f"Internal documents of {HN}: authorised-prescriber list; approved "
       "abbreviation list; medication-order records.")

    # Disclaimer
    h(doc, 1, "Disclaimer")
    p(doc,
      "This policy reorganises the supplied MOM.5 objective-element wording and "
      "Guidebook interpretation into plain-language policy format. The modal strength "
      "of the source has been preserved. Optional examples and mechanisms have not "
      "been converted into mandatory requirements. The MBBS-minimum qualification "
      "requirement, the single uniform location with patient name and unique "
      "identification number, the prohibition on strike-through or overwrite, and "
      "the multi-drug strength exception limited to vitamin and/or mineral "
      "combinations only have all been retained verbatim.")

    save_and_verify(doc, "HCO_MOM_5_v2_REWRITE_DRAFT.docx")


def gen_mom6():
    """MOM.6 — Safe Dispensing of Medications (stop-work Section 6; sections 1-14 + Disclaimer)"""
    doc = Document()

    # Title
    h(doc, 0, "Policy on Safe Dispensing of Medications")
    p(doc, HN)

    # Document control
    h(doc, 1, "Document control")
    doc_ctrl(doc, "HCO/MOM/POL/06", "Pharmacy In-Charge")
    p(doc, "A blank marked ________ must be completed before issue.")

    # Statement of intent
    h(doc, 1, "Statement of intent")
    p(doc,
      "Medications are dispensed only against a valid prescription, checked before "
      "dispensing, correctly labelled, verified before a high-risk order is dispensed, "
      "and recalled or expired medications never reach a patient.")

    # 1. Purpose
    h(doc, 1, "1. Purpose")
    p(doc,
      f"This policy explains how {HN} dispenses medications safely, handles recalls "
      "and near-expiry stock, labels dispensed medications, verifies high-risk orders "
      "before dispensing, and manages medication returns.")
    p(doc,
      "This policy does not cover storage or prescription of medications in detail "
      "— those are covered in other hospital policies.")

    # 2. Scope
    h(doc, 1, "2. Scope")
    p(doc,
      f"This policy applies to all pharmacy staff involved in dispensing medications "
      f"at {HN}.")

    # 3. Policy standards
    h(doc, 1, "3. Policy standards")
    p(doc,
      f"{HN} dispenses medications only against a valid prescription or order, "
      "checked for generic composition, formulation, expiry, and strength, across "
      "both bulk and retail pharmacy. Medication recalls and near-expiry stock are "
      "handled through defined systems, every dispensed medication is labelled, "
      "high-risk orders are verified before dispensing, and medication returns "
      "follow written guidance.")
    p(doc, "Staff follow the written guidance below and keep the records it requires.")

    # 4. Non-negotiable rules
    h(doc, 1, "4. Non-negotiable rules")
    lb(doc,
       "Do not dispense a medication without a valid prescription or medication "
       "order, except for over-the-counter drugs, and do not dispense without "
       "checking generic composition, formulation, expiry date, and strength where "
       "applicable. Do not sell physicians' samples.")
    lb(doc,
       "Do not leave a recalled medication in usable stock, and do not skip "
       "reporting an internally identified recall to the appropriate regulatory "
       "authority.")
    lb(doc,
       "Do not let a beyond-expiry-date medication remain available in clinical "
       "stock.")
    lb(doc,
       "Do not dispense a medication that is unlabelled, recalled, expired, or a "
       "high-risk order that has not been verified for dose, frequency, and route.")
    lb(doc,
       "Do not accept a returned medication without written guidance on which "
       "medications are accepted and the minimum conditions they must meet.")

    # 5. What we do
    h(doc, 1, "5. What we do")

    h(doc, 2, "5.1 Dispense medications safely")
    p(doc,
      "Written guidance governs the safe dispensing of medications. Medications are "
      "dispensed only against a valid prescription or medication order, except for "
      "over-the-counter drugs. Before dispensing, medications are checked for generic "
      "composition, formulation, expiry date, and — where applicable — strength. "
      "This applies to both bulk and retail pharmacy. Physicians' samples are never "
      "sold.")

    h(doc, 2, "5.2 Handle medication recalls")
    p(doc,
      f"{HN} has an established system for medication recalls, whether triggered by "
      "a regulatory authority, the manufacturer, or internal feedback — for example, "
      "noticing a visible contaminant in an IV fluid bottle. Where a recall arises "
      "from internal feedback, the appropriate regulatory authority is also informed. "
      "A record is kept whenever a recall occurs.")

    h(doc, 2, "5.3 Handle near-expiry medications")
    p(doc,
      f"{HN} defines what counts as \"near expiry\" — for example, three months "
      "before the expiry date — and has a mechanism to withdraw near-expiry stock "
      "before it goes past that date. No beyond-expiry-date medication is available "
      "in usable clinical stock.")

    h(doc, 2, "5.4 Label dispensed medications")
    p(doc,
      "At minimum, every label includes dosage instructions the patient can "
      "understand. Where medicines are dispensed as cut strips or from bulk "
      "containers, the label also includes the drug name, strength, and expiry date. "
      "This applies to both inpatients and outpatients, and to reconstituted drugs, "
      f"such as chemotherapy medications. {HN} could use technology like QR codes "
      "on individual medicines to strengthen this process.")

    h(doc, 2, "5.5 Verify high-risk orders before dispensing")
    p(doc,
      "High-risk medications are dispensed only against a written order, verified "
      "by staff before dispensing, and in line with applicable statutory "
      "requirements.")

    h(doc, 2, "5.6 Manage medication returns")
    p(doc,
      "Written guidance directs how medications are returned to the pharmacy, at "
      "minimum covering which medications are accepted for return — defined by "
      "inclusion or exclusion — and the minimum conditions for accepting a return, "
      "such as matching drug name, strength, batch number, and expiry date to the "
      "bill, and no visible damage. It's preferable not to accept a return of any "
      "medication with a specific temperature storage requirement once it has left "
      "the hospital's premises.")

    # 6. Stop-work authority
    h(doc, 1, "6. Stop-work authority")
    p(doc,
      "Do not dispense a medication that is unlabelled, recalled, expired, or a "
      "high-risk order that has not been verified for dose, frequency and route.")
    p(doc,
      "Stop-work applies to the dispense. Immediate life-saving issue from floor "
      "stock in an emergency follows the documented after-hours or emergency-dispense "
      "rules and is recorded.")
    p(doc,
      "The person who stops tells the Pharmacy In-Charge the same shift. Refusing "
      "an unsafe dispense is not a disciplinary matter.")

    # 7. Governance and responsibility
    h(doc, 1, "7. Governance and responsibility")
    gov_tbl(doc, [
        ("Medical Superintendent",
         "Accountable for safe dispensing implementation across the organisation."),
        ("Pharmacy In-Charge",
         "Owns day-to-day dispensing, recall handling, labelling compliance, "
         "high-risk order verification, and medication returns. Receives stop-work "
         "escalations the same shift."),
        ("Medication Safety Officer",
         "Coordinates notification of internally identified recalls to the "
         "appropriate regulatory authority; brings dispensing incidents to the DTC."),
        ("Quality Coordinator",
         "Audits adherence to this policy; holds training records."),
    ])

    # 8. Quality monitoring
    h(doc, 1, "8. Quality monitoring")
    mon_tbl(doc, [
        ("Safe dispensing",
         "Valid prescription or order checked before every dispense; generic "
         "composition, formulation, expiry, and strength checked; applies to bulk "
         "and retail pharmacy; no physicians' samples sold."),
        ("Medication recalls",
         "Established recall system in place; internally identified recalls reported "
         "to the appropriate regulatory authority the same shift; records kept on "
         "occurrence."),
        ("Near-expiry and beyond-expiry stock",
         "Near-expiry threshold defined; withdrawal mechanism in place before the "
         "defined date; no beyond-expiry-date medication available in clinical stock."),
        ("Labelling",
         "All dispensed medications labelled with dosage instructions; cut strips "
         "and bulk-container labels include drug name, strength, and expiry date; "
         "applies to inpatients and outpatients; reconstituted drugs (e.g. "
         "chemotherapy) labelled."),
        ("High-risk order verification",
         "Written order present for every high-risk dispense; staff verification "
         "completed before dispensing; statutory requirements met."),
        ("Medication returns",
         "Written returns guidance in place; inclusion/exclusion list defined; "
         "minimum conditions documented; return records kept."),
        ("Stop-work events",
         "Stop-work events logged with trigger, action taken, and outcome."),
    ])

    # 9. Training and staff acknowledgement
    h(doc, 1, "9. Training and staff acknowledgement")
    p(doc,
      "All pharmacy staff involved in dispensing medications shall be familiar with "
      "the safe-dispensing guidance, recall and near-expiry procedures, labelling "
      "requirements, high-risk order verification, medication-return guidance, and "
      "stop-work authority in this policy.")
    p(doc,
      f"I have read the Policy on Safe Dispensing of Medications of {HN}. "
      "I will follow the processes described.")
    sig_tbl(doc)

    # 10. Distribution
    h(doc, 1, "10. Distribution")
    p(doc,
      "This policy shall be available to all pharmacy staff, the Pharmacy In-Charge, "
      "the Medical Superintendent, the Medication Safety Officer, and the Quality "
      "Coordinator.")

    # 11. Abbreviations
    h(doc, 1, "11. Abbreviations")
    abbrev_tbl(doc, [
        ("DTC",  "Drug and Therapeutics Committee"),
        ("IV",   "Intravenous"),
        ("MOM",  "Management of Medication (NABH Hospitals chapter)"),
        ("MSO",  "Medication Safety Officer"),
        ("NABH", "National Accreditation Board for Hospitals and Healthcare Providers"),
        ("OTC",  "Over-the-counter"),
        ("QR",   "Quick Response (code)"),
    ])

    # 12. Traceability table
    h(doc, 1, "12. Traceability table")
    p(doc,
      "This table is an index. It is not how the policy is organised. An asterisk "
      "in the Level column means documentation of the process is required.")
    tr = tbl(doc, 7, 3)
    for ci, hdr in enumerate(("Objective Element", "Level", "Traceability to this policy")):
        tr.cell(0, ci).text = hdr
    trace_rows = [
        ("MOM.6.a", "Commitment*",
         "Sections 3 and 5.1 address the written safe-dispensing guidance, valid "
         "prescription or order requirement, pre-dispense checks (generic "
         "composition, formulation, expiry, strength), applicability to bulk and "
         "retail pharmacy, and the prohibition on selling physicians' samples."),
        ("MOM.6.b", "Commitment*",
         "Section 5.2 addresses the established recall system, the requirement to "
         "inform the appropriate regulatory authority when a recall is internally "
         "identified, and the record-keeping requirement on every recall occurrence."),
        ("MOM.6.c", "Commitment*",
         "Section 5.3 addresses the definition of 'near expiry,' the withdrawal "
         "mechanism for near-expiry stock, and the absolute requirement that no "
         "beyond-expiry-date medication is available in usable clinical stock — "
         "near-expiry and beyond-expiry are treated as distinct stages."),
        ("MOM.6.d", "CORE*",
         "Sections 3 and 5.4 address universal labelling scope — both inpatients "
         "and outpatients, and reconstituted drugs such as chemotherapy medications. "
         "Cut strips and bulk containers carry drug name, strength, and expiry date "
         "in addition to dosage instructions. This resolves the Guidebook internal "
         "inconsistency in favour of universal scope (not outpatient-only). This OE "
         "is also a stop-work trigger in Section 6."),
        ("MOM.6.e", "CORE",
         "Sections 3 and 5.5 address the written-order requirement for high-risk "
         "medications, staff verification before dispensing, and statutory "
         "requirements. This is also a stop-work trigger in Section 6."),
        ("MOM.6.f", "Commitment*",
         "Section 5.6 addresses the written medication-return guidance, the "
         "inclusion/exclusion definition of accepted medications, and the minimum "
         "conditions for accepting a return."),
    ]
    for ri, (oe, lvl, txt) in enumerate(trace_rows, 1):
        tr.cell(ri, 0).text = oe
        tr.cell(ri, 1).text = lvl
        tr.cell(ri, 2).text = txt

    # 13. Required Records/Evidence Checklist
    h(doc, 1, "13. Required Records/Evidence Checklist")

    h(doc, 2, "Safe dispensing — MOM.6.a (Commitment*)")
    lb(doc, "Written safe-dispensing guidance in place.")
    lb(doc,
       "Dispense records showing valid prescription or order checked before "
       "every dispense.")
    lb(doc,
       "Pre-dispense check records covering generic composition, formulation, "
       "expiry date, and strength.")
    lb(doc, "Evidence the process applies to both bulk and retail pharmacy.")
    lb(doc, "Confirmation physicians' samples are not sold.")

    h(doc, 2, "Medication recalls — MOM.6.b (Commitment*)")
    lb(doc, "Established recall system documentation.")
    lb(doc,
       "Recall file with trigger, batches affected, quarantine action, and "
       "recovery steps.")
    lb(doc,
       "Regulatory-authority notification records for any internally identified "
       "recalls.")
    lb(doc, "Records kept on every recall occurrence.")

    h(doc, 2, "Near-expiry and beyond-expiry — MOM.6.c (Commitment*)")
    lb(doc, "Organisation's defined 'near-expiry' threshold (e.g. three months).")
    lb(doc, "Withdrawal records for near-expiry stock.")
    lb(doc,
       "Evidence that no beyond-expiry-date medication is available in usable "
       "clinical stock.")

    h(doc, 2, "Labelling — MOM.6.d (CORE*)")
    lb(doc,
       "Sample labelled dispensed medications showing dosage instructions the "
       "patient can understand.")
    lb(doc,
       "Cut-strip and bulk-container labels showing drug name, strength, expiry "
       "date, and dosage instructions.")
    lb(doc, "Evidence that labelling applies to both inpatients and outpatients.")
    lb(doc,
       "Evidence that labelling applies to reconstituted drugs (e.g. chemotherapy "
       "medications).")

    h(doc, 2, "High-risk order verification — MOM.6.e (CORE)")
    lb(doc, "Written-order records for every high-risk medication dispense.")
    lb(doc,
       "Staff-verification records before dispensing (second-person check or "
       "equivalent).")
    lb(doc, "Statutory compliance records where applicable.")

    h(doc, 2, "Medication returns — MOM.6.f (Commitment*)")
    lb(doc, "Written medication-return guidance.")
    lb(doc, "Inclusion/exclusion list of medications accepted for return.")
    lb(doc,
       "Minimum-conditions record (drug name, strength, batch, expiry matching "
       "bill; no visible damage).")
    lb(doc, "Return records.")

    # 14. References
    h(doc, 1, "14. References")
    ln(doc,
       "National Accreditation Board for Hospitals and Healthcare Providers. NABH "
       "Accreditation Standards for Hospitals, 6th Edition. MOM.6.")
    ln(doc, "Guidebook interpretation supplied for MOM.6.a through MOM.6.f.")
    ln(doc,
       f"Internal documents of {HN}: safe-dispensing guidance; recall system "
       "records; near-expiry withdrawal records; medication-return guidance.")

    # Disclaimer
    h(doc, 1, "Disclaimer")
    p(doc,
      "This policy reorganises the supplied MOM.6 objective-element wording and "
      "Guidebook interpretation into plain-language policy format. The modal strength "
      "of the source has been preserved. Optional examples and mechanisms have not "
      "been converted into mandatory requirements. The labelling scope has been "
      "resolved as universal — applying to both inpatients and outpatients and to "
      "reconstituted drugs — consistent with the operative statement in the Guidebook. "
      "The CORE stop-work trigger for unlabelled, recalled, expired, or unverified "
      "high-risk orders has been retained verbatim. The near-expiry versus "
      "beyond-expiry distinction, and the requirement to report internally identified "
      "recalls to the appropriate regulatory authority, have been retained verbatim.")

    save_and_verify(doc, "HCO_MOM_6_v2_REWRITE_DRAFT.docx")


# ══════════════════════════════════════════════════════════════════════════════
# MOM.7 — Safe Administration of Medications   (HAS stop-work: Section 6)
# Content: mom7_content.txt (approved).
# Structure: Document control, Sec 3 standards, Sec 4 non-negotiables,
#            Sec 5 (11 subsections), Sec 6 Stop-work, Sec 7 Governance,
#            Sec 8 Monitoring, Sec 9 Training, Sec 10 Distribution,
#            Sec 11 Abbreviations, Sec 12 Traceability, Sec 13 Records,
#            Sec 14 References, Disclaimer.
# COREs: d, h | Stars: h*, j*, k* | Achievement: none | Excellence: none
# Statute: YES (P2) — applicable laws governing who may administer medications.
# Exact items verified:
#   5.1: "registered nurse or a doctor with at minimum an MBBS qualification" exact
#   5.3: two identifiers — unique identification number + full name; bed-number alone excluded
#   5.4: five-parameter completeness check (name/strength/route/frequency/time) with
#        deferral — not adjustment — where any is missing; two-staff independent
#        documented check for high-risk medications
#   5.8: IV-extension-tube absolute prohibitions — epidurals/irrigation/drains/
#        central lines/enteric feeding tubes; line-trace before every connection
#   5.9: each dose documented separately — not batched at shift-end; actual
#        administration (not original order) reflected
#   Sec 6: four-trigger stop-work (identification / medication·strength·route·timing
#           verification / physical inspection / permitted-person) + tubing mis-
#           connection trigger; plain-text role names as specified
# ══════════════════════════════════════════════════════════════════════════════
def gen_mom7():
    """MOM.7 — Safe Administration of Medications (stop-work Section 6; sections 1-14 + Disclaimer)"""
    doc = Document()

    # Title
    h(doc, 0, "Policy on Safe Administration of Medications")
    p(doc, HN)

    # Document control
    h(doc, 1, "Document control")
    doc_ctrl(doc, "HCO/MOM/POL/07", "Nursing Superintendent")
    p(doc, "A blank marked ________ must be completed before issue.")

    # Statement of intent
    h(doc, 1, "Statement of intent")
    p(doc,
      "Medications are administered only by permitted personnel, only after the "
      "patient's identity and the medication order are fully verified, with tubing "
      "connections traced to prevent misconnection, and every dose documented accurately.")

    # 1. Purpose
    h(doc, 1, "1. Purpose")
    p(doc,
      f"This policy explains how {HN} ensures only permitted personnel administer "
      "medications, verifies patient identity and the order before every dose, "
      "prevents tubing and catheter mis-connections, documents every administration, "
      "and governs self-administration and outside medications.")
    p(doc,
      "This policy does not cover prescription, dispensing, or storage of medications "
      "in detail — those are covered in other hospital policies.")

    # 2. Scope
    h(doc, 1, "2. Scope")
    p(doc, f"This policy applies to all staff who administer medications at {HN}.")

    # 3. Policy standards
    h(doc, 1, "3. Policy standards")
    p(doc,
      f"{HN} ensures medications are administered only by those permitted by law, "
      "with the patient identified using at least two identifiers and the medication, "
      "strength, route, and timing verified against the order and physically inspected "
      "before every dose. Measures prevent catheter and tubing mis-connections. Every "
      "administration is documented in a uniform location, and where self-administration "
      "or outside medications are permitted, written guidance governs how.")
    p(doc, "Staff follow the written guidance below and keep the records it requires.")

    # 4. Non-negotiable rules
    h(doc, 1, "4. Non-negotiable rules")
    lb(doc,
       "Do not let anyone administer medication who isn't a registered nurse or doctor "
       "with at minimum an MBBS qualification, unless another staff category is "
       "specifically backed by legislation or a government order.")
    lb(doc,
       "Do not prepare a second drug before labelling the first one, wherever more "
       "than one drug is prepared and loaded.")
    lb(doc,
       "Do not administer medication without identifying the patient using at least "
       "two identifiers, one of which is the unique identification number and one of "
       "which is the patient's full name.")
    lb(doc,
       "Do not administer medication without verifying it against the order and "
       "physically inspecting it — checking general appearance and expiry date — and "
       "do not proceed if any order parameter (name, strength, route, frequency, or "
       "time) is missing or incomplete; defer instead, pending verification with the "
       "treating team.")
    lb(doc,
       "Do not treat a verbal confirmation of an order as anything other than a verbal "
       "order — follow the organisation's verbal-order process, including "
       "countersignature within 24 hours.")
    lb(doc,
       "Do not administer a high-risk medication without independent verification by "
       "at least two staff members, documented.")
    lb(doc,
       "Do not skip verifying strength, route, or timing against the order before "
       "administration — where a discrepancy is found, defer administration.")
    lb(doc,
       "Do not use an IV extension tube for an epidural, irrigation, a drain, a "
       "central line, or to extend an enteric feeding tube, and do not make any "
       "connection or reconnection without first tracing the line from its origin to "
       "the connection port.")
    lb(doc,
       "Do not document medication administration anywhere other than the uniform "
       "location, and do not batch-chart multiple doses at shift-end instead of "
       "documenting each dose separately as it happens.")
    lb(doc,
       "Do not permit patient self-administration, or accept medications brought from "
       "outside the organisation, without written guidance governing it.")

    # 5. What we do
    h(doc, 1, "5. What we do")

    h(doc, 2, "5.1 Administer medications only by permitted personnel")
    p(doc,
      "Only a registered nurse or a doctor with at minimum an MBBS qualification "
      "administers medication. Where another staff category is authorised to administer "
      "medication, that authorisation is backed by legislation or a government order, "
      f"not an informal internal decision. {HN} keeps an authorised-personnel list.")

    h(doc, 2, "5.2 Label medications when preparing more than one")
    p(doc,
      "Wherever more than one drug is prepared and loaded — for example anaesthetic "
      "drugs in the OT, or chemotherapy drugs — the first medication is labelled "
      "before the second is prepared. This applies throughout the organisation, not "
      "just in these example areas.")

    h(doc, 2, "5.3 Identify the patient before administering")
    p(doc,
      "At least two identifiers are used to confirm the patient before administration "
      "— one is the unique identification number (for example the hospital or IP "
      "number), and one is the patient's full name. Matching only by bed number, or "
      "relying on an attendant's confirmation alone, doesn't meet this requirement. "
      "(Related requirements are covered in the hospital's other policies.)")

    h(doc, 2, "5.4 Verify and inspect the medication before administering")
    p(doc,
      "Staff verify the medication order and physically inspect the medication — "
      "checking its general appearance (for example, unusual melting or clumping) "
      "and its expiry date — before administration. If any parameter of the order "
      "(name, strength, route, frequency, or time) is missing or incomplete, "
      "administration is deferred pending verification with the treating team. Any "
      "verbal confirmation obtained is treated as a verbal order, following the "
      "organisation's verbal-order process, including countersignature within 24 "
      "hours. For high-risk medications, verification is done independently by at "
      "least two staff members — nurse and nurse, or nurse and doctor — and "
      "documented. Nurses are trained on high-risk medications and empowered to "
      "raise concerns about a prescription while verifying it.")

    h(doc, 2, "5.5 Verify strength before administering")
    p(doc,
      "The strength of the medication is verified from the order before "
      "administration. Where a discrepancy is found, administration is deferred "
      "rather than adjusted at the bedside. (Related requirements are covered in "
      "the hospital's other policies.)")

    h(doc, 2, "5.6 Verify route before administering")
    p(doc,
      "Where applicable, the site or route of administration is verified from the "
      "order before administration. (Related requirements are covered in the "
      "hospital's other policies.)")

    h(doc, 2, "5.7 Verify timing before administering")
    p(doc,
      f"Where a medication order doesn't state an explicit administration time — "
      "for example, standard frequency codes like once-daily-three-times or "
      f"twice-daily — {HN} maintains documentation supporting the correct timing, "
      "and these suggested timings are followed. The ISMP Acute Care Guidelines "
      "for Timely Administration of Scheduled Medications, which classify medications "
      "as time-critical or non-time-critical, are a useful reference "
      f"{HN} could adopt or adapt.")

    h(doc, 2, "5.8 Prevent catheter and tubing mis-connections")
    p(doc,
      f"{HN} ensures inadvertent administration of a drug through the wrong route "
      "is avoided. IV extension tubes are never used for epidurals, irrigation, "
      "drains, or central lines, and never used to extend enteric feeding tubes. "
      "Functionally dissimilar tubes used in patient care are positioned away from "
      "one another. Staff trace every line from its origin to the connection port "
      "to verify the correct attachment before making any connection or reconnection, "
      "or before administering any medication, solution, or other product.")

    h(doc, 2, "5.9 Document every administration")
    p(doc,
      "Documentation of medication administration happens in a uniform location. "
      "It includes the medication's name, strength, route, timing, and the name "
      "(or employee ID) and signature of the person who administered it. Each dose "
      "is documented separately — not batched together at shift-end. Documentation "
      "reflects what was actually administered, not what was originally ordered — "
      "for example, if a different brand of the same generic drug was given, or "
      "half a 500 mg tablet was given in place of a 250 mg tablet, the record "
      "shows what was actually given. For infusions, the start time, rate or volume, "
      "and end time are captured; for continuous infusions, the drop rate or volume "
      "is documented and the total volume infused calculated for each shift.")

    h(doc, 2, "5.10 Govern patient self-administration")
    p(doc,
      f"{HN} decides, and documents, whether it permits patients to self-administer "
      "medications — for example, self-administration of insulin. Where permitted, "
      "written guidance defines which medications a patient may self-administer, "
      "it's preferable to also have a method for assisting self-administration, and "
      "the patient is reminded to take each dose, with reminders documented.")

    h(doc, 2, "5.11 Govern medications brought from outside")
    p(doc,
      f"{HN} decides whether it permits patients to use medications they've brought "
      "from outside the organisation, and implements the necessary measures either "
      "way. Where permitted, written guidance sets out the pre-requisites such "
      "medications must meet — for example, a clear label showing name, strength, "
      "expiry date, and batch number.")

    # 6. Stop-work authority
    h(doc, 1, "6. Stop-work authority")
    p(doc,
      "Do not administer a medication when the patient has not been identified with "
      "the organisation's identifiers, when the medication, strength, route, or "
      "timing has not been verified against the order, when the product fails "
      f"physical inspection, or when you are not on the list of persons permitted "
      f"by law and by {HN} to administer that medication.")
    p(doc,
      "Do not connect a catheter or tubing for medication administration until the "
      "line has been traced from the patient to the source.")
    p(doc,
      "Stop-work applies to the administration start. Immediate life-saving "
      "administration in a crash continues with the best available permitted staff "
      "and is documented afterward.")
    p(doc,
      "The person who stops tells the Nurse In-Charge and the treating doctor the "
      "same shift. Refusing an unsafe administration is not a disciplinary matter.")

    # 7. Governance and responsibility
    h(doc, 1, "7. Governance and responsibility")
    gov_tbl(doc, [
        ("Medical Superintendent",
         "Accountable for ensuring safe medication administration is resourced and "
         "implemented; maintains the authorised-administration-personnel list "
         "together with the Nursing Superintendent."),
        ("Nursing Superintendent",
         "Owns day-to-day implementation of this policy; maintains the authorised-"
         "administration-personnel list; ensures identification, verification, "
         "documentation, and line-tracing practices are followed; receives stop-work "
         "escalations."),
        ("Medication Safety Officer",
         "Coordinates medication-safety processes; brings administration incidents "
         "and audit findings to the Drug and Therapeutics Committee."),
        ("Treating doctors",
         "Provide clear, complete medication orders; countersign verbal orders "
         "within 24 hours; respond to stop-work escalations from administering staff."),
        ("Pharmacy In-Charge",
         "Ensures medications are correctly labelled and verified before reaching "
         "the administering nurse or doctor; supports the high-risk medication "
         "verification process."),
        ("Quality Coordinator",
         "Audits this policy; holds training records and staff acknowledgements."),
    ])

    # 8. Quality monitoring
    h(doc, 1, "8. Quality monitoring")
    mon_tbl(doc, [
        ("Authorised administration personnel",
         "Only registered nurses, MBBS-minimum doctors, or legislation-backed staff "
         "administer medications; authorised-personnel list current."),
        ("Medication preparation labelling",
         "First medication labelled before a second is prepared wherever more than "
         "one drug is prepared and loaded; applies throughout the organisation."),
        ("Patient identification",
         "At least two identifiers — unique identification number plus full name — "
         "confirmed before every administration; bed number or attendant alone "
         "not accepted."),
        ("Medication verification and physical inspection",
         "Order verified (name, strength, route, frequency, time) and medication "
         "physically inspected (appearance and expiry) before every dose; "
         "administration deferred when any parameter is missing or incomplete; "
         "two-staff independent documented check for high-risk medications."),
        ("Strength, route, and timing verification",
         "Strength, route (where applicable), and timing verified from the order "
         "before every administration; discrepancies result in deferral."),
        ("Catheter and tubing mis-connection prevention",
         "IV extension tubes never used for epidurals, irrigation, drains, central "
         "lines, or enteric feeding tubes; line traced before every connection or "
         "reconnection; training records current."),
        ("Administration documentation",
         "Documentation in a uniform location for every dose — name, strength, "
         "route, timing, administrator identity; each dose documented separately, "
         "not batched; actual administration reflected, not original order alone "
         "where different; infusion records complete."),
        ("Patient self-administration",
         "Decision documented; written guidance in place where permitted; patient "
         "reminders documented per dose."),
        ("Outside medications",
         "Decision documented; written pre-requisites guidance in place where "
         "permitted."),
        ("Stop-work events",
         "Stop-work events logged with trigger, action taken, and outcome."),
    ])

    # 9. Training and staff acknowledgement
    h(doc, 1, "9. Training and staff acknowledgement")
    p(doc,
      "All staff who administer medications shall be familiar with the authorised-"
      "personnel requirements, patient identification, medication verification and "
      "inspection, line-tracing and IV-extension-tube prohibition, administration "
      "documentation, and stop-work authority in this policy.")
    p(doc,
      f"I have read the Policy on Safe Administration of Medications of {HN}. "
      "I will follow the processes described.")
    sig_tbl(doc)

    # 10. Distribution
    h(doc, 1, "10. Distribution")
    p(doc,
      "This policy shall be available to all staff who administer medications, "
      "the Nursing Superintendent, treating doctors, pharmacy staff, the Medication "
      "Safety Officer, and the Quality Coordinator.")

    # 11. Abbreviations
    h(doc, 1, "11. Abbreviations")
    abbrev_tbl(doc, [
        ("ADR",  "Adverse drug reaction"),
        ("CORE", "Core objective element (NABH)"),
        ("ICU",  "Intensive Care Unit"),
        ("ISMP", "Institute for Safe Medication Practices"),
        ("IV",   "Intravenous"),
        ("MBBS", "Bachelor of Medicine, Bachelor of Surgery"),
        ("MOM",  "Management of Medication (NABH Hospitals chapter)"),
        ("NABH", "National Accreditation Board for Hospitals and Healthcare Providers"),
        ("OT",   "Operation Theatre"),
    ])

    # 12. Traceability table
    h(doc, 1, "12. Traceability table")
    p(doc,
      "This table is an index. It is not how the policy is organised. An asterisk "
      "in the Level column means documentation of the process is required.")
    tr = tbl(doc, 12, 3)
    for ci, hdr in enumerate(("Objective Element", "Level", "Traceability to this policy")):
        tr.cell(0, ci).text = hdr
    trace_rows = [
        ("MOM.7.a", "Commitment",
         "Sections 3 and 5.1 address the registered-nurse or MBBS-minimum-doctor "
         "requirement and the legislation-or-government-order basis for any other "
         "authorised category. This is also a stop-work trigger in Section 6."),
        ("MOM.7.b", "Commitment",
         "Section 5.2 addresses the labelling requirement before preparing a second "
         "drug, wherever more than one drug is prepared and loaded."),
        ("MOM.7.c", "Commitment",
         "Sections 3 and 5.3 address two-identifier patient identification — unique "
         "identification number and full name — before every administration; bed-"
         "number matching alone is explicitly excluded. This is also a stop-work "
         "trigger in Section 6."),
        ("MOM.7.d", "CORE",
         "Sections 3 and 5.4 address medication-order verification and physical "
         "inspection before administration, the five-parameter completeness check "
         "(name, strength, route, frequency, time) with deferral — not adjustment "
         "— where any parameter is missing, and two-staff independent documented "
         "verification for high-risk medications. This is also a stop-work trigger "
         "in Section 6."),
        ("MOM.7.e", "Commitment",
         "Section 5.5 addresses strength verification from the order before "
         "administration, and deferral — not bedside adjustment — where a "
         "discrepancy is found."),
        ("MOM.7.f", "Commitment",
         "Section 5.6 addresses route (and site, where applicable) verification "
         "from the order before administration."),
        ("MOM.7.g", "Commitment",
         "Section 5.7 addresses timing verification — including documentation "
         "supporting the correct time where the order uses a frequency code only."),
        ("MOM.7.h", "CORE*",
         "Sections 3 and 5.8 address the absolute prohibition on using IV extension "
         "tubes for epidurals, irrigation, drains, central lines, or enteric feeding "
         "tubes, the line-trace requirement before every connection or reconnection, "
         "and positioning of dissimilar tubes away from each other. This is also a "
         "stop-work trigger in Section 6."),
        ("MOM.7.i", "Commitment",
         "Section 5.9 addresses documentation of every administration in a uniform "
         "location, each dose separately (not batched), reflecting actual "
         "administration, with full infusion detail."),
        ("MOM.7.j", "Commitment*",
         "Section 5.10 addresses the decision on whether patient self-administration "
         "is permitted, written guidance where it is, and documented dose reminders."),
        ("MOM.7.k", "Commitment*",
         "Section 5.11 addresses the decision on whether outside medications are "
         "permitted, and written pre-requisites guidance where they are."),
    ]
    for ri, (oe, lvl, txt) in enumerate(trace_rows, 1):
        tr.cell(ri, 0).text = oe
        tr.cell(ri, 1).text = lvl
        tr.cell(ri, 2).text = txt

    # 13. Required Records/Evidence Checklist
    h(doc, 1, "13. Required Records/Evidence Checklist")

    h(doc, 2, "Authorised administration personnel — MOM.7.a")
    lb(doc,
       "Authorised-administration-personnel list held by the Nursing Superintendent "
       "and Medical Superintendent, naming registered nurses, MBBS-minimum doctors, "
       "and any other category with its legislative or government-order backing.")
    lb(doc, "Documented supervision records where student staff administer under supervision.")
    lb(doc, "Confirmation no unlisted person administered a medication.")

    h(doc, 2, "Medication preparation labelling — MOM.7.b")
    lb(doc,
       "Observation or audit records confirming the first medication is labelled "
       "before a second drug is prepared, across OT, ICU, emergency, and wards.")
    lb(doc,
       "Training records on the labelling rule for areas where more than one drug "
       "is prepared and loaded.")

    h(doc, 2, "Patient identification — MOM.7.c")
    lb(doc,
       "Administration records showing at least two identifiers used — unique "
       "identification number and full name — before each administration.")
    lb(doc, "Training records confirming bed-number-only matching is not accepted.")
    lb(doc, "Records of administrations stopped where identity could not be confirmed.")

    h(doc, 2, "Medication verification and physical inspection — MOM.7.d (CORE)")
    lb(doc,
       "Pre-administration verification records showing medication matched against "
       "order and physically inspected for appearance and expiry.")
    lb(doc,
       "Deferral records where any order parameter (name, strength, route, frequency, "
       "or time) was missing or incomplete, pending treating-team verification.")
    lb(doc,
       "Two-staff independent documented verification records for high-risk "
       "medications before administration.")
    lb(doc,
       "Training records confirming nurses are trained on high-risk medications "
       "and empowered to raise concerns.")

    h(doc, 2, "Strength, route, and timing verification — MOM.7.e, MOM.7.f, MOM.7.g")
    lb(doc,
       "Strength verification records from order before administration; deferral "
       "records for discrepancies.")
    lb(doc, "Route verification records (where applicable) before administration.")
    lb(doc,
       "Timing documentation showing correct administration time where the order "
       "uses a frequency code only.")

    h(doc, 2, "Catheter and tubing mis-connection prevention — MOM.7.h (CORE*)")
    lb(doc,
       "Written measures in place describing the IV-extension-tube prohibition "
       "(epidurals, irrigation, drains, central lines, enteric feeding tubes) and "
       "the line-trace requirement before every connection or reconnection.")
    lb(doc,
       "Line-tracing records or observation evidence before connections for IV, "
       "epidural, and enteral medication administration.")
    lb(doc,
       "Training records for ICU, OT, emergency, and ward staff on catheter and "
       "tubing mis-connection prevention.")

    h(doc, 2, "Administration documentation — MOM.7.i")
    lb(doc,
       "Medication chart or equivalent uniform-location records showing name, "
       "strength, route, timing, and administrator identity for each dose.")
    lb(doc,
       "Confirmation each dose is documented separately — no shift-end batch-charting.")
    lb(doc,
       "Records reflecting actual administration where it differed from the original "
       "order (brand substitution, half-tablet, infusion rate or volume adjustment).")
    lb(doc,
       "Infusion records showing start time, rate or volume, and end time; "
       "continuous-infusion records showing drop rate and total volume per shift.")

    h(doc, 2, "Patient self-administration — MOM.7.j (Commitment*)")
    lb(doc,
       "Documented organisational decision on whether patient self-administration "
       "is permitted.")
    lb(doc,
       "Written guidance defining which medications a patient may self-administer, "
       "where permitted.")
    lb(doc, "Dose-reminder documentation records, where self-administration is permitted.")

    h(doc, 2, "Outside medications — MOM.7.k (Commitment*)")
    lb(doc,
       "Documented organisational decision on whether patients may bring their own "
       "medications.")
    lb(doc,
       "Written pre-requisites guidance (for example, clear label showing name, "
       "strength, expiry date, and batch number), where outside medications are "
       "permitted.")
    lb(doc,
       "Records showing outside medications identified by pharmacy or the treating "
       "doctor before use.")

    # 14. References
    h(doc, 1, "14. References")
    ln(doc,
       "National Accreditation Board for Hospitals and Healthcare Providers. NABH "
       "Accreditation Standards for Hospitals, 6th Edition. MOM.7.")
    ln(doc, "Guidebook interpretation supplied for MOM.7.a through MOM.7.k.")
    ln(doc,
       "Institute for Safe Medication Practices (ISMP). Acute Care Guidelines for "
       "Timely Administration of Scheduled Medications.")
    ln(doc,
       f"Internal documents of {HN}: authorised-administration-personnel list; "
       "high-risk medication list; medication administration records; line-tracing "
       "procedures; self-administration and outside-medication guidance.")

    # Disclaimer
    h(doc, 1, "Disclaimer")
    p(doc,
      "This policy reorganises the supplied MOM.7 objective-element wording and "
      "Guidebook interpretation into plain-language policy format. The modal strength "
      "of the source has been preserved. Optional examples and mechanisms have not "
      "been converted into mandatory requirements. The exact requirements of the "
      "registered-nurse or MBBS-minimum-doctor qualification, two-identifier patient "
      "identification (unique identification number plus full name; bed-number alone "
      "explicitly excluded), the five-parameter completeness check (name, strength, "
      "route, frequency, time) with deferral — not adjustment — for any missing "
      "parameter, two-staff independent documented verification for high-risk "
      "medications (nurse-nurse or nurse-doctor), the absolute IV-extension-tube "
      "prohibitions (epidurals, irrigation, drains, central lines, enteric feeding "
      "tubes), line-tracing before every connection or reconnection, each-dose "
      "documentation separately and not batched at shift-end, and the stop-work "
      "text with all four administration triggers (patient identification, "
      "medication/strength/route/timing verification, physical inspection, permitted-"
      "person check) plus the tubing mis-connection trigger have been retained "
      "verbatim. Role names in the stop-work section are in plain text as specified.")

    save_and_verify(doc, "HCO_MOM_7_v2_REWRITE_DRAFT.docx")


# ══════════════════════════════════════════════════════════════════════════════
# MOM.8 — Monitoring after Medication Administration   (NO stop-work)
# Content: mom8_content.txt (approved).
# Structure: Document control, Sec 3 standards, Sec 4 non-negotiables,
#            Sec 5 (6 subsections), Sec 6 Governance, Sec 7 Monitoring,
#            Sec 8 Training, Sec 9 Distribution, Sec 10 Abbreviations,
#            Sec 11 Traceability, Sec 12 Records, Sec 13 References, Disclaimer.
# COREs: c | Stars: a*, c*, d* | Achievement: none | Excellence: none
# Exact items verified:
#   5.1: three named high-risk patient categories — dialysis / ICU / elderly —
#        retained as examples, not exhaustive list
#   5.3: five mandatory capture steps — identifying / documenting / reporting /
#        analysing / acting on — treated as exhaustive mandatory list, not a pick-list
#   5.4: reporting timeframe = organisation defines it then adheres to it;
#        NO specific hours or days substituted
#   MOM.8.c CORE does NOT create a stop-work section (confirmed: MOM.8 not in
#        stop-work proposals)
# ══════════════════════════════════════════════════════════════════════════════
def gen_mom8():
    """MOM.8 — Monitoring after Medication Administration (no stop-work; sections 1-13 + Disclaimer)"""
    doc = Document()

    # Title
    h(doc, 0, "Policy on Monitoring after Medication Administration")
    p(doc, HN)

    # Document control
    h(doc, 1, "Document control")
    doc_ctrl(doc, "HCO/MOM/POL/08", "Medication Safety Officer")
    p(doc, "A blank marked ________ must be completed before issue.")

    # Statement of intent
    h(doc, 1, "Statement of intent")
    p(doc,
      "Every patient is monitored after medication administration to confirm the "
      "intended effect, high-risk patients are monitored regularly, and near misses, "
      "medication errors, and adverse drug reactions are captured, reported, analysed, "
      "and acted on.")

    # 1. Purpose
    h(doc, 1, "1. Purpose")
    p(doc,
      f"This policy explains how {HN} monitors patients after medication "
      "administration, adjusts medication based on that monitoring, and captures, "
      "reports, analyses, and acts on near misses, medication errors, and adverse "
      "drug reactions.")
    p(doc,
      "This policy does not cover medication administration or prescription in detail "
      "— those are covered in other hospital policies.")

    # 2. Scope
    h(doc, 1, "2. Scope")
    p(doc,
      f"This policy applies to all clinical staff involved in monitoring patients "
      f"after medication administration at {HN}.")

    # 3. Policy standards
    h(doc, 1, "3. Policy standards")
    p(doc,
      f"{HN} monitors patients collaboratively after medication administration to "
      "confirm the intended effect, defines which situations and medications need "
      "closer monitoring, monitors high-risk patients — including those on dialysis, "
      "in the ICU, and the elderly — regularly, and changes medication based on "
      "monitoring where appropriate. Near misses, medication errors, and adverse "
      "drug reactions are defined, captured under written guidance, reported within "
      "a defined time frame, analysed by the multidisciplinary committee, and "
      "followed up with corrective or preventive action where appropriate.")
    p(doc, "Staff follow the written guidance below and keep the records it requires.")

    # 4. Non-negotiable rules
    h(doc, 1, "4. Non-negotiable rules")
    lb(doc,
       "Do not skip monitoring a patient after medication administration, and do not "
       "leave situations and medications requiring more frequent monitoring undefined.")
    lb(doc,
       "Do not leave high-risk patients — those on dialysis, in the ICU, or the "
       "elderly — without regular monitoring of medication effects.")
    lb(doc,
       "Do not operate without written guidance defining near misses, medication "
       "errors, and adverse drug reactions, and without a defined process for "
       "identifying, documenting, reporting, analysing, and acting on them.")
    lb(doc,
       "Do not leave the reporting time frame for near misses, medication errors, "
       "or adverse drug reactions undefined, and do not miss the defined time frame "
       "once it's set.")
    lb(doc,
       "Do not skip multidisciplinary committee analysis of collected near misses, "
       "medication errors, and adverse drug reactions, or let that analysis run "
       "without a defined completion time frame.")

    # 5. What we do
    h(doc, 1, "5. What we do")

    h(doc, 2, "5.1 Monitor patients after medication administration")
    p(doc,
      "Monitoring after administration is done collaboratively, to confirm the "
      "medicine is having its intended effect — this could include monitoring "
      "through laboratory results. Monitoring also helps identify near misses, "
      f"medication errors, and adverse drug reactions. {HN} defines the situations "
      "and medications that call for more frequent monitoring — for example, "
      "administration of high-risk medicines. Medication effects in high-risk "
      "patients — those on dialysis, in the ICU, and the elderly — are monitored "
      "on a regular basis; these are examples of high-risk groups, not an "
      "exhaustive list.")

    h(doc, 2, "5.2 Change medication based on monitoring")
    p(doc,
      "Where appropriate, medications are changed based on clinical response and "
      "any adverse drug reactions identified through monitoring.")

    h(doc, 2, "5.3 Capture near misses, medication errors, and adverse drug reactions")
    p(doc,
      f"Near misses, medication errors, and adverse drug reactions are defined by "
      f"{HN}, in line with best practice. Written guidance sets out the process for "
      "capturing them, covering five steps: identifying, documenting, reporting, "
      "analysing, and taking action.")

    h(doc, 2, "5.4 Report within a defined time frame")
    p(doc,
      f"{HN} defines the time frame for reporting a near miss, medication error, "
      "or adverse drug reaction once it occurs, and adheres to that time frame.")

    h(doc, 2, "5.5 Collect and analyse")
    p(doc,
      "Details of near misses, medication errors, and adverse drug reactions are "
      "collected and analysed by the multidisciplinary committee, within a defined "
      "time frame. It's preferable for a clinical pharmacologist or clinical "
      "pharmacist to take part in this analysis.")

    h(doc, 2, "5.6 Take corrective and preventive action")
    p(doc,
      "Where appropriate, corrective or preventive action is taken based on the "
      "analysis, and records of these actions are kept. It's preferable for "
      "corrective and preventive action to be based on root-cause analysis.")

    # 6. Governance and responsibility
    h(doc, 1, "6. Governance and responsibility")
    gov_tbl(doc, [
        ("Medical Superintendent",
         "Accountable for ensuring post-administration monitoring and the "
         "near-miss/error/ADR capture system are resourced and implemented."),
        ("Medication Safety Officer",
         "Owns day-to-day implementation of this policy; coordinates the capture, "
         "reporting, and analysis of near misses, medication errors, and ADRs; "
         "brings findings to the multidisciplinary committee."),
        ("Multidisciplinary committee (DTC)",
         "Analyses near misses, medication errors, and ADRs within the defined "
         "time frame; takes or approves corrective and preventive action."),
        ("Treating doctors / clinical staff",
         "Monitor patients collaboratively after medication administration; change "
         "medications based on monitoring where appropriate; report near misses, "
         "medication errors, and ADRs within the defined time frame."),
        ("Nursing staff",
         "Monitor patients after administration; capture and report near misses, "
         "medication errors, and ADRs per the defined process and time frame."),
        ("Quality Coordinator",
         "Audits this policy; holds training records and staff acknowledgements."),
    ])

    # 7. Quality monitoring
    h(doc, 1, "7. Quality monitoring")
    mon_tbl(doc, [
        ("Post-administration monitoring",
         "Done collaboratively; situations and medications needing more frequent "
         "monitoring defined; high-risk patients (those on dialysis, in the ICU, "
         "and the elderly) monitored regularly; monitoring records present."),
        ("Medication changes from monitoring",
         "Changes based on clinical response and ADRs where appropriate; "
         "rewritten orders present."),
        ("Near-miss/error/ADR definitions and capture guidance",
         "Definitions in place and in line with best practice; written guidance "
         "covering all five steps (identifying, documenting, reporting, analysing, "
         "acting on); MOM.8.c is CORE — absence of definitions or a working "
         "capture system is accreditation-blocking; no stop-work section applies."),
        ("Reporting time frame",
         "Time frame defined by the organisation; adherence evidenced by comparing "
         "event date to report date on a sample of incident forms."),
        ("Multidisciplinary committee analysis",
         "Analysis conducted within the defined time frame; DTC record shows "
         "collection and analysis, not just individual review."),
        ("Corrective and preventive action",
         "CAPA records maintained where appropriate; open actions tracked."),
    ])

    # 8. Training and staff acknowledgement
    h(doc, 1, "8. Training and staff acknowledgement")
    p(doc,
      "All clinical staff involved in monitoring patients after medication "
      "administration shall be familiar with the monitoring requirements, the "
      "near-miss/error/ADR definitions, the capture and reporting process, the "
      "defined reporting time frame, and the roles in analysis and action in "
      "this policy.")
    p(doc,
      f"I have read the Policy on Monitoring after Medication Administration of "
      f"{HN}. I will follow the processes described.")
    sig_tbl(doc)

    # 9. Distribution
    h(doc, 1, "9. Distribution")
    p(doc,
      "This policy shall be available to all clinical staff who monitor patients "
      "after medication administration, the Medication Safety Officer, the "
      "multidisciplinary committee, and the Quality Coordinator.")

    # 10. Abbreviations
    h(doc, 1, "10. Abbreviations")
    abbrev_tbl(doc, [
        ("ADR",  "Adverse drug reaction"),
        ("CAPA", "Corrective and Preventive Action"),
        ("CORE", "Core objective element (NABH)"),
        ("DTC",  "Drug and Therapeutics Committee (the organisation's "
                  "multidisciplinary medication management committee)"),
        ("ICU",  "Intensive Care Unit"),
        ("MOM",  "Management of Medication (NABH Hospitals chapter)"),
        ("MSO",  "Medication Safety Officer"),
        ("NABH", "National Accreditation Board for Hospitals and Healthcare Providers"),
    ])

    # 11. Traceability table
    h(doc, 1, "11. Traceability table")
    p(doc,
      "This table is an index. It is not how the policy is organised. An asterisk "
      "in the Level column means documentation of the process is required.")
    tr = tbl(doc, 7, 3)
    for ci, hdr in enumerate(("Objective Element", "Level", "Traceability to this policy")):
        tr.cell(0, ci).text = hdr
    trace_rows = [
        ("MOM.8.a", "Commitment*",
         "Sections 3 and 5.1 address collaborative post-administration monitoring, "
         "the organisation's definition of situations and medications requiring more "
         "frequent monitoring, and the three named high-risk patient categories: "
         "those on dialysis, in the ICU, and the elderly."),
        ("MOM.8.b", "Commitment",
         "Section 5.2 addresses medication changes based on clinical response and "
         "adverse drug reactions identified through monitoring, where appropriate."),
        ("MOM.8.c", "CORE*",
         "Sections 3 and 5.3 address the definition of near misses, medication "
         "errors, and adverse drug reactions in line with best practice, and written "
         "guidance covering all five mandatory capture steps: identifying, documenting, "
         "reporting, analysing, and acting on them. MOM.8.c carries CORE status — "
         "this policy does not contain a stop-work section, which is correct, as "
         "MOM.8 is not in the MOM stop-work proposals."),
        ("MOM.8.d", "Commitment*",
         "Sections 3 and 5.4 address the definition of the reporting time frame and "
         "the requirement to adhere to it. The specific time frame is set by "
         f"{HN} — this policy does not substitute a specific number of hours or days."),
        ("MOM.8.e", "Commitment",
         "Section 5.5 addresses collection and analysis by the multidisciplinary "
         "committee within a defined time frame, with a preference for clinical "
         "pharmacologist or pharmacist involvement."),
        ("MOM.8.f", "Commitment",
         "Section 5.6 addresses corrective or preventive action taken where "
         "appropriate, with records kept, and a preference for root-cause analysis "
         "as the basis."),
    ]
    for ri, (oe, lvl, txt) in enumerate(trace_rows, 1):
        tr.cell(ri, 0).text = oe
        tr.cell(ri, 1).text = lvl
        tr.cell(ri, 2).text = txt

    # 12. Required Records/Evidence Checklist
    h(doc, 1, "12. Required Records/Evidence Checklist")

    h(doc, 2, "Post-administration monitoring — MOM.8.a (Commitment*)")
    lb(doc,
       "Written monitoring guidance naming what to watch for and when to escalate, "
       "matched to the drug and clinical setting.")
    lb(doc,
       "Record of situations and medications where more frequent monitoring is "
       "required, as defined by the organisation.")
    lb(doc,
       "Recorded monitoring for a sample of patients — including high-risk patients "
       "(those on dialysis, in the ICU, and the elderly) — after medication "
       "administration.")

    h(doc, 2, "Medication changes from monitoring — MOM.8.b")
    lb(doc,
       "Medication-change records based on monitoring findings — drug changed, held, "
       "or dose adjusted based on clinical response or ADR.")
    lb(doc,
       "Rewritten-order records for any medication change.")
    lb(doc,
       "Confirmation the treating doctor was informed and the situation was not "
       "silently continued.")

    h(doc, 2, "Near-miss/error/ADR capture — MOM.8.c (CORE*)")
    lb(doc,
       "Written definitions of near misses, medication errors, and adverse drug "
       "reactions, in line with best practice.")
    lb(doc,
       "Written guidance setting out the five mandatory capture steps: identifying, "
       "documenting, reporting, analysing, and acting on.")
    lb(doc,
       "Sample of captured near-miss and incident reports demonstrating the process "
       "is working.")

    h(doc, 2, "Reporting time frame — MOM.8.d (Commitment*)")
    lb(doc,
       "Defined reporting time frame for near misses, medication errors, and adverse "
       "drug reactions — as set by the organisation.")
    lb(doc,
       "Adherence evidence: comparison of event date/time to report date/time on a "
       "sample of incident forms.")

    h(doc, 2, "Multidisciplinary committee analysis — MOM.8.e")
    lb(doc,
       "DTC analysis records showing details of near misses, medication errors, and "
       "ADRs collected and analysed as a committee.")
    lb(doc,
       "Defined analysis completion time frame, evidenced in committee minutes or "
       "terms of reference.")

    h(doc, 2, "Corrective and preventive action — MOM.8.f")
    lb(doc,
       "CAPA records where corrective or preventive action was taken — with owner "
       "and due date.")
    lb(doc, "Open-action tracking records.")
    lb(doc,
       "Decision record where CAPA was considered but judged not required, where "
       "that is the case.")

    # 13. References
    h(doc, 1, "13. References")
    ln(doc,
       "National Accreditation Board for Hospitals and Healthcare Providers. NABH "
       "Accreditation Standards for Hospitals, 6th Edition. MOM.8.")
    ln(doc, "Guidebook interpretation supplied for MOM.8.a through MOM.8.f.")
    ln(doc,
       f"Internal documents of {HN}: post-administration monitoring guidance; "
       "near-miss/error/ADR definitions and written guidance; incident reporting "
       "system; DTC analysis records; CAPA records.")

    # Disclaimer
    h(doc, 1, "Disclaimer")
    p(doc,
      "This policy reorganises the supplied MOM.8 objective-element wording and "
      "Guidebook interpretation into plain-language policy format. The modal strength "
      "of the source has been preserved. Optional examples and mechanisms have not "
      "been converted into mandatory requirements. The three named high-risk patient "
      "categories (those on dialysis, in the ICU, and the elderly) have been retained "
      "verbatim as examples of high-risk groups, not converted to an exhaustive "
      "mandatory list. The five mandatory capture-process steps (identifying, "
      "documenting, reporting, analysing, and acting on) have been retained as an "
      "exhaustive mandatory list, not a pick-list. The reporting-timeframe requirement "
      "is stated as the organisation defines it and then adheres to it — no specific "
      "number of hours or days has been substituted. MOM.8.c carries CORE status; "
      "this policy does not contain a stop-work section — this is correct, as MOM.8 "
      "is not in the MOM stop-work proposals.")

    save_and_verify(doc, "HCO_MOM_8_v2_REWRITE_DRAFT.docx")


# ══════════════════════════════════════════════════════════════════════════════
# MOM.9 — Narcotics, Psychotropics, Chemotherapy and Radio-pharmaceuticals
#         (HAS stop-work: Section 6)
# Content: mom9_content.txt (approved).
# Structure: Document control, Sec 3 standards, Sec 4 non-negotiables,
#            Sec 5 (5 subsections), Sec 6 Stop-work, Sec 7 Governance,
#            Sec 8 Monitoring, Sec 9 Training, Sec 10 Distribution,
#            Sec 11 Abbreviations, Sec 12 Traceability, Sec 13 Records,
#            Sec 14 References, Disclaimer.
# COREs: none | Stars: a* | Statute: YES (P2)
# Exact items verified:
#   5.2: three distinct prescriber-authorisation tiers —
#        narcotics: designated/privileged medical officer;
#        chemotherapy: demonstrated knowledge/competency (preferably oncologist);
#        radio-pharmaceuticals: statutorily authorised caregiver
#   5.4: class II biosafety cabinet (preferably IIA) + appropriate PPE for
#        chemotherapy preparation — explicit requirement
#   Sec 6: two trigger clauses (chemo/radio-pharma preparation + narcotics
#           security/issue); emergency-controlled-drug carve-out; plain-text
#           role names "Pharmacy In-Charge" and "Medical Superintendent"
#   No NDPS Act section numbers, AERB rule numbers, or other regulation-specific
#   citations invented — only generic "applicable regulations/law" language used
# ══════════════════════════════════════════════════════════════════════════════
def gen_mom9():
    """MOM.9 — Narcotics, Psychotropics, Chemotherapy, Radio-pharmaceuticals
    (stop-work Section 6; sections 1-14 + Disclaimer)"""
    doc = Document()

    # Title
    h(doc, 0, "Policy on Narcotics, Psychotropics, Chemotherapy and Radio-pharmaceuticals")
    p(doc, HN)

    # Document control
    h(doc, 1, "Document control")
    doc_ctrl(doc, "HCO/MOM/POL/09", "Pharmacy In-Charge")
    p(doc, "A blank marked ________ must be completed before issue.")

    # Statement of intent
    h(doc, 1, "Statement of intent")
    p(doc,
      "Narcotic drugs, psychotropic substances, chemotherapeutic agents, and "
      "radio-pharmaceuticals are prescribed only by appropriate caregivers, stored "
      "securely, prepared only by qualified personnel, and every use is recorded "
      "and disposed of in line with applicable law.")

    # 1. Purpose
    h(doc, 1, "1. Purpose")
    p(doc,
      f"This policy explains how {HN} handles narcotic drugs, psychotropic "
      "substances, chemotherapeutic agents, and radio-pharmaceuticals safely — "
      "how they're prescribed, stored, prepared, administered, recorded, and "
      "disposed of.")
    p(doc,
      "This policy does not cover general medication prescription, storage, or "
      "dispensing in detail — those are covered in other hospital policies.")

    # 2. Scope
    h(doc, 1, "2. Scope")
    p(doc,
      f"This policy applies to all staff involved in prescribing, storing, "
      f"preparing, administering, or disposing of narcotic drugs, psychotropic "
      f"substances, chemotherapeutic agents, and radio-pharmaceuticals at {HN}.")

    # 3. Policy standards
    h(doc, 1, "3. Policy standards")
    p(doc,
      f"{HN} handles narcotic drugs, psychotropic substances, chemotherapeutic "
      "agents, and radio-pharmaceuticals under written guidance developed in line "
      "with applicable law, prescribed only by appropriately authorised caregivers, "
      "stored securely, and prepared and administered only by qualified personnel. "
      "Usage, administration, and disposal are recorded and comply with applicable "
      "statutory requirements.")
    p(doc, "Staff follow the written guidance below and keep the records it requires.")

    # 4. Non-negotiable rules
    h(doc, 1, "4. Non-negotiable rules")
    lb(doc,
       "Do not operate without written guidance, developed in line with applicable "
       "law, covering how narcotic drugs, psychotropic substances, chemotherapeutic "
       "agents, and radio-pharmaceuticals are handled.")
    lb(doc,
       "Do not let anyone other than a designated or privileged medical officer "
       "prescribe a narcotic drug, do not let anyone without demonstrated knowledge "
       "and competency prescribe chemotherapy, and do not let anyone other than a "
       "statutorily authorised caregiver prescribe a radio-pharmaceutical.")
    lb(doc,
       "Do not leave narcotic drugs or psychotropic substances unsecured, and do not "
       "issue them without a register entry and an authorised prescriber.")
    lb(doc,
       "Do not leave chemotherapeutic agents accessible to unauthorised personnel, "
       "and do not store or label radio-pharmaceuticals outside applicable regulatory "
       "guidelines.")
    lb(doc,
       "Do not prepare or administer chemotherapeutic agents or radio-pharmaceuticals "
       "without qualified, specially trained personnel and the required preparation "
       "conditions, including a class II biosafety cabinet with appropriate personal "
       "protective equipment for chemotherapy preparation.")
    lb(doc,
       "Do not skip strict inventory control, or skip recording usage, administration, "
       "wastage, and disposal of these medications in line with applicable statutory "
       "requirements.")

    # 5. What we do
    h(doc, 1, "5. What we do")

    h(doc, 2, "5.1 Follow written guidance developed in line with applicable law")
    p(doc,
      "Written guidance, developed in consonance with applicable local and national "
      "regulations and guidelines, is implemented for the safe use of narcotic drugs, "
      "psychotropic substances, chemotherapeutic agents, and radio-pharmaceuticals. "
      "It could address all the requirements in this policy in one document.")

    h(doc, 2, "5.2 Ensure prescription by appropriate caregivers")
    p(doc,
      "Narcotic drugs are prescribed only by a designated or privileged medical "
      "officer. Chemotherapy is prescribed only by someone with the knowledge to "
      "monitor and treat its adverse effects — preferably a medical oncologist, or "
      "another doctor trained and competent in chemotherapy. Radio-pharmaceuticals "
      "are prescribed only by a caregiver authorised under applicable regulations.")

    h(doc, 2, "5.3 Store these medications securely")
    p(doc,
      "Narcotic drugs are stored securely in line with applicable regulatory "
      "requirements, with security measures that prevent diversion and abuse. "
      "Chemotherapeutic agents are accessible only to authorised personnel. "
      "Radio-pharmaceuticals are labelled and stored following applicable regulatory "
      "guidelines. It's preferable to store all of these separately from other "
      "medications.")

    h(doc, 2, "5.4 Prepare and administer only through qualified personnel")
    p(doc,
      "Personnel preparing and administering chemotherapeutic drugs have received "
      "special training for it. A class II biosafety cabinet — preferably class IIA "
      "— with appropriate personal protective equipment is used for preparing or "
      "mixing chemotherapeutic drugs. Radio-pharmaceuticals are prepared and "
      "administered only by an authorised caregiver.")

    h(doc, 2, "5.5 Keep proper records of usage and disposal")
    p(doc,
      "Strict inventory control is kept for narcotic drugs, psychotropic substances, "
      "chemotherapeutic agents, and radio-pharmaceuticals. Records of the usage, "
      "administration, wastage, and disposal of narcotic drugs are kept in line with "
      "applicable statutory requirements. Disposal of all these medications follows "
      "applicable statutory requirements and, where relevant, the manufacturer's "
      "recommendation.")

    # 6. Stop-work authority
    h(doc, 1, "6. Stop-work authority")
    p(doc,
      "Do not prepare or administer chemotherapeutic agents or radio-pharmaceuticals "
      "without qualified personnel and the required preparation conditions.")
    p(doc,
      "Do not leave narcotic drugs or psychotropic substances unsecured, or issue "
      "them without the required register entry and authorised prescriber.")
    p(doc,
      "Stop-work applies to preparation, issue, and administration of these classes. "
      "Immediate life-saving analgesia or anaesthesia using a controlled drug follows "
      "the documented emergency-controlled-drug rules and is entered in the register "
      "the same shift.")
    p(doc,
      "The person who stops tells the Pharmacy In-Charge and the Medical "
      "Superintendent the same shift. Refusing unsafe handling of these agents is "
      "not a disciplinary matter.")

    # 7. Governance and responsibility
    h(doc, 1, "7. Governance and responsibility")
    gov_tbl(doc, [
        ("Medical Superintendent",
         "Accountable for ensuring this policy is resourced and implemented; "
         "receives stop-work escalations from the Pharmacy In-Charge."),
        ("Pharmacy In-Charge",
         "Owns day-to-day implementation; maintains the narcotic and psychotropic "
         "register; ensures secure storage, issue controls, and inventory records; "
         "receives stop-work escalations and escalates to the Medical Superintendent."),
        ("Medical Oncologist / Trained Chemotherapy Doctor",
         "Prescribes chemotherapy; has demonstrated knowledge to monitor and treat "
         "adverse effects; oversees chemotherapy preparation safety."),
        ("Authorised Caregiver — Radio-pharmaceuticals",
         "Prescribes, prepares, and administers radio-pharmaceuticals under "
         "applicable regulatory authorisation; ensures regulatory storage and "
         "labelling requirements are met."),
        ("Quality Coordinator",
         "Audits this policy; holds training records and staff acknowledgements."),
    ])

    # 8. Quality monitoring
    h(doc, 1, "8. Quality monitoring")
    mon_tbl(doc, [
        ("Written guidance",
         "In place, developed in line with applicable law, covering prescription, "
         "storage, preparation, records, and disposal for all four drug classes."),
        ("Narcotic and psychotropic prescription authorisation",
         "Narcotics prescribed only by a designated or privileged medical officer; "
         "prescriber category documented."),
        ("Chemotherapy prescriber authorisation",
         "Chemotherapy prescribed only by a doctor with demonstrated knowledge and "
         "competency — preferably a medical oncologist; training or privileging "
         "records current."),
        ("Radio-pharmaceutical prescriber authorisation",
         "Radio-pharmaceuticals prescribed only by a statutorily authorised caregiver; "
         "authorisation records current."),
        ("Narcotic and psychotropic security and issue controls",
         "Drugs stored securely; narcotic register maintained; issue only with "
         "register entry and authorised prescriber."),
        ("Chemotherapy preparation — class II biosafety cabinet and PPE",
         "Class II biosafety cabinet (preferably IIA) in use for preparation; "
         "appropriate PPE used; preparation staff specially trained and records "
         "current."),
        ("Radio-pharmaceutical storage and preparation",
         "Labelled and stored per applicable regulatory guidelines; prepared and "
         "administered only by authorised caregiver."),
        ("Inventory, wastage, and disposal records",
         "Strict inventory control; usage, administration, wastage, and disposal "
         "records in line with applicable statutory requirements."),
        ("Stop-work events",
         "Stop-work events logged with trigger, action taken, and outcome."),
    ])

    # 9. Training and staff acknowledgement
    h(doc, 1, "9. Training and staff acknowledgement")
    p(doc,
      "All staff involved in prescribing, storing, preparing, administering, or "
      "disposing of narcotic drugs, psychotropic substances, chemotherapeutic agents, "
      "or radio-pharmaceuticals shall be familiar with the written guidance, "
      "prescriber-authorisation requirements, security and issue controls, "
      "preparation conditions (including class II biosafety cabinet and PPE for "
      "chemotherapy), inventory and disposal requirements, and stop-work authority "
      "in this policy.")
    p(doc,
      f"I have read the Policy on Narcotics, Psychotropics, Chemotherapy and "
      f"Radio-pharmaceuticals of {HN}. I will follow the processes described.")
    sig_tbl(doc)

    # 10. Distribution
    h(doc, 1, "10. Distribution")
    p(doc,
      "This policy shall be available to the Pharmacy In-Charge, pharmacists, "
      "nurses and doctors authorised to prescribe or administer these agents, "
      "the Medical Oncologist, the authorised radio-pharmaceutical caregiver, "
      "the Medical Superintendent, and the Quality Coordinator.")

    # 11. Abbreviations
    h(doc, 1, "11. Abbreviations")
    abbrev_tbl(doc, [
        ("AERB",  "Atomic Energy Regulatory Board"),
        ("CORE",  "Core objective element (NABH)"),
        ("MOM",   "Management of Medication (NABH Hospitals chapter)"),
        ("NABH",  "National Accreditation Board for Hospitals and Healthcare Providers"),
        ("NDPS",  "Narcotic Drugs and Psychotropic Substances"),
        ("OE",    "Objective Element (NABH)"),
        ("PPE",   "Personal Protective Equipment"),
    ])

    # 12. Traceability table
    h(doc, 1, "12. Traceability table")
    p(doc,
      "This table is an index. It is not how the policy is organised. An asterisk "
      "in the Level column means documentation of the process is required.")
    tr = tbl(doc, 6, 3)
    for ci, hdr in enumerate(("Objective Element", "Level", "Traceability to this policy")):
        tr.cell(0, ci).text = hdr
    trace_rows = [
        ("MOM.9.a", "Commitment*",
         "Sections 3 and 5.1 address written guidance developed in consonance with "
         "applicable local and national regulations and guidelines, covering all "
         "handling requirements for narcotic drugs, psychotropic substances, "
         "chemotherapeutic agents, and radio-pharmaceuticals."),
        ("MOM.9.b", "Commitment",
         "Section 5.2 addresses the three distinct prescriber-authorisation tiers: "
         "narcotics (designated or privileged medical officer), chemotherapy "
         "(demonstrated knowledge and competency, preferably medical oncologist), "
         "and radio-pharmaceuticals (caregiver authorised under applicable regulations). "
         "No NDPS Act section numbers or AERB rule citations are stated."),
        ("MOM.9.c", "Commitment",
         "Section 5.3 addresses secure narcotic and psychotropic storage preventing "
         "diversion and abuse, authorised-personnel-only access for chemotherapy, "
         "and regulatory-compliant labelling and storage for radio-pharmaceuticals."),
        ("MOM.9.d", "Commitment",
         "Section 5.4 addresses the class II biosafety cabinet (preferably class IIA) "
         "with appropriate PPE for chemotherapy preparation, special training of "
         "preparation personnel, and authorised-caregiver requirement for "
         "radio-pharmaceutical preparation and administration."),
        ("MOM.9.e", "Commitment",
         "Section 5.5 addresses strict inventory control and records of usage, "
         "administration, wastage, and disposal for narcotic drugs in line with "
         "applicable statutory requirements, and statutory-compliant disposal for "
         "all four drug classes."),
    ]
    for ri, (oe, lvl, txt) in enumerate(trace_rows, 1):
        tr.cell(ri, 0).text = oe
        tr.cell(ri, 1).text = lvl
        tr.cell(ri, 2).text = txt

    # 13. Required Records/Evidence Checklist
    h(doc, 1, "13. Required Records/Evidence Checklist")

    h(doc, 2, "Written guidance — MOM.9.a (Commitment*)")
    lb(doc,
       "Written guidance in place, developed in line with applicable local and "
       "national regulations and guidelines, covering prescription, storage, "
       "preparation, administration, record-keeping, and disposal of narcotic "
       "drugs, psychotropic substances, chemotherapeutic agents, and "
       "radio-pharmaceuticals.")
    lb(doc,
       "Date of last review and evidence that the guidance reflects current "
       "applicable statutory requirements.")

    h(doc, 2, "Prescription by appropriate caregivers — MOM.9.b")
    lb(doc,
       "List of designated or privileged medical officers authorised to prescribe "
       "narcotic drugs; narcotic prescriptions bearing authorised-prescriber identity.")
    lb(doc,
       "Records confirming the chemotherapy prescriber's demonstrated knowledge and "
       "competency — for example, training records, privileging documents, or "
       "oncology credentials.")
    lb(doc,
       "Authorisation document confirming the radio-pharmaceutical prescriber is "
       "authorised under applicable regulations.")

    h(doc, 2, "Secure storage — MOM.9.c")
    lb(doc,
       "Physical security measures for narcotic and psychotropic drug storage "
       "(for example, locked cabinet, restricted key access) that prevent diversion "
       "and abuse; in line with applicable regulatory requirements.")
    lb(doc,
       "Access controls confirming chemotherapeutic agents are accessible only to "
       "authorised personnel.")
    lb(doc,
       "Storage and labelling records for radio-pharmaceuticals confirming compliance "
       "with applicable regulatory guidelines.")

    h(doc, 2, "Qualified preparation and administration — MOM.9.d")
    lb(doc,
       "Training records for personnel preparing and administering chemotherapeutic "
       "drugs confirming special training.")
    lb(doc,
       "Evidence of class II biosafety cabinet (preferably class IIA) in use for "
       "chemotherapy preparation, with appropriate PPE available and used.")
    lb(doc,
       "Authorisation records confirming radio-pharmaceuticals are prepared and "
       "administered only by an authorised caregiver.")

    h(doc, 2, "Records of usage and disposal — MOM.9.e")
    lb(doc,
       "Narcotic register showing usage, administration, wastage, and disposal in "
       "line with applicable statutory requirements; cross-checked against inventory.")
    lb(doc,
       "Strict inventory control records for narcotic drugs, psychotropic substances, "
       "chemotherapeutic agents, and radio-pharmaceuticals.")
    lb(doc,
       "Disposal records for all four drug classes, following applicable statutory "
       "requirements and manufacturer recommendations where relevant.")

    # 14. References
    h(doc, 1, "14. References")
    ln(doc,
       "National Accreditation Board for Hospitals and Healthcare Providers. NABH "
       "Accreditation Standards for Hospitals, 6th Edition. MOM.9.")
    ln(doc, "Guidebook interpretation supplied for MOM.9.a through MOM.9.e.")
    ln(doc,
       "Applicable local and national regulations and guidelines for narcotic drugs, "
       "psychotropic substances, chemotherapeutic agents, and radio-pharmaceuticals "
       "(including, but not limited to, NDPS regulations and AERB guidelines where "
       "applicable — specific instruments identified by the organisation).")
    ln(doc,
       f"Internal documents of {HN}: written narcotic/psychotropic/chemotherapy/"
       "radio-pharmaceutical handling guidance; narcotic register; class II "
       "biosafety cabinet procedures; chemotherapy preparation records.")

    # Disclaimer
    h(doc, 1, "Disclaimer")
    p(doc,
      "This policy reorganises the supplied MOM.9 objective-element wording and "
      "Guidebook interpretation into plain-language policy format. The modal strength "
      "of the source has been preserved. Optional examples and mechanisms have not "
      "been converted into mandatory requirements. The three distinct prescriber-"
      "authorisation tiers (narcotic drugs: designated or privileged medical officer; "
      "chemotherapy: demonstrated knowledge and competency, preferably medical "
      "oncologist; radio-pharmaceuticals: statutorily authorised caregiver), the "
      "class II biosafety cabinet with appropriate PPE requirement for chemotherapy "
      "preparation, and the stop-work text with both trigger clauses (chemotherapy "
      "and radio-pharmaceutical preparation; narcotic security and issue) have been "
      "retained verbatim. No specific NDPS Act section numbers, AERB rule numbers, "
      "or other regulation-specific citations have been invented — only generic "
      "applicable-regulations language is used throughout. Role names in the "
      "stop-work section are in plain text as specified.")

    save_and_verify(doc, "HCO_MOM_9_v2_REWRITE_DRAFT.docx")


# ══════════════════════════════════════════════════════════════════════════════
# MOM.10 — Implantable Prosthesis and Medical Devices   (NO stop-work)
# Content: mom10_content.txt (approved).
# Structure: Document control, Sec 3 standards, Sec 4 non-negotiables,
#            Sec 5 (5 subsections), Sec 6 Governance, Sec 7 Monitoring,
#            Sec 8 Training, Sec 9 Distribution, Sec 10 Abbreviations,
#            Sec 11 Traceability, Sec 12 Records, Sec 13 References, Disclaimer.
# COREs: none | Stars: b*, e* | Achievement: e | Statute: none
# Exact items verified:
#   5.2: four lifecycle-stage elements in written guidance —
#        procurement / storage / issuing / use — all four named explicitly
#   5.4: three-location batch/serial number —
#        medical record / master logbook / discharge summary — all three named
#   5.5: internal-feedback recall triggers BOTH regulatory authority AND
#        manufacturer notification (not just one of them)
#   No stop-work section created (MOM.10 not in stop-work proposals)
# ══════════════════════════════════════════════════════════════════════════════
def gen_mom10():
    """MOM.10 — Implantable Prosthesis and Medical Devices
    (no stop-work; sections 1-13 + Disclaimer)"""
    doc = Document()

    # Title
    h(doc, 0, "Policy on Implantable Prosthesis and Medical Devices")
    p(doc, HN)

    # Document control
    h(doc, 1, "Document control")
    doc_ctrl(doc, "HCO/MOM/POL/10", "OT In-Charge")
    p(doc, "A blank marked ________ must be completed before issue.")

    # Statement of intent
    h(doc, 1, "Statement of intent")
    p(doc,
      "Implantable prosthesis and medical devices are used only when backed by "
      "scientific evidence and recognised regulatory approval, tracked from "
      "procurement through to the patient's discharge summary, and recalled "
      "effectively when needed.")

    # 1. Purpose
    h(doc, 1, "1. Purpose")
    p(doc,
      f"This policy explains how {HN} selects and approves implantable prosthesis "
      "and medical devices, manages their procurement and use, counsels patients "
      "on precautions, records batch and serial numbers, and handles recalls.")
    p(doc,
      "This policy does not cover general informed consent procedures or infection "
      "prevention and control in detail — those are covered in other hospital "
      "policies.")

    # 2. Scope
    h(doc, 1, "2. Scope")
    p(doc,
      f"This policy applies to all staff involved in selecting, procuring, using, "
      f"or recalling implantable prosthesis and medical devices at {HN}.")

    # 3. Policy standards
    h(doc, 1, "3. Policy standards")
    p(doc,
      f"{HN} selects implantable prosthesis and medical devices only where "
      "sufficient scientific data and recognised national or international "
      "regulatory approval exist, with the multidisciplinary committee approving "
      "each item. Written guidance governs procurement, storage, issuing, and use, "
      "and follows infection prevention and control requirements. Patients and "
      "families are counselled on precautions, batch and serial numbers are "
      "recorded in the patient's medical record, the master logbook, and the "
      "discharge summary, and recalls are handled effectively.")
    p(doc, "Staff follow the written guidance below and keep the records it requires.")

    # 4. Non-negotiable rules
    h(doc, 1, "4. Non-negotiable rules")
    lb(doc,
       "Do not select an implantable prosthesis or medical device without sufficient "
       "scientific data and recognised national or international regulatory approval, "
       "and without the multidisciplinary committee's approval for that item.")
    lb(doc,
       "Do not procure, store, issue, or use an implantable prosthesis or medical "
       "device without written guidance covering all four of those stages, addressing "
       "statutory regulations and manufacturer recommendations.")
    lb(doc,
       "Do not skip documenting counselling given to the patient or family about "
       "the device, including any precautions.")
    lb(doc,
       "Do not skip recording the batch and serial number of an implant in all "
       "three places: the patient's medical record, the master logbook, and the "
       "discharge summary.")
    lb(doc,
       "Do not respond to a recall — whether from a regulatory authority, the "
       "manufacturer, or internal feedback — without maintaining a record of it, "
       "and where the recall arises from internal feedback, without informing "
       "the appropriate regulatory authority and the manufacturer.")

    # 5. What we do
    h(doc, 1, "5. What we do")

    h(doc, 2, "5.1 Select devices based on evidence and approval")
    p(doc,
      f"{HN} ensures relevant and sufficient scientific data are available before "
      "selecting an implant or device, and checks for applicable national or "
      "international regulatory approval — for example US-FDA notification, or "
      "Central Drugs Standard Control Organisation notification under the Drugs "
      "and Cosmetics Act. These are examples; other relevant national or "
      "international approvals also satisfy this requirement. The multidisciplinary "
      "committee approves the use of each particular implant.")

    h(doc, 2, "5.2 Manage the full lifecycle under written guidance")
    p(doc,
      "Written guidance directs the procurement, storage, issuing, and use of "
      "implantable prosthesis and medical devices, addressing applicable statutory "
      "regulations and manufacturer recommendations. Infection prevention and "
      "control requirements related to using these devices are implemented.")

    h(doc, 2, "5.3 Counsel patients and families on precautions")
    p(doc,
      "The patient and family are counselled about using the implant or device, "
      "including any precautions — for example, avoiding specific drugs, or "
      "reporting particular symptoms to the hospital. These are examples of "
      "precautions; the actual content depends on the device and clinical "
      "situation. Details of this counselling are documented in the informed "
      "consent record.")

    h(doc, 2, "5.4 Record batch and serial numbers")
    p(doc,
      "The batch and serial number of every implantable prosthesis and medical "
      "device are recorded in three places: the patient's medical record, the "
      "master logbook, and the discharge summary. Where an implant doesn't come "
      f"with a pre-labelled sticker, {HN} has a mechanism to identify it — for "
      "example, by manufacturer, type, size, batch number, and serial number.")

    h(doc, 2, "5.5 Handle recalls effectively")
    p(doc,
      "A recall could arise from a regulatory authority, the manufacturer, or "
      f"internal feedback. Where a recall arises from internal feedback, {HN} "
      "also informs the appropriate regulatory authority and the manufacturer. "
      "A record is maintained whenever a recall occurs.")

    # 6. Governance and responsibility
    h(doc, 1, "6. Governance and responsibility")
    gov_tbl(doc, [
        ("Medical Superintendent",
         "Accountable for ensuring this policy is resourced and implemented; "
         "approves or delegates approval of implant and device selection with "
         "the multidisciplinary committee."),
        ("OT In-Charge",
         "Owns day-to-day implementation of this policy; ensures written guidance "
         "covers procurement, storage, issuing, and use; coordinates batch and "
         "serial number recording; manages recall response."),
        ("Multidisciplinary Committee",
         "Approves the use of each particular implantable prosthesis and medical "
         "device on the basis of scientific data and regulatory approval."),
        ("Treating doctors",
         "Counsel patients and families on precautions; document counselling in "
         "the informed consent record; record batch and serial numbers in the "
         "medical record."),
        ("Quality Coordinator",
         "Audits this policy; holds training records and staff acknowledgements; "
         "tracks open recall actions."),
    ])

    # 7. Quality monitoring
    h(doc, 1, "7. Quality monitoring")
    mon_tbl(doc, [
        ("Device selection — scientific data and regulatory approval",
         "Scientific data available before selection; national or international "
         "regulatory approval confirmed; multidisciplinary committee approval "
         "documented for each item."),
        ("Lifecycle written guidance — all four stages",
         "Written guidance covering procurement, storage, issuing, and use; "
         "addresses statutory regulations and manufacturer recommendations; "
         "infection prevention and control requirements implemented."),
        ("Patient and family counselling",
         "Counselling on precautions documented in the informed consent record "
         "for each implant or device; no gaps."),
        ("Batch and serial number recording — three locations",
         "Batch and serial number in the patient's medical record, the master "
         "logbook, and the discharge summary — all three confirmed for every "
         "procedure; identification mechanism in place where no pre-labelled "
         "sticker exists."),
        ("Recall management",
         "Recall records maintained for every recall event; internal-feedback "
         "recalls trigger both regulatory authority and manufacturer notification; "
         "open actions tracked."),
    ])

    # 8. Training and staff acknowledgement
    h(doc, 1, "8. Training and staff acknowledgement")
    p(doc,
      "All staff involved in selecting, procuring, using, or recalling implantable "
      "prosthesis and medical devices shall be familiar with the selection criteria, "
      "lifecycle written guidance, counselling documentation, batch and serial number "
      "recording requirements, and recall handling in this policy.")
    p(doc,
      f"I have read the Policy on Implantable Prosthesis and Medical Devices of "
      f"{HN}. I will follow the processes described.")
    sig_tbl(doc)

    # 9. Distribution
    h(doc, 1, "9. Distribution")
    p(doc,
      "This policy shall be available to the OT In-Charge, treating doctors, the "
      "multidisciplinary committee, the procurement team, and the Quality "
      "Coordinator.")

    # 10. Abbreviations
    h(doc, 1, "10. Abbreviations")
    abbrev_tbl(doc, [
        ("CDSCO", "Central Drugs Standard Control Organisation"),
        ("FDA",   "Food and Drug Administration (United States)"),
        ("MOM",   "Management of Medication (NABH Hospitals chapter)"),
        ("NABH",  "National Accreditation Board for Hospitals and Healthcare Providers"),
        ("OT",    "Operation Theatre"),
    ])

    # 11. Traceability table
    h(doc, 1, "11. Traceability table")
    p(doc,
      "This table is an index. It is not how the policy is organised. An asterisk "
      "in the Level column means documentation of the process is required.")
    tr = tbl(doc, 6, 3)
    for ci, hdr in enumerate(("Objective Element", "Level", "Traceability to this policy")):
        tr.cell(0, ci).text = hdr
    trace_rows = [
        ("MOM.10.a", "Commitment",
         "Section 5.1 addresses selection of implants and devices only where "
         "relevant and sufficient scientific data are available and applicable "
         "national or international regulatory approval is confirmed, with "
         "multidisciplinary committee approval for each item."),
        ("MOM.10.b", "Commitment*",
         "Sections 3 and 5.2 address written guidance covering all four lifecycle "
         "stages — procurement, storage, issuing, and use — addressing statutory "
         "regulations and manufacturer recommendations, and implementing infection "
         "prevention and control requirements."),
        ("MOM.10.c", "Commitment",
         "Section 5.3 addresses counselling of the patient and family on "
         "precautions, with documentation in the informed consent record for "
         "each device."),
        ("MOM.10.d", "Commitment",
         "Sections 3 and 5.4 address the three-location batch and serial number "
         "recording requirement: the patient's medical record, the master logbook, "
         "and the discharge summary — all three for every implant — and the "
         "identification mechanism where no pre-labelled sticker exists."),
        ("MOM.10.e", "Achievement*",
         "Sections 3 and 5.5 address recall management: records maintained for "
         "every recall; internal-feedback recalls trigger notification of both the "
         "appropriate regulatory authority and the manufacturer. This is an "
         "Achievement-level OE — it requires evidence of an effective recall "
         "system, not just a documented process."),
    ]
    for ri, (oe, lvl, txt) in enumerate(trace_rows, 1):
        tr.cell(ri, 0).text = oe
        tr.cell(ri, 1).text = lvl
        tr.cell(ri, 2).text = txt

    # 12. Required Records/Evidence Checklist
    h(doc, 1, "12. Required Records/Evidence Checklist")

    h(doc, 2, "Device selection — MOM.10.a")
    lb(doc,
       "Scientific data on file for each approved implant or device, reviewed "
       "before selection.")
    lb(doc,
       "Regulatory approval records (national or international) for each approved "
       "item — for example, CDSCO notification or US-FDA notification.")
    lb(doc,
       "Multidisciplinary committee approval records for each particular implant "
       "or device in use.")

    h(doc, 2, "Lifecycle written guidance — MOM.10.b (Commitment*)")
    lb(doc,
       "Written guidance covering procurement, storage, issuing, and use of "
       "implantable prosthesis and medical devices — all four stages present.")
    lb(doc,
       "Evidence that statutory regulations and manufacturer recommendations are "
       "addressed in the guidance.")
    lb(doc,
       "Evidence that infection prevention and control requirements for using these "
       "devices are implemented.")

    h(doc, 2, "Patient and family counselling — MOM.10.c")
    lb(doc,
       "Informed consent records documenting counselling on precautions for each "
       "patient receiving an implant or device — no gaps.")
    lb(doc,
       "Examples of precautions documented (drugs to avoid, symptoms to report, etc.).")

    h(doc, 2, "Batch and serial number recording — MOM.10.d")
    lb(doc,
       "Patient's medical record showing batch and serial number for each implant "
       "or device used.")
    lb(doc,
       "Master logbook entries showing batch and serial number for each procedure.")
    lb(doc,
       "Discharge summary showing batch and serial number for each implant or device.")
    lb(doc,
       "Identification mechanism in place for implants without pre-labelled stickers "
       "(recording manufacturer, type, size, batch number, and serial number manually).")

    h(doc, 2, "Recall management — MOM.10.e (Achievement*)")
    lb(doc,
       "Recall log maintained, with a record for every recall event — including "
       "source (regulatory authority, manufacturer, or internal feedback).")
    lb(doc,
       "For recalls arising from internal feedback: evidence that the appropriate "
       "regulatory authority was notified and the manufacturer was notified.")
    lb(doc, "Open recall actions tracked to closure.")

    # 13. References
    h(doc, 1, "13. References")
    ln(doc,
       "National Accreditation Board for Hospitals and Healthcare Providers. NABH "
       "Accreditation Standards for Hospitals, 6th Edition. MOM.10.")
    ln(doc, "Guidebook interpretation supplied for MOM.10.a through MOM.10.e.")
    ln(doc,
       f"Internal documents of {HN}: implant and device approved-list; "
       "multidisciplinary committee minutes; lifecycle written guidance; "
       "batch and serial number master logbook; recall log.")

    # Disclaimer
    h(doc, 1, "Disclaimer")
    p(doc,
      "This policy reorganises the supplied MOM.10 objective-element wording and "
      "Guidebook interpretation into plain-language policy format. The modal "
      "strength of the source has been preserved. Optional examples and mechanisms "
      "have not been converted into mandatory requirements. The four lifecycle-stage "
      "elements in written guidance (procurement, storage, issuing, and use) have "
      "been retained as an exhaustive mandatory list. The three-location batch and "
      "serial number requirement (medical record, master logbook, and discharge "
      "summary) has been retained verbatim. The internal-feedback recall requirement "
      "triggering notification of both the appropriate regulatory authority and the "
      "manufacturer has been retained verbatim. This policy does not contain a "
      "stop-work section — this is correct, as MOM.10 is not in the MOM stop-work "
      "proposals.")

    save_and_verify(doc, "HCO_MOM_10_v2_REWRITE_DRAFT.docx")


# ══════════════════════════════════════════════════════════════════════════════
# MOM.11 — Storage and Availability of Medical Supplies and Consumables
#          (NO stop-work)
# Content: mom11_content.txt (approved).
# Structure: Document control, Sec 3 standards, Sec 4 non-negotiables,
#            Sec 5 (5 subsections), Sec 6 Governance, Sec 7 Monitoring,
#            Sec 8 Training, Sec 9 Distribution, Sec 10 Abbreviations,
#            Sec 11 Traceability, Sec 12 Records, Sec 13 References, Disclaimer.
# COREs: none | Stars: a* | Statute: none
# Exact items verified:
#   Purpose: scope-clarifying note — "medical supplies and consumables" means
#            items used in patient care OTHER THAN medications and implants;
#            both exclusions stated explicitly
#   5.1: five acquisition-process elements —
#        vendor selection / vendor evaluation / indenting / purchase order
#        generation / receipt of goods — all five named
#   5.3: "including wards" universal storage scope stated explicitly
#   5.4: no specific inventory method mandated — examples given (ABC, VED,
#        FSN, FEFO, lead-time) but requirement is "sound inventory control
#        practice" demonstrably in use, not any particular method
# ══════════════════════════════════════════════════════════════════════════════
def gen_mom11():
    """MOM.11 — Storage and Availability of Medical Supplies and Consumables
    (no stop-work; sections 1-13 + Disclaimer)"""
    doc = Document()

    # Title
    h(doc, 0, "Policy on Storage and Availability of Medical Supplies and Consumables")
    p(doc, HN)

    # Document control
    h(doc, 1, "Document control")
    doc_ctrl(doc, "HCO/MOM/POL/11", "Stores In-Charge")
    p(doc, "A blank marked ________ must be completed before issue.")

    # Statement of intent
    h(doc, 1, "Statement of intent")
    p(doc,
      "Medical supplies and consumables are acquired through a defined process, "
      "stored safely and cleanly, controlled through sound inventory practice, "
      "and checked for condition before every use.")

    # 1. Purpose
    h(doc, 1, "1. Purpose")
    p(doc,
      f"This policy explains how {HN} acquires medical supplies and consumables, "
      "stores them safely, controls inventory soundly, and verifies their condition "
      "before they're used. In this policy, \"medical supplies and consumables\" "
      "means items used in patient care other than medications and implants — "
      "those are covered in other hospital policies.")
    p(doc,
      "This policy does not cover medication or implant storage in detail — those "
      "are covered in other hospital policies.")

    # 2. Scope
    h(doc, 1, "2. Scope")
    p(doc,
      f"This policy applies to Stores staff and all clinical staff who acquire, "
      f"store, or use medical supplies and consumables at {HN}.")

    # 3. Policy standards
    h(doc, 1, "3. Policy standards")
    p(doc,
      f"{HN} acquires medical supplies and consumables through a defined process, "
      "uses them safely with precautions to maintain sterility and integrity where "
      "appropriate, stores them in a clean, safe, and secure environment following "
      "manufacturer recommendations, applies sound inventory control practices, "
      "and verifies their condition before dispensing and use.")
    p(doc, "Staff follow the written guidance below and keep the records it requires.")

    # 4. Non-negotiable rules
    h(doc, 1, "4. Non-negotiable rules")
    lb(doc,
       "Do not acquire medical supplies or consumables outside the defined process "
       "covering vendor selection, vendor evaluation, indenting, purchase order "
       "generation, and receipt of goods.")
    lb(doc,
       "Do not open or use a medical supply or consumable without the precautions "
       "needed to maintain its sterility and integrity, where appropriate.")
    lb(doc,
       "Do not store medical supplies or consumables outside the manufacturer's "
       "recommendations, in any area including wards, and do not leave hazardous "
       "materials unidentified or unsafely stored.")
    lb(doc,
       "Do not store medical supplies or consumables without following a sound "
       "inventory control practice.")
    lb(doc,
       "Do not dispense or use a medical supply or consumable without first "
       "checking its condition.")

    # 5. What we do
    h(doc, 1, "5. What we do")

    h(doc, 2, "5.1 Acquire supplies through a defined process")
    p(doc,
      "The acquisition process for medical supplies and consumables addresses "
      "five things: vendor selection, vendor evaluation, the indenting process, "
      "purchase order generation, and receipt of goods.")

    h(doc, 2, "5.2 Use supplies safely")
    p(doc,
      "Where appropriate, medical supplies and consumables are opened and used "
      "with the precautions needed to maintain their sterility and integrity.")

    h(doc, 2, "5.3 Store supplies safely")
    p(doc,
      "Storage requirements specified by the manufacturer are followed, in every "
      "area supplies are stored, including wards. Supplies are protected from loss "
      "or theft, storage areas are kept clean, and hazardous materials are "
      "identified and kept safely.")

    h(doc, 2, "5.4 Apply sound inventory control")
    p(doc,
      f"{HN} follows or demonstrates sound inventory control practices — for "
      "example ABC analysis, VED (vital/essential/desirable), FSN (fast/slow/"
      "non-moving), First-Expiry-First-Out, or lead-time analysis. These are "
      f"examples; {HN} may use any combination of practices it can demonstrate "
      "is actually in use, not just documented.")

    h(doc, 2, "5.5 Verify condition before use")
    p(doc,
      "Medical supplies and consumables are kept in a condition suitable for safe "
      "use, and that condition is checked before dispensing and use — for example "
      "checking for an opened package, a damp cotton roll, physical damage, or "
      "discolouration. These are examples of defects to check for, not an "
      "exhaustive list.")

    # 6. Governance and responsibility
    h(doc, 1, "6. Governance and responsibility")
    gov_tbl(doc, [
        ("Medical Superintendent",
         "Accountable for ensuring this policy is resourced and implemented."),
        ("Stores In-Charge",
         "Owns day-to-day implementation of this policy; manages the defined "
         "acquisition process; maintains inventory control; ensures storage "
         "conditions are met in the main stores area."),
        ("Clinical staff — ward and department heads",
         "Ensure manufacturer storage requirements are followed in their areas, "
         "including wards; check condition of supplies before use; ensure safe "
         "use precautions are applied where appropriate."),
        ("Quality Coordinator",
         "Audits this policy; holds training records and staff acknowledgements."),
    ])

    # 7. Quality monitoring
    h(doc, 1, "7. Quality monitoring")
    mon_tbl(doc, [
        ("Defined acquisition process — all five elements",
         "Process in place covering vendor selection, vendor evaluation, indenting, "
         "purchase order generation, and receipt of goods; all five elements "
         "present and working."),
        ("Safe use — sterility and integrity precautions",
         "Precautions applied when opening and using supplies where appropriate; "
         "no instances of inappropriate opening technique or sterility break."),
        ("Storage compliance — including wards",
         "Manufacturer storage requirements followed in all storage areas, "
         "including wards; supplies protected from loss or theft; storage areas "
         "clean; hazardous materials identified and safely stored."),
        ("Inventory control",
         "Sound inventory control practice demonstrated as actually in use — not "
         "just documented; method(s) employed consistent with those declared."),
        ("Condition check before use",
         "Condition checked before dispensing and use; defective supplies "
         "(opened, damaged, discoloured, damp) identified and removed."),
    ])

    # 8. Training and staff acknowledgement
    h(doc, 1, "8. Training and staff acknowledgement")
    p(doc,
      "All Stores staff and clinical staff who acquire, store, or use medical "
      "supplies and consumables shall be familiar with the defined acquisition "
      "process, safe-use precautions, storage requirements (including in wards), "
      "inventory control expectations, and condition-checking requirements in "
      "this policy.")
    p(doc,
      f"I have read the Policy on Storage and Availability of Medical Supplies "
      f"and Consumables of {HN}. I will follow the processes described.")
    sig_tbl(doc)

    # 9. Distribution
    h(doc, 1, "9. Distribution")
    p(doc,
      "This policy shall be available to the Stores In-Charge, Stores staff, "
      "ward and department heads, and the Quality Coordinator.")

    # 10. Abbreviations
    h(doc, 1, "10. Abbreviations")
    abbrev_tbl(doc, [
        ("ABC",  "Activity-Based Classification — value analysis method for inventory"),
        ("FEFO", "First-Expiry-First-Out"),
        ("FSN",  "Fast/Slow/Non-moving — inventory turnover classification"),
        ("MOM",  "Management of Medication (NABH Hospitals chapter)"),
        ("NABH", "National Accreditation Board for Hospitals and Healthcare Providers"),
        ("VED",  "Vital/Essential/Desirable — inventory criticality classification"),
    ])

    # 11. Traceability table
    h(doc, 1, "11. Traceability table")
    p(doc,
      "This table is an index. It is not how the policy is organised. An asterisk "
      "in the Level column means documentation of the process is required.")
    tr = tbl(doc, 6, 3)
    for ci, hdr in enumerate(("Objective Element", "Level", "Traceability to this policy")):
        tr.cell(0, ci).text = hdr
    trace_rows = [
        ("MOM.11.a", "Commitment*",
         "Sections 3 and 5.1 address the defined acquisition process, covering "
         "all five elements: vendor selection, vendor evaluation, the indenting "
         "process, purchase order generation, and receipt of goods. The Purpose "
         "section also clarifies that this policy covers items used in patient "
         "care other than medications and implants — those are covered elsewhere."),
        ("MOM.11.b", "Commitment",
         "Section 5.2 addresses the safe opening and use of medical supplies and "
         "consumables with precautions to maintain sterility and integrity, "
         "where appropriate."),
        ("MOM.11.c", "Commitment",
         "Section 5.3 addresses storage per manufacturer requirements in every "
         "storage area, including wards — the \"including wards\" scope is stated "
         "explicitly. Also addresses protection from loss or theft, clean storage "
         "areas, and identification and safe storage of hazardous materials."),
        ("MOM.11.d", "Commitment",
         "Section 5.4 addresses sound inventory control practice — demonstrated "
         "as actually in use. Examples are given (ABC, VED, FSN, FEFO, lead-time "
         "analysis) but no specific method is mandated; the requirement is "
         "demonstrable practice, not any particular system."),
        ("MOM.11.e", "Commitment",
         "Section 5.5 addresses verification of condition before dispensing and "
         "use — for example checking for an opened package, a damp cotton roll, "
         "physical damage, or discolouration. Examples are not an exhaustive list."),
    ]
    for ri, (oe, lvl, txt) in enumerate(trace_rows, 1):
        tr.cell(ri, 0).text = oe
        tr.cell(ri, 1).text = lvl
        tr.cell(ri, 2).text = txt

    # 12. Required Records/Evidence Checklist
    h(doc, 1, "12. Required Records/Evidence Checklist")

    h(doc, 2, "Defined acquisition process — MOM.11.a (Commitment*)")
    lb(doc,
       "Documented acquisition process covering all five elements: vendor "
       "selection, vendor evaluation, indenting, purchase order generation, and "
       "receipt of goods.")
    lb(doc,
       "Working evidence for each element — for example, approved vendor list, "
       "vendor-evaluation records, indent forms, purchase orders, and goods-receipt "
       "notes.")
    lb(doc,
       "Confirmation that supplies are not acquired outside this defined process.")

    h(doc, 2, "Safe use — MOM.11.b")
    lb(doc,
       "Training records or SOPs confirming staff know the precautions for "
       "maintaining sterility and integrity of the supplies they use.")
    lb(doc,
       "Audit or observation evidence that appropriate precautions are followed "
       "when opening and using supplies.")

    h(doc, 2, "Storage compliance including wards — MOM.11.c")
    lb(doc,
       "Manufacturer storage-requirement records on file for each major category "
       "of supply.")
    lb(doc,
       "Audit evidence that storage requirements are followed in all areas, "
       "including wards — temperature, humidity, light, and packaging requirements "
       "where applicable.")
    lb(doc,
       "Hazardous materials identified and safely stored; inventory of hazardous "
       "materials on file.")

    h(doc, 2, "Inventory control — MOM.11.d")
    lb(doc,
       "Declaration of the inventory control method(s) in use (from the documented "
       "practice, not just from the policy).")
    lb(doc,
       "Working records demonstrating the declared method is actually applied — "
       "for example ABC classification outputs, VED analysis records, FEFO stock "
       "rotation records, or lead-time analysis sheets.")

    h(doc, 2, "Condition check before use — MOM.11.e")
    lb(doc,
       "SOP or checklist for condition inspection before dispensing and use, "
       "naming the defects to check for.")
    lb(doc, "Audit evidence that condition checks happen before dispensing and use.")
    lb(doc, "Records of defective items identified, removed, and replaced.")

    # 13. References
    h(doc, 1, "13. References")
    ln(doc,
       "National Accreditation Board for Hospitals and Healthcare Providers. NABH "
       "Accreditation Standards for Hospitals, 6th Edition. MOM.11.")
    ln(doc, "Guidebook interpretation supplied for MOM.11.a through MOM.11.e.")
    ln(doc,
       f"Internal documents of {HN}: acquisition process documentation; approved "
       "vendor list; inventory control records; storage SOPs; condition-check "
       "checklists.")

    # Disclaimer
    h(doc, 1, "Disclaimer")
    p(doc,
      "This policy reorganises the supplied MOM.11 objective-element wording and "
      "Guidebook interpretation into plain-language policy format. The modal "
      "strength of the source has been preserved. Optional examples and mechanisms "
      "have not been converted into mandatory requirements. The scope-clarifying "
      "note in the Purpose section (\"medical supplies and consumables\" means "
      "items used in patient care other than medications and implants) has been "
      "retained verbatim. All five acquisition-process elements (vendor selection, "
      "vendor evaluation, indenting, purchase order generation, and receipt of "
      "goods) have been retained as an exhaustive mandatory list. The \"including "
      "wards\" universal storage scope has been stated explicitly. No specific "
      "inventory method has been mandated — the requirement is sound inventory "
      "control practice demonstrably in use, not any particular system. This "
      "policy does not contain a stop-work section — this is correct, as MOM.11 "
      "is not in the MOM stop-work proposals.")

    save_and_verify(doc, "HCO_MOM_11_v2_REWRITE_DRAFT.docx")


# ══════════════════════════════════════════════════════════════════════════════
if __name__ == "__main__":
    gen_mom9()
    print("\nMOM.9 draft generated.")
    gen_mom10()
    print("\nMOM.10 draft generated.")
    gen_mom11()
    print("\nMOM.11 draft generated.")
