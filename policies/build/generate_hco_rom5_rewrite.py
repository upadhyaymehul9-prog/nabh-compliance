# -*- coding: utf-8 -*-
"""
Generate HCO_ROM_5_v2_REWRITE_DRAFT.docx
Plain-language rewrite per HCO_Policy_Rewrite_Rulebook.pdf
No stop-work: Governance=6, QM=7, Training=8, Distribution=9,
Abbreviations=10, Traceability=11, Records=12, References=13, Disclaimer (unnumbered)
CORE: none. Achievement: d. Excellence: f. Asterisked: b, e, f.
"""
from docx import Document

OUT = "policies/build/rewrite_reference/HCO_ROM_5_v2_REWRITE_DRAFT.docx"

HN = "«Hospital Name»"
EM = "—"
doc = Document()


def h(level, text):
    styles = {0: "Title", 1: "Heading 1", 2: "Heading 2"}
    return doc.add_paragraph(text, style=styles[level])


def p(text):
    return doc.add_paragraph(text, style="Normal")


def ln(text):
    return doc.add_paragraph(text, style="List Number")


def lb(text):
    return doc.add_paragraph(text, style="List Bullet")


def tbl(rows, cols):
    t = doc.add_table(rows=rows, cols=cols)
    try:
        t.style = "Table Grid"
    except KeyError:
        pass
    return t


# ── Title ─────────────────────────────────────────────────────────────────────
h(0, "Policy on Professional Functioning of the Organisation")
p(HN)

# ── Document control ──────────────────────────────────────────────────────────
h(1, "Document control")

dc = tbl(6, 4)
for ri, (a, b, c, d) in enumerate([
    ("Document No.", "HCO/ROM/POL/05", "Version", "2.0"),
    ("Issue No.", "01", "Review due", "One year from implementation"),
    ("Date created", "________", "Date of implementation", "________"),
]):
    dc.cell(ri, 0).text = a
    dc.cell(ri, 1).text = b
    dc.cell(ri, 2).text = c
    dc.cell(ri, 3).text = d

for ri, (label, content) in enumerate([
    ("Prepared by",
     "Quality Coordinator Name: ________ Signature: ________"),
    ("Reviewed by",
     "Medical Superintendent Name: ________ Signature: ________"),
    ("Approved by",
     "Medical Superintendent Name: ________ Signature: ________"),
], start=3):
    dc.cell(ri, 0).text = label
    dc.cell(ri, 1).text = content
    dc.cell(ri, 1).merge(dc.cell(ri, 3))

p("A blank marked ________ must be completed before issue.")

# ── Statement of intent ───────────────────────────────────────────────────────
h(1, "Statement of intent")
p(f"{HN} functions professionally. It maintains written strategic and "
  f"operational plans aligned to its vision, mission and values, developed in "
  f"consultation with stakeholders. It coordinates its functioning with "
  f"departments and external agencies, monitors progress toward defined goals, "
  f"prepares and reviews an annual budget, reviews its committees for "
  f"effectiveness, tracks measurable service standards, and has written guidance "
  f"for managing change.")

# ── 1. Purpose ────────────────────────────────────────────────────────────────
h(1, "1. Purpose")
p(f"This policy explains how {HN} develops and maintains strategic and "
  f"operational plans aligned to its vision, mission and values in consultation "
  f"with stakeholders; coordinates and monitors progress toward those plans with "
  f"departments and external agencies; prepares and reviews an annual budget; "
  f"reviews the effectiveness of its committees; documents and monitors measurable "
  f"service standards; and maintains written guidance for change management.")
p("This policy does not cover patient assessment, clinical care, medication "
  "management, patient rights, infection control, or quality and safety "
  "monitoring — those are covered in other hospital policies. The other "
  "Responsibilities of Management standards have their own policies too.")

# ── 2. Scope ──────────────────────────────────────────────────────────────────
h(1, "2. Scope")
p(f"This policy applies to those responsible for governance (as identified "
  f"under ROM.1), the Medical Superintendent, the Quality Coordinator, "
  f"departmental leaders, committee chairs, and staff whose roles involve "
  f"strategic planning, budgeting, or committee participation at {HN}.")

# ── 3. Policy standards ───────────────────────────────────────────────────────
h(1, "3. Policy standards")
p(f"{HN} maintains documented strategic and operational plans with measurable "
  f"goals; coordinates their implementation with departments and external "
  f"agencies; prepares an annual budget that covers those plans and includes an "
  f"adequate allocation for IPC and quality improvement activities; reviews the "
  f"effectiveness of each committee it maintains; documents and monitors "
  f"measurable service standards; and keeps written guidance for managing change.")
p("Staff follow the written guidance below, keep the records it requires, and "
  "raise concerns through the channels it describes.")

# ── 4. Non-negotiable rules ───────────────────────────────────────────────────
h(1, "4. Non-negotiable rules")
ln(f"The hospital must have written strategic and operational plans {EM} a "
   f"verbal commitment or a presentation slide is not a plan.")
ln(f"Goals and objectives drawn from those plans must have measurable outcomes "
   f"and defined time frames {EM} goals without metrics cannot be monitored.")
ln(f"The annual budget must include an adequate allocation for IPC activities "
   f"and quality improvement activities {EM} a budget that makes no provision "
   f"for these is incomplete.")
ln(f"The leadership must formally review the effectiveness of each committee "
   f"the hospital maintains and document the outcome {EM} a committee that has "
   f"never been reviewed does not comply with this standard.")
ln(f"Minutes must be maintained for each committee meeting {EM} a meeting with "
   f"no minutes has no record of decisions or actions.")
ln(f"Service standards must be documented as measurable benchmarks and "
   f"monitored at a defined frequency {EM} standards that are neither measurable "
   f"nor actively monitored are not compliant.")
ln("Staff who see any of these rules broken report it the same shift to the "
   "department in-charge or the Medical Superintendent.")

# ── 5. What we do ─────────────────────────────────────────────────────────────
h(1, "5. What we do")

h(2, "5.1 Develop and maintain strategic and operational plans")
p(f"The leader of {HN} defines and develops the process for producing strategic "
  f"and operational plans. The strategic plan is aligned to the organisation's "
  f"vision and mission statement and is consistent with its values. It is "
  f"developed in consultation with all stakeholders, including the community "
  f"the organisation serves.")
p(f"The strategic plan takes into account both an external and internal scan "
  f"{EM} for example using SWOT analysis (Strengths, Weaknesses, Opportunities, "
  f"Threats). Inputs considered in finalising the plan include the findings from "
  f"the hospital's risk management programme, patient safety goals, and the "
  f"results of facility rounds. Both the strategic plan and each operational plan "
  f"have defined goals and objectives that are measurable and carry defined "
  f"time frames.")
p("The operational plan is produced on an annual basis. Both plans are "
  "documented and retained on file.")

h(2, "5.2 Coordinate functioning and monitor progress")
p("Goals and objectives are drawn from the strategic and operational plans. "
  "They are consistent with the hospital's mission and values and have "
  "measurable outcomes.")
p(f"{HN} coordinates its functioning with its departments and with external "
  f"agencies as applicable. Progress toward the defined goals and objectives "
  f"is monitored at regular intervals. Where a goal has not been achieved, the "
  f"reasons are analysed and appropriate action is taken. This may be done "
  f"through periodic reviews and/or formal management review meetings, which "
  f"are documented.")
p("A record is kept of the analysis and the action taken for each goal that "
  "was not achieved.")

h(2, "5.3 Plan and budget annually")
p(f"{HN} plans and budgets for its activities each year. Budget development is "
  f"guided by the hospital's strategic and operational plans. The budget "
  f"development and review process includes a comparison of budgeted expenditure "
  f"against actual expenditure. Responsibilities for budget development and "
  f"review are clearly defined and followed.")
p(f"An adequate budget is allocated for infection prevention and control "
  f"activities (see also IPC.2.b) and for quality improvement activities "
  f"(see also PSQ.6.d). What constitutes an adequate allocation is for the "
  f"hospital to determine in light of its own IPC programme and quality "
  f"improvement activities {EM} the source requirement is that the allocation "
  f"is adequate, not that it meets a prescribed percentage or amount. The "
  f"hospital may organise its budget on a calendar year (January{EM}December) "
  f"or a financial year (April{EM}March) basis.")
p(f"The budget is documented and approved. The form it takes is for the "
  f"hospital to define {EM} the obligation is that it covers the hospital's "
  f"activities and includes provision for IPC and quality improvement.")

h(2, "5.4 Review committee effectiveness")
p(f"The leadership of {HN} reviews the functioning of each committee for "
  f"effectiveness. The review considers: whether the purpose for which the "
  f"committee was established is being met; whether the committee is meeting "
  f"at its prescribed frequency; whether it is identifying issues and suggesting "
  f"remedial measures; and whether corrective and preventive actions it has "
  f"recommended are being adequately monitored by way of risk mitigation within "
  f"the committee's scope.")
p(f"Minutes are maintained for each committee meeting. The Guidebook notes "
  f"that, to support an effective review, the organisation may document the "
  f"scope of every committee, the roles and responsibilities assigned to its "
  f"members, and the prescribed meeting frequency {EM} the organisation decides "
  f"whether and how to do this.")

h(2, "5.5 Document and monitor service standards")
p(f"{HN} develops measurable benchmarks for the different aspects of the "
  f"services it provides. The benchmarks may be derived from the organisation's "
  f"values, from national guidelines, or from the hospital's own experience. "
  f"The hospital defines the person responsible for monitoring each benchmark "
  f"and the frequency at which monitoring is done.")
p("Service standards are documented. Monitoring results are recorded.")

h(2, "5.6 Written guidance for change management")
p(f"{HN} maintains written guidance for managing change within the organisation. "
  f"This guidance covers operational, financial, and departmental changes, and "
  f"addresses succession planning and changes in leadership.")
p(f"The written guidance addresses: how the change will be communicated across "
  f"the organisation; who owns the change; and how the cultural aspects of the "
  f"transition will be managed.")
p(f"This is an Excellence-level objective element, assessed during "
  f"re-accreditation. The hospital keeps a record of significant changes managed "
  f"under this guidance.")

# ── 6. Governance and responsibility ──────────────────────────────────────────
h(1, "6. Governance and responsibility")
gt = tbl(6, 2)
for ri, (role, resp) in enumerate([
    ("Role", "Responsibility"),
    ("Medical Superintendent",
     "Accountable that this policy is followed and properly resourced."),
    ("Medical Superintendent",
     "Owns the strategic plan, operational plans, annual budget, committee "
     "records and day-to-day records for this policy."),
    ("Those responsible for governance",
     "Approve the strategic plan and annual budget; receive and act on "
     "committee effectiveness reviews."),
    ("Quality Coordinator",
     "Audits this policy periodically; keeps training records; coordinates "
     "management review meetings; supports governance packs this policy "
     "requires."),
    ("Departmental leaders",
     "Participate in strategic and operational planning; manage their "
     "department's contribution to defined goals; participate in committees "
     "as assigned."),
]):
    gt.cell(ri, 0).text = role
    gt.cell(ri, 1).text = resp

# ── 7. Quality monitoring ─────────────────────────────────────────────────────
h(1, "7. Quality monitoring")
p("The Quality Coordinator audits this policy periodically, checking a sample "
  "of records against each of the What-we-do steps.")
p("Documentary evidence is on file for each asterisked objective element "
  "(marked * in the Level column of Section 11).")
p("If an audit finds a gap that is still open after a defined timeframe, the hospital "
  "carries out a root-cause analysis and records the findings and corrective "
  "actions.")
p("This policy is reviewed every year, and sooner if the strategic plan, "
  "committee structure, or budget process changes significantly.")

# ── 8. Training and staff acknowledgement ─────────────────────────────────────
h(1, "8. Training and staff acknowledgement")
p("Staff covered by this policy are trained when they join, and again every "
  "year after that. Training covers the What-we-do steps and the non-negotiable "
  "rules.")
p("I have read the Policy on Professional Functioning of the Organisation of "
  "«Hospital Name». I will follow the processes described.")
p("Name: ___________________________ Designation: ___________________________")
p("Department / floor: ____________________ Date: ____________")
p("Signature: ___________________________")
p("(One row per staff member. The Quality Coordinator holds signed "
  "acknowledgements with the induction record.)")

# ── 9. Distribution ───────────────────────────────────────────────────────────
h(1, "9. Distribution")
p("Medical Superintendent; those responsible for governance; Quality "
  "Coordinator; departmental leaders; committee chairs; staff covered by ROM.5")

# ── 10. Abbreviations ─────────────────────────────────────────────────────────
h(1, "10. Abbreviations")
at = tbl(9, 2)
for ri, (abbr, meaning) in enumerate([
    ("Abbreviation", "Meaning"),
    ("CAPA", "Corrective and Preventive Action"),
    ("HCO", "Hospital (Full Accreditation programme under NABH Hospitals "
             "6th Edition)"),
    ("IPC", "Infection Prevention and Control"),
    ("NABH", "National Accreditation Board for Hospitals and Healthcare "
              "Providers"),
    ("OE", "Objective Element"),
    ("PSQ", "Patient Safety and Quality Improvement (NABH Hospitals "
             "6th Edition chapter)"),
    ("ROM", "Responsibilities of Management (NABH Hospitals 6th Edition "
             "chapter)"),
    ("SWOT", "Strengths, Weaknesses, Opportunities, Threats"),
]):
    at.cell(ri, 0).text = abbr
    at.cell(ri, 1).text = meaning

# ── 11. Traceability ──────────────────────────────────────────────────────────
h(1, "11. Traceability to NABH HCO Full Accreditation 6th Edition ROM.5")
p("This table is an index. It is not how the policy is organised. An asterisk "
  "in the Level column means documentation of the process is required.")

tt = tbl(7, 5)
for ci, txt in enumerate(("OE", "Level", "Requirement",
                           "Where addressed", "Responsible")):
    tt.cell(0, ci).text = txt

for ri, (oe, level, req, where, who) in enumerate([
    ("ROM.5.a", "Commitment",
     "The organisation has strategic and operational plans, including long-term "
     "and short-term goals commensurate to the organisation's vision, mission "
     "and values in consultation with the various stakeholders.",
     "Section 3; 5.1", "Medical Superintendent"),
    ("ROM.5.b", "Commitment*",
     "The organisation coordinates the functioning with departments and external "
     "agencies and monitors the progress in achieving the defined goals and "
     "objectives.",
     "Section 3; 5.2", "Medical Superintendent"),
    ("ROM.5.c", "Commitment",
     "The organisation plans and budgets for its activities annually.",
     "Section 3; 5.3", "Medical Superintendent"),
    ("ROM.5.d", "Achievement",
     "The functioning of committees is reviewed for their effectiveness.",
     "Section 3; 5.4", "Medical Superintendent"),
    ("ROM.5.e", "Commitment*",
     "The organisation documents the service standards that are measurable "
     "and monitors them.",
     "Section 3; 5.5", "Medical Superintendent"),
    ("ROM.5.f", "Excellence*",
     "Systems and processes are in place for change management.",
     "Section 3; 5.6", "Medical Superintendent"),
], start=1):
    tt.cell(ri, 0).text = oe
    tt.cell(ri, 1).text = level
    tt.cell(ri, 2).text = req
    tt.cell(ri, 3).text = where
    tt.cell(ri, 4).text = who

# ── 12. Required Records / Evidence Checklist ─────────────────────────────────
h(1, "12. Required Records / Evidence Checklist")
p("Records the hospital holds under this policy, listed by objective element.")

h(2, "ROM.5.a — Strategic and operational plans")
lb("Written strategic plan (vision/mission aligned, stakeholder consultation "
   "record, measurable goals with time frames).")
lb("Annual operational plan(s) drawn from the strategic plan.")

h(2, "ROM.5.b — Goal coordination and monitoring")
lb("Progress-review records against strategic and operational plan goals "
   "(periodic reviews and/or management review meeting minutes).")
lb("Analysis record for each goal not achieved, with the action taken.")

h(2, "ROM.5.c — Annual budget")
lb("Annual budget document covering the hospital's activities.")
lb("Budget-vs-actual review record.")
lb("Evidence of an adequate allocation for IPC activities and quality "
   "improvement within the budget.")

h(2, "ROM.5.d — Committee effectiveness")
lb("Minutes for each committee meeting.")
lb("Leadership review record assessing each committee's effectiveness.")

h(2, "ROM.5.e — Service standards")
lb("Documented service-standard benchmarks (measurable, with responsible "
   "person and monitoring frequency defined).")
lb("Monitoring records for each benchmark.")

h(2, "ROM.5.f — Change management")
lb("Written change-management guidance document.")
lb("Record of significant changes managed under the guidance.")

# ── 13. References ────────────────────────────────────────────────────────────
h(1, "13. References")
lb("National Accreditation Board for Hospitals and Healthcare Providers "
   "(NABH), Accreditation Standards for Hospitals, 6th Edition "
   "(January 2025) — Responsibilities of Management, standard ROM.5.")
lb("NABH Guidebook to Accreditation Standards for Hospitals, 6th Edition "
   "— ROM.5 interpretations.")
lb(f"Internal documents of {HN}: strategic plan, operational plans, annual "
   "budget, management review meeting minutes, committee minutes, "
   "service-standard monitoring records, change-management guidance.")

# ── Disclaimer ────────────────────────────────────────────────────────────────
h(1, "Disclaimer")
p(f"This document is a template prepared for the guidance of {HN} and must be "
  f"reviewed, adapted and formally approved by {HN} before use. Every entry "
  f"marked ________ must be completed before the document is issued.")
p("The requirements in this document are accreditation requirements of the "
  "NABH Accreditation Standards for Hospitals, 6th Edition, not duties under "
  "a named Act of Parliament. This policy does not import the Consumer "
  "Protection Act, 2019, the Clinical Establishments Act, 2010, or the Mental "
  "Healthcare Act, 2017 as a checklist. Statutory duties that arise under other "
  "documents of «Hospital Name» remain those documents. "
  "«Hospital Name» is responsible for verifying any statutory duty "
  "that applies to it; this document does not constitute legal advice.")
p("«Hospital Name» remains responsible for verifying that it is "
  "current and consistent with the edition of the accreditation standard "
  "against which it is being assessed. The clinical and technical content "
  "reflects recognised national and international guidance current at the date "
  "of preparation.")
p("This document is not issued by, endorsed by, or affiliated with NABH, the "
  "World Health Organization, the National Centre for Disease Control, the Food "
  "Safety and Standards Authority of India, any Pollution Control Board, or "
  "any other body named in it. Wording is original; no text has been reproduced "
  "from the standards, rules or guidelines referenced.")

# ── Verify styles ─────────────────────────────────────────────────────────────
print("=== Paragraph style verification ===")
for i, para in enumerate(doc.paragraphs):
    print(f"{i:3d}  {para.style.name!r:32s}  {para.text[:55]!r}")

print(f"\nTotal paragraphs: {len(doc.paragraphs)}")

# ── Save ──────────────────────────────────────────────────────────────────────
doc.save(OUT)
print(f"\nSaved: {OUT}")
