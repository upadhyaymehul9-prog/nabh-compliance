# -*- coding: utf-8 -*-
"""
Generate HCO_ROM_3_v2_REWRITE_DRAFT.docx
Plain-language rewrite per HCO_Policy_Rewrite_Rulebook.pdf
No stop-work: Governance=6, QM=7, Training=8, Distribution=9,
Abbreviations=10, Traceability=11, Records=12, References=13, Disclaimer (unnumbered)
"""
from docx import Document

OUT = "policies/build/rewrite_reference/HCO_ROM_3_v2_REWRITE_DRAFT.docx"

HN = "«Hospital Name»"   # «Hospital Name»
EM = "—"                        # —
doc = Document()


def h(level, text):
    styles = {0: "Title", 1: "Heading 1", 2: "Heading 2"}
    return doc.add_paragraph(text, style=styles[level])


def p(text):
    return doc.add_paragraph(text, style="Normal")


def ln(text):
    return doc.add_paragraph(text, style="List Number")


def lb(text):
    return doc.add_paragraph(text, style="List Bullet")


def tbl(rows, cols):
    t = doc.add_table(rows=rows, cols=cols)
    try:
        t.style = "Table Grid"
    except KeyError:
        pass
    return t


# ── Title ─────────────────────────────────────────────────────────────────────
h(0, "Policy on Sustainability, Social Responsibility and Staff Well-being")
p(HN)

# ── Document control ──────────────────────────────────────────────────────────
h(1, "Document control")

dc = tbl(6, 4)
for ri, (a, b, c, d) in enumerate([
    ("Document No.", "HCO/ROM/POL/03", "Version", "2.0"),
    ("Issue No.", "01", "Review due", "One year from implementation"),
    ("Date created", "________", "Date of implementation", "________"),
]):
    dc.cell(ri, 0).text = a
    dc.cell(ri, 1).text = b
    dc.cell(ri, 2).text = c
    dc.cell(ri, 3).text = d

for ri, (label, content) in enumerate([
    ("Prepared by",
     "Medical Superintendent Name: ________ Signature: ________"),
    ("Reviewed by",
     "Quality Coordinator Name: ________ Signature: ________"),
    ("Approved by",
     "Medical Superintendent Name: ________ Signature: ________"),
], start=3):
    dc.cell(ri, 0).text = label
    dc.cell(ri, 1).text = content
    dc.cell(ri, 1).merge(dc.cell(ri, 3))

p("A blank marked ________ must be completed before issue.")

# ── Statement of intent ───────────────────────────────────────────────────────
h(1, "Statement of intent")
p(f"{HN} is a responsible organisation. It runs a written sustainability "
  f"programme, takes practical steps to reduce its environmental impact, meets "
  f"its obligations to the community, looks after its staff’s well-being, "
  f"and monitors its financial position. These are active commitments with "
  f"records and regular review {EM} not statements of aspiration.")

# ── 1. Purpose ────────────────────────────────────────────────────────────────
h(1, "1. Purpose")
p(f"This policy explains how {HN} manages its environmental, social and "
  f"governance (ESG) sustainability programme; acts to become a more "
  f"energy-efficient and environmentally friendly hospital; fulfils its social "
  f"responsibility to the community it serves; promotes the well-being of its "
  f"staff; applies sustainability criteria in procurement; encourages staff to "
  f"use common or public transport; and ensures its long-term financial "
  f"sustainability.")
p("This policy does not cover patient assessment, clinical care, medication "
  "management, patient rights, infection control, or quality and safety "
  "monitoring — those are covered in other hospital policies. The other "
  "Responsibilities of Management standards have their own policies too.")

# ── 2. Scope ──────────────────────────────────────────────────────────────────
h(1, "2. Scope")
p(f"This policy applies to those responsible for governance (as identified "
  f"under ROM.1), the Medical Superintendent, departmental leaders, the Quality "
  f"Coordinator, and all staff at {HN} whose roles involve sustainability "
  f"programmes, procurement, staff welfare, or financial planning and review.")

# ── 3. Policy standards ───────────────────────────────────────────────────────
h(1, "3. Policy standards")
p(f"{HN} maintains a documented ESG sustainability programme, takes active "
  f"steps to reduce its environmental footprint, meets its social-responsibility "
  f"obligations, supports the well-being of its staff, applies sustainability "
  f"criteria when procuring goods and services, encourages environmentally "
  f"responsible commuting, and reviews its financial sustainability regularly.")
p("Staff follow the written guidance below, keep the records it requires, and "
  "raise concerns through the channels it describes.")

# ── 4. Non-negotiable rules ───────────────────────────────────────────────────
h(1, "4. Non-negotiable rules")
ln("The sustainability programme must be written — a verbal commitment "
   "or a meeting agenda item is not enough.")
ln("Do not leave energy-efficiency and environmental initiatives undocumented "
   "— record what was done, when it was done, and what result it achieved.")
ln("Do not treat social responsibility as fulfilled by a policy statement alone "
   "— the programme must be implemented and its activity must be documented.")
ln("Do not leave staff without accessible facilities and a working channel for "
   "raising workload concerns — these must be in place, not just promised.")
ln("Do not encourage staff to use public or common transport without a record "
   "showing the encouragement was actually communicated to them.")
ln("Do not treat financial sustainability as a declaration — the hospital "
   "must have a working process for reviewing how it balances resource "
   "allocation, costs and revenue, and must act on what it finds.")
ln("Staff who see any of these rules broken report it the same shift to the "
   "department in-charge or the Medical Superintendent.")

# ── 5. What we do ─────────────────────────────────────────────────────────────
h(1, "5. What we do")

h(2, "5.1 Address the ESG sustainability programme")
p(f"Those responsible for governance maintain a written ESG sustainability "
  f"programme for {HN}. The programme is a written document, reviewed at least "
  f"annually and tabled at a governance meeting.")
p("The Guidebook defines the three dimensions of the programme. Environmental "
  "sustainability covers energy usage and efficiency, climate change strategy, "
  "waste reduction, biodiversity loss, and greenhouse gas emissions. Social "
  "sustainability covers fair pay, equal employment opportunity, employee "
  "benefits, workplace health and safety, community engagement, and responsible "
  "supply-chain partnerships. Governance sustainability covers corporate "
  "governance, risk management, compliance, ethical business practices, and "
  "accounting integrity.")
p("A slogan or a PowerPoint is not this programme. It is a document.")

h(2, "5.2 Take initiatives toward an energy-efficient and environmentally "
     "friendly hospital")
p(f"{HN} focuses on the efficient and sustainable use of energy, water and "
  f"other utilities. It takes active measures to create awareness among staff "
  f"and patients about saving electricity and water.")
p("The hospital documents its initiatives — what was done, when it "
  "started, and what was achieved. Initiatives may include energy-efficient "
  "lighting, rainwater harvesting, solar or wind energy, battery-operated or "
  "electric vehicles, reuse of treated water for gardening or flushing, "
  "reducing plastic use, and green building materials — these are "
  "examples, not a mandatory checklist. The hospital records the ones it has "
  "taken and the results.")

h(2, "5.3 Address social responsibility")
p(f"Those responsible for governance and the leadership develop and implement "
  f"a written social-responsibility programme. The programme is documented, and "
  f"activity against it is reported at least annually.")
p("The programme addresses the hospital’s obligations to the community it "
  "serves. Examples include free health camps and outreach programmes. At a "
  "minimum, the hospital meets all applicable regulatory requirements, including "
  "those relating to corporate social responsibility.")

h(2, "5.4 Promote staff well-being")
p(f"{HN} takes proactive steps to promote the well-being of its staff. The "
  f"hospital defines work hours, monitors workload, provides scheduled breaks, "
  f"and maintains accessible facilities where staff can seek support and advice "
  f"when they need it.")
p("Examples of well-being measures include healthy lifestyle programmes, stress "
  "management support, access to dining facilities, recognition schemes, and "
  "staff engagement activities — the hospital uses the ones that fit its "
  "context. Monitoring absenteeism or overtime can help identify stress and "
  "fatigue as a management tool, as can staff satisfaction surveys; these are "
  "options, not requirements. What is required is that the support facilities "
  "and the workload-concerns channel are in place and working.")

h(2, "5.5 Follow sustainable procurement practices")
p(f"{HN} applies environmental and social criteria when evaluating tenders and "
  f"selecting suppliers. The hospital’s purchase policy names those "
  f"criteria. Completed procurement files show that the criteria were applied "
  f"when a purchase was made.")
p("Working with suppliers to prioritise environmentally friendly and socially "
  "responsible products is one approach the hospital may use — the "
  "purchase policy specifies how.")

h(2, "5.6 Encourage staff to use common or public transport")
p(f"{HN} actively encourages employees to use common or public transport to "
  f"reduce the environmental impact of commuting. That encouragement is "
  f"documented and communicated to staff.")
p("Examples of encouragement include providing information about carpooling, "
  "cycling and public transport options, and any incentive scheme the hospital "
  "chooses to operate. The obligation is on the hospital to encourage and to "
  "show that the message reached staff — it is not a requirement that every "
  "individual must commute by public transport. Environmentally friendly "
  "logistics and transport management within the hospital can further contribute "
  "to sustainability.")

h(2, "5.7 Ensure financial sustainability")
p(f"{HN} ensures financial sustainability by actively balancing the financial "
  f"aspects of healthcare delivery — efficient resource allocation, cost "
  f"management and revenue generation. The hospital has a defined process for "
  f"reviewing whether this balance is being maintained, and it acts on the "
  f"findings.")
p("The Guidebook notes that this balance is crucial for long-term "
  "sustainability. The specific form the review takes is for the hospital to "
  "define; the obligation is that it happens.")

# ── 6. Governance and responsibility ──────────────────────────────────────────
h(1, "6. Governance and responsibility")
gt = tbl(6, 2)
for ri, (role, resp) in enumerate([
    ("Role", "Responsibility"),
    ("Medical Superintendent",
     "Accountable that this policy is followed and properly resourced."),
    ("Medical Superintendent",
     "Owns the day-to-day records and implementation for this policy."),
    ("Those responsible for governance",
     "Ensure the ESG programme, social-responsibility programme and "
     "financial-sustainability review are produced, tabled and acted upon."),
    ("Quality Coordinator",
     "Audits this policy periodically; keeps training records; supports "
     "governance packs and committee reviews this policy requires."),
    ("Departmental leaders",
     "Carry out the department-level duties this policy names."),
]):
    gt.cell(ri, 0).text = role
    gt.cell(ri, 1).text = resp

# ── 7. Quality monitoring ─────────────────────────────────────────────────────
h(1, "7. Quality monitoring")
p("The Quality Coordinator audits this policy periodically, checking a sample "
  "of records against each of the What-we-do steps.")
p("Documentary evidence is on file for each asterisked objective element "
  "(marked * in the Level column of Section 11).")
p("If an audit finds a gap that is still open after a defined timeframe, the hospital "
  "carries out a root-cause analysis and records the findings and corrective "
  "actions.")
p("This policy is reviewed every year, and sooner if the ESG programme, "
  "social-responsibility commitments, or the hospital’s financial position "
  "change significantly.")

# ── 8. Training and staff acknowledgement ─────────────────────────────────────
h(1, "8. Training and staff acknowledgement")
p("Staff covered by this policy are trained when they join, and again every "
  "year after that. Training covers the What-we-do steps and the non-negotiable "
  "rules.")
p("I have read the Policy on Sustainability, Social Responsibility and Staff "
  "Well-being of «Hospital Name». I will follow the processes "
  "described.")
p("Name: ___________________________ Designation: ___________________________")
p("Department / floor: ____________________ Date: ____________")
p("Signature: ___________________________")
p("(One row per staff member. The Quality Coordinator holds signed "
  "acknowledgements with the induction record.)")

# ── 9. Distribution ───────────────────────────────────────────────────────────
h(1, "9. Distribution")
p("Medical Superintendent; those responsible for governance; Quality "
  "Coordinator; departmental leaders; staff covered by ROM.3")

# ── 10. Abbreviations ─────────────────────────────────────────────────────────
h(1, "10. Abbreviations")
at = tbl(6, 2)
for ri, (abbr, meaning) in enumerate([
    ("Abbreviation", "Meaning"),
    ("ESG", "Environment, Social and Governance"),
    ("HCO", "Hospital (Full Accreditation programme under NABH Hospitals "
             "6th Edition)"),
    ("NABH", "National Accreditation Board for Hospitals and Healthcare "
              "Providers"),
    ("OE", "Objective Element"),
    ("ROM", "Responsibilities of Management (NABH Hospitals 6th Edition "
             "chapter)"),
]):
    at.cell(ri, 0).text = abbr
    at.cell(ri, 1).text = meaning

# ── 11. Traceability ──────────────────────────────────────────────────────────
h(1, "11. Traceability to NABH HCO Full Accreditation 6th Edition ROM.3")
p("This table is an index. It is not how the policy is organised. An asterisk "
  "in the Level column means documentation of the process is required.")

tt = tbl(8, 5)
for ci, txt in enumerate(("OE", "Level", "Requirement",
                           "Where addressed", "Responsible")):
    tt.cell(0, ci).text = txt

for ri, (oe, level, req, where, who) in enumerate([
    ("ROM.3.a", "Commitment",
     "Those responsible for governance address the organisation’s "
     "sustainability programme in terms of Environment Social and Governance "
     "(ESG) responsibility.",
     "Section 3; 5.1", "Medical Superintendent"),
    ("ROM.3.b", "Commitment*",
     "The organisation takes initiatives towards an energy-efficient and "
     "environmentally friendly hospital.",
     "Section 3; 5.2", "Medical Superintendent"),
    ("ROM.3.c", "Commitment",
     "Those responsible for governance address the organisations social "
     "responsibility.",
     "Section 3; 5.3", "Medical Superintendent"),
    ("ROM.3.d", "Commitment",
     "Staff well-being is promoted.",
     "Section 3; 5.4", "Medical Superintendent"),
    ("ROM.3.e", "Excellence",
     "The organisation follows sustainable procurement practices.",
     "Section 3; 5.5", "Medical Superintendent"),
    ("ROM.3.f", "Achievement",
     "Hospitals shall encourage employees to use common / public "
     "transportation to reduce the environmental impact of commuting and "
     "carbon footprint.",
     "Section 3; 5.6", "Medical Superintendent"),
    ("ROM.3.g", "Achievement",
     "The organisation ensures financial sustainability of the hospital by "
     "balancing the financial aspects of healthcare delivery.",
     "Section 3; 5.7", "Medical Superintendent"),
], start=1):
    tt.cell(ri, 0).text = oe
    tt.cell(ri, 1).text = level
    tt.cell(ri, 2).text = req
    tt.cell(ri, 3).text = where
    tt.cell(ri, 4).text = who

# ── 12. Required Records / Evidence Checklist ─────────────────────────────────
h(1, "12. Required Records / Evidence Checklist")
p("Records the hospital holds under this policy, listed by objective element.")

h(2, "ROM.3.a — Sustainability programme (ESG)")
lb("Written ESG sustainability programme.")
lb("Governance-meeting agenda or minutes tabling the programme.")
lb("Annual review record of the programme.")

h(2, "ROM.3.b — Energy-efficiency and environmental initiatives")
lb("Dated log of energy-efficiency and environmental initiatives and their "
   "results.")
lb("Cross-reference to the engineering file (facility-management controls and "
   "this initiative record are held separately).")
lb("Before/after data or savings evidence where available.")

h(2, "ROM.3.c — Social responsibility")
lb("Documented social-responsibility programme (community access, charity "
   "care, local health initiatives as the hospital defines).")
lb("Annual activity report for the programme.")

h(2, "ROM.3.d — Staff well-being")
lb("Documented work-hour monitoring method for staff.")
lb("Record of healthy-lifestyle support and scheduled-break arrangements.")
lb("Record of the channel for raising workload concerns and any concerns "
   "logged through it.")

h(2, "ROM.3.e — Sustainable procurement")
lb("Purchase policy naming environmental and social tender criteria.")
lb("Completed procurement file showing those criteria were applied.")
lb("At least one tender evaluation reflecting the sustainable-procurement "
   "criteria.")

h(2, "ROM.3.f — Sustainable commuting")
lb("Current staff communication promoting common or public transport for "
   "commuting.")
lb("Documentation of any incentive scheme offered.")
lb("Record that the communication was issued to staff (not that every "
   "individual complied).")

h(2, "ROM.3.g — Financial sustainability")
lb("Evidence that the hospital reviews how it balances resource allocation, "
   "costs and revenue (form defined by the hospital).")
lb("Evidence that the review findings were acted on or noted.")

# ── 13. References ────────────────────────────────────────────────────────────
h(1, "13. References")
lb("National Accreditation Board for Hospitals and Healthcare Providers "
   "(NABH), Accreditation Standards for Hospitals, 6th Edition "
   "(January 2025) — Responsibilities of Management, standard ROM.3.")
lb("NABH Guidebook to Accreditation Standards for Hospitals, 6th Edition "
   "— ROM.3 interpretations.")
lb(f"Internal documents of {HN}: ESG sustainability programme, "
   "social-responsibility programme, purchase policy, staff well-being "
   "records, financial-sustainability reviews.")

# ── Disclaimer ────────────────────────────────────────────────────────────────
h(1, "Disclaimer")
p(f"This document is a template prepared for the guidance of {HN} and must be "
  f"reviewed, adapted and formally approved by {HN} before use. Every entry "
  f"marked ________ must be completed before the document is issued.")
p("The requirements in this document are accreditation requirements of the "
  "NABH Accreditation Standards for Hospitals, 6th Edition, not duties under "
  "a named Act of Parliament. This policy does not import the Consumer "
  "Protection Act, 2019, the Clinical Establishments Act, 2010, or the Mental "
  "Healthcare Act, 2017 as a checklist. Statutory duties that arise under other "
  "documents of «Hospital Name» remain those documents. "
  "«Hospital Name» is responsible for verifying any statutory duty "
  "that applies to it; this document does not constitute legal advice.")
p("«Hospital Name» remains responsible for verifying that it is "
  "current and consistent with the edition of the accreditation standard "
  "against which it is being assessed. The clinical and technical content "
  "reflects recognised national and international guidance current at the date "
  "of preparation.")
p("This document is not issued by, endorsed by, or affiliated with NABH, the "
  "World Health Organization, the National Centre for Disease Control, the Food "
  "Safety and Standards Authority of India, any Pollution Control Board, or "
  "any other body named in it. Wording is original; no text has been reproduced "
  "from the standards, rules or guidelines referenced.")

# ── Verify styles ─────────────────────────────────────────────────────────────
print("=== Paragraph style verification (first 40) ===")
for i, para in enumerate(doc.paragraphs[:40]):
    print(f"{i:3d}  {para.style.name!r:32s}  {para.text[:55]!r}")

print(f"\nTotal paragraphs: {len(doc.paragraphs)}")

# ── Save ──────────────────────────────────────────────────────────────────────
doc.save(OUT)
print(f"\nSaved: {OUT}")
