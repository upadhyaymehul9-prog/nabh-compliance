# -*- coding: utf-8 -*-
"""
generate_hco_hrm_rewrites.py
Generates HCO HRM.1-5 v2 rewrite-reference DOCX files.

Pipeline : python-docx, identical to generate_hco_fms_rewrites.py.
Output   : policies/build/rewrite_reference/HCO_HRM_N_v2_REWRITE_DRAFT.docx
Source   : hrm1_content.txt … hrm5_content.txt (Downloads) — content
           embedded directly as string literals below.
"""
import os
from docx import Document

HN  = "«Hospital Name»"
OUT = "policies/build/rewrite_reference"
os.makedirs(OUT, exist_ok=True)


# ── Helpers (identical to generate_hco_fms_rewrites.py) ───────────────────────

def h(doc, lv, txt):
    return doc.add_paragraph(txt, style={0: "Title", 1: "Heading 1", 2: "Heading 2"}[lv])

def p(doc, txt=""):
    return doc.add_paragraph(txt, style="Normal")

def ln(doc, txt):
    return doc.add_paragraph(txt, style="List Number")

def lb(doc, txt):
    return doc.add_paragraph(txt, style="List Bullet")

def tbl(doc, rows, cols):
    t = doc.add_table(rows=rows, cols=cols)
    try:
        t.style = "Table Grid"
    except KeyError:
        pass
    return t

def doc_ctrl(doc, no, prep, appr="Medical Superintendent"):
    dc = tbl(doc, 6, 4)
    for ri, (a, b, c, d) in enumerate([
        ("Document No.", no,        "Version",                "1.0"),
        ("Issue No.",    "01",       "Review due",             "One year from implementation"),
        ("Date created", "________", "Date of implementation", "________"),
    ]):
        dc.cell(ri, 0).text = a; dc.cell(ri, 1).text = b
        dc.cell(ri, 2).text = c; dc.cell(ri, 3).text = d
    for ri, (lbl, txt) in enumerate([
        ("Prepared by", f"{prep}  Name: ________  Signature: ________"),
        ("Reviewed by", "Quality Coordinator  Name: ________  Signature: ________"),
        ("Approved by", f"{appr}  Name: ________  Signature: ________"),
    ], start=3):
        dc.cell(ri, 0).text = lbl
        c1 = dc.cell(ri, 1); c1.text = txt; c1.merge(dc.cell(ri, 3))

def gov_tbl(doc, rows):
    t = tbl(doc, len(rows) + 1, 2)
    t.cell(0, 0).text = "Role"; t.cell(0, 1).text = "Responsibility"
    for ri, (role, resp) in enumerate(rows, 1):
        t.cell(ri, 0).text = role; t.cell(ri, 1).text = resp

def abbrev_tbl(doc, rows):
    t = tbl(doc, len(rows) + 1, 2)
    t.cell(0, 0).text = "Abbreviation"; t.cell(0, 1).text = "Meaning"
    for ri, (a, m) in enumerate(rows, 1):
        t.cell(ri, 0).text = a; t.cell(ri, 1).text = m

def trace_tbl(doc, rows):
    t = tbl(doc, len(rows) + 1, 4)
    for ci, hdr in enumerate(("OE", "Level", "Where addressed", "Responsible")):
        t.cell(0, ci).text = hdr
    for ri, row in enumerate(rows, 1):
        for ci, v in enumerate(row):
            t.cell(ri, ci).text = v

def sig_tbl(doc):
    t = tbl(doc, 4, 4)
    for ci, hdr in enumerate(("Staff name", "Designation", "Signature", "Date")):
        t.cell(0, ci).text = hdr
    for ri in range(1, 4):
        for ci in range(4):
            t.cell(ri, ci).text = "________"

def hrm_disclaimer(doc):
    p(doc,
      f"This document is a template prepared for the guidance of {HN} and must be reviewed, "
      f"adapted and formally approved by {HN} before use. Every entry marked ________ "
      f"must be completed before the document is issued.")
    p(doc,
      f"The requirements in this document are accreditation requirements of the NABH Full "
      f"Accreditation Standards for Hospitals, 6th Edition, not duties under a named Act of "
      f"Parliament. Statutory duties that arise under other documents of {HN} remain those "
      f"documents. {HN} is responsible for verifying any statutory duty that applies to it; "
      f"this document does not constitute legal advice.")
    p(doc,
      f"The clinical and technical content reflects recognised national and international guidance "
      f"current at the date of preparation. {HN} remains responsible for verifying that it is "
      f"current and consistent with the edition of the accreditation standard against which it "
      f"is being assessed.")
    p(doc,
      "This document is not issued by, endorsed by, or affiliated with NABH or any other body "
      "named in it. Wording is original; no text has been reproduced from the standards, rules "
      "or guidelines referenced.")

def save_and_verify(doc, fname):
    import sys
    out = sys.stdout
    def pr(s):
        try:
            out.write(s + "\n")
        except UnicodeEncodeError:
            out.write(s.encode("ascii", "replace").decode() + "\n")
    pr(f"\n=== {fname} ===")
    pr(f"  Total paras: {len(doc.paragraphs)}")
    path = os.path.join(OUT, fname)
    doc.save(path)
    pr(f"  Saved: {path}")

# Shared constants
_HR  = "HR In-Charge / Personnel Officer"
_MS  = "Medical Superintendent"
_QC  = "Quality Coordinator"

HRM_ABBREVS_BASE = [
    ("HCO",  "Hospital (Full Accreditation programme under NABH Hospitals 6th Edition)"),
    ("HR",   "Human Resources"),
    ("HRM",  "Human Resource Management (NABH Hospitals 6th Edition chapter)"),
    ("NABH", "National Accreditation Board for Hospitals and Healthcare Providers"),
    ("OE",   "Objective Element"),
]


# ══════════════════════════════════════════════════════════════════════════════
# HRM.1 — Human Resource Planning and Governance   (no stop-work)
# ══════════════════════════════════════════════════════════════════════════════
def gen_hrm1():
    doc = Document()
    h(doc, 0, "Policy on Human Resource Planning and Governance")
    p(doc, HN)
    h(doc, 1, "Document control")
    doc_ctrl(doc, "HCO/HRM/POL/01", _HR)
    p(doc, "A blank marked ________ must be completed before issue.")
    h(doc, 1, "Statement of intent")
    p(doc, f"{HN} plans its workforce so that the right number and mix of skilled staff is "
           "available now and in the future, not just when a crisis makes the gap obvious.")
    h(doc, 1, "1. Purpose")
    p(doc, f"This policy explains how {HN} plans staffing to meet patient care needs, defines "
           "jobs and reporting lines, checks new staff before they start, and learns from staff who leave.")
    p(doc, "This policy does not cover recruitment steps, induction training, or staff appraisal "
           "— those are covered in other hospital policies.")
    h(doc, 1, "2. Scope")
    p(doc, f"This policy applies to the HR In-Charge / Personnel Officer, department heads, and "
           f"the Medical Superintendent, for every category of staff — full-time, part-time, "
           f"employed, honorary, voluntary and temporary.")
    h(doc, 1, "3. Policy standards")
    p(doc, f"{HN} prepares an annual workforce plan against patient volume, services and technology "
           "in use; keeps enough staff, in the right mix, to meet patient needs; has a tested "
           "contingency plan for staff shortages; defines a job description for every staff category; "
           "background-checks new staff; documents reporting lines for every category of staff; and "
           "uses exit interviews to improve HR practice.")
    h(doc, 1, "4. Non-negotiable rules")
    ln(doc, "Do not let the workforce plan go a year without being compared against actual patient "
            "volume, services and staffing.")
    ln(doc, "Do not leave a staffing shortfall unaddressed — trigger the contingency plan.")
    ln(doc, "Do not let any staff category go without a defined job description and specification.")
    ln(doc, "Do not skip the background check on a new staff member, or delay it beyond one month of joining.")
    ln(doc, "Do not leave a staff role without a documented reporting line.")
    ln(doc, "Do not let exit-interview findings sit unread — they must feed back into HR practice.")
    ln(doc, f"Staff who see any of these rules broken report it the same shift to the "
            f"HR In-Charge / Personnel Officer or the Medical Superintendent.")
    h(doc, 1, "5. What we do")
    h(doc, 2, "5.1 Plan the workforce every year")
    p(doc, f"The HR In-Charge / Personnel Officer prepares a workforce plan every year, comparing "
           "current staff numbers and skill mix against the hospital's mission, patient volume and mix, "
           "services offered, and medical technology in use, with input from department heads. Recognised "
           "staffing methods are used to set levels against the strategic and operational plan. Where the "
           "year's actual staffing varies from the plan, the corrective action taken is recorded and "
           "carried into the next year's plan.")
    h(doc, 2, "5.2 Keep enough staff, in the right mix")
    p(doc, f"{HN} keeps a number and mix of staff commensurate with workload and clinical need. "
           "Nursing staffing follows a published guideline (for example, WHO's Workload Indicators of "
           "Staffing Need method, or another recognised nursing staffing guideline). The "
           "HR In-Charge / Personnel Officer compares sanctioned against actual strength by department "
           "on a regular schedule and escalates an unresolved shortfall to the Medical Superintendent. "
           "Any shortfall triggers the contingency plan in 5.3.")
    h(doc, 2, "5.3 Plan for staff shortages before they happen")
    p(doc, f"The hospital keeps a written contingency plan for shift-by-shift, short-term and "
           "long-term workforce shortages, including unplanned ones. The plan may include reprioritising "
           "tasks, reallocating tasks across available staff, and drawing on a pool of filler staff "
           "such as previous employees or agency-sourced casual staff. Every shortage event is logged "
           "with its cause, the measure used, and the outcome. The plan itself is tested at a regular "
           "interval to confirm it still works.")
    h(doc, 2, "5.4 Define every job")
    p(doc, f"Every staff category at {HN} — full-time, part-time, employed, honorary, voluntary or "
           "temporary — has a job description and specification laying down the job's content and the "
           "qualifications, skills and experience needed. A role requiring the skills of a doctor or "
           "nurse carries a minimum qualification of MBBS or GNM respectively, unless a government or "
           "statutory body has granted an exemption. The HR In-Charge / Personnel Officer holds the "
           "current set of job descriptions; a new hire signs to confirm they have received and "
           "understood theirs.")
    h(doc, 2, "5.5 Check new staff before they start")
    p(doc, f"{HN} performs a background check on every new staff member, using a defined method, "
           "either before the person joins or within one month of joining. The "
           "HR In-Charge / Personnel Officer keeps a register recording the method used, the date "
           "completed, and the outcome for each new hire. A staff member with no background check on "
           "file is escalated to the Medical Superintendent — this is not left for later.")
    h(doc, 2, "5.6 Document who reports to whom")
    p(doc, "Reporting relationships are defined for every category of staff, documented as an "
           "organisation chart showing hierarchy, line of control and functions at each level. The "
           "chart is disseminated to all stakeholders, and reporting relationships are also defined "
           "at department or service level. The HR In-Charge / Personnel Officer holds the current chart.")
    h(doc, 2, "5.7 Learn from staff who leave")
    p(doc, f"{HN} conducts exit interviews and uses them to improve HR practice. A personal interview "
           "is the default method; taking part is voluntary for the departing staff member. The "
           "HR In-Charge / Personnel Officer compiles findings at a regular interval into a trend "
           "report for the Medical Superintendent and, where the hospital has one, the governing body, "
           "with proposed HR improvements tracked to closure.")
    h(doc, 1, "6. Governance and responsibility")
    gov_tbl(doc, [
        (_HR,
         "Owns the workforce plan, staffing comparison, contingency plan, job descriptions, "
         "background checks, organisation chart, and exit-interview trend reports."),
        ("Department heads",
         "Provide input to the workforce plan; confirm reporting lines and job content for their department."),
        (_MS,
         "Accountable that this policy is followed; receives escalations on unresolved staffing "
         "shortfalls and exit-interview trend reports."),
    ])
    h(doc, 1, "7. Quality monitoring")
    p(doc, "The HR In-Charge / Personnel Officer reviews a sample of records under this policy "
           "periodically against the steps in Section 5, confirming the workforce plan, contingency-plan "
           "tests, job descriptions, background checks, the organisation chart, and exit-interview trend "
           "reports are current. This policy itself is reviewed every year, and sooner if there is a "
           "major change in hospital services or staffing structure.")
    h(doc, 1, "8. Training and staff acknowledgement")
    p(doc, f"Staff covered by this policy — the HR In-Charge / Personnel Officer and department heads "
           "— are trained on this policy when they take up the role, and again every year after that. "
           "Training covers what's in Section 5 and the non-negotiable rules.")
    p(doc, f"I have read the Policy on Human Resource Planning and Governance of {HN}. I will follow the processes described.")
    sig_tbl(doc)
    h(doc, 1, "9. Distribution")
    p(doc, "HR In-Charge / Personnel Officer; department heads; Medical Superintendent; governing body, "
           "where applicable.")
    h(doc, 1, "10. Abbreviations")
    abbrev_tbl(doc, [
        ("GNM",  "General Nursing and Midwifery"),
    ] + HRM_ABBREVS_BASE + [
        ("MBBS", "Bachelor of Medicine, Bachelor of Surgery"),
        ("WISN", "Workload Indicators of Staffing Need (WHO method)"),
    ])
    h(doc, 1, "11. Traceability to NABH HCO Full Accreditation 6th Edition HRM.1")
    p(doc, "This table is an index. It is not how the policy is organised. An asterisk in the "
           "Level column means documentation of the process is required.")
    trace_tbl(doc, [
        ("HRM.1.a", "Commitment",
         "Section 3; 5.1", _HR),
        ("HRM.1.b", "CORE*",
         "Section 3; 5.2", _HR),
        ("HRM.1.c", "Achievement",
         "Section 3; 5.3", _HR),
        ("HRM.1.d", "Commitment",
         "Section 3; 5.4", _HR),
        ("HRM.1.e", "Commitment",
         "Section 3; 5.5", _HR),
        ("HRM.1.f", "Commitment*",
         "Section 3; 5.6", _HR),
        ("HRM.1.g", "Achievement",
         "Section 3; 5.7", _HR),
    ])
    h(doc, 1, "12. Required Records / Evidence Checklist")
    h(doc, 2, "HRM.1.a — Human resource planning supports current and future patient needs.")
    lb(doc, "Annual workforce plan comparing current and projected staffing against services and patient volume.")
    lb(doc, "Department-head input record into the plan.")
    lb(doc, "Corrective-action record for a variance found during the year.")
    h(doc, 2, "HRM.1.b — Adequate number and mix of staff maintained.")
    lb(doc, "Sanctioned-versus-actual staffing comparison record by department.")
    lb(doc, "Staffing-norm reference used for nursing.")
    lb(doc, "Escalation record for an unresolved shortfall.")
    h(doc, 2, "HRM.1.c — Contingency plans for workforce shortages.")
    lb(doc, "Written contingency plan for long- and short-term workforce shortages, including unplanned shortages.")
    lb(doc, "Shortage-event log with cause, measure used and outcome.")
    lb(doc, "Test record of the contingency plan.")
    h(doc, 2, "HRM.1.d — Job specification and job description defined.")
    lb(doc, "Job description on file for each staff category, including qualification, skill and experience requirements.")
    lb(doc, "Signed acknowledgement record from a new hire receiving their job description.")
    lb(doc, "Minimum-qualification exemption record, where applicable.")
    h(doc, 2, "HRM.1.e — Background check of new staff.")
    lb(doc, "Background-check register recording method, date and outcome per new hire.")
    lb(doc, "Completion record before or within one month of joining.")
    lb(doc, "Escalation record for any staff member with no background check on file.")
    h(doc, 2, "HRM.1.f — Reporting relationships defined.")
    lb(doc, "Current organisation structure or chart showing hierarchy and reporting lines.")
    lb(doc, "Department- or service-level reporting-relationship record.")
    lb(doc, "Dissemination record to stakeholders.")
    h(doc, 2, "HRM.1.g — Exit interviews conducted and used.")
    lb(doc, "Completed exit-interview record for a departing staff member.")
    lb(doc, "Trend report compiled from exit interviews.")
    lb(doc, "HR-improvement action record from exit-interview findings.")
    h(doc, 1, "13. References")
    lb(doc, "NABH Accreditation Standards for Hospitals, 6th Edition — standard HRM.1.")
    lb(doc, "NABH Guidebook to Accreditation Standards for Hospitals, 6th Edition — HRM.1 interpretations.")
    h(doc, 1, "Disclaimer")
    hrm_disclaimer(doc)
    save_and_verify(doc, "HCO_HRM_1_v2_REWRITE_DRAFT.docx")


# ══════════════════════════════════════════════════════════════════════════════
# HRM.2 — Staff Recruitment   (no stop-work)
# ══════════════════════════════════════════════════════════════════════════════
def gen_hrm2():
    doc = Document()
    h(doc, 0, "Policy on Staff Recruitment")
    p(doc, HN)
    h(doc, 1, "Document control")
    doc_ctrl(doc, "HCO/HRM/POL/02", _HR)
    p(doc, "A blank marked ________ must be completed before issue.")
    h(doc, 1, "Statement of intent")
    p(doc, f"{HN} recruits staff through a transparent, documented process that checks fitness, "
           "character and conduct before someone starts patient-facing work.")
    h(doc, 1, "1. Purpose")
    p(doc, f"This policy explains how {HN} recruits staff, checks their fitness to work, sets "
           "standards of conduct, and documents its administrative procedures.")
    p(doc, "This policy does not cover induction training, job descriptions, or staff appraisal "
           "— those are covered in other hospital policies.")
    h(doc, 1, "2. Scope")
    p(doc, f"This policy applies to the HR In-Charge / Personnel Officer and every category of "
           f"staff being recruited at {HN}.")
    h(doc, 1, "3. Policy standards")
    p(doc, f"{HN} recruits against written, transparent guidance; conducts a lawful, consent-based "
           "pre-employment medical examination; defines and implements a code of conduct that protects "
           "patient confidentiality; and documents its administrative HR procedures.")
    h(doc, 1, "4. Non-negotiable rules")
    ln(doc, "Do not recruit any staff category without written, documented guidance covering that category.")
    ln(doc, "Do not perform any pre-employment test without the candidate's consent and without it being lawful.")
    ln(doc, "Do not let a staff member start work without having signed the code of conduct.")
    ln(doc, "Do not leave administrative procedures — attendance, leave, conduct, replacement — undocumented.")
    ln(doc, f"Staff who see any of these rules broken report it the same shift to the "
            f"HR In-Charge / Personnel Officer or the Medical Superintendent.")
    h(doc, 1, "5. What we do")
    h(doc, 2, "5.1 Recruit against written, transparent guidance")
    p(doc, f"Written guidance governs recruitment at {HN}, based on defined criteria for each staff "
           "category, ensuring an adequate number and skill mix to provide the hospital's services. "
           "The procedure confirms a candidate has the necessary registration, qualifications, skills "
           "and experience before appointment, and follows statutory requirements where they apply. "
           "The process is documented and carried out transparently. The HR In-Charge / Personnel "
           "Officer maintains a recruitment register logging the vacancy, candidates considered, "
           "selection rationale and fill date.")
    h(doc, 2, "5.2 Confirm fitness to work, lawfully")
    p(doc, "A pre-employment medical examination is conducted to confirm a candidate is fit to provide "
           "safe care. The scope of testing is guided by the nature of the role, but any test performed "
           "follows the law of the land — for example, pre-employment HIV testing without the candidate's "
           "consent is illegal and is not this hospital's practice. The HR In-Charge / Personnel Officer "
           "holds the examination record with the personnel file.")
    h(doc, 2, "5.3 Set and sign a code of conduct")
    p(doc, f"{HN} defines and implements a code of conduct outlining the do's and don'ts of workplace "
           "behaviour, aligned with the hospital's values and ethics framework and including protection "
           "of patient confidentiality. Staff sign the code at the time of joining; it may form part "
           "of the hospital's service rules. The HR In-Charge / Personnel Officer holds signed "
           "acknowledgements with the personnel file.")
    h(doc, 2, "5.4 Document HR administrative procedures")
    p(doc, "Administrative procedures for human resource management are documented — at minimum "
           "attendance, leave, conduct and replacement. The HR In-Charge / Personnel Officer maintains "
           "the current procedure set and covers it during induction.")
    h(doc, 1, "6. Governance and responsibility")
    gov_tbl(doc, [
        (_HR,
         "Owns the recruitment register, medical examination records, code-of-conduct "
         "acknowledgements, and administrative procedure documentation."),
        (_MS, "Accountable that this policy is followed."),
    ])
    h(doc, 1, "7. Quality monitoring")
    p(doc, "The HR In-Charge / Personnel Officer reviews a sample of recruitment files periodically "
           "against the steps in Section 5, confirming registration/qualification checks, medical "
           "examination and consent records, and signed codes of conduct are on file. This policy "
           "itself is reviewed every year.")
    h(doc, 1, "8. Training and staff acknowledgement")
    p(doc, "Staff covered by this policy are trained on it when they take up the role, and again "
           "every year after that.")
    p(doc, f"I have read the Policy on Staff Recruitment of {HN}. I will follow the processes described.")
    sig_tbl(doc)
    h(doc, 1, "9. Distribution")
    p(doc, "HR In-Charge / Personnel Officer; Medical Superintendent; department heads.")
    h(doc, 1, "10. Abbreviations")
    abbrev_tbl(doc, HRM_ABBREVS_BASE)
    h(doc, 1, "11. Traceability to NABH HCO Full Accreditation 6th Edition HRM.2")
    p(doc, "This table is an index. It is not how the policy is organised. An asterisk in the "
           "Level column means documentation of the process is required.")
    trace_tbl(doc, [
        ("HRM.2.a", "CORE*",   "Section 3; 5.1", _HR),
        ("HRM.2.b", "Commitment", "Section 3; 5.2", _HR),
        ("HRM.2.c", "CORE",    "Section 3; 5.3", _HR),
        ("HRM.2.d", "Commitment*", "Section 3; 5.4", _HR),
    ])
    h(doc, 1, "12. Required Records / Evidence Checklist")
    h(doc, 2, "HRM.2.a — Written guidance governs recruitment.")
    lb(doc, "Written recruitment guidance document.")
    lb(doc, "Recruitment register — vacancy, candidates, selection rationale, fill date.")
    lb(doc, "Statutory-requirement compliance record where applicable.")
    h(doc, 2, "HRM.2.b — Pre-employment medical examination.")
    lb(doc, "Pre-employment medical examination record on file.")
    lb(doc, "Consent record for any testing performed, confirming no non-consensual testing.")
    lb(doc, "Fitness-to-work determination record.")
    h(doc, 2, "HRM.2.c — Code of conduct defined and implemented.")
    lb(doc, "Written code of conduct document.")
    lb(doc, "Signed staff acknowledgement at joining.")
    lb(doc, "Confidentiality-protection clause record within the code.")
    h(doc, 2, "HRM.2.d — Administrative HR procedures documented.")
    lb(doc, "Documented administrative procedures — attendance, leave, conduct, replacement.")
    lb(doc, "Current-version record held by HR.")
    lb(doc, "Induction coverage record for these procedures.")
    h(doc, 1, "13. References")
    lb(doc, "NABH Accreditation Standards for Hospitals, 6th Edition — standard HRM.2.")
    lb(doc, "NABH Guidebook to Accreditation Standards for Hospitals, 6th Edition — HRM.2 interpretations.")
    h(doc, 1, "Disclaimer")
    hrm_disclaimer(doc)
    save_and_verify(doc, "HCO_HRM_2_v2_REWRITE_DRAFT.docx")


# ══════════════════════════════════════════════════════════════════════════════
# HRM.3 — Staff Induction Training   (no stop-work)
# ══════════════════════════════════════════════════════════════════════════════
def gen_hrm3():
    doc = Document()
    h(doc, 0, "Policy on Staff Induction Training")
    p(doc, HN)
    h(doc, 1, "Document control")
    doc_ctrl(doc, "HCO/HRM/POL/03", _HR)
    p(doc, "A blank marked ________ must be completed before issue.")
    h(doc, 1, "Statement of intent")
    p(doc, f"{HN} orients every new person — staff, volunteer, student or trainee — to the hospital "
           "and to their specific role before they are left to work unsupervised.")
    h(doc, 1, "1. Purpose")
    p(doc, f"This policy explains what induction training covers at {HN} and when it is completed.")
    p(doc, "This policy does not cover ongoing professional training, job-specific clinical training, "
           "or performance appraisal — those are covered in other hospital policies.")
    h(doc, 1, "2. Scope")
    p(doc, f"This policy applies to the HR In-Charge / Personnel Officer, department heads, and "
           f"every staff member, consultant, outsourced worker, volunteer, student and trainee "
           f"joining {HN}.")
    h(doc, 1, "3. Policy standards")
    p(doc, f"{HN} provides induction training within one month of joining, covering the hospital's "
           "vision and values, staff and patient rights, safety, CPR, infection control, service "
           "standards, administrative procedures, department-level policies, and information-systems use.")
    h(doc, 1, "4. Non-negotiable rules")
    ln(doc, "Do not let a new staff member, consultant, outsourced worker, volunteer, student or "
            "trainee go without induction training within one month of joining.")
    ln(doc, "Do not skip any of the required induction topics: vision/mission/values, staff and patient "
            "rights, safety, CPR, infection control, service standards, administrative procedures, "
            "department-level orientation, and information-systems training.")
    ln(doc, "Do not let department-level orientation be skipped in favour of organisation-level "
            "induction alone.")
    ln(doc, f"Staff who see any of these rules broken report it the same shift to the "
            f"HR In-Charge / Personnel Officer or the Medical Superintendent.")
    h(doc, 1, "5. What we do")
    h(doc, 2, "5.1 Provide induction within one month")
    p(doc, f"{HN} provides induction training to its staff, including doctors, consultants "
           "(including visiting), outsourced staff, volunteers, students and trainees, oriented to "
           "the hospital and to their specific assignment. Induction is completed within one month "
           "of joining and covers every topic in 5.2 through 5.9. Contents may be issued as a "
           "booklet; separate induction may run at organisational and departmental level. The "
           "HR In-Charge / Personnel Officer keeps the training record.")
    h(doc, 2, "5.2 Orient to vision, mission and values")
    p(doc, "Induction includes orientation to the hospital's vision, mission and values, so that "
           "staff — including outsourced staff — are aware of and can correctly interpret them.")
    h(doc, 2, "5.3 Cover staff and patient rights")
    p(doc, "Induction includes awareness of staff rights and responsibilities and of patient rights "
           "and responsibilities, so that staff can comprehend the implications of both and can "
           "identify and report a violation of patient rights when it occurs.")
    h(doc, 2, "5.4 Train on safety")
    p(doc, "Induction includes training on safety — patient, visitor and staff safety, including "
           "the hospital's emergency codes.")
    h(doc, 2, "5.5 Train on CPR")
    p(doc, "Induction includes training on cardio-pulmonary resuscitation. At minimum, doctors, "
           "nursing staff, technologists and rehabilitation staff are trained to at least basic life "
           "support; doctors and nurses in intensive care or high-dependency units undergo appropriate "
           "advanced training (for example ACLS, PALS or NRP, or an equivalent). A staff member with "
           "a valid existing training certificate does not need to repeat it.")
    h(doc, 2, "5.6 Train on infection prevention and control")
    p(doc, "Induction includes training in hospital infection prevention and control — the policies, "
           "procedures and practices of the infection prevention and control programme.")
    h(doc, 2, "5.7 Orient to service standards")
    p(doc, "Induction includes orientation to the hospital's service standards, so that staff are "
           "trained to implement them.")
    h(doc, 2, "5.8 Orient to administrative and department-level procedures")
    p(doc, "Induction includes an orientation on administrative procedures — attendance, leave, "
           "conduct and similar matters — and awareness of organisation-wide policies. It also "
           "includes an orientation on the policies and procedures of the specific department, unit, "
           "service or programme the staff member will work in, delivered at that level.")
    h(doc, 2, "5.9 Train on information systems")
    p(doc, "Staff are trained on information systems, information security, information use and "
           "management, according to their job responsibility, job description and data and "
           "information needs. Where the hospital uses electronic health records, staff who access, "
           "review or document in the EMR are trained to ensure it is used correctly.")
    h(doc, 1, "6. Governance and responsibility")
    gov_tbl(doc, [
        (_HR,
         "Owns the induction schedule, content and training records at organisational level."),
        ("Department heads",
         "Deliver department/unit/service-level induction content."),
        (_MS, "Accountable that this policy is followed."),
    ])
    h(doc, 1, "7. Quality monitoring")
    p(doc, "The HR In-Charge / Personnel Officer reviews a sample of induction records periodically "
           "against the steps in Section 5, confirming every required topic was covered within one "
           "month of joining. This policy itself is reviewed every year.")
    h(doc, 1, "8. Training and staff acknowledgement")
    p(doc, "Staff covered by this policy are trained on it when they take up the role, and again "
           "every year after that.")
    p(doc, f"I have read the Policy on Staff Induction Training of {HN}. I will follow the processes described.")
    sig_tbl(doc)
    h(doc, 1, "9. Distribution")
    p(doc, "HR In-Charge / Personnel Officer; department heads; Medical Superintendent.")
    h(doc, 1, "10. Abbreviations")
    abbrev_tbl(doc, [
        ("ACLS", "Advanced Cardiac Life Support"),
        ("CPR",  "Cardio-Pulmonary Resuscitation"),
        ("EMR",  "Electronic Medical Record"),
    ] + HRM_ABBREVS_BASE + [
        ("NRP",  "Neonatal Resuscitation Program"),
        ("PALS", "Paediatric Advanced Life Support"),
    ])
    h(doc, 1, "11. Traceability to NABH HCO Full Accreditation 6th Edition HRM.3")
    p(doc, "This table is an index. It is not how the policy is organised. An asterisk in the "
           "Level column means documentation of the process is required.")
    trace_tbl(doc, [
        ("HRM.3.a", "CORE",       "Section 3; 5.1", _HR),
        ("HRM.3.b", "Commitment", "Section 3; 5.2", _HR),
        ("HRM.3.c", "Commitment", "Section 3; 5.3", _HR),
        ("HRM.3.d", "Commitment", "Section 3; 5.4", _HR),
        ("HRM.3.e", "Commitment", "Section 3; 5.5", _HR),
        ("HRM.3.f", "Commitment", "Section 3; 5.6", _HR),
        ("HRM.3.g", "Commitment", "Section 3; 5.7", _HR),
        ("HRM.3.h", "Commitment", "Section 3; 5.8", _HR),
        ("HRM.3.i", "Commitment", "Section 3; 5.8", "Department heads"),
        ("HRM.3.j", "Commitment", "Section 3; 5.9", _HR),
    ])
    h(doc, 1, "12. Required Records / Evidence Checklist")
    h(doc, 2, "HRM.3.a — Induction training provided.")
    lb(doc, "Induction-training record within one month of joining.")
    lb(doc, "Attendance record covering doctors, consultants, outsourced staff, volunteers, students and trainees.")
    lb(doc, "Induction-content record covering HRM.3.b–j.")
    h(doc, 2, "HRM.3.b — Vision, mission and values orientation.")
    lb(doc, "Vision, mission and values orientation record within induction.")
    lb(doc, "Staff-awareness confirmation record.")
    h(doc, 2, "HRM.3.c — Staff and patient rights and responsibilities.")
    lb(doc, "Staff-rights and patient-rights training record.")
    lb(doc, "Violation-identification-and-reporting awareness confirmation.")
    h(doc, 2, "HRM.3.d — Safety training.")
    lb(doc, "Safety training record — patient, visitor, staff safety, emergency codes.")
    lb(doc, "Emergency-code awareness confirmation.")
    h(doc, 2, "HRM.3.e — CPR training.")
    lb(doc, "CPR training record at the appropriate level (BLS minimum; advanced for ICU/HDU staff).")
    lb(doc, "Valid-certificate exemption record, where applicable.")
    h(doc, 2, "HRM.3.f — Infection prevention and control training.")
    lb(doc, "Hospital infection prevention and control training record within induction.")
    h(doc, 2, "HRM.3.g — Service standards orientation.")
    lb(doc, "Service-standards orientation record within induction.")
    h(doc, 2, "HRM.3.h — Administrative procedures orientation.")
    lb(doc, "Administrative-procedures orientation record — attendance, leave, conduct.")
    lb(doc, "Organisation-wide policy-awareness record.")
    h(doc, 2, "HRM.3.i — Department/unit-level orientation.")
    lb(doc, "Department, unit, service or programme-level policy-and-procedure orientation record.")
    lb(doc, "Delivery-location record confirming it was given at that level.")
    h(doc, 2, "HRM.3.j — Information systems training.")
    lb(doc, "Information-systems, security and data-use training record.")
    lb(doc, "EMR-access training record, where the hospital uses electronic health records.")
    h(doc, 1, "13. References")
    lb(doc, "NABH Accreditation Standards for Hospitals, 6th Edition — standard HRM.3.")
    lb(doc, "NABH Guidebook to Accreditation Standards for Hospitals, 6th Edition — HRM.3 interpretations.")
    h(doc, 1, "Disclaimer")
    hrm_disclaimer(doc)
    save_and_verify(doc, "HCO_HRM_3_v2_REWRITE_DRAFT.docx")


# ══════════════════════════════════════════════════════════════════════════════
# HRM.4 — Professional Training and Development   (no stop-work)
# ══════════════════════════════════════════════════════════════════════════════
def gen_hrm4():
    doc = Document()
    h(doc, 0, "Policy on Professional Training and Development")
    p(doc, HN)
    h(doc, 1, "Document control")
    doc_ctrl(doc, "HCO/HRM/POL/04", _HR)
    p(doc, "A blank marked ________ must be completed before issue.")
    h(doc, 1, "Statement of intent")
    p(doc, f"{HN} keeps its staff's skills current through an ongoing, documented training programme, "
           "and checks that training actually changes what happens at the workplace.")
    h(doc, 1, "1. Purpose")
    p(doc, f"This policy explains how {HN} plans, records, and evaluates staff professional training "
           "and development, and how it supports continuing learning.")
    p(doc, "This policy does not cover induction training or job-specific clinical training — those "
           "are covered in other hospital policies.")
    h(doc, 1, "2. Scope")
    p(doc, f"This policy applies to the HR In-Charge / Personnel Officer, department heads, and "
           f"every category of staff, including doctors and outsourced staff where applicable.")
    h(doc, 1, "3. Policy standards")
    p(doc, f"{HN} runs training and development against written guidance; keeps a training record "
           "for every session; retrains staff when responsibilities or equipment change; collects "
           "feedback to improve the programme; evaluates whether training actually worked; and "
           "supports staff in continuing professional development.")
    h(doc, 1, "4. Non-negotiable rules")
    ln(doc, "Do not run training and development without written guidance covering training-needs "
            "identification, methodology, documentation, assessment and a training calendar.")
    ln(doc, "Do not leave a training session unrecorded — title, trainer, date, duration and "
            "trainee list with signatures are mandatory.")
    ln(doc, "Do not skip training when a staff member's job responsibilities change or new "
            "equipment is introduced.")
    ln(doc, "Do not run the training programme without a feedback mechanism covering course "
            "material, facilities and trainer capability.")
    ln(doc, "Do not treat training as complete without evaluating whether it worked, both "
            "immediately and after time has passed.")
    ln(doc, f"Staff who see any of these rules broken report it the same shift to the "
            f"HR In-Charge / Personnel Officer or the Medical Superintendent.")
    h(doc, 1, "5. What we do")
    h(doc, 2, "5.1 Run training against written guidance")
    p(doc, f"Written guidance governs training and development at {HN}: a training manual covering "
           "identification of training needs, methodology, documentation, assessment, impact "
           "evaluation and a training calendar. At minimum, staff are trained on occupational safety "
           "and soft skills, and educated on patient-centred care — respecting patient preferences, "
           "shared decision-making and integrated care. Training covers all staff categories, "
           "including doctors and outsourced staff where applicable.")
    h(doc, 2, "5.2 Keep the training record")
    p(doc, "The HR In-Charge / Personnel Officer maintains the training record for every session: "
           "at minimum the title, trainer(s), date, duration and list of trainees with signatures. "
           "Contents are captured where possible; records may be kept digitally.")
    h(doc, 2, "5.3 Retrain when the job or the equipment changes")
    p(doc, "Training also occurs when job responsibilities change or new equipment is introduced, "
           "focused on the revised responsibilities or the newly introduced equipment and technology. "
           "For new equipment, operating staff are trained on both operational use and daily "
           "maintenance before working with it independently.")
    h(doc, 2, "5.4 Collect feedback to improve the programme")
    p(doc, "Feedback mechanisms are in place for improving the training and development programme, "
           "covering both internal and external training — appropriateness of course material, "
           "training facilities and trainer capability.")
    h(doc, 2, "5.5 Evaluate whether training worked")
    p(doc, f"{HN} evaluates training effectiveness immediately after training (for example a pre- "
           "and post-test) and again after a defined period has lapsed, to confirm the training "
           "improved workplace competency. Incident reports and assessment non-conformities are "
           "useful inputs for the later check. The evaluation covers knowledge, skills and attitude; "
           "retraining is provided where the evaluation shows it is needed.")
    h(doc, 2, "5.6 Support continuing professional development")
    p(doc, f"{HN} supports continuing professional development and learning, so staff can keep up "
           "with advancements in their field — encouraging and resourcing attendance at courses or "
           "conferences, and providing access to distance learning or e-learning. The hospital "
           "specifies minimum mandatory training hours every staff member attends each year.")
    h(doc, 1, "6. Governance and responsibility")
    gov_tbl(doc, [
        (_HR,
         "Owns the training manual, training records, feedback mechanism, effectiveness "
         "evaluations, and CPD support arrangements."),
        ("Department heads",
         "Identify training needs arising from job or equipment changes in their department."),
        (_MS,
         "Accountable that this policy is followed; sets minimum mandatory annual training hours."),
    ])
    h(doc, 1, "7. Quality monitoring")
    p(doc, "The HR In-Charge / Personnel Officer reviews a sample of training records periodically "
           "against the steps in Section 5, confirming sessions are documented, feedback is "
           "collected, and effectiveness evaluations and any resulting retraining are on file. "
           "This policy itself is reviewed every year.")
    h(doc, 1, "8. Training and staff acknowledgement")
    p(doc, "Staff covered by this policy are trained on it when they take up the role, and again "
           "every year after that.")
    p(doc, f"I have read the Policy on Professional Training and Development of {HN}. I will follow the processes described.")
    sig_tbl(doc)
    h(doc, 1, "9. Distribution")
    p(doc, "HR In-Charge / Personnel Officer; department heads; Medical Superintendent.")
    h(doc, 1, "10. Abbreviations")
    abbrev_tbl(doc, [
        ("CPD", "Continuing Professional Development"),
    ] + HRM_ABBREVS_BASE)
    h(doc, 1, "11. Traceability to NABH HCO Full Accreditation 6th Edition HRM.4")
    p(doc, "This table is an index. It is not how the policy is organised. An asterisk in the "
           "Level column means documentation of the process is required.")
    trace_tbl(doc, [
        ("HRM.4.a", "CORE*",       "Section 3; 5.1", _HR),
        ("HRM.4.b", "Commitment",  "Section 3; 5.2", _HR),
        ("HRM.4.c", "Commitment",  "Section 3; 5.3", _HR),
        ("HRM.4.d", "Commitment",  "Section 3; 5.4", _HR),
        ("HRM.4.e", "Achievement", "Section 3; 5.5", _HR),
        ("HRM.4.f", "Achievement", "Section 3; 5.6", _HR),
    ])
    h(doc, 1, "12. Required Records / Evidence Checklist")
    h(doc, 2, "HRM.4.a — Written guidance for training and development.")
    lb(doc, "Written training and development policy or manual.")
    lb(doc, "Training-needs-identification, methodology, assessment and calendar record.")
    lb(doc, "Coverage record for all staff categories, including doctors and outsourced staff.")
    h(doc, 2, "HRM.4.b — Training record maintained.")
    lb(doc, "Training record with title, trainer, date, duration and trainee list with signatures.")
    lb(doc, "Content-capture record where possible.")
    h(doc, 2, "HRM.4.c — Training on job/equipment change.")
    lb(doc, "Training record triggered by a job-responsibility change or new-equipment introduction.")
    lb(doc, "Operational-and-maintenance training record for new equipment.")
    lb(doc, "Training-completion confirmation before independent use.")
    h(doc, 2, "HRM.4.d — Feedback mechanism.")
    lb(doc, "Feedback-mechanism record for training-programme improvement.")
    lb(doc, "Feedback data on course material, facilities and trainer capability.")
    h(doc, 2, "HRM.4.e — Evaluation of training effectiveness.")
    lb(doc, "Immediate post-training evaluation record — pre/post-test.")
    lb(doc, "Later workplace-effectiveness evaluation record.")
    lb(doc, "Retraining record where the evaluation showed a need.")
    h(doc, 2, "HRM.4.f — Continuing professional development.")
    lb(doc, "CPD support record — courses, conferences, e-learning access.")
    lb(doc, "Minimum mandatory annual training-hours specification record.")
    lb(doc, "Staff-participation record against the mandatory hours.")
    h(doc, 1, "13. References")
    lb(doc, "NABH Accreditation Standards for Hospitals, 6th Edition — standard HRM.4.")
    lb(doc, "NABH Guidebook to Accreditation Standards for Hospitals, 6th Edition — HRM.4 interpretations.")
    h(doc, 1, "Disclaimer")
    hrm_disclaimer(doc)
    save_and_verify(doc, "HCO_HRM_4_v2_REWRITE_DRAFT.docx")


# ══════════════════════════════════════════════════════════════════════════════
# HRM.5 — Job-Specific Staff Training   (no stop-work)
# ══════════════════════════════════════════════════════════════════════════════
def gen_hrm5():
    doc = Document()
    h(doc, 0, "Policy on Job-Specific Staff Training")
    p(doc, HN)
    h(doc, 1, "Document control")
    doc_ctrl(doc, "HCO/HRM/POL/05", _HR)
    p(doc, "A blank marked ________ must be completed before issue.")
    h(doc, 1, "Statement of intent")
    p(doc, f"{HN} trains staff on the specific skills their job requires — blood handling, "
           "vulnerable-patient care, restraint, communication, resuscitation and infection control "
           "— not just general induction.")
    h(doc, 1, "1. Purpose")
    p(doc, f"This policy explains the job-specific training {HN} provides beyond induction, based "
           "on what a staff member's role actually requires.")
    p(doc, "This policy does not cover the clinical procedures themselves for blood transfusion, "
           "vulnerable-patient care, control and restraint, or patient communication — those are "
           "covered in other hospital policies. This policy is the staff training layer for those procedures.")
    h(doc, 1, "2. Scope")
    p(doc, f"This policy applies to the HR In-Charge / Personnel Officer, department heads, and "
           "every staff member whose role requires blood-handling, vulnerable-patient care, restraint, "
           "communication, CPR, or infection-control skills.")
    h(doc, 1, "3. Policy standards")
    p(doc, f"{HN} trains relevant staff in blood and blood product handling, vulnerable-patient care, "
           "control and restraint techniques, healthcare communication, periodic CPR for "
           "direct-patient-care staff, and infection prevention and control.")
    h(doc, 1, "4. Non-negotiable rules")
    ln(doc, "Do not let staff handling blood or blood products work without training in safe "
            "transport, consent, documentation and transfusion-reaction handling.")
    ln(doc, "Do not let staff caring for vulnerable patients work without training in identifying "
            "and caring for them.")
    ln(doc, "Do not let staff use control and restraint techniques without training in their "
            "appropriate use.")
    ln(doc, "Do not let CPR training lapse beyond two years for direct-patient-care staff, or "
            "beyond the interval required after a protocol change.")
    ln(doc, "Do not let infection-prevention-and-control training lapse beyond a year for any "
            "staff member.")
    ln(doc, f"Staff who see any of these rules broken report it the same shift to the "
            f"HR In-Charge / Personnel Officer or the Medical Superintendent.")
    h(doc, 1, "5. What we do")
    h(doc, 2, "5.1 Train staff on blood and blood product handling")
    p(doc, "Staff involved in blood transfusion services — doctors, nurses, technicians and staff "
           "transporting blood from the blood bank or storage unit — are trained in handling blood "
           "and blood products: safe transport, obtaining informed consent, required documentation, "
           "identifying and handling transfusion reactions, and educating the patient and family on "
           "donation. Blood transfusion service practice itself is covered in other hospital "
           "policies; this element is the staff training layer.")
    h(doc, 2, "5.2 Train staff on vulnerable-patient care")
    p(doc, "Relevant staff are trained in identifying and rendering care to vulnerable patients, "
           "per the hospital's written guidance. The vulnerable-patient care process itself is "
           "covered in other hospital policies; this element is the staff training layer.")
    h(doc, 2, "5.3 Train staff on control and restraint")
    p(doc, "Relevant staff are trained in the appropriate use of control and restraint techniques, "
           "per the hospital's written guidance. The control-and-restraint process itself is covered "
           "in other hospital policies; this element is the staff training layer.")
    h(doc, 2, "5.4 Train staff on healthcare communication")
    p(doc, "Staff are trained in healthcare communication techniques, including handling challenging "
           "situations and good communication practice. Training needs may be identified from patient "
           "complaints, incident reports, appraisals and employee feedback. Patient-facing "
           "communication practice itself is covered in other hospital policies; this element is "
           "the staff training layer.")
    h(doc, 2, "5.5 Train staff on CPR periodically")
    p(doc, "Staff involved in direct patient care are trained on cardio-pulmonary resuscitation "
           "periodically, at the level appropriate to their role. Doctors, nurses and rehabilitation "
           "staff refresh at least once in two years, or sooner if protocol changes; staff in "
           "emergency, intensive care or high-dependency units undergo appropriate advanced training "
           "(for example ACLS, ATLS, PALS or NRP, or an equivalent).")
    h(doc, 2, "5.6 Train staff on infection prevention and control")
    p(doc, f"{HN} provides staff training on infection prevention and control through in-service "
           "sessions at least once a year, including antimicrobial policy and antimicrobial "
           "stewardship content for medical professionals, infection-prevention-and-control nurses, "
           "the clinical pharmacist and support staff.")
    h(doc, 1, "6. Governance and responsibility")
    gov_tbl(doc, [
        (_HR,
         "Owns training records for all job-specific training under this policy."),
        ("Department heads",
         "Identify which staff in their department are \"relevant staff\" for each training requirement."),
        (_MS, "Accountable that this policy is followed."),
    ])
    h(doc, 1, "7. Quality monitoring")
    p(doc, "The HR In-Charge / Personnel Officer reviews a sample of training records periodically "
           "against the steps in Section 5, confirming CPR refreshers are current and "
           "infection-control training has not lapsed beyond a year. This policy itself is "
           "reviewed every year.")
    h(doc, 1, "8. Training and staff acknowledgement")
    p(doc, "Staff covered by this policy are trained on it when they take up the role, and again "
           "every year after that.")
    p(doc, f"I have read the Policy on Job-Specific Staff Training of {HN}. I will follow the processes described.")
    sig_tbl(doc)
    h(doc, 1, "9. Distribution")
    p(doc, "HR In-Charge / Personnel Officer; department heads; Medical Superintendent.")
    h(doc, 1, "10. Abbreviations")
    abbrev_tbl(doc, [
        ("ACLS", "Advanced Cardiac Life Support"),
        ("ATLS", "Advanced Trauma Life Support"),
        ("CPR",  "Cardio-Pulmonary Resuscitation"),
    ] + HRM_ABBREVS_BASE + [
        ("NRP",  "Neonatal Resuscitation Program"),
        ("PALS", "Paediatric Advanced Life Support"),
    ])
    h(doc, 1, "11. Traceability to NABH HCO Full Accreditation 6th Edition HRM.5")
    p(doc, "This table is an index. It is not how the policy is organised. An asterisk in the "
           "Level column means documentation of the process is required.")
    trace_tbl(doc, [
        ("HRM.5.a", "Commitment", "Section 3; 5.1", _HR),
        ("HRM.5.b", "Commitment", "Section 3; 5.2", _HR),
        ("HRM.5.c", "Commitment", "Section 3; 5.3", _HR),
        ("HRM.5.d", "Commitment", "Section 3; 5.4", _HR),
        ("HRM.5.e", "CORE",       "Section 3; 5.5", _HR),
        ("HRM.5.f", "Commitment", "Section 3; 5.6", _HR),
    ])
    h(doc, 1, "12. Required Records / Evidence Checklist")
    h(doc, 2, "HRM.5.a — Blood and blood product handling training.")
    lb(doc, "Blood-and-blood-product handling training record for relevant staff (doctors, nurses, technicians, transport staff).")
    lb(doc, "Training-content record — safe transport, informed consent, documentation, transfusion-reaction handling.")
    h(doc, 2, "HRM.5.b — Vulnerable-patient training.")
    lb(doc, "Vulnerable-patient identification-and-care training record.")
    lb(doc, "Relevant-staff coverage record.")
    h(doc, 2, "HRM.5.c — Control and restraint training.")
    lb(doc, "Control-and-restraint-technique training record.")
    lb(doc, "Relevant-staff coverage record.")
    h(doc, 2, "HRM.5.d — Healthcare communication training.")
    lb(doc, "Healthcare-communication-technique training record.")
    lb(doc, "Training-needs source record — complaints, incident reports, appraisals, feedback.")
    h(doc, 2, "HRM.5.e — Periodic CPR training.")
    lb(doc, "Periodic CPR-training record for direct-patient-care staff, at least once in two years or sooner after protocol change.")
    lb(doc, "Advanced-training record for emergency, ICU or high-dependency staff.")
    lb(doc, "Refresher-schedule tracking record.")
    h(doc, 2, "HRM.5.f — Infection prevention and control training.")
    lb(doc, "Infection-prevention-and-control training record at least annually.")
    lb(doc, "Antimicrobial-stewardship-content record for medical professionals, IPC nurses, clinical pharmacist and support staff.")
    h(doc, 1, "13. References")
    lb(doc, "NABH Accreditation Standards for Hospitals, 6th Edition — standard HRM.5.")
    lb(doc, "NABH Guidebook to Accreditation Standards for Hospitals, 6th Edition — HRM.5 interpretations.")
    h(doc, 1, "Disclaimer")
    hrm_disclaimer(doc)
    save_and_verify(doc, "HCO_HRM_5_v2_REWRITE_DRAFT.docx")


# ══════════════════════════════════════════════════════════════════════════════
# HRM.6 — Safety and Quality-Related Staff Training   (no stop-work)
# ══════════════════════════════════════════════════════════════════════════════
def gen_hrm6():
    doc = Document()
    h(doc, 0, "Policy on Safety and Quality-Related Staff Training")
    p(doc, HN)
    h(doc, 1, "Document control")
    doc_ctrl(doc, "HCO/HRM/POL/06", _HR)
    p(doc, "A blank marked ________ must be completed before issue.")
    h(doc, 1, "Statement of intent")
    p(doc, f"{HN} trains staff in safety and quality-related aspects of their work so that they can "
           "protect patients, themselves, and each other, respond correctly to incidents and "
           "emergencies, and contribute to the hospital's quality improvement programme.")
    h(doc, 1, "1. Purpose")
    p(doc, f"This policy sets out how {HN} trains staff in the organisation's safety programme, in "
           "detecting and handling risks in the work environment, in the procedure to follow during "
           "an incident, in occupational safety, in the disaster management plan, in handling fire "
           "and non-fire emergencies, and in the organisation's quality improvement programme.")
    p(doc, "This policy does not cover recruitment, induction, job-specific clinical training, or "
           "performance appraisal — those are covered in other hospital policies. The other HRM "
           "standards have their own policies too.")
    h(doc, 1, "2. Scope")
    p(doc, f"This policy applies to all staff of {HN}, across departments and categories. Staff "
           "working in laboratory, imaging, emergency, intensive care, the blood centre and surgical "
           "services receive additional department-specific safety and quality assurance training as "
           "described in Section 5.")
    h(doc, 1, "3. Policy standards")
    p(doc, f"{HN} trains staff in the organisation's safety programme, including patient safety, "
           "with additional safety training for staff working in laboratory and imaging services.")
    p(doc, f"{HN} trains staff to detect, handle, minimise and eliminate identified risks — "
           "physical, chemical, environmental and process-related — in the work environment, "
           "including practical demonstration of key actions.")
    p(doc, f"{HN} makes staff aware of the procedure to follow when an incident occurs.")
    p(doc, f"{HN} trains staff in occupational safety aspects for areas with identified "
           "occupational hazards.")
    p(doc, f"{HN} trains staff in the disaster management plan and their specific role in managing "
           "an internal or external disaster.")
    p(doc, f"{HN} trains staff in handling fire and non-fire emergencies, including their specific "
           "role in each.")
    p(doc, f"{HN} trains staff in the organisation's quality improvement programme, including their "
           "own role in it, with additional department-specific quality assurance training where "
           "applicable.")
    h(doc, 1, "4. Non-negotiable rules")
    ln(doc, "Do not put a staff member to work without training them in the organisation's safety "
            "programme, including patient safety.")
    ln(doc, "Do not send staff into the work environment without training them to detect, handle, "
            "minimise and eliminate the identified physical, chemical, environmental and "
            "process-related risks in that environment.")
    ln(doc, "Do not leave staff unaware of the procedure to follow when an incident occurs.")
    ln(doc, "Do not skip occupational safety training for staff working in areas with identified "
            "occupational hazards.")
    ln(doc, "Do not leave a staff member untrained in the disaster management plan and their "
            "specific role in it.")
    ln(doc, "Do not leave a staff member untrained in handling fire and non-fire emergencies.")
    ln(doc, "Do not leave staff unaware of the quality improvement programme's structure and their "
            "own role in it.")
    ln(doc, f"Staff who see a rule under this policy broken report it the same shift to the "
            f"HR In-Charge / Personnel Officer or the Medical Superintendent.")
    h(doc, 1, "5. What we do")
    h(doc, 2, "5.1 Safety programme training")
    p(doc, f"Staff at {HN} are trained in the organisation's safety programme, including patient "
           "safety, through a regular training programme or printed materials. Staff working in "
           "laboratory and imaging services are additionally trained in their respective safety "
           "programmes. The safety programme itself is covered in other hospital policies; this "
           "policy covers the staff training layer.")
    h(doc, 2, "5.2 Risk detection and handling")
    p(doc, f"Staff at {HN} are trained in detecting, handling, minimising and eliminating identified "
           "risks in the organisation's environment — physical (poor lighting, slippery floors, blind "
           "corners, open electrical points, exposed wiring), chemical (mishandling, spills, "
           "aerosolisation), environmental (noise, smoke, dampness, heat) and process-related "
           "(needle-stick injury, blood and body-fluid exposure, cytotoxic drugs, soiled linen). "
           "Staff can practically demonstrate actions such as managing a blood spill or handling "
           "hazardous materials.")
    h(doc, 2, "5.3 Incident procedure awareness")
    p(doc, f"Staff at {HN} are made aware of the procedure to follow in the event of an incident, "
           "and are able to describe the sequence of events they will undertake if one occurs.")
    h(doc, 2, "5.4 Occupational safety training")
    p(doc, f"Staff at {HN} are trained in occupational safety aspects for the areas with identified "
           "occupational hazards — for example needle-stick injury and blood/body-fluid exposure, "
           "radiation exposure, laser exposure, medical-gas exposure, chemotherapy exposure and "
           "noise in utility areas — and in the preventive actions to avoid each risk. "
           "Occupational-exposure practice itself is covered in other hospital policies; this "
           "policy covers the staff training layer.")
    h(doc, 2, "5.5 Disaster management plan training")
    p(doc, f"Staff at {HN} are trained in the organisation's disaster management plan, including "
           "their specific role in managing an internal or external disaster.")
    h(doc, 2, "5.6 Fire and non-fire emergency training")
    p(doc, f"Staff at {HN} are trained in handling fire and non-fire emergencies: classes of fire, "
           "use of fire extinguishers, evacuation plans and fire procedures, plus the hospital's "
           "identified non-fire emergencies and each staff member's specific role in them.")
    h(doc, 2, "5.7 Quality improvement programme training")
    p(doc, f"Staff at {HN} are trained in the organisation's quality improvement programme — its "
           "structure and their own role in contributing to it. Staff working in laboratory, "
           "imaging, emergency, intensive care, the blood centre and surgical services are "
           "additionally trained on their respective quality assurance programmes.")
    h(doc, 1, "6. Governance and responsibility")
    gov_tbl(doc, [
        (_HR,
         "Owns this policy. Plans and organises safety and quality-related training. Maintains "
         "training records."),
        ("Department heads / Unit In-Charges",
         "Identify department-specific hazards and quality assurance training needs. Ensure staff "
         "under them attend required training."),
        (_MS,
         "Oversees implementation of this policy. Receives escalations under Section 4."),
        ("All Staff",
         "Attend and complete required training. Report identified risks and rule breaches."),
    ])
    h(doc, 1, "7. Quality monitoring")
    p(doc, f"The HR In-Charge / Personnel Officer reviews training records to confirm staff have "
           "completed the training required under this policy. Documentary evidence is on file for "
           "each CORE objective element. This policy itself is reviewed every year.")
    h(doc, 1, "8. Training and staff acknowledgement")
    p(doc, "All staff covered by this policy complete the training described in Section 5 and sign "
           "the acknowledgement below.")
    p(doc, f"I have read the Policy on Safety and Quality-Related Staff Training of {HN}. I will follow the processes described.")
    sig_tbl(doc)
    h(doc, 1, "9. Distribution")
    p(doc, f"HR department; all department heads / unit in-charges; all staff through the hospital "
           "intranet and department policy folders.")
    h(doc, 1, "10. Abbreviations")
    abbrev_tbl(doc, HRM_ABBREVS_BASE)
    h(doc, 1, "11. Traceability to NABH HCO Full Accreditation 6th Edition HRM.6")
    p(doc, "This table is an index. It is not how the policy is organised. An asterisk in the "
           "Level column means documentation of the process is required.")
    trace_tbl(doc, [
        ("HRM.6.a", "Commitment", "Section 3; 5.1", _HR),
        ("HRM.6.b", "Commitment", "Section 3; 5.2", _HR),
        ("HRM.6.c", "Commitment", "Section 3; 5.3", _HR),
        ("HRM.6.d", "Commitment", "Section 3; 5.4", _HR),
        ("HRM.6.e", "CORE",       "Section 3; 5.5", _HR),
        ("HRM.6.f", "CORE",       "Section 3; 5.6", _HR),
        ("HRM.6.g", "Commitment", "Section 3; 5.7", _HR),
    ])
    h(doc, 1, "12. Required Records / Evidence Checklist")
    h(doc, 2, "HRM.6.a — Safety programme training.")
    lb(doc, "Safety-programme training record for all staff, including patient-safety content.")
    lb(doc, "Laboratory- or imaging-specific safety-training record, where applicable.")
    lb(doc, "Attendance record for all safety-programme training sessions.")
    h(doc, 2, "HRM.6.b — Risk detection and handling.")
    lb(doc, "Risk-detection-and-handling training record — physical, chemical, environmental and process-related risks.")
    lb(doc, "Practical-demonstration record — blood-spill management, hazardous-material handling.")
    h(doc, 2, "HRM.6.c — Incident procedure awareness.")
    lb(doc, "Incident-procedure awareness record and staff confirmation record of procedure steps.")
    h(doc, 2, "HRM.6.d — Occupational safety training.")
    lb(doc, "Occupational-safety-aspect training record — needle-stick, radiation, laser, medical-gas, chemotherapy, noise exposure.")
    lb(doc, "Preventive-action awareness record for each identified occupational hazard.")
    h(doc, 2, "HRM.6.e — Disaster management plan training.")
    lb(doc, "Disaster-management-plan training record, including specific-role training for internal and external disaster scenarios.")
    h(doc, 2, "HRM.6.f — Fire and non-fire emergency training.")
    lb(doc, "Fire-and-non-fire-emergency training record, including fire-extinguisher use and evacuation-procedure demonstration record.")
    lb(doc, "Role-specific emergency-training record for each identified non-fire emergency.")
    h(doc, 2, "HRM.6.g — Quality improvement programme training.")
    lb(doc, "Quality-improvement-programme training record, including role-in-programme awareness record.")
    lb(doc, "Department-specific quality-assurance training record, where applicable — laboratory, imaging, emergency, intensive care, blood centre, surgical services.")
    h(doc, 1, "13. References")
    lb(doc, "NABH Accreditation Standards for Hospitals, 6th Edition — standard HRM.6.")
    lb(doc, "NABH Guidebook to Accreditation Standards for Hospitals, 6th Edition — HRM.6 interpretations.")
    h(doc, 1, "Disclaimer")
    hrm_disclaimer(doc)
    save_and_verify(doc, "HCO_HRM_6_v2_REWRITE_DRAFT.docx")


# ══════════════════════════════════════════════════════════════════════════════
# HRM.7 — Staff Performance Appraisal   (no stop-work)
# ══════════════════════════════════════════════════════════════════════════════
def gen_hrm7():
    doc = Document()
    h(doc, 0, "Policy on Staff Performance Appraisal")
    p(doc, HN)
    h(doc, 1, "Document control")
    doc_ctrl(doc, "HCO/HRM/POL/07", _HR)
    p(doc, "A blank marked ________ must be completed before issue.")
    h(doc, 1, "Statement of intent")
    p(doc, f"{HN} runs a performance appraisal system as an integral part of human resource "
           "management, so that staff performance is evaluated fairly, against pre-determined "
           "criteria, and used to support staff development.")
    h(doc, 1, "1. Purpose")
    p(doc, f"This policy sets out how {HN} appraises staff performance — covering who is appraised, "
           "how staff learn about the system, the criteria used, how appraisal supports development, "
           "and how often appraisal happens.")
    p(doc, "This policy does not cover recruitment, induction, safety and quality training, or "
           "disciplinary and grievance handling — those are covered in other hospital policies. "
           "The other HRM standards have their own policies too.")
    h(doc, 1, "2. Scope")
    p(doc, f"This policy applies to all categories of staff at {HN}, starting with the person "
           "heading the organisation and including doctors. For outsourced staff, the appraisal "
           "may be conducted by the contractor.")
    h(doc, 1, "3. Policy standards")
    p(doc, f"{HN} conducts performance appraisal for all categories of staff, including competency "
           "assessment where appropriate.")
    p(doc, f"{HN} makes staff aware of the appraisal system at the time of induction.")
    p(doc, f"{HN} evaluates performance against pre-determined criteria — key performance indicators "
           "or key result areas derived from the job description.")
    p(doc, f"{HN} uses the appraisal system as a tool for staff development, including identifying "
           "training needs and managing underperformance.")
    p(doc, f"{HN} carries out performance appraisal at defined intervals, at least once a year, "
           "and documents it.")
    h(doc, 1, "4. Non-negotiable rules")
    ln(doc, "Do not skip performance appraisal for any category of staff, including the person "
            "heading the organisation and doctors.")
    ln(doc, "Do not leave staff unaware of the appraisal system at the time of induction.")
    ln(doc, "Do not evaluate performance without pre-determined criteria derived from the job "
            "description.")
    ln(doc, "Do not run the appraisal system as a paperwork exercise disconnected from staff "
            "development — training needs and underperformance must be addressed through it.")
    ln(doc, "Do not let a staff member's performance appraisal go more than a year without being "
            "carried out and documented.")
    ln(doc, f"Staff who see a rule under this policy broken report it the same shift to the "
            f"HR In-Charge / Personnel Officer or the Medical Superintendent.")
    h(doc, 1, "5. What we do")
    h(doc, 2, "5.1 Who is appraised")
    p(doc, f"Performance appraisal is done for all categories of staff at {HN}, starting with the "
           "person heading the organisation and including doctors, and includes competency assessment "
           "where appropriate. For outsourced staff, the appraisal may be done by the contractor.")
    h(doc, 2, "5.2 Awareness at induction")
    p(doc, f"Staff at {HN} are made aware of the appraisal system at the time of induction — for "
           "example through the service booklet and as part of induction training, which is covered "
           "in other hospital policies.")
    h(doc, 2, "5.3 Appraisal criteria")
    p(doc, f"Performance is evaluated against pre-determined criteria at {HN}, based on key "
           "performance indicators or key result areas derived from the job description.")
    h(doc, 2, "5.4 Appraisal as a development tool")
    p(doc, f"{HN} uses the appraisal system as a tool for further development — identifying "
           "training requirements and providing for them where possible, with key result areas set "
           "for each staff member and a training-need assessment done. Written guidance covers "
           "effective management of underperformance.")
    h(doc, 2, "5.5 Frequency and documentation")
    p(doc, f"Performance appraisal at {HN} is carried out at a defined interval and documented, "
           "at least once a year.")
    h(doc, 1, "6. Governance and responsibility")
    gov_tbl(doc, [
        (_HR,
         "Owns this policy. Coordinates the annual appraisal cycle. Maintains appraisal records."),
        ("Department heads / Unit In-Charges",
         "Conduct appraisals for staff reporting to them, against defined criteria. Identify "
         "training needs from appraisal outcomes."),
        (_MS,
         "Oversees implementation of this policy. Receives escalations under Section 4."),
        ("All Staff",
         "Participate in the appraisal process. Report identified risks and rule breaches."),
    ])
    h(doc, 1, "7. Quality monitoring")
    p(doc, f"The HR In-Charge / Personnel Officer reviews appraisal records annually to confirm "
           "every staff member has been appraised within the defined interval. Documentary evidence "
           "is on file for each asterisked objective element. This policy itself is reviewed every year.")
    h(doc, 1, "8. Training and staff acknowledgement")
    p(doc, "All staff covered by this policy are made aware of the appraisal system at induction "
           "and sign the acknowledgement below.")
    p(doc, f"I have read the Policy on Staff Performance Appraisal of {HN}. I will follow the processes described.")
    sig_tbl(doc)
    h(doc, 1, "9. Distribution")
    p(doc, "HR department; all department heads / unit in-charges; all staff through the hospital "
           "intranet and department policy folders.")
    h(doc, 1, "10. Abbreviations")
    abbrev_tbl(doc, [
        ("KPI", "Key Performance Indicator"),
        ("KRA", "Key Result Area"),
    ] + HRM_ABBREVS_BASE)
    h(doc, 1, "11. Traceability to NABH HCO Full Accreditation 6th Edition HRM.7")
    p(doc, "This table is an index. It is not how the policy is organised. An asterisk in the "
           "Level column means documentation of the process is required.")
    trace_tbl(doc, [
        ("HRM.7.a", "Commitment*", "Section 3; 5.1", _HR),
        ("HRM.7.b", "Commitment",  "Section 3; 5.2", _HR),
        ("HRM.7.c", "Commitment",  "Section 3; 5.3", _HR),
        ("HRM.7.d", "Commitment",  "Section 3; 5.4", _HR),
        ("HRM.7.e", "Commitment",  "Section 3; 5.5", _HR),
    ])
    h(doc, 1, "12. Required Records / Evidence Checklist")
    h(doc, 2, "HRM.7.a — Performance appraisal done for all staff categories.")
    lb(doc, "Performance-appraisal record for all staff categories, including the organisation head and doctors.")
    lb(doc, "Competency-assessment record where appropriate.")
    lb(doc, "Contractor-conducted appraisal record for outsourced staff.")
    h(doc, 2, "HRM.7.b — Staff made aware of appraisal system at induction.")
    lb(doc, "Appraisal-system-awareness record at induction.")
    lb(doc, "Service-booklet or induction-material reference and staff-acknowledgement record.")
    h(doc, 2, "HRM.7.c — Performance evaluated on pre-determined criteria.")
    lb(doc, "Pre-determined-criteria document — key performance indicators or key result areas derived from the job description.")
    lb(doc, "Evaluation record against those criteria, with job-description cross-reference.")
    h(doc, 2, "HRM.7.d — Appraisal system used as a development tool.")
    lb(doc, "Development-action record from the appraisal — training requirement identified, key result areas and training-need assessment.")
    lb(doc, "Underperformance-management written-guidance record.")
    h(doc, 2, "HRM.7.e — Appraisal carried out at defined intervals, documented.")
    lb(doc, "Dated performance-appraisal record, at least annually, with a cycle-tracking record confirming no missed year.")
    h(doc, 1, "13. References")
    lb(doc, "NABH Accreditation Standards for Hospitals, 6th Edition — standard HRM.7.")
    lb(doc, "NABH Guidebook to Accreditation Standards for Hospitals, 6th Edition — HRM.7 interpretations.")
    h(doc, 1, "Disclaimer")
    hrm_disclaimer(doc)
    save_and_verify(doc, "HCO_HRM_7_v2_REWRITE_DRAFT.docx")


# ══════════════════════════════════════════════════════════════════════════════
# HRM.8 — Disciplinary and Grievance Handling   (no stop-work)
# ══════════════════════════════════════════════════════════════════════════════
def gen_hrm8():
    doc = Document()
    h(doc, 0, "Policy on Disciplinary and Grievance Handling")
    p(doc, HN)
    h(doc, 1, "Document control")
    doc_ctrl(doc, "HCO/HRM/POL/08", _HR)
    p(doc, "A blank marked ________ must be completed before issue.")
    h(doc, 1, "Statement of intent")
    p(doc, f"{HN} defines and implements a process for disciplinary and grievance handling, so that "
           "staff issues are addressed fairly, lawfully, and with a real path to redress.")
    h(doc, 1, "1. Purpose")
    p(doc, f"This policy sets out how {HN} handles disciplinary matters and staff grievances — the "
           "written guidance governing the mechanism, how staff know about it, the principles it "
           "follows, its legal basis, the appeal provision, and how grievances are redressed.")
    p(doc, "This policy does not cover recruitment, induction, performance appraisal, or staff "
           "health and safety — those are covered in other hospital policies. The other HRM "
           "standards have their own policies too.")
    h(doc, 1, "2. Scope")
    p(doc, f"This policy applies to all categories of staff at {HN}. Grievances covered include "
           "workplace issues such as bullying and harassment.")
    h(doc, 1, "3. Policy standards")
    p(doc, f"{HN} has written guidance governing its disciplinary and grievance handling mechanisms, "
           "covering how cases are handled fairly, the legal basis followed, and the appeal "
           "provision — including workplace grievances such as bullying and harassment.")
    p(doc, f"{HN} ensures the disciplinary and grievance handling mechanism is known to all "
           "categories of staff.")
    p(doc, f"{HN} bases its disciplinary policy and procedure on the principles of natural justice.")
    p(doc, f"{HN} keeps its disciplinary and grievance procedure in consonance with the prevailing "
           "laws applicable to its staff.")
    p(doc, f"{HN} provides for appeals in all disciplinary cases, through an appellate authority "
           "higher than the disciplinary authority.")
    p(doc, f"{HN} takes action to redress grievances and documents and communicates that action to "
           "the aggrieved staff member.")
    h(doc, 1, "4. Non-negotiable rules")
    ln(doc, "Do not run disciplinary or grievance handling without written guidance covering how "
            "cases are handled, the legal basis, and the appeal provision.")
    ln(doc, "Do not leave any category of staff unaware of the disciplinary and grievance handling "
            "mechanism.")
    ln(doc, "Do not decide a disciplinary case without allowing both the employee and the hospital "
            "to present their case.")
    ln(doc, "Do not run a disciplinary or grievance procedure that is out of step with the labour "
            "laws applicable to this hospital's staff.")
    ln(doc, "Do not deny a staff member the right to appeal a disciplinary decision to a higher "
            "authority than the one that made it.")
    ln(doc, "Do not leave a grievance unaddressed — action taken must be documented and communicated "
            "to the aggrieved staff member.")
    ln(doc, f"Staff who see a rule under this policy broken report it the same shift to the "
            f"HR In-Charge / Personnel Officer or the Medical Superintendent.")
    h(doc, 1, "5. What we do")
    h(doc, 2, "5.1 Written guidance")
    p(doc, f"Written guidance governs disciplinary and grievance handling mechanisms at {HN}, "
           "covering how cases falling under Sections 5.3 to 5.5 of this policy are handled, and "
           "including workplace grievances such as bullying and harassment.")
    h(doc, 2, "5.2 Staff awareness")
    p(doc, f"The disciplinary and grievance handling mechanism at {HN} is known to all categories "
           "of staff, who are aware of the procedure to follow if they feel aggrieved.")
    h(doc, 2, "5.3 Natural justice")
    p(doc, f"The disciplinary policy and procedure at {HN} are based on the principles of natural "
           "justice — both parties (employee and employer) are allowed to present their case before "
           "a decision is taken.")
    h(doc, 2, "5.4 Legal compliance")
    p(doc, f"The disciplinary and grievance procedure at {HN} is kept in consonance with the "
           "prevailing laws applicable to this hospital's staff category. An Internal Complaints "
           "Committee is established to handle complaints of sexual harassment.")
    h(doc, 2, "5.5 Appeals")
    p(doc, f"{HN} provides for appeals in all disciplinary cases, through a designated appellate "
           "authority higher than the disciplinary authority that issued the original decision.")
    h(doc, 2, "5.6 Grievance redress")
    p(doc, f"{HN} takes action to redress grievances through the written redress procedure; actions "
           "taken are documented and communicated to the aggrieved staff member.")
    h(doc, 1, "6. Governance and responsibility")
    gov_tbl(doc, [
        (_HR,
         "Owns this policy. Maintains the written guidance and grievance redress records."),
        ("Disciplinary Authority",
         "Conducts disciplinary proceedings in line with natural justice and applicable law."),
        ("Appellate Authority",
         "Hears appeals against disciplinary decisions; is senior to the disciplinary authority."),
        ("Internal Complaints Committee",
         "Handles complaints of sexual harassment."),
        (_MS,
         "Oversees implementation of this policy. Receives escalations under Section 4."),
        ("All Staff",
         "Know the mechanism and their right to raise a grievance or appeal. Report identified "
         "risks and rule breaches."),
    ])
    h(doc, 1, "7. Quality monitoring")
    p(doc, f"The HR In-Charge / Personnel Officer reviews disciplinary and grievance case records "
           "to confirm the mechanism is being followed. Documentary evidence is on file for the "
           "CORE objective element and the asterisked objective element. This policy itself is "
           "reviewed every year.")
    h(doc, 1, "8. Training and staff acknowledgement")
    p(doc, "All staff covered by this policy are made aware of the disciplinary and grievance "
           "handling mechanism and sign the acknowledgement below.")
    p(doc, f"I have read the Policy on Disciplinary and Grievance Handling of {HN}. I will follow the processes described.")
    sig_tbl(doc)
    h(doc, 1, "9. Distribution")
    p(doc, "HR department; all department heads / unit in-charges; all staff through the hospital "
           "intranet and department policy folders.")
    h(doc, 1, "10. Abbreviations")
    abbrev_tbl(doc, [
        ("ICC", "Internal Complaints Committee"),
    ] + HRM_ABBREVS_BASE)
    h(doc, 1, "11. Traceability to NABH HCO Full Accreditation 6th Edition HRM.8")
    p(doc, "This table is an index. It is not how the policy is organised. An asterisk in the "
           "Level column means documentation of the process is required.")
    trace_tbl(doc, [
        ("HRM.8.a", "Commitment*", "Section 3; 5.1", _HR),
        ("HRM.8.b", "Commitment",  "Section 3; 5.2", _HR),
        ("HRM.8.c", "Commitment",  "Section 3; 5.3", _HR),
        ("HRM.8.d", "CORE",        "Section 3; 5.4", _HR),
        ("HRM.8.e", "Commitment",  "Section 3; 5.5", _HR),
        ("HRM.8.f", "Commitment",  "Section 3; 5.6", _HR),
    ])
    h(doc, 1, "12. Required Records / Evidence Checklist")
    h(doc, 2, "HRM.8.a — Written guidance governs the mechanism.")
    lb(doc, "Written disciplinary and grievance handling guidance, covering how cases are handled, the legal basis, and the appeal provision.")
    lb(doc, "Workplace-issue inclusion record — bullying, harassment.")
    h(doc, 2, "HRM.8.b — Mechanism known to all categories of staff.")
    lb(doc, "Staff-awareness record of the disciplinary and grievance mechanism, across all staff categories.")
    h(doc, 2, "HRM.8.c — Disciplinary policy based on natural justice.")
    lb(doc, "Natural-justice-principle record — both parties heard before a decision — and disciplinary case record demonstrating the principle applied.")
    h(doc, 2, "HRM.8.d — Procedure in consonance with prevailing laws.")
    lb(doc, "Labour-law compliance record for the disciplinary and grievance procedure, with a legal-currency review record.")
    lb(doc, "Internal Complaints Committee constitution record for sexual-harassment complaints.")
    h(doc, 2, "HRM.8.e — Provision for appeals in all disciplinary cases.")
    lb(doc, "Appellate-authority designation record, higher than the disciplinary authority, and appeal-case record showing the provision was used.")
    h(doc, 2, "HRM.8.f — Actions taken to redress the grievance.")
    lb(doc, "Grievance-redress action record, with documentation and communication record to the aggrieved staff member, and closure-tracking record.")
    h(doc, 1, "13. References")
    lb(doc, "NABH Accreditation Standards for Hospitals, 6th Edition — standard HRM.8.")
    lb(doc, "NABH Guidebook to Accreditation Standards for Hospitals, 6th Edition — HRM.8 interpretations.")
    h(doc, 1, "Disclaimer")
    hrm_disclaimer(doc)
    save_and_verify(doc, "HCO_HRM_8_v2_REWRITE_DRAFT.docx")


# ══════════════════════════════════════════════════════════════════════════════
# HRM.9 — Staff Health and Safety   (no stop-work)
# ══════════════════════════════════════════════════════════════════════════════
def gen_hrm9():
    doc = Document()
    h(doc, 0, "Policy on Staff Health and Safety")
    p(doc, HN)
    h(doc, 1, "Document control")
    doc_ctrl(doc, "HCO/HRM/POL/09", _HR)
    p(doc, "A blank marked ________ must be completed before issue.")
    h(doc, 1, "Statement of intent")
    p(doc, f"{HN} addresses the health and safety needs of its staff, so that staff are protected "
           "from occupational health hazards, workplace injury, and workplace violence, and are "
           "cared for when harm does occur.")
    h(doc, 1, "1. Purpose")
    p(doc, f"This policy sets out how {HN} addresses staff health problems and occupational health "
           "hazards, conducts periodic health checks for staff in direct patient care, provides "
           "treatment for workplace-related injuries, and prevents and handles workplace violence.")
    p(doc, "This policy does not cover recruitment, induction, performance appraisal, or "
           "disciplinary and grievance handling — those are covered in other hospital policies. "
           "The other HRM standards have their own policies too.")
    h(doc, 1, "2. Scope")
    p(doc, f"This policy applies to all staff of {HN}, with periodic health checks specifically "
           "covering staff who deal with direct patient care.")
    h(doc, 1, "3. Policy standards")
    p(doc, f"{HN} takes care of the health problems of staff, including occupational health hazards, "
           "in accordance with its staff health and safety policy — covering physical and mental "
           "health, safe working conditions, vaccination and immunisation, personal protective "
           "equipment, and support for staff involved in unanticipated adverse events.")
    p(doc, f"{HN} conducts health checks for staff dealing with direct patient care at least once "
           "a year and documents the findings.")
    p(doc, f"{HN} provides treatment to staff who sustain workplace-related injuries, including "
           "injuries from workplace violence.")
    p(doc, f"{HN} has measures in place to prevent and handle workplace violence.")
    h(doc, 1, "4. Non-negotiable rules")
    ln(doc, "Do not leave a staff member's health problem, including an occupational health hazard, "
            "unaddressed under the hospital's staff health and safety policy.")
    ln(doc, "Do not let a year pass without a documented health check for a staff member in direct "
            "patient care.")
    ln(doc, "Do not deny treatment to a staff member who sustains a workplace-related injury.")
    ln(doc, "Do not leave the hospital without measures in place to prevent and handle workplace "
            "violence.")
    ln(doc, f"Staff who see a rule under this policy broken report it the same shift to the "
            f"HR In-Charge / Personnel Officer or the Medical Superintendent.")
    h(doc, 1, "5. What we do")
    h(doc, 2, "5.1 Staff health and safety policy")
    p(doc, f"{HN} has written guidance on staff health and safety addressing physical and mental "
           "health and safe working conditions across all shifts, in consonance with the law of "
           "the land and good work practice. The hospital runs a staff vaccination and immunisation "
           "programme, provides appropriate personal protective equipment with training on its use, "
           "and supports staff (as 'second victim') involved in unanticipated adverse events, "
           "medical error or patient-related injury.")
    h(doc, 2, "5.2 Periodic health checks")
    p(doc, f"Health checks for staff dealing with direct patient care are done at least once a year "
           f"at {HN}, with findings and results documented in the personal file. Parameters may "
           "differ by staff category; competent individuals may be identified to perform the checks. "
           "The staff member is not charged for the check; more frequent checks may be done where "
           "needed.")
    h(doc, 2, "5.3 Treatment for workplace injuries")
    p(doc, f"{HN} provides treatment to staff who sustain workplace-related injuries — for example "
           "needle-stick injuries, back injuries from patient transport, or noise-related hearing "
           "impairment — including counselling where appropriate. Injuries from workplace violence "
           "are included.")
    h(doc, 2, "5.4 Workplace violence prevention and handling")
    p(doc, f"{HN} has measures in place for prevention and handling of workplace violence, using an "
           "integrative and participative approach: workplace risk assessment identifying situations "
           "of special risk, workplace interventions (information and communication), environmental "
           "interventions (signage, security, restricted access) and individual interventions "
           "(training). A mechanism handles these situations, including liaison with law enforcement "
           "where applicable and counselling for affected staff, as part of the hospital's written "
           "security guidance, which is covered in other hospital policies.")
    h(doc, 1, "6. Governance and responsibility")
    gov_tbl(doc, [
        (_HR,
         "Owns this policy. Maintains staff health records and coordinates annual health checks."),
        ("Occupational Health In-Charge",
         "Runs the vaccination and immunisation programme, PPE provision and training, and "
         "second-victim support."),
        ("Security In-Charge",
         "Implements workplace-violence prevention measures and the incident-handling mechanism."),
        (_MS,
         "Oversees implementation of this policy. Receives escalations under Section 4."),
        ("All Staff",
         "Attend health checks and use provided PPE. Report identified risks and rule breaches."),
    ])
    h(doc, 1, "7. Quality monitoring")
    p(doc, f"The HR In-Charge / Personnel Officer reviews staff health records to confirm annual "
           "health checks are completed on schedule. Documentary evidence is on file for the CORE "
           "objective element and the asterisked objective elements. This policy itself is reviewed "
           "every year.")
    h(doc, 1, "8. Training and staff acknowledgement")
    p(doc, "All staff covered by this policy are made aware of the staff health and safety measures "
           "described in Section 5 and sign the acknowledgement below.")
    p(doc, f"I have read the Policy on Staff Health and Safety of {HN}. I will follow the processes described.")
    sig_tbl(doc)
    h(doc, 1, "9. Distribution")
    p(doc, "HR department; all department heads / unit in-charges; all staff through the hospital "
           "intranet and department policy folders.")
    h(doc, 1, "10. Abbreviations")
    abbrev_tbl(doc, HRM_ABBREVS_BASE + [
        ("PPE",  "Personal Protective Equipment"),
    ])
    h(doc, 1, "11. Traceability to NABH HCO Full Accreditation 6th Edition HRM.9")
    p(doc, "This table is an index. It is not how the policy is organised. An asterisk in the "
           "Level column means documentation of the process is required.")
    trace_tbl(doc, [
        ("HRM.9.a", "Commitment*", "Section 3; 5.1", _HR),
        ("HRM.9.b", "Commitment",  "Section 3; 5.2", _HR),
        ("HRM.9.c", "Commitment",  "Section 3; 5.3", _HR),
        ("HRM.9.d", "CORE*",       "Section 3; 5.4", _HR),
    ])
    h(doc, 1, "12. Required Records / Evidence Checklist")
    h(doc, 2, "HRM.9.a — Health problems, including occupational hazards, taken care of.")
    lb(doc, "Written staff health and safety policy covering physical and mental health.")
    lb(doc, "Staff vaccination and immunisation programme record.")
    lb(doc, "PPE-provision and second-victim-support record.")
    h(doc, 2, "HRM.9.b — Annual health checks for direct-patient-care staff, documented.")
    lb(doc, "Annual health-check record for direct-patient-care staff, with findings and results documentation in the personal file, and no-charge-to-staff confirmation.")
    h(doc, 2, "HRM.9.c — Treatment provided for workplace-related injuries.")
    lb(doc, "Workplace-injury treatment record — needle-stick, patient-transport injury, noise-related, etc. — including counselling record where appropriate, and workplace-violence-injury inclusion record.")
    h(doc, 2, "HRM.9.d — Measures for prevention and handling of workplace violence.")
    lb(doc, "Workplace-violence risk-assessment record.")
    lb(doc, "Written security guidance covering workplace-violence prevention and handling.")
    lb(doc, "Law-enforcement-liaison and counselling record for affected staff.")
    h(doc, 1, "13. References")
    lb(doc, "NABH Accreditation Standards for Hospitals, 6th Edition — standard HRM.9.")
    lb(doc, "NABH Guidebook to Accreditation Standards for Hospitals, 6th Edition — HRM.9 interpretations.")
    h(doc, 1, "Disclaimer")
    hrm_disclaimer(doc)
    save_and_verify(doc, "HCO_HRM_9_v2_REWRITE_DRAFT.docx")


# ══════════════════════════════════════════════════════════════════════════════
# HRM.10 — Staff Personal Information and Records   (no stop-work)
# ══════════════════════════════════════════════════════════════════════════════
def gen_hrm10():
    doc = Document()
    h(doc, 0, "Policy on Staff Personal Information and Records")
    p(doc, HN)
    h(doc, 1, "Document control")
    doc_ctrl(doc, "HCO/HRM/POL/10", _HR)
    p(doc, "A blank marked ________ must be completed before issue.")
    h(doc, 1, "Statement of intent")
    p(doc, f"{HN} maintains documented personal information for each staff member, so that a "
           "complete, confidential, and current record is available for every category of staff.")
    h(doc, 1, "1. Purpose")
    p(doc, f"This policy sets out how {HN} maintains staff personal files — what they contain, "
           "how confidentiality is protected, how in-service training is recorded, and how "
           "evaluation results and remarks are documented.")
    p(doc, "This policy does not cover recruitment, induction, performance appraisal, or staff "
           "health and safety — those are covered in other hospital policies. The other HRM "
           "standards have their own policies too.")
    h(doc, 1, "2. Scope")
    p(doc, f"This policy applies to the personal file of every staff member at {HN}.")
    h(doc, 1, "3. Policy standards")
    p(doc, f"{HN} maintains personal files for all staff and ensures their confidentiality.")
    p(doc, f"{HN} ensures personal files contain each staff member's qualification, job "
           "description, credential verification, and health status.")
    p(doc, f"{HN} maintains records of in-service training and education for staff.")
    p(doc, f"{HN} ensures personal files contain the results of all evaluations and remarks.")
    h(doc, 1, "4. Non-negotiable rules")
    ln(doc, "Do not maintain a staff personal file without ensuring its confidentiality and "
            "restricting access to it.")
    ln(doc, "Do not leave a personal file without the staff member's qualification, job "
            "description, credential verification and health status on record.")
    ln(doc, "Do not skip maintaining a record of in-service training and education for a staff "
            "member.")
    ln(doc, "Do not leave a personal file without the results of all evaluations and remarks on "
            "record.")
    ln(doc, f"Staff who see a rule under this policy broken report it the same shift to the "
            f"HR In-Charge / Personnel Officer or the Medical Superintendent.")
    h(doc, 1, "5. What we do")
    h(doc, 2, "5.1 Personal file maintenance and confidentiality")
    p(doc, f"Personal files are maintained for all staff at {HN}, kept current and updated "
           "(electronic format is acceptable), with confidentiality maintained and access "
           "restricted.")
    h(doc, 2, "5.2 Personal file content")
    p(doc, f"Each staff member's personal file at {HN} contains their qualification, job "
           "description, verification of credentials and health status.")
    h(doc, 2, "5.3 In-service training records")
    p(doc, f"Records of in-service training and education are maintained for staff at {HN}. "
           "For internal training, an annual summary may be filed, with a supporting document "
           "verifying attendance. Where training records are held elsewhere, the personal file "
           "carries traceability to them; electronic training records are acceptable.")
    h(doc, 2, "5.4 Evaluation results and remarks")
    p(doc, f"Personal files at {HN} contain the results of all evaluations and remarks — "
           "performance appraisals, training assessments, health-check outcomes, and records of "
           "achievement, appreciation, complaint, warning or memo.")
    h(doc, 1, "6. Governance and responsibility")
    gov_tbl(doc, [
        (_HR,
         "Owns this policy. Maintains personal files and controls access to them."),
        ("Department heads / Unit In-Charges",
         "Provide evaluation, appraisal, and training-attendance inputs for personal files."),
        (_MS,
         "Oversees implementation of this policy. Receives escalations under Section 4."),
        ("All Staff",
         "Provide accurate information for their personal file. Report identified risks and "
         "rule breaches."),
    ])
    h(doc, 1, "7. Quality monitoring")
    p(doc, f"The HR In-Charge / Personnel Officer periodically reviews personal files for "
           "completeness and currency. This policy itself is reviewed every year.")
    h(doc, 1, "8. Training and staff acknowledgement")
    p(doc, "All staff covered by this policy are made aware of what their personal file contains "
           "and sign the acknowledgement below.")
    p(doc, f"I have read the Policy on Staff Personal Information and Records of {HN}. I will follow the processes described.")
    sig_tbl(doc)
    h(doc, 1, "9. Distribution")
    p(doc, "HR department; all department heads / unit in-charges; all staff through the hospital "
           "intranet and department policy folders.")
    h(doc, 1, "10. Abbreviations")
    abbrev_tbl(doc, HRM_ABBREVS_BASE)
    h(doc, 1, "11. Traceability to NABH HCO Full Accreditation 6th Edition HRM.10")
    p(doc, "This table is an index. It is not how the policy is organised. An asterisk in the "
           "Level column means documentation of the process is required.")
    trace_tbl(doc, [
        ("HRM.10.a", "Commitment", "Section 3; 5.1", _HR),
        ("HRM.10.b", "Commitment", "Section 3; 5.2", _HR),
        ("HRM.10.c", "Commitment", "Section 3; 5.3", _HR),
        ("HRM.10.d", "Commitment", "Section 3; 5.4", _HR),
    ])
    h(doc, 1, "12. Required Records / Evidence Checklist")
    h(doc, 2, "HRM.10.a — Personal files maintained, confidentiality ensured.")
    lb(doc, "Personal-file maintenance record for all staff, current and updated, with a confidentiality and access-restriction record.")
    h(doc, 2, "HRM.10.b — Personal file contains qualification, job description, credential verification, health status.")
    lb(doc, "Personal-file content record — qualification, job description, credential verification, health status — with a completeness-check record.")
    h(doc, 2, "HRM.10.c — Records of in-service training and education maintained.")
    lb(doc, "In-service training and education record maintained in, or traceable from, the personal file, with an annual training-summary record and attendance-verification supporting document.")
    h(doc, 2, "HRM.10.d — Personal files contain evaluation results and remarks.")
    lb(doc, "Personal-file record of evaluation results and remarks — appraisals, training assessment, health-check outcome, achievement / complaint / warning / memo — with a confidential-handling record.")
    h(doc, 1, "13. References")
    lb(doc, "NABH Accreditation Standards for Hospitals, 6th Edition — standard HRM.10.")
    lb(doc, "NABH Guidebook to Accreditation Standards for Hospitals, 6th Edition — HRM.10 interpretations.")
    h(doc, 1, "Disclaimer")
    hrm_disclaimer(doc)
    save_and_verify(doc, "HCO_HRM_10_v2_REWRITE_DRAFT.docx")


# ══════════════════════════════════════════════════════════════════════════════
# Shared helper — stop-work section for credentialing standards (HRM.11-13)
# ══════════════════════════════════════════════════════════════════════════════
def _cred_stop_work(doc, prof):
    """Insert the stop-work authority section for a credentialing policy.

    prof — the professional type phrase, e.g. 'medical professional'.
    """
    p(doc, f"Do not let a {prof} provide patient care without supervision unless the "
           f"HR In-Charge / Personnel Officer and the credentialing file show that professional "
           f"identified and privileged for that care under this policy.")
    p(doc, f"Stop-work applies to starting or continuing unsupervised care by that professional. "
           "Care already under way is handed to a privileged professional or brought under "
           "supervision; immediate life-saving care is not withdrawn while that handover happens.")
    p(doc, f"The person who stops tells the HR In-Charge / Personnel Officer and the Medical "
           "Superintendent the same shift. Refusing to let an uncredentialed or unprivileged "
           f"professional practise unsupervised is not a disciplinary matter.")


# ══════════════════════════════════════════════════════════════════════════════
# HRM.11 — Credentialing and Privileging of Medical Professionals
#           Stop-work: HRM.11.a and HRM.11.d ONLY
# ══════════════════════════════════════════════════════════════════════════════
def gen_hrm11():
    doc = Document()
    h(doc, 0, "Policy on Credentialing and Privileging of Medical Professionals")
    p(doc, HN)
    h(doc, 1, "Document control")
    doc_ctrl(doc, "HCO/HRM/POL/11", _HR)
    p(doc, "A blank marked ________ must be completed before issue.")
    h(doc, 1, "Statement of intent")
    p(doc, f"{HN} runs a process for credentialing and privileging of medical professionals "
           "permitted to provide patient care without supervision, so that only identified, "
           "qualified, and privileged professionals provide unsupervised care.")
    h(doc, 1, "1. Purpose")
    p(doc, f"This policy sets out how {HN} identifies medical professionals permitted to provide "
           "unsupervised patient care, documents and verifies their credentials, grants and "
           "communicates their privileges, and ensures they practise within those privileges.")
    p(doc, "This policy does not cover recruitment, induction, safety and quality training, or "
           "staff health and safety — those are covered in other hospital policies. The other "
           "HRM standards have their own policies too.")
    h(doc, 1, "2. Scope")
    p(doc, f"This policy applies to all medical professionals at {HN} permitted by law, "
           "regulation and the organisation to provide patient care without supervision.")
    h(doc, 1, "3. Policy standards")
    p(doc, f"{HN} identifies medical professionals permitted to provide patient care without "
           "supervision.")
    p(doc, f"{HN} documents and periodically updates the education, registration, training and "
           "experience of identified medical professionals.")
    p(doc, f"{HN} verifies information about medical professionals when possible.")
    p(doc, f"{HN} grants medical professionals privileges to admit and care for patients in "
           "consonance with their qualification, training, experience and registration, reviewed "
           "at least annually.")
    p(doc, f"{HN} ensures the requisite services a medical professional is authorised to provide "
           "are known to that professional and to relevant departments.")
    p(doc, f"{HN} ensures medical professionals admit and care for patients as per their "
           "privileging.")
    h(doc, 1, "4. Non-negotiable rules")
    ln(doc, "Do not let a medical professional provide unsupervised patient care without first "
            "identifying them on the permitted list.")
    ln(doc, "Do not leave a medical professional's education, registration, training or "
            "experience undocumented or out of date.")
    ln(doc, "Do not skip verifying a medical professional's credentials with the awarding "
            "organisation when possible.")
    ln(doc, "Do not grant or allow privileges that are not in consonance with a medical "
            "professional's qualification, training, experience and registration.")
    ln(doc, "Do not leave a medical professional or the relevant departments unaware of the "
            "services that professional is authorised to provide.")
    ln(doc, "Do not let a medical professional admit or care for a patient outside their granted "
            "privileges.")
    ln(doc, "Do not bypass the stop-work authority in Section 6 when the trigger conditions are "
            "met.")
    ln(doc, f"Staff who see a rule under this policy broken report it the same shift to the "
            f"HR In-Charge / Personnel Officer or the Medical Superintendent.")
    h(doc, 1, "5. What we do")
    h(doc, 2, "5.1 Identification of permitted professionals")
    p(doc, f"{HN} identifies medical professionals permitted by law, regulation and the "
           "organisation to provide patient care without supervision — individuals with the "
           "required qualification(s), training and experience, in consonance with the law. "
           "Providing unsupervised care outside this identified list is a stop-work trigger "
           "(Section 6).")
    h(doc, 2, "5.2 Credential documentation")
    p(doc, f"The education, registration, training and experience of identified medical "
           f"professionals at {HN} are documented and updated periodically — after acquisition "
           "of new skills or qualification — and maintained in each professional's personal file.")
    h(doc, 2, "5.3 Credential verification")
    p(doc, f"Information about medical professionals at {HN} is appropriately verified when "
           "possible, by checking with the organisation that awarded the qualification or "
           "training. The National Medical Commission's website is a useful reference for "
           "verification.")
    h(doc, 2, "5.4 Granting of privileges")
    p(doc, f"{HN} grants medical professionals privileges to admit and care for patients in "
           "consonance with their qualification, training, experience and registration — "
           "identifying the clinical services each is authorised to perform (for example, "
           "radiotherapy only by a radiation oncologist). Privileges are reviewed at least "
           "annually and revised where necessary. Granting or exercising privileges outside "
           "this record is a stop-work trigger (Section 6).")
    h(doc, 2, "5.5 Communication of privileges")
    p(doc, f"The requisite services a medical professional at {HN} is authorised to provide "
           "are known to that professional and to the relevant departments — for example OP "
           "consultation rights, admission rights and rights to specific procedures or surgeries "
           "(inclusion or exclusion). Concerned departments (for example, front desk for "
           "admission rights, the operation theatre for surgical rights) are informed of the "
           "relevant privileging.")
    h(doc, 2, "5.6 Practising within privileges")
    p(doc, f"Medical professionals at {HN} admit and care for patients as per their privileging, "
           "using a standardised format applied uniformly. New faculty may work under "
           "proctorship until independent privileges are granted; the hospital maintains a "
           "mechanism confirming professionals provide only the services they are privileged for.")
    h(doc, 1, "6. Stop-work authority")
    _cred_stop_work(doc, "medical professional")
    h(doc, 1, "7. Governance and responsibility")
    gov_tbl(doc, [
        (_HR,
         "Owns this policy. Maintains the credentialing file and the identified-professional "
         "list."),
        ("Credentialing Committee",
         "Reviews credentials and grants privileges; conducts the annual privilege review."),
        (_MS,
         "Oversees implementation of this policy. Receives stop-work escalations."),
        ("All Staff",
         "Report any medical professional providing unsupervised care outside the "
         "identified or privileged list."),
    ])
    h(doc, 1, "8. Quality monitoring")
    p(doc, f"The HR In-Charge / Personnel Officer reviews credentialing files to confirm every "
           "medical professional providing unsupervised care is identified and privileged. "
           "Documentary evidence is on file for each CORE objective element. Compliance with "
           "the stop-work authority in Section 6 is monitored through the credentialing file "
           "review. This policy itself is reviewed every year.")
    h(doc, 1, "9. Training and staff acknowledgement")
    p(doc, "All staff covered by this policy are made aware of the credentialing and privileging "
           "process, including the stop-work authority in Section 6, and sign the acknowledgement "
           "below.")
    p(doc, f"I have read the Policy on Credentialing and Privileging of Medical Professionals of {HN}. I will follow the processes described.")
    sig_tbl(doc)
    h(doc, 1, "10. Distribution")
    p(doc, "HR department; Medical Superintendent's office; all department heads / unit "
           "in-charges; all staff through the hospital intranet and department policy folders.")
    h(doc, 1, "11. Abbreviations")
    abbrev_tbl(doc, HRM_ABBREVS_BASE + [
        ("NMC", "National Medical Commission"),
    ])
    h(doc, 1, "12. Traceability to NABH HCO Full Accreditation 6th Edition HRM.11")
    p(doc, "This table is an index. It is not how the policy is organised. An asterisk in the "
           "Level column means documentation of the process is required. Stop-work applies ONLY "
           "to the OEs marked 'Section 6 Stop-work' below.")
    trace_tbl(doc, [
        ("HRM.11.a", "CORE",       "Section 3; 5.1; Section 6 Stop-work", _HR),
        ("HRM.11.b", "Commitment", "Section 3; 5.2", _HR),
        ("HRM.11.c", "Commitment", "Section 3; 5.3", _HR),
        ("HRM.11.d", "CORE",       "Section 3; 5.4; Section 6 Stop-work", _HR),
        ("HRM.11.e", "Commitment", "Section 3; 5.5", _HR),
        ("HRM.11.f", "Commitment", "Section 3; 5.6", _HR),
    ])
    h(doc, 1, "13. Required Records / Evidence Checklist")
    h(doc, 2, "HRM.11.a — Permitted medical professionals identified. [Stop-work trigger]")
    lb(doc, "Identified-medical-professional list permitted to provide unsupervised patient care, with a qualification, training and experience verification record.")
    lb(doc, "Cross-reference to the stop-work authority in Section 6 for unlisted practice.")
    h(doc, 2, "HRM.11.b — Credentials documented and updated periodically.")
    lb(doc, "Education, registration, training and experience documentation record, updated periodically, with a personal-file record of each update.")
    h(doc, 2, "HRM.11.c — Credentials verified when possible.")
    lb(doc, "Verification record with the awarding organisation, including a National Medical Commission or equivalent reference-check record.")
    h(doc, 2, "HRM.11.d — Privileges granted in consonance with qualification, training, experience and registration. [Stop-work trigger]")
    lb(doc, "Granted-privilege record naming the clinical services each professional is authorised for, with an annual privilege-review record.")
    lb(doc, "Cross-reference to the stop-work authority in Section 6 for privileging outside this record.")
    h(doc, 2, "HRM.11.e — Requisite services known to professional and relevant departments.")
    lb(doc, "Communicated-service record to the professional and to relevant departments, including admission-rights or surgical-rights notification records (front desk, operation theatre, etc.).")
    h(doc, 2, "HRM.11.f — Professionals admit and care as per privileging.")
    lb(doc, "Standardised privileging-format record, applied uniformly.")
    lb(doc, "Proctorship record for new faculty until independent privileges are granted.")
    lb(doc, "Mechanism-confirmation record that professionals provide only privileged services.")
    h(doc, 1, "14. References")
    lb(doc, "NABH Accreditation Standards for Hospitals, 6th Edition — standard HRM.11.")
    lb(doc, "NABH Guidebook to Accreditation Standards for Hospitals, 6th Edition — HRM.11 interpretations.")
    h(doc, 1, "Disclaimer")
    hrm_disclaimer(doc)
    save_and_verify(doc, "HCO_HRM_11_v2_REWRITE_DRAFT.docx")


# ══════════════════════════════════════════════════════════════════════════════
# HRM.12 — Credentialing and Privileging of Nursing Professionals
#           Stop-work: HRM.12.a and HRM.12.d ONLY
#           Statute: Indian Nursing Council Act, 1947 (explicitly sourced)
# ══════════════════════════════════════════════════════════════════════════════
def gen_hrm12():
    doc = Document()
    h(doc, 0, "Policy on Credentialing and Privileging of Nursing Professionals")
    p(doc, HN)
    h(doc, 1, "Document control")
    doc_ctrl(doc, "HCO/HRM/POL/12", _HR)
    p(doc, "A blank marked ________ must be completed before issue.")
    h(doc, 1, "Statement of intent")
    p(doc, f"{HN} runs a process for credentialing and privileging of nursing professionals "
           "permitted to provide patient care without supervision, so that only identified, "
           "qualified, and privileged nurses provide unsupervised care.")
    h(doc, 1, "1. Purpose")
    p(doc, f"This policy sets out how {HN} identifies nursing staff permitted to provide "
           "unsupervised patient care, documents and verifies their credentials, grants and "
           "communicates their privileges, and ensures they practise within those privileges.")
    p(doc, "This policy does not cover recruitment, induction, safety and quality training, or "
           "staff health and safety — those are covered in other hospital policies. The other "
           "HRM standards have their own policies too.")
    h(doc, 1, "2. Scope")
    p(doc, f"This policy applies to all nursing staff at {HN} permitted by law, regulation and "
           "the organisation to provide patient care without supervision.")
    h(doc, 1, "3. Policy standards")
    p(doc, f"{HN} identifies nursing staff permitted to provide patient care without supervision, "
           "in consonance with the Indian Nursing Council Act, 1947.")
    p(doc, f"{HN} documents and periodically updates the education, registration, training and "
           "experience of nursing staff.")
    p(doc, f"{HN} verifies information about nursing staff when possible.")
    p(doc, f"{HN} grants nursing staff privileges in consonance with their qualification, "
           "training, experience and registration, reviewed at least annually.")
    p(doc, f"{HN} ensures the requisite services a nursing professional is authorised to provide "
           "are known to that professional and to relevant departments.")
    p(doc, f"{HN} ensures nursing professionals care for patients as per their privileging.")
    h(doc, 1, "4. Non-negotiable rules")
    ln(doc, "Do not let a nursing professional provide unsupervised patient care without first "
            "identifying them on the permitted list, in consonance with the Indian Nursing "
            "Council Act, 1947.")
    ln(doc, "Do not leave a nursing professional's education, registration, training or "
            "experience undocumented or out of date.")
    ln(doc, "Do not skip verifying a nursing professional's credentials with the awarding "
            "organisation when possible.")
    ln(doc, "Do not grant or allow privileges that are not in consonance with a nursing "
            "professional's qualification, training, experience and registration.")
    ln(doc, "Do not leave a nursing professional or the relevant departments unaware of the "
            "services that professional is authorised to provide.")
    ln(doc, "Do not let a nursing professional care for a patient outside their granted "
            "privileges.")
    ln(doc, "Do not bypass the stop-work authority in Section 6 when the trigger conditions are "
            "met.")
    ln(doc, f"Staff who see a rule under this policy broken report it the same shift to the "
            f"HR In-Charge / Personnel Officer or the Medical Superintendent.")
    h(doc, 1, "5. What we do")
    h(doc, 2, "5.1 Identification of permitted professionals")
    p(doc, f"{HN} identifies nursing staff permitted by law, regulation and the organisation "
           "to provide patient care without supervision — individuals with the required "
           "qualification(s), training and experience, in consonance with the law (Indian "
           "Nursing Council Act, 1947). Providing unsupervised care outside this identified "
           "list is a stop-work trigger (Section 6).")
    h(doc, 2, "5.2 Credential documentation")
    p(doc, f"The education, registration, training and experience of nursing staff at {HN} are "
           "documented and updated periodically, after acquisition of new skills or qualification.")
    h(doc, 2, "5.3 Credential verification")
    p(doc, f"Information about nursing staff at {HN} is appropriately verified when possible, "
           "by checking with the organisation that awarded the qualification or training.")
    h(doc, 2, "5.4 Granting of privileges")
    p(doc, f"{HN} grants nursing staff privileges in consonance with their qualification, "
           "training, experience and registration — identifying what each nurse is authorised "
           "to do (for example, an infection-prevention-and-control nurse needs the requisite "
           "in-house or external training, experience, aptitude and knowledge for that role). "
           "Privileges are reviewed at least annually and revised where necessary. Granting or "
           "exercising privileges outside this record is a stop-work trigger (Section 6).")
    h(doc, 2, "5.5 Communication of privileges")
    p(doc, f"The requisite services a nursing professional at {HN} is authorised to provide "
           "are known to that professional, to nursing services and to the concerned departments, "
           "communicated internally.")
    h(doc, 2, "5.6 Practising within privileges")
    p(doc, f"Nursing professionals at {HN} care for patients as per their privileging. New "
           "staff may work under supervision until independent privileges are granted; the "
           "hospital maintains a mechanism confirming nursing professionals provide only the "
           "services they are privileged for.")
    h(doc, 1, "6. Stop-work authority")
    _cred_stop_work(doc, "nursing professional")
    h(doc, 1, "7. Governance and responsibility")
    gov_tbl(doc, [
        (_HR,
         "Owns this policy. Maintains the credentialing file and the identified-professional "
         "list."),
        ("Nursing Superintendent",
         "Reviews nursing credentials and grants privileges; conducts the annual privilege "
         "review."),
        (_MS,
         "Oversees implementation of this policy. Receives stop-work escalations."),
        ("All Staff",
         "Report any nursing professional providing unsupervised care outside the "
         "identified or privileged list."),
    ])
    h(doc, 1, "8. Quality monitoring")
    p(doc, f"The HR In-Charge / Personnel Officer reviews credentialing files to confirm every "
           "nursing professional providing unsupervised care is identified and privileged. "
           "Documentary evidence is on file for each CORE objective element. Compliance with "
           "the stop-work authority in Section 6 is monitored through the credentialing file "
           "review. This policy itself is reviewed every year.")
    h(doc, 1, "9. Training and staff acknowledgement")
    p(doc, "All staff covered by this policy are made aware of the credentialing and privileging "
           "process, including the stop-work authority in Section 6, and sign the acknowledgement "
           "below.")
    p(doc, f"I have read the Policy on Credentialing and Privileging of Nursing Professionals of {HN}. I will follow the processes described.")
    sig_tbl(doc)
    h(doc, 1, "10. Distribution")
    p(doc, "HR department; Nursing Superintendent's office; all department heads / unit "
           "in-charges; all staff through the hospital intranet and department policy folders.")
    h(doc, 1, "11. Abbreviations")
    abbrev_tbl(doc, HRM_ABBREVS_BASE + [
        ("INC", "Indian Nursing Council"),
    ])
    h(doc, 1, "12. Traceability to NABH HCO Full Accreditation 6th Edition HRM.12")
    p(doc, "This table is an index. It is not how the policy is organised. An asterisk in the "
           "Level column means documentation of the process is required. Stop-work applies ONLY "
           "to the OEs marked 'Section 6 Stop-work' below.")
    trace_tbl(doc, [
        ("HRM.12.a", "CORE",       "Section 3; 5.1; Section 6 Stop-work", _HR),
        ("HRM.12.b", "Commitment", "Section 3; 5.2", _HR),
        ("HRM.12.c", "Commitment", "Section 3; 5.3", _HR),
        ("HRM.12.d", "CORE",       "Section 3; 5.4; Section 6 Stop-work", _HR),
        ("HRM.12.e", "Commitment", "Section 3; 5.5", _HR),
        ("HRM.12.f", "Commitment", "Section 3; 5.6", _HR),
    ])
    h(doc, 1, "13. Required Records / Evidence Checklist")
    h(doc, 2, "HRM.12.a — Permitted nursing staff identified. [Stop-work trigger]")
    lb(doc, "Identified-nursing-professional list permitted to provide unsupervised patient care, with a qualification, training and experience verification record referencing the Indian Nursing Council Act, 1947.")
    lb(doc, "Cross-reference to the stop-work authority in Section 6 for unlisted practice.")
    h(doc, 2, "HRM.12.b — Credentials documented and updated periodically.")
    lb(doc, "Education, registration, training and experience documentation record, updated periodically, with a personal-file record of each update.")
    h(doc, 2, "HRM.12.c — Credentials verified when possible.")
    lb(doc, "Verification record with the awarding organisation, with a verification-completeness record.")
    h(doc, 2, "HRM.12.d — Privileges granted in consonance with qualification, training, experience and registration. [Stop-work trigger]")
    lb(doc, "Granted-privilege record naming what each nurse is authorised to do, with an annual privilege-review record.")
    lb(doc, "Cross-reference to the stop-work authority in Section 6 for privileging outside this record.")
    h(doc, 2, "HRM.12.e — Requisite services known to nurse and relevant departments.")
    lb(doc, "Communicated-service record to the nurse and to nursing services and concerned departments, with an internal-communication record.")
    h(doc, 2, "HRM.12.f — Nursing professionals care as per privileging.")
    lb(doc, "Supervision record for new staff until independent privileges are granted.")
    lb(doc, "Mechanism-confirmation and privileging-compliance spot-check record.")
    h(doc, 1, "14. References")
    lb(doc, "NABH Accreditation Standards for Hospitals, 6th Edition — standard HRM.12.")
    lb(doc, "NABH Guidebook to Accreditation Standards for Hospitals, 6th Edition — HRM.12 interpretations.")
    lb(doc, "Indian Nursing Council Act, 1947.")
    h(doc, 1, "Disclaimer")
    hrm_disclaimer(doc)
    save_and_verify(doc, "HCO_HRM_12_v2_REWRITE_DRAFT.docx")


# ══════════════════════════════════════════════════════════════════════════════
# HRM.13 — Credentialing and Privileging of Para-Clinical Professionals
#           Stop-work: HRM.13.a and HRM.13.c ONLY
#           No annual review frequency for privileges (not stated in source)
# ══════════════════════════════════════════════════════════════════════════════
def gen_hrm13():
    doc = Document()
    h(doc, 0, "Policy on Credentialing and Privileging of Para-Clinical Professionals")
    p(doc, HN)
    h(doc, 1, "Document control")
    doc_ctrl(doc, "HCO/HRM/POL/13", _HR)
    p(doc, "A blank marked ________ must be completed before issue.")
    h(doc, 1, "Statement of intent")
    p(doc, f"{HN} runs a process for credentialing and privileging of para-clinical professionals "
           "permitted to provide patient care without supervision, so that only identified, "
           "qualified, and privileged professionals provide unsupervised care.")
    h(doc, 1, "1. Purpose")
    p(doc, f"This policy sets out how {HN} identifies para-clinical professionals permitted to "
           "provide unsupervised patient care, documents and verifies their credentials, grants "
           "and communicates their privileges, and ensures they practise within those privileges.")
    p(doc, "This policy does not cover recruitment, induction, safety and quality training, or "
           "staff health and safety — those are covered in other hospital policies. The other "
           "HRM standards have their own policies too.")
    h(doc, 1, "2. Scope")
    p(doc, f"This policy applies to all para-clinical professionals at {HN} — for example "
           "physiotherapists, rehabilitation therapists, dieticians, pharmacists, clinical "
           "pharmacists and technologists — permitted by law, regulation and the organisation "
           "to provide patient care without supervision.")
    h(doc, 1, "3. Policy standards")
    p(doc, f"{HN} identifies para-clinical professionals permitted to provide patient care "
           "without supervision.")
    p(doc, f"{HN} documents and periodically updates the education, registration, training and "
           "experience of para-clinical professionals.")
    p(doc, f"{HN} grants para-clinical professionals privileges in consonance with their "
           "qualification, training, experience and registration.")
    p(doc, f"{HN} ensures the requisite services a para-clinical professional is authorised to "
           "provide are known to that professional and to relevant departments.")
    p(doc, f"{HN} ensures para-clinical professionals care for patients as per their privileging.")
    h(doc, 1, "4. Non-negotiable rules")
    ln(doc, "Do not let a para-clinical professional provide unsupervised patient care without "
            "first identifying them on the permitted list.")
    ln(doc, "Do not leave a para-clinical professional's education, registration, training or "
            "experience unverified, undocumented, or out of date.")
    ln(doc, "Do not grant or allow privileges that are not in consonance with a para-clinical "
            "professional's qualification, training, experience and registration.")
    ln(doc, "Do not leave a para-clinical professional or the relevant departments unaware of "
            "the services that professional is authorised to provide.")
    ln(doc, "Do not let a para-clinical professional care for a patient outside their granted "
            "privileges.")
    ln(doc, "Do not bypass the stop-work authority in Section 6 when the trigger conditions are "
            "met.")
    ln(doc, f"Staff who see a rule under this policy broken report it the same shift to the "
            f"HR In-Charge / Personnel Officer or the Medical Superintendent.")
    h(doc, 1, "5. What we do")
    h(doc, 2, "5.1 Identification of permitted professionals")
    p(doc, f"{HN} identifies para-clinical professionals (for example physiotherapist, "
           "rehabilitation therapist, dietician, pharmacist, clinical pharmacist, technologist) "
           "permitted by law, regulation and the organisation to provide patient care without "
           "supervision — individuals with the required qualification(s), training and experience, "
           "in consonance with the law. Providing unsupervised care outside this identified list "
           "is a stop-work trigger (Section 6).")
    h(doc, 2, "5.2 Credential documentation and verification")
    p(doc, f"The education, registration, training and experience of para-clinical professionals "
           f"at {HN} are appropriately verified, documented and updated periodically, after "
           "acquisition of new skills or qualification, by checking with the organisation that "
           "awarded the qualification or training.")
    h(doc, 2, "5.3 Granting of privileges")
    p(doc, f"{HN} grants para-clinical professionals privileges in consonance with their "
           "qualification, training, experience and registration — specifying what each is "
           "authorised to do, with the requisite registration or licence held where applicable. "
           "Granting or exercising privileges outside this record is a stop-work trigger "
           "(Section 6).")
    h(doc, 2, "5.4 Communication of privileges")
    p(doc, f"The requisite services a para-clinical professional at {HN} is authorised to "
           "provide are known to that professional and to the concerned departments, communicated "
           "internally.")
    h(doc, 2, "5.5 Practising within privileges")
    p(doc, f"Para-clinical professionals at {HN} care for patients as per their privileging. "
           "New staff may work under supervision until independent privileges are granted; the "
           "hospital maintains a mechanism confirming para-clinical professionals provide only "
           "the services they are privileged for.")
    h(doc, 1, "6. Stop-work authority")
    _cred_stop_work(doc, "para-clinical professional")
    h(doc, 1, "7. Governance and responsibility")
    gov_tbl(doc, [
        (_HR,
         "Owns this policy. Maintains the credentialing file and the identified-professional "
         "list."),
        ("Department heads / Unit In-Charges",
         "Verify credentials and confirm para-clinical professionals practise within their "
         "granted privileges."),
        (_MS,
         "Oversees implementation of this policy. Receives stop-work escalations."),
        ("All Staff",
         "Report any para-clinical professional providing unsupervised care outside the "
         "identified or privileged list."),
    ])
    h(doc, 1, "8. Quality monitoring")
    p(doc, f"The HR In-Charge / Personnel Officer reviews credentialing files to confirm every "
           "para-clinical professional providing unsupervised care is identified and privileged. "
           "Documentary evidence is on file for each CORE objective element. Compliance with "
           "the stop-work authority in Section 6 is monitored through the credentialing file "
           "review. This policy itself is reviewed every year.")
    h(doc, 1, "9. Training and staff acknowledgement")
    p(doc, "All staff covered by this policy are made aware of the credentialing and privileging "
           "process, including the stop-work authority in Section 6, and sign the acknowledgement "
           "below.")
    p(doc, f"I have read the Policy on Credentialing and Privileging of Para-Clinical Professionals of {HN}. I will follow the processes described.")
    sig_tbl(doc)
    h(doc, 1, "10. Distribution")
    p(doc, "HR department; all department heads / unit in-charges; all staff through the hospital "
           "intranet and department policy folders.")
    h(doc, 1, "11. Abbreviations")
    abbrev_tbl(doc, HRM_ABBREVS_BASE)
    h(doc, 1, "12. Traceability to NABH HCO Full Accreditation 6th Edition HRM.13")
    p(doc, "This table is an index. It is not how the policy is organised. An asterisk in the "
           "Level column means documentation of the process is required. Stop-work applies ONLY "
           "to the OEs marked 'Section 6 Stop-work' below.")
    trace_tbl(doc, [
        ("HRM.13.a", "CORE",       "Section 3; 5.1; Section 6 Stop-work", _HR),
        ("HRM.13.b", "Commitment", "Section 3; 5.2", _HR),
        ("HRM.13.c", "CORE",       "Section 3; 5.3; Section 6 Stop-work", _HR),
        ("HRM.13.d", "Commitment", "Section 3; 5.4", _HR),
        ("HRM.13.e", "Commitment", "Section 3; 5.5", _HR),
    ])
    h(doc, 1, "13. Required Records / Evidence Checklist")
    h(doc, 2, "HRM.13.a — Permitted para-clinical professionals identified. [Stop-work trigger]")
    lb(doc, "Identified-para-clinical-professional list permitted to provide unsupervised patient care, with a qualification, training and experience verification record.")
    lb(doc, "Cross-reference to the stop-work authority in Section 6 for unlisted practice.")
    h(doc, 2, "HRM.13.b — Credentials verified, documented, updated periodically.")
    lb(doc, "Education, registration, training and experience verification and documentation record, updated periodically, with a personal-file record of each update.")
    lb(doc, "Verification-with-awarding-organisation record.")
    h(doc, 2, "HRM.13.c — Privileges granted in consonance with qualification, training, experience and registration. [Stop-work trigger]")
    lb(doc, "Granted-privilege record naming what each para-clinical professional is authorised to do, with a registration or licence-on-file record where applicable.")
    lb(doc, "Cross-reference to the stop-work authority in Section 6 for privileging outside this record.")
    h(doc, 2, "HRM.13.d — Requisite services known to professional and relevant departments.")
    lb(doc, "Communicated-service record to the professional and to concerned departments, with an internal-communication and confirmation-of-awareness record.")
    h(doc, 2, "HRM.13.e — Para-clinical professionals care as per privileging.")
    lb(doc, "Supervision record for new staff until independent privileges are granted.")
    lb(doc, "Mechanism-confirmation and privileging-compliance spot-check record.")
    h(doc, 1, "14. References")
    lb(doc, "NABH Accreditation Standards for Hospitals, 6th Edition — standard HRM.13.")
    lb(doc, "NABH Guidebook to NABH Accreditation Standards for Hospitals, 6th Edition — HRM.13 interpretations.")
    h(doc, 1, "Disclaimer")
    hrm_disclaimer(doc)
    save_and_verify(doc, "HCO_HRM_13_v2_REWRITE_DRAFT.docx")


# ══════════════════════════════════════════════════════════════════════════════
# Main
# ══════════════════════════════════════════════════════════════════════════════
if __name__ == "__main__":
    gen_hrm1()
    gen_hrm2()
    gen_hrm3()
    gen_hrm4()
    gen_hrm5()
    gen_hrm6()
    gen_hrm7()
    gen_hrm8()
    gen_hrm9()
    gen_hrm10()
    gen_hrm11()
    gen_hrm12()
    gen_hrm13()
    print("\nAll 13 HRM rewrite drafts generated.")
