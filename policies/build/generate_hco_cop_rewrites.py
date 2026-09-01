# -*- coding: utf-8 -*-
"""
generate_hco_cop_rewrites.py
Generates HCO COP chapter v2 rewrite-reference DOCX files.

Pipeline : python-docx, identical to generate_hco_psq_rewrites.py.
Output   : policies/build/rewrite_reference/HCO_COP_N_v2_REWRITE_DRAFT.docx
Source   : ChatGPT first-draft content (approved) + policies/build/cop_raw_dump_*.txt
"""
import os
from docx import Document

HN  = "«Hospital Name»"
OUT = "policies/build/rewrite_reference"
os.makedirs(OUT, exist_ok=True)


# ── Helpers ───────────────────────────────────────────────────────────────────

def h(doc, lv, txt):
    return doc.add_paragraph(txt, style={0: "Title", 1: "Heading 1", 2: "Heading 2"}[lv])

def p(doc, txt=""):
    return doc.add_paragraph(txt, style="Normal")

def ln(doc, txt):
    return doc.add_paragraph(txt, style="List Number")

def lb(doc, txt):
    return doc.add_paragraph(txt, style="List Bullet")

def tbl(doc, rows, cols):
    t = doc.add_table(rows=rows, cols=cols)
    try:
        t.style = "Table Grid"
    except KeyError:
        pass
    return t

def doc_ctrl(doc, no, prep, appr="Medical Superintendent"):
    dc = tbl(doc, 6, 4)
    for ri, (a, b, c, d) in enumerate([
        ("Document No.", no,          "Version",                "2.0"),
        ("Issue No.",    "01",         "Review due",             "One year from implementation"),
        ("Date created", "________",   "Date of implementation", "________"),
    ]):
        dc.cell(ri, 0).text = a; dc.cell(ri, 1).text = b
        dc.cell(ri, 2).text = c; dc.cell(ri, 3).text = d
    for ri, (lbl, txt) in enumerate([
        ("Prepared by", f"{prep}  Name: ________  Signature: ________"),
        ("Reviewed by", "Quality Coordinator  Name: ________  Signature: ________"),
        ("Approved by", f"{appr}  Name: ________  Signature: ________"),
    ], start=3):
        dc.cell(ri, 0).text = lbl
        c1 = dc.cell(ri, 1); c1.text = txt; c1.merge(dc.cell(ri, 3))

def abbrev_tbl(doc, rows):
    t = tbl(doc, len(rows) + 1, 2)
    t.cell(0, 0).text = "Abbreviation"; t.cell(0, 1).text = "Meaning"
    for ri, (a, m) in enumerate(rows, 1):
        t.cell(ri, 0).text = a; t.cell(ri, 1).text = m

def gov_tbl(doc, rows):
    t = tbl(doc, len(rows) + 1, 2)
    t.cell(0, 0).text = "Role"; t.cell(0, 1).text = "Responsibility"
    for ri, (role, resp) in enumerate(rows, 1):
        t.cell(ri, 0).text = role; t.cell(ri, 1).text = resp

def mon_tbl(doc, rows):
    t = tbl(doc, len(rows) + 1, 2)
    t.cell(0, 0).text = "Monitoring area"; t.cell(0, 1).text = "What is monitored"
    for ri, (area, what) in enumerate(rows, 1):
        t.cell(ri, 0).text = area; t.cell(ri, 1).text = what

def sig_tbl(doc):
    sig = tbl(doc, 4, 4)
    for ci, hdr in enumerate(("Staff name", "Designation", "Signature", "Date")):
        sig.cell(0, ci).text = hdr
    for ri in range(1, 4):
        for ci in range(4):
            sig.cell(ri, ci).text = "________"

def save_and_verify(doc, fname):
    import sys
    out = sys.stdout

    def pr(s):
        try:
            out.write(s + "\n")
        except UnicodeEncodeError:
            out.write(s.encode("ascii", "replace").decode() + "\n")

    pr(f"\n=== {fname} ===")
    for i, para in enumerate(doc.paragraphs[:80]):
        sn = para.style.name if para.style else "(None)"
        pr(f"{i:3d}  {sn!r:30s}  {para.text[:60]!r}")
    counts = {}
    for para in doc.paragraphs:
        sn = para.style.name if para.style else "(None)"
        counts[sn] = counts.get(sn, 0) + 1
    pr("  Style inventory:")
    for sn, n in sorted(counts.items()):
        pr(f"    {sn}: {n}")
    pr(f"  Total paras: {len(doc.paragraphs)}")
    path = os.path.join(OUT, fname)
    doc.save(path)
    pr(f"  Saved: {path}")


# ══════════════════════════════════════════════════════════════════════════════
# COP.1 — Uniform Care Guided by Written Guidance   (no stop-work)
# Content: ChatGPT final draft (approved, COP 01.pdf).
# Structure: Document control table, Governance table, Section 12 bullet list.
# COREs: b | Stars: a*, f* | Excellence: d, e
# Exact quantities: "at least two identifiers" (3.2), "two clinical care pathways every year" (3.4)
# Telemedicine conditionality: COP.1.f — conditional on service being provided
# ══════════════════════════════════════════════════════════════════════════════
def gen_cop1():
    doc = Document()

    # Title
    h(doc, 0, "Policy on Uniform Patient Care and Clinical Care Pathways")
    p(doc, HN)

    # Document control
    h(doc, 1, "Document control")
    doc_ctrl(doc, "HCO/COP/POL/01", "Quality Coordinator")
    p(doc, "A blank marked ________ must be completed before issue.")

    # Statement of intent
    h(doc, 1, "Statement of intent")
    p(doc,
      f"{HN} ensures that patients with the same clinical condition and care needs receive "
      "the same quality of health care throughout the organisation. Care is provided uniformly "
      "across settings and ward categories and does not depend on whether the patient is paying "
      "or non-paying or on the source of payment.")
    p(doc,
      "The organisation uses patient identification processes, clinical practice guidance, "
      "clinical care pathways and multidisciplinary care planning to support uniform patient care.")
    p(doc,
      "Where telemedicine services are provided, they are delivered safely and securely based "
      "on written guidance.")

    # 1. Purpose
    h(doc, 1, "1. Purpose")
    p(doc,
      "The purpose of this policy is to establish requirements for uniform patient care, patient "
      "identification, clinical practice guidance, clinical care pathways, multidisciplinary and "
      "multi-speciality care, and telemedicine where provided.")
    p(doc, "Those requirements are covered in the hospital's other policies.")

    # 2. Scope
    h(doc, 1, "2. Scope")
    p(doc, f"This policy applies to outpatient, day-care, inpatient and emergency care provided by {HN}.")
    p(doc,
      "It applies across different categories of wards and settings of care and to patients "
      "irrespective of whether they are paying or non-paying or supported by a governmental or "
      "private insurance scheme.")
    p(doc, "This section applies where the organisation provides telemedicine services.")

    # 3. Policy standards
    h(doc, 1, "3. Policy standards")

    h(doc, 2, "3.1 Uniform patient care")
    p(doc,
      "The organisation needs to implement mechanisms to ensure that care received in outpatient, "
      "day-care, inpatient and emergency settings is uniform.")
    p(doc, "The level of care provided to patients shall be the same throughout the organisation.")
    p(doc,
      "Patients with the same clinical condition and care needs shall receive the same quality "
      "of health care throughout the organisation, irrespective of the setting and category of "
      "ward and whether the patient is paying or non-paying or supported by a governmental or "
      "private insurance scheme or not.")
    p(doc,
      "Care is based on the clinical needs of the patient and not on the class or category of "
      "the patient.")
    p(doc,
      "Where the organisation has separate outpatient departments for different categories of "
      "patients, the methodology for care delivery shall be uniform in all outpatient departments.")

    h(doc, 2, "3.2 Patient identification")
    p(doc, "The mechanism for identification of patients shall be uniform across the organisation.")
    p(doc, "For any care-related aspect, at least two identifiers shall be used.")
    p(doc,
      "One of the identifiers shall be the unique identification number generated at the time "
      "of registration.")
    p(doc,
      "The patient identification process shall be used in imaging, diagnostic and therapeutic "
      "procedures, blood transfusion, day care and nutrition or diet.")

    h(doc, 2, "3.3 Evidence-based clinical practice guidelines and clinical protocols")
    p(doc,
      "The organisation implements evidence-based clinical practice guidelines and/or clinical "
      "protocols to guide uniform patient care.")
    p(doc,
      "Clinical practice guidelines brought out by national and international professional "
      "organisations may be used.")
    p(doc,
      "Standard treatment guidelines brought out by the Government of India are a good starting "
      "point.")
    p(doc,
      "In the absence of evidence-based clinical practice guidelines, or where adapting the "
      "clinical practice guidelines is not feasible, sound clinical practices shall guide the "
      "delivery of care.")
    p(doc,
      "For the definitions of evidence-based medicine and clinical practice guidelines, refer "
      "to the Glossary.")

    h(doc, 2, "3.4 Clinical care pathways")
    p(doc,
      "Clinical care pathways shall be multidisciplinary and based on evidence and/or best "
      "clinical practices.")
    p(doc, "They provide detailed guidance at various stages of care.")
    p(doc, "At a minimum, the organisation shall develop two clinical care pathways every year.")
    p(doc,
      "The clinical care pathways shall be followed consistently across all settings of care "
      "and shall be reviewed and updated annually.")

    h(doc, 2, "3.5 Multidisciplinary and multi-speciality care")
    p(doc,
      "Whenever the patient's clinical condition warrants care from multiple disciplines and/or "
      "multiple specialties, a multidisciplinary and multi-speciality care plan shall be charted "
      "out based on established best clinical practices or clinical practice guidelines for the "
      "specific clinical condition.")
    p(doc,
      "An integrated care plan, including medical professional, nursing, nutritional and supportive "
      "care professionals, shall be developed and implemented appropriately by representatives of "
      "the various disciplines and specialities.")
    p(doc,
      "Examples of situations where multidisciplinary care may be required include care of a "
      "cancer patient determined by a multi-disciplinary tumour board, care of transplant patients, "
      "diabetes patients cared for by multiple specialists to prevent and manage end-organ damage, "
      "palliative or end-of-life care for terminally ill patients, and long-stay patients.")

    h(doc, 2, "3.6 Telemedicine")
    p(doc, "This section applies where the organisation provides telemedicine services.")
    p(doc,
      "Whenever the telemedicine facility is provided, the organisation shall develop written "
      "guidance in consonance with prevailing laws and guidelines and implement the same.")
    p(doc,
      "The written guidance shall include protection of the patient's identity and confidentiality.")
    p(doc,
      "The limitations of information and communication technologies shall be explicitly addressed.")
    p(doc, "The organisation shall have a mechanism for appropriate data storage and retrieval.")
    p(doc, "The organisation shall have an MoU for its telemedicine services if outsourced.")

    # 4. Non-negotiable rules
    h(doc, 1, "4. Non-negotiable rules")
    lb(doc,
       "Do not provide different levels of care to patients with the same clinical condition "
       "and care needs because of payment status, source of payment, ward category or care setting.")
    lb(doc,
       "Do not use different care-delivery methods in separate outpatient departments for "
       "different categories of patients.")
    lb(doc,
       "Do not identify a patient for a care-related aspect using fewer than at least two "
       "identifiers.")
    lb(doc,
       "Do not omit the unique identification number generated at registration from the required "
       "patient identifiers.")
    lb(doc,
       "Do not use a non-uniform patient-identification process across the organisation.")
    lb(doc,
       "Do not leave care without evidence-based clinical practice guidelines or feasible "
       "adaptation of such guidelines without using sound clinical practices to guide delivery "
       "of care.")
    lb(doc, "Do not develop fewer than two clinical care pathways every year.")
    lb(doc, "Do not leave clinical care pathways without annual review and update.")
    lb(doc,
       "Do not omit multidisciplinary and multi-speciality care planning when the patient's "
       "clinical condition warrants care from multiple disciplines or specialties.")
    lb(doc,
       "Where telemedicine services are provided, do not provide them without written guidance "
       "developed in accordance with prevailing laws and guidelines.")
    lb(doc,
       "Where telemedicine services are provided, do not operate without protection of patient "
       "identity and confidentiality, explicit consideration of information and communication "
       "technology limitations, and a mechanism for appropriate data storage and retrieval.")
    lb(doc,
       "Where telemedicine services are outsourced, do not provide the outsourced telemedicine "
       "services without an MoU.")

    # 5. What we do
    h(doc, 1, "5. What we do")

    h(doc, 2, "5.1 Provide uniform care")
    p(doc,
      "Care is provided uniformly across outpatient, day-care, inpatient and emergency settings.")
    p(doc,
      "Patients with the same clinical condition and care needs receive the same quality of "
      "health care throughout the organisation.")
    p(doc,
      "Care is based on clinical need and is not influenced by payment status, source of payment, "
      "ward category or care setting.")

    h(doc, 2, "5.2 Identify patients consistently")
    p(doc, "The same patient-identification mechanism is used throughout the organisation.")
    p(doc,
      "At least two identifiers are used for any care-related aspect, with the registration-"
      "generated unique identification number as one identifier.")
    p(doc,
      "The process is used for imaging, diagnostic and therapeutic procedures, blood transfusion, "
      "day care and nutrition or diet.")

    h(doc, 2, "5.3 Use clinical practice guidance")
    p(doc,
      "Evidence-based clinical practice guidelines and/or clinical protocols guide uniform "
      "patient care.")
    p(doc, "National and international professional organisation guidelines may be used.")
    p(doc, "Government of India standard treatment guidelines are a good starting point.")
    p(doc,
      "Where evidence-based guidelines are absent or cannot feasibly be adapted, sound clinical "
      "practices guide care.")
    p(doc,
      "The Glossary is used for the definitions of evidence-based medicine and clinical practice "
      "guidelines.")

    h(doc, 2, "5.4 Develop and maintain clinical care pathways")
    p(doc,
      "Clinical care pathways are multidisciplinary and based on evidence and/or best clinical "
      "practices.")
    p(doc, "At least two clinical care pathways are developed every year.")
    p(doc,
      "Clinical care pathways are followed consistently across all settings of care and reviewed "
      "and updated annually.")

    h(doc, 2, "5.5 Plan multidisciplinary and multi-speciality care")
    p(doc,
      "When the patient's clinical condition warrants care from multiple disciplines or "
      "specialties, a multidisciplinary and multi-speciality care plan is prepared based on "
      "established best clinical practices or clinical practice guidelines.")
    p(doc,
      "An integrated plan covering medical, nursing, nutritional and supportive care is "
      "developed and implemented appropriately by representatives of the relevant disciplines "
      "and specialities.")

    h(doc, 2, "5.6 Provide telemedicine safely where applicable")
    p(doc,
      "Where the organisation provides telemedicine services, written guidance is developed "
      "and implemented in accordance with prevailing laws and guidelines.")
    p(doc, "Patient identity and confidentiality are protected.")
    p(doc, "Limitations of information and communication technologies are addressed.")
    p(doc, "Appropriate data storage and retrieval are provided.")
    p(doc, "Where telemedicine services are outsourced, an MoU is maintained.")

    # 6. Governance and responsibility — proper table
    h(doc, 1, "6. Governance and responsibility")
    gov_tbl(doc, [
        ("Management",
         "Management ensures mechanisms for uniform patient care and supports implementation "
         "of the requirements in this policy."),
        ("Clinical leadership",
         "Clinical leadership supports consistent care delivery, use of clinical practice "
         "guidance, development of clinical care pathways and multidisciplinary and "
         "multi-speciality care where clinically warranted."),
        ("Clinical and nursing personnel",
         "Clinical and nursing personnel follow the applicable patient identification "
         "processes, clinical guidance, clinical care pathways and multidisciplinary care plans."),
        ("Telemedicine responsibility",
         "Where telemedicine is provided, responsible personnel implement the written guidance "
         "and requirements for patient identity, confidentiality, technology limitations, data "
         "storage and retrieval and outsourced services."),
    ])

    # 7. Quality monitoring — table
    h(doc, 1, "7. Quality monitoring")
    mon_tbl(doc, [
        ("Uniform patient care",
         "Same quality of care for patients with the same clinical condition and care needs "
         "across settings and ward categories"),
        ("Patient identification",
         "Uniform process and use of at least two identifiers, including the registration-"
         "generated unique identification number"),
        ("Clinical guidance",
         "Use of evidence-based clinical practice guidelines and/or clinical protocols, or "
         "sound clinical practices where stated"),
        ("Clinical care pathways",
         "Development of two clinical care pathways every year and annual review and update"),
        ("Multidisciplinary care",
         "Care plans when the patient's clinical condition warrants multiple disciplines or "
         "specialties"),
        ("Telemedicine",
         "Written guidance and required safeguards where telemedicine services are provided"),
    ])

    # 8. Training and staff acknowledgement
    h(doc, 1, "8. Training and staff acknowledgement")
    p(doc,
      "Staff involved in patient care shall be familiar with the requirements for uniform care, "
      "patient identification, clinical practice guidance, clinical care pathways and "
      "multidisciplinary care applicable to their work.")
    p(doc,
      "Personnel providing telemedicine services, where applicable, shall be familiar with the "
      "written telemedicine guidance.")
    p(doc,
      f"I have read the Policy on Uniform Patient Care and Clinical Care Pathways of {HN}. "
      "I will follow the processes described.")
    sig_tbl(doc)

    # 9. Distribution
    h(doc, 1, "9. Distribution")
    p(doc,
      "This policy shall be available to personnel involved in patient care and the "
      "implementation of the processes described in this policy.")
    p(doc,
      "Where telemedicine services are provided, the applicable written telemedicine guidance "
      "shall be available to the personnel responsible for those services.")

    # 10. Abbreviations
    h(doc, 1, "10. Abbreviations")
    abbrev_tbl(doc, [
        ("MoU",   "Memorandum of Understanding"),
        ("ICMR",  "Indian Council of Medical Research"),
        ("MoHFW", "Ministry of Health and Family Welfare"),
    ])

    # 11. Traceability table
    h(doc, 1, "11. Traceability table")
    p(doc,
      "This table is an index. It is not how the policy is organised. An asterisk in the Level "
      "column means documentation of the process is required.")
    tr = tbl(doc, 7, 3)
    for ci, hdr in enumerate(("Objective Element", "Level", "Traceability to this policy")):
        tr.cell(0, ci).text = hdr
    trace_rows = [
        ("COP.1.a", "Commitment*",
         "Sections 3.1 and 5.1 address uniform care irrespective of setting, ward category, "
         "payment status or source of payment."),
        ("COP.1.b", "CORE",
         "Sections 3.2 and 5.2 address the uniform identification process, at least two "
         "identifiers and the registration-generated unique identification number."),
        ("COP.1.c", "Commitment",
         "Sections 3.3 and 5.3 address clinical practice guidelines, clinical protocols, sound "
         "clinical practices where required and reference to the Glossary."),
        ("COP.1.d", "Excellence",
         "Sections 3.4 and 5.4 address multidisciplinary clinical care pathways, two clinical "
         "care pathways every year, consistent use and annual review and update."),
        ("COP.1.e", "Excellence",
         "Sections 3.5 and 5.5 address clinically triggered multidisciplinary and multi-"
         "speciality care planning and integrated care."),
        ("COP.1.f", "Commitment*",
         "Sections 3.6 and 5.6 address telemedicine requirements conditionally where "
         "telemedicine services are provided."),
    ]
    for ri, (oe, lvl, txt) in enumerate(trace_rows, 1):
        tr.cell(ri, 0).text = oe
        tr.cell(ri, 1).text = lvl
        tr.cell(ri, 2).text = txt

    # 12. Required Records/Evidence Checklist — bulleted list
    h(doc, 1, "12. Required Records/Evidence Checklist")

    h(doc, 2, "Uniform patient care")
    lb(doc,
       "Written guidance or mechanisms showing uniform care across outpatient, day-care, "
       "inpatient and emergency settings.")
    lb(doc,
       "Care-delivery processes showing the same quality of health care for patients with "
       "the same clinical condition and care needs.")
    lb(doc,
       "Processes showing that care is not influenced by payment status, source of payment "
       "or ward category.")
    lb(doc,
       "Uniform care-delivery methodology used across separate outpatient departments where "
       "different categories of patients are served.")

    h(doc, 2, "Patient identification")
    lb(doc, "Uniform patient-identification process used across the organisation.")
    lb(doc,
       "Patient identification records showing use of at least two identifiers for care-"
       "related aspects.")
    lb(doc,
       "Patient identification process showing use of the unique identification number "
       "generated at registration as one identifier.")
    lb(doc,
       "Patient identification process applied to imaging, diagnostic and therapeutic "
       "procedures, blood transfusion, day care and nutrition or diet.")

    h(doc, 2, "Clinical practice guidance")
    lb(doc,
       "Evidence-based clinical practice guidelines and/or clinical protocols used to guide "
       "uniform patient care.")
    lb(doc,
       "National or international professional organisation guidelines where the organisation "
       "has chosen to use them.")
    lb(doc,
       "Government of India standard treatment guidelines where used as a starting point.")
    lb(doc,
       "Clinical practices used to guide care where evidence-based clinical practice guidelines "
       "are absent or cannot feasibly be adapted.")
    lb(doc,
       "Access to the Glossary for the definitions of evidence-based medicine and clinical "
       "practice guidelines.")

    h(doc, 2, "Clinical care pathways")
    lb(doc, "At least two clinical care pathways developed every year.")
    lb(doc,
       "Clinical care pathways showing multidisciplinary development and basis in evidence "
       "and/or best clinical practices.")
    lb(doc, "Clinical care pathways showing detailed guidance at various stages of care.")
    lb(doc, "Records showing clinical care pathways are followed across all settings of care.")
    lb(doc, "Records showing clinical care pathways are reviewed and updated annually.")

    h(doc, 2, "Multidisciplinary and multi-speciality care")
    lb(doc,
       "Multidisciplinary and multi-speciality care plans for patients whose clinical condition "
       "warrants care from multiple disciplines or specialties.")
    lb(doc,
       "Integrated care plans covering medical, nursing, nutritional and supportive care where "
       "appropriate.")
    lb(doc,
       "Documentation showing participation by representatives of relevant disciplines and "
       "specialities.")
    lb(doc,
       "Care plans based on established best clinical practices or clinical practice guidelines "
       "for the specific clinical condition.")

    h(doc, 2, "Telemedicine")
    lb(doc,
       "Written telemedicine guidance where the organisation provides telemedicine services.")
    lb(doc, "Telemedicine guidance aligned with prevailing laws and guidelines.")
    lb(doc, "Safeguards protecting patient identity and confidentiality.")
    lb(doc,
       "Written consideration of limitations of information and communication technologies.")
    lb(doc, "Mechanism for appropriate telemedicine data storage and retrieval.")
    lb(doc, "MoU for outsourced telemedicine services where applicable.")

    # 13. References
    h(doc, 1, "13. References")
    ln(doc,
       "National Accreditation Board for Hospitals and Healthcare Providers. NABH Accreditation "
       "Standards for Hospitals, 6th Edition. COP.1.")
    ln(doc, "Guidebook interpretation supplied for COP.1.a through COP.1.f.")
    ln(doc,
       "NABH Glossary for the definitions of evidence-based medicine and clinical practice "
       "guidelines.")
    ln(doc, "Telemedicine Practice Guidelines, ICMR, 2020, MoHFW.")

    # Disclaimer
    h(doc, 1, "Disclaimer")
    p(doc,
      "This policy reorganises the supplied COP.1 objective-element wording and Guidebook "
      "interpretation into plain-language policy format. The modal strength of the source has "
      "been preserved. Optional examples and mechanisms have not been converted into mandatory "
      "requirements. The exact requirements of at least two identifiers, two clinical care "
      "pathways every year, and annual review and update have been retained. Telemedicine "
      "requirements are conditional and apply where the organisation provides telemedicine "
      "services.")

    save_and_verify(doc, "HCO_COP_1_v2_REWRITE_DRAFT.docx")


# ══════════════════════════════════════════════════════════════════════════════
# COP.2 — Emergency Services   (HAS stop-work: Section 6)
# Content: ChatGPT final draft (approved, COP 02.pdf).
# Structure: Document control table, Stop-work Section 6, Governance table,
#            Section 13 bullet checklist, Section 14 References.
# COREs: c | Stars: c*, d*, h*, i* | Stop-work trigger: identified area + triage competence + overcrowding
# Death management: 5 BID elements + 4 post-resus elements = 9 total (all mandatory)
# ══════════════════════════════════════════════════════════════════════════════
def gen_cop2():
    doc = Document()

    # Title
    h(doc, 0, "Policy on Emergency Care Services")
    p(doc, HN)

    # Document control
    h(doc, 1, "Document control")
    doc_ctrl(doc, "HCO/COP/POL/02", "Emergency In-Charge")
    p(doc, "A blank marked ________ must be completed before issue.")

    # Statement of intent
    h(doc, 1, "Statement of intent")
    p(doc,
      f"{HN} provides emergency patients with accessible emergency care, appropriate resources, "
      "written emergency guidance, qualified triage, reassessment, documented disposition, "
      "discharge or transfer information, quality assurance, and defined processes for patients "
      "found dead on arrival or who die within a few minutes of arrival.")

    # 1. Purpose
    h(doc, 1, "1. Purpose")
    p(doc,
      "The purpose of this policy is to establish requirements for receiving and managing "
      "emergency patients, triage, reassessment, disposition, transfer, emergency quality "
      "assurance, and management of patients found dead on arrival or who die within a few "
      "minutes of arrival.")
    p(doc, "Those requirements are covered in the hospital's other policies.")

    # 2. Scope
    h(doc, 1, "2. Scope")
    p(doc,
      f"This policy applies to the emergency department and personnel involved in emergency "
      f"patient care at {HN}.")

    # 3. Policy standards
    h(doc, 1, "3. Policy standards")

    h(doc, 2, "3.1 Emergency area and resources")
    p(doc,
      "The identified area to treat emergency patients shall be easily accessible for initiation "
      "of care.")
    p(doc, "There shall be understandable signage and directions leading to the emergency area.")
    p(doc,
      "Emergency services preferably have designated triage, resuscitation treatment and patient "
      "holding areas.")
    p(doc,
      "The organisation shall specify the minimum number of beds based on its scope of services.")
    p(doc,
      "At a minimum, adequate and appropriate human resources, basic resuscitation equipment, "
      "equipment for monitoring vital parameters, appropriate consumables, and life-saving and "
      "emergency care drugs shall be available.")
    p(doc,
      "Personnel operating the emergency area shall be privileged to work in this area and have "
      "access to ongoing training.")

    h(doc, 2, "3.2 Overcrowding and crowd management")
    p(doc,
      "Prevention of patient overcrowding includes monitoring footfall trends and developing "
      "strategies to manage overcrowding and violence.")
    p(doc,
      "High footfall times are expected to be anticipated, with adequate human resources assigned "
      "and strategies to manage overcrowding put in place.")
    p(doc,
      "Crowd management shall include an appropriate policy for patient relatives, attendants "
      "and visitors.")

    h(doc, 2, "3.3 Emergency care and medico-legal cases")
    p(doc,
      "Written guidance shall include guidelines, SOPs or protocols for general emergency care "
      "and management of specific conditions, for example poisoning, road traffic accidents, "
      "acute stroke and coronary disease.")
    p(doc, "It shall address both adult and paediatric patients.")
    p(doc,
      "The procedure shall incorporate at a minimum identification, assessment and provision "
      "of care.")
    p(doc,
      "Where emergency services are outside the scope of the organisation, or facilities for "
      "appropriate emergency care of a given clinical condition are unavailable, at a minimum "
      "the patient shall be provided first aid before transfer to another organisation.")
    p(doc, "Processes shall be in place to ensure patient safety.")
    p(doc, "The organisation shall define what constitutes a medico-legal case (MLC).")
    p(doc,
      "Care provided, especially documentation and intimation to appropriate authorities, shall "
      "be in accordance with statutory requirements.")
    p(doc,
      "The organisation shall have a policy on management of suspected sexual assault and "
      "guidance on storage of samples of MLC patients.")

    h(doc, 2, "3.4 Triage")
    p(doc, "Triage shall be done only by qualified/trained personnel.")
    p(doc,
      "Written guidance based on evidence or sound clinical practices shall guide triage.")
    p(doc,
      "Triage shall be part of routine day-to-day functioning of the emergency department and "
      "not only disaster management.")
    p(doc,
      "Criteria could be separate for trauma and non-trauma patients and adults and children, "
      "for example PAT or POPS.")
    p(doc,
      "If several clients are waiting to be triaged, a visual triage assessment may be "
      "conducted.")

    h(doc, 2, "3.5 Reassessment")
    p(doc,
      "Patients waiting in the emergency are reassessed as appropriate for a change in status.")
    p(doc,
      "The findings of reassessment shall be documented in the patient's medical record.")

    h(doc, 2, "3.6 Admission, discharge and transfer")
    p(doc,
      "The organisation shall maintain documentation indicating whether an emergency patient "
      "was sent home after initial care, admitted for further care, admitted in an emergency "
      "for a short stay and then discharged, or transferred to another organisation.")
    p(doc,
      "Staff shall have a clear understanding of the scope of the organisation's activities "
      "and the procedure for referral and transfer to another appropriate centre for patients "
      "who cannot be cared for in-house, after due first-aid or emergency care.")

    h(doc, 2, "3.7 Discharge and transfer note")
    p(doc,
      "In case of discharge to home or transfer to another organisation, a discharge or "
      "transfer note shall be given to the patient.")
    p(doc,
      "The note shall contain salient clinical findings, investigations done, treatment given "
      "and condition at discharge or transfer.")
    p(doc, "The reasons for discharge or transfer shall be documented.")

    h(doc, 2, "3.8 Quality assurance")
    p(doc, "The organisation shall implement a quality assurance programme.")
    p(doc,
      "Written quality assurance guidance may be developed individually or could be part of "
      "the organisation's overall quality-improvement programme.")
    p(doc,
      "The quality assurance programme shall involve all aspects of the functioning of the "
      "Emergency department.")
    p(doc,
      "The Emergency department shall collect data on key performance indicators, including "
      "care outcomes, as part of its quality improvement programme.")
    p(doc,
      "An example of a care outcome measure is return to the emergency department for the "
      "same complaint.")
    p(doc,
      "Collected data shall be collated, analysed and used for further improvements.")
    p(doc, "Improvements shall be monitored for sustenance.")

    h(doc, 2, "3.9 Patients found dead on arrival")
    p(doc,
      "There shall be written guidance for managing a patient found dead on arrival, conforming "
      "to relevant local laws.")
    p(doc,
      "The guidance shall address maintaining a logbook of patients found dead on arrival; "
      "the decision on whether to perform a post-mortem; the decision regarding issue of a "
      "medical certificate of cause of death; temporary storage of the body in appropriate "
      "conditions; and what to do in case of unclaimed or unaccompanied bodies.")

    h(doc, 2, "3.10 Patients who die within a few minutes of arrival")
    p(doc,
      "There shall be written guidance for managing a patient who dies within a few minutes "
      "of arrival after a failed attempt at resuscitation, conforming to relevant local laws.")
    p(doc,
      "The guidance shall address registration of such patients and recording the entire "
      "resuscitation events; the decision on whether to perform a post-mortem; temporary "
      "storage of the body in appropriate conditions; issue of a medical certificate of cause "
      "of death; and handing over of the body.")

    # 4. Non-negotiable rules
    h(doc, 1, "4. Non-negotiable rules")
    lb(doc,
       "Do not receive and treat an emergency patient in an area that is not the identified "
       "emergency area, or when triage-competent personnel and basic resuscitation equipment "
       "are not available.")
    lb(doc,
       "Do not leave overcrowding unmanaged when it blocks triage, resuscitation or safe "
       "holding.")
    lb(doc,
       "Do not provide emergency care without written guidance incorporating at minimum "
       "identification, assessment and provision of care.")
    lb(doc,
       "Do not transfer a patient without first providing first aid when appropriate emergency "
       "care is outside the organisation's scope or unavailable for the given clinical condition.")
    lb(doc,
       "Do not perform triage through personnel who are not qualified/trained.")
    lb(doc,
       "Do not treat triage as only a disaster-management activity; it shall be part of "
       "routine day-to-day emergency functioning.")
    lb(doc,
       "Do not omit documentation of reassessment findings in the patient's medical record.")
    lb(doc,
       "Do not discharge a patient home or transfer a patient without giving the patient the "
       "required discharge or transfer note.")
    lb(doc,
       "Do not omit any of the required contents of the discharge or transfer note or the "
       "reasons for discharge or transfer.")
    lb(doc,
       "Do not omit the required management processes for either a patient found dead on "
       "arrival or a patient who dies within a few minutes after arrival following failed "
       "resuscitation.")

    # 5. What we do
    h(doc, 1, "5. What we do")

    h(doc, 2, "5.1 Emergency intake and resources")
    p(doc,
      "Maintain an accessible identified emergency area with understandable signage and "
      "directions.")
    p(doc, "Specify the minimum number of beds based on scope of services.")
    p(doc,
      "Maintain adequate and appropriate human resources, basic resuscitation equipment, "
      "vital-parameter monitoring equipment, consumables and life-saving/emergency care drugs.")
    p(doc,
      "Ensure emergency personnel are privileged to work in the area and have access to "
      "ongoing training.")

    h(doc, 2, "5.2 Crowd management")
    p(doc,
      "Monitor footfall trends and develop strategies for overcrowding and violence.")
    p(doc,
      "Anticipate high footfall times, assign adequate human resources and put strategies "
      "in place.")
    p(doc,
      "Maintain an appropriate policy for patient relatives, attendants and visitors.")

    h(doc, 2, "5.3 Emergency care")
    p(doc,
      "Use written guidelines, SOPs or protocols for general emergency care and specific "
      "conditions.")
    p(doc, "Cover adult and paediatric patients.")
    p(doc,
      "Use identification, assessment and provision of care as the minimum procedure elements.")
    p(doc,
      "Provide first aid before transfer when appropriate emergency care is outside scope or "
      "unavailable.")

    h(doc, 2, "5.4 Triage")
    p(doc, "Ensure triage is performed only by qualified/trained personnel.")
    p(doc, "Use written guidance based on evidence or sound clinical practices.")
    p(doc, "Keep triage as part of routine day-to-day emergency functioning.")
    p(doc,
      "Separate criteria or visual triage assessment may be used as described in the Guidebook.")

    h(doc, 2, "5.5 Reassessment")
    p(doc,
      "Reassess patients waiting in emergency as appropriate for a change in status.")
    p(doc, "Document reassessment findings in the patient's medical record.")

    h(doc, 2, "5.6 Disposition and transfer")
    p(doc,
      "Document whether patients are sent home, admitted, kept for a short emergency stay "
      "and discharged, or transferred.")
    p(doc,
      "Ensure staff understand the scope of activities and referral and transfer procedures.")

    h(doc, 2, "5.7 Discharge and transfer note")
    p(doc,
      "Give the patient a discharge or transfer note when discharged home or transferred.")
    p(doc,
      "Include salient clinical findings, investigations done, treatment given, condition at "
      "discharge or transfer, and reasons for discharge or transfer.")

    h(doc, 2, "5.8 Quality assurance")
    p(doc,
      "Implement a quality assurance programme covering all aspects of Emergency department "
      "functioning.")
    p(doc,
      "Collect, collate and analyse key performance indicator data, including care outcomes.")
    p(doc,
      "Use collected data for further improvements and monitor improvements for sustenance.")
    p(doc,
      "The quality assurance programme may be standalone or could be part of the overall "
      "quality-improvement programme.")

    h(doc, 2, "5.9 Patients found dead on arrival")
    p(doc, "Use written guidance conforming to relevant local laws.")
    p(doc,
      "Address the logbook, post-mortem decision, medical certificate of cause of death "
      "decision, temporary body storage and unclaimed/unaccompanied bodies.")

    h(doc, 2, "5.10 Patients who die within a few minutes of arrival")
    p(doc, "Use written guidance conforming to relevant local laws.")
    p(doc,
      "Address registration and complete recording of resuscitation events, post-mortem "
      "decision, temporary body storage, medical certificate of cause of death and handing "
      "over of the body.")

    # 6. Stop-work authority
    h(doc, 1, "6. Stop-work authority")
    p(doc,
      "Do not receive and treat an emergency patient in an area that is not the identified "
      "emergency area, or when triage-competent personnel and basic resuscitation equipment "
      "are not available. Do not leave overcrowding unmanaged when it blocks triage, "
      "resuscitation or safe holding. Stop-work applies to unsafe emergency intake, not to "
      "immediate life-saving measures already under way — those continue while escalation "
      "happens. The person who stops tells the Emergency In-Charge and the Medical "
      "Superintendent the same shift. Refusing unsafe emergency intake is not a disciplinary "
      "matter.")

    # 7. Governance and responsibility — proper table
    h(doc, 1, "7. Governance and responsibility")
    gov_tbl(doc, [
        ("Management",
         "Management ensures the identified emergency area, required resources, written "
         "guidance and quality assurance arrangements."),
        ("Emergency In-Charge",
         "The Emergency In-Charge oversees emergency operations and escalation of unsafe "
         "emergency intake."),
        ("Qualified/trained personnel",
         "Qualified/trained personnel perform triage."),
        ("Emergency personnel",
         "Emergency personnel follow applicable emergency, reassessment, referral, transfer "
         "and documentation processes."),
        ("Quality personnel",
         "Quality personnel support quality assurance data collection, collation, analysis "
         "and improvement monitoring."),
    ])

    # 8. Quality monitoring — proper table
    h(doc, 1, "8. Quality monitoring")
    mon_tbl(doc, [
        ("Emergency footfall and overcrowding",
         "Emergency department footfall and overcrowding trends"),
        ("Emergency care KPIs",
         "Emergency care key performance indicators, including care outcomes"),
        ("Quality data",
         "Quality data collation and analysis"),
        ("Improvements",
         "Further improvements and monitoring for sustenance"),
        ("Reassessment and disposition",
         "Reassessment documentation and patient disposition documentation"),
        ("Discharge and transfer note",
         "Discharge and transfer note requirements"),
    ])

    # 9. Training and staff acknowledgement
    h(doc, 1, "9. Training and staff acknowledgement")
    p(doc,
      "Personnel operating the emergency area shall have access to ongoing training. Triage "
      "shall be performed only by qualified/trained personnel.")
    p(doc,
      "Staff shall have a clear understanding of the scope of the organisation's activities "
      "and applicable referral and transfer procedures.")
    p(doc,
      f"I have read the Policy on Emergency Care Services of {HN}. I will follow the "
      "processes described.")
    sig_tbl(doc)

    # 10. Distribution
    h(doc, 1, "10. Distribution")
    p(doc,
      "This policy shall be available to personnel involved in emergency care and emergency "
      "department operations.")

    # 11. Abbreviations
    h(doc, 1, "11. Abbreviations")
    abbrev_tbl(doc, [
        ("MLC",  "Medico-legal case"),
        ("SOP",  "Standard Operating Procedure"),
        ("PAT",  "Paediatric Assessment Triangle"),
        ("POPS", "Paediatric Observation Priority Score"),
    ])

    # 12. Traceability table
    h(doc, 1, "12. Traceability table")
    p(doc,
      "This table is an index. It is not how the policy is organised. An asterisk in the Level "
      "column means documentation of the process is required.")
    tr = tbl(doc, 10, 3)
    for ci, hdr in enumerate(("Objective Element", "Level", "Traceability to this policy")):
        tr.cell(0, ci).text = hdr
    trace_rows = [
        ("COP.2.a", "Commitment",
         "Sections 3.1 and 5.1 address the identified accessible emergency area, signage and "
         "directions, minimum beds, resources, privileged personnel and ongoing training."),
        ("COP.2.b", "Achievement",
         "Sections 3.2 and 5.2 address prevention of overcrowding and crowd management, "
         "including the required policy for patient relatives, attendants and visitors."),
        ("COP.2.c", "CORE*",
         "Sections 3.3 and 5.3 address emergency care in accordance with statutory requirements "
         "and written guidance, including minimum identification, assessment and provision of care, "
         "MLC definition, statutory documentation, sexual assault policy and MLC sample storage."),
        ("COP.2.d", "Commitment*",
         "Sections 3.4 and 5.4 address triage by qualified/trained personnel using written "
         "guidance as routine emergency functioning."),
        ("COP.2.e", "Commitment",
         "Sections 3.5 and 5.5 address appropriate reassessment for change in status and "
         "documentation of findings."),
        ("COP.2.f", "Commitment",
         "Sections 3.6 and 5.6 address documentation of admission, discharge and transfer and "
         "the clear referral/transfer process."),
        ("COP.2.g", "Commitment",
         "Sections 3.7 and 5.7 address the discharge/transfer note with salient clinical "
         "findings, investigations done, treatment given, condition at discharge/transfer and "
         "reasons."),
        ("COP.2.h", "Achievement*",
         "Sections 3.8 and 5.8 address the quality assurance programme, Emergency department "
         "coverage, KPI data, analysis, improvement and sustenance monitoring."),
        ("COP.2.i", "Commitment*",
         "Sections 3.9, 3.10, 5.9 and 5.10 address systems for management of patients found "
         "dead on arrival (five mandatory elements) and patients who die within a few minutes "
         "of arrival after failed resuscitation (four mandatory elements)."),
    ]
    for ri, (oe, lvl, txt) in enumerate(trace_rows, 1):
        tr.cell(ri, 0).text = oe
        tr.cell(ri, 1).text = lvl
        tr.cell(ri, 2).text = txt

    # 13. Required Records/Evidence Checklist — bulleted list
    h(doc, 1, "13. Required Records/Evidence Checklist")

    h(doc, 2, "Emergency area and resources")
    lb(doc, "Identified emergency area with signage and understandable directions.")
    lb(doc, "Minimum number of beds specified based on scope of services.")
    lb(doc,
       "Human resources, basic resuscitation equipment, vital-parameter monitoring equipment, "
       "consumables and life-saving/emergency drugs available.")
    lb(doc, "Emergency personnel privileging and ongoing training arrangements.")

    h(doc, 2, "Crowd management")
    lb(doc, "Footfall trend information and overcrowding/violence management strategies.")
    lb(doc, "Planning for high-footfall times and allocation of adequate human resources.")
    lb(doc, "Policy for patient relatives, attendants and visitors.")

    h(doc, 2, "Emergency care and MLC")
    lb(doc,
       "Written emergency guidelines, SOPs or protocols for general care and specific "
       "conditions.")
    lb(doc, "Adult and paediatric emergency guidance.")
    lb(doc, "Emergency procedure covering identification, assessment and provision of care.")
    lb(doc,
       "First-aid and transfer arrangements where appropriate emergency care is outside "
       "scope or unavailable.")
    lb(doc, "Patient-safety processes.")
    lb(doc, "Definition of MLC and statutory documentation/intimation process.")
    lb(doc,
       "Policy for suspected sexual assault management and guidance for storage of MLC "
       "samples.")

    h(doc, 2, "Triage and reassessment")
    lb(doc, "Qualification/training arrangements for personnel performing triage.")
    lb(doc, "Written triage guidance based on evidence or sound clinical practices.")
    lb(doc, "Routine day-to-day triage process.")
    lb(doc, "Medical records containing documented reassessment findings.")

    h(doc, 2, "Disposition and transfer")
    lb(doc,
       "Emergency records showing whether patients were sent home, admitted, kept for short "
       "emergency stay and discharged, or transferred.")
    lb(doc,
       "Referral and transfer procedure for patients who cannot be cared for in-house after "
       "due first-aid/emergency care.")
    lb(doc,
       "Discharge/transfer notes containing salient clinical findings, investigations done, "
       "treatment given, condition at discharge/transfer and reasons for discharge/transfer.")

    h(doc, 2, "Quality assurance")
    lb(doc,
       "Written quality assurance guidance, standalone or part of the overall "
       "quality-improvement programme.")
    lb(doc, "Emergency department quality assurance covering all aspects of functioning.")
    lb(doc,
       "Key performance indicator and care-outcome data with collation and analysis.")
    lb(doc, "Improvement actions and monitoring for sustenance.")

    h(doc, 2, "Found dead on arrival")
    lb(doc, "Written guidance conforming to relevant local laws.")
    lb(doc, "Logbook for patients found dead on arrival.")
    lb(doc,
       "Post-mortem decision and medical certificate of cause of death decision processes.")
    lb(doc, "Temporary body storage arrangements.")
    lb(doc, "Process for unclaimed or unaccompanied bodies.")

    h(doc, 2, "Death within a few minutes of arrival")
    lb(doc, "Written guidance conforming to relevant local laws.")
    lb(doc, "Registration process and complete recording of resuscitation events.")
    lb(doc, "Post-mortem decision process.")
    lb(doc, "Temporary body storage arrangements.")
    lb(doc, "Medical certificate of cause of death process and body-handover process.")

    # 14. References
    h(doc, 1, "14. References")
    ln(doc,
       "National Accreditation Board for Hospitals and Healthcare Providers. NABH Accreditation "
       "Standards for Hospitals, 6th Edition. COP.2.")
    ln(doc, "Guidebook interpretation supplied for COP.2.a through COP.2.i.")
    ln(doc,
       "Relevant local laws for management of patients found dead on arrival or who die within "
       "a few minutes of arrival.")

    # Disclaimer
    h(doc, 1, "Disclaimer")
    p(doc,
      "This policy is a plain-language reorganisation of the supplied COP.2 objective elements "
      "and Guidebook interpretation. Mandatory requirements, stated examples, preferable measures "
      "and optional mechanisms have been kept at their stated level of strength.")

    save_and_verify(doc, "HCO_COP_2_v2_REWRITE_DRAFT.docx")


# ══════════════════════════════════════════════════════════════════════════════
# ══════════════════════════════════════════════════════════════════════════════
# COP.3 — Ambulance Services and Safe Transportation   (HAS stop-work: Section 6)
# Content: ChatGPT final draft (approved, COP%2003.pdf).
# Structure: Document control table, Stop-work Section 6, Governance table,
#            Section 13 bullet checklist, Section 14 References.
# COREs: none | Stars: f* | Stop-work: dispatch triggers only (NOT overcrowding)
# FIX: The sentence "Do not leave overcrowding unmanaged when it blocks triage,
#      resuscitation or safe holding." belongs to COP.2 and is excluded here.
# ══════════════════════════════════════════════════════════════════════════════
def gen_cop3():
    doc = Document()

    # Title
    h(doc, 0, "Policy on Ambulance Services")
    p(doc, HN)

    # Document control
    h(doc, 1, "Document control")
    doc_ctrl(doc, "HCO/COP/POL/03", "Ambulance/Transport In-Charge")
    p(doc, "A blank marked ________ must be completed before issue.")

    # Statement of intent
    h(doc, 1, "Statement of intent")
    p(doc,
      f"{HN} shall have access to ambulance services commensurate with its scope of services. "
      "Ambulance services shall support safe patient transport with appropriate care, adequate "
      "access and space, appropriate equipment and trained personnel, reliable communication, "
      "and safe transfer of clinical information.")

    # 1. Purpose
    h(doc, 1, "1. Purpose")
    p(doc,
      "The purpose of this policy is to establish requirements for ambulance access, space, "
      "statutory fitness, equipment, personnel, daily and post-trip checks, communication and "
      "support for patients while in transit.")
    p(doc, "Those requirements are covered in the hospital's other policies.")

    # 2. Scope
    h(doc, 1, "2. Scope")
    p(doc,
      f"This policy applies to ambulance services used by {HN}, whether provided in-house or "
      "through an outsourced service.")

    # 3. Policy standards
    h(doc, 1, "3. Policy standards")

    h(doc, 2, "3.1 Ambulance access and level")
    p(doc,
      "The organisation may provide an in-house or outsourced ambulance service for safe "
      "patient transport with appropriate care.")
    p(doc,
      "The organisation shall decide the appropriate level of ambulance based on the National "
      "Ambulance Code AIS-125.")

    h(doc, 2, "3.2 Ambulance space and access")
    p(doc,
      "The organisation shall demarcate a proper space for ambulances, keeping easy "
      "accessibility for receiving patients and enabling ambulances to exit quickly.")
    p(doc,
      "Adequate and prominent signage shall guide ambulance drivers to the ambulance entry "
      "and route to the emergency department.")
    p(doc, "It is preferable that the organisation has an ambulance parking bay.")

    h(doc, 2, "3.3 Fitness and equipment")
    p(doc,
      "The vehicle used as an ambulance shall adhere to applicable statutory requirements. "
      "Examples include registration as an ambulance under the Motor Vehicle Act, a valid "
      "fitness certificate, pollution control certificate and insurance. These are examples "
      "and not an exhaustive list of statutory requirements.")
    p(doc,
      "An ambulance shall have at least basic life support equipment for both adult and "
      "paediatric patients.")
    p(doc,
      "Based on the organisation's scope, additional equipment, for example monitoring and "
      "resuscitative equipment, shall be available in the ambulance.")
    p(doc, "Reference: National Ambulance Code AIS-125.")

    h(doc, 2, "3.4 Personnel")
    p(doc, "The ambulance shall be operated by a driver with a valid driving licence.")
    p(doc,
      "Personnel in the ambulance shall have training in basic life support and basic "
      "cardiopulmonary resuscitation.")
    p(doc,
      "A technician, nurse and/or doctor may be present depending on the situation and "
      "scope of the ambulance.")

    h(doc, 2, "3.5 Daily and post-trip checks")
    p(doc,
      "The daily check shall indicate the functioning status of ambulance systems like "
      "lights, siren, beacon lights, fuel and tyres.")
    p(doc, "Medical equipment shall be checked using a documented checklist.")
    p(doc,
      "Emergency medications shall be available in the ambulance during patient transport.")
    p(doc,
      "A daily check shall ensure availability and expiry dates of emergency medications "
      "and shall be documented.")
    p(doc,
      "After every trip, medications used shall be topped up if used and the same verified "
      "and documented.")
    p(doc,
      "Medications shall be stored in a safe environment as per organisational policy.")
    p(doc,
      "If the organisation follows a system of sealing the emergency medication kit, the "
      "check shall be carried out after each use of the kit.")

    h(doc, 2, "3.6 Communication")
    p(doc,
      "The ambulance shall be connected with the organisation/control room by "
      "wireless/mobile phones.")
    p(doc,
      "The communication system shall encompass the whole process of patient transport.")
    p(doc,
      "There shall be written guidance on how a call for patient transport is received, "
      "who is expected to respond and who organises the transport.")
    p(doc,
      "Communication shall ensure that the ambulance leaves the hospital within a "
      "predefined timeframe based upon the patient's needs.")

    h(doc, 2, "3.7 Care while in transit")
    p(doc,
      "From first communication with the patient or attendant, a file is created to record "
      "appropriate information.")
    p(doc,
      "Attempts are made to gather important clinical information such as age, weight, "
      "provisional diagnosis and ongoing treatment at the referring organisation.")
    p(doc,
      "This information is used by ambulance personnel of the receiving organisation to be "
      "better prepared to assess, initiate emergency care/interventions during transit and "
      "transport the patient safely.")
    p(doc,
      "During transit, when required, there is an exchange of information between ambulance "
      "personnel and the medical professional at the receiving organisation.")
    p(doc,
      "When the patient is shifted by an external agency, wherever possible, the doctor of "
      "the receiving organisation attempts to communicate with the ambulance personnel of the "
      "external agency to ascertain the clinical situation and make appropriate suggestions.")
    p(doc,
      "The medical professional in the ambulance is responsible for decision-making regarding "
      "care/interventions during transit.")

    # 4. Non-negotiable rules
    h(doc, 1, "4. Non-negotiable rules")
    lb(doc,
       "Do not dispatch or continue an ambulance transfer when the vehicle fails its readiness "
       "check, required equipment is non-functional, the driver is not licensed for the vehicle "
       "class, or the organisation/control-room communication link is down.")
    lb(doc,
       "Do not operate an ambulance without at least basic life support equipment for both "
       "adult and paediatric patients.")
    lb(doc,
       "Do not operate an ambulance without a driver with a valid driving licence.")
    lb(doc,
       "Do not allow ambulance personnel to provide service without basic life support and "
       "basic cardiopulmonary resuscitation training.")
    lb(doc,
       "Do not omit the daily check of vehicle systems, medical equipment and emergency "
       "medication availability and expiry, or the required documentation.")
    lb(doc,
       "Do not omit the post-trip top-up of medications used, verification and documentation.")
    lb(doc,
       "Where a sealed emergency medication kit is used, do not omit the check after each use.")
    lb(doc,
       "Do not operate the ambulance without the required organisation/control-room "
       "communication link, whole-process communication coverage and written call-receipt/"
       "response guidance.")
    lb(doc,
       "Do not dispatch an ambulance outside the predefined timeframe based on the patient's "
       "needs.")

    # 5. What we do
    h(doc, 1, "5. What we do")

    h(doc, 2, "5.1 Access and ambulance level")
    p(doc,
      "Provide in-house or outsourced ambulance service as decided by the organisation.")
    p(doc,
      "Select the appropriate ambulance level based on National Ambulance Code AIS-125.")

    h(doc, 2, "5.2 Space and access")
    p(doc, "Demarcate ambulance space for easy patient receipt and quick exit.")
    p(doc,
      "Maintain prominent signage to the ambulance entry and emergency department route.")
    p(doc, "An ambulance parking bay may be provided.")

    h(doc, 2, "5.3 Vehicle fitness and equipment")
    p(doc, "Maintain statutory compliance applicable to the ambulance.")
    p(doc, "Provide at least basic life support equipment for adult and paediatric patients.")
    p(doc, "Provide additional monitoring and resuscitative equipment based on scope.")

    h(doc, 2, "5.4 Personnel")
    p(doc, "Use a driver with a valid driving licence.")
    p(doc,
      "Ensure ambulance personnel have basic life support and basic cardiopulmonary "
      "resuscitation training.")
    p(doc,
      "A technician, nurse and/or doctor may be included depending on situation and scope.")

    h(doc, 2, "5.5 Checks and medication")
    p(doc,
      "Perform and document the daily vehicle, medical equipment and medication checks.")
    p(doc, "After every trip, top up medications used, verify and document this.")
    p(doc, "Where a sealed kit system is used, check the kit after each use.")
    p(doc, "Store medications safely according to organisational policy.")

    h(doc, 2, "5.6 Communication")
    p(doc,
      "Maintain connection with the organisation/control room by wireless or mobile phones.")
    p(doc, "Ensure communication covers the whole transport process.")
    p(doc, "Follow written call-receipt and response guidance.")
    p(doc,
      "Ensure ambulance departure within the predefined timeframe based on patient needs.")

    h(doc, 2, "5.7 In-transit care")
    p(doc,
      "Create a file from first communication and gather appropriate clinical information.")
    p(doc,
      "Use available information to prepare for assessment, emergency care/interventions "
      "and safe transport.")
    p(doc,
      "Exchange information with the receiving medical professional when required.")
    p(doc,
      "Where an external agency transports the patient, wherever possible, the receiving "
      "doctor attempts communication with its ambulance personnel.")
    p(doc,
      "The medical professional in the ambulance makes decisions regarding "
      "care/interventions during transit.")

    # 6. Stop-work authority
    # NOTE: "Do not leave overcrowding unmanaged when it blocks triage, resuscitation or
    # safe holding." is a COP.2 sentence that leaked into earlier drafts — it is excluded here.
    h(doc, 1, "6. Stop-work authority")
    p(doc,
      "Do not dispatch or continue an ambulance transfer when the vehicle fails its readiness "
      "check, required equipment is non-functional, the driver is not licensed for the vehicle "
      "class, or the organisation/control-room communication link is down. Stop-work applies "
      "to the transfer dispatch, not to on-scene life-saving measures. The person who stops "
      "tells the Ambulance / Transport In-Charge and the Medical Superintendent the same shift. "
      "Refusing an unsafe ambulance dispatch is not a disciplinary matter.")

    # 7. Governance and responsibility — proper table
    h(doc, 1, "7. Governance and responsibility")
    gov_tbl(doc, [
        ("Ambulance/Transport In-Charge",
         "Ambulance/Transport In-Charge oversees ambulance readiness, dispatch, communication "
         "and transport processes."),
        ("Management",
         "Management determines the appropriate ambulance level and whether services are "
         "provided in-house or outsourced."),
        ("Ambulance personnel",
         "Ambulance personnel maintain required checks, equipment, medications and "
         "communication."),
        ("Medical professional in ambulance",
         "The medical professional in the ambulance is responsible for decision-making "
         "regarding care/interventions during transit."),
    ])

    # 8. Quality monitoring — proper table
    h(doc, 1, "8. Quality monitoring")
    mon_tbl(doc, [
        ("Daily readiness checks",
         "Daily ambulance readiness checks and documented findings"),
        ("Post-trip checks",
         "Post-trip medication top-up, verification and documentation"),
        ("Equipment and medication",
         "Equipment functionality and medication availability/expiry"),
        ("Communication",
         "Communication-system availability and predefined response timeframe"),
        ("Access and space",
         "Ambulance access, space and signage"),
        ("Statutory compliance",
         "Compliance with applicable statutory requirements"),
    ])

    # 9. Training and staff acknowledgement
    h(doc, 1, "9. Training and staff acknowledgement")
    p(doc,
      "Ambulance personnel shall have basic life support and basic cardiopulmonary "
      "resuscitation training. The ambulance driver shall have a valid driving licence.")
    p(doc,
      f"I have read the Policy on Ambulance Services of {HN}. I will follow the processes "
      "described.")
    sig_tbl(doc)

    # 10. Distribution
    h(doc, 1, "10. Distribution")
    p(doc,
      "This policy shall be available to ambulance/transport personnel and personnel "
      "responsible for emergency patient transport.")

    # 11. Abbreviations
    h(doc, 1, "11. Abbreviations")
    abbrev_tbl(doc, [
        ("AIS", "Automotive Industry Standard"),
        ("CPR", "Cardiopulmonary Resuscitation"),
    ])

    # 12. Traceability table
    h(doc, 1, "12. Traceability table")
    p(doc,
      "This table is an index. It is not how the policy is organised. An asterisk in the Level "
      "column means documentation of the process is required.")
    tr = tbl(doc, 8, 3)
    for ci, hdr in enumerate(("Objective Element", "Level", "Traceability to this policy")):
        tr.cell(0, ci).text = hdr
    trace_rows = [
        ("COP.3.a", "Commitment",
         "Sections 3.1 and 5.1 address access to ambulance services commensurate with scope "
         "and ambulance level selection based on National Ambulance Code AIS-125."),
        ("COP.3.b", "Commitment",
         "Sections 3.2 and 5.2 address adequate ambulance access and space, quick exit and "
         "prominent signage."),
        ("COP.3.c", "Commitment",
         "Sections 3.3 and 5.3 address statutory compliance, at least basic life support "
         "equipment for adult and paediatric patients, and scope-based additional equipment."),
        ("COP.3.d", "Commitment",
         "Sections 3.4 and 5.4 address valid driving licence and basic life support/CPR "
         "training; additional crew may be present depending on situation and scope."),
        ("COP.3.e", "Commitment",
         "Sections 3.5 and 5.5 address daily vehicle/equipment/medication checks, post-trip "
         "medication top-up, verification and documentation, and sealed-kit checks where "
         "applicable."),
        ("COP.3.f", "Commitment*",
         "Sections 3.6 and 5.6 address the proper communication system covering the whole "
         "patient-transport process, written call/response guidance and predefined timeframe "
         "based on patient needs."),
        ("COP.3.g", "Achievement",
         "Sections 3.7 and 5.7 address early treatment opportunities while in transit, "
         "clinical information exchange and responsibility of the medical professional in "
         "the ambulance for transit care decisions."),
    ]
    for ri, (oe, lvl, txt) in enumerate(trace_rows, 1):
        tr.cell(ri, 0).text = oe
        tr.cell(ri, 1).text = lvl
        tr.cell(ri, 2).text = txt

    # 13. Required Records/Evidence Checklist — bulleted list
    h(doc, 1, "13. Required Records/Evidence Checklist")

    h(doc, 2, "Ambulance access and statutory compliance")
    lb(doc,
       "Demarcated ambulance space with accessible patient-receiving area and quick exit.")
    lb(doc, "Prominent signage to ambulance entry and emergency department route.")
    lb(doc,
       "Applicable ambulance statutory documents, such as registration, fitness certificate, "
       "pollution control certificate and insurance.")
    lb(doc,
       "Documentation showing the ambulance level selected with reference to National "
       "Ambulance Code AIS-125.")

    h(doc, 2, "Equipment and personnel")
    lb(doc, "Basic life support equipment for both adult and paediatric patients.")
    lb(doc, "Scope-based additional monitoring and resuscitative equipment.")
    lb(doc, "Valid driving licence for the ambulance driver.")
    lb(doc, "Basic life support and CPR training records for ambulance personnel.")

    h(doc, 2, "Daily and post-trip checks")
    lb(doc,
       "Documented daily vehicle-system checks covering lights, siren, beacon lights, fuel "
       "and tyres.")
    lb(doc, "Documented medical-equipment checks.")
    lb(doc, "Daily emergency-medication availability and expiry checks.")
    lb(doc, "Post-trip medication top-up, verification and documentation.")
    lb(doc,
       "Sealed emergency medication-kit checks after each use where that system is used.")
    lb(doc, "Medication storage arrangements according to organisational policy.")

    h(doc, 2, "Communication and transport")
    lb(doc,
       "Connection between ambulance and organisation/control room by wireless or mobile "
       "phones.")
    lb(doc, "Written call-receipt, response and transport-organisation guidance.")
    lb(doc, "Communication coverage for the whole patient-transport process.")
    lb(doc, "Predefined timeframe for ambulance departure based on patient needs.")

    h(doc, 2, "In-transit care")
    lb(doc, "File created from first communication with the patient or attendant.")
    lb(doc,
       "Available clinical information such as age, weight, provisional diagnosis and "
       "ongoing treatment at the referring organisation.")
    lb(doc,
       "Information exchange between ambulance personnel and receiving medical professional "
       "when required.")
    lb(doc,
       "Attempts, wherever possible, by the receiving doctor to communicate with an external "
       "ambulance agency when applicable.")
    lb(doc,
       "Documentation identifying the medical professional in the ambulance responsible for "
       "care/intervention decisions during transit.")

    # 14. References
    h(doc, 1, "14. References")
    ln(doc, "National Ambulance Code AIS-125.")
    ln(doc,
       "Motor Vehicle Act and applicable statutory requirements for ambulance vehicles.")
    ln(doc, "Guidebook interpretation supplied for COP.3.a through COP.3.g.")

    # Disclaimer
    h(doc, 1, "Disclaimer")
    p(doc,
      "This policy reorganises the supplied COP.3 objective elements and Guidebook "
      "interpretation into plain language. Mandatory requirements and their stated modal "
      "strength have been retained. Optional, preferable, illustrative and scope-dependent "
      "provisions have not been converted into mandatory requirements.")

    save_and_verify(doc, "HCO_COP_3_v2_REWRITE_DRAFT.docx")


# ══════════════════════════════════════════════════════════════════════════════
# ══════════════════════════════════════════════════════════════════════════════
# COP.4 — Community Emergencies, Epidemics and Other Disasters  (NO stop-work)
# Content: ChatGPT final draft (approved, COP%2004.pdf).
# Structure: Document control table, NO Stop-work, Governance Section 6 (table),
#            Quality monitoring Section 7 (table), Checklist Section 12, Refs Section 13.
# OEs: a★, b★, c, d (all Commitment). Prepared by: Medical Superintendent.
# 14 plan elements: 7 "must" + 7 "shall" in Section 3.2 (all List Bullet).
# Frequencies: "at least twice a year" + "at least one mock drill once in 12 months".
# ══════════════════════════════════════════════════════════════════════════════
def gen_cop4():
    doc = Document()

    # Title
    h(doc, 0, "Policy on Community Emergencies, Epidemics and Other Disasters")
    p(doc, HN)

    # Document control
    h(doc, 1, "Document control")
    doc_ctrl(doc, "HCO/COP/POL/04", "Medical Superintendent")
    p(doc, "A blank marked ________ must be completed before issue.")

    # Statement of intent
    h(doc, 1, "Statement of intent")
    p(doc,
      "The organisation shall identify potential community emergencies, epidemics and other "
      "disasters and manage them through a documented plan. Resources shall be available "
      "according to threat perception and expected workload. The plan shall be tested at "
      "least twice a year, with at least one mock drill once in 12 months.")

    # 1. Purpose
    h(doc, 1, "1. Purpose")
    p(doc,
      "This policy establishes how the organisation identifies and manages community "
      "emergencies, epidemics and other disasters that may cause a sudden rush of victims.")
    p(doc, "Those requirements are covered in the hospital's other policies.")

    # 2. Scope
    h(doc, 1, "2. Scope")
    p(doc,
      "This policy applies to the organisation's planning and preparedness for community "
      "emergencies, epidemics and other disasters relevant to its geographical location and "
      "the community served.")

    # 3. Policy standards
    h(doc, 1, "3. Policy standards")

    h(doc, 2, "3.1 Identification of potential emergencies")
    p(doc,
      "The organisation shall identify potential community emergencies, epidemics and other "
      "disasters likely to cause a sudden rush of victims.")
    p(doc,
      "The potential emergencies shall be identified based on geographical location and the "
      "community served by the organisation.")
    p(doc,
      "Examples include earthquake, floods, train accident, civil unrest outside the "
      "organisation's premises, major fire and outbreak of disease/epidemics. These are "
      "examples, not a fixed list.")
    p(doc,
      "For example, an organisation in an industrial town shall identify the industrial "
      "hazards that may occur in its vicinity.")

    h(doc, 2, "3.2 Documented disaster, community emergency and epidemic plan")
    p(doc, "The disaster, community emergency and epidemic plan must incorporate:")
    lb(doc, "alert code")
    lb(doc, "information and communication")
    lb(doc, "action cards for each of the staff")
    lb(doc,
       "availability and earmarking of resources including adequacy of medical supplies, "
       "equipment, materials and trained personnel")
    lb(doc, "establishment of command nucleus")
    lb(doc, "training and mock drills")
    lb(doc, "managing clinical activities during the event")
    p(doc, "The plan shall also include:")
    lb(doc, "activating and deactivating the plan")
    lb(doc, "receiving, identifying and triaging casualties")
    lb(doc, "defined areas for reception and treatment for casualties")
    lb(doc, "transportation aids")
    lb(doc, "communication aids")
    lb(doc, "managing visitors and controlling the movement of individuals and vehicles")
    lb(doc, "relocating/discharging admitted patients wherever needed")
    p(doc,
      "The plans shall conform to the relevant local laws and national plans on disaster "
      "management.")
    p(doc,
      "The emergency room could follow triage policy according to the National Disaster "
      "Management Authority (NDMA) guidelines. A good reference is NDMA guidelines.")

    h(doc, 2, "3.3 Medical supplies, equipment and materials")
    p(doc, "Resource availability shall be according to threat perception.")
    p(doc,
      "The number of resources, including medical consumables and equipment, shall be "
      "commensurate with the expected workload.")

    h(doc, 2, "3.4 Testing of the plan")
    p(doc,
      "Testing twice a year is only the minimum frequency, and this may be increased.")
    p(doc,
      "In case the organisation has different plans for different disasters, each of the "
      "plans shall be tested at least twice a year.")
    p(doc, "The plan can be tested using a table-top exercise, or a mock drill.")
    p(doc,
      "At a minimum, at least one mock drill shall be held once in 12 months.")
    p(doc,
      "This shall test all the components of the plan and not just awareness.")
    p(doc,
      "In the case of a mock drill, simulated patients, not real patients, shall be used.")
    p(doc,
      "After every table-top exercise/mock drill, the variations are identified, the reason "
      "for the same is analysed, debriefing conducted and where appropriate the necessary "
      "corrective and/or preventive actions are taken.")

    # 4. Non-negotiable rules
    h(doc, 1, "4. Non-negotiable rules")
    lb(doc,
       "Do not leave potential community emergencies, epidemics or other disasters "
       "unidentified when they are relevant to the organisation's geographical location and "
       "the community served.")
    lb(doc,
       "Do not manage a community emergency, epidemic or other disaster without the "
       "documented plan incorporating the required alert code, information and communication, "
       "action cards, resources, command nucleus, training and mock drills, and management "
       "of clinical activities during the event.")
    lb(doc,
       "Do not use a documented disaster, community emergency and epidemic plan that omits "
       "activation and deactivation, casualty reception/identification/triage, defined "
       "reception and treatment areas, transportation aids, communication aids, visitor and "
       "movement management, and relocation/discharge of admitted patients wherever needed.")
    lb(doc,
       "Do not use the plan without conformity to relevant local laws and national plans on "
       "disaster management.")
    lb(doc,
       "Do not provide medical supplies, equipment and materials at levels that are not "
       "according to threat perception and commensurate with expected workload.")
    lb(doc, "Do not treat twice-yearly plan testing as optional.")
    lb(doc,
       "Where different plans exist for different disasters, do not leave any of those plans "
       "untested at least twice a year.")
    lb(doc,
       "Do not count testing that covers only awareness as sufficient; testing shall cover "
       "all components of the plan.")
    lb(doc,
       "Do not use real patients in a mock drill; simulated patients shall be used.")
    lb(doc,
       "Do not close a table-top exercise or mock drill without identifying variations, "
       "analysing the reason for the variations, conducting debriefing and, where "
       "appropriate, taking the necessary corrective and/or preventive actions.")

    # 5. What we do
    h(doc, 1, "5. What we do")

    h(doc, 2, "5.1 Identify relevant emergencies")
    p(doc,
      "Identify potential community emergencies, epidemics and other disasters likely to "
      "cause a sudden rush of victims.")
    p(doc, "Base identification on geographical location and the community served.")
    p(doc,
      "Identify industrial hazards that may occur in the vicinity when relevant to an "
      "organisation in an industrial town.")

    h(doc, 2, "5.2 Maintain the documented plan")
    p(doc,
      "Maintain the alert code, information and communication arrangements, action cards "
      "for each staff member, resource availability and earmarking, command nucleus, "
      "training and mock drills, and arrangements for managing clinical activities during "
      "the event.")
    p(doc,
      "Include activation and deactivation, casualty reception/identification/triage, "
      "defined casualty reception and treatment areas, transportation aids, communication "
      "aids, visitor and movement management, and relocation/discharge of admitted patients "
      "wherever needed.")
    p(doc,
      "Ensure conformity with relevant local laws and national plans on disaster management.")
    p(doc,
      "The emergency room could follow triage policy according to NDMA guidelines.")

    h(doc, 2, "5.3 Maintain resources")
    p(doc, "Make resources available according to threat perception.")
    p(doc,
      "Ensure the number of medical consumables, equipment and other resources is "
      "commensurate with expected workload.")

    h(doc, 2, "5.4 Test and improve the plan")
    p(doc, "Test the plan at least twice a year.")
    p(doc,
      "Where different plans exist for different disasters, test each plan at least twice "
      "a year.")
    p(doc, "A table-top exercise or mock drill can be used for testing.")
    p(doc, "Hold at least one mock drill once in 12 months.")
    p(doc, "Test all components of the plan and not just awareness.")
    p(doc, "Use simulated patients, not real patients, for mock drills.")
    p(doc,
      "After every table-top exercise/mock drill, identify variations, analyse the reason, "
      "conduct debriefing and, where appropriate, take corrective and/or preventive actions.")

    # 6. Governance and responsibility — proper table (NO stop-work section above)
    h(doc, 1, "6. Governance and responsibility")
    gov_tbl(doc, [
        ("Organisation",
         "The organisation is responsible for identifying relevant community emergencies, "
         "epidemics and other disasters, maintaining the documented plan, making required "
         "resources available and ensuring the plan is tested."),
        ("All staff",
         "The plan includes an established command nucleus and action cards for each staff "
         "member, defining individual responsibilities during the event."),
        ("Emergency Room",
         "The emergency room could follow triage policy according to NDMA guidelines where "
         "the organisation chooses to adopt it."),
    ])

    # 7. Quality monitoring — proper table
    h(doc, 1, "7. Quality monitoring")
    mon_tbl(doc, [
        ("Plan testing frequency",
         "Completion of plan testing at least twice a year"),
        ("Disaster-specific plan testing",
         "Testing of each disaster-specific plan at least twice a year where different "
         "plans exist"),
        ("Mock drill frequency",
         "Completion of at least one mock drill once in 12 months"),
        ("Testing coverage",
         "Coverage of all plan components during testing"),
        ("Mock-drill patients",
         "Use of simulated patients rather than real patients in mock drills"),
        ("Post-exercise actions",
         "Identification of variations, analysis of reasons, debriefing and "
         "corrective/preventive action where appropriate after every exercise"),
    ])

    # 8. Training and staff acknowledgement
    h(doc, 1, "8. Training and staff acknowledgement")
    p(doc,
      "Training and mock drills are incorporated into the documented disaster, community "
      "emergency and epidemic plan.")
    p(doc,
      f"I have read the Policy on Community Emergencies, Epidemics and Other Disasters of "
      f"{HN}. I will follow the processes described.")
    sig_tbl(doc)

    # 9. Distribution
    h(doc, 1, "9. Distribution")
    p(doc,
      "The documented disaster, community emergency and epidemic plan, including action "
      "cards for each staff member, shall be available to the personnel involved in "
      "managing the event.")

    # 10. Abbreviations
    h(doc, 1, "10. Abbreviations")
    abbrev_tbl(doc, [
        ("NDMA", "National Disaster Management Authority"),
    ])

    # 11. Traceability table
    h(doc, 1, "11. Traceability table")
    p(doc,
      "This table is an index. It is not how the policy is organised. An asterisk in the "
      "Level column means documentation of the process is required.")
    tr = tbl(doc, 5, 3)
    for ci, hdr in enumerate(("Objective Element", "Level", "Traceability to this policy")):
        tr.cell(0, ci).text = hdr
    trace_rows = [
        ("COP.4.a", "Commitment*",
         "Sections 3.1 and 5.1 address identification of potential community emergencies, "
         "epidemics and other disasters based on geographical location and the community "
         "served."),
        ("COP.4.b", "Commitment*",
         "Sections 3.2 and 5.2 address the documented disaster, community emergency and "
         "epidemic plan with all 14 required plan elements and conformity to relevant laws "
         "and national plans."),
        ("COP.4.c", "Commitment",
         "Sections 3.3 and 5.3 address availability of medical supplies, equipment and "
         "materials according to threat perception and commensurate with expected workload."),
        ("COP.4.d", "Commitment",
         "Sections 3.4 and 5.4 address testing at least twice a year, at least one mock "
         "drill once in 12 months, all-component coverage, simulated patients and four "
         "mandatory post-exercise steps."),
    ]
    for ri, (oe, lvl, txt) in enumerate(trace_rows, 1):
        tr.cell(ri, 0).text = oe
        tr.cell(ri, 1).text = lvl
        tr.cell(ri, 2).text = txt

    # 12. Required Records/Evidence Checklist — bulleted list
    h(doc, 1, "12. Required Records/Evidence Checklist")

    h(doc, 2, "Emergency identification")
    lb(doc,
       "Documented list of potential community emergencies, epidemics and other disasters "
       "identified for the organisation's geographical location and community.")
    lb(doc,
       "Identification of relevant hazards in the vicinity, including industrial hazards "
       "where relevant.")

    h(doc, 2, "Documented disaster, community emergency and epidemic plan")
    lb(doc,
       "Current documented plan containing the alert code; information and communication; "
       "action cards for each staff member; availability and earmarking of medical supplies, "
       "equipment, materials and trained personnel; command nucleus; training and mock drills; "
       "and management of clinical activities during the event.")
    lb(doc,
       "Plan covering activation and deactivation; casualty reception, identification and "
       "triage; defined casualty reception and treatment areas; transportation aids; "
       "communication aids; visitor management and control of individual and vehicle movement; "
       "and relocation/discharge of admitted patients wherever needed.")
    lb(doc,
       "Plan showing conformity with relevant local laws and national plans on disaster "
       "management.")
    lb(doc,
       "NDMA guidelines retained as a reference; any emergency-room triage policy based on "
       "NDMA guidelines where the organisation chooses to follow it.")

    h(doc, 2, "Resources")
    lb(doc,
       "Arrangements showing availability and earmarking of medical supplies, equipment, "
       "materials and trained personnel according to threat perception and expected workload.")

    h(doc, 2, "Testing and exercises")
    lb(doc,
       "Testing records showing the plan was tested at least twice a year.")
    lb(doc,
       "Where different disaster plans exist, testing records for each plan showing testing "
       "at least twice a year.")
    lb(doc,
       "Record of at least one mock drill once in 12 months.")
    lb(doc,
       "Exercise records showing all components of the plan were tested, not just awareness.")
    lb(doc,
       "Mock-drill records showing simulated patients, not real patients, were used.")
    lb(doc,
       "After-exercise records showing variations identified, reasons analysed, debriefing "
       "conducted and corrective and/or preventive actions taken where appropriate.")

    # 13. References
    h(doc, 1, "13. References")
    ln(doc, "National Disaster Management Authority (NDMA) guidelines.")
    ln(doc, "Relevant local laws and national plans on disaster management.")

    # Disclaimer
    h(doc, 1, "Disclaimer")
    p(doc,
      "This policy reorganises the supplied COP.4 objective elements and Guidebook "
      "interpretation into plain language. Mandatory requirements and their stated modal "
      "strength have been retained. Illustrative examples, optional approaches and references "
      "have not been converted into mandatory requirements.")

    save_and_verify(doc, "HCO_COP_4_v2_REWRITE_DRAFT.docx")


# ══════════════════════════════════════════════════════════════════════════════
# COP.5 — Cardio-pulmonary Resuscitation Services   (STOP-WORK YES)
# Content: ChatGPT final draft (approved, COP%2005.pdf).
# Structure: Stop-work Section 6, Governance Section 7, Quality monitoring Section 8.
# All OEs: Commitment | No stars | No COREs
# Exact frequency: "at least once a quarter" (3.5/5.5)
# 9 coverage areas (3.3/5.3) | 5 record items (3.4/5.4) | 5 analysis foci (3.5/5.5)
# 4 committee roles (3.5/5.5)
# ══════════════════════════════════════════════════════════════════════════════
def gen_cop5():
    doc = Document()

    # Title
    h(doc, 0, "Policy on Cardio-pulmonary Resuscitation Services")
    p(doc, HN)

    # Document control
    h(doc, 1, "Document control")
    doc_ctrl(doc, "HCO/COP/POL/05", "Medical Superintendent")
    p(doc, "A blank marked ________ must be completed before issue.")

    # Statement of intent
    h(doc, 1, "Statement of intent")
    p(doc,
      "Cardio-pulmonary resuscitation services shall be available and provided to patients "
      f"at all times. {HN} shall maintain a documented CPR procedure, appropriate equipment "
      "and medications, trained team roles, event recording, post-event analysis and corrective "
      "and preventive measures.")

    # 1. Purpose
    h(doc, 1, "1. Purpose")
    p(doc,
      "This policy establishes requirements for providing cardio-pulmonary resuscitation "
      "services, maintaining CPR readiness, recording CPR events and mock drills, analysing "
      "outcomes, and implementing corrective and preventive measures. Those requirements are "
      "covered in the hospital's other policies.")

    # 2. Scope
    h(doc, 1, "2. Scope")
    p(doc,
      f"This policy applies across all areas of {HN} where cardio-pulmonary resuscitation "
      "may be required, including the patient care areas identified in this policy.")

    # 3. Policy standards
    h(doc, 1, "3. Policy standards")

    h(doc, 2, "3.1 CPR procedure and immediate response")
    p(doc,
      "The organisation shall document the procedure for cardio-pulmonary resuscitation (CPR) "
      "for adults across all areas in the organisation. This shall be in consonance with "
      "accepted practices.")
    p(doc,
      "Where appropriate, it shall also address obstetric, paediatric and neonatal patients.")
    p(doc,
      "The organisation shall ensure that medical equipment for resuscitation and medications "
      "for basic and advanced life support are provided in standardised manner.")
    p(doc,
      "Basic life support shall be initiated as soon as a condition requiring CPR is identified. "
      "This shall be implemented in all areas of the organisation.")
    p(doc,
      "The protocols could be displayed prominently in all critical areas such as emergency, "
      "ICU, OT and all crash carts.")

    h(doc, 2, "3.2 CPR team roles and responsibilities")
    p(doc,
      "CPR team members shall have a clear understanding of their roles and responsibilities "
      "during resuscitation and shall comply with their assigned roles and responsibilities "
      "to effectively function as a team.")

    h(doc, 2, "3.3 Equipment and medications")
    p(doc,
      "At a minimum, emergency medications and equipment for intubation based on the needs "
      "of the patients served shall be available in patient care areas including the blood "
      "centre, radiology, day care, dialysis, chemo ward, OPD, rehabilitation services areas, "
      "endoscopy, and in areas where any invasive procedure is performed.")
    p(doc,
      "Other equipment like defibrillator shall be easily accessible to ensure that there is "
      "no delay in cardio-pulmonary resuscitation.")
    p(doc,
      "It is preferable that the minimum emergency medication is standardised across the "
      "organisation.")

    h(doc, 2, "3.4 Recording CPR events and mock drills")
    p(doc,
      "In the actual event of cardio-pulmonary resuscitation, or a mock drill of the same, "
      "all the activities along with the personnel attended shall be recorded.")
    p(doc,
      "At the minimum, it will include timeliness of response, availability of human resources, "
      "equipment, drugs, and barriers if any.")
    p(doc,
      "The recording could be done using the pre-defined procedural checklist and by monitoring "
      "whether the prescribed activity has been performed properly and in the right sequence.")
    p(doc,
      "It is a good practice to debrief team members regarding the necessary immediate "
      "corrective and preventive action.")

    h(doc, 2, "3.5 Post-event analysis")
    p(doc,
      "The frequency of the committee meeting shall be at least once a quarter.")
    p(doc,
      "The analysis shall focus on the initiation of CPR, time of arrival of the team, "
      "availability of required resources, recording of the sequence of events during CPR "
      "(including technique) and the overall coordination.")
    p(doc,
      "The organisation shall monitor outcome of CPR and identify areas for improvement.")
    p(doc,
      "The multidisciplinary committee shall be independent and include at least one "
      "physician/cardiologist, one anaesthesiologist, one member from the code blue team "
      "and one nurse.")
    p(doc,
      "The analysis shall be completed within a defined time frame.")

    h(doc, 2, "3.6 Corrective and preventive measures")
    p(doc,
      "Corrective and preventive measures shall be completed within a defined time frame.")
    p(doc,
      "The findings of the post-event analysis shall be communicated to the personnel who "
      "participated in the CPR.")
    p(doc,
      "Any lapses shall be discussed, with the view to improve the outcomes in future.")
    p(doc,
      "During subsequent resuscitations, it is preferable that implementation of these actions "
      "is noted and training be modified, if necessary.")

    # 4. Non-negotiable rules
    h(doc, 1, "4. Non-negotiable rules")
    lb(doc,
       "Do not start CPR without initiating basic life support as soon as a condition "
       "requiring CPR is identified.")
    lb(doc,
       "Do not start a planned resuscitation response without a named CPR team with clear "
       "roles and without the minimum emergency medications and equipment available at the "
       "location.")
    lb(doc,
       "Do not use a CPR process without the documented CPR procedure and standardised "
       "resuscitation equipment and medications required by this policy.")
    lb(doc,
       "Do not omit emergency medications and intubation equipment from the specified "
       "patient care areas.")
    lb(doc,
       "Do not conduct CPR where the defibrillator cannot be accessed without delay.")
    lb(doc,
       "Do not omit recording of an actual CPR event or mock drill, including timeliness "
       "of response, availability of human resources, equipment, drugs, and barriers if any.")
    lb(doc,
       "Do not complete a post-event analysis without the multidisciplinary committee meeting "
       "at least once a quarter and analysing the specified CPR performance areas.")
    lb(doc,
       "Do not omit the independent multidisciplinary committee composition specified in "
       "this policy.")
    lb(doc,
       "Do not leave corrective and preventive measures incomplete beyond the defined time "
       "frame.")
    lb(doc,
       "Do not withhold post-event analysis findings from personnel who participated in "
       "the CPR.")

    # 5. What we do
    h(doc, 1, "5. What we do")

    h(doc, 2, "5.1 CPR procedure and readiness")
    lb(doc,
       "Maintain a documented adult CPR procedure across all areas in consonance with "
       "accepted practices.")
    lb(doc, "Where appropriate, address obstetric, paediatric and neonatal patients.")
    lb(doc,
       "Provide resuscitation equipment and basic and advanced life-support medications in "
       "a standardised manner.")
    lb(doc,
       "Initiate basic life support as soon as a condition requiring CPR is identified in "
       "all areas.")
    lb(doc,
       "Protocols could be displayed prominently in critical areas such as emergency, ICU, "
       "OT and crash carts.")

    h(doc, 2, "5.2 CPR team")
    lb(doc,
       "Ensure CPR team members understand and comply with their assigned roles and "
       "responsibilities.")

    h(doc, 2, "5.3 Equipment and medication coverage")
    lb(doc,
       "Keep at least the required emergency medications and intubation equipment based on "
       "patient needs in the blood centre, radiology, day care, dialysis, chemo ward, OPD, "
       "rehabilitation services areas, endoscopy, and areas where an invasive procedure is "
       "performed.")
    lb(doc, "Keep the defibrillator easily accessible so there is no delay in CPR.")
    lb(doc, "Standardisation of the minimum emergency medication is preferable.")

    h(doc, 2, "5.4 Event recording")
    lb(doc, "Record all activities and personnel involved in actual CPR events and mock drills.")
    lb(doc,
       "Record timeliness of response, availability of human resources, equipment, drugs, "
       "and barriers if any.")
    lb(doc, "A predefined procedural checklist could be used.")
    lb(doc,
       "Debriefing team members regarding immediate corrective and preventive action is "
       "good practice.")

    h(doc, 2, "5.5 Post-event analysis")
    lb(doc, "Have the multidisciplinary committee meet at least once a quarter.")
    lb(doc,
       "Analyse initiation of CPR, time of arrival of the team, availability of required "
       "resources, recording of the sequence of events during CPR including technique, and "
       "overall coordination.")
    lb(doc, "Monitor CPR outcomes and identify areas for improvement.")
    lb(doc,
       "Maintain an independent multidisciplinary committee with at least one "
       "physician/cardiologist, one anaesthesiologist, one member from the code blue team "
       "and one nurse.")
    lb(doc, "Complete the analysis within a defined time frame.")

    h(doc, 2, "5.6 Corrective and preventive measures")
    lb(doc, "Complete corrective and preventive measures within a defined time frame.")
    lb(doc,
       "Communicate post-event analysis findings to personnel who participated in the CPR.")
    lb(doc, "Discuss lapses to improve outcomes in future.")
    lb(doc,
       "Preferably note implementation of actions during subsequent resuscitations and "
       "modify training if necessary.")

    # 6. Stop-work authority
    h(doc, 1, "6. Stop-work authority")
    p(doc,
      "Do not start a planned resuscitation response (or mock drill counted as CPR "
      "competence evidence) without a named CPR team with clear roles and without the "
      "minimum emergency medications and equipment available at the location. Stop-work "
      "does not block an unexpected cardiac arrest already in progress — start CPR "
      "with available staff and escalate for equipment/team immediately. The person "
      "responsible tells the CPR Committee chair or Emergency In-Charge the same shift. "
      "Refusing to run a hollow CPR response is not a disciplinary matter.")
    p(doc,
      "The post-event analysis committee is a separate governance and quality loop, not "
      "part of the stop-work trigger. Stop-work concerns CPR readiness at the moment of "
      "need, not analysis-timeline compliance.")

    # 7. Governance and responsibility — proper table
    h(doc, 1, "7. Governance and responsibility")
    gov_tbl(doc, [
        ("CPR team members",
         "CPR team members are responsible for understanding and complying with their "
         "assigned roles and responsibilities during resuscitation."),
        ("Multidisciplinary CPR committee",
         "Responsible for post-event analysis, CPR outcome monitoring and identification "
         "of areas for improvement. The committee shall be independent and include at least "
         "one physician/cardiologist, one anaesthesiologist, one member from the code blue "
         "team and one nurse."),
    ])

    # 8. Quality monitoring — proper table
    h(doc, 1, "8. Quality monitoring")
    mon_tbl(doc, [
        ("CPR event recording",
         "CPR events and mock drills, including the required minimum record content"),
        ("Committee meetings",
         "Committee meetings at least once a quarter"),
        ("Analysis coverage",
         "Analysis of initiation of CPR, time of arrival of the team, availability of "
         "required resources, recording of the sequence of events during CPR including "
         "technique, and overall coordination"),
        ("CPR outcomes",
         "CPR outcomes and areas for improvement"),
        ("CAPA completion",
         "Completion of corrective and preventive measures within the defined time frame"),
        ("Post-event communication",
         "Communication of post-event findings to CPR participants"),
    ])

    # 9. Training and staff acknowledgement
    h(doc, 1, "9. Training and staff acknowledgement")
    p(doc,
      "Training and mock drills form part of the documented CPR arrangements. CPR team "
      "members shall understand their assigned roles and responsibilities.")
    p(doc,
      f"I have read the Policy on Cardio-pulmonary Resuscitation Services of {HN}. "
      "I will follow the processes described.")
    sig_tbl(doc)

    # 10. Distribution
    h(doc, 1, "10. Distribution")
    p(doc,
      "The policy shall be available to personnel involved in CPR services, including CPR "
      "team members and personnel working in patient care areas.")

    # 11. Abbreviations
    h(doc, 1, "11. Abbreviations")
    abbrev_tbl(doc, [
        ("BLS",  "Basic Life Support"),
        ("ALS",  "Advanced Life Support"),
        ("CPR",  "Cardio-pulmonary Resuscitation"),
        ("CAPA", "Corrective and Preventive Action"),
        ("ICU",  "Intensive Care Unit"),
        ("OT",   "Operation Theatre"),
        ("OPD",  "Outpatient Department"),
    ])

    # 12. Traceability table
    h(doc, 1, "12. Traceability table")
    p(doc,
      "This table is an index. It is not how the policy is organised. An asterisk in the "
      "Level column means documentation of the process is required.")
    tr = tbl(doc, 7, 3)
    for ci, hdr in enumerate(("Objective Element", "Level", "Traceability to this policy")):
        tr.cell(0, ci).text = hdr
    trace_rows = [
        ("COP.5.a", "Commitment",
         "Sections 3.1 and 5.1 address the documented CPR procedure for adults across all "
         "areas, standardised resuscitation equipment and medications, and immediate BLS "
         "initiation."),
        ("COP.5.b", "Commitment",
         "Sections 3.2 and 5.2 address CPR team compliance with assigned roles and "
         "responsibilities."),
        ("COP.5.c", "Commitment",
         "Sections 3.3 and 5.3 address emergency medications and intubation equipment in "
         "the nine specified areas and areas where invasive procedures are performed, and "
         "defibrillator accessibility."),
        ("COP.5.d", "Commitment",
         "Sections 3.4 and 5.4 address recording of actual CPR events and mock drills, "
         "including the five minimum content items."),
        ("COP.5.e", "Commitment",
         "Sections 3.5 and 5.5 address post-event analysis: quarterly committee meetings, "
         "five analysis areas, CPR outcome monitoring, independent committee with minimum "
         "composition and analysis within a defined time frame."),
        ("COP.5.f", "Commitment",
         "Sections 3.6 and 5.6 address CAPA within a defined time frame, findings "
         "communicated to CPR participants, and lapses discussed for future improvement."),
    ]
    for ri, (oe, lvl, txt) in enumerate(trace_rows, 1):
        tr.cell(ri, 0).text = oe
        tr.cell(ri, 1).text = lvl
        tr.cell(ri, 2).text = txt

    # 13. Required Records/Evidence Checklist
    h(doc, 1, "13. Required Records/Evidence Checklist")

    h(doc, 2, "CPR procedure and readiness")
    lb(doc,
       "Documented adult CPR procedure across all areas in consonance with accepted practices.")
    lb(doc,
       "Written provisions, where appropriate, for obstetric, paediatric and neonatal CPR.")
    lb(doc,
       "Standardised resuscitation equipment and BLS/ALS medications across the organisation.")
    lb(doc,
       "Evidence that BLS is initiated immediately on identification of a CPR-requiring "
       "condition in all areas.")

    h(doc, 2, "Equipment and medication coverage")
    lb(doc,
       "Emergency medications and intubation equipment available in the blood centre, "
       "radiology, day care, dialysis, chemo ward, OPD, rehabilitation services areas, "
       "endoscopy, and all areas where invasive procedures are performed.")
    lb(doc, "Defibrillator accessible without delay in CPR locations.")

    h(doc, 2, "CPR event and mock drill recording")
    lb(doc, "Records of actual CPR events and mock drills including personnel attended.")
    lb(doc,
       "Minimum record content: timeliness of response, availability of human resources, "
       "equipment, drugs, and barriers if any.")

    h(doc, 2, "Post-event analysis")
    lb(doc,
       "Evidence that the multidisciplinary committee meets at least once a quarter.")
    lb(doc,
       "Analysis records covering: initiation of CPR, time of arrival of the team, "
       "availability of required resources, recording of event sequence including technique, "
       "and overall coordination.")
    lb(doc, "CPR outcome records and identified areas for improvement.")
    lb(doc,
       "Evidence of independent multidisciplinary committee with at least one "
       "physician/cardiologist, one anaesthesiologist, one member from the code blue team "
       "and one nurse.")
    lb(doc, "Evidence of analysis completed within the defined time frame.")

    h(doc, 2, "CAPA and communication")
    lb(doc, "CAPA records showing measures completed within the defined time frame.")
    lb(doc,
       "Evidence that post-event analysis findings were communicated to CPR participants.")
    lb(doc, "Records of lapses discussed and actions taken for future improvement.")

    # 14. References
    h(doc, 1, "14. References")
    ln(doc,
       "National Accreditation Board for Hospitals and Healthcare Providers. NABH "
       "Accreditation Standards for Hospitals, 6th Edition. COP.5.")
    ln(doc, "Guidebook interpretation supplied for COP.5.a through COP.5.f.")

    # Disclaimer
    h(doc, 1, "Disclaimer")
    p(doc,
      "This policy reorganises the supplied COP.5 objective elements and Guidebook "
      "interpretation into plain language. Mandatory requirements and their stated modal "
      "strength have been retained. Illustrative examples, optional approaches and references "
      "have not been converted into mandatory requirements.")

    save_and_verify(doc, "HCO_COP_5_v2_REWRITE_DRAFT.docx")


# ══════════════════════════════════════════════════════════════════════════════
# COP.6 — Nursing Care   (NO stop-work)
# Content: raw source dump cop_raw_dump_5-8.txt (no ChatGPT PDF — rebuilt from source)
# Structure: Governance Section 6, Quality monitoring Section 7 (no stop-work shift)
# Stars: a★, d★ | Achievement: c | All other OEs: Commitment
# Exact frequency: "reviewed annually at the minimum" (3.1/5.1)
# Five nursing care plan components (3.4/5.4): assessment, plan of care,
#   implementation of care, evaluation, modification of plan of care as may be required
# ══════════════════════════════════════════════════════════════════════════════
def gen_cop6():
    doc = Document()

    # Title
    h(doc, 0, "Policy on Nursing Care")
    p(doc, HN)

    # Document control
    h(doc, 1, "Document control")
    doc_ctrl(doc, "HCO/COP/POL/06", "Nursing Superintendent")
    p(doc, "A blank marked ________ must be completed before issue.")

    # Statement of intent
    h(doc, 1, "Statement of intent")
    p(doc,
      f"{HN} ensures that nursing care is provided to patients in consonance with clinical "
      "protocols. Nursing care follows written guidance, is assigned based on clinical need and "
      "competence, is documented in an individualised nursing care plan, and is supported by "
      "appropriate equipment. Nurses are empowered to make patient care decisions within their "
      "defined scope of practice.")

    # 1. Purpose
    h(doc, 1, "1. Purpose")
    p(doc,
      "This policy establishes requirements for providing nursing care through written guidance "
      "and clinical practice guidelines, assigning patient care based on clinical requirements "
      "and nursing competence, implementing acuity-based staffing, maintaining individualised "
      "nursing care plans, providing appropriate nursing equipment, and defining the scope of "
      "nursing practice decisions.")

    # 2. Scope
    h(doc, 1, "2. Scope")
    p(doc,
      f"This policy applies to all nursing staff providing patient care across all wards, "
      f"departments and care settings of {HN}.")

    # 3. Policy standards
    h(doc, 1, "3. Policy standards")

    h(doc, 2, "3.1 Written guidance and clinical practice guidelines")
    p(doc,
      "Nursing care is provided in accordance with written guidance. The written guidance could "
      "be in the form of a nursing manual or standard operating procedures incorporating various "
      "basic nursing practices and procedures.")
    p(doc,
      "Care of patients in specific clinical situations shall be guided by nursing clinical "
      "practice guidelines based on best clinical practices.")
    p(doc,
      "Nursing clinical care guidelines and pathways shall be reviewed annually at the minimum, "
      "and revised as appropriate.")
    p(doc,
      "Examples of nursing clinical practice guidelines include prevention of fall, prevention "
      "of development of pressure ulcers in an in-patient, and deep venous thrombosis risk "
      "assessment and prevention. These examples are illustrative; the organisation selects "
      "guidelines appropriate to its clinical scope.")

    h(doc, 2, "3.2 Assignment of patient care")
    p(doc,
      "Assignment of patient care shall be based on the patient's clinical requirements and "
      "the competence of the nursing staff, and shall align with the guidelines laid down by "
      "regulatory and professional bodies.")

    h(doc, 2, "3.3 Acuity-based staffing")
    p(doc,
      "The organisation implements acuity-based staffing — matching both the number and "
      "competence of nursing personnel to patient acuity — to improve patient outcomes.")
    p(doc,
      "Patient outcomes linked to acuity-based staffing may include incidence of pressure "
      "sores, falls, medication administration errors, and ventilator-associated pneumonia. "
      "The organisation selects relevant outcome indicators for its own scope of services.")

    h(doc, 2, "3.4 Nursing care plan")
    p(doc,
      "Care shall be provided as per the nursing care plan, which shall be individualised as "
      "per the clinical needs of each patient. Where a patient care plan has been developed, "
      "the nursing care plan shall be aligned with it. Uniformity and continuity of care shall "
      "be practised.")
    p(doc,
      "Components of the nursing care plan include: Assessment; Plan of care; Implementation "
      "of care; Evaluation; and Modification of plan of care as may be required.")
    p(doc,
      "Documentation includes all nursing-related care and not just monitoring of vitals and "
      "documentation of medication administration. Nursing progress shall be documented in a "
      "timely manner for each patient.")

    h(doc, 2, "3.5 Nursing equipment")
    p(doc,
      "There shall be an adequate number of basic nursing equipment and gadgets necessary for "
      "functioning in each designated area. Examples include nebuliser machines, glucometers, "
      "sphygmomanometers, thermometers and weighing scales — the organisation determines the "
      "appropriate equipment for its scope and patient population.")
    p(doc,
      "The equipment shall be appropriate for the area. For example, BP cuffs in a paediatric "
      "area shall be of appropriate size for the patient population served.")

    h(doc, 2, "3.6 Scope of nursing practice decisions")
    p(doc,
      "The organisation shall define the patient care decisions that come under the scope of "
      "nursing practice.")
    p(doc,
      "Nurses shall be aware of their defined scope and shall be able to make appropriate "
      "nursing-related decisions in a timely manner.")

    # 4. Non-negotiable rules
    h(doc, 1, "4. Non-negotiable rules")
    lb(doc,
       "Do not provide nursing care for specific clinical situations without nursing clinical "
       "practice guidelines based on best clinical practices.")
    lb(doc,
       "Do not let nursing clinical care guidelines and pathways go more than one year without "
       "review.")
    lb(doc,
       "Do not assign patient care without considering the patient's clinical requirements and "
       "the competence of the nursing staff.")
    lb(doc,
       "Do not omit the nursing care plan for any patient where one is required; the plan shall "
       "be individualised as per the patient's clinical needs.")
    lb(doc,
       "Do not omit any of the five nursing care plan components: Assessment; Plan of care; "
       "Implementation of care; Evaluation; and Modification of plan of care as may be required.")
    lb(doc,
       "Do not leave a ward or department without an adequate number of nursing equipment "
       "appropriate for the area and patient population.")
    lb(doc,
       "Do not leave the scope of nursing practice decisions undefined or leave nurses unaware "
       "of their defined scope.")

    # 5. What we do
    h(doc, 1, "5. What we do")

    h(doc, 2, "5.1 Written guidance and clinical practice guidelines")
    lb(doc,
       "Provide written guidance for nursing care — this could be a nursing manual or SOPs "
       "covering basic nursing practices and procedures.")
    lb(doc,
       "Guide care of patients in specific clinical situations through nursing clinical practice "
       "guidelines based on best clinical practices.")
    lb(doc,
       "Review nursing clinical care guidelines and pathways annually at the minimum and revise "
       "as appropriate.")
    lb(doc,
       "Examples of applicable guidelines include fall prevention, pressure ulcer prevention "
       "and DVT risk assessment — the organisation selects relevant guidelines for its scope.")

    h(doc, 2, "5.2 Assignment of patient care")
    lb(doc,
       "Assign patient care based on the patient's clinical requirements and the competence of "
       "the nursing staff, aligned with regulatory and professional body guidelines.")

    h(doc, 2, "5.3 Acuity-based staffing")
    lb(doc,
       "Implement acuity-based staffing — match the number and competence of nursing personnel "
       "to patient acuity — to improve patient outcomes.")
    lb(doc,
       "Monitor patient outcomes linked to acuity-based staffing and use results for "
       "improvement.")

    h(doc, 2, "5.4 Nursing care plan")
    lb(doc,
       "Provide nursing care as per an individualised nursing care plan based on each patient's "
       "clinical needs; where a patient care plan exists, align the nursing care plan with it.")
    lb(doc, "Practise uniformity and continuity of care.")
    lb(doc,
       "Ensure every nursing care plan covers all five components: Assessment; Plan of care; "
       "Implementation of care; Evaluation; and Modification of plan of care as may be required.")
    lb(doc,
       "Document all nursing-related care — not just monitoring of vitals and medication "
       "administration — in a timely manner for each patient.")

    h(doc, 2, "5.5 Nursing equipment")
    lb(doc,
       "Maintain an adequate number of basic nursing equipment appropriate for each ward and "
       "department and the patient population it serves.")
    lb(doc,
       "Ensure equipment is appropriate for the area — e.g., appropriate-size BP cuffs in "
       "paediatric areas.")

    h(doc, 2, "5.6 Scope of nursing practice decisions")
    lb(doc,
       "Define the patient care decisions that come under the scope of nursing practice and "
       "ensure all nurses are aware of their defined scope.")
    lb(doc,
       "Ensure nurses are able to make appropriate nursing-related decisions in a timely manner.")

    # 6. Governance and responsibility — proper table (no stop-work, so Section 6)
    h(doc, 1, "6. Governance and responsibility")
    gov_tbl(doc, [
        ("Nursing Superintendent",
         "Owns day-to-day implementation of nursing care processes; ensures nursing clinical "
         "practice guidelines are reviewed annually at the minimum; oversees individualised "
         "nursing care plan documentation and compliance."),
        ("Medical Superintendent",
         "Accountable that this policy is followed and appropriately resourced."),
        ("Quality Coordinator",
         "Supports audit of nursing care documentation, CPG review records and nursing "
         "equipment adequacy."),
    ])

    # 7. Quality monitoring — proper table
    h(doc, 1, "7. Quality monitoring")
    mon_tbl(doc, [
        ("Written guidance and CPG review",
         "Nursing written guidance in place; nursing CPGs/pathways reviewed annually at the "
         "minimum and revised as appropriate"),
        ("Nursing care plans",
         "Individualised nursing care plans present and documented, covering all five "
         "components: assessment, plan of care, implementation of care, evaluation, and "
         "modification of plan of care as may be required"),
        ("Nursing progress documentation",
         "Nursing progress documented in a timely manner for each patient, covering all "
         "nursing-related care and not only vitals and medication administration"),
        ("Care assignment",
         "Patient care assignment based on clinical requirements and nursing staff competence, "
         "aligned with regulatory and professional body guidelines"),
        ("Nursing equipment",
         "Adequate and area-appropriate nursing equipment present in all wards and departments"),
        ("Scope of practice",
         "Defined scope of nursing practice decisions; nurses aware and making timely decisions"),
    ])

    # 8. Training and staff acknowledgement
    h(doc, 1, "8. Training and staff acknowledgement")
    p(doc,
      "Nursing staff shall have access to ongoing training on nursing care processes, clinical "
      "practice guidelines and care plan documentation.")
    p(doc,
      f"I have read the Policy on Nursing Care of {HN}. I will follow the processes described.")
    sig_tbl(doc)

    # 9. Distribution
    h(doc, 1, "9. Distribution")
    p(doc,
      "This policy shall be available to all nursing staff, the Nursing Superintendent, "
      "ward in-charges and the Quality Coordinator.")

    # 10. Abbreviations
    h(doc, 1, "10. Abbreviations")
    abbrev_tbl(doc, [
        ("CPG",  "Clinical Practice Guideline"),
        ("DVT",  "Deep Venous Thrombosis"),
        ("SOP",  "Standard Operating Procedure"),
        ("VAP",  "Ventilator-Associated Pneumonia"),
    ])

    # 11. Traceability table
    h(doc, 1, "11. Traceability table")
    p(doc,
      "This table is an index. It is not how the policy is organised. An asterisk in the "
      "Level column means documentation of the process is required.")
    tr = tbl(doc, 7, 3)
    for ci, hdr in enumerate(("Objective Element", "Level", "Traceability to this policy")):
        tr.cell(0, ci).text = hdr
    trace_rows = [
        ("COP.6.a", "Commitment*",
         "Sections 3.1 and 5.1 address nursing care in accordance with written guidance, "
         "nursing CPGs based on best clinical practices for specific clinical situations, and "
         "annual review of nursing CPGs/pathways at the minimum."),
        ("COP.6.b", "Commitment",
         "Sections 3.2 and 5.2 address assignment of patient care based on clinical "
         "requirements and nursing staff competence, aligned with regulatory and professional "
         "body guidelines."),
        ("COP.6.c", "Achievement",
         "Sections 3.3 and 5.3 address implementation of acuity-based staffing — matching "
         "numbers and competence to patient acuity — to improve patient outcomes."),
        ("COP.6.d", "Commitment*",
         "Sections 3.4 and 5.4 address the individualised nursing care plan aligned with the "
         "overall patient care plan where one exists; uniformity and continuity of care; all "
         "five nursing care plan components; and timely documentation of all nursing-related "
         "care."),
        ("COP.6.e", "Commitment",
         "Sections 3.5 and 5.5 address adequate nursing equipment in adequate numbers, "
         "appropriate for each area and patient population."),
        ("COP.6.f", "Commitment",
         "Sections 3.6 and 5.6 address the defined scope of nursing practice decisions and "
         "nurses' ability to make timely nursing decisions within that scope."),
    ]
    for ri, (oe, lvl, txt) in enumerate(trace_rows, 1):
        tr.cell(ri, 0).text = oe
        tr.cell(ri, 1).text = lvl
        tr.cell(ri, 2).text = txt

    # 12. Required Records/Evidence Checklist
    h(doc, 1, "12. Required Records/Evidence Checklist")

    h(doc, 2, "Written guidance and CPG review")
    lb(doc,
       "Nursing written guidance (nursing manual, SOPs, or clinical practice guidelines) "
       "covering basic nursing practices and procedures.")
    lb(doc,
       "Nursing clinical practice guidelines for specific clinical situations, based on best "
       "clinical practices.")
    lb(doc,
       "Evidence of annual review and revision of nursing CPGs and pathways at the minimum.")

    h(doc, 2, "Nursing care plan and documentation")
    lb(doc,
       "Individualised nursing care plans covering all five components: assessment, plan of "
       "care, implementation of care, evaluation, and modification of plan of care as may be "
       "required.")
    lb(doc,
       "Evidence that the nursing care plan is aligned with the overall patient care plan "
       "where one exists.")
    lb(doc,
       "Nursing progress records documenting all nursing-related care in a timely manner, "
       "not limited to vitals and medication administration.")

    h(doc, 2, "Care assignment and staffing")
    lb(doc,
       "Evidence that patient care assignment is based on clinical requirements and nursing "
       "staff competence, aligned with regulatory and professional body guidelines.")
    lb(doc,
       "Acuity-based staffing arrangements — numbers and competence matched to patient acuity "
       "— and patient outcome monitoring.")

    h(doc, 2, "Nursing equipment")
    lb(doc,
       "Adequate number of basic nursing equipment present in each ward and department, "
       "appropriate for the area and patient population served.")

    h(doc, 2, "Scope of nursing practice")
    lb(doc,
       "Written definition of patient care decisions within the scope of nursing practice.")
    lb(doc,
       "Evidence that nurses are aware of their defined scope and make timely nursing decisions.")

    # 13. References
    h(doc, 1, "13. References")
    ln(doc,
       "National Accreditation Board for Hospitals and Healthcare Providers. NABH "
       "Accreditation Standards for Hospitals, 6th Edition. COP.6.")
    ln(doc, "Guidebook interpretation supplied for COP.6.a through COP.6.f.")

    # Disclaimer
    h(doc, 1, "Disclaimer")
    p(doc,
      "This policy reorganises the supplied COP.6 objective elements and Guidebook "
      "interpretation into plain language. Mandatory requirements and their stated modal "
      "strength have been retained. Illustrative examples, optional approaches and references "
      "have not been converted into mandatory requirements.")

    save_and_verify(doc, "HCO_COP_6_v2_REWRITE_DRAFT.docx")


# ══════════════════════════════════════════════════════════════════════════════
# COP.7 — Clinical Procedures Performed Safely   (HAS stop-work: Section 6)
# Content: Approved plain-language draft (cop7_content.txt), checked against source.
# Structure: Document control table, Stop-work Section 6, Governance Section 7,
#            Quality monitoring Section 8, Traceability Section 12, Checklist Section 13.
# OEs: a: Commitment | b: Commitment* | c: Commitment | d: CORE* | e: Commitment |
#      f: Commitment | g: Commitment
# Exact quantities: two identifiers (3.4/5.4) | three monitoring parameters: pulse,
#   blood pressure, respiratory rate (3.6/5.6) | five documentation elements (3.7/5.7):
#   procedure name, who performed it, salient steps, key findings, post-procedure care |
#   four meta-elements: name, date, time, signature (3.7/5.7)
# ══════════════════════════════════════════════════════════════════════════════
def gen_cop7():
    doc = Document()

    # Title
    h(doc, 0, "Policy on Clinical Procedures Performed Safely")
    p(doc, HN)

    # Document control
    h(doc, 1, "Document control")
    doc_ctrl(doc, "HCO/COP/POL/07", "Medical Superintendent")
    p(doc, "A blank marked ________ must be completed before issue.")

    # Statement of intent
    h(doc, 1, "Statement of intent")
    p(doc,
      "Clinical procedures are performed based on genuine clinical need, by qualified "
      "personnel, with the right patient, procedure, and site verified every time.")

    # 1. Purpose
    h(doc, 1, "1. Purpose")
    p(doc,
      f"This policy explains how {HN} decides when a clinical procedure is needed, makes "
      "sure only qualified personnel perform or assist it, verifies the correct patient, "
      "procedure, and site before starting, and documents what was done.")
    p(doc,
      "This policy does not cover informed consent content, patient assessment, or "
      "privileging criteria in detail — those are covered in other hospital policies.")

    # 2. Scope
    h(doc, 1, "2. Scope")
    p(doc,
      f"This policy applies to all clinical staff who order, plan, perform, or assist "
      f"in performing clinical procedures at {HN}.")

    # 3. Policy standards
    h(doc, 1, "3. Policy standards")
    p(doc,
      f"{HN} bases every clinical procedure on genuine clinical need, uses only privileged "
      "personnel to perform or assist, follows a documented checklist to verify the correct "
      "patient, procedure, and site before starting, takes informed consent where applicable, "
      "monitors patients during and after the procedure, and documents every procedure "
      "accurately.")
    p(doc, "Staff follow the written guidance below and keep the records it requires.")

    # 4. Non-negotiable rules
    h(doc, 1, "4. Non-negotiable rules")
    lb(doc,
       "Do not perform a procedure unless it is based on the patient's clinical need and "
       "consistent with standard treatment guidelines or sound clinical practice.")
    lb(doc,
       "Do not perform any procedure — diagnostic, therapeutic, or supportive — without "
       "written guidance covering who does it, pre-procedure instructions, the conduct of "
       "the procedure, and post-procedure care.")
    lb(doc,
       "Do not skip the pre-procedure assessment, which must include at least the patient's "
       "vitals.")
    lb(doc,
       "Do not let anyone who is not privileged for a procedure perform it or assist in it.")
    lb(doc,
       "Do not proceed with any procedure without completing the documented "
       "patient/procedure/site verification checklist, using at least two identifiers, one "
       "of which is the unique identification number.")
    lb(doc,
       "Do not skip the attempt to verify patient/procedure/site in an emergency — where "
       "full verification isn't possible, document the exception in the medical record.")
    lb(doc,
       "Do not proceed with a procedure requiring informed consent without obtaining it from "
       "the person performing the procedure or a doctor from the treating team.")
    lb(doc,
       "Do not skip the minimum monitoring parameters — pulse, blood pressure, and "
       "respiratory rate — during and after an invasive procedure.")
    lb(doc,
       "Do not document a procedure without including the procedure name, who performed it, "
       "the salient steps, key findings, and post-procedure care, along with name, date, "
       "time, and signature.")

    # 5. What we do
    h(doc, 1, "5. What we do")

    h(doc, 2, "5.1 Base every procedure on clinical need")
    p(doc,
      "The decision to perform a procedure is based on the patient's clinical needs, "
      "following standard treatment guidelines or sound clinical practice.")
    p(doc, "A qualified medical practitioner decides whether the procedure is indicated.")
    p(doc,
      "Where more than one procedure option exists, the choice is based on the likely best "
      "outcome, taking the patient's wishes and safety into account.")
    p(doc,
      f"{HN} could audit its procedures over time to help achieve the best outcomes.")

    h(doc, 2, "5.2 Follow written guidance for every type of procedure")
    p(doc,
      "Written guidance covers all procedures — diagnostic, therapeutic, and supportive.")
    p(doc,
      "The guidance states who performs the procedure, pre-procedure instructions where "
      "applicable, how the procedure is carried out, and post-procedure instructions and "
      "care where applicable.")
    p(doc, "Equipment used for procedures is set up and used safely.")
    p(doc,
      "Before any procedure, a brief assessment is done that includes at least the "
      "patient's vitals.")

    h(doc, 2, "5.3 Use only privileged personnel")
    p(doc,
      "Anyone who orders, plans, performs, or assists in a procedure is privileged to "
      "do so.")

    h(doc, 2, "5.4 Verify the correct patient, procedure, and site every time")
    p(doc,
      "Before any procedure, staff use a documented checklist to prevent wrong-patient, "
      "wrong-procedure, and wrong-site events.")
    p(doc,
      "At least two identifiers are used to confirm the patient, one of which is the "
      "unique identification number generated at registration.")
    p(doc,
      "Every team member shares responsibility; the person performing the procedure carries "
      "the final responsibility. Where a trainee performs the procedure, the supervising "
      "clinician carries the final responsibility.")
    p(doc,
      "In an emergency, staff still attempt verification; exceptions are documented in the "
      "medical record. Patients or their relatives are involved in site marking whenever "
      "possible.")

    h(doc, 2, "5.5 Take informed consent")
    p(doc,
      "Where applicable, informed consent is taken by the person performing the procedure "
      "or a doctor from the treating team.")
    p(doc,
      "If a trainee performs the procedure, this is stated in the consent documentation "
      "and the treating doctor supervises.")

    h(doc, 2, "5.6 Monitor patients during and after the procedure")
    p(doc,
      "For invasive procedures, monitoring includes at minimum pulse, blood pressure, and "
      "respiratory rate, plus any other parameter the patient's condition calls for.")
    p(doc,
      "The extent and duration of monitoring depends on procedure complexity and the "
      "patient's other health conditions.")

    h(doc, 2, "5.7 Document every procedure accurately")
    p(doc,
      "Every procedure is documented with the procedure name, who performed it, the key "
      "steps taken, the findings, and the post-procedure care given.")
    p(doc, "Every entry includes name, date, time, and signature.")

    # 6. Stop-work authority
    h(doc, 1, "6. Stop-work authority")
    p(doc,
      "Do not perform a clinical procedure that is not based on the patient's clinical "
      "need, that lacks written guidance where required, or that is ordered, performed, or "
      "assisted by personnel not qualified for that procedure.")
    p(doc,
      "Stop-work applies to the elective or non-emergent procedure start. Emergent "
      "life-saving procedures continue with the best available qualified help and are "
      "documented afterward.")
    p(doc,
      "The person who stops tells the treating doctor and the Medical Superintendent the "
      "same shift. Refusing an unsafe procedure is not a disciplinary matter.")

    # 7. Governance and responsibility — proper table
    h(doc, 1, "7. Governance and responsibility")
    gov_tbl(doc, [
        ("Medical Superintendent",
         "Accountable that this policy is followed and appropriately resourced; receives "
         "stop-work notifications the same shift."),
        ("Treating doctor / clinical team",
         "Decides clinical indication, ensures the correct patient/procedure/site "
         "verification checklist is completed, takes or supervises informed consent, and "
         "is notified of stop-work the same shift."),
        ("Person performing the procedure",
         "Carries final responsibility for the verification checklist and for accurate "
         "procedure documentation."),
        ("Supervising clinician (where a trainee performs)",
         "Carries final responsibility for the verification checklist and supervises "
         "consent documentation when a person in training performs the procedure."),
        ("All clinical staff involved",
         "Follow the applicable written guidance; participate in pre-procedure "
         "patient/procedure/site verification; share responsibility for correct "
         "verification."),
    ])

    # 8. Quality monitoring — proper table
    h(doc, 1, "8. Quality monitoring")
    mon_tbl(doc, [
        ("Clinical indication",
         "Procedures based on clinical need and consistent with standard treatment "
         "guidelines or sound clinical practice"),
        ("Written guidance",
         "Written guidance in place for all types of procedure performed"),
        ("Patient/procedure/site verification",
         "Documented checklist completed before every procedure, with at least two "
         "identifiers including the unique identification number"),
        ("Emergency verification exceptions",
         "Exceptions documented in the medical record where full verification was not "
         "possible in an emergency"),
        ("Invasive procedure monitoring",
         "Minimum monitoring parameters — pulse, blood pressure, and respiratory rate — "
         "recorded for invasive procedures"),
        ("Procedure documentation completeness",
         "Procedure records containing all five documentation elements and all four "
         "meta-elements"),
    ])

    # 9. Training and staff acknowledgement
    h(doc, 1, "9. Training and staff acknowledgement")
    p(doc,
      "Staff involved in clinical procedures shall be familiar with the written guidance "
      "for the procedures they perform or assist, the patient/procedure/site verification "
      "process, the minimum monitoring requirements, and the procedure documentation "
      "requirements applicable to their work.")
    p(doc,
      f"I have read the Policy on Clinical Procedures Performed Safely of {HN}. "
      "I will follow the processes described.")
    sig_tbl(doc)

    # 10. Distribution
    h(doc, 1, "10. Distribution")
    p(doc,
      "This policy shall be available to all clinical staff who order, plan, perform, or "
      "assist in performing clinical procedures.")

    # 11. Abbreviations
    h(doc, 1, "11. Abbreviations")
    abbrev_tbl(doc, [
        ("SOP", "Standard Operating Procedure"),
        ("WHO", "World Health Organisation"),
    ])

    # 12. Traceability table
    h(doc, 1, "12. Traceability table")
    p(doc,
      "This table is an index. It is not how the policy is organised. An asterisk in the "
      "Level column means documentation of the process is required.")
    tr = tbl(doc, 8, 3)
    for ci, hdr in enumerate(("Objective Element", "Level", "Traceability to this policy")):
        tr.cell(0, ci).text = hdr
    trace_rows = [
        ("COP.7.a", "Commitment",
         "Section 5.1 addresses basing every procedure on clinical need, practitioner "
         "decision on indication, outcome-based choice where options exist, and optional "
         "auditing of procedures."),
        ("COP.7.b", "Commitment*",
         "Section 5.2 addresses written guidance for all types of procedure, the "
         "required content of that guidance, safe equipment set-up and use, and the minimum "
         "pre-procedure assessment including vitals."),
        ("COP.7.c", "Commitment",
         "Section 5.3 addresses use of only privileged personnel to order, plan, "
         "perform, or assist in any procedure."),
        ("COP.7.d", "CORE*",
         "Section 5.4 addresses the documented patient/procedure/site verification "
         "checklist, at least two identifiers including the unique identification number, "
         "responsibility allocation, emergency exception documentation, and patient/relative "
         "involvement in site marking."),
        ("COP.7.e", "Commitment",
         "Section 5.5 addresses informed consent by the person performing the "
         "procedure or a treating-team doctor, and trainee-performer consent documentation "
         "and supervision."),
        ("COP.7.f", "Commitment",
         "Section 5.6 addresses the three minimum monitoring parameters — pulse, "
         "blood pressure, and respiratory rate — for invasive procedures, plus additional "
         "monitoring as clinically indicated."),
        ("COP.7.g", "Commitment",
         "Section 5.7 addresses the five documentation elements — procedure name, "
         "who performed it, key steps, findings, post-procedure care — and the four "
         "meta-elements — name, date, time, signature."),
    ]
    for ri, (oe, lvl, txt) in enumerate(trace_rows, 1):
        tr.cell(ri, 0).text = oe
        tr.cell(ri, 1).text = lvl
        tr.cell(ri, 2).text = txt

    # 13. Required Records/Evidence Checklist
    h(doc, 1, "13. Required Records/Evidence Checklist")

    h(doc, 2, "Clinical indication and written guidance")
    lb(doc,
       "Evidence that every procedure performed is based on the patient's clinical need "
       "and consistent with standard treatment guidelines or sound clinical practice.")
    lb(doc,
       "Written guidance for all types of procedure performed — diagnostic, therapeutic, "
       "and supportive — covering who performs the procedure, pre-procedure instructions, "
       "procedure conduct, and post-procedure care where applicable.")
    lb(doc, "Evidence of safe equipment set-up and use for procedures.")
    lb(doc,
       "Pre-procedure assessment records showing at least the patient's vitals were "
       "assessed before every procedure.")

    h(doc, 2, "Personnel privileging")
    lb(doc,
       "Evidence that anyone who orders, plans, performs, or assists in a procedure is "
       "privileged to do so.")

    h(doc, 2, "Patient/procedure/site verification")
    lb(doc,
       "Documented patient/procedure/site verification checklist completed before every "
       "procedure.")
    lb(doc,
       "Verification records showing use of at least two identifiers, one of which is the "
       "unique identification number generated at registration.")
    lb(doc,
       "Records showing the person performing the procedure carries final verification "
       "responsibility, or the supervising clinician where a trainee performs the procedure.")
    lb(doc,
       "Medical record entries documenting exceptions where full verification could not be "
       "completed in an emergency.")
    lb(doc,
       "Evidence of patient or relative involvement in site marking whenever possible.")

    h(doc, 2, "Informed consent")
    lb(doc,
       "Informed consent documentation signed by the person performing the procedure or "
       "a doctor from the treating team.")
    lb(doc,
       "Consent documentation stating that a person in training performed the procedure "
       "and that the treating doctor supervised, where applicable.")

    h(doc, 2, "Monitoring")
    lb(doc,
       "Monitoring records for invasive procedures showing at minimum pulse, blood pressure, "
       "and respiratory rate, with additional parameters recorded as clinically indicated.")

    h(doc, 2, "Procedure documentation")
    lb(doc,
       "Procedure records containing all five documentation elements: procedure name, who "
       "performed it, key steps, key findings, and post-procedure care.")
    lb(doc,
       "Procedure records containing all four meta-elements: name, date, time, and "
       "signature.")

    # 14. References
    h(doc, 1, "14. References")
    ln(doc,
       "National Accreditation Board for Hospitals and Healthcare Providers. NABH "
       "Accreditation Standards for Hospitals, 6th Edition. COP.7.")
    ln(doc, "Guidebook interpretation supplied for COP.7.a through COP.7.g.")
    ln(doc,
       "WHO Surgical Safety Checklist (reference for patient/procedure/site verification).")

    # Disclaimer
    h(doc, 1, "Disclaimer")
    p(doc,
      "This policy reorganises the supplied COP.7 objective elements and approved "
      "plain-language content into policy format. The stop-work text is reproduced exactly "
      "as supplied. Mandatory requirements and their stated modal strength have been "
      "retained. The exact two-identifier quantity, three minimum monitoring parameters, "
      "five procedure documentation elements, and four meta-elements have been preserved "
      "throughout.")

    save_and_verify(doc, "HCO_COP_7_v2_REWRITE_DRAFT.docx")


# ══════════════════════════════════════════════════════════════════════════════
# COP.8 — Transfusion Services   (HAS stop-work: Section 6)
# Content: Approved plain-language draft (cop8_content.txt), checked against source.
# Structure: Document control table, Stop-work Section 6, Governance Section 7,
#            Quality monitoring Section 8, Traceability Section 12, Checklist Section 13.
# OEs: a: Commitment | b: Commitment* | c: Commitment | d: CORE* | e: Commitment* |
#      f: Achievement* | g: Achievement*
# Exact quantities:
#   Nine COP.8.b blood-centre written-guidance activities (5.2):
#     donor selection; TTI screening; blood collection incl apheresis; storage;
#     compatibility testing; distribution; transfusion in clinical areas;
#     discard of sero-positive/unutilised; donor/patient family education.
#   Eight COP.8.d transfusion written-guidance elements (5.4):
#     patient consent; safe procurement/transportation; patient identification;
#     verification of order; safe storage prior to transfusion; blood administration;
#     patient monitoring; identification of/response to transfusion reactions.
#   HBTC composition: clinical departments, blood transfusion officer, nursing,
#     Quality Coordinator, management — preferably chaired by a clinician.
#   Stop-work: "minimum two identifiers" exact.
#   Haemovigilance Programme participation: mandatory (not advisory).
#   Emergency blood availability time frame: organisation-defined — no invented number.
# ══════════════════════════════════════════════════════════════════════════════
def gen_cop8():
    doc = Document()

    # Title
    h(doc, 0, "Policy on Transfusion Services")
    p(doc, HN)

    # Document control
    h(doc, 1, "Document control")
    doc_ctrl(doc, "HCO/COP/POL/08", "Blood Transfusion Officer")
    p(doc, "A blank marked ________ must be completed before issue.")

    # Statement of intent
    h(doc, 1, "Statement of intent")
    p(doc,
      "Blood and blood components are collected, tested, stored, and given to patients "
      "safely, with rational use and complete pre-transfusion verification every time.")

    # 1. Purpose
    h(doc, 1, "1. Purpose")
    p(doc,
      f"This policy explains how {HN} keeps its blood supply available and safe, from "
      "donor collection through to giving blood to a patient, verifies the correct patient "
      "and correct blood before every transfusion, tracks and responds to transfusion "
      "reactions, and runs a rational-use committee for blood.")
    p(doc,
      "This policy does not cover general informed consent procedures or infection control "
      "in detail — those are covered in other hospital policies.")

    # 2. Scope
    h(doc, 1, "2. Scope")
    p(doc,
      f"This policy applies to blood centre staff, transfusion services staff, and clinical "
      f"staff involved in ordering, administering, or monitoring blood transfusions at {HN}.")

    # 3. Policy standards — two summary paragraphs only
    h(doc, 1, "3. Policy standards")
    p(doc,
      f"{HN} keeps blood and blood components available from a registered blood centre — "
      "in-house or outsourced — collects, tests, stores, and distributes them under "
      "written guidance, ensures safe and rational use through a Hospital Blood Transfusion "
      "Committee, defines how quickly blood must be available in an emergency, and tracks "
      "and responds to transfusion reactions.")
    p(doc, "Staff follow the written guidance below and keep the records it requires.")

    # 4. Non-negotiable rules
    h(doc, 1, "4. Non-negotiable rules")
    lb(doc,
       "Do not use an unregistered blood centre — blood and blood components come only "
       "from a registered in-house or outsourced blood centre, with an MoU if outsourced, "
       "and patient care must not suffer for want of blood or components.")
    lb(doc,
       "Do not release blood before all required pre-transfusion tests, including "
       "Transfusion Transmissible Infections (TTI) testing and Red Cell Serology, are "
       "complete using approved methods.")
    lb(doc,
       "Do not take a quarantined unit into inventory before re-testing is complete.")
    lb(doc,
       "Do not mix untested and tested blood units in storage, and do not label a unit's "
       "serology status before TTI screening reports are released.")
    lb(doc,
       "Do not start a transfusion without completing all eight elements of the written "
       "guidance: patient consent for blood administration, safe procurement and "
       "transportation maintaining the cold chain, patient identification, verification of "
       "the blood administration order, safe storage prior to transfusion, blood "
       "administration, monitoring of the patient, and identification of and response to "
       "potential transfusion reactions.")
    lb(doc,
       "Do not start transfusion of blood or blood components when patient identification "
       "(minimum two identifiers) is incomplete, compatibility checks are missing, or "
       "required informed consent has not been obtained, except where the organisation's "
       "documented emergency-transfusion guidance applies.")
    lb(doc,
       "Do not leave the emergency-use time frame for blood availability undefined, or "
       "leave routine and emergency turnaround times undefined or unmonitored.")
    lb(doc,
       "Do not skip detection, reporting, evaluation, or analysis of a suspected adverse "
       "transfusion reaction, and do not skip corrective and preventive action based on "
       "that analysis.")
    lb(doc,
       "Do not operate without participating in the Haemovigilance Programme of India.")

    # 5. What we do
    h(doc, 1, "5. What we do")

    h(doc, 2, "5.1 Keep blood available from a registered blood centre")
    p(doc,
      "Blood and blood components are available from either an in-house or an outsourced "
      "registered blood centre. The blood centre has adequate infrastructure, staff, and "
      "equipment for its workload and scope, so operations run smoothly without delays.")
    p(doc,
      "Where the blood centre is outsourced, there is a signed MoU, and patient care does "
      "not suffer for want of blood or components — blood is transported from the external "
      "blood centre safely and properly.")
    p(doc, "The NABH standards for blood centres are a useful reference.")

    h(doc, 2, "5.2 Follow written guidance for blood centre activities")
    p(doc,
      "Written guidance covers nine blood centre activities: blood donor selection; blood "
      "screening for transfusion-transmissible diseases; blood collection, including "
      "apheresis procedures; blood storage; compatibility testing; blood distribution; "
      "transfusion of blood and products in clinical areas; discard of sero-positive or "
      "unutilised blood; and education of families of potential donors and patients "
      "regarding blood donation.")
    p(doc,
      "Blood and apheresis products are collected only from voluntary, non-remunerated, "
      "low-risk, safe, and healthy donors, with pre- and post-donation counselling and "
      "informed consent. Apheresis procedures follow relevant guidelines. Blood donation "
      "drives follow national and statutory guidelines. Drugs and equipment for treating "
      "donor reactions are available, and staff are trained to use them. Adverse donor "
      "reactions are identified, managed, and reported in line with the National "
      "Haemovigilance Programme of India.")
    p(doc,
      "All pre-transfusion tests — including Transfusion Transmissible Infections (TTI) "
      "testing and Red Cell Serology — are carried out using approved methods before blood "
      "is released. Recipient and donor blood samples are retained and stored after each "
      "transfusion, as per guidelines. If a quality-control failure occurs during TTI "
      "testing, quarantined units are not taken into inventory until re-testing is "
      "complete. There is a process for retesting, recall, and referral of reactive blood "
      "donors, in line with national guidelines. Whole blood or components from any unit "
      "that tests positive or reactive are discarded as per written guidance.")

    h(doc, 2, "5.3 Store blood safely from collection to transfusion")
    p(doc,
      "The blood centre prepares components according to its scope. The temperature of "
      "whole blood, its components, and the storage environment is maintained and recorded "
      "as per guidelines.")
    p(doc,
      "Untested blood and products are quarantined, with separate, designated storage for "
      "untested and tested units based on serology status. Labelling — including serology "
      "status — happens only after TTI screening reports are released, so that untested "
      "and tested blood cannot get mixed up.")

    h(doc, 2, "5.4 Ensure safe and rational transfusion")
    p(doc,
      "Transfusion of blood and blood components follows written guidance covering eight "
      "things: patient consent for blood administration; safe procurement and "
      "transportation, maintaining the cold chain and correct source; patient "
      "identification; verification of blood administration orders; safe storage prior to "
      "transfusion; blood administration; monitoring of the patient; and identification of "
      "and response to potential transfusion reactions.")
    p(doc,
      "Measures are in place to prevent mismatched transfusion under all circumstances — "
      "for example, confirming patient identity at cross-matching, attaching the "
      "compatibility testing label to the unit at issue, confirming patient identity and "
      "blood group again at the time of transfusion, keeping a portion of the integral "
      "tubing with its segment number attached to the blood bag at issue, and keeping all "
      "labels attached to the blood bag until the transfusion is complete, so any adverse "
      "reaction can be properly investigated. These are examples of how the prevention "
      "principle is put into practice, not a fixed checklist.")
    p(doc,
      "Written guidance also governs the indications for using blood and blood components, "
      "based on standard practice guidelines or sound clinical practice from national and "
      "international professional bodies, and covers inventory and ordering schedules, both "
      "planned and unplanned. Protocols could address specific clinical situations relevant "
      "to the hospital's scope — for example, neonatal transfusions, autoimmune haemolytic "
      "anaemia, incompatible transfusion, or massive transfusion.")
    p(doc,
      "A Hospital Blood Transfusion Committee (HBTC) ensures the rational use of blood by "
      "adopting or adapting rational-use guidelines. It is preferably chaired by a "
      "clinician, and includes representation from clinical departments, the blood "
      "transfusion officer, nursing, the Quality Coordinator, and management. The committee "
      "meets periodically and monitors the availability of blood and components, transfusion "
      "practices and audits, adverse transfusion reactions, and wastage of blood or "
      "products. The committee could also seek periodic feedback from clinicians on the "
      "availability and effectiveness of blood and components.")

    h(doc, 2, "5.5 Define how fast blood must be available")
    p(doc,
      f"{HN} defines what counts as an emergency use of blood — covering both actual and "
      "anticipated need — and puts procedures in place to make sure blood is available "
      "for it.")
    p(doc,
      f"The time frame within which blood must be available in an emergency is defined by "
      f"{HN}. Turnaround time for both routine and emergency blood issues is defined and "
      "monitored. This applies even if the hospital does not have its own in-house blood "
      "centre.")

    h(doc, 2, "5.6 Track and respond to transfusion reactions")
    p(doc,
      "Suspected adverse transfusion reactions are detected, reported, evaluated, and "
      "analysed, with corrective and preventive action taken afterward. It is preferable "
      "to capture feedback on every transfusion, so no reaction goes unnoticed — analysis "
      f"can be done by an individual or a committee, as {HN} decides. A record of "
      "transfusion reactions is maintained.")
    p(doc,
      f"{HN} participates in the Haemovigilance Programme of India. (Related requirements "
      "are covered in the hospital's other policies.)")

    h(doc, 2, "5.7 Run a quality assurance programme for the blood centre")
    p(doc,
      "The quality assurance programme — whether standalone or part of the hospital's "
      "overall quality improvement programme — covers every aspect of the blood centre's "
      "work.")
    p(doc,
      "The blood centre takes part in an External Quality Assurance Scheme (EQAS) or "
      "Proficiency Testing (PT) programme, or a suitable interlaboratory comparison if "
      "EQAS/PT is not available. Results from internal quality control, EQAS, PT, or "
      "interlaboratory testing are monitored, with corrective and preventive action taken "
      "and records kept. Blood, products, and reagents are sampled at a defined frequency "
      "for internal quality control, following guidelines.")
    p(doc,
      "The blood centre collects data on key performance indicators as part of its quality "
      "improvement work — for example, wastage of blood and products, or the rate of "
      "transfusion reactions. This data is collated, analysed, and used to drive further "
      "improvements, which are monitored to make sure they hold.")

    # 6. Stop-work authority — text exactly as given in cop8_content.txt
    h(doc, 1, "6. Stop-work authority")
    p(doc,
      "Do not start transfusion of blood or blood components when patient identification "
      "(minimum two identifiers) is incomplete, compatibility checks are missing, or "
      "required informed consent has not been obtained, except where the organisation's "
      "documented emergency-transfusion guidance applies.")
    p(doc,
      "Stop-work applies to the transfusion start. Life-saving emergency transfusion "
      "follows the organisation's emergency-transfusion written guidance and is documented.")
    p(doc,
      "The person who stops tells the Transfusion / Blood Bank In-Charge and the treating "
      "doctor the same shift. Refusing an unsafe transfusion is not a disciplinary matter.")

    # 7. Governance and responsibility
    h(doc, 1, "7. Governance and responsibility")
    gov_tbl(doc, [
        ("Medical Superintendent",
         "Accountable that this policy is followed and appropriately resourced."),
        ("Blood Transfusion Officer",
         "Oversees blood centre operations, written guidance for all nine activities, "
         "pre-transfusion testing, storage, and quality assurance; receives stop-work "
         "notifications the same shift."),
        ("Hospital Blood Transfusion Committee (HBTC)",
         "Ensures rational use of blood; preferably chaired by a clinician; includes "
         "representation from clinical departments, the blood transfusion officer, nursing, "
         "the Quality Coordinator, and management; meets periodically."),
        ("Clinical staff (ordering and administering transfusions)",
         "Follow all eight transfusion written-guidance elements; complete pre-transfusion "
         "verification; monitor patients; report suspected adverse reactions."),
        ("Blood centre staff",
         "Follow written guidance for all nine blood centre activities; maintain storage "
         "and quarantine requirements; carry out pre-transfusion testing using approved "
         "methods."),
    ])

    # 8. Quality monitoring
    h(doc, 1, "8. Quality monitoring")
    mon_tbl(doc, [
        ("Registered blood centre",
         "Use of a registered blood centre; MoU in place where outsourced"),
        ("Pre-transfusion testing",
         "TTI testing and Red Cell Serology complete using approved methods before blood "
         "is released"),
        ("Storage and quarantine",
         "Separate storage of untested and tested units; labelling only after TTI reports "
         "released"),
        ("Transfusion verification",
         "All eight written-guidance elements completed before every transfusion"),
        ("Blood availability",
         "Defined emergency and routine turnaround times; monitoring against the "
         "organisation's defined time frames"),
        ("Adverse transfusion reactions",
         "Detection, reporting, evaluation, analysis, and CAPA for suspected adverse "
         "reactions; record of transfusion reactions maintained"),
        ("Haemovigilance",
         "Active participation in the Haemovigilance Programme of India"),
        ("Quality assurance programme",
         "EQAS/PT participation; internal quality control at defined frequency; KPI data "
         "collected, analysed, and used for improvement"),
    ])

    # 9. Training and staff acknowledgement
    h(doc, 1, "9. Training and staff acknowledgement")
    p(doc,
      "Blood centre staff and clinical staff involved in transfusion shall be familiar "
      "with the written guidance applicable to their work, including pre-transfusion "
      "testing, storage requirements, the eight-element transfusion process, adverse "
      "reaction reporting, and Haemovigilance Programme requirements.")
    p(doc,
      f"I have read the Policy on Transfusion Services of {HN}. I will follow the "
      "processes described.")
    sig_tbl(doc)

    # 10. Distribution
    h(doc, 1, "10. Distribution")
    p(doc,
      "This policy shall be available to blood centre staff, transfusion services staff, "
      "and clinical staff involved in ordering, administering, or monitoring blood "
      "transfusions.")

    # 11. Abbreviations
    h(doc, 1, "11. Abbreviations")
    abbrev_tbl(doc, [
        ("CAPA", "Corrective and Preventive Action"),
        ("EQAS", "External Quality Assurance Scheme"),
        ("HBTC", "Hospital Blood Transfusion Committee"),
        ("MoU",  "Memorandum of Understanding"),
        ("PT",   "Proficiency Testing"),
        ("TTI",  "Transfusion Transmissible Infections"),
    ])

    # 12. Traceability table
    h(doc, 1, "12. Traceability table")
    p(doc,
      "This table is an index. It is not how the policy is organised. An asterisk in the "
      "Level column means documentation of the process is required.")
    tr = tbl(doc, 8, 3)
    for ci, hdr in enumerate(("Objective Element", "Level", "Traceability to this policy")):
        tr.cell(0, ci).text = hdr
    trace_rows = [
        ("COP.8.a", "Commitment",
         "Section 5.1 addresses availability of blood and blood components from a "
         "registered in-house or outsourced blood centre, with an MoU where outsourced, "
         "and adequate infrastructure, staff, and equipment."),
        ("COP.8.b", "Commitment*",
         "Section 5.2 addresses written guidance for all nine blood centre activities: "
         "blood donor selection; TTI screening; blood collection including apheresis; "
         "blood storage; compatibility testing; blood distribution; transfusion in clinical "
         "areas; discard of sero-positive or unutilised blood; and donor/patient family "
         "education."),
        ("COP.8.c", "Commitment",
         "Section 5.3 addresses safe storage from collection to transfusion: component "
         "preparation, temperature maintenance and recording, quarantine of untested units, "
         "separate designated storage by serology status, and labelling only after TTI "
         "reports are released."),
        ("COP.8.d", "CORE*",
         "Section 5.4 addresses written guidance covering all eight transfusion elements: "
         "patient consent; safe procurement and transportation; patient identification; "
         "verification of blood administration orders; safe storage prior to transfusion; "
         "blood administration; patient monitoring; and identification of and response to "
         "potential transfusion reactions. Also addresses mismatch prevention measures, "
         "rational-use indications, and the HBTC."),
        ("COP.8.e", "Commitment*",
         "Section 5.5 addresses the organisation-defined emergency blood availability time "
         "frame and defined, monitored turnaround times for routine and emergency blood "
         "issues."),
        ("COP.8.f", "Achievement*",
         "Section 5.6 addresses detection, reporting, evaluation, analysis, and CAPA for "
         "suspected adverse transfusion reactions; maintenance of a transfusion reaction "
         "record; and mandatory participation in the Haemovigilance Programme of India."),
        ("COP.8.g", "Achievement*",
         "Section 5.7 addresses the quality assurance programme covering all blood centre "
         "activities, EQAS/PT participation or interlaboratory comparison, internal quality "
         "control at a defined frequency, and KPI data collection, analysis, and "
         "improvement monitoring."),
    ]
    for ri, (oe, lvl, txt) in enumerate(trace_rows, 1):
        tr.cell(ri, 0).text = oe
        tr.cell(ri, 1).text = lvl
        tr.cell(ri, 2).text = txt

    # 13. Required Records/Evidence Checklist
    h(doc, 1, "13. Required Records/Evidence Checklist")

    h(doc, 2, "Registered blood centre and outsourcing")
    lb(doc,
       "Evidence that blood and blood components come from a registered in-house or "
       "outsourced blood centre.")
    lb(doc, "Signed MoU with the outsourced blood centre where applicable.")
    lb(doc,
       "Evidence that patient care does not suffer for want of blood or components.")

    h(doc, 2, "Written guidance — nine blood centre activities")
    lb(doc, "Written guidance for blood donor selection.")
    lb(doc, "Written guidance for blood screening for transfusion-transmissible diseases.")
    lb(doc, "Written guidance for blood collection, including apheresis procedures.")
    lb(doc, "Written guidance for blood storage.")
    lb(doc, "Written guidance for compatibility testing.")
    lb(doc, "Written guidance for blood distribution.")
    lb(doc, "Written guidance for transfusion of blood and products in clinical areas.")
    lb(doc, "Written guidance for discard of sero-positive or unutilised blood.")
    lb(doc,
       "Written guidance for education of families of potential donors and patients "
       "regarding blood donation.")

    h(doc, 2, "Donor selection and blood collection")
    lb(doc,
       "Records showing blood and apheresis products collected only from voluntary, "
       "non-remunerated, low-risk, safe, and healthy donors.")
    lb(doc, "Evidence of pre- and post-donation counselling and informed consent.")
    lb(doc,
       "Drugs and equipment for treating donor reactions available; staff trained to "
       "use them.")
    lb(doc,
       "Adverse donor reaction records reported in line with the National Haemovigilance "
       "Programme of India.")

    h(doc, 2, "Pre-transfusion testing and quarantine")
    lb(doc,
       "Records showing TTI testing and Red Cell Serology carried out using approved "
       "methods before blood is released.")
    lb(doc,
       "Records of recipient and donor blood sample retention and storage after each "
       "transfusion.")
    lb(doc,
       "Evidence that quarantined units are not taken into inventory until re-testing is "
       "complete after a quality-control failure.")
    lb(doc,
       "Process records for retesting, recall, and referral of reactive blood donors in "
       "line with national guidelines.")
    lb(doc,
       "Records showing whole blood or components from reactive units discarded as per "
       "written guidance.")

    h(doc, 2, "Storage")
    lb(doc,
       "Temperature records for whole blood, components, and the storage environment.")
    lb(doc,
       "Evidence of separate, designated storage for untested and tested units by serology "
       "status.")
    lb(doc,
       "Records showing labelling — including serology status — done only after TTI "
       "screening reports are released.")

    h(doc, 2, "Transfusion — eight written-guidance elements")
    lb(doc, "Patient consent for blood administration.")
    lb(doc,
       "Records showing safe procurement and transportation maintaining the cold chain "
       "and correct source.")
    lb(doc, "Patient identification completed before transfusion.")
    lb(doc, "Verification of blood administration orders.")
    lb(doc, "Safe storage prior to transfusion.")
    lb(doc, "Blood administration records.")
    lb(doc, "Patient monitoring records during and after transfusion.")
    lb(doc,
       "Records of identification of and response to potential transfusion reactions.")

    h(doc, 2, "Mismatch prevention and rational use")
    lb(doc,
       "Evidence of mismatch-prevention measures in place under all circumstances.")
    lb(doc,
       "Written guidance on indications for blood and blood components based on standard "
       "practice guidelines or sound clinical practice.")
    lb(doc, "Inventory and ordering schedules, both planned and unplanned.")

    h(doc, 2, "Hospital Blood Transfusion Committee")
    lb(doc,
       "Evidence that an HBTC is in place, preferably chaired by a clinician, with "
       "representation from clinical departments, the blood transfusion officer, nursing, "
       "the Quality Coordinator, and management.")
    lb(doc,
       "Records of periodic HBTC meetings monitoring blood availability, transfusion "
       "practices and audits, adverse reactions, and blood/product wastage.")

    h(doc, 2, "Emergency blood availability")
    lb(doc,
       "Documented definition of emergency use of blood, covering actual and anticipated "
       "need.")
    lb(doc,
       f"Defined emergency blood availability time frame (set by {HN}), with evidence of "
       "monitoring against it.")
    lb(doc,
       "Defined and monitored turnaround times for both routine and emergency blood "
       "issues.")

    h(doc, 2, "Adverse reactions and Haemovigilance")
    lb(doc,
       "Records showing suspected adverse transfusion reactions detected, reported, "
       "evaluated, and analysed.")
    lb(doc, "CAPA records for adverse transfusion reactions.")
    lb(doc, "Record of transfusion reactions maintained.")
    lb(doc,
       "Evidence of active participation in the Haemovigilance Programme of India.")

    h(doc, 2, "Quality assurance programme")
    lb(doc,
       "Quality assurance programme covering all blood centre activities — standalone or "
       "part of the overall quality improvement programme.")
    lb(doc,
       "Evidence of EQAS or PT participation, or interlaboratory comparison where "
       "EQAS/PT is not available.")
    lb(doc,
       "Records showing internal quality control, EQAS/PT/interlaboratory results "
       "monitored, with CAPA and records kept.")
    lb(doc,
       "Evidence of blood, product, and reagent sampling at a defined frequency for "
       "internal quality control.")
    lb(doc,
       "KPI data (for example, blood/product wastage, transfusion reaction rates) "
       "collected, analysed, and used for improvement; improvements monitored.")

    # 14. References
    h(doc, 1, "14. References")
    ln(doc,
       "National Accreditation Board for Hospitals and Healthcare Providers. NABH "
       "Accreditation Standards for Hospitals, 6th Edition. COP.8.")
    ln(doc, "Guidebook interpretation supplied for COP.8.a through COP.8.g.")
    ln(doc, "National Haemovigilance Programme of India guidelines.")
    ln(doc,
       "NABH Standards for Blood Centres (reference for blood centre requirements).")

    # Disclaimer
    h(doc, 1, "Disclaimer")
    p(doc,
      "This policy reorganises the supplied COP.8 objective elements and approved "
      "plain-language content into policy format. The stop-work text is reproduced exactly "
      "as supplied. Mandatory requirements and their stated modal strength have been "
      "retained. The nine COP.8.b blood-centre written-guidance activities, the eight "
      "COP.8.d transfusion written-guidance elements, the HBTC mandatory composition, "
      "mandatory Haemovigilance Programme participation, and the minimum-two-identifiers "
      "stop-work quantity have been preserved. No numerical time frame for emergency blood "
      "availability has been invented; the time frame is defined by the organisation.")

    save_and_verify(doc, "HCO_COP_8_v2_REWRITE_DRAFT.docx")


# ══════════════════════════════════════════════════════════════════════════════
# COP.9 — Intensive Care and High Dependency Units   (stop-work YES)
# Content: approved plain-language content (cop9_content.txt).
# Structure: Section 3 two summary paragraphs; full detail Section 5.
# Stars: a*, b*, d*, e*, f* | Achievement: f | CORE: none
# Exact quantities: four-direction criteria (admission/transfer-in/discharge/transfer-out);
#   DAMA mechanism separate; "at least once a day" counselling exact.
# ══════════════════════════════════════════════════════════════════════════════

def gen_cop9():
    doc = Document()

    # Title
    h(doc, 0, "Policy on Intensive Care and High Dependency Units")
    p(doc, HN)

    # Document control
    h(doc, 1, "Document control")
    doc_ctrl(doc, "HCO/COP/POL/09", "ICU/HDU In-Charge")
    p(doc, "A blank marked ________ must be completed before issue.")

    # Statement of intent
    h(doc, 1, "Statement of intent")
    p(doc,
      "Care in intensive care and high dependency units follows written guidance, "
      "uses defined admission and discharge criteria, and is backed by adequate "
      "staff, equipment, and infection control.")

    # 1. Purpose
    h(doc, 1, "1. Purpose")
    p(doc,
      f"This policy explains how {HN} provides ICU/HDU care based on written "
      "guidance, applies defined criteria for admitting, transferring, and "
      "discharging patients, keeps adequate trained staff and equipment available, "
      "handles bed shortages, follows infection prevention and control practices, "
      "and counsels patients and families.")
    p(doc,
      "This policy does not cover general infection control procedures or nursing "
      "staffing ratios in detail — those are covered in other hospital policies.")

    # 2. Scope
    h(doc, 1, "2. Scope")
    p(doc,
      f"This policy applies to all clinical and nursing staff working in the "
      f"intensive care and high dependency units at {HN}.")

    # 3. Policy standards — two summary paragraphs only
    h(doc, 1, "3. Policy standards")
    p(doc,
      f"{HN} provides ICU/HDU care under written guidance that, at a minimum, "
      "covers everything in this policy, applies defined admission, transfer-in, "
      "discharge, and transfer-out criteria, keeps adequate trained staff and "
      "necessary equipment available, has a procedure for bed shortages, follows "
      "infection prevention and control practices, runs a quality assurance "
      "programme, and counsels patients and families at least once a day.")
    p(doc, "Staff follow the written guidance below and keep the records it requires.")

    # 4. Non-negotiable rules
    h(doc, 1, "4. Non-negotiable rules")
    lb(doc,
       "Do not provide ICU/HDU care without written guidance based on standard "
       "treatment guidelines or sound clinical practice — at a minimum, this "
       "guidance must cover everything else in this policy.")
    lb(doc,
       "Do not admit, transfer in, discharge, or transfer out an ICU/HDU patient "
       "without applying the hospital's defined criteria for that decision.")
    lb(doc,
       "Do not leave a mechanism for \"discharge against medical advice\" undefined.")
    lb(doc,
       "Do not staff or equip the ICU/HDU with anything less than the necessary "
       "lifesaving and monitoring equipment and competent, trained staff, or without "
       "following applicable statutory requirements.")
    lb(doc,
       "Do not leave a bed-shortage situation without following the hospital's "
       "defined procedure for it.")
    lb(doc,
       "Do not depart from infection prevention and control practices consistent "
       "with standard treatment guidelines or sound clinical practice.")
    lb(doc,
       "Do not skip counselling the patient or family at least once a day, and "
       "whenever the patient's condition changes significantly, by a doctor of the "
       "treating team.")
    lb(doc,
       "Do not counsel a patient or family on a significant change without "
       "documenting it in writing or by audio-video recording.")

    # 5. What we do
    h(doc, 1, "5. What we do")

    h(doc, 2, "5.1 Follow written guidance covering the whole standard")
    p(doc,
      "Care in the ICU/HDU is based on written guidance drawn from standard "
      "treatment guidelines or sound clinical practice for intensive/critical care. "
      "At a minimum, this guidance addresses everything else in this policy — "
      "admission and discharge criteria, staffing and equipment, bed-shortage "
      f"procedure, infection control, quality assurance, and counselling. {HN} "
      "could adapt care-bundles where appropriate.")

    h(doc, 2, "5.2 Apply defined admission, transfer, and discharge criteria")
    p(doc,
      f"{HN} develops and adheres to criteria — based on physiologic and/or "
      "diagnostic parameters — covering four situations: admission, transfer-in, "
      "discharge, and transfer-out of ICU/HDU patients. There is also a defined "
      "mechanism for discharge against medical advice. Staff are trained to apply "
      "these criteria. National and international critical care society guidelines "
      "are a useful starting point.")

    h(doc, 2, "5.3 Keep adequate staff and equipment")
    p(doc,
      "The ICU/HDU is equipped with all necessary lifesaving and monitoring "
      "equipment, handled by trained staff. Staff working in the ICU/HDU are "
      "competent, based on qualification and/or training. Applicable statutory "
      f"requirements are followed. The exact equipment and staffing requirements are "
      f"{HN}'s own determination, based on the scope and complexity of its services "
      "and guided by standard treatment guidelines or best clinical practice. The "
      "Indian Nursing Council's recommendations are a useful reference for nursing "
      "workforce planning.")

    h(doc, 2, "5.4 Have a procedure for bed shortages")
    p(doc,
      "Whenever there are no vacant ICU beds and a patient needs one, "
      f"{HN} follows a detailed, pre-established policy and procedure to address "
      "the situation — not something improvised at the time.")

    h(doc, 2, "5.5 Follow infection prevention and control practices")
    p(doc,
      "Infection prevention and control practices in the ICU/HDU are consistent "
      "with standard treatment guidelines or sound clinical practice. Written "
      "guidance for this could stand alone or be part of the hospital's overall "
      "infection prevention and control manual. (Related requirements are covered "
      "in the hospital's other policies.)")

    h(doc, 2, "5.6 Run a quality assurance programme")
    p(doc,
      "The quality assurance programme covers every aspect of ICU/HDU functioning, "
      "consistent with standard treatment guidelines or sound clinical practice. "
      "Care outcomes are monitored — for example, risk-adjusted standardised "
      "mortality rate, infection rates, re-admission rates, or re-intubation rates. "
      f"These are examples; {HN} chooses the outcomes most relevant to its patients. "
      "National and international critical care society guidelines are a useful "
      "starting point for this programme, and the written guidance for it could "
      "stand alone or be part of the hospital's overall quality-improvement "
      "programme.")

    h(doc, 2, "5.7 Counsel patients and families")
    p(doc,
      "A doctor from the treating team counsels the patient and/or family at least "
      "once a day, and whenever there's a significant change in the patient's "
      "condition. Counselling covers significant events since the last session, "
      "expected outcomes, and the family's queries about the patient's changing "
      "condition. Counselling is documented in writing, by audio-video recording, "
      "or both. It's preferable to also obtain an acknowledgement.")

    # 6. Stop-work authority — text exactly as given in cop9_content.txt
    h(doc, 1, "6. Stop-work authority")
    p(doc,
      "Do not admit, transfer, or continue managing an ICU/HDU patient when "
      "defined admission/discharge criteria are not applied, required minimum "
      "staffing or equipment is absent, or the bed-shortage escalation procedure "
      "has not been followed when no bed is available.")
    p(doc,
      "Stop-work applies to non-emergent admission/transfer decisions and elective "
      "bed allocation. A patient already in the ICU/HDU or in an active emergency "
      "continues receiving care while the gap is escalated.")
    p(doc,
      "The person who stops tells the ICU In-Charge and the Medical Superintendent "
      "the same shift. Refusing an unsafe ICU/HDU admission or continuation is not "
      "a disciplinary matter.")

    # 7. Governance and responsibility
    h(doc, 1, "7. Governance and responsibility")
    gov_tbl(doc, [
        ("Medical Superintendent",
         "Accountable that this policy is followed and appropriately resourced; "
         "receives stop-work notifications the same shift."),
        ("ICU/HDU In-Charge",
         "Oversees ICU/HDU operations; ensures written guidance, defined criteria "
         "for all four admission/transfer/discharge/transfer-out situations, "
         "bed-shortage procedure, IPC practices, and QA programme are in place; "
         "receives stop-work notifications the same shift."),
        ("Treating doctors (ICU/HDU)",
         "Apply defined criteria for every admission, transfer-in, discharge, and "
         "transfer-out decision; counsel patients and families at least once a day "
         "and whenever the patient's condition changes significantly; document "
         "counselling in writing or by audio-video recording."),
        ("ICU/HDU nursing staff",
         "Maintain care standards; follow IPC practices; assist with daily "
         "documentation; support quality assurance monitoring."),
    ])

    # 8. Quality monitoring
    h(doc, 1, "8. Quality monitoring")
    mon_tbl(doc, [
        ("Written guidance",
         "Written guidance covering all required elements is in place and current"),
        ("Criteria application",
         "Defined criteria applied for all four situations (admission, transfer-in, "
         "discharge, transfer-out); DAMA mechanism defined and in use"),
        ("Staffing and equipment",
         "ICU/HDU staffed with competent personnel and equipped with required "
         "lifesaving and monitoring equipment; statutory requirements followed"),
        ("Bed-shortage procedure",
         "Pre-established procedure followed whenever ICU beds are unavailable; "
         "deviations escalated and documented"),
        ("Infection prevention and control",
         "IPC practices in ICU/HDU consistent with standard treatment guidelines "
         "or sound clinical practice"),
        ("Quality assurance",
         "Defined care outcomes monitored and analysed (e.g., mortality rates, "
         "infection rates, re-admission or re-intubation rates)"),
        ("Counselling documentation",
         "Daily counselling by treating team doctor documented; counselling on "
         "significant changes documented in writing or audio-video recording"),
    ])

    # 9. Training and staff acknowledgement
    h(doc, 1, "9. Training and staff acknowledgement")
    p(doc,
      "All clinical and nursing staff working in the ICU/HDU shall be familiar "
      "with the written guidance applicable to their work, including admission and "
      "discharge criteria for all four situations, the bed-shortage procedure, "
      "infection prevention and control practices, the quality assurance "
      "programme, and daily counselling requirements.")
    p(doc,
      f"I have read the Policy on Intensive Care and High Dependency Units of "
      f"{HN}. I will follow the processes described.")
    sig_tbl(doc)

    # 10. Distribution
    h(doc, 1, "10. Distribution")
    p(doc,
      "This policy shall be available to all clinical and nursing staff working "
      "in the intensive care and high dependency units.")

    # 11. Abbreviations
    h(doc, 1, "11. Abbreviations")
    abbrev_tbl(doc, [
        ("DAMA", "Discharge Against Medical Advice"),
        ("HDU",  "High Dependency Unit"),
        ("ICU",  "Intensive Care Unit"),
        ("IPC",  "Infection Prevention and Control"),
        ("QA",   "Quality Assurance"),
    ])

    # 12. Traceability table
    h(doc, 1, "12. Traceability table")
    p(doc,
      "This table is an index. It is not how the policy is organised. An asterisk "
      "in the Level column means documentation of the process is required.")
    tr = tbl(doc, 8, 3)
    for ci, hdr in enumerate(("Objective Element", "Level", "Traceability to this policy")):
        tr.cell(0, ci).text = hdr
    trace_rows = [
        ("COP.9.a", "Commitment*",
         "Section 5.1 addresses written guidance for ICU/HDU care drawn from "
         "standard treatment guidelines or sound clinical practice, covering "
         "at a minimum all other elements of this standard."),
        ("COP.9.b", "Commitment*",
         "Section 5.2 addresses defined criteria covering four situations: "
         "admission, transfer-in, discharge, and transfer-out of ICU/HDU "
         "patients. Also addresses the defined mechanism for discharge against "
         "medical advice (DAMA) as a separate requirement."),
        ("COP.9.c", "Commitment",
         "Section 5.3 addresses adequacy of lifesaving and monitoring equipment, "
         "trained and competent ICU/HDU staff based on qualification and/or "
         "training, and compliance with applicable statutory requirements."),
        ("COP.9.d", "Commitment*",
         "Section 5.4 addresses the pre-established, detailed bed-shortage "
         "procedure followed whenever ICU beds are unavailable."),
        ("COP.9.e", "Commitment*",
         "Section 5.5 addresses infection prevention and control practices in "
         "the ICU/HDU consistent with standard treatment guidelines or sound "
         "clinical practice."),
        ("COP.9.f", "Achievement*",
         "Section 5.6 addresses the quality assurance programme covering all "
         "aspects of ICU/HDU functioning, with defined outcomes monitored "
         "(e.g., mortality rate, infection rates, re-admission and re-intubation "
         "rates)."),
        ("COP.9.g", "Commitment",
         "Section 5.7 addresses counselling by a treating team doctor at least "
         "once a day and on every significant condition change, with documentation "
         "in writing or by audio-video recording."),
    ]
    for ri, (oe, lvl, txt) in enumerate(trace_rows, 1):
        tr.cell(ri, 0).text = oe
        tr.cell(ri, 1).text = lvl
        tr.cell(ri, 2).text = txt

    # 13. Required Records/Evidence Checklist
    h(doc, 1, "13. Required Records/Evidence Checklist")

    h(doc, 2, "Written guidance")
    lb(doc,
       "Written guidance for ICU/HDU care based on standard treatment guidelines "
       "or sound clinical practice, covering all required elements.")

    h(doc, 2, "Admission, transfer, and discharge criteria")
    lb(doc,
       "Defined criteria (based on physiologic and/or diagnostic parameters) "
       "covering admission, transfer-in, discharge, and transfer-out.")
    lb(doc,
       "Defined mechanism for discharge against medical advice (DAMA).")
    lb(doc, "Evidence that staff are trained to apply the criteria.")

    h(doc, 2, "Staffing and equipment")
    lb(doc,
       "Evidence of lifesaving and monitoring equipment in the ICU/HDU.")
    lb(doc,
       "Evidence that ICU/HDU staff are competent based on qualification and/or "
       "training.")
    lb(doc, "Evidence of compliance with applicable statutory requirements.")

    h(doc, 2, "Bed-shortage procedure")
    lb(doc,
       "Documented bed-shortage procedure; records of application when beds are "
       "unavailable.")

    h(doc, 2, "Infection prevention and control")
    lb(doc,
       "Written guidance for IPC practices in the ICU/HDU (standalone or part of "
       "the hospital's IPC manual).")

    h(doc, 2, "Quality assurance programme")
    lb(doc,
       "Quality assurance programme documentation covering all aspects of ICU/HDU "
       "functioning.")
    lb(doc,
       "Records of defined care outcomes monitored and analysed.")

    h(doc, 2, "Counselling records")
    lb(doc,
       "Daily counselling entries in patient records (in writing or audio-video "
       "recording).")
    lb(doc,
       "Counselling records for every significant change in patient condition.")

    # 14. References
    h(doc, 1, "14. References")
    ln(doc,
       "National Accreditation Board for Hospitals and Healthcare Providers. NABH "
       "Accreditation Standards for Hospitals, 6th Edition. COP.9.")
    ln(doc, "Guidebook interpretation supplied for COP.9.a through COP.9.g.")
    ln(doc,
       "National and international critical care society guidelines (reference for "
       "written guidance, criteria development, and quality assurance programme).")
    ln(doc,
       "Indian Nursing Council recommendations (reference for nursing workforce "
       "planning in ICU/HDU).")

    # Disclaimer
    h(doc, 1, "Disclaimer")
    p(doc,
      "This policy reorganises the supplied COP.9 objective elements and approved "
      "plain-language content into policy format. The stop-work text is reproduced "
      "exactly as supplied. Mandatory requirements and their stated modal strength "
      "have been retained. The four-direction criteria (admission, transfer-in, "
      "discharge, transfer-out) and the separate DAMA mechanism are preserved "
      "exactly as stated. The exact phrase \"at least once a day\" for counselling "
      "frequency is preserved. No specific staffing ratios, equipment lists, or "
      "outcome thresholds have been invented; these are the hospital's own "
      "determination.")

    save_and_verify(doc, "HCO_COP_9_v2_REWRITE_DRAFT.docx")


# ══════════════════════════════════════════════════════════════════════════════
# COP.10 — Obstetric Care   (stop-work YES)
# Content: approved plain-language content (cop10_content.txt).
# Structure: Section 3 two summary paragraphs; full detail Section 5.
# Stars: a*, d* | Achievement: e | CORE: none
# Exact quantities: three minimum written-guidance areas (5.1); four minimum
#   ante-natal areas (5.4); three minimum monitoring items (5.8);
#   six ART scope elements (5.10); COP.10.j and COP.10.k scope-conditional.
# ══════════════════════════════════════════════════════════════════════════════

def gen_cop10():
    doc = Document()

    # Title
    h(doc, 0, "Policy on Obstetric Care")
    p(doc, HN)

    # Document control
    h(doc, 1, "Document control")
    doc_ctrl(doc, "HCO/COP/POL/10", "Obstetric In-Charge")
    p(doc, "A blank marked ________ must be completed before issue.")

    # Statement of intent
    h(doc, 1, "Statement of intent")
    p(doc,
      "Obstetric care is organised and delivered safely, with high-risk cases "
      "identified and cared for by competent staff, ante-natal services following "
      "written guidance, and mothers treated with dignity and privacy throughout.")

    # 1. Purpose
    h(doc, 1, "1. Purpose")
    p(doc,
      f"This policy explains how {HN} organises and delivers safe obstetric care, "
      "identifies and manages high-risk obstetric cases, runs ante-natal services "
      "under written guidance, welcomes birth companions, treats mothers "
      "respectfully and confidentially, explains danger signs, monitors mother and "
      "baby through labour and after birth, and — where these services are within "
      "scope — provides neonatal intensive care and Assisted Reproductive "
      "Technology services in line with the law.")
    p(doc,
      "This policy does not cover general patient assessment, informed consent "
      "procedures, or newborn care in detail — those are covered in other hospital "
      "policies.")

    # 2. Scope
    h(doc, 1, "2. Scope")
    p(doc,
      f"This policy applies to all obstetric, nursing, and neonatal staff at {HN}. "
      f"Sections on neonatal intensive care (5.9) and Assisted Reproductive "
      f"Technology (5.10) apply only if {HN} provides those specific services.")

    # 3. Policy standards — two summary paragraphs only
    h(doc, 1, "3. Policy standards")
    p(doc,
      f"{HN} delivers obstetric care under written guidance covering "
      "pregnant-patient assessment, ante-natal, peri-natal, and post-natal care, "
      "identifies and either manages or refers high-risk obstetric cases using "
      "competent staff, runs ante-natal services with a complete ante-natal card "
      "for every patient, welcomes birth companions, treats mothers with privacy, "
      "respect, and confidentiality, explains danger signs, monitors labour and "
      "the post-natal period for key risks, and — where in scope — provides NICU "
      "care and follows Assisted Reproductive Technology law.")
    p(doc, "Staff follow the written guidance below and keep the records it requires.")

    # 4. Non-negotiable rules
    h(doc, 1, "4. Non-negotiable rules")
    lb(doc,
       "Do not deliver obstetric care without written guidance covering, at "
       "minimum, pregnant-patient assessment (nutrition, immunisation, education), "
       "ante-natal care, and peri-natal and post-natal care.")
    lb(doc,
       f"Do not leave whether {HN} provides high-risk obstetric care undefined, "
       "or leave what counts as \"high-risk\" undefined.")
    lb(doc,
       "Do not manage a high-risk obstetric case without competent workforce and "
       "facilities for both mother and neonate — refer proactively to another "
       "centre if high-risk care is out of scope.")
    lb(doc,
       "Do not let anyone other than a competent doctor or nurse — qualified, "
       "experienced, or trained — care for a high-risk obstetric case.")
    lb(doc,
       "Do not run ante-natal services without written guidance covering "
       "assessment, immunisation, diet counselling, and frequency of visits, and "
       "without a complete ante-natal card (or equivalent) for every patient.")
    lb(doc,
       "Do not skip monitoring foetal heart rate during labour, the progression "
       "of labour, and post-natal monitoring for post-partum haemorrhage.")
    lb(doc,
       f"Where {HN} provides high-risk obstetric care, do not do so without a "
       "NICU staffed with competent, qualified paediatric/neonatal personnel.")
    lb(doc,
       f"Where {HN} provides Assisted Reproductive Technology services, do not "
       "provide them without the required infrastructure, ART lab, operation "
       "theatre, competent team, consent process, and screening/SOPs, and without "
       "complying with the Assisted Reproductive Technology (Regulation) Act as "
       "amended.")

    # 5. What we do
    h(doc, 1, "5. What we do")

    h(doc, 2, "5.1 Organise obstetric care under written guidance")
    p(doc,
      "Written guidance based on standard treatment guidelines or sound clinical "
      "practice governs the organisation and delivery of obstetric care. At a "
      "minimum it covers three things: assessment of pregnant patients, including "
      "nutrition, immunisation, and education; ante-natal care guidelines; and "
      "peri-natal and post-natal care guidelines. "
      f"{HN} states in this guidance whether it provides care for high-risk "
      "obstetric cases, and defines what counts as high-risk, in line with best "
      "clinical practice. Staff are trained in managing obstetric emergencies, "
      "based on evidence-based clinical practice guidelines.")

    h(doc, 2, "5.2 Identify and manage high-risk obstetric cases")
    p(doc,
      "Ante-natal examinations guide the early identification of high-risk "
      "obstetric cases, which have more frequent check-ups where appropriate. "
      f"Where high-risk obstetric care isn't in {HN}'s scope, patients are "
      f"proactively referred to an appropriate centre. Where {HN} does provide "
      "high-risk care, it has a competent workforce and the facilities to care "
      "for both mothers and neonates in these cases.")

    h(doc, 2, "5.3 Use competent staff for high-risk cases")
    p(doc,
      "Persons caring for high-risk obstetric cases are competent — this includes "
      "both doctors and nursing staff, not doctors alone. Competency is based on "
      "qualification, experience, or training, or a combination of these.")

    h(doc, 2, "5.4 Provide ante-natal services under written guidance")
    p(doc,
      "Written guidance governs ante-natal services, covering at minimum four "
      "things: assessment, immunisation, diet counselling, and frequency of "
      "visits. Every patient has an ante-natal card (or an equivalent record), "
      "complete with her medical and obstetric history.")

    h(doc, 2, "5.5 Welcome birth companions")
    p(doc,
      "Staff encourage and welcome the presence of a birth companion during "
      "labour.")

    h(doc, 2, "5.6 Treat mothers with dignity and privacy")
    p(doc,
      "The Labour/Delivery Room has provisions for privacy — for example, "
      "curtains, partitions between tables, or non-see-through windows. Staff "
      "treat the pregnant woman and her companion cordially and respectfully, "
      "and her confidentiality is protected throughout her stay.")

    h(doc, 2, "5.7 Explain danger signs")
    p(doc,
      "A treating doctor, or a doctor from the treating team, explains danger "
      "signs and important care activities to the expecting mother and her "
      "companion.")

    h(doc, 2, "5.8 Assess and monitor through pregnancy and birth")
    p(doc,
      "Maternal nutrition is part of every obstetric patient's assessment, "
      "documented in her ante-natal card. It's preferable for a dietician to "
      "carry out this assessment. Appropriate peri-natal and post-natal "
      "monitoring is performed — at a minimum, this covers foetal heart rate "
      "during labour, the progression of labour, and post-natal monitoring for "
      "post-partum haemorrhage.")

    h(doc, 2, "5.9 Provide neonatal care for high-risk cases, where this service is offered")
    p(doc,
      f"Where {HN} provides care for high-risk obstetric cases, it also has a "
      "NICU — level I, II, or III, depending on scope — with appropriate "
      "equipment and staff, including a competent paediatrician or neonatologist "
      "and nurses, qualified and trained for neonatal care. This section does not "
      f"apply where high-risk obstetric care is outside {HN}'s scope.")

    h(doc, 2, "5.10 Follow Assisted Reproductive Technology law, where this service is offered")
    p(doc,
      f"Where {HN} provides Assisted Reproductive Technology (ART) services, it "
      "maintains the required infrastructure, an ART lab, an operation theatre, "
      "a competent ART team with trained counsellors and counselling techniques, "
      "a consent process, and patient screening and SOPs following national "
      "guidelines. All of this complies with the Assisted Reproductive Technology "
      f"(Regulation) Act as amended. This section does not apply where ART "
      f"services are outside {HN}'s scope.")

    # 6. Stop-work authority — text exactly as given in cop10_content.txt
    h(doc, 1, "6. Stop-work authority")
    p(doc,
      "Do not manage a high-risk obstetric emergency without trained staff or "
      "without applying defined admission/referral criteria; do not continue "
      "labour without foetal heart rate monitoring; do not provide Assisted "
      "Reproductive Technology services in breach of the ART Act as amended.")
    p(doc,
      "Stop-work applies to non-emergent obstetric care decisions and ART service "
      "delivery. A mother already in active labour or an obstetric emergency "
      "already in progress continues receiving the best available qualified care "
      "while the gap is escalated.")
    p(doc,
      "The person who stops tells the Obstetric In-Charge and the Medical "
      "Superintendent the same shift. Refusing unsafe obstetric or ART care is "
      "not a disciplinary matter.")

    # 7. Governance and responsibility
    h(doc, 1, "7. Governance and responsibility")
    gov_tbl(doc, [
        ("Medical Superintendent",
         "Accountable that this policy is followed and appropriately resourced; "
         "receives stop-work notifications the same shift."),
        ("Obstetric In-Charge",
         "Oversees delivery of obstetric care; ensures written guidance, trained "
         "staff, and required facilities are in place; receives stop-work "
         "notifications the same shift."),
        ("Treating obstetricians and doctors",
         "Apply high-risk identification criteria; explain danger signs to mothers "
         "and companions; document assessment and monitoring in ante-natal cards."),
        ("Obstetric and nursing staff",
         "Follow written guidance for ante-natal, peri-natal, and post-natal care; "
         "monitor foetal heart rate during labour, labour progression, and "
         "post-natal post-partum haemorrhage; support birth companions; maintain "
         "privacy and dignity in the Labour/Delivery Room."),
        ("ART team (where ART services are in scope)",
         "Follows ART Act requirements and national ART guidelines; maintains "
         "consent process, patient screening, and SOPs."),
        ("NICU staff (where high-risk obstetric care is in scope)",
         "Provides competent paediatric/neonatal care using qualified and trained "
         "personnel; maintains appropriate equipment."),
    ])

    # 8. Quality monitoring
    h(doc, 1, "8. Quality monitoring")
    mon_tbl(doc, [
        ("Written guidance",
         "Written guidance covering minimum three areas (pregnant-patient "
         "assessment, ante-natal care, peri-natal and post-natal care) in place "
         "and current"),
        ("High-risk identification",
         "High-risk cases identified via ante-natal examination; proactive referral "
         "where high-risk care is out of scope; \"high-risk\" criteria defined in "
         "written guidance"),
        ("Ante-natal written guidance",
         "Written guidance for ante-natal services covering all four minimum areas "
         "(assessment, immunisation, diet counselling, frequency of visits)"),
        ("Ante-natal card",
         "Complete ante-natal card (or equivalent) maintained for every patient, "
         "with medical and obstetric history including maternal nutrition "
         "assessment"),
        ("Labour and post-natal monitoring",
         "Foetal heart rate during labour, labour progression, and post-natal "
         "PPH monitoring performed and documented for every patient"),
        ("Privacy and dignity",
         "Privacy provisions in Labour/Delivery Room; mothers and companions "
         "treated respectfully; confidentiality maintained"),
        ("Danger signs",
         "Danger signs and important care activities explained to mother and "
         "companion by a treating doctor; documented"),
        ("NICU (where in scope)",
         "NICU with appropriate equipment and competent paediatric/neonatal staff "
         "in place where high-risk obstetric care is provided"),
        ("ART compliance (where in scope)",
         "Required ART infrastructure, lab, OT, team, consent, and screening/SOPs "
         "in line with the ART Act where ART services are provided"),
    ])

    # 9. Training and staff acknowledgement
    h(doc, 1, "9. Training and staff acknowledgement")
    p(doc,
      "All obstetric, nursing, and neonatal staff shall be familiar with the "
      "written guidance applicable to their work, including high-risk obstetric "
      "criteria, ante-natal care requirements, labour and post-natal monitoring "
      "requirements, privacy and dignity obligations, and — where applicable — "
      "NICU care standards and ART Act requirements.")
    p(doc,
      f"I have read the Policy on Obstetric Care of {HN}. I will follow the "
      "processes described.")
    sig_tbl(doc)

    # 10. Distribution
    h(doc, 1, "10. Distribution")
    p(doc,
      "This policy shall be available to all obstetric, nursing, and neonatal "
      f"staff at {HN}.")

    # 11. Abbreviations
    h(doc, 1, "11. Abbreviations")
    abbrev_tbl(doc, [
        ("ART",  "Assisted Reproductive Technology"),
        ("NICU", "Neonatal Intensive Care Unit"),
        ("OT",   "Operation Theatre"),
        ("PPH",  "Post-Partum Haemorrhage"),
        ("SOP",  "Standard Operating Procedure"),
    ])

    # 12. Traceability table
    h(doc, 1, "12. Traceability table")
    p(doc,
      "This table is an index. It is not how the policy is organised. An asterisk "
      "in the Level column means documentation of the process is required.")
    tr = tbl(doc, 12, 3)
    for ci, hdr in enumerate(("Objective Element", "Level", "Traceability to this policy")):
        tr.cell(0, ci).text = hdr
    trace_rows = [
        ("COP.10.a", "Commitment*",
         "Section 5.1 addresses written guidance covering at minimum three areas: "
         "assessment of pregnant patients including nutrition, immunisation, and "
         "education; ante-natal care guidelines; and peri-natal and post-natal "
         "care guidelines."),
        ("COP.10.b", "Commitment",
         "Sections 5.1 and 5.2 address whether high-risk obstetric care is "
         "provided, the definition of high-risk, early identification via "
         "ante-natal examination, more frequent check-ups for high-risk cases, "
         "proactive referral where high-risk care is out of scope, and staff "
         "training in obstetric emergency management."),
        ("COP.10.c", "Commitment",
         "Section 5.3 addresses competency requirements for persons caring for "
         "high-risk obstetric cases — both doctors and nursing staff — based on "
         "qualification, experience, or training; and proactive referral when "
         "high-risk care is outside scope."),
        ("COP.10.d", "Commitment*",
         "Section 5.4 addresses written guidance for ante-natal services covering "
         "at minimum four areas: assessment, immunisation, diet counselling, and "
         "frequency of visits. Also addresses the mandatory ante-natal card (or "
         "equivalent) with complete medical and obstetric history for every "
         "patient."),
        ("COP.10.e", "Achievement",
         "Section 5.5 addresses welcoming and encouraging the presence of a birth "
         "companion during labour."),
        ("COP.10.f", "Commitment",
         "Section 5.6 addresses privacy provisions in the Labour/Delivery Room "
         "(e.g., curtains, partitions, non-see-through windows), cordial and "
         "respectful treatment of the pregnant woman and companion, and "
         "confidentiality throughout her stay."),
        ("COP.10.g", "Commitment",
         "Section 5.7 addresses explanation of danger signs and important care "
         "activities to the expecting mother and companion by the treating doctor "
         "or a doctor from the treating team."),
        ("COP.10.h", "Commitment",
         "Section 5.8 addresses maternal nutrition as a mandatory component of "
         "every obstetric patient's assessment, documented in the ante-natal card, "
         "with a dietician performing the assessment where possible."),
        ("COP.10.i", "Commitment",
         "Section 5.8 addresses appropriate peri-natal and post-natal monitoring, "
         "covering at minimum three items: foetal heart rate during labour, the "
         "progression of labour, and post-natal monitoring for post-partum "
         "haemorrhage."),
        ("COP.10.j", "Commitment",
         "Section 5.9 addresses NICU provision (level I, II, or III) with "
         "appropriate equipment and competent paediatric/neonatal personnel where "
         "high-risk obstetric care is in scope. Explicitly scope-conditional: does "
         "not apply where high-risk obstetric care is outside scope."),
        ("COP.10.k", "Commitment",
         "Section 5.10 addresses ART service provision including required "
         "infrastructure, ART lab, operation theatre, competent ART team with "
         "trained counsellors, consent process, patient screening, and SOPs "
         "following national guidelines, in compliance with the Assisted "
         "Reproductive Technology (Regulation) Act as amended. Explicitly "
         "scope-conditional: does not apply where ART services are outside scope."),
    ]
    for ri, (oe, lvl, txt) in enumerate(trace_rows, 1):
        tr.cell(ri, 0).text = oe
        tr.cell(ri, 1).text = lvl
        tr.cell(ri, 2).text = txt

    # 13. Required Records/Evidence Checklist
    h(doc, 1, "13. Required Records/Evidence Checklist")

    h(doc, 2, "Written guidance — obstetric care")
    lb(doc,
       "Written guidance based on standard treatment guidelines or sound clinical "
       "practice covering pregnant-patient assessment (nutrition, immunisation, "
       "education), ante-natal care guidelines, and peri-natal and post-natal "
       "care guidelines.")
    lb(doc,
       "Written guidance stating whether high-risk obstetric care is provided and "
       "defining what counts as high-risk.")
    lb(doc,
       "Evidence of staff training in managing obstetric emergencies.")

    h(doc, 2, "High-risk identification and management")
    lb(doc,
       "Records showing high-risk cases identified via ante-natal examinations.")
    lb(doc,
       "Evidence of more frequent check-ups for high-risk cases where appropriate.")
    lb(doc,
       "Referral records for high-risk cases where high-risk care is out of scope.")
    lb(doc,
       "Evidence of competent workforce and facilities for high-risk care "
       "where provided.")

    h(doc, 2, "Ante-natal services")
    lb(doc,
       "Written guidance for ante-natal services covering assessment, "
       "immunisation, diet counselling, and frequency of visits.")
    lb(doc,
       "Complete ante-natal card (or equivalent) for every patient, with medical "
       "and obstetric history including maternal nutrition assessment.")

    h(doc, 2, "Privacy, dignity, and danger signs")
    lb(doc,
       "Evidence of privacy provisions in the Labour/Delivery Room (e.g., "
       "curtains, partitions, non-see-through windows).")
    lb(doc,
       "Records of danger signs explained to mothers and companions by a treating "
       "doctor.")

    h(doc, 2, "Labour and post-natal monitoring")
    lb(doc,
       "Records of foetal heart rate monitoring during labour.")
    lb(doc,
       "Records of labour progression monitoring.")
    lb(doc,
       "Records of post-natal monitoring for post-partum haemorrhage.")

    h(doc, 2, "NICU (where high-risk obstetric care is in scope)")
    lb(doc,
       "Evidence of NICU provision (level I, II, or III) with appropriate "
       "equipment and competent paediatric/neonatal staff.")

    h(doc, 2, "ART services (where in scope)")
    lb(doc,
       "Evidence of required ART infrastructure, ART lab, and operation theatre.")
    lb(doc,
       "Evidence of competent ART team including trained counsellors.")
    lb(doc,
       "ART consent process documentation and patient screening records.")
    lb(doc,
       "SOPs following national guidelines for ART services.")
    lb(doc,
       "Evidence of compliance with the Assisted Reproductive Technology "
       "(Regulation) Act as amended.")

    # 14. References
    h(doc, 1, "14. References")
    ln(doc,
       "National Accreditation Board for Hospitals and Healthcare Providers. NABH "
       "Accreditation Standards for Hospitals, 6th Edition. COP.10.")
    ln(doc, "Guidebook interpretation supplied for COP.10.a through COP.10.k.")
    ln(doc,
       "Assisted Reproductive Technology (Regulation) Act as amended (reference "
       "for COP.10.k ART compliance requirements).")

    # Disclaimer
    h(doc, 1, "Disclaimer")
    p(doc,
      "This policy reorganises the supplied COP.10 objective elements and approved "
      "plain-language content into policy format. The stop-work text is reproduced "
      "exactly as supplied. Mandatory requirements and their stated modal strength "
      "have been retained. The three minimum written-guidance areas (COP.10.a), "
      "four minimum ante-natal written-guidance areas (COP.10.d), three minimum "
      "monitoring items (COP.10.i), and six ART scope elements (infrastructure, "
      "ART lab, operation theatre, competent team, consent process, "
      "screening/SOPs) are preserved exactly as stated. COP.10.j (NICU) and "
      "COP.10.k (ART) are explicitly presented as scope-conditional, not universal "
      "requirements.")

    save_and_verify(doc, "HCO_COP_10_v2_REWRITE_DRAFT.docx")


# ══════════════════════════════════════════════════════════════════════════════
# COP.11 — Paediatric Services   (stop-work YES)
# Content: approved plain-language content (cop11_content.txt).
# Structure: Section 3 two summary paragraphs; full detail Section 5.
# Stars: a*, b*, f* | Excellence: h (scope-conditional)
# CORE: none
# Exact quantities: four assessment components (growth, developmental,
#   immunisation, nutritional); "actively promotes breastfeeding" exact phrase;
#   testing interval for rapid-response is hospital-defined — NOT invented;
#   four mandatory nutrition education sub-topics.
# ══════════════════════════════════════════════════════════════════════════════

def gen_cop11():
    doc = Document()

    # Title
    h(doc, 0, "Policy on Paediatric Services")
    p(doc, HN)

    # Document control
    h(doc, 1, "Document control")
    doc_ctrl(doc, "HCO/COP/POL/11", "Paediatric In-Charge")
    p(doc, "A blank marked ________ must be completed before issue.")

    # Statement of intent
    h(doc, 1, "Statement of intent")
    p(doc,
      "Paediatric services are organised and delivered safely, by staff with "
      "age-specific competency, with active safeguards against child and neonate "
      "abduction and abuse.")

    # 1. Purpose
    h(doc, 1, "1. Purpose")
    p(doc,
      f"This policy explains how {HN} organises and delivers safe paediatric "
      "care, provides neonatal care in line with national and international "
      "guidelines, ensures staff have age-specific competency, prevents child and "
      "neonate abduction and abuse, assesses children's growth and development, "
      "educates families, and — where in scope — provides adolescent-friendly "
      "health services.")
    p(doc,
      "This policy does not cover general infection control, general patient "
      "assessment, or informed consent in detail — those are covered in other "
      "hospital policies.")

    # 2. Scope
    h(doc, 1, "2. Scope")
    p(doc,
      f"This policy applies to all doctors, nursing staff, and other personnel "
      f"caring for children and neonates at {HN}. If {HN} does not provide "
      "paediatric services, this policy does not apply.")

    # 3. Policy standards — two summary paragraphs only
    h(doc, 1, "3. Policy standards")
    p(doc,
      f"{HN} delivers paediatric care under written guidance that defines its "
      "scope, provides neonatal care in line with national and international "
      "guidelines, ensures doctors and nurses have age-specific competency, gives "
      "children adequate amenities, assesses growth, development, immunisation, "
      "and nutrition, actively prevents child and neonate abduction and abuse, "
      "educates families, and — where offered — provides adolescent-friendly "
      "health services.")
    p(doc, "Staff follow the written guidance below and keep the records it requires.")

    # 4. Non-negotiable rules
    h(doc, 1, "4. Non-negotiable rules")
    lb(doc,
       "Do not provide paediatric services without written guidance covering, "
       "at minimum, patient assessment, organisation of care, and addressing "
       "special needs, and defining the scope of paediatric services offered.")
    lb(doc,
       "Do not deliver neonatal care that isn't in line with national and "
       "international guidelines, and do not fail to actively promote "
       "breastfeeding.")
    lb(doc,
       "Do not let anyone without age-specific competency — by qualification, "
       "experience, or training — care for children, whether doctor or nurse.")
    lb(doc,
       "Do not skip any of the four paediatric assessment components: growth, "
       "development, immunisation, and nutrition.")
    lb(doc,
       "Do not operate without a defined, tested rapid-response process for child "
       "or neonate abduction and abuse, adequate security and surveillance, and "
       "trained staff.")
    lb(doc,
       "Do not educate families about nutrition, immunisation, or safe parenting "
       "in a language they don't understand, or skip any of the mandatory "
       "nutrition sub-topics — breastfeeding, weaning, protein-energy "
       "malnutrition, and childhood obesity.")

    # 5. What we do
    h(doc, 1, "5. What we do")

    h(doc, 2, "5.1 Organise paediatric care under written guidance")
    p(doc,
      "Written guidance based on standard treatment guidelines or sound clinical "
      "practice governs paediatric care. At a minimum, it covers three things: "
      "assessment of paediatric patients, organisation of care, and addressing "
      f"special needs. It also defines the scope of {HN}'s paediatric services — "
      "for example, well-baby clinics, different NICU levels, or a PICU. These "
      f"are examples; {HN} defines its own scope.")

    h(doc, 2, "5.2 Deliver neonatal care to guideline standards")
    p(doc,
      "Written guidance for neonatal care is in line with guidelines from national "
      f"and international bodies. {HN} actively promotes breastfeeding practices.")

    h(doc, 2, "5.3 Use age-specific competency")
    p(doc,
      "Both doctors and nursing staff caring for children have age-specific "
      "competency, based on qualification, experience, or training, or a "
      "combination of these.")

    h(doc, 2, "5.4 Provide adequate amenities")
    p(doc,
      "Adequate amenities for the care of infants and children are available — "
      "for example, a breastfeeding room or a play area. These are examples; "
      f"{HN} determines what's appropriate for its paediatric population.")

    h(doc, 2, "5.5 Carry out a complete paediatric assessment")
    p(doc,
      "Paediatric assessment covers four things: growth, developmental, "
      "immunisation, and nutritional assessment, done using appropriate tools — "
      "preferably validated ones — and documented. Growth charts and immunisation "
      "records are maintained and kept up to date.")

    h(doc, 2, "5.6 Prevent child and neonate abduction and abuse")
    p(doc,
      f"Written guidance directs how {HN} prevents child and neonate abduction "
      "and abuse. Adequate security and surveillance is in place — for example, "
      "CCTV cameras — and there's a defined rapid-response process for any "
      "eventuality. This process is tested at intervals "
      f"{HN} defines, using a table-top exercise or a mock drill. Staff are "
      "trained in prevention and rapid response, and know how to recognise and "
      "escalate suspected child abuse.")

    h(doc, 2, "5.7 Educate families")
    p(doc,
      "Families are educated about nutrition, immunisation, and safe parenting, "
      "in a language they understand — using educational material if helpful. "
      "Nutrition education covers breastfeeding, weaning, and aspects of "
      "malnutrition, including protein-energy malnutrition and childhood obesity. "
      "Parents are also given appropriate at-home growth-monitoring information.")

    h(doc, 2, "5.8 Provide adolescent-friendly health care, where this service is offered")
    p(doc,
      f"Where {HN} offers adolescent health services, these are need-based — "
      "preventive, curative, and counselling — delivered with a respectful, "
      "non-judgemental attitude throughout registration, reception, consultation, "
      "and treatment, with confidentiality maintained throughout. Services could "
      "include managing physical and mental health conditions, counselling "
      "(including premarital counselling and anti-substance-abuse counselling), "
      "preventive health check-ups, immunisation such as HPV vaccination, and "
      "health education on topics like development, nutrition, sleep, and safe "
      "sexual practices. These are examples; the specific service mix is "
      f"{HN}'s own decision based on need.")

    # 6. Stop-work authority — text exactly as given in cop11_content.txt
    h(doc, 1, "6. Stop-work authority")
    p(doc,
      "Do not care for a child or neonate without an active, tested process to "
      "prevent abduction or abuse in place, or without staff who have age-specific "
      "competency for that patient.")
    p(doc,
      "Stop-work applies to routine paediatric admission and ongoing care "
      "arrangements. A child already receiving emergency care continues to receive "
      "it while the gap is escalated.")
    p(doc,
      "The person who stops tells the Paediatric In-Charge and the Medical "
      "Superintendent the same shift. Refusing to place a child in an unsafe care "
      "arrangement is not a disciplinary matter.")

    # 7. Governance and responsibility
    h(doc, 1, "7. Governance and responsibility")
    gov_tbl(doc, [
        ("Medical Superintendent",
         "Accountable that this policy is followed and appropriately resourced; "
         "receives stop-work notifications the same shift."),
        ("Paediatric In-Charge",
         "Oversees paediatric services; ensures written guidance, age-specific "
         "competency, four-component assessment, abduction/abuse prevention, and "
         "family education are in place; receives stop-work notifications the "
         "same shift."),
        ("Paediatric doctors and nursing staff",
         "Apply age-specific competency; carry out all four paediatric assessment "
         "components (growth, developmental, immunisation, nutritional); follow "
         "neonatal care guidelines; actively promote breastfeeding; educate "
         "families in a language they understand; recognise and escalate suspected "
         "child abuse."),
        ("Security/surveillance staff",
         "Maintain CCTV and other security measures; respond to rapid-response "
         "process activations for abduction or abuse; participate in mock drills "
         "and table-top exercises at intervals defined by the organisation."),
    ])

    # 8. Quality monitoring
    h(doc, 1, "8. Quality monitoring")
    mon_tbl(doc, [
        ("Written guidance",
         "Written guidance covering assessment, organisation of care, and "
         "addressing special needs; scope of paediatric services defined"),
        ("Neonatal care guideline alignment",
         "Written guidance for neonatal care in line with national and "
         "international guidelines; active promotion of breastfeeding practices"),
        ("Age-specific competency",
         "Doctors and nursing staff caring for children have documented "
         "age-specific competency by qualification, experience, or training"),
        ("Paediatric assessment completeness",
         "All four assessment components completed and documented for every "
         "child: growth, developmental, immunisation, and nutritional assessment"),
        ("Abduction/abuse prevention",
         "Written guidance in place; adequate security and surveillance; "
         "rapid-response process defined and tested at organisation-defined "
         "intervals; staff trained in prevention, response, and abuse recognition"),
        ("Family education",
         "Families educated in a language they understand; all four mandatory "
         "nutrition sub-topics covered (breastfeeding, weaning, protein-energy "
         "malnutrition, childhood obesity); at-home growth-monitoring information "
         "provided"),
        ("Adolescent services (where in scope)",
         "Need-based adolescent health services delivered with respectful, "
         "non-judgemental attitude and confidentiality where offered"),
    ])

    # 9. Training and staff acknowledgement
    h(doc, 1, "9. Training and staff acknowledgement")
    p(doc,
      "All doctors, nursing staff, and other personnel caring for children and "
      "neonates shall be familiar with the written guidance applicable to their "
      "work, including neonatal care guidelines, age-specific competency "
      "requirements, the four-component paediatric assessment, abduction and abuse "
      "prevention procedures, the rapid-response process, and family education "
      "requirements.")
    p(doc,
      f"I have read the Policy on Paediatric Services of {HN}. I will follow "
      "the processes described.")
    sig_tbl(doc)

    # 10. Distribution
    h(doc, 1, "10. Distribution")
    p(doc,
      "This policy shall be available to all doctors, nursing staff, and other "
      f"personnel caring for children and neonates at {HN}.")

    # 11. Abbreviations
    h(doc, 1, "11. Abbreviations")
    abbrev_tbl(doc, [
        ("CCTV", "Closed-Circuit Television"),
        ("HPV",  "Human Papillomavirus"),
        ("NICU", "Neonatal Intensive Care Unit"),
        ("PICU", "Paediatric Intensive Care Unit"),
    ])

    # 12. Traceability table
    h(doc, 1, "12. Traceability table")
    p(doc,
      "This table is an index. It is not how the policy is organised. An asterisk "
      "in the Level column means documentation of the process is required.")
    tr = tbl(doc, 9, 3)
    for ci, hdr in enumerate(("Objective Element", "Level", "Traceability to this policy")):
        tr.cell(0, ci).text = hdr
    trace_rows = [
        ("COP.11.a", "Commitment*",
         "Section 5.1 addresses written guidance for paediatric care based on "
         "standard treatment guidelines or sound clinical practice, covering at "
         "minimum assessment, organisation of care, and addressing special needs; "
         "and defining the scope of paediatric services offered."),
        ("COP.11.b", "Commitment*",
         "Section 5.2 addresses written guidance for neonatal care in line with "
         "national and international guidelines, and the active promotion of "
         "breastfeeding practices. The phrase \"actively promotes\" is preserved "
         "exactly from the source — not softened."),
        ("COP.11.c", "Commitment",
         "Section 5.3 addresses age-specific competency for both doctors and "
         "nursing staff caring for children, based on qualification, experience, "
         "or training, or a combination of these."),
        ("COP.11.d", "Commitment",
         "Section 5.4 addresses adequate amenities for infants and children "
         "(e.g., breastfeeding room, play area); specific amenities are the "
         "organisation's own determination."),
        ("COP.11.e", "Commitment",
         "Section 5.5 addresses the four-component paediatric assessment: growth, "
         "developmental, immunisation, and nutritional assessment, using "
         "appropriate (preferably validated) tools and documented. Growth charts "
         "and immunisation records maintained and updated."),
        ("COP.11.f", "Commitment*",
         "Section 5.6 addresses written guidance for preventing child and neonate "
         "abduction and abuse, adequate security and surveillance (e.g., CCTV), a "
         "defined rapid-response process tested at intervals the organisation "
         "defines using table-top exercise or mock drill, and trained staff who "
         "can recognise and escalate suspected child abuse. No specific testing "
         "interval has been invented."),
        ("COP.11.g", "Commitment",
         "Section 5.7 addresses family education on nutrition, immunisation, and "
         "safe parenting in a language the family understands. Nutrition education "
         "covers four mandatory sub-topics: breastfeeding, weaning, protein-energy "
         "malnutrition, and childhood obesity. At-home growth-monitoring "
         "information is also provided."),
        ("COP.11.h", "Excellence",
         "Section 5.8 addresses adolescent-friendly health services where offered: "
         "need-based preventive, curative, and counselling services with a "
         "respectful, non-judgemental attitude and maintained confidentiality. "
         "Scope-conditional: applies only where the organisation offers adolescent "
         "health services."),
    ]
    for ri, (oe, lvl, txt) in enumerate(trace_rows, 1):
        tr.cell(ri, 0).text = oe
        tr.cell(ri, 1).text = lvl
        tr.cell(ri, 2).text = txt

    # 13. Required Records/Evidence Checklist
    h(doc, 1, "13. Required Records/Evidence Checklist")

    h(doc, 2, "Written guidance — paediatric services")
    lb(doc,
       "Written guidance for paediatric care covering assessment, organisation "
       "of care, and addressing special needs.")
    lb(doc,
       "Defined scope of paediatric services offered by the organisation.")

    h(doc, 2, "Neonatal care and breastfeeding")
    lb(doc,
       "Written guidance for neonatal care in line with national and "
       "international guidelines.")
    lb(doc,
       "Evidence that breastfeeding is actively promoted.")

    h(doc, 2, "Age-specific competency")
    lb(doc,
       "Evidence of age-specific competency for doctors and nursing staff caring "
       "for children (by qualification, experience, or training).")

    h(doc, 2, "Amenities")
    lb(doc,
       "Evidence of adequate amenities for infants and children appropriate to "
       "the organisation's paediatric population.")

    h(doc, 2, "Paediatric assessment records")
    lb(doc,
       "Assessment records documenting all four components for every child: "
       "growth, developmental, immunisation, and nutritional assessment.")
    lb(doc, "Growth charts maintained and kept up to date.")
    lb(doc, "Immunisation records maintained and kept up to date.")

    h(doc, 2, "Abduction and abuse prevention")
    lb(doc,
       "Written guidance for preventing child and neonate abduction and abuse.")
    lb(doc,
       "Evidence of adequate security and surveillance (e.g., CCTV).")
    lb(doc,
       "Defined rapid-response process; records of testing at "
       "organisation-defined intervals (table-top exercise or mock drill).")
    lb(doc,
       "Evidence of staff training in prevention, rapid response, and recognising "
       "and escalating suspected child abuse.")

    h(doc, 2, "Family education records")
    lb(doc,
       "Records of family education on nutrition, immunisation, and safe "
       "parenting, in a language the family understands.")
    lb(doc,
       "Evidence that all four mandatory nutrition sub-topics were covered: "
       "breastfeeding, weaning, protein-energy malnutrition, and childhood "
       "obesity.")
    lb(doc,
       "Evidence that at-home growth-monitoring information was provided to "
       "parents.")

    h(doc, 2, "Adolescent services (where in scope)")
    lb(doc,
       "Evidence of need-based adolescent health services with a respectful, "
       "non-judgemental approach and maintained confidentiality where offered.")

    # 14. References
    h(doc, 1, "14. References")
    ln(doc,
       "National Accreditation Board for Hospitals and Healthcare Providers. NABH "
       "Accreditation Standards for Hospitals, 6th Edition. COP.11.")
    ln(doc, "Guidebook interpretation supplied for COP.11.a through COP.11.h.")
    ln(doc,
       "National and international guidelines for neonatal care (reference for "
       "COP.11.b written guidance).")

    # Disclaimer
    h(doc, 1, "Disclaimer")
    p(doc,
      "This policy reorganises the supplied COP.11 objective elements and approved "
      "plain-language content into policy format. The stop-work text is reproduced "
      "exactly as supplied. Mandatory requirements and their stated modal strength "
      "have been retained. The phrase \"actively promotes breastfeeding\" is "
      "preserved exactly — it has not been softened to \"encourages\" or any "
      "weaker formulation. The four paediatric assessment components (growth, "
      "developmental, immunisation, nutritional) are preserved exactly. The "
      "testing interval for the abduction/abuse rapid-response process is "
      "hospital-defined — no specific interval has been invented. The four "
      "mandatory family-education nutrition sub-topics (breastfeeding, weaning, "
      "protein-energy malnutrition, childhood obesity) are preserved. "
      "COP.11.h (adolescent services, Excellence level) is presented as "
      "scope-conditional.")

    save_and_verify(doc, "HCO_COP_11_v2_REWRITE_DRAFT.docx")


# ══════════════════════════════════════════════════════════════════════════════
# COP.12 — Procedural Sedation   (stop-work YES)
# Content: approved plain-language content (cop12_content.txt).
# Structure: Section 3 two summary paragraphs; full detail Section 5.
# Stars: a*, g* | CORE: none
# Exact quantities preserved:
#   5.1: six written-guidance areas (areas/orders/pre-assessment/intra/post/dc)
#   5.3: technician NEVER administers — absolute prohibition, no exception
#   5.4: monitor ALWAYS different from procedure performer — absolute, no exception
#   5.5: six intra-procedure parameters (incl. cardiac rhythm)
#   5.6: five post-sedation parameters (NO cardiac rhythm — different list)
#   5.8: five emergency equipment items
# ══════════════════════════════════════════════════════════════════════════════

def gen_cop12():
    doc = Document()

    # Title
    h(doc, 0, "Policy on Procedural Sedation")
    p(doc, HN)

    # Document control
    h(doc, 1, "Document control")
    doc_ctrl(doc, "HCO/COP/POL/12", "Anaesthesia/Sedation In-Charge")
    p(doc, "A blank marked ________ must be completed before issue.")

    # Statement of intent
    h(doc, 1, "Statement of intent")
    p(doc,
      "Procedural sedation is administered consistently across the hospital, by "
      "competent staff, with continuous monitoring and equipment ready for a deeper "
      "level of sedation than planned.")

    # 1. Purpose
    h(doc, 1, "1. Purpose")
    p(doc,
      f"This policy explains how {HN} administers procedural sedation "
      "consistently, obtains informed consent, ensures only competent staff "
      "administer it, keeps the person monitoring separate from the person "
      "performing the procedure, monitors patients during and after sedation, "
      "applies discharge criteria, and keeps equipment ready for deeper-than-"
      "intended sedation.")
    p(doc,
      "This policy does not cover general informed consent procedures or general "
      "anaesthesia in detail — those are covered in other hospital policies.")

    # 2. Scope
    h(doc, 1, "2. Scope")
    p(doc,
      f"This policy applies to all staff who administer, monitor, or are otherwise "
      f"involved in procedural sedation at {HN}.")

    # 3. Policy standards — two summary paragraphs only
    h(doc, 1, "3. Policy standards")
    p(doc,
      f"{HN} administers procedural sedation under written guidance that is "
      "standardised across the organisation, obtains informed consent covering "
      "risks, benefits, and alternatives, uses only competent staff to administer "
      "sedation with the monitoring role kept separate from the procedure role, "
      "monitors a defined set of parameters during and after the procedure, applies "
      "documented discharge criteria, and keeps emergency equipment ready in case a "
      "patient goes into a deeper level of sedation than intended.")
    p(doc, "Staff follow the written guidance below and keep the records it requires.")

    # 4. Non-negotiable rules
    h(doc, 1, "4. Non-negotiable rules")
    lb(doc,
       "Do not administer procedural sedation without written guidance covering, "
       "at minimum, where sedation is given, how orders are written, pre-procedure "
       "assessment, during-procedure monitoring, after-procedure monitoring, and "
       "discharge/transfer-out criteria.")
    lb(doc,
       "Do not administer sedation without a documented pre-sedation assessment.")
    lb(doc,
       "Do not administer sedation without informed consent covering risks, "
       "benefits, and alternatives, taken by the person administering sedation or "
       "a doctor on the sedation team.")
    lb(doc,
       "Do not let a technician administer procedural sedation, under any "
       "circumstances.")
    lb(doc,
       "Where the parenteral route is used, do not let a nurse administer sedation "
       "without a doctor's supervision.")
    lb(doc,
       "Do not let the same person both perform the procedure and monitor the "
       "sedation.")
    lb(doc,
       "Do not skip any of the six required intra-procedure monitoring parameters: "
       "heart rate, cardiac rhythm, respiratory rate, blood pressure, oxygen "
       "saturation, and level of sedation.")
    lb(doc,
       "Do not skip any of the five required post-sedation monitoring parameters: "
       "heart rate, respiratory rate, blood pressure, oxygen saturation, and level "
       "of sedation.")
    lb(doc,
       "Do not discharge a patient from the observation/recovery area without a "
       "qualified individual applying documented discharge criteria.")
    lb(doc,
       "Do not administer sedation in a room without all five emergency equipment "
       "items in working order: emergency resuscitation equipment, suction, "
       "advanced airway equipment, positive pressure ventilation, and supplemental "
       "oxygen.")

    # 5. What we do
    h(doc, 1, "5. What we do")

    h(doc, 2, "5.1 Follow written guidance for consistent sedation")
    p(doc,
      "Written guidance based on standard treatment guidelines or sound clinical "
      "practice governs procedural sedation. At a minimum, it covers six things: "
      "identification of the areas where sedation is administered and procedures "
      "performed, the mechanism for writing orders, pre-procedure assessment, "
      "monitoring during the procedure, monitoring after the procedure, and "
      "discharge or transfer-out criteria. A pre-sedation assessment is performed "
      "and documented for every patient, evaluating the risk and appropriateness "
      "of sedation, with its scope and content based on professional guidelines. "
      f"Sedation is administered in a standardised way across the whole "
      "organisation.")

    h(doc, 2, "5.2 Obtain informed consent")
    p(doc,
      "Informed consent for procedural sedation is taken by the person "
      "administering sedation, or by a doctor on the sedation team. The patient "
      "is informed of the risks, benefits, and alternatives, and the patient and "
      "family are educated about post-procedural pain relief. (Related "
      "requirements are covered in the hospital's other policies.)")

    h(doc, 2, "5.3 Use only competent, trained personnel")
    p(doc,
      "Where the parenteral route is used, sedation is administered by a doctor, "
      "or by a nurse under a doctor's supervision. A technician never administers "
      "sedation, without exception.")

    h(doc, 2, "5.4 Keep the monitoring role separate from the procedure role")
    p(doc,
      "The person monitoring a sedated patient is always different from the person "
      "performing the procedure, without exception. The monitoring person is "
      "trained to detect abnormalities in monitoring parameters and to recognise "
      "apnoea and airway obstruction.")

    h(doc, 2, "5.5 Monitor patients during the procedure")
    p(doc,
      "Intra-procedure monitoring covers, at a minimum, six parameters: heart "
      "rate, cardiac rhythm, respiratory rate, blood pressure, oxygen saturation, "
      "and level of sedation. Additional parameters may be monitored case by case. "
      "Monitored parameters are documented; routine cardiac rhythm need not be "
      "documented continuously on a monitor, but any rhythm abnormality is "
      "documented.")

    h(doc, 2, "5.6 Monitor patients after the procedure")
    p(doc,
      "After sedation, the patient's vitals are monitored at intervals "
      f"{HN} decides, until the patient has fully recovered. At a minimum, this "
      "covers five parameters: heart rate, respiratory rate, blood pressure, "
      "oxygen saturation, and level of sedation. The extent and duration of "
      "monitoring can be adjusted for the complexity of the procedure and the "
      "patient's other health conditions. Level of sedation could be monitored "
      "using a checklist covering mild, moderate, and deep sedation. Monitoring "
      "is documented.")

    h(doc, 2, "5.7 Apply discharge criteria")
    p(doc,
      "Discharge or transfer-out from the observation/recovery area follows "
      f"criteria {HN} has developed and documented, in line with physiologic "
      "parameters and sound clinical practice. A qualified individual applies "
      "these criteria, and that application is documented.")

    h(doc, 2, "5.8 Keep equipment ready for deeper sedation than intended")
    p(doc,
      "The room where procedural sedation is given has, in working order: "
      "equipment for emergency resuscitation, suction, advanced airway equipment, "
      "positive pressure ventilation, and supplemental oxygen. Emergency equipment "
      "and supplies are immediately available and suited to the type of sedation, "
      "the patient's age, and the patient's medical condition. A person trained in "
      f"airway management, or an anaesthesiologist, is available in {HN} and able "
      "to respond quickly to the area.")

    # 6. Stop-work authority — text exactly as given in cop12_content.txt
    h(doc, 1, "6. Stop-work authority")
    p(doc,
      "Do not start procedural sedation without consent, without the required "
      "monitoring in place, or without a competent person qualified to administer "
      "it; do not administer sedation with a technician, or without a separate "
      "person monitoring, or where the sedation room lacks emergency equipment for "
      "a deeper level of sedation than intended.")
    p(doc,
      "Stop-work applies to the start of a planned sedation procedure. Where "
      "sedation is already under way and a patient's condition changes "
      "unexpectedly, the response continues with the best available qualified help "
      "while the gap is escalated.")
    p(doc,
      "The person who stops tells the Anaesthesia/Sedation In-Charge and the "
      "Medical Superintendent the same shift. Refusing to proceed with unsafe "
      "sedation is not a disciplinary matter.")

    # 7. Governance and responsibility
    h(doc, 1, "7. Governance and responsibility")
    gov_tbl(doc, [
        ("Medical Superintendent",
         "Accountable that this policy is followed and appropriately resourced; "
         "receives stop-work notifications the same shift."),
        ("Anaesthesia/Sedation In-Charge",
         "Oversees the sedation programme; ensures written guidance, competency "
         "requirements, role-separation rule, and emergency equipment standards "
         "are in place; receives stop-work notifications the same shift."),
        ("Doctors administering sedation",
         "Perform and document pre-sedation assessment; obtain informed consent "
         "covering risks, benefits, and alternatives; administer sedation (or "
         "supervise nurses where parenteral route is used); apply and document "
         "discharge criteria."),
        ("Monitoring personnel",
         "Maintain role separation from procedure performer at all times; monitor "
         "all six intra-procedure and five post-sedation parameters; detect "
         "abnormalities, apnoea, and airway obstruction; document monitored "
         "parameters."),
        ("Nursing staff",
         "May administer parenteral sedation only under direct doctor supervision; "
         "never administer sedation independently via the parenteral route; a "
         "technician never administers sedation."),
    ])

    # 8. Quality monitoring
    h(doc, 1, "8. Quality monitoring")
    mon_tbl(doc, [
        ("Written guidance",
         "Written guidance covering all six required areas in place and current; "
         "pre-sedation assessment performed and documented for every patient"),
        ("Informed consent",
         "Consent records covering risks, benefits, and alternatives; taken by "
         "administering doctor or sedation team doctor"),
        ("Staff competency — no technician",
         "Parenteral sedation administered only by a doctor or nurse under "
         "supervision; no technician administers sedation under any circumstances"),
        ("Role separation",
         "Person monitoring is always different from procedure performer; "
         "monitoring person's training in abnormality detection documented"),
        ("Intra-procedure monitoring",
         "All six parameters monitored and documented for every procedure: heart "
         "rate, cardiac rhythm, respiratory rate, blood pressure, oxygen "
         "saturation, level of sedation"),
        ("Post-sedation monitoring",
         "All five parameters monitored at organisation-defined intervals until "
         "full recovery: heart rate, respiratory rate, blood pressure, oxygen "
         "saturation, level of sedation; monitoring documented"),
        ("Discharge criteria",
         "Documented discharge/transfer-out criteria applied by a qualified "
         "individual; application documented for every patient"),
        ("Emergency equipment",
         "All five items in working order checked and recorded: emergency "
         "resuscitation equipment, suction, advanced airway equipment, positive "
         "pressure ventilation, supplemental oxygen"),
    ])

    # 9. Training and staff acknowledgement
    h(doc, 1, "9. Training and staff acknowledgement")
    p(doc,
      "All staff who administer, monitor, or are otherwise involved in procedural "
      "sedation shall be familiar with the written guidance applicable to their "
      "work, including pre-sedation assessment, informed consent requirements, the "
      "absolute prohibition on technicians administering sedation, the role-"
      "separation rule, the six intra-procedure and five post-sedation monitoring "
      "parameters, discharge criteria, and emergency equipment requirements.")
    p(doc,
      f"I have read the Policy on Procedural Sedation of {HN}. I will follow "
      "the processes described.")
    sig_tbl(doc)

    # 10. Distribution
    h(doc, 1, "10. Distribution")
    p(doc,
      "This policy shall be available to all staff who administer, monitor, or "
      f"are otherwise involved in procedural sedation at {HN}.")

    # 11. Abbreviations
    h(doc, 1, "11. Abbreviations")
    abbrev_tbl(doc, [
        ("BP",   "Blood pressure"),
        ("SpO2", "Peripheral oxygen saturation"),
    ])

    # 12. Traceability table
    h(doc, 1, "12. Traceability table")
    p(doc,
      "This table is an index. It is not how the policy is organised. An asterisk "
      "in the Level column means documentation of the process is required.")
    tr = tbl(doc, 9, 3)
    for ci, hdr in enumerate(("Objective Element", "Level", "Traceability to this policy")):
        tr.cell(0, ci).text = hdr
    trace_rows = [
        ("COP.12.a", "Commitment*",
         "Section 5.1 addresses written guidance covering at minimum six areas: "
         "identification of sedation areas and procedure locations, order-writing "
         "mechanism, pre-procedure assessment, intra-procedure monitoring, "
         "post-procedure monitoring, and discharge/transfer-out criteria. Also "
         "addresses the pre-sedation assessment performed and documented for every "
         "patient."),
        ("COP.12.b", "Commitment",
         "Section 5.2 addresses informed consent for procedural sedation covering "
         "risks, benefits, and alternatives, taken by the administering doctor or "
         "a sedation team doctor. Also addresses patient/family education on "
         "post-procedural pain relief."),
        ("COP.12.c", "Commitment",
         "Section 5.3 addresses competency requirements: parenteral sedation by a "
         "doctor or nurse under doctor supervision only. States the absolute "
         "prohibition: a technician never administers sedation, without exception."),
        ("COP.12.d", "Commitment",
         "Section 5.4 addresses the absolute rule that the monitoring person is "
         "always different from the procedure performer, without exception. Also "
         "addresses training of the monitoring person in detecting abnormalities, "
         "apnoea, and airway obstruction."),
        ("COP.12.e", "Commitment",
         "Section 5.5 addresses intra-procedure monitoring of at minimum six "
         "parameters: heart rate, cardiac rhythm, respiratory rate, blood "
         "pressure, oxygen saturation, and level of sedation. Notes that routine "
         "cardiac rhythm need not be documented continuously, but rhythm "
         "abnormalities must be documented."),
        ("COP.12.f", "Commitment",
         "Section 5.6 addresses post-sedation monitoring of at minimum five "
         "parameters: heart rate, respiratory rate, blood pressure, oxygen "
         "saturation, and level of sedation. Cardiac rhythm is NOT in the "
         "post-sedation list. Monitoring at organisation-defined intervals until "
         "full recovery; monitoring documented."),
        ("COP.12.g", "Commitment*",
         "Section 5.7 addresses discharge/transfer-out criteria developed and "
         "documented by the organisation, in line with physiologic parameters and "
         "sound clinical practice. A qualified individual applies the criteria, "
         "and the application is documented."),
        ("COP.12.h", "Commitment",
         "Section 5.8 addresses the five emergency equipment items that must be "
         "in working order in the sedation room: emergency resuscitation "
         "equipment, suction, advanced airway equipment, positive pressure "
         "ventilation, and supplemental oxygen. Also addresses availability of a "
         "trained airway manager or anaesthesiologist able to respond quickly."),
    ]
    for ri, (oe, lvl, txt) in enumerate(trace_rows, 1):
        tr.cell(ri, 0).text = oe
        tr.cell(ri, 1).text = lvl
        tr.cell(ri, 2).text = txt

    # 13. Required Records/Evidence Checklist
    h(doc, 1, "13. Required Records/Evidence Checklist")

    h(doc, 2, "Written guidance and pre-sedation assessment")
    lb(doc,
       "Written guidance for procedural sedation covering all six minimum areas: "
       "sedation areas, order-writing, pre-procedure assessment, intra-procedure "
       "monitoring, post-procedure monitoring, discharge/transfer-out criteria.")
    lb(doc,
       "Pre-sedation assessment documented for every patient, covering risk and "
       "appropriateness of sedation.")

    h(doc, 2, "Informed consent")
    lb(doc,
       "Informed consent records covering risks, benefits, and alternatives for "
       "every sedation procedure.")
    lb(doc,
       "Evidence that consent was taken by the administering doctor or a sedation "
       "team doctor.")
    lb(doc,
       "Records of patient/family education on post-procedural pain relief.")

    h(doc, 2, "Staff competency and role separation")
    lb(doc,
       "Evidence that parenteral sedation is administered only by a doctor or a "
       "nurse under direct doctor supervision — no technician administers sedation "
       "under any circumstances.")
    lb(doc,
       "Evidence that the monitoring role is kept separate from the procedure role "
       "for every sedation.")
    lb(doc,
       "Evidence of monitoring person's training in detecting abnormalities in "
       "monitoring parameters and recognising apnoea and airway obstruction.")

    h(doc, 2, "Intra-procedure monitoring records")
    lb(doc,
       "Intra-procedure monitoring records documenting all six parameters: heart "
       "rate, cardiac rhythm, respiratory rate, blood pressure, oxygen saturation, "
       "and level of sedation.")
    lb(doc,
       "Records of any cardiac rhythm abnormalities documented per episode.")

    h(doc, 2, "Post-sedation monitoring and discharge records")
    lb(doc,
       "Post-sedation monitoring records documenting at minimum five parameters: "
       "heart rate, respiratory rate, blood pressure, oxygen saturation, and level "
       "of sedation — at organisation-defined intervals until full recovery.")
    lb(doc,
       "Discharge/transfer-out records showing a qualified individual applied "
       "documented criteria; application documented for every patient.")

    h(doc, 2, "Emergency equipment")
    lb(doc,
       "Equipment check records confirming all five items in working order: "
       "emergency resuscitation equipment, suction, advanced airway equipment, "
       "positive pressure ventilation, and supplemental oxygen.")
    lb(doc,
       "Evidence that a trained airway manager or anaesthesiologist is available "
       f"in {HN} and able to respond quickly to the sedation area.")

    # 14. References
    h(doc, 1, "14. References")
    ln(doc,
       "National Accreditation Board for Hospitals and Healthcare Providers. NABH "
       "Accreditation Standards for Hospitals, 6th Edition. COP.12.")
    ln(doc, "Guidebook interpretation supplied for COP.12.a through COP.12.h.")
    ln(doc,
       "Standard treatment guidelines and professional society guidelines on "
       "procedural sedation (reference for written guidance, pre-sedation "
       "assessment, and monitoring parameters).")

    # Disclaimer
    h(doc, 1, "Disclaimer")
    p(doc,
      "This policy reorganises the supplied COP.12 objective elements and approved "
      "plain-language content into policy format. The stop-work text is reproduced "
      "exactly as supplied. Mandatory requirements and their stated modal strength "
      "have been retained. The six written-guidance areas (COP.12.a), the absolute "
      "prohibition on technicians administering sedation (COP.12.c — \"without "
      "exception\"), the absolute role-separation rule (COP.12.d — \"always "
      "different\", \"without exception\"), the six intra-procedure monitoring "
      "parameters (COP.12.e, including cardiac rhythm), the five post-sedation "
      "monitoring parameters (COP.12.f — cardiac rhythm is absent from this list), "
      "and the five emergency equipment items (COP.12.h) are all preserved "
      "exactly. No monitoring interval has been invented for post-sedation "
      "monitoring; the interval is organisation-defined.")

    save_and_verify(doc, "HCO_COP_12_v2_REWRITE_DRAFT.docx")


# ══════════════════════════════════════════════════════════════════════════════
# COP.13 — Anaesthesia Services   (stop-work YES)
# Content: approved plain-language content (cop13_content.txt).
# Structure: Section 3 two summary paragraphs; full detail Section 5.
# Stars: a*, g* | CORE: b (pre-anaesthesia assessment/plan), e (monitoring)
# Exact quantities preserved:
#   Local anaesthesia excluded from scope (Purpose, Scope, 5.1)
#   5.2: five anaesthesia plan elements
#   5.4: anaesthesia consent ALWAYS separate from surgery consent
#   5.5: seven intra-anaesthesia monitoring parameters (incl. temperature);
#        regional-anaesthesia ETCO2 exception; anaesthesiologist present
#        throughout (absolute statement)
#   5.8: five anaesthesia chart elements
# ══════════════════════════════════════════════════════════════════════════════

def gen_cop13():
    doc = Document()

    # Title
    h(doc, 0, "Policy on Anaesthesia Services")
    p(doc, HN)

    # Document control
    h(doc, 1, "Document control")
    doc_ctrl(doc, "HCO/COP/POL/13", "Anaesthesia In-Charge")
    p(doc, "A blank marked ________ must be completed before issue.")

    # Statement of intent
    h(doc, 1, "Statement of intent")
    p(doc,
      "Anaesthesia services are provided consistently and safely, with a documented "
      "pre-anaesthesia assessment and plan for every patient, continuous "
      "intra-anaesthesia monitoring, and defined criteria for recovery-area "
      "transfer.")

    # 1. Purpose
    h(doc, 1, "1. Purpose")
    p(doc,
      f"This policy explains how {HN} provides anaesthesia consistently, assesses "
      "patients before anaesthesia and before induction, obtains separate "
      "anaesthesia consent, monitors patients during and after anaesthesia, applies "
      "transfer criteria from the recovery area, documents anaesthesia given, "
      "follows infection prevention and control guidelines, and tracks "
      "intra-operative adverse anaesthesia events.")
    p(doc,
      "This policy does not cover local anaesthesia, general informed consent "
      "procedures, or infection control in detail — those are covered in other "
      "hospital policies. This policy does not apply to local anaesthesia.")

    # 2. Scope
    h(doc, 1, "2. Scope")
    p(doc,
      f"This policy applies to all anaesthesiologists and supporting staff involved "
      f"in administering, monitoring, or documenting anaesthesia at {HN}, "
      "excluding local anaesthesia.")

    # 3. Policy standards — two summary paragraphs only
    h(doc, 1, "3. Policy standards")
    p(doc,
      f"{HN} provides anaesthesia under written guidance applied consistently "
      "across the organisation, with a documented pre-anaesthesia assessment and "
      "plan and a pre-induction assessment for every patient, separate informed "
      "consent for anaesthesia, continuous monitoring of a defined set of "
      "parameters during the procedure, post-anaesthesia monitoring until complete "
      "recovery, defined transfer criteria from the recovery area, complete "
      "documentation in the anaesthesia chart, compliance with infection prevention "
      "and control guidelines, and tracking of intra-operative adverse anaesthesia "
      "events.")
    p(doc, "Staff follow the written guidance below and keep the records it requires.")

    # 4. Non-negotiable rules
    h(doc, 1, "4. Non-negotiable rules")
    lb(doc,
       "Do not provide anaesthesia without written guidance documenting the "
       "indications, the type of anaesthesia, and the procedure — this policy does "
       "not apply to local anaesthesia.")
    lb(doc,
       "Do not wheel a patient into the OT complex without a completed "
       "pre-anaesthesia assessment and documented plan by a qualified "
       "anaesthesiologist — this applies to both routine and emergency cases.")
    lb(doc,
       "Do not skip any of the five mandatory anaesthesia plan elements: "
       "pre-medications, type of anaesthesia, anaesthesia agent and dose, special "
       "requirements, and anticipated post-anaesthesia care including post-operative "
       "pain management where appropriate.")
    lb(doc,
       "Do not induce anaesthesia without a pre-induction assessment by an "
       "anaesthesiologist, done immediately before induction.")
    lb(doc,
       "Do not obtain anaesthesia consent from anyone other than the "
       "anaesthesiologist, and do not combine anaesthesia consent with surgery "
       "consent — they must be separate.")
    lb(doc,
       "Do not administer anaesthesia without monitoring all seven required "
       "parameters: temperature, heart rate, cardiac rhythm, respiratory rate, "
       "blood pressure, oxygen saturation, and end-tidal carbon dioxide — except "
       "for regional anaesthesia, where end-tidal carbon dioxide is replaced by "
       "continual observation of ventilation adequacy.")
    lb(doc,
       "Do not leave the anaesthesiologist absent at any point during the "
       "procedure.")
    lb(doc,
       "Do not transfer a patient from the recovery area without the "
       "anaesthesiologist applying documented transfer criteria.")
    lb(doc,
       "Do not administer anaesthesia in breach of infection prevention and "
       "control guidelines.")

    # 5. What we do
    h(doc, 1, "5. What we do")

    h(doc, 2, "5.1 Provide anaesthesia under consistent written guidance")
    p(doc,
      "Written guidance based on standard treatment guidelines or sound clinical "
      f"practice governs anaesthesia services across {HN}. The indications, type "
      "of anaesthesia, and procedure are documented for every patient. This policy "
      "does not apply to local anaesthesia; \"anaesthesia\" is defined as in the "
      "NABH glossary.")

    h(doc, 2, "5.2 Complete a pre-anaesthesia assessment and plan")
    p(doc,
      "Every patient has a pre-anaesthesia assessment by a qualified "
      "anaesthesiologist, done before being wheeled into the OT complex. This "
      "applies to both routine and emergency cases. For elective surgery, the "
      "assessment may be done before admission; for emergency or obstetric "
      "patients, it may be done shortly before the procedure. The resulting plan "
      "covers five things: pre-medications, type of anaesthesia, anaesthesia agent "
      "and dose, special requirements, and anticipated post-anaesthesia care, "
      "including post-operative pain management where appropriate. The "
      "anaesthesiologist reviews the patient's current medications. The plan is "
      "documented in the patient's medical record.")

    h(doc, 2, "5.3 Complete a pre-induction assessment")
    p(doc,
      "Immediately before inducing anaesthesia, an anaesthesiologist re-evaluates "
      "the patient in a pre-induction assessment. Any changes to the anaesthesia "
      "plan are documented. In urgent cases, the pre-anaesthesia and pre-induction "
      "assessments may be done one after another or simultaneously, but each is "
      "documented separately in the patient's medical record.")

    h(doc, 2, "5.4 Obtain separate anaesthesia consent")
    p(doc,
      "The anaesthesiologist obtains informed consent for anaesthesia, educating "
      "the patient and/or family on the risks, benefits, and alternatives. "
      "Anaesthesia consent is always separate from surgery consent. (Related "
      "requirements are covered in the hospital's other policies.)")

    h(doc, 2, "5.5 Monitor patients throughout anaesthesia")
    p(doc,
      "During anaesthesia, monitoring covers seven parameters at minimum: "
      "temperature, heart rate, cardiac rhythm, respiratory rate, blood pressure, "
      "oxygen saturation, and end-tidal carbon dioxide. For regional anaesthesia, "
      "end-tidal carbon dioxide is replaced by continual observation of qualitative "
      "clinical signs showing ventilation is adequate. Other parameters may be "
      "monitored case by case. Routine cardiac rhythm need not be documented "
      "continuously, but any rhythm abnormality is documented. The anaesthesiologist "
      "is present throughout the procedure. Monitoring results are documented in "
      "the patient's medical record.")

    h(doc, 2, "5.6 Monitor patients after anaesthesia")
    p(doc,
      "Post-anaesthesia status is monitored in the recovery area or OT and "
      "documented, at minimum covering vitals, until the patient recovers "
      "completely from anaesthesia. This is done by an anaesthesiologist. If the "
      "patient is unstable and needs ICU care, monitoring continues there. Where a "
      "patient is transferred directly from OT to ICU, monitoring and "
      "documentation match what would be required in the recovery room.")

    h(doc, 2, "5.7 Apply defined transfer criteria")
    p(doc,
      "The anaesthesiologist applies criteria — documented by "
      f"{HN}, based on physiologic parameters and sound clinical practice — to "
      "decide when a patient can be transferred from the recovery area.")

    h(doc, 2, "5.8 Document anaesthesia given")
    p(doc,
      "The anaesthesia chart records five things for every patient: the name of "
      "the anaesthesiologist who performed the procedure, the names and "
      "designations of anyone who assisted, the type of anaesthesia used, the "
      "medications used, and the date, time, and signature of the person making "
      "the entry. (Related requirements are covered in the hospital's other "
      "policies.)")

    h(doc, 2, "5.9 Follow infection prevention and control guidelines")
    p(doc,
      "Anaesthesia procedures comply with infection prevention and control "
      "guidelines to prevent cross-infection between patients — for example, in "
      "how circuits are managed. These guidelines are documented, either separately "
      "or as part of the infection prevention and control manual. (Related "
      "requirements are covered in the hospital's other policies.)")

    h(doc, 2, "5.10 Track intra-operative adverse anaesthesia events")
    p(doc,
      f"{HN} defines the intra-operative adverse anaesthesia events relevant to "
      "its scope — for example, cardiac arrest, anaesthesia gas failure, "
      "endotracheal tube slippage, or air embolism. These are examples; "
      f"{HN} decides what's relevant to it. There's a mechanism to make sure "
      "every such event is captured — for instance, a dedicated heading in the "
      "anaesthesia record — documented and monitored so corrective and preventive "
      "action can follow.")

    # 6. Stop-work authority — text exactly as given in cop13_content.txt
    h(doc, 1, "6. Stop-work authority")
    p(doc,
      "Do not wheel a patient into the OT without a completed pre-anaesthesia "
      "assessment and plan by a qualified anaesthesiologist; do not administer "
      "anaesthesia without the seven required monitoring parameters in place (or "
      "the regional-anaesthesia alternative); do not proceed without anaesthesia "
      "consent.")
    p(doc,
      "Stop-work applies to elective and non-emergent anaesthesia starts. A "
      "genuine life-threatening emergency continues with the best available "
      "qualified anaesthesia support while the gap is escalated and documented.")
    p(doc,
      "The person who stops tells the Anaesthesia In-Charge and the Medical "
      "Superintendent the same shift. Refusing to proceed with unsafe anaesthesia "
      "is not a disciplinary matter.")

    # 7. Governance and responsibility
    h(doc, 1, "7. Governance and responsibility")
    gov_tbl(doc, [
        ("Medical Superintendent",
         "Accountable that this policy is followed and appropriately resourced; "
         "receives stop-work notifications the same shift."),
        ("Anaesthesia In-Charge",
         "Oversees the anaesthesia programme; ensures written guidance, "
         "assessment/plan completion, monitoring compliance, chart documentation, "
         "and adverse event tracking are in place; receives stop-work notifications "
         "the same shift."),
        ("Anaesthesiologists",
         "Perform and document pre-anaesthesia assessment and plan (five elements) "
         "before OT, for all patients including emergencies; perform pre-induction "
         "assessment immediately before induction; obtain anaesthesia consent "
         "separately from surgery consent; monitor all seven required parameters "
         "throughout; remain present throughout the procedure; apply and document "
         "transfer criteria; complete anaesthesia chart (five elements) for every "
         "patient; follow IPC guidelines."),
        ("Supporting anaesthesia staff",
         "Assist anaesthesiologists during procedures; follow IPC guidelines; "
         "support adverse anaesthesia event capture."),
    ])

    # 8. Quality monitoring
    h(doc, 1, "8. Quality monitoring")
    mon_tbl(doc, [
        ("Written guidance",
         "Written guidance covering indications, type of anaesthesia, and "
         "procedure; excludes local anaesthesia explicitly"),
        ("Pre-anaesthesia assessment and plan",
         "Assessment completed before OT complex for all patients (routine and "
         "emergency); plan covers all five elements; patient's current medications "
         "reviewed; documented in medical record"),
        ("Pre-induction assessment",
         "Completed immediately before induction by an anaesthesiologist; changes "
         "to plan documented; separately documented in urgent cases"),
        ("Anaesthesia consent",
         "Informed consent obtained by the anaesthesiologist, always separate from "
         "surgery consent; covers risks, benefits, and alternatives"),
        ("Intra-anaesthesia monitoring",
         "All seven parameters monitored throughout: temperature, heart rate, "
         "cardiac rhythm, respiratory rate, blood pressure, oxygen saturation, "
         "ETCO2 (or continual qualitative observation for regional anaesthesia); "
         "anaesthesiologist present throughout; documented"),
        ("Post-anaesthesia monitoring",
         "Vitals monitored until complete recovery by anaesthesiologist; ICU "
         "monitoring continues if patient is unstable; documented"),
        ("Transfer criteria",
         "Criteria applied by anaesthesiologist and documented for every transfer "
         "from the recovery area"),
        ("Anaesthesia chart",
         "Five elements documented for every patient: anaesthesiologist name, "
         "assistant names and designations, type of anaesthesia, medications, "
         "date/time/signature"),
        ("Adverse anaesthesia events",
         "Defined events captured through a named mechanism; documented and "
         "monitored; CAPA applied"),
    ])

    # 9. Training and staff acknowledgement
    h(doc, 1, "9. Training and staff acknowledgement")
    p(doc,
      "All anaesthesiologists and supporting staff shall be familiar with the "
      "written guidance applicable to their work, including the pre-anaesthesia "
      "assessment and five-element plan requirements, pre-induction assessment, "
      "separate anaesthesia consent, the seven intra-anaesthesia monitoring "
      "parameters (including the regional-anaesthesia ETCO2 alternative), the "
      "requirement for the anaesthesiologist to be present throughout, transfer "
      "criteria, anaesthesia chart requirements, IPC guidelines, and adverse event "
      "capture.")
    p(doc,
      f"I have read the Policy on Anaesthesia Services of {HN}. I will follow "
      "the processes described.")
    sig_tbl(doc)

    # 10. Distribution
    h(doc, 1, "10. Distribution")
    p(doc,
      "This policy shall be available to all anaesthesiologists and supporting "
      f"staff involved in administering, monitoring, or documenting anaesthesia "
      f"at {HN}.")

    # 11. Abbreviations
    h(doc, 1, "11. Abbreviations")
    abbrev_tbl(doc, [
        ("CAPA",  "Corrective and Preventive Action"),
        ("ETCO2", "End-tidal carbon dioxide"),
        ("ICU",   "Intensive Care Unit"),
        ("IPC",   "Infection Prevention and Control"),
        ("OT",    "Operation Theatre"),
    ])

    # 12. Traceability table
    h(doc, 1, "12. Traceability table")
    p(doc,
      "This table is an index. It is not how the policy is organised. An asterisk "
      "in the Level column means documentation of the process is required.")
    tr = tbl(doc, 11, 3)
    for ci, hdr in enumerate(("Objective Element", "Level", "Traceability to this policy")):
        tr.cell(0, ci).text = hdr
    trace_rows = [
        ("COP.13.a", "Commitment*",
         "Section 5.1 addresses written guidance for anaesthesia services based on "
         "standard treatment guidelines or sound clinical practice, covering "
         "indications, type of anaesthesia, and procedure documented for every "
         "patient. Explicitly excludes local anaesthesia."),
        ("COP.13.b", "CORE",
         "Section 5.2 addresses the mandatory pre-anaesthesia assessment and plan "
         "by a qualified anaesthesiologist, completed before the patient is wheeled "
         "into the OT complex, for all patients including emergencies. The plan "
         "covers five elements: pre-medications, type of anaesthesia, anaesthesia "
         "agent and dose, special requirements, and anticipated post-anaesthesia "
         "care including post-operative pain management where appropriate."),
        ("COP.13.c", "Commitment",
         "Section 5.3 addresses the pre-induction assessment by an "
         "anaesthesiologist done immediately before induction; changes to the plan "
         "documented. In urgent cases, assessments may be sequential or "
         "simultaneous but each documented separately."),
        ("COP.13.d", "Commitment",
         "Section 5.4 addresses informed consent for anaesthesia, obtained by the "
         "anaesthesiologist and always separate from surgery consent. Covers risks, "
         "benefits, and alternatives."),
        ("COP.13.e", "CORE",
         "Section 5.5 addresses intra-anaesthesia monitoring of seven minimum "
         "parameters: temperature, heart rate, cardiac rhythm, respiratory rate, "
         "blood pressure, oxygen saturation, and end-tidal carbon dioxide. For "
         "regional anaesthesia, ETCO2 is replaced by continual observation of "
         "qualitative clinical signs showing ventilation is adequate. The "
         "anaesthesiologist is present throughout the procedure (absolute). "
         "Documented in the medical record."),
        ("COP.13.f", "Commitment",
         "Section 5.6 addresses post-anaesthesia monitoring of vitals until "
         "complete recovery, by an anaesthesiologist, in the recovery area or OT. "
         "If patient is unstable, monitoring continues in ICU. ICU-direct "
         "transfers are documented to recovery-room standard."),
        ("COP.13.g", "Commitment*",
         "Section 5.7 addresses the anaesthesiologist applying organisation-"
         "defined transfer criteria based on physiologic parameters and sound "
         "clinical practice; application documented for every transfer."),
        ("COP.13.h", "Commitment",
         "Section 5.8 addresses the anaesthesia chart recording five elements for "
         "every patient: name of the anaesthesiologist, names and designations of "
         "assistants, type of anaesthesia, medications used, and date/time/"
         "signature of entry."),
        ("COP.13.i", "Commitment",
         "Section 5.9 addresses compliance with IPC guidelines to prevent "
         "cross-infection between patients (e.g., circuit management); guidelines "
         "documented standalone or within the IPC manual."),
        ("COP.13.j", "Achievement",
         "Section 5.10 addresses organisation-defined intra-operative adverse "
         "anaesthesia events, a capture mechanism, documentation, monitoring, and "
         "corrective and preventive action."),
    ]
    for ri, (oe, lvl, txt) in enumerate(trace_rows, 1):
        tr.cell(ri, 0).text = oe
        tr.cell(ri, 1).text = lvl
        tr.cell(ri, 2).text = txt

    # 13. Required Records/Evidence Checklist
    h(doc, 1, "13. Required Records/Evidence Checklist")

    h(doc, 2, "Written guidance")
    lb(doc,
       "Written guidance for anaesthesia services documenting indications, type of "
       "anaesthesia, and procedure; explicitly excluding local anaesthesia.")

    h(doc, 2, "Pre-anaesthesia assessment and plan")
    lb(doc,
       "Pre-anaesthesia assessment completed before OT for every patient (routine "
       "and emergency), by a qualified anaesthesiologist.")
    lb(doc,
       "Anaesthesia plan covering all five elements: pre-medications, type, agent "
       "and dose, special requirements, anticipated post-anaesthesia care.")
    lb(doc,
       "Review of patient's current medications documented.")
    lb(doc, "Plan documented in the patient's medical record.")

    h(doc, 2, "Pre-induction assessment")
    lb(doc,
       "Pre-induction assessment by an anaesthesiologist, done immediately before "
       "induction; changes to plan documented.")
    lb(doc,
       "In urgent cases: each assessment documented separately even if done "
       "sequentially or simultaneously.")

    h(doc, 2, "Anaesthesia consent")
    lb(doc,
       "Informed consent records for anaesthesia obtained by the anaesthesiologist, "
       "separate from surgery consent.")

    h(doc, 2, "Intra-anaesthesia monitoring")
    lb(doc,
       "Monitoring records covering all seven parameters for every patient: "
       "temperature, heart rate, cardiac rhythm, respiratory rate, blood pressure, "
       "oxygen saturation, ETCO2 (or continual qualitative ventilation observation "
       "for regional anaesthesia).")
    lb(doc,
       "Evidence that the anaesthesiologist was present throughout every procedure.")

    h(doc, 2, "Post-anaesthesia monitoring and transfer")
    lb(doc,
       "Post-anaesthesia monitoring records (vitals until complete recovery) in "
       "the recovery area, OT, or ICU as appropriate.")
    lb(doc,
       "Transfer criteria documented; anaesthesiologist's application of criteria "
       "recorded for every transfer from recovery.")

    h(doc, 2, "Anaesthesia chart")
    lb(doc,
       "Anaesthesia chart for every patient recording all five elements: "
       "anaesthesiologist's name, assistants' names and designations, type of "
       "anaesthesia, medications, date/time/signature.")

    h(doc, 2, "IPC compliance and adverse events")
    lb(doc,
       "IPC guidelines for anaesthesia procedures documented (standalone or in IPC "
       "manual); evidence of compliance.")
    lb(doc,
       "Organisation-defined list of intra-operative adverse anaesthesia events; "
       "capture mechanism in place; records of events documented; CAPA applied.")

    # 14. References
    h(doc, 1, "14. References")
    ln(doc,
       "National Accreditation Board for Hospitals and Healthcare Providers. NABH "
       "Accreditation Standards for Hospitals, 6th Edition. COP.13.")
    ln(doc, "Guidebook interpretation supplied for COP.13.a through COP.13.j.")
    ln(doc,
       "NABH Glossary (reference for the definition of \"anaesthesia\" as "
       "applicable to this policy).")

    # Disclaimer
    h(doc, 1, "Disclaimer")
    p(doc,
      "This policy reorganises the supplied COP.13 objective elements and approved "
      "plain-language content into policy format. The stop-work text is reproduced "
      "exactly as supplied. Mandatory requirements and their stated modal strength "
      "have been retained. Local anaesthesia is explicitly excluded from scope in "
      "the Purpose, Scope, and Section 5.1. The five anaesthesia plan elements "
      "(COP.13.b) are preserved exactly. The seven intra-anaesthesia monitoring "
      "parameters (COP.13.e) including temperature — the anaesthesia-specific "
      "addition — are preserved exactly, as is the regional-anaesthesia ETCO2 "
      "exception (continual observation of qualitative clinical signs). The "
      "statement that the anaesthesiologist is present throughout the procedure is "
      "preserved as an absolute. The five anaesthesia chart elements (COP.13.h) "
      "are preserved exactly. Anaesthesia consent is stated as always separate "
      "from surgery consent. No specific time frames have been invented for "
      "assessment completion in emergency cases.")

    save_and_verify(doc, "HCO_COP_13_v2_REWRITE_DRAFT.docx")


# ══════════════════════════════════════════════════════════════════════════════
# COP.14 — Surgical Services   (stop-work YES)
# Content: approved plain-language content (cop14_content.txt).
# Structure: Section 3 two summary paragraphs; full detail Section 5.
# Stars: a*, d* (CORE), i*, j* | Achievement: i, j | CORE: d
# Exact quantities preserved:
#   5.3: fresh consent except life-saving exception
#   5.4: three site-marking triggers (laterality/multiple structures/levels);
#        time-out immediately before with three mandatory agreements
#   5.5: eight operative-note elements; before transfer from recovery
#   5.8: seven OT facility areas (+recovery room where applicable)
#   5.10: two distinct monitoring frequencies — daily (humidity/pressure/temp)
#         vs. at least every six months (filter integrity) — kept separate
# ══════════════════════════════════════════════════════════════════════════════

def gen_cop14():
    doc = Document()

    # Title
    h(doc, 0, "Policy on Surgical Services")
    p(doc, HN)

    # Document control
    h(doc, 1, "Document control")
    doc_ctrl(doc, "HCO/COP/POL/14", "OT In-Charge")
    p(doc, "A blank marked ________ must be completed before issue.")

    # Statement of intent
    h(doc, 1, "Statement of intent")
    p(doc,
      "Surgical services are provided consistently and safely, with wrong-site, "
      "wrong-patient, and wrong-surgery events prevented through mandatory site "
      "marking and time-out, and complete documentation from pre-operative "
      "assessment through to the operative note.")

    # 1. Purpose
    h(doc, 1, "1. Purpose")
    p(doc,
      f"This policy explains how {HN} provides surgical services under written "
      "guidance, assesses patients and obtains consent before surgery, prevents "
      "wrong-site, wrong-patient, and wrong-surgery events, documents every "
      "operation, plans post-operative care, keeps the operating theatre safe and "
      "properly equipped, and runs a quality assurance programme including "
      "environmental surveillance.")
    p(doc,
      "This policy does not cover general informed consent procedures or infection "
      "control in detail — those are covered in other hospital policies. If "
      f"{HN} does not provide surgical services, this policy does not apply.")

    # 2. Scope
    h(doc, 1, "2. Scope")
    p(doc,
      f"This policy applies to all surgeons, anaesthesiologists, and operating "
      f"theatre staff involved in surgical services at {HN}.")

    # 3. Policy standards — two summary paragraphs only
    h(doc, 1, "3. Policy standards")
    p(doc,
      f"{HN} provides surgical services under written guidance listing the "
      "procedures offered and the competency needed for each, assesses and obtains "
      "consent from every surgical patient, prevents wrong-site, wrong-patient, "
      "and wrong-surgery events through mandatory site marking and time-out, "
      "documents every operation completely before the patient leaves recovery, "
      "plans post-operative care, keeps the operating theatre free of "
      "sterile/unsterile mix-ups and properly equipped, and runs a quality "
      "assurance programme that includes environmental surveillance of the "
      "operating theatre.")
    p(doc, "Staff follow the written guidance below and keep the records it requires.")

    # 4. Non-negotiable rules
    h(doc, 1, "4. Non-negotiable rules")
    lb(doc,
       "Do not perform a surgical procedure that isn't on the hospital's "
       "documented list, or perform it with personnel below the competency level "
       "the guidance defines for that procedure.")
    lb(doc,
       "Do not proceed with elective surgery without a pre-operative assessment, "
       "a documented pre-operative diagnosis, and pre-operative instructions given "
       "to the patient and family — do the same for emergency cases whenever "
       "possible, and document when it wasn't possible.")
    lb(doc,
       "Do not proceed with surgery without informed consent taken by the "
       "operating surgeon or a doctor on the operating team.")
    lb(doc,
       "Do not perform a new or additional procedure that wasn't planned or "
       "consented to, without fresh consent — except where the procedure is "
       "life-saving.")
    lb(doc,
       "Do not begin a surgical or invasive procedure without completing site "
       "marking, where the procedure involves laterality, multiple structures, or "
       "multiple levels.")
    lb(doc,
       "Do not begin a surgical procedure without holding a time-out immediately "
       "beforehand, with every team member present, confirming the correct patient "
       "identity, the correct procedure, and the correct surgical site.")
    lb(doc,
       "Do not transfer a patient out of recovery before the operative note is "
       "completed, covering all eight required elements.")
    lb(doc,
       "Do not let an operating theatre run without resuscitation equipment, "
       "without instruments in verified working condition, or in a way that mixes "
       "sterile and unsterile materials.")

    # 5. What we do
    h(doc, 1, "5. What we do")

    h(doc, 2, "5.1 Provide surgical services under written guidance")
    p(doc,
      "Written guidance based on standard treatment guidelines or sound clinical "
      f"practice governs surgical services. It includes the list of surgical "
      f"procedures {HN} offers and the competency level required to perform each "
      "one.")

    h(doc, 2, "5.2 Assess and prepare patients before surgery")
    p(doc,
      "Every surgical patient has a pre-operative assessment, a documented "
      "pre-operative diagnosis, and pre-operative instructions given to everyone "
      "concerned, including the patient and family. This applies to all elective "
      "cases, and to emergency cases whenever possible — where it isn't possible "
      "for an emergency case, that's documented too. This is done by the operating "
      "surgeon or a doctor on the operating team.")

    h(doc, 2, "5.3 Obtain informed consent")
    p(doc,
      "Consent is taken by the operating surgeon or a doctor on the operating "
      "team. If the patient's clinical status or expected outcome changes after "
      "consent but before surgery, this is explained to the patient and family and "
      "documented. If a new or additional procedure comes up that wasn't planned or "
      "consented to, fresh consent is obtained — except where the procedure is "
      "life-saving, in which case the situation is documented instead. (Related "
      "requirements are covered in the hospital's other policies.)")

    h(doc, 2, "5.4 Prevent wrong-site, wrong-patient, and wrong-surgery events")
    p(doc,
      f"Written guidance sets out how {HN} prevents these events — for example, "
      "through identification tags, badges, cross-checks, and time-outs.")
    p(doc,
      "Site marking is done wherever a procedure involves laterality, multiple "
      "structures (like fingers, toes, or lesions), or multiple levels (like the "
      "spine). The mark is instantly recognisable, unambiguous, and used the same "
      "way throughout the hospital. It's made by the person — or a team member — "
      "who will perform the entire procedure and stay with the patient throughout. "
      "The patient is actively involved in marking their own site whenever "
      "possible. Marking can happen any time before the procedure starts, but the "
      "mark stays visible once prepping and draping are complete.")
    p(doc,
      "A time-out is held immediately before the procedure starts, with every team "
      "member present, to confirm three things together: the correct patient "
      "identity, the correct procedure, and the correct surgical site.")

    h(doc, 2, "5.5 Document the operation")
    p(doc,
      "Before the patient is transferred out of recovery, an operative note is "
      "completed by the surgeon or a doctor on the operating team. If someone "
      "other than the chief surgeon writes it, the chief surgeon countersigns it. "
      "At a minimum, it covers eight things: the surgery performed, the name(s) of "
      "the surgeon(s), the name(s) of the anaesthesiologist(s), the post-operative "
      "diagnosis, any peri-operative complications, the amount of blood loss, the "
      "salient steps of the procedure, and the key intra-operative findings.")

    h(doc, 2, "5.6 Plan post-operative care")
    p(doc,
      "Post-operative care follows a documented plan, written by the operating "
      "surgeon or a member of the operating team. The plan addresses whatever is "
      "clinically relevant — for example, IV fluids, medication, wound care, "
      "nursing care, or watching for complications.")

    h(doc, 2, "5.7 Keep the operating theatre safe from cross-contamination")
    p(doc,
      "The layout and practices of the operating theatre prevent sterile and "
      "unsterile materials from mixing. It's preferable to use separate closed "
      "trolleys for sterile and unsterile goods, though any method that achieves "
      "the same outcome is acceptable.")

    h(doc, 2, "5.8 Keep the operating theatre properly equipped")
    p(doc,
      "The operating theatre complex has, at minimum, seven facility areas: a "
      "pre-op holding area, changing rooms, hand-washing facilities, operating "
      "rooms, a waiting area for relatives, a storage area, and a collection area "
      "for waste and linen — plus a recovery room, where applicable. Beyond "
      "anaesthesia and surgical equipment, there's resuscitation equipment, and "
      "radiation protection equipment where applicable. Instruments are kept in "
      "working condition, checked through a defined verification mechanism — "
      "feedback from surgeons or anaesthesiologists is one way to do this. "
      "Supplies match the scope and complexity of the surgery performed.")

    h(doc, 2, "5.9 Run a quality assurance programme")
    p(doc,
      "The quality assurance programme monitors care-related outcomes — for "
      "example, intra-operative mishaps like cautery burns, patient falls, or "
      "position-related nerve injuries, and peri-operative events like surgical "
      "site infections or deep venous thrombosis. These are examples; "
      f"{HN} decides which outcomes matter most for its own surgical practice. "
      "The programme also covers pre-operative preparation, antimicrobial "
      "prophylaxis, and adherence to procedures that prevent adverse events.")

    h(doc, 2, "5.10 Monitor the operating theatre environment")
    p(doc,
      "The quality assurance programme includes surveillance of the operating "
      "theatre environment — daily monitoring of humidity, pressure differential, "
      f"and temperature, and monitoring of filter integrity at least every six "
      f"months. {HN} follows the applicable NABH guidelines for OT "
      "air-conditioning.")

    # 6. Stop-work authority — text exactly as given in cop14_content.txt
    h(doc, 1, "6. Stop-work authority")
    p(doc,
      "Do not begin or continue a surgical or invasive procedure without completed "
      "site marking (where laterality, multiple structures, or multiple levels "
      "apply) and a completed time-out confirming correct patient, procedure, and "
      "site with the whole team present.")
    p(doc,
      "Stop-work applies to the start of the procedure. A genuine life-threatening "
      "emergency continues with the best available qualified team while the gap is "
      "escalated and documented.")
    p(doc,
      "The person who stops tells the OT In-Charge and the Medical Superintendent "
      "the same shift. Refusing to proceed without site marking or time-out is not "
      "a disciplinary matter.")

    # 7. Governance and responsibility
    h(doc, 1, "7. Governance and responsibility")
    gov_tbl(doc, [
        ("Medical Superintendent",
         "Accountable that this policy is followed and appropriately resourced; "
         "receives stop-work notifications the same shift."),
        ("OT In-Charge",
         "Oversees OT operations; ensures written guidance, site-marking and "
         "time-out compliance, operative note completion, OT safety, equipment "
         "readiness, and environmental surveillance are in place; receives "
         "stop-work notifications the same shift."),
        ("Operating surgeons",
         "Perform or supervise pre-operative assessment; obtain consent (including "
         "fresh consent for unplanned procedures except life-saving); perform site "
         "marking; hold time-out; complete operative note before patient leaves "
         "recovery (eight elements); write post-operative care plan; countersign "
         "operative note where applicable."),
        ("Anaesthesiologists",
         "Participate in time-out; comply with COP.13 for anaesthesia-specific "
         "obligations."),
        ("OT nursing and supporting staff",
         "Maintain sterile/unsterile separation; verify instrument working "
         "condition; support time-out; assist with daily and periodic "
         "environmental monitoring."),
    ])

    # 8. Quality monitoring
    h(doc, 1, "8. Quality monitoring")
    mon_tbl(doc, [
        ("Written guidance",
         "List of surgical procedures and competency levels documented and "
         "current; procedures performed only by staff at the required competency "
         "level"),
        ("Pre-operative assessment",
         "Pre-operative assessment, documented diagnosis, and pre-operative "
         "instructions completed for all elective cases; documented for emergency "
         "cases when possible, with documented reason when not"),
        ("Informed consent",
         "Consent taken by operating surgeon or team; fresh consent for new/"
         "additional procedures; exception for life-saving procedures documented"),
        ("Site marking",
         "Site marking completed for all procedures involving laterality, multiple "
         "structures, or multiple levels; mark stays visible after draping"),
        ("Time-out",
         "Time-out held immediately before every procedure with all team members "
         "present, confirming correct patient identity, procedure, and surgical "
         "site; documented"),
        ("Operative note",
         "Operative note completed with all eight elements before patient leaves "
         "recovery; countersigned by chief surgeon where applicable"),
        ("Post-operative care plan",
         "Documented post-operative care plan written by operating surgeon or "
         "team for every patient"),
        ("OT sterile/unsterile separation",
         "No mixing of sterile and unsterile materials; defined method in place"),
        ("OT facility and equipment",
         "All seven facility areas present; resuscitation equipment available; "
         "instruments verified working; supplies matched to surgical scope"),
        ("QA programme — outcomes",
         "Defined care outcomes monitored and analysed (e.g., SSI, DVT, "
         "intra-operative mishaps); CAPA applied"),
        ("Environmental surveillance — daily",
         "Daily monitoring of humidity, pressure differential, and temperature "
         "in the OT complex; results recorded"),
        ("Environmental surveillance — filter integrity",
         "Filter integrity monitored at least every six months; results recorded"),
    ])

    # 9. Training and staff acknowledgement
    h(doc, 1, "9. Training and staff acknowledgement")
    p(doc,
      "All surgeons, anaesthesiologists, and operating theatre staff shall be "
      "familiar with the written guidance applicable to their work, including "
      "pre-operative assessment requirements, informed consent obligations "
      "(including the fresh-consent-except-life-saving rule), site marking for "
      "laterality/multiple structures/multiple levels, the time-out process, "
      "operative note requirements (eight elements), OT sterile/unsterile "
      "separation, instrument verification, and environmental monitoring "
      "frequencies.")
    p(doc,
      f"I have read the Policy on Surgical Services of {HN}. I will follow "
      "the processes described.")
    sig_tbl(doc)

    # 10. Distribution
    h(doc, 1, "10. Distribution")
    p(doc,
      "This policy shall be available to all surgeons, anaesthesiologists, and "
      f"operating theatre staff involved in surgical services at {HN}.")

    # 11. Abbreviations
    h(doc, 1, "11. Abbreviations")
    abbrev_tbl(doc, [
        ("CAPA", "Corrective and Preventive Action"),
        ("DVT",  "Deep Venous Thrombosis"),
        ("OT",   "Operation Theatre"),
        ("SSI",  "Surgical Site Infection"),
    ])

    # 12. Traceability table
    h(doc, 1, "12. Traceability table")
    p(doc,
      "This table is an index. It is not how the policy is organised. An asterisk "
      "in the Level column means documentation of the process is required.")
    tr = tbl(doc, 11, 3)
    for ci, hdr in enumerate(("Objective Element", "Level", "Traceability to this policy")):
        tr.cell(0, ci).text = hdr
    trace_rows = [
        ("COP.14.a", "Commitment*",
         "Section 5.1 addresses written guidance for surgical services including "
         "the list of procedures offered and the competency level required for "
         "each procedure."),
        ("COP.14.b", "Commitment",
         "Section 5.2 addresses pre-operative assessment, documented pre-operative "
         "diagnosis, and pre-operative instructions for all elective and, where "
         "possible, emergency cases; documents when not possible for emergencies."),
        ("COP.14.c", "Commitment",
         "Section 5.3 addresses informed consent by operating surgeon or team; "
         "documenting changes in patient status after consent; obtaining fresh "
         "consent for new or additional unplanned procedures — except where the "
         "procedure is life-saving, in which case the situation is documented."),
        ("COP.14.d", "CORE*",
         "Section 5.4 addresses wrong-site/patient/surgery prevention. Site "
         "marking required for three triggers: laterality, multiple structures "
         "(e.g., fingers, toes, lesions), and multiple levels (e.g., spine). Mark "
         "is recognisable, unambiguous, consistent, made by the procedure performer "
         "or a team member who stays with the patient, and remains visible after "
         "draping. Time-out held immediately before the procedure with all team "
         "members present, confirming three things: correct patient identity, "
         "correct procedure, and correct surgical site."),
        ("COP.14.e", "Commitment",
         "Section 5.5 addresses the operative note completed before the patient "
         "leaves recovery, covering eight elements: surgery performed, "
         "surgeon name(s), anaesthesiologist name(s), post-operative diagnosis, "
         "peri-operative complications, blood loss, salient steps, and key "
         "intra-operative findings. Chief surgeon countersigns if not primary "
         "author."),
        ("COP.14.f", "Commitment",
         "Section 5.6 addresses a documented post-operative care plan written by "
         "the operating surgeon or a member of the operating team, covering "
         "clinically relevant elements."),
        ("COP.14.g", "Commitment",
         "Section 5.7 addresses prevention of sterile/unsterile mix-ups in the "
         "OT through layout and practices; preferred method is separate closed "
         "trolleys, though any equivalent method is acceptable."),
        ("COP.14.h", "Commitment",
         "Section 5.8 addresses OT complex facility requirements: at minimum seven "
         "areas (pre-op holding, changing rooms, hand-washing, operating rooms, "
         "relatives waiting, storage, waste and linen collection) plus recovery "
         "room where applicable. Also addresses resuscitation equipment, radiation "
         "protection where applicable, instrument verification, and supply adequacy."),
        ("COP.14.i", "Achievement*",
         "Section 5.9 addresses the quality assurance programme monitoring "
         "care-related outcomes (e.g., SSI, DVT, intra-operative mishaps), "
         "pre-operative preparation, antimicrobial prophylaxis, and adherence to "
         "adverse-event prevention procedures."),
        ("COP.14.j", "Achievement*",
         "Section 5.10 addresses OT environmental surveillance at two distinct "
         "frequencies: daily monitoring of humidity, pressure differential, and "
         "temperature; and monitoring of filter integrity at least every six months. "
         "These two frequencies are separate requirements. NABH OT "
         "air-conditioning guidelines are followed."),
    ]
    for ri, (oe, lvl, txt) in enumerate(trace_rows, 1):
        tr.cell(ri, 0).text = oe
        tr.cell(ri, 1).text = lvl
        tr.cell(ri, 2).text = txt

    # 13. Required Records/Evidence Checklist
    h(doc, 1, "13. Required Records/Evidence Checklist")

    h(doc, 2, "Written guidance and procedure list")
    lb(doc,
       "Written guidance for surgical services including the list of procedures "
       "offered and the competency level required for each.")

    h(doc, 2, "Pre-operative assessment and consent")
    lb(doc,
       "Pre-operative assessment, documented diagnosis, and pre-operative "
       "instructions for every elective surgical patient.")
    lb(doc,
       "Documentation when pre-operative assessment was not possible for an "
       "emergency case.")
    lb(doc,
       "Informed consent records for every surgical patient, taken by the "
       "operating surgeon or team.")
    lb(doc,
       "Fresh consent records for unplanned procedures; or documentation of "
       "life-saving circumstances where fresh consent was not obtained.")

    h(doc, 2, "Site marking and time-out")
    lb(doc,
       "Site marking records for all procedures involving laterality, multiple "
       "structures, or multiple levels; mark documented as visible after draping.")
    lb(doc,
       "Time-out records for every procedure, confirming correct patient identity, "
       "correct procedure, and correct surgical site with all team members present.")

    h(doc, 2, "Operative note and post-operative care")
    lb(doc,
       "Operative note completed before patient leaves recovery, covering all "
       "eight required elements; chief surgeon's countersignature where applicable.")
    lb(doc,
       "Documented post-operative care plan for every surgical patient.")

    h(doc, 2, "OT safety and equipment")
    lb(doc,
       "Evidence of sterile/unsterile separation method in use in the OT.")
    lb(doc,
       "OT facility checklist confirming all seven areas are present and "
       "functional; recovery room where applicable.")
    lb(doc,
       "Evidence of resuscitation equipment and radiation protection equipment "
       "where applicable.")
    lb(doc,
       "Instrument verification records (defined mechanism; working condition "
       "confirmed).")

    h(doc, 2, "Quality assurance and environmental surveillance")
    lb(doc,
       "QA programme documentation: defined outcome metrics, monitoring records, "
       "and CAPA for surgical care.")
    lb(doc,
       "Daily OT environmental monitoring records: humidity, pressure differential, "
       "and temperature.")
    lb(doc,
       "Filter integrity monitoring records: at least every six months.")

    # 14. References
    h(doc, 1, "14. References")
    ln(doc,
       "National Accreditation Board for Hospitals and Healthcare Providers. NABH "
       "Accreditation Standards for Hospitals, 6th Edition. COP.14.")
    ln(doc, "Guidebook interpretation supplied for COP.14.a through COP.14.j.")
    ln(doc,
       "NABH guidelines for OT air-conditioning (reference for COP.14.j "
       "environmental surveillance).")

    # Disclaimer
    h(doc, 1, "Disclaimer")
    p(doc,
      "This policy reorganises the supplied COP.14 objective elements and approved "
      "plain-language content into policy format. The stop-work text — which "
      "corresponds to COP.14.d (CORE*) — is reproduced exactly as supplied. "
      "Mandatory requirements and their stated modal strength have been retained. "
      "The three site-marking triggers (laterality, multiple structures, multiple "
      "levels), the time-out's three mandatory agreements (correct patient identity, "
      "correct procedure, correct surgical site), the eight operative-note elements, "
      "the seven OT facility areas, and the two distinct environmental monitoring "
      "frequencies (daily: humidity/pressure differential/temperature; at least "
      "every six months: filter integrity) are all preserved exactly and kept "
      "separate. The fresh-consent-except-life-saving exception (COP.14.c) is "
      "preserved exactly. The policy is scope-conditional: it does not apply if "
      "surgical services are not provided.")

    save_and_verify(doc, "HCO_COP_14_v2_REWRITE_DRAFT.docx")


# ══════════════════════════════════════════════════════════════════════════════
if __name__ == "__main__":
    gen_cop1()
    print("\nCOP.1 draft generated.")
    gen_cop2()
    print("\nCOP.2 draft generated.")
    gen_cop3()
    print("\nCOP.3 draft generated.")
    gen_cop4()
    print("\nCOP.4 draft generated.")
    gen_cop5()
    print("\nCOP.5 draft generated.")
    gen_cop6()
    print("\nCOP.6 draft generated.")
    gen_cop7()
    print("\nCOP.7 draft generated.")
    gen_cop8()
    print("\nCOP.8 draft generated.")
    gen_cop9()
    print("\nCOP.9 draft generated.")
    gen_cop10()
    print("\nCOP.10 draft generated.")
    gen_cop11()
    print("\nCOP.11 draft generated.")
    gen_cop12()
    print("\nCOP.12 draft generated.")
    gen_cop13()
    print("\nCOP.13 draft generated.")
    gen_cop14()
    print("\nCOP.14 draft generated.")
