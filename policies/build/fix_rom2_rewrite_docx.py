# -*- coding: utf-8 -*-
"""
Fix HCO_ROM_2_v2_REWRITE_DRAFT.docx:
  - Title style for the main title (was Heading 1)
  - Heading 1 for all top-level sections (was Heading 2)
  - Heading 2 for 5.x subsections (was Heading 3)
  - Prepared/Reviewed/Approved rows moved into Document control table
    (matching ROM.1's 6-row x 4-col layout with merged cells in rows 3-5)
  - Stray Prepared/Reviewed/Approved paragraphs removed
"""
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

SRC = "policies/build/rewrite_reference/HCO_ROM_2_v2_REWRITE_DRAFT.docx"

doc = Document(SRC)

# ── 1. Fix paragraph styles ──────────────────────────────────────────────────

TOP_LEVEL = {
    "Document control",
    "Statement of intent",
    "1. Purpose",
    "2. Scope",
    "3. Policy standards",
    "4. Non-negotiable rules",
    "5. What we do",
    "6. Governance and responsibility",
    "7. Quality monitoring",
    "8. Training and staff acknowledgement",
    "9. Distribution",
    "10. Abbreviations",
    "11. Traceability to NABH HCO Full Accreditation 6th Edition ROM.2",
    "12. Required Records / Evidence Checklist",
    "13. References",
    "Disclaimer",
}

stray = []

for p in doc.paragraphs:
    text = p.text
    sname = p.style.name

    if text == "Policy on Ethical Management of the Organisation":
        p.style = doc.styles["Title"]
    elif sname == "Heading 2" and text in TOP_LEVEL:
        p.style = doc.styles["Heading 1"]
    elif sname == "Heading 3" and text.startswith("5."):
        p.style = doc.styles["Heading 2"]
    elif text.startswith(("Prepared by:", "Reviewed by:", "Approved by:")):
        stray.append(p)

for p in stray:
    p._element.getparent().remove(p._element)

# ── 2. Rebuild Document control table ────────────────────────────────────────
# Target layout (matching ROM.1):
#   Row 0: Document No. | HCO/ROM/POL/02        | Version                | 2.0
#   Row 1: Issue No.    | 01                     | Review due             | one year from implementation
#   Row 2: Date created | ________               | Date of implementation | ________
#   Row 3: Prepared by  | Medical Superintendent   Name: ________   Signature: ________  [cols 1-3 merged]
#   Row 4: Reviewed by  | Quality Coordinator   Name: ________   Signature: ________     [cols 1-3 merged]
#   Row 5: Approved by  | Medical Superintendent   Name: ________   Signature: ________  [cols 1-3 merged]

body = doc.element.body
old_tbl_elem = doc.tables[0]._element
old_pos = list(body).index(old_tbl_elem)

# Build new table at end (we'll move it)
new_tbl = doc.add_table(rows=6, cols=4)
try:
    new_tbl.style = "Table Grid"
except KeyError:
    pass  # style may not exist; leave default

# Rows 0-2: four cells each
meta = [
    ("Document No.", "HCO/ROM/POL/02", "Version", "2.0"),
    ("Issue No.", "01", "Review due", "one year from implementation"),
    ("Date created", "________", "Date of implementation", "________"),
]
for ri, (a, b, c, d) in enumerate(meta):
    new_tbl.cell(ri, 0).text = a
    new_tbl.cell(ri, 1).text = b
    new_tbl.cell(ri, 2).text = c
    new_tbl.cell(ri, 3).text = d

# Rows 3-5: label in col 0; content in cols 1-3 (merged)
sign_rows = [
    ("Prepared by", "Medical Superintendent   Name: ________   Signature: ________"),
    ("Reviewed by", "Quality Coordinator   Name: ________   Signature: ________"),
    ("Approved by", "Medical Superintendent   Name: ________   Signature: ________"),
]
for ri, (label, content) in enumerate(sign_rows, start=3):
    new_tbl.cell(ri, 0).text = label
    new_tbl.cell(ri, 1).text = content
    new_tbl.cell(ri, 1).merge(new_tbl.cell(ri, 3))

# Move new table element to old table's position, then remove old table
new_tbl_elem = new_tbl._element
body.remove(new_tbl_elem)
body.insert(old_pos, new_tbl_elem)
body.remove(old_tbl_elem)

# ── 3. Verify: print first 20 paragraph styles ──────────────────────────────
print("=== First 20 paragraphs (style + text) ===")
for i, p in enumerate(doc.paragraphs[:20]):
    print(f"{i:2d}  {p.style.name!r:30s}  {p.text[:65]!r}")

# ── 4. Save ──────────────────────────────────────────────────────────────────
doc.save(SRC)
print("\nSaved:", SRC)
