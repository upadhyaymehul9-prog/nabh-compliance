# -*- coding: utf-8 -*-
"""
generate_hco_fms_rewrites.py
Generates HCO FMS.1-7 v2 rewrite-reference DOCX files.

Pipeline : python-docx, identical to generate_hco_aac_rewrites.py.
Output   : policies/build/rewrite_reference/HCO_FMS_N_v2_REWRITE_DRAFT.docx
Source   : policies/drafts_hco/hco_fmsN_v2_draft.json +
           policies/build/hco_fms_v2_methods.py
"""
import os
from docx import Document

HN  = "«Hospital Name»"   # guillemet-wrapped placeholder
OUT = "policies/build/rewrite_reference"
os.makedirs(OUT, exist_ok=True)


# ── Helpers (identical to generate_hco_aac_rewrites.py) ──────────────────────

def h(doc, lv, txt):
    return doc.add_paragraph(txt, style={0: "Title", 1: "Heading 1", 2: "Heading 2"}[lv])

def p(doc, txt=""):
    return doc.add_paragraph(txt, style="Normal")

def ln(doc, txt):
    return doc.add_paragraph(txt, style="List Number")

def lb(doc, txt):
    return doc.add_paragraph(txt, style="List Bullet")

def tbl(doc, rows, cols):
    t = doc.add_table(rows=rows, cols=cols)
    try:
        t.style = "Table Grid"
    except KeyError:
        pass
    return t

def doc_ctrl(doc, no, prep, appr="Medical Superintendent"):
    dc = tbl(doc, 6, 4)
    for ri, (a, b, c, d) in enumerate([
        ("Document No.", no,         "Version",              "2.0"),
        ("Issue No.",    "01",        "Review due",           "One year from implementation"),
        ("Date created", "________",  "Date of implementation", "________"),
    ]):
        dc.cell(ri, 0).text = a; dc.cell(ri, 1).text = b
        dc.cell(ri, 2).text = c; dc.cell(ri, 3).text = d
    for ri, (lbl, txt) in enumerate([
        ("Prepared by", f"{prep}  Name: ________  Signature: ________"),
        ("Reviewed by", "Quality Coordinator  Name: ________  Signature: ________"),
        ("Approved by", f"{appr}  Name: ________  Signature: ________"),
    ], start=3):
        dc.cell(ri, 0).text = lbl
        c1 = dc.cell(ri, 1); c1.text = txt; c1.merge(dc.cell(ri, 3))

def gov_tbl(doc, rows):
    t = tbl(doc, len(rows) + 1, 2)
    t.cell(0, 0).text = "Role"; t.cell(0, 1).text = "Responsibility"
    for ri, (role, resp) in enumerate(rows, 1):
        t.cell(ri, 0).text = role; t.cell(ri, 1).text = resp

def abbrev_tbl(doc, rows):
    t = tbl(doc, len(rows) + 1, 2)
    t.cell(0, 0).text = "Abbreviation"; t.cell(0, 1).text = "Meaning"
    for ri, (a, m) in enumerate(rows, 1):
        t.cell(ri, 0).text = a; t.cell(ri, 1).text = m

def trace_tbl(doc, rows):
    t = tbl(doc, len(rows) + 1, 5)
    for ci, hdr in enumerate(("OE", "Level", "Requirement", "Where addressed", "Responsible")):
        t.cell(0, ci).text = hdr
    for ri, row in enumerate(rows, 1):
        for ci, v in enumerate(row):
            t.cell(ri, ci).text = v

def disclaimer(doc, p2=None):
    p(doc,
      f"This document is a template prepared for the guidance of {HN} and must be reviewed, "
      f"adapted and formally approved by {HN} before use. Every entry marked ________ "
      f"must be completed before the document is issued.")
    p(doc, p2 or (
        f"The requirements in this document are accreditation requirements of the NABH Accreditation "
        f"Standards for Hospitals, 6th Edition, not duties under a named Act of Parliament. This "
        f"policy does not import the Consumer Protection Act, 2019, the Clinical Establishments Act, "
        f"2010, or the Mental Healthcare Act, 2017 as a checklist. Statutory duties that arise under "
        f"other documents of {HN} remain those documents. {HN} is responsible for verifying any "
        f"statutory duty that applies to it; this document does not constitute legal advice."))
    p(doc,
      f"{HN} remains responsible for verifying that it is current and consistent with the edition "
      f"of the accreditation standard against which it is being assessed. The clinical and technical "
      f"content reflects recognised national and international guidance current at the date of preparation.")
    p(doc,
      "This document is not issued by, endorsed by, or affiliated with NABH, the World Health "
      "Organization, the National Centre for Disease Control, the Food Safety and Standards Authority "
      "of India, any Pollution Control Board, or any other body named in it. Wording is original; no "
      "text has been reproduced from the standards, rules or guidelines referenced.")

def save_and_verify(doc, fname):
    import sys
    out = sys.stdout
    def pr(s):
        try:
            out.write(s + "\n")
        except UnicodeEncodeError:
            out.write(s.encode("ascii", "replace").decode() + "\n")
    pr(f"\n=== {fname} ===")
    for i, para in enumerate(doc.paragraphs[:50]):
        sn = para.style.name if para.style else "(None)"
        pr(f"{i:3d}  {sn!r:30s}  {para.text[:60]!r}")
    pr(f"  Total paras: {len(doc.paragraphs)}")
    path = os.path.join(OUT, fname)
    doc.save(path)
    print(f"  Saved: {path}")


# ── Shared FMS constants ──────────────────────────────────────────────────────

def fms_gov(std):
    return [
        ("Medical Superintendent",
         f"Accountable that {std} is resourced and followed."),
        ("Engineering In-Charge",
         "Owns day-to-day implementation and records for this standard."),
        ("Quality Coordinator",
         "Audits this policy quarterly; holds training acknowledgements."),
        ("departmental leaders",
         "Run the department-level duties this standard names."),
    ]

FMS_ABBREVS = [
    ("AHU",  "Air Handling Unit"),
    ("CAPA", "Corrective and Preventive Action"),
    ("CORE", "Core objective element (NABH)"),
    ("DG",   "Diesel Generator"),
    ("ELV",  "Extra Low Voltage"),
    ("FMS",  "Facility Management and Safety (NABH Hospitals 6th Edition chapter)"),
    ("HCO",  "Hospital (Full Accreditation programme under NABH Hospitals 6th Edition)"),
    ("HVAC", "Heating, Ventilation and Air Conditioning"),
    ("MSDS", "Material Safety Data Sheet"),
    ("NABH", "National Accreditation Board for Hospitals and Healthcare Providers"),
    ("OE",   "Objective Element"),
    ("PPE",  "Personal Protective Equipment"),
    ("RO",   "Reverse Osmosis"),
    ("STP",  "Sewage Treatment Plant"),
]

# Plain-text shortcuts (rulebook 3.2: guillemets only for «Hospital Name»)
_ENG = "Engineering In-Charge"
_QC  = "Quality Coordinator"
_MS  = "Medical Superintendent"
_MON = "monthly"
_QTR = "quarterly"
_YR  = "annually"
_OY  = "once a year"


# ══════════════════════════════════════════════════════════════════════════════
# FMS.1 — Safe and Secure Environment   (no stop-work)
# ══════════════════════════════════════════════════════════════════════════════
def gen_fms1():
    doc = Document()
    h(doc, 0, "Policy on Safe and Secure Environment")
    p(doc, HN)
    h(doc, 1, "Document control")
    doc_ctrl(doc, "HCO/FMS/POL/01", "Engineering In-Charge")
    p(doc, "A blank marked ________ must be completed before issue.")
    h(doc, 1, "Statement of intent")
    p(doc, "The organisation has a system in place to provide a safe and secure environment.")
    h(doc, 1, "1. Purpose")
    p(doc,
      f"This policy defines how {HN} installs and inspects patient-safety devices and infrastructure, "
      f"provides accessible facilities for differently-abled persons, conducts monthly safety-inspection "
      f"rounds, documents findings and acts on them, and carries out risk assessments before any "
      f"construction, renovation or expansion of the facility.")
    p(doc,
      "This policy covers safe and secure environment specifically. Related duties — like patient "
      "assessment, clinical care, medication management, patient rights, infection control, quality "
      "and safety monitoring, or hospital governance — are covered in the hospital's other policies, "
      "not repeated here. Other FMS standards have their own policies too.")
    p(doc, "Words marked like this are defaults. A blank marked ________ must be filled before issue.")
    h(doc, 1, "2. Scope")
    p(doc,
      f"This policy applies to engineering, biomedical, nursing and departmental leaders, and staff "
      f"who run facilities, utilities, medical gases, fire and non-fire emergencies at {HN}, including "
      f"the {_ENG}, the {_MS}, departmental leaders and the Quality Coordinator.")
    p(doc,
      "This policy covers safe and secure environment specifically. Related duties — like patient "
      "assessment, clinical care, medication management, patient rights, infection control, quality "
      "and safety monitoring, or hospital governance — are covered in the hospital's other policies, "
      "not repeated here. Other FMS standards have their own policies too.")
    h(doc, 1, "3. Policy standards")
    p(doc,
      f"Patient-safety devices and infrastructure are installed and periodically inspected across {HN}. "
      "Accessibility facilities meet regulatory minimums for differently-abled persons. Monthly "
      "facility-inspection rounds identify and monitor safety, security-risk and restricted areas. "
      "Every finding is documented, acted on and reviewed by the safety committee. No construction, "
      "renovation or expansion begins without a completed risk assessment covering noise, vibration "
      "and infection prevention and control.")
    p(doc, "Staff follow written guidance and keep the records listed in the traceability table.")
    h(doc, 1, "4. Non-negotiable rules")
    ln(doc,
       "Do not run a patient-care area that lacks required patient-safety devices (grab bars, bed rails, "
       "call bells, alarms, warning signs and fire-safety devices as applicable to that area) or for "
       "which no periodic inspection record exists.")
    ln(doc,
       "Do not operate without providing at minimum the accessibility facilities for differently-abled "
       "persons that applicable regulatory requirements mandate — a wheelchair-accessible entrance and "
       "an adapted toilet at minimum.")
    ln(doc,
       "Do not let a calendar month pass without a completed, checklist-based facility-inspection round; "
       "a round without a completed checklist is not a round for this purpose.")
    ln(doc,
       "Do not leave a finding from a facility inspection round without a documented corrective and "
       "preventive action and a safety-committee review within the same calendar month.")
    ln(doc,
       "Do not start construction, renovation or expansion of the existing hospital without a completed "
       "risk assessment covering noise, vibration and infection prevention and control in place before "
       "work begins.")
    ln(doc,
       "Staff who see a FMS.1 rule broken report it the same shift to the Engineering In-Charge or "
       "the Medical Superintendent.")
    h(doc, 1, "5. What we do")
    h(doc, 2, "5.1 Patient safety devices and infrastructure are installed across the...")
    p(doc,
      f"Patient-safety devices and infrastructure (grab bars, bed rails, stretcher and wheelchair "
      f"belts, call bells, alarms, radiation or biohazard warning signs, fire-safety devices as listed "
      f"for each area) are installed across {HN} and inspected at a defined interval (default {_MON}).")
    p(doc, f"The {_ENG} holds the current list and last inspection. A missing device in a care area "
            "that is in use is escalated the same shift.")
    h(doc, 2, "5.2 The organisation has facilities for the differently-abled")
    p(doc,
      f"{HN} has facilities for the differently-abled (at a minimum a wheelchair-accessible entrance "
      f"and an adapted toilet, as regulatory requirement and this hospital's building allow).")
    p(doc, f"The {_ENG} keeps the current access list. A step-only entrance with no documented "
            "alternative does not satisfy this element.")
    h(doc, 2, "5.3 Facility inspection rounds to ensure safety are conducted at least once...")
    p(doc,
      "Facility inspection rounds to ensure safety are conducted at least once a month using a "
      "checklist. Potential safety and security-risk / restricted areas are identified and monitored.")
    p(doc, f"The {_ENG} owns the round calendar. A month without a completed round is a gap.")
    h(doc, 2, "5.4 Inspection reports of facility rounds are documented and corrective and...")
    p(doc,
      "Inspection reports of facility rounds are documented. Corrective and preventive measures are "
      "undertaken. The safety committee reviews the reports monthly. Pre- and post-correction evidence "
      "is kept for at least one accreditation cycle.")
    p(doc, f"The {_QC} files the reviewed reports with the {_ENG}.")
    h(doc, 2, "5.5 Before construction, renovation and expansion of the existing hospital...")
    p(doc,
      "Before construction, renovation or expansion of the existing hospital, a risk-assessment is "
      "carried out covering noise, vibration and infection prevention and control. IPC.4 remains the "
      "home of construction-infection controls; this element is that the assessment happens before "
      "work starts.")
    p(doc, f"The {_ENG} holds the dated assessment. Work started with no assessment does not satisfy "
            "this element.")
    # No stop-work — section 6 = Governance
    h(doc, 1, "6. Governance and responsibility")
    gov_tbl(doc, fms_gov("FMS.1"))
    h(doc, 1, "7. Quality monitoring (RCA → CAPA)")
    p(doc, "The Quality Coordinator audits this policy quarterly. The audit reviews:")
    lb(doc, "Records for a sample of this standard's objective elements, checked against the What-we-do steps.")
    lb(doc, "CORE objective elements show no critical gaps in the sample.")
    lb(doc, "Stop-work events (if any) are logged with outcome.")
    p(doc, "Root-cause analysis is required when a gap found in this audit remains open beyond 90 days.")
    p(doc, "This policy is reviewed annually, and sooner after a related facility change, "
           "utility failure, equipment recall or fire-plan change.")
    h(doc, 1, "8. Training and staff acknowledgement")
    p(doc, f"Staff covered by this policy are trained at induction and {_OY} after that. Training "
           "covers the What-we-do steps, non-negotiables and stop-work (if present).")
    p(doc, "Staff acknowledgement")
    p(doc, f"I have read the Policy on Safe and Secure Environment of {HN}. I will follow the processes described.")
    p(doc, "Name: ___________________________    Designation: ___________________________")
    p(doc, "Department / floor: ____________________    Date: ____________")
    p(doc, "Signature: ___________________________")
    p(doc, f"(One row per staff member. The {_QC} holds signed acknowledgements with the induction record.)")
    h(doc, 1, "9. Distribution")
    p(doc, "Medical Superintendent; Engineering In-Charge; Quality Coordinator; departmental leaders; "
           "staff covered by FMS.1")
    h(doc, 1, "10. Abbreviations")
    abbrev_tbl(doc, FMS_ABBREVS)
    h(doc, 1, "11. Traceability to NABH HCO Full Accreditation 6th Edition FMS.1")
    p(doc, "This table is an index. It is not how the policy is organised. An asterisk in the Level "
           "column means documentation of the process is required.")
    trace_tbl(doc, [
        ("FMS.1.a", "CORE",
         "Patient safety devices and infrastructure are installed across the organisation and inspected periodically.",
         "Section 3; 5.1", "Engineering In-Charge"),
        ("FMS.1.b", "Commitment",
         "The organisation has facilities for the differently-abled.",
         "Section 3; 5.2", "Engineering In-Charge"),
        ("FMS.1.c", "CORE",
         "Facility inspection rounds to ensure safety are conducted at least once a month.",
         "Section 3; 5.3", "Engineering In-Charge"),
        ("FMS.1.d", "Commitment",
         "Inspection reports of facility rounds are documented and corrective and preventive measures are undertaken.",
         "Section 3; 5.4", "Engineering In-Charge"),
        ("FMS.1.e", "Commitment",
         "Before construction, renovation and expansion of the existing hospital, risk-assessment is carried out.",
         "Section 3; 5.5", "Engineering In-Charge"),
    ])
    h(doc, 1, "12. Required Records / Evidence Checklist")
    p(doc, "Records the hospital holds under this policy, listed by objective element.")
    h(doc, 2, "FMS.1.a — Patient safety devices and infrastructure are installed across the "
              "organisation and inspected periodically.")
    lb(doc, "Current inventory of patient-safety devices and infrastructure (grab bars, bed rails, call bells, fire-safety devices) by area.")
    lb(doc, "Inspection log showing the last-inspection date for each device or area.")
    lb(doc, "Escalation record for any missing device found in an in-use care area.")
    h(doc, 2, "FMS.1.b — The organisation has facilities for the differently-abled.")
    lb(doc, "Facility list showing the wheelchair-accessible entrance and toilet location(s).")
    lb(doc, "Regulatory-requirement reference for the accessibility provision made.")
    lb(doc, "Documented alternative arrangement for any area with step-only access.")
    h(doc, 2, "FMS.1.c — Facility inspection rounds to ensure safety are conducted at least once a month.")
    lb(doc, "Monthly facility-inspection-round checklist, completed and dated.")
    lb(doc, "Round calendar showing no missed month.")
    lb(doc, "List of identified safety, security-risk or restricted areas being monitored.")
    h(doc, 2, "FMS.1.d — Inspection reports of facility rounds are documented and corrective "
              "and preventive measures are undertaken.")
    lb(doc, "Documented inspection reports of facility rounds.")
    lb(doc, "Corrective and preventive action record for each finding.")
    lb(doc, "Safety-committee review record of the reports.")
    h(doc, 2, "FMS.1.e — Before construction, renovation and expansion of the existing hospital, "
              "risk-assessment is carried out.")
    lb(doc, "Dated risk-assessment record covering noise, vibration and infection prevention, completed before work started.")
    lb(doc, "Engineering In-Charge sign-off before construction, renovation or expansion began.")
    lb(doc, "Cross-reference to the IPC.4 construction-infection-control measures applied alongside.")
    h(doc, 1, "13. References")
    lb(doc, "National Accreditation Board for Hospitals and Healthcare Providers (NABH), Accreditation "
            "Standards for Hospitals, 6th Edition (January 2025) — Facility Management and Safety, standard FMS.1.")
    lb(doc, "NABH Guidebook to Accreditation Standards for Hospitals, 6th Edition — FMS.1 interpretations.")
    lb(doc, f"Internal documents of {HN}: facility-inspection records, as-built drawings, utility and "
            "medical-equipment logs, medical-gas records, fire and non-fire plans named for FMS.1.")
    h(doc, 1, "Disclaimer")
    disclaimer(doc)
    save_and_verify(doc, "HCO_FMS_1_v2_REWRITE_DRAFT.docx")


# ══════════════════════════════════════════════════════════════════════════════
# FMS.2 — Planned Facilities and Environment-Friendly Measures   (stop-work: FMS.2.d)
# ══════════════════════════════════════════════════════════════════════════════
def gen_fms2():
    doc = Document()
    h(doc, 0, "Policy on Planned Facilities and Environment-Friendly Measures")
    p(doc, HN)
    h(doc, 1, "Document control")
    doc_ctrl(doc, "HCO/FMS/POL/02", "Engineering In-Charge")
    p(doc, "A blank marked ________ must be completed before issue.")
    h(doc, 1, "Statement of intent")
    p(doc, "The organisation's environment and facilities operate in a planned manner and promotes "
           "environment-friendly measures.")
    h(doc, 1, "1. Purpose")
    p(doc,
      f"This policy defines how {HN} ensures facilities and space match services, keeps as-built "
      "drawings current, maintains comprehensible signage, provides potable water and electricity "
      "around the clock, and provides and regularly tests backup sources for both.")
    p(doc,
      "This policy covers planned facilities and environment-friendly measures specifically. Related "
      "duties — like patient assessment, clinical care, medication management, patient rights, "
      "infection control, quality and safety monitoring, or hospital governance — are covered in "
      "the hospital's other policies, not repeated here. Other FMS standards have their own policies too.")
    p(doc, "Words marked like this are defaults. A blank marked ________ must be filled before issue.")
    h(doc, 1, "2. Scope")
    p(doc,
      f"This policy applies to engineering, biomedical, nursing and departmental leaders, and staff "
      f"who run facilities, utilities, medical gases, fire and non-fire emergencies at {HN}, including "
      f"the {_ENG}, the {_MS}, departmental leaders and the Quality Coordinator.")
    p(doc,
      "This policy covers planned facilities and environment-friendly measures specifically. Related "
      "duties — like patient assessment, clinical care, medication management, patient rights, "
      "infection control, quality and safety monitoring, or hospital governance — are covered in "
      "the hospital's other policies, not repeated here. Other FMS standards have their own policies too.")
    h(doc, 1, "3. Policy standards")
    p(doc,
      f"Facilities and space at {HN} match the services offered. As-built and updated drawings are "
      "maintained by a named custodian. Internal and external signage is in a form patients, families "
      "and the community can understand, and meets statutory posting requirements. Potable water and "
      "electricity are available round the clock with tested backup sources for any failure.")
    p(doc, "Staff follow written guidance and keep the records listed in the traceability table.")
    h(doc, 1, "4. Non-negotiable rules")
    ln(doc,
       "Do not operate a clinical service without matching facility space; any service without a "
       "corresponding documented space is not available from this hospital for accreditation purposes.")
    ln(doc,
       "Do not let the as-built and updated drawing set become incomplete or without a named custodian; "
       "a drawing that does not reflect the current facility is not an as-built drawing.")
    ln(doc,
       "Do not allow signage that cannot be understood by patients, families or the community, or that "
       "does not meet applicable statutory posting requirements.")
    ln(doc,
       "Do not leave a care area without potable water or electricity; test potable-water quality "
       "biochemically at least once in three months and microbiologically at least once a month, "
       "collected at the tap.")
    ln(doc,
       "Do not operate without identified backup electricity and water sources available for every "
       "critical area.")
    ln(doc,
       "Do not count an alternate source as available if it has not been tested at the defined "
       "frequency with documented results.")
    ln(doc, "Do not bypass the stop-work authority in section 6 when the trigger conditions are met.")
    ln(doc,
       "Staff who see a FMS.2 rule broken report it the same shift to the Engineering In-Charge or "
       "the Medical Superintendent.")
    h(doc, 1, "5. What we do")
    h(doc, 2, "5.1 Facilities and space provisions are appropriate to the scope of services")
    p(doc,
      "Facilities and space provisions match the scope of services this hospital actually offers, "
      "using national or international guidance (including Atomic Energy Regulatory Board guidance "
      "where radiation services are in scope).")
    p(doc,
      f"The {_ENG} and {_MS} keep the current space-vs-services map. A service listed without a "
      "matching space is escalated.")
    h(doc, 2, "5.2 As built and updated drawings are maintained as per statutory...")
    p(doc,
      "As-built and updated drawings are maintained as required by the applicable registering "
      "authority for this site: site layout, floor plans, floor-wise fire-evacuation plans, and "
      "separate civil, electrical, extra-low-voltage, plumbing, heating-ventilation-air-conditioning, "
      "piped medical-gas and information-technology drawings.")
    p(doc,
      f"A named person under the {_ENG} holds the current set. This element does not import a named "
      "Act as a checklist for the whole standard.")
    h(doc, 2, "5.3 There are internal and external sign postings in the organisation in a...")
    p(doc,
      "Internal and external sign posting is in a manner patients, families and the community "
      "understand (language and/or pictorial; bilingual where this hospital defines). Signs meet "
      "applicable statutory posting rules for this site.")
    p(doc, f"The {_ENG} walks the signs {_QTR}.")
    h(doc, 2, "5.4 Potable water and electricity are available round the clock")
    p(doc,
      "Potable water and electricity are available round the clock. Potable-water quality is tested "
      "at the tap: biochemical at least once in three months and microbiological at least once a month, "
      "against the current IS 10500. Dialysis reverse-osmosis inlet water is tested for endotoxin "
      "every month where dialysis is in scope.")
    p(doc, "A care area without water or power is a stop-work trigger (section 6).")
    h(doc, 2, "5.5 Alternate sources for electricity and water are provided as a backup...")
    p(doc,
      "Alternate sources for electricity and water are provided as backup for any failure or shortage "
      "(diesel generator, solar, uninterruptible power supply; bore or tanker or extra tanks as this "
      "hospital names). Electric load matches demand. National Building Code is a reference for water "
      "quantity, not a paste of every clause.")
    p(doc, f"The {_ENG} owns the backup list. Critical areas have a named continuity action when supply fails.")
    h(doc, 2, "5.6 The organisation tests the functioning of these alternate sources at a...")
    p(doc,
      f"The organisation tests these alternate sources at a predefined frequency (default {_MON} for "
      "the diesel generator and a documented water-acceptance test when an emergency source is used). "
      "Results are recorded. Refer to FMS.2.d for water quality.")
    p(doc, "A backup that is never tested does not satisfy this element.")
    h(doc, 1, "6. Stop-work authority")
    p(doc,
      "Do not start or continue clinical care in an area that has no potable water or no electricity, "
      "unless the organisation has declared a FMS.7 emergency and the written service-continuity plan "
      "for that failure is running.")
    p(doc,
      "Stop-work applies to starting or continuing routine care in that area. Immediate life-saving "
      "care continues while water or power is restored or the patient is moved.")
    p(doc,
      f"The person who stops tells the {_ENG} and the {_MS} the same shift. Refusing to run care "
      "without water or power is not a disciplinary matter.")
    h(doc, 1, "7. Governance and responsibility")
    gov_tbl(doc, fms_gov("FMS.2"))
    h(doc, 1, "8. Quality monitoring (RCA → CAPA)")
    p(doc, "The Quality Coordinator audits this policy quarterly. The audit reviews:")
    lb(doc, "Records for a sample of this standard's objective elements, checked against the What-we-do steps.")
    lb(doc, "CORE objective elements show no critical gaps in the sample.")
    lb(doc, "Stop-work events (if any) are logged with outcome.")
    p(doc, "Root-cause analysis is required when a gap found in this audit remains open beyond 90 days.")
    p(doc, "This policy is reviewed annually, and sooner after a related facility change, "
           "utility failure, equipment recall or fire-plan change.")
    h(doc, 1, "9. Training and staff acknowledgement")
    p(doc, f"Staff covered by this policy are trained at induction and {_OY} after that. Training "
           "covers the What-we-do steps, non-negotiables and stop-work (if present).")
    p(doc, "Staff acknowledgement")
    p(doc, f"I have read the Policy on Planned Facilities and Environment-Friendly Measures of {HN}. "
           "I will follow the processes described.")
    p(doc, "Name: ___________________________    Designation: ___________________________")
    p(doc, "Department / floor: ____________________    Date: ____________")
    p(doc, "Signature: ___________________________")
    p(doc, f"(One row per staff member. The {_QC} holds signed acknowledgements with the induction record.)")
    h(doc, 1, "10. Distribution")
    p(doc, "Medical Superintendent; Engineering In-Charge; Quality Coordinator; departmental leaders; "
           "staff covered by FMS.2")
    h(doc, 1, "11. Abbreviations")
    abbrev_tbl(doc, FMS_ABBREVS)
    h(doc, 1, "12. Traceability to NABH HCO Full Accreditation 6th Edition FMS.2")
    p(doc, "This table is an index. It is not how the policy is organised. An asterisk in the Level "
           "column means documentation of the process is required.")
    trace_tbl(doc, [
        ("FMS.2.a", "Commitment",
         "Facilities and space provisions are appropriate to the scope of services.",
         "Section 3; 5.1", "Engineering In-Charge"),
        ("FMS.2.b", "Commitment",
         "As built and updated drawings are maintained as per statutory requirements.",
         "Section 3; 5.2", "Engineering In-Charge"),
        ("FMS.2.c", "CORE",
         "There are internal and external sign postings in the organisation in a manner understood by the patient, families and community.",
         "Section 3; 5.3", "Engineering In-Charge"),
        ("FMS.2.d", "CORE",
         "Potable water and electricity are available round the clock.",
         "Section 3; 5.4; Section 6 Stop-work", "Engineering In-Charge"),
        ("FMS.2.e", "Commitment",
         "Alternate sources for electricity and water are provided as a backup for any failure / shortage.",
         "Section 3; 5.5", "Engineering In-Charge"),
        ("FMS.2.f", "Commitment",
         "The organisation tests the functioning of these alternate sources at a predefined frequency.",
         "Section 3; 5.6", "Engineering In-Charge"),
    ])
    h(doc, 1, "13. Required Records / Evidence Checklist")
    p(doc, "Records the hospital holds under this policy, listed by objective element.")
    h(doc, 2, "FMS.2.a — Facilities and space provisions are appropriate to the scope of services.")
    lb(doc, "Current space-versus-services map for each department.")
    lb(doc, "National or international guidance reference used (for example AERB guidance for a radiation area).")
    lb(doc, "Escalation record for any service listed without matching space.")
    h(doc, 2, "FMS.2.b — As built and updated drawings are maintained as per statutory requirements.")
    lb(doc, "Current as-built and updated drawing set: site layout, floor plans, floor-wise evacuation plans, and separate civil, electrical, ELV, plumbing, HVAC, medical-gas and IT drawings.")
    lb(doc, "Named custodian record for the drawing set.")
    lb(doc, "Date of last update against a facility change.")
    h(doc, 2, "FMS.2.c — There are internal and external sign postings in the organisation in a "
              "manner understood by the patient, families and community.")
    lb(doc, "Signage inventory or walk log by area, noting language and/or pictorial form and bilingual coverage.")
    lb(doc, "Statutory posting-requirement reference applicable to this site.")
    lb(doc, "Quarterly signage-walk record.")
    h(doc, 2, "FMS.2.d — Potable water and electricity are available round the clock.")
    lb(doc, "Potable-water test results: biochemical at least quarterly and microbiological at least monthly, against IS 10500.")
    lb(doc, "Dialysis reverse-osmosis inlet-water endotoxin test record, where dialysis is in scope.")
    lb(doc, "Continuity log confirming no unplanned water or power outage in a care area without the FMS.7 continuity plan activating.")
    h(doc, 2, "FMS.2.e — Alternate sources for electricity and water are provided as a backup "
              "for any failure / shortage.")
    lb(doc, "Backup-source inventory: diesel generator, solar, uninterruptible power supply, bore/tanker/extra tanks as named.")
    lb(doc, "Electric-load calculation matching demand.")
    lb(doc, "Continuity-action record naming each critical area's backup path.")
    h(doc, 2, "FMS.2.f — The organisation tests the functioning of these alternate sources at a predefined frequency.")
    lb(doc, "Diesel-generator test log at the defined frequency (default monthly).")
    lb(doc, "Water-acceptance test record for any occasion an emergency water source was used.")
    lb(doc, "Corrective-action record for any failed alternate-source test.")
    h(doc, 1, "14. References")
    lb(doc, "National Accreditation Board for Hospitals and Healthcare Providers (NABH), Accreditation "
            "Standards for Hospitals, 6th Edition (January 2025) — Facility Management and Safety, standard FMS.2.")
    lb(doc, "NABH Guidebook to Accreditation Standards for Hospitals, 6th Edition — FMS.2 interpretations.")
    lb(doc, f"Internal documents of {HN}: facility-inspection records, as-built drawings, utility and "
            "medical-equipment logs, medical-gas records, fire and non-fire plans named for FMS.2.")
    h(doc, 1, "Disclaimer")
    disclaimer(doc)
    save_and_verify(doc, "HCO_FMS_2_v2_REWRITE_DRAFT.docx")


# ══════════════════════════════════════════════════════════════════════════════
# FMS.3 — Safety of Patients, Families, Staff and Visitors   (stop-work: FMS.3.e)
# ══════════════════════════════════════════════════════════════════════════════
def gen_fms3():
    doc = Document()
    h(doc, 0, "Policy on Safety of Patients, Families, Staff and Visitors")
    p(doc, HN)
    h(doc, 1, "Document control")
    doc_ctrl(doc, "HCO/FMS/POL/03", "Engineering In-Charge")
    p(doc, "A blank marked ________ must be completed before issue.")
    h(doc, 1, "Statement of intent")
    p(doc, "The organisation's environment and facilities operate to ensure the safety of patients, "
           "their families, staff and visitors.")
    h(doc, 1, "1. Purpose")
    p(doc,
      f"This policy defines how {HN} controls access to high-security areas, identifies and safely "
      "manages hazardous materials, implements spill-response plans, conducts electrical safety "
      "audits, and manages material not in use.")
    p(doc,
      "This policy covers safety of patients, families, staff and visitors specifically. Related "
      "duties — like patient assessment, clinical care, medication management, patient rights, "
      "infection control, quality and safety monitoring, or hospital governance — are covered in "
      "the hospital's other policies, not repeated here. Other FMS standards have their own policies too.")
    p(doc, "Words marked like this are defaults. A blank marked ________ must be filled before issue.")
    h(doc, 1, "2. Scope")
    p(doc,
      f"This policy applies to engineering, biomedical, nursing and departmental leaders, and staff "
      f"who run facilities, utilities, medical gases, fire and non-fire emergencies at {HN}, including "
      f"the {_ENG}, the {_MS}, departmental leaders and the Quality Coordinator.")
    p(doc,
      "This policy covers safety of patients, families, staff and visitors specifically. Related "
      "duties — like patient assessment, clinical care, medication management, patient rights, "
      "infection control, quality and safety monitoring, or hospital governance — are covered in "
      "the hospital's other policies, not repeated here. Other FMS standards have their own policies too.")
    h(doc, 1, "3. Policy standards")
    p(doc,
      f"{HN} defines extra-security areas and controls access for staff, patients and visitors. "
      "Hazardous materials are identified, documented and handled safely at every stage. Spill plans "
      "are implemented with floor-accessible summaries and kits. Electrical safety audits are "
      "conducted annually. Material not in use is systematically identified and disposed of.")
    p(doc, "Staff follow written guidance and keep the records listed in the traceability table.")
    h(doc, 1, "4. Non-negotiable rules")
    ln(doc,
       "Do not operate critical areas (operating theatre, ICUs including NICU, labour room, emergency) "
       "without a written security plan defining access for staff, patients and visitors; designated "
       "extra-security areas must have documented controls such as CCTV.")
    ln(doc,
       "Do not start or continue using a hazardous material that has not been identified and documented, "
       "or for which a sorting, storage, handling, transport and disposal procedure does not exist.")
    ln(doc,
       "Do not use a hazardous material in an area where the summarised Material Safety Data Sheet is "
       "not accessible to floor staff and the hazardous-materials spill kit is not reachable.")
    ln(doc,
       "Do not condemn or dispose of material not in use without following the written "
       "identification-and-disposal procedure.")
    ln(doc,
       "Do not allow a calendar year to pass without a completed electrical safety audit of the "
       "facility with documented actions.")
    ln(doc, "Do not bypass the stop-work authority in section 6 when the trigger conditions are met.")
    ln(doc,
       "Staff who see a FMS.3 rule broken report it the same shift to the Engineering In-Charge or "
       "the Medical Superintendent.")
    h(doc, 1, "5. What we do")
    h(doc, 2, "5.1 Operational planning identifies areas which need to have extra security...")
    p(doc,
      "Operational planning identifies extra-security areas and describes access for staff, patients "
      "and visitors (at a minimum operating theatre, intensive-care units including neonatal if in "
      "scope, labour room and emergency). Vulnerable spots (dark areas, long corridors, critical-area "
      "doors) have a defined control such as closed-circuit television.")
    p(doc, "Written security guidance exists.")
    h(doc, 2, "5.2 Patient safety aspects in terms of structural safety of hospitals...")
    p(doc,
      "When this hospital plans, designs or constructs a new building, or re-plans or retrofits an "
      "existing one, patient-safety structural aspects of critical areas are considered. Indian Seismic "
      "Code IS 1893 (Part 1), latest version, is the minimum structural reference named in the Guidebook.")
    p(doc,
      f"The {_ENG} files what was applied. This element is evidenced on actual building work, not by "
      "a policy sentence alone.")
    h(doc, 2, "5.3 The organisation conducts electrical safety audits for the facility")
    p(doc,
      "The organisation conducts electrical safety audits of the facility at least once a year to "
      "reduce risk to people and property and to prevent fire from short-circuiting. National Electrical "
      "Code of India 2023 is a reference. Thermal imaging may be used.")
    p(doc, f"The {_ENG} holds the last audit and actions.")
    h(doc, 2, "5.4 There is a procedure which addresses the identification and disposal of...")
    p(doc,
      "There is a written procedure for identifying and disposing of material not in use "
      "(non-functioning items, excess stock, general scrap). Condemnation records sit with the "
      f"{_ENG}.")
    h(doc, 2, "5.5 Hazardous materials are identified and used safely within the...")
    p(doc,
      "Hazardous materials used here are identified and used safely: sorting, storage, handling, "
      "transport and disposal, using Material Safety Data Sheets. Common examples the Guidebook names "
      "include chemicals, blood and cultures, mercury, nuclear isotopes, medical gases, liquefied "
      "petroleum gas, steam and ethylene oxide.")
    p(doc, "Using an unidentified hazardous material is a stop-work trigger (section 6).")
    h(doc, 2, "5.6 The plan for managing spills of hazardous materials is implemented")
    p(doc,
      "The spill plan for hazardous materials is implemented: a summarised Material Safety Data Sheet "
      "the floor can read, a hazardous-materials kit where those materials are stored, and trained handlers.")
    p(doc, "A kit in a locked office the floor cannot reach does not satisfy this element.")
    h(doc, 1, "6. Stop-work authority")
    p(doc,
      "Do not use a hazardous material that has not been identified, or for which the spill plan is "
      "not implemented (no Material Safety Data Sheet path, no kit, or staff not trained for that material).")
    p(doc,
      "Stop-work applies to using that material. Immediate life-saving care that already depends on "
      "a stocked clinical product continues while the spill path is restored.")
    p(doc,
      f"The person who stops tells the {_ENG} and the {_MS} the same shift. Refusing an unidentified "
      "hazardous material is not a disciplinary matter.")
    h(doc, 1, "7. Governance and responsibility")
    gov_tbl(doc, fms_gov("FMS.3"))
    h(doc, 1, "8. Quality monitoring (RCA → CAPA)")
    p(doc, "The Quality Coordinator audits this policy quarterly. The audit reviews:")
    lb(doc, "Records for a sample of this standard's objective elements, checked against the What-we-do steps.")
    lb(doc, "Documentary evidence is on file for each asterisked objective element in the sample.")
    lb(doc, "CORE objective elements show no critical gaps in the sample.")
    lb(doc, "Stop-work events (if any) are logged with outcome.")
    p(doc, "Root-cause analysis is required when a gap found in this audit remains open beyond 90 days.")
    p(doc, "This policy is reviewed annually, and sooner after a related facility change, "
           "utility failure, equipment recall or fire-plan change.")
    h(doc, 1, "9. Training and staff acknowledgement")
    p(doc, f"Staff covered by this policy are trained at induction and {_OY} after that. Training "
           "covers the What-we-do steps, non-negotiables and stop-work (if present).")
    p(doc, "Staff acknowledgement")
    p(doc, f"I have read the Policy on Safety of Patients, Families, Staff and Visitors of {HN}. "
           "I will follow the processes described.")
    p(doc, "Name: ___________________________    Designation: ___________________________")
    p(doc, "Department / floor: ____________________    Date: ____________")
    p(doc, "Signature: ___________________________")
    p(doc, f"(One row per staff member. The {_QC} holds signed acknowledgements with the induction record.)")
    h(doc, 1, "10. Distribution")
    p(doc, "Medical Superintendent; Engineering In-Charge; Quality Coordinator; departmental leaders; "
           "staff covered by FMS.3")
    h(doc, 1, "11. Abbreviations")
    abbrev_tbl(doc, FMS_ABBREVS)
    h(doc, 1, "12. Traceability to NABH HCO Full Accreditation 6th Edition FMS.3")
    p(doc, "This table is an index. It is not how the policy is organised. An asterisk in the Level "
           "column means documentation of the process is required.")
    trace_tbl(doc, [
        ("FMS.3.a", "Commitment*",
         "Operational planning identifies areas which need to have extra security and describes access to different areas in the hospital by staff, patients, and visitors.",
         "Section 3; 5.1", "Engineering In-Charge"),
        ("FMS.3.b", "Excellence",
         "Patient safety aspects in terms of structural safety of hospitals, especially of critical areas are considered while planning, design and construction of new hospitals and re-planning, assessment and retrofitting of existing hospitals.",
         "Section 3; 5.2", "Engineering In-Charge"),
        ("FMS.3.c", "Commitment",
         "The organisation conducts electrical safety audits for the facility.",
         "Section 3; 5.3", "Engineering In-Charge"),
        ("FMS.3.d", "Commitment*",
         "There is a procedure which addresses the identification and disposal of material(s) not in use in the organisation.",
         "Section 3; 5.4", "Engineering In-Charge"),
        ("FMS.3.e", "CORE*",
         "Hazardous materials are identified and used safely within the organisation.",
         "Section 3; 5.5; Section 6 Stop-work", "Engineering In-Charge"),
        ("FMS.3.f", "Commitment*",
         "The plan for managing spills of hazardous materials is implemented.",
         "Section 3; 5.6", "Engineering In-Charge"),
    ])
    h(doc, 1, "13. Required Records / Evidence Checklist")
    p(doc, "Records the hospital holds under this policy, listed by objective element.")
    h(doc, 2, "FMS.3.a — Operational planning identifies areas which need to have extra security "
              "and describes access to different areas in the hospital by staff, patients, and visitors.")
    lb(doc, "Written security operational plan naming extra-security areas (operating theatre, ICU/NICU, labour room, emergency) and access rules for staff, patients and visitors.")
    lb(doc, "Vulnerable-spot control record — closed-circuit television coverage or equivalent for dark areas and long corridors.")
    lb(doc, "Review or update record of the security plan.")
    h(doc, 2, "FMS.3.b — Patient safety aspects in terms of structural safety of hospitals, "
              "especially of critical areas are considered while planning, design and construction of "
              "new hospitals and re-planning, assessment and retrofitting of existing hospitals.")
    lb(doc, "Structural-safety design reference actually applied (for example Indian Seismic Code IS 1893 Part 1) for the specific building project.")
    lb(doc, "Engineering sign-off record for that construction, re-planning or retrofit project.")
    lb(doc, "Evidence of what was applied to real building work, not a policy statement alone.")
    h(doc, 2, "FMS.3.c — The organisation conducts electrical safety audits for the facility.")
    lb(doc, "Completed electrical safety audit report, at least annually.")
    lb(doc, "Action-taken record against audit findings.")
    lb(doc, "Reference to the National Electrical Code or thermal-imaging method used.")
    h(doc, 2, "FMS.3.d — There is a procedure which addresses the identification and disposal "
              "of material(s) not in use in the organisation.")
    lb(doc, "Written procedure for identifying and disposing of material not in use.")
    lb(doc, "Condemnation register listing items, dates and disposal method.")
    lb(doc, "Engineering In-Charge sign-off on the register.")
    h(doc, 2, "FMS.3.e — Hazardous materials are identified and used safely within the organisation.")
    lb(doc, "Current hazardous-materials inventory (for example chemicals, blood/cultures, mercury, isotopes, medical gases, LPG, steam, ethylene oxide).")
    lb(doc, "Material Safety Data Sheet on file for each identified material.")
    lb(doc, "Handling, storage, transport and disposal record for a sample of materials.")
    h(doc, 2, "FMS.3.f — The plan for managing spills of hazardous materials is implemented.")
    lb(doc, "Summarised Material Safety Data Sheet accessible on the floor where the material is stored.")
    lb(doc, "Hazardous-materials spill-kit inventory and location log.")
    lb(doc, "Training record for staff handling that material.")
    h(doc, 1, "14. References")
    lb(doc, "National Accreditation Board for Hospitals and Healthcare Providers (NABH), Accreditation "
            "Standards for Hospitals, 6th Edition (January 2025) — Facility Management and Safety, standard FMS.3.")
    lb(doc, "NABH Guidebook to Accreditation Standards for Hospitals, 6th Edition — FMS.3 interpretations.")
    lb(doc, f"Internal documents of {HN}: facility-inspection records, as-built drawings, utility and "
            "medical-equipment logs, medical-gas records, fire and non-fire plans named for FMS.3.")
    h(doc, 1, "Disclaimer")
    disclaimer(doc)
    save_and_verify(doc, "HCO_FMS_3_v2_REWRITE_DRAFT.docx")


# ══════════════════════════════════════════════════════════════════════════════
# FMS.4 — Facility, Engineering Support and Utility Systems   (stop-work: FMS.4.c)
# ══════════════════════════════════════════════════════════════════════════════
def gen_fms4():
    doc = Document()
    h(doc, 0, "Policy on Facility, Engineering Support and Utility Systems")
    p(doc, HN)
    h(doc, 1, "Document control")
    doc_ctrl(doc, "HCO/FMS/POL/04", "Engineering In-Charge")
    p(doc, "A blank marked ________ must be completed before issue.")
    h(doc, 1, "Statement of intent")
    p(doc, "The organisation has a programme for the facility, engineering support services and utility system.")
    h(doc, 1, "1. Purpose")
    p(doc,
      f"This policy defines how {HN} plans, inventories, operates and maintains utility and "
      "engineering equipment, keeps competent maintenance personnel available round the clock, and "
      "guides equipment replacement and disposal.")
    p(doc,
      "This policy covers facility, engineering support and utility systems specifically. Related "
      "duties — like patient assessment, clinical care, medication management, patient rights, "
      "infection control, quality and safety monitoring, or hospital governance — are covered in "
      "the hospital's other policies, not repeated here. Other FMS standards have their own policies too.")
    p(doc, "Words marked like this are defaults. A blank marked ________ must be filled before issue.")
    h(doc, 1, "2. Scope")
    p(doc,
      f"This policy applies to engineering, biomedical, nursing and departmental leaders, and staff "
      f"who run facilities, utilities, medical gases, fire and non-fire emergencies at {HN}, including "
      f"the {_ENG}, the {_MS}, departmental leaders and the Quality Coordinator.")
    p(doc,
      "This policy covers facility, engineering support and utility systems specifically. Related "
      "duties — like patient assessment, clinical care, medication management, patient rights, "
      "infection control, quality and safety monitoring, or hospital governance — are covered in "
      "the hospital's other policies, not repeated here. Other FMS standards have their own policies too.")
    h(doc, 1, "3. Policy standards")
    p(doc,
      f"{HN} plans utility and engineering equipment against services and the strategic plan. All "
      "equipment is inventoried with unique identifiers. Implemented operational and maintenance plans "
      "cover every system. Calibration is kept current. Competent personnel are available for every "
      "shift. Maintenance is contactable round the clock. Downtime on critical equipment is tracked "
      "from complaint to completion. Equipment replacement and disposal follows written guidance.")
    p(doc, "Staff follow written guidance and keep the records listed in the traceability table.")
    h(doc, 1, "4. Non-negotiable rules")
    ln(doc, "Do not run utility or engineering equipment that is not in the current equipment inventory with a unique identifier.")
    ln(doc,
       "Do not run critical utility equipment (diesel generator, lifts, uninterruptible power supply, "
       "fire-related equipment, dialysis reverse-osmosis plant, water pumps) without an implemented, "
       "documented operational and maintenance plan.")
    ln(doc, "Do not let utility equipment become overdue for calibration without a corrective action in place.")
    ln(doc, f"Do not operate a shift without a named competent person available for each plant system that is running.")
    ln(doc, "Do not leave the maintenance escalation matrix unavailable at the nursing station and departments during any shift.")
    ln(doc, "Do not condemn or dispose of utility or engineering equipment without following the written equipment-replacement and disposal guidance.")
    ln(doc, "Do not bypass the stop-work authority in section 6 when the trigger conditions are met.")
    ln(doc,
       "Staff who see a FMS.4 rule broken report it the same shift to the Engineering In-Charge or "
       "the Medical Superintendent.")
    h(doc, 1, "5. What we do")
    h(doc, 2, "5.1 The organisation plans for utility and engineering equipment in...")
    p(doc,
      "The organisation plans utility and engineering equipment to match services and the strategic "
      "plan, including future needs (diesel generator, chiller). Plans are implemented and reviewed "
      "at a defined interval. Selection, rental, update or upgrade is collaborative (end-user, "
      "management, finance, engineering).")
    p(doc, f"The {_ENG} holds the current plan.")
    h(doc, 2, "5.2 Equipment is inventoried, and proper logs are maintained as required")
    p(doc,
      "Equipment is inventoried and logs are maintained. Each item has a unique identifier. "
      "Quality-conformance certificates, factory test certificates and installation reports are kept "
      "where they apply.")
    p(doc, "A plant item with no identity number does not satisfy this element.")
    h(doc, 2, "5.3 The documented operational and maintenance (preventive and breakdown)...")
    p(doc,
      "The documented operational and maintenance (preventive and breakdown) plan is implemented for "
      "utility and engineering equipment, electrical systems, water, heating-ventilation-air-conditioning, "
      "facility and furniture — including transformers, low-tension and high-tension panels, lifts, "
      "tanks, reverse-osmosis and sewage-treatment if present, chillers, air-handling units and filters.")
    p(doc, "Running critical utility equipment with no implemented plan is a stop-work trigger (section 6).")
    h(doc, 2, "5.4 Utility equipment, are periodically inspected and calibrated (wherever...")
    p(doc,
      "Utility equipment is periodically inspected and calibrated where applicable (for example "
      "steam-steriliser pressure gauges, medication-refrigerator temperature gauges), in-house or "
      "outsourced, with traceability to prescribed standards.")
    p(doc, f"The {_ENG} holds the calibration schedule.")
    h(doc, 2, "5.5 Competent personnel operate, inspect, test and maintain equipment and...")
    p(doc,
      "Competent personnel (qualification, experience or training) operate, inspect, test and maintain "
      "equipment and utility systems. Enough supervisors and tradespeople, including fire-safety and "
      "electrical-safety trained staff, plus tools and personal protective equipment, are available.")
    p(doc, f"A plant running with no competent person named for that shift is escalated to the {_ENG}.")
    h(doc, 2, "5.6 Maintenance staff is contactable round the clock for emergency repairs")
    p(doc,
      "Maintenance staff are contactable round the clock for emergency repairs. An escalation matrix "
      "(who to call if the person on duty cannot complete the job) is at the nursing station and departments.")
    p(doc, "A night call with no answer path does not satisfy this element.")
    h(doc, 2, "5.7 Downtime for critical equipment breakdowns is monitored from reporting...")
    p(doc,
      "Downtime for critical engineering and utility equipment breakdowns is monitored from reporting "
      "to inspection and corrective action. At a minimum the critical list includes diesel generator, "
      "lifts, uninterruptible power supply, fire-related equipment, dialysis reverse-osmosis and water "
      "pumps. A complaint register records receipt, job allotment and user-ratified completion.")
    p(doc, "Start of downtime is complaint time; end is user-ratified completion.")
    h(doc, 2, "5.8 Written guidance supports equipment replacement, identification of...")
    p(doc,
      "Written guidance supports equipment replacement, identification of unwanted material and "
      "disposal. Unusable utility and engineering equipment is condemned in a systematic way. Records are kept.")
    h(doc, 1, "6. Stop-work authority")
    p(doc,
      "Do not run critical utility equipment (as this hospital names it: at a minimum diesel generator, "
      "lifts, uninterruptible power supply, fire-related utility, reverse-osmosis plant for dialysis, "
      "water pumps) when there is no implemented operational and maintenance plan, or when the equipment "
      "is known unsafe.")
    p(doc,
      "Stop-work applies to starting or continuing that equipment. Life-saving care continues under "
      "the FMS.7 continuity plan.")
    p(doc,
      f"The person who stops tells the {_ENG} and the {_MS} the same shift. Refusing unsafe utility "
      "equipment is not a disciplinary matter.")
    h(doc, 1, "7. Governance and responsibility")
    gov_tbl(doc, fms_gov("FMS.4"))
    h(doc, 1, "8. Quality monitoring (RCA → CAPA)")
    p(doc, "The Quality Coordinator audits this policy quarterly. The audit reviews:")
    lb(doc, "Records for a sample of this standard's objective elements, checked against the What-we-do steps.")
    lb(doc, "Documentary evidence is on file for each asterisked objective element in the sample.")
    lb(doc, "CORE objective elements show no critical gaps in the sample.")
    lb(doc, "Stop-work events (if any) are logged with outcome.")
    p(doc, "Root-cause analysis is required when a gap found in this audit remains open beyond 90 days.")
    p(doc, "This policy is reviewed annually, and sooner after a related facility change, "
           "utility failure, equipment recall or fire-plan change.")
    h(doc, 1, "9. Training and staff acknowledgement")
    p(doc, f"Staff covered by this policy are trained at induction and {_OY} after that. Training "
           "covers the What-we-do steps, non-negotiables and stop-work (if present).")
    p(doc, "Staff acknowledgement")
    p(doc, f"I have read the Policy on Facility, Engineering Support and Utility Systems of {HN}. "
           "I will follow the processes described.")
    p(doc, "Name: ___________________________    Designation: ___________________________")
    p(doc, "Department / floor: ____________________    Date: ____________")
    p(doc, "Signature: ___________________________")
    p(doc, f"(One row per staff member. The {_QC} holds signed acknowledgements with the induction record.)")
    h(doc, 1, "10. Distribution")
    p(doc, "Medical Superintendent; Engineering In-Charge; Quality Coordinator; departmental leaders; "
           "staff covered by FMS.4")
    h(doc, 1, "11. Abbreviations")
    abbrev_tbl(doc, FMS_ABBREVS)
    h(doc, 1, "12. Traceability to NABH HCO Full Accreditation 6th Edition FMS.4")
    p(doc, "This table is an index. It is not how the policy is organised. An asterisk in the Level "
           "column means documentation of the process is required.")
    trace_tbl(doc, [
        ("FMS.4.a", "Commitment",
         "The organisation plans for utility and engineering equipment in accordance with its services and strategic plan.",
         "Section 3; 5.1", "Engineering In-Charge"),
        ("FMS.4.b", "Commitment",
         "Equipment is inventoried, and proper logs are maintained as required.",
         "Section 3; 5.2", "Engineering In-Charge"),
        ("FMS.4.c", "CORE*",
         "The documented operational and maintenance (preventive and breakdown) plan is implemented.",
         "Section 3; 5.3; Section 6 Stop-work", "Engineering In-Charge"),
        ("FMS.4.d", "Commitment",
         "Utility equipment, are periodically inspected and calibrated (wherever applicable) for their proper functioning.",
         "Section 3; 5.4", "Engineering In-Charge"),
        ("FMS.4.e", "Commitment",
         "Competent personnel operate, inspect, test and maintain equipment and utility systems.",
         "Section 3; 5.5", "Engineering In-Charge"),
        ("FMS.4.f", "Commitment",
         "Maintenance staff is contactable round the clock for emergency repairs.",
         "Section 3; 5.6", "Engineering In-Charge"),
        ("FMS.4.g", "Achievement",
         "Downtime for critical equipment breakdowns is monitored from reporting to inspection and implementation of corrective actions.",
         "Section 3; 5.7", "Engineering In-Charge"),
        ("FMS.4.h", "Commitment*",
         "Written guidance supports equipment replacement, identification of unwanted material and disposal.",
         "Section 3; 5.8", "Engineering In-Charge"),
    ])
    h(doc, 1, "13. Required Records / Evidence Checklist")
    p(doc, "Records the hospital holds under this policy, listed by objective element.")
    h(doc, 2, "FMS.4.a — The organisation plans for utility and engineering equipment in "
              "accordance with its services and strategic plan.")
    lb(doc, "Written utility and engineering equipment plan matched to services and the strategic plan, including future needs (for example diesel generator, chiller).")
    lb(doc, "Periodic review record of the plan.")
    lb(doc, "Record showing collaborative selection (end-user, management, finance, engineering) for equipment decisions.")
    h(doc, 2, "FMS.4.b — Equipment is inventoried, and proper logs are maintained as required.")
    lb(doc, "Equipment inventory with a unique identifier per item.")
    lb(doc, "Quality-conformance, factory test or installation certificate on file where applicable.")
    lb(doc, "Log-completeness check confirming no item without an identity.")
    h(doc, 2, "FMS.4.c — The documented operational and maintenance (preventive and breakdown) plan is implemented.")
    lb(doc, "Documented operational and maintenance (preventive and breakdown) plan for utility and engineering equipment.")
    lb(doc, "Preventive-maintenance schedule and completion log for the critical-equipment list.")
    lb(doc, "Breakdown-response record showing the plan was actually followed.")
    h(doc, 2, "FMS.4.d — Utility equipment, are periodically inspected and calibrated "
              "(wherever applicable) for their proper functioning.")
    lb(doc, "Calibration schedule for utility equipment (for example steam-steriliser pressure gauges, medication-refrigerator temperature gauges).")
    lb(doc, "Calibration certificate with traceability to a prescribed standard.")
    lb(doc, "In-house or outsourced calibration record.")
    h(doc, 2, "FMS.4.e — Competent personnel operate, inspect, test and maintain equipment and utility systems.")
    lb(doc, "Competency or training record for staff operating utility equipment.")
    lb(doc, "Roster showing a named competent person for each shift.")
    lb(doc, "Escalation record for any shift with no competent person named.")
    h(doc, 2, "FMS.4.f — Maintenance staff is contactable round the clock for emergency repairs.")
    lb(doc, "Escalation matrix displayed at the nursing station and departments.")
    lb(doc, "On-call roster for maintenance staff, round the clock.")
    lb(doc, "Night-call response-time log.")
    h(doc, 2, "FMS.4.g — Downtime for critical equipment breakdowns is monitored from reporting "
              "to inspection and implementation of corrective actions.")
    lb(doc, "Complaint register logging receipt, job allotment and user-ratified completion for critical equipment (diesel generator, lifts, UPS, fire-related, dialysis RO, water pumps).")
    lb(doc, "Downtime-duration record from complaint time to ratified completion.")
    lb(doc, "Trend report reviewed against the critical-equipment list.")
    h(doc, 2, "FMS.4.h — Written guidance supports equipment replacement, identification of "
              "unwanted material and disposal.")
    lb(doc, "Written equipment-replacement and disposal guidance.")
    lb(doc, "Condemnation record for utility and engineering equipment.")
    lb(doc, "Systematic disposal log.")
    h(doc, 1, "14. References")
    lb(doc, "National Accreditation Board for Hospitals and Healthcare Providers (NABH), Accreditation "
            "Standards for Hospitals, 6th Edition (January 2025) — Facility Management and Safety, standard FMS.4.")
    lb(doc, "NABH Guidebook to Accreditation Standards for Hospitals, 6th Edition — FMS.4 interpretations.")
    lb(doc, f"Internal documents of {HN}: facility-inspection records, as-built drawings, utility and "
            "medical-equipment logs, medical-gas records, fire and non-fire plans named for FMS.4.")
    h(doc, 1, "Disclaimer")
    disclaimer(doc)
    save_and_verify(doc, "HCO_FMS_4_v2_REWRITE_DRAFT.docx")


# ══════════════════════════════════════════════════════════════════════════════
# FMS.5 — Medical Equipment Management   (stop-work: FMS.5.c, FMS.5.d, FMS.5.g)
# ══════════════════════════════════════════════════════════════════════════════
_FMS5_P2 = (
    f"Several requirements in this document are statutory rather than advisory — in particular "
    f"those arising under the Medical Devices Rules as cited in NABH FMS.5 (Gazette of India "
    f"GSR 78(E) 2023) including adverse-event monitoring, hazard notices and recalls. Statutory "
    f"requirements change, and State authorities impose additional or stricter conditions. "
    f"«Hospital Name» is responsible for verifying the current text of any rule cited here "
    f"and the conditions attached to its own authorisations and licences; this document does not "
    f"constitute legal advice."
)

def gen_fms5():
    doc = Document()
    h(doc, 0, "Policy on Medical Equipment Management")
    p(doc, HN)
    h(doc, 1, "Document control")
    doc_ctrl(doc, "HCO/FMS/POL/05", "Engineering In-Charge")
    p(doc, "A blank marked ________ must be completed before issue.")
    h(doc, 1, "Statement of intent")
    p(doc, "The organisation has a programme for medical equipment management.")
    h(doc, 1, "1. Purpose")
    p(doc,
      f"This policy defines how {HN} plans medical-equipment procurement, inventories and identifies "
      "every device, implements operational and maintenance plans, keeps calibration current, ensures "
      "only qualified personnel operate and maintain equipment, monitors adverse events and recalls, "
      "and guides replacement and disposal.")
    p(doc,
      "This policy covers medical equipment management specifically. Related duties — like patient "
      "assessment, clinical care, medication management, patient rights, infection control, quality "
      "and safety monitoring, or hospital governance — are covered in the hospital's other "
      "policies, not repeated here. Other FMS standards have their own policies too.")
    p(doc, "Words marked like this are defaults. A blank marked ________ must be filled before issue.")
    h(doc, 1, "2. Scope")
    p(doc,
      f"This policy applies to engineering, biomedical, nursing and departmental leaders, and staff "
      f"who run facilities, utilities, medical gases, fire and non-fire emergencies at {HN}, including "
      f"the {_ENG}, the {_MS}, departmental leaders and the Quality Coordinator.")
    p(doc,
      "This policy covers medical equipment management specifically. Related duties — like patient "
      "assessment, clinical care, medication management, patient rights, infection control, quality "
      "and safety monitoring, or hospital governance — are covered in the hospital's other "
      "policies, not repeated here. Other FMS standards have their own policies too.")
    h(doc, 1, "3. Policy standards")
    p(doc,
      f"{HN} plans medical equipment against services and the strategic plan. Every device is "
      "inventoried, classified by risk and given a unique identifier. Implemented operational and "
      "maintenance plans cover operator training, daily checks and breakdown response. Calibration is "
      "current before commissioning and after every repair. Operators and maintainers are qualified. "
      "Adverse events and recalls are monitored and acted on without delay. Disposal follows written guidance.")
    p(doc, "Staff follow written guidance and keep the records listed in the traceability table.")
    h(doc, 1, "4. Non-negotiable rules")
    ln(doc, "Do not put medical equipment into clinical use without a unique identifier and an entry in the medical-equipment inventory.")
    ln(doc, "Do not use medical equipment that does not have an implemented, documented operational and maintenance plan.")
    ln(doc,
       "Do not use medical equipment that measures patient parameters past its scheduled calibration "
       "due date; remove it from clinical use until calibration is complete.")
    ln(doc,
       "Do not allow an operator to use a medical device they have not been trained for; do not allow "
       "maintenance by personnel who are not a biomedical or instrumentation engineer or technologist "
       "with relevant training and experience.")
    ln(doc, "Do not continue clinical use of any medical device subject to an open manufacturer or regulatory recall or hazard notice.")
    ln(doc, "Do not condemn or dispose of medical equipment without following the written replacement-and-disposal guidance.")
    ln(doc, "Do not bypass the stop-work authority in section 6 when the trigger conditions are met.")
    ln(doc,
       "Staff who see a FMS.5 rule broken report it the same shift to the Engineering In-Charge or "
       "the Medical Superintendent.")
    h(doc, 1, "5. What we do")
    h(doc, 2, "5.1 The organisation plans for medical equipment in accordance with its...")
    p(doc,
      "The organisation plans medical equipment to match services and the strategic plan, including "
      "future needs. Indian Public Health Standards are a reference for a minimum set. Selection, "
      "rental, update or upgrade is collaborative (end-user, management, finance, engineering, biomedical).")
    p(doc, f"The {_ENG} holds the current medical-equipment plan with the {_MS}.")
    h(doc, 2, "5.2 Medical equipment is inventoried and proper logs are maintained as...")
    p(doc,
      "Medical equipment is inventoried and logs are maintained. Items are classified by medical-device "
      "risk. Each has a unique identifier, including rental and demonstration items. Factory test and "
      "conformance certificates are kept.")
    p(doc, "An in-use device with no identity does not satisfy this element.")
    h(doc, 2, "5.3 The documented operational and maintenance (preventive and breakdown)...")
    p(doc,
      "The documented operational and maintenance (preventive and breakdown) plan for medical equipment "
      "is implemented: operator training, daily operating checks, preventive-maintenance tracker, and "
      "breakdown response including nights and weekends.")
    p(doc, "Using equipment with no implemented plan is a stop-work trigger (section 6).")
    h(doc, 2, "5.4 Medical equipment is periodically inspected and calibrated for their...")
    p(doc,
      "Medical equipment that measures is inspected and calibrated on a weekly, monthly or annual "
      "schedule as the manufacturer and this hospital define, in-house or outsourced, with traceability. "
      "Conformance is checked before commissioning and again after repair.")
    p(doc, "A measuring device past its calibration due date is a stop-work trigger (section 6).")
    h(doc, 2, "5.5 Qualified and trained personnel operate and maintain medical equipment")
    p(doc,
      "Qualified and trained personnel operate and maintain medical equipment. Operators are trained "
      "for the devices they use (for example blood-gas analyser, electrocardiograph, syringe pump). "
      "Maintenance is by a biomedical engineer or technologist, or an instrumentation engineer or "
      "technologist, with relevant training and experience.")
    h(doc, 2, "5.6 Written guidance supports medical equipment replacement and disposal")
    p(doc, "Written guidance supports medical equipment replacement and disposal. Condemnation is systematic and recorded.")
    h(doc, 2, "5.7 There is monitoring of medical equipment and medical devices related to...")
    p(doc,
      "Medical equipment and devices are monitored for adverse events. Hazard notices and recalls "
      "from the manufacturer or regulator are acted on at once; the device is not returned to clinical "
      "use until the issue is closed. The Guidebook names Gazette of India GSR 78(E) 2023 / Medical "
      "Devices Rules 2023 and participation in the Materiovigilance Programme of India. MOM.10 remains "
      "the home of medication-supply recall; this element is equipment and devices.")
    p(doc, "An open recall left in use is a stop-work trigger (section 6).")
    h(doc, 2, "5.8 Downtime for critical equipment breakdown is monitored from reporting...")
    p(doc,
      "Downtime for critical medical-equipment breakdown is monitored from reporting to inspection and "
      "corrective action. At a minimum the critical list includes ventilators, X-ray, magnetic resonance "
      "imaging, catheterisation laboratory, computed tomography, anaesthesia machines, monitors, "
      "laboratory and ultrasound — especially where there is no alternative. A complaint register "
      "records receipt, job allotment and user-ratified completion.")
    p(doc, "Start of downtime is complaint time; end is user-ratified completion.")
    h(doc, 1, "6. Stop-work authority")
    p(doc,
      "Do not use medical equipment that is past the due date for preventive maintenance or calibration, "
      "or that is under an open manufacturer or regulatory recall or hazard notice.")
    p(doc,
      "Stop-work applies to putting that equipment into clinical use. Immediate life-saving care uses "
      "the next safe alternative and is documented.")
    p(doc,
      f"The person who stops tells the {_ENG} and the {_MS} the same shift. Refusing overdue or "
      "recalled equipment is not a disciplinary matter.")
    h(doc, 1, "7. Governance and responsibility")
    gov_tbl(doc, fms_gov("FMS.5"))
    h(doc, 1, "8. Quality monitoring (RCA → CAPA)")
    p(doc, "The Quality Coordinator audits this policy quarterly. The audit reviews:")
    lb(doc, "Records for a sample of this standard's objective elements, checked against the What-we-do steps.")
    lb(doc, "Documentary evidence is on file for each asterisked objective element in the sample.")
    lb(doc, "CORE objective elements show no critical gaps in the sample.")
    lb(doc, "Stop-work events (if any) are logged with outcome.")
    p(doc, "Root-cause analysis is required when a gap found in this audit remains open beyond 90 days.")
    p(doc, "This policy is reviewed annually, and sooner after a related facility change, "
           "utility failure, equipment recall or fire-plan change.")
    h(doc, 1, "9. Training and staff acknowledgement")
    p(doc, f"Staff covered by this policy are trained at induction and {_OY} after that. Training "
           "covers the What-we-do steps, non-negotiables and stop-work (if present).")
    p(doc, "Staff acknowledgement")
    p(doc, f"I have read the Policy on Medical Equipment Management of {HN}. I will follow the processes described.")
    p(doc, "Name: ___________________________    Designation: ___________________________")
    p(doc, "Department / floor: ____________________    Date: ____________")
    p(doc, "Signature: ___________________________")
    p(doc, f"(One row per staff member. The {_QC} holds signed acknowledgements with the induction record.)")
    h(doc, 1, "10. Distribution")
    p(doc, "Medical Superintendent; Engineering In-Charge; Quality Coordinator; departmental leaders; "
           "staff covered by FMS.5")
    h(doc, 1, "11. Abbreviations")
    abbrev_tbl(doc, FMS_ABBREVS)
    h(doc, 1, "12. Traceability to NABH HCO Full Accreditation 6th Edition FMS.5")
    p(doc, "This table is an index. It is not how the policy is organised. An asterisk in the Level "
           "column means documentation of the process is required.")
    trace_tbl(doc, [
        ("FMS.5.a", "Commitment",
         "The organisation plans for medical equipment in accordance with its services and strategic plan.",
         "Section 3; 5.1", "Engineering In-Charge"),
        ("FMS.5.b", "Commitment",
         "Medical equipment is inventoried and proper logs are maintained as required.",
         "Section 3; 5.2", "Engineering In-Charge"),
        ("FMS.5.c", "CORE*",
         "The documented operational and maintenance (preventive and breakdown) plan for medical equipment is implemented.",
         "Section 3; 5.3; Section 6 Stop-work", "Engineering In-Charge"),
        ("FMS.5.d", "Commitment",
         "Medical equipment is periodically inspected and calibrated for their proper functioning.",
         "Section 3; 5.4; Section 6 Stop-work", "Engineering In-Charge"),
        ("FMS.5.e", "Commitment",
         "Qualified and trained personnel operate and maintain medical equipment.",
         "Section 3; 5.5", "Engineering In-Charge"),
        ("FMS.5.f", "Commitment*",
         "Written guidance supports medical equipment replacement and disposal.",
         "Section 3; 5.6", "Engineering In-Charge"),
        ("FMS.5.g", "Commitment*",
         "There is monitoring of medical equipment and medical devices related to adverse events, and compliance hazard notices on recalls.",
         "Section 3; 5.7; Section 6 Stop-work", "Engineering In-Charge"),
        ("FMS.5.h", "Achievement",
         "Downtime for critical equipment breakdown is monitored from reporting to inspection and implementation of corrective actions.",
         "Section 3; 5.8", "Engineering In-Charge"),
    ])
    h(doc, 1, "13. Required Records / Evidence Checklist")
    p(doc, "Records the hospital holds under this policy, listed by objective element.")
    h(doc, 2, "FMS.5.a — The organisation plans for medical equipment in accordance with its services and strategic plan.")
    lb(doc, "Written medical-equipment plan matched to services and the strategic plan, referencing the Indian Public Health Standards minimum set.")
    lb(doc, "Collaborative-selection record (end-user, management, finance, engineering, biomedical).")
    lb(doc, "Periodic review record of the plan.")
    h(doc, 2, "FMS.5.b — Medical equipment is inventoried and proper logs are maintained as required.")
    lb(doc, "Medical-equipment inventory classified by device risk, with a unique identifier including rental and demonstration items.")
    lb(doc, "Factory test or conformance certificate on file.")
    lb(doc, "In-use device check confirming no device without an identity.")
    h(doc, 2, "FMS.5.c — The documented operational and maintenance (preventive and breakdown) "
              "plan for medical equipment is implemented.")
    lb(doc, "Documented operational and maintenance (preventive and breakdown) plan for medical equipment.")
    lb(doc, "Operator-training record and daily operating-check log.")
    lb(doc, "Breakdown-response record, including nights and weekends.")
    h(doc, 2, "FMS.5.d — Medical equipment is periodically inspected and calibrated for their proper functioning.")
    lb(doc, "Calibration schedule (weekly, monthly or annual as the manufacturer defines) with traceability.")
    lb(doc, "Pre-commissioning and post-repair conformance-check record.")
    lb(doc, "Calibration-due tracking log confirming no overdue device in clinical use.")
    h(doc, 2, "FMS.5.e — Qualified and trained personnel operate and maintain medical equipment.")
    lb(doc, "Operator-training record per device type (for example blood-gas analyser, electrocardiograph, syringe pump).")
    lb(doc, "Biomedical or instrumentation engineer/technologist qualification record for maintenance staff.")
    lb(doc, "Training-currency check for operators and maintainers.")
    h(doc, 2, "FMS.5.f — Written guidance supports medical equipment replacement and disposal.")
    lb(doc, "Written medical-equipment replacement and disposal guidance.")
    lb(doc, "Condemnation record, systematically applied.")
    lb(doc, "Disposal log.")
    h(doc, 2, "FMS.5.g — There is monitoring of medical equipment and medical devices related to "
              "adverse events, and compliance hazard notices on recalls.")
    lb(doc, "Adverse-event and hazard-notice/recall log for medical equipment and devices.")
    lb(doc, "Record showing a recalled device was withdrawn from clinical use until the issue closed.")
    lb(doc, "Materiovigilance Programme of India participation record, where applicable.")
    h(doc, 2, "FMS.5.h — Downtime for critical equipment breakdown is monitored from reporting "
              "to inspection and implementation of corrective actions.")
    lb(doc, "Complaint register for critical medical equipment (ventilators, X-ray, MRI, cath lab, CT, anaesthesia machines, monitors, laboratory, ultrasound).")
    lb(doc, "Downtime-duration record from reporting to corrective action.")
    lb(doc, "Alternative-equipment-use record where no backup device exists.")
    h(doc, 1, "14. References")
    lb(doc, "National Accreditation Board for Hospitals and Healthcare Providers (NABH), Accreditation "
            "Standards for Hospitals, 6th Edition (January 2025) — Facility Management and Safety, standard FMS.5.")
    lb(doc, "NABH Guidebook to Accreditation Standards for Hospitals, 6th Edition — FMS.5 interpretations.")
    lb(doc, f"Internal documents of {HN}: facility-inspection records, as-built drawings, utility and "
            "medical-equipment logs, medical-gas records, fire and non-fire plans named for FMS.5.")
    h(doc, 1, "Disclaimer")
    disclaimer(doc, _FMS5_P2)
    save_and_verify(doc, "HCO_FMS_5_v2_REWRITE_DRAFT.docx")


# ══════════════════════════════════════════════════════════════════════════════
# FMS.6 — Medical Gases, Vacuum and Compressed Air   (stop-work: FMS.6.b, FMS.6.d)
# ══════════════════════════════════════════════════════════════════════════════
_FMS6_P2 = (
    "Several requirements in this document are statutory rather than advisory — in particular "
    "those arising under statutory provisions for medical gases as named in NABH FMS.6, including "
    "the Explosives Act, Gas Cylinder Rules and Static and Mobile Pressure Vessels (Unfired) Rules "
    "where they apply to this hospital's gases. Statutory requirements change, and State authorities "
    "impose additional or stricter conditions. «Hospital Name» is responsible for verifying "
    "the current text of any rule cited here and the conditions attached to its own authorisations "
    "and licences; this document does not constitute legal advice."
)

def gen_fms6():
    doc = Document()
    h(doc, 0, "Policy on Medical Gases, Vacuum and Compressed Air")
    p(doc, HN)
    h(doc, 1, "Document control")
    doc_ctrl(doc, "HCO/FMS/POL/06", "Engineering In-Charge")
    p(doc, "A blank marked ________ must be completed before issue.")
    h(doc, 1, "Statement of intent")
    p(doc, "The organisation has a programme for medical gases, vacuum and compressed air.")
    h(doc, 1, "1. Purpose")
    p(doc,
      f"This policy defines how {HN} governs the procurement, handling, storage, distribution, use "
      "and replenishment of medical gases, compressed air and vacuum; maintains safety measures at "
      "every level; keeps operational and maintenance plans current; and provides and tests backup sources.")
    p(doc,
      "This policy covers medical gases, vacuum and compressed air specifically. Related duties — "
      "like patient assessment, clinical care, medication management, patient rights, infection control, "
      "quality and safety monitoring, or hospital governance — are covered in the hospital's other "
      "policies, not repeated here. Other FMS standards have their own policies too.")
    p(doc, "Words marked like this are defaults. A blank marked ________ must be filled before issue.")
    h(doc, 1, "2. Scope")
    p(doc,
      f"This policy applies to engineering, biomedical, nursing and departmental leaders, and staff "
      f"who run facilities, utilities, medical gases, fire and non-fire emergencies at {HN}, including "
      f"the {_ENG}, the {_MS}, departmental leaders and the Quality Coordinator.")
    p(doc,
      "This policy covers medical gases, vacuum and compressed air specifically. Related duties — "
      "like patient assessment, clinical care, medication management, patient rights, infection control, "
      "quality and safety monitoring, or hospital governance — are covered in the hospital's other "
      "policies, not repeated here. Other FMS standards have their own policies too.")
    h(doc, 1, "3. Policy standards")
    p(doc,
      f"Written guidance governs every stage of medical-gas management at {HN}. Gases are handled, "
      "stored and distributed with standardised colour coding, alarms, valve boxes, pin-indexed outlets "
      "and automatic changeover. An operational, inspection, testing and maintenance plan follows the "
      "manufacturer. Backup sources are in place and tested regularly.")
    p(doc, "Staff follow written guidance and keep the records listed in the traceability table.")
    h(doc, 1, "4. Non-negotiable rules")
    ln(doc,
       "Do not procure, store, distribute or use any medical gas without written guidance covering "
       "colour coding, signage, handling and replenishment in place for that gas.")
    ln(doc,
       "Do not use a medical-gas outlet or manifold that lacks the required colour coding, alarm, "
       "valve box, pin-indexed outlet or automatic changeover to the alternate source.")
    ln(doc, "Do not silence a plant-room alarm for a medical-gas system without a documented reason and corrective action.")
    ln(doc, "Do not operate piped medical gas, compressed air or vacuum without an implemented operational, inspection, testing and maintenance plan following the manufacturer.")
    ln(doc,
       "Do not operate without a required alternate source (stand-by compressor, stand-by vacuum pump, "
       "stand-by manifold or bulk cylinders) for each gas, compressed air and vacuum line in use.")
    ln(doc, "Do not count an alternate source as available if it has not been tested at the defined frequency with documented results.")
    ln(doc, "Do not bypass the stop-work authority in section 6 when the trigger conditions are met.")
    ln(doc,
       "Staff who see a FMS.6 rule broken report it the same shift to the Engineering In-Charge or "
       "the Medical Superintendent.")
    h(doc, 1, "5. What we do")
    h(doc, 2, "5.1 Written guidance governs the implementation of procurement, handling...")
    p(doc,
      "Written guidance governs procurement, handling, storage, distribution, usage and replenishment "
      "of all medical gases used here, including colour coding and full/empty signage. HTM 02-01 or "
      "the National Fire Protection Association medical-gas handbook is a reference.")
    p(doc,
      "The Guidebook names the Explosives Act, Gas Cylinder Rules and Static and Mobile Pressure "
      "Vessels (Unfired) Rules. Those duties stay with the named statute; this policy does not paste the Act.")
    h(doc, 2, "5.2 Medical gases are handled, stored, distributed and used in a safe manner")
    p(doc,
      "Medical gases are handled, stored, distributed and used safely: colour-coded cylinders and "
      "pipelines; alarms and valve boxes; twenty-four-hour monitoring of plant alarms; pin-indexed "
      "outlets; automatic change-over to the alternate source.")
    p(doc, "A live leak or a silenced required alarm is a stop-work trigger (section 6).")
    h(doc, 2, "5.3 There is an operational, inspection, testing and maintenance plan for...")
    p(doc,
      "There is an operational, inspection, testing and maintenance plan for piped medical gas, "
      "compressed air and vacuum, following the manufacturer. Compressed-air purity is checked at a "
      "terminal outlet at least once a year, at least one terminal in the operating theatre and one "
      "in intensive care if those areas exist.")
    h(doc, 2, "5.4 Alternate sources for medical gases, vacuum and compressed air are...")
    p(doc,
      "Alternate sources for medical gases, vacuum and compressed air are provided for failure "
      "(stand-by compressor and vacuum pump; stand-by manifold or bulk cylinders).")
    p(doc, "No required alternate source is a stop-work trigger (section 6).")
    h(doc, 2, "5.5 The organisation regularly tests the functioning of these alternate...")
    p(doc, f"The organisation tests these alternate sources at a defined frequency (default {_MON}). Results are documented.")
    h(doc, 1, "6. Stop-work authority")
    p(doc,
      "Do not use a medical-gas outlet or manifold when there is a live leak, a required alarm that "
      "has been silenced without a recorded reason, or no required alternate source for that gas, "
      "vacuum or compressed air.")
    p(doc, "Stop-work applies to using that outlet or manifold. Immediate life-saving care uses cylinders or the written alternate source.")
    p(doc,
      f"The person who stops tells the {_ENG} and the {_MS} the same shift. Refusing an unsafe gas "
      "outlet is not a disciplinary matter.")
    h(doc, 1, "7. Governance and responsibility")
    gov_tbl(doc, fms_gov("FMS.6"))
    h(doc, 1, "8. Quality monitoring (RCA → CAPA)")
    p(doc, "The Quality Coordinator audits this policy quarterly. The audit reviews:")
    lb(doc, "Records for a sample of this standard's objective elements, checked against the What-we-do steps.")
    lb(doc, "Documentary evidence is on file for each asterisked objective element in the sample.")
    lb(doc, "CORE objective elements show no critical gaps in the sample.")
    lb(doc, "Stop-work events (if any) are logged with outcome.")
    p(doc, "Root-cause analysis is required when a gap found in this audit remains open beyond 90 days.")
    p(doc, "This policy is reviewed annually, and sooner after a related facility change, "
           "utility failure, equipment recall or fire-plan change.")
    h(doc, 1, "9. Training and staff acknowledgement")
    p(doc, f"Staff covered by this policy are trained at induction and {_OY} after that. Training "
           "covers the What-we-do steps, non-negotiables and stop-work (if present).")
    p(doc, "Staff acknowledgement")
    p(doc, f"I have read the Policy on Medical Gases, Vacuum and Compressed Air of {HN}. I will follow the processes described.")
    p(doc, "Name: ___________________________    Designation: ___________________________")
    p(doc, "Department / floor: ____________________    Date: ____________")
    p(doc, "Signature: ___________________________")
    p(doc, f"(One row per staff member. The {_QC} holds signed acknowledgements with the induction record.)")
    h(doc, 1, "10. Distribution")
    p(doc, "Medical Superintendent; Engineering In-Charge; Quality Coordinator; departmental leaders; "
           "staff covered by FMS.6")
    h(doc, 1, "11. Abbreviations")
    abbrev_tbl(doc, FMS_ABBREVS)
    h(doc, 1, "12. Traceability to NABH HCO Full Accreditation 6th Edition FMS.6")
    p(doc, "This table is an index. It is not how the policy is organised. An asterisk in the Level "
           "column means documentation of the process is required.")
    trace_tbl(doc, [
        ("FMS.6.a", "Commitment*",
         "Written guidance governs the implementation of procurement, handling, storage, distribution, usage and replenishment of medical gases.",
         "Section 3; 5.1", "Engineering In-Charge"),
        ("FMS.6.b", "CORE",
         "Medical gases are handled, stored, distributed and used in a safe manner.",
         "Section 3; 5.2; Section 6 Stop-work", "Engineering In-Charge"),
        ("FMS.6.c", "Commitment*",
         "There is an operational, inspection, testing and maintenance plan for piped medical gas, compressed air and vacuum installation.",
         "Section 3; 5.3", "Engineering In-Charge"),
        ("FMS.6.d", "CORE",
         "Alternate sources for medical gases, vacuum and compressed air are provided for, in case of failure.",
         "Section 3; 5.4; Section 6 Stop-work", "Engineering In-Charge"),
        ("FMS.6.e", "Commitment",
         "The organisation regularly tests the functioning of these alternate sources.",
         "Section 3; 5.5", "Engineering In-Charge"),
    ])
    h(doc, 1, "13. Required Records / Evidence Checklist")
    p(doc, "Records the hospital holds under this policy, listed by objective element.")
    h(doc, 2, "FMS.6.a — Written guidance governs the implementation of procurement, handling, "
              "storage, distribution, usage and replenishment of medical gases.")
    lb(doc, "Written medical-gas procurement, handling, storage, distribution, usage and replenishment guidance, including colour-coding and signage.")
    lb(doc, "Reference used (for example HTM 02-01 or the NFPA medical-gas handbook).")
    lb(doc, "Statutory-provision record (Explosives Act, Gas Cylinder Rules, Static and Mobile Pressure Vessels Rules) where those apply to this hospital's gases.")
    h(doc, 2, "FMS.6.b — Medical gases are handled, stored, distributed and used in a safe manner.")
    lb(doc, "Colour-coded cylinder and pipeline inventory with alarm and valve-box log.")
    lb(doc, "Twenty-four-hour plant-alarm monitoring record.")
    lb(doc, "Pin-indexed outlet and automatic-changeover verification record.")
    h(doc, 2, "FMS.6.c — There is an operational, inspection, testing and maintenance plan for "
              "piped medical gas, compressed air and vacuum installation.")
    lb(doc, "Documented operational, inspection, testing and maintenance plan for piped medical gas, compressed air and vacuum.")
    lb(doc, "Compressed-air purity test record, at least annually, at operating-theatre and intensive-care terminal outlets.")
    lb(doc, "Manufacturer-following maintenance log.")
    h(doc, 2, "FMS.6.d — Alternate sources for medical gases, vacuum and compressed air are "
              "provided for, in case of failure.")
    lb(doc, "Stand-by compressor/vacuum-pump and stand-by manifold or bulk-cylinder inventory.")
    lb(doc, "Automatic-changeover test record.")
    lb(doc, "Alternate-source readiness confirmation log.")
    h(doc, 2, "FMS.6.e — The organisation regularly tests the functioning of these alternate sources.")
    lb(doc, "Test log for alternate medical-gas sources at the defined interval (default monthly).")
    lb(doc, "Manufacturer or written-plan interval reference used.")
    lb(doc, "Corrective-action record for any failed test.")
    h(doc, 1, "14. References")
    lb(doc, "National Accreditation Board for Hospitals and Healthcare Providers (NABH), Accreditation "
            "Standards for Hospitals, 6th Edition (January 2025) — Facility Management and Safety, standard FMS.6.")
    lb(doc, "NABH Guidebook to Accreditation Standards for Hospitals, 6th Edition — FMS.6 interpretations.")
    lb(doc, f"Internal documents of {HN}: facility-inspection records, as-built drawings, utility and "
            "medical-equipment logs, medical-gas records, fire and non-fire plans named for FMS.6.")
    h(doc, 1, "Disclaimer")
    disclaimer(doc, _FMS6_P2)
    save_and_verify(doc, "HCO_FMS_6_v2_REWRITE_DRAFT.docx")


# ══════════════════════════════════════════════════════════════════════════════
# FMS.7 — Fire and Non-Fire Emergencies   (stop-work: FMS.7.a, FMS.7.c)
# ══════════════════════════════════════════════════════════════════════════════
def gen_fms7():
    doc = Document()
    h(doc, 0, "Policy on Fire and Non-Fire Emergencies")
    p(doc, HN)
    h(doc, 1, "Document control")
    doc_ctrl(doc, "HCO/FMS/POL/07", "Engineering In-Charge")
    p(doc, "A blank marked ________ must be completed before issue.")
    h(doc, 1, "Statement of intent")
    p(doc, "The organisation has plans for fire and non-fire emergencies within the facilities.")
    h(doc, 1, "1. Purpose")
    p(doc,
      f"This policy defines how {HN} plans and maintains provisions for fire and non-fire emergencies "
      "— including fire detection, abatement, containment and evacuation; non-fire emergency "
      "identification and management; displayed exit plans; regular mock drills; and maintenance of "
      "fire-related equipment and infrastructure.")
    p(doc,
      "This policy covers fire and non-fire emergencies specifically. Related duties — like patient "
      "assessment, clinical care, medication management, patient rights, infection control, quality "
      "and safety monitoring, or hospital governance — are covered in the hospital's other "
      "policies, not repeated here. Other FMS standards have their own policies too.")
    p(doc, "Words marked like this are defaults. A blank marked ________ must be filled before issue.")
    h(doc, 1, "2. Scope")
    p(doc,
      f"This policy applies to engineering, biomedical, nursing and departmental leaders, and staff "
      f"who run facilities, utilities, medical gases, fire and non-fire emergencies at {HN}, including "
      f"the {_ENG}, the {_MS}, departmental leaders and the Quality Coordinator.")
    p(doc,
      "This policy covers fire and non-fire emergencies specifically. Related duties — like patient "
      "assessment, clinical care, medication management, patient rights, infection control, quality "
      "and safety monitoring, or hospital governance — are covered in the hospital's other "
      "policies, not repeated here. Other FMS standards have their own policies too.")
    h(doc, 1, "3. Policy standards")
    p(doc,
      f"{HN} has implemented fire plans covering detection, abatement, containment and evacuation, "
      "with qualified personnel and current NABH fire-safety measures. Non-fire emergencies are "
      "identified and planned for, with NDMA/State/District guidelines as a reference. Exit plans "
      "are documented and displayed on every floor. Mock drills are held at least twice a year. "
      "Fire-related equipment and infrastructure are maintained under an active plan.")
    p(doc, "Staff follow written guidance and keep the records listed in the traceability table.")
    h(doc, 1, "4. Non-negotiable rules")
    ln(doc,
       "Do not occupy a patient-care floor without an implemented fire plan covering detection, "
       "abatement, containment and evacuation — with qualified personnel, current NABH "
       "fire-safety measures, smoke-control provisions and emergency illumination in place.")
    ln(doc,
       "Do not occupy a patient-care floor without a documented and displayed exit plan on that floor, "
       "including near lifts and inside enclosed rooms and laboratories; exit doors must remain open "
       "or have push bars.")
    ln(doc,
       "Do not operate without written plans for the non-fire emergencies this hospital has identified "
       "(at minimum earthquake, flood, structural collapse, utility failure and toxic leak), developed "
       "with reference to NDMA/State/District guidelines.")
    ln(doc,
       "Do not let six months pass without at least one mock drill testing the full fire or non-fire "
       "emergency plan; each drill uses simulated, not real, patients and is followed by a debrief "
       "and corrective action.")
    ln(doc,
       "Do not operate fire-related equipment and infrastructure without an implemented maintenance "
       "plan covering inspection, testing, preventive and breakdown maintenance.")
    ln(doc, "Do not bypass the stop-work authority in section 6 when the trigger conditions are met.")
    ln(doc,
       "Staff who see a FMS.7 rule broken report it the same shift to the Engineering In-Charge or "
       "the Medical Superintendent.")
    h(doc, 1, "5. What we do")
    h(doc, 2, "5.1 The organisation has plans and provisions for early detection...")
    p(doc,
      f"{HN} has plans and provisions for early detection, abatement, containment of fire and "
      "evacuation: a fire plan (inflammable items, explosion, short-circuit, negligence or "
      "incompetence), qualified personnel, current NABH minimum fire-safety measures, smoke control, "
      "training, mock-drill schedules including table-top, drill records, displayed exits, evacuation "
      "of patients, staff and visitors, and emergency illumination. The hospital establishes liaison "
      "with civil and police authorities and the fire brigade as required by law.")
    p(doc, "Occupying a patient-care floor without required detection, abatement or evacuation provision is a stop-work trigger (section 6).")
    h(doc, 2, "5.2 The organisation has plans and provisions for identification, and...")
    p(doc,
      "The organisation has plans and provisions for identification and management of non-fire "
      "emergencies (examples the Guidebook names include earthquake, flood, toxic leak, structural "
      "collapse, utility failure, boiler burst, violence, stray animals). National Disaster Management "
      "Authority / State / District guidelines are referred to. Liaison with civil, police and fire "
      "authorities is as required by law.")
    p(doc, "Portal wording “identification, and management” is kept as printed.")
    h(doc, 2, "5.3 The organisation has a documented and displayed exit plan in case of...")
    p(doc,
      "A documented exit plan is displayed on each floor, especially near lifts and inside enclosed "
      "rooms and laboratories. Exit doors remain open or have push bars. Fire signage follows the fire "
      "service and/or National Building Code. Refuge areas are signed and maintained where they apply.")
    p(doc, "A patient-care floor with no displayed exit plan is a stop-work trigger (section 6).")
    h(doc, 2, "5.4 Mock drills are held at least twice a year")
    p(doc,
      "Mock drills are held at least twice a year (minimum; more often if this hospital defines). "
      "This covers fire and the important non-fire events this hospital names. A table-top exercise "
      "or a mock drill may be used; at a minimum one mock drill every six months tests the whole plan, "
      "not only awareness. Simulated patients, not real patients, are used. Variations are debriefed "
      "and corrected.")
    p(doc, "A year with fewer than two drills is a gap under this element.")
    h(doc, 2, "5.5 There is a maintenance plan for fire-related equipment and...")
    p(doc,
      "There is a maintenance plan for fire-related equipment and infrastructure: inspection, testing, "
      "preventive and breakdown maintenance, following the manufacturer and applicable statutory recommendations.")
    p(doc, f"The {_ENG} holds the plan and last service evidence.")
    h(doc, 1, "6. Stop-work authority")
    p(doc,
      "Do not occupy a patient-care floor when required fire detection, abatement or evacuation "
      "provision is not in place, or when the exit plan for that floor is not displayed.")
    p(doc,
      "Stop-work applies to placing or keeping patients on that floor as a planned location of care. "
      "Immediate life-saving evacuation and life-saving care continue.")
    p(doc,
      f"The person who stops tells the {_ENG} and the {_MS} the same shift. Refusing to occupy an "
      "unprotected floor is not a disciplinary matter.")
    h(doc, 1, "7. Governance and responsibility")
    gov_tbl(doc, fms_gov("FMS.7"))
    h(doc, 1, "8. Quality monitoring (RCA → CAPA)")
    p(doc, "The Quality Coordinator audits this policy quarterly. The audit reviews:")
    lb(doc, "Records for a sample of this standard's objective elements, checked against the What-we-do steps.")
    lb(doc, "Documentary evidence is on file for each asterisked objective element in the sample.")
    lb(doc, "CORE objective elements show no critical gaps in the sample.")
    lb(doc, "Stop-work events (if any) are logged with outcome.")
    p(doc, "Root-cause analysis is required when a gap found in this audit remains open beyond 90 days.")
    p(doc, "This policy is reviewed annually, and sooner after a related facility change, "
           "utility failure, equipment recall or fire-plan change.")
    h(doc, 1, "9. Training and staff acknowledgement")
    p(doc, f"Staff covered by this policy are trained at induction and {_OY} after that. Training "
           "covers the What-we-do steps, non-negotiables and stop-work (if present).")
    p(doc, "Staff acknowledgement")
    p(doc, f"I have read the Policy on Fire and Non-Fire Emergencies of {HN}. I will follow the processes described.")
    p(doc, "Name: ___________________________    Designation: ___________________________")
    p(doc, "Department / floor: ____________________    Date: ____________")
    p(doc, "Signature: ___________________________")
    p(doc, f"(One row per staff member. The {_QC} holds signed acknowledgements with the induction record.)")
    h(doc, 1, "10. Distribution")
    p(doc, "Medical Superintendent; Engineering In-Charge; Quality Coordinator; departmental leaders; "
           "staff covered by FMS.7")
    h(doc, 1, "11. Abbreviations")
    abbrev_tbl(doc, FMS_ABBREVS)
    h(doc, 1, "12. Traceability to NABH HCO Full Accreditation 6th Edition FMS.7")
    p(doc, "This table is an index. It is not how the policy is organised. An asterisk in the Level "
           "column means documentation of the process is required.")
    trace_tbl(doc, [
        ("FMS.7.a", "CORE*",
         "The organisation has plans and provisions for early detection, abatement, containment of fire and evacuation in the event of fire emergencies.",
         "Section 3; 5.1; Section 6 Stop-work", "Engineering In-Charge"),
        ("FMS.7.b", "CORE*",
         "The organisation has plans and provisions for identification, and management of non-fire emergencies.",
         "Section 3; 5.2", "Engineering In-Charge"),
        ("FMS.7.c", "Commitment",
         "The organisation has a documented and displayed exit plan in case of fire and non-fire emergencies.",
         "Section 3; 5.3; Section 6 Stop-work", "Engineering In-Charge"),
        ("FMS.7.d", "Commitment",
         "Mock drills are held at least twice a year.",
         "Section 3; 5.4", "Engineering In-Charge"),
        ("FMS.7.e", "Commitment*",
         "There is a maintenance plan for fire-related equipment and infrastructure.",
         "Section 3; 5.5", "Engineering In-Charge"),
    ])
    h(doc, 1, "13. Required Records / Evidence Checklist")
    p(doc, "Records the hospital holds under this policy, listed by objective element.")
    h(doc, 2, "FMS.7.a — The organisation has plans and provisions for early detection, abatement, "
              "containment of fire and evacuation in the event of fire emergencies.")
    lb(doc, "Written fire-safety plan covering detection, abatement, containment and evacuation, naming qualified personnel and current NABH minimum fire-safety measures.")
    lb(doc, "Mock-drill schedule and drill record, including table-top exercises.")
    lb(doc, "Displayed evacuation plan and emergency-illumination check record.")
    h(doc, 2, "FMS.7.b — The organisation has plans and provisions for identification, and "
              "management of non-fire emergencies.")
    lb(doc, "Written non-fire-emergency plan (for example earthquake, flood, toxic leak, structural collapse, utility failure, boiler burst, violence, stray animals).")
    lb(doc, "NDMA, State or District guideline reference used.")
    lb(doc, "Liaison record with civil, police and fire authorities.")
    h(doc, 2, "FMS.7.c — The organisation has a documented and displayed exit plan in case of "
              "fire and non-fire emergencies.")
    lb(doc, "Exit plan displayed on each floor, near lifts and inside enclosed rooms and laboratories.")
    lb(doc, "Exit-door check record (open or push-bar) and fire-signage reference (fire service or National Building Code).")
    lb(doc, "Refuge-area signage and maintenance record, where applicable.")
    h(doc, 2, "FMS.7.d — Mock drills are held at least twice a year.")
    lb(doc, "Mock-drill record showing at least two drills a year covering fire and the named non-fire events.")
    lb(doc, "Debrief and corrective-action record for variations found.")
    lb(doc, "Confirmation that simulated, not real, patients were used.")
    h(doc, 2, "FMS.7.e — There is a maintenance plan for fire-related equipment and infrastructure.")
    lb(doc, "Written maintenance plan for fire-related equipment and infrastructure.")
    lb(doc, "Inspection, testing, preventive and breakdown maintenance log following manufacturer and statutory recommendations.")
    lb(doc, "Last-service evidence on file.")
    h(doc, 1, "14. References")
    lb(doc, "National Accreditation Board for Hospitals and Healthcare Providers (NABH), Accreditation "
            "Standards for Hospitals, 6th Edition (January 2025) — Facility Management and Safety, standard FMS.7.")
    lb(doc, "NABH Guidebook to Accreditation Standards for Hospitals, 6th Edition — FMS.7 interpretations.")
    lb(doc, f"Internal documents of {HN}: facility-inspection records, as-built drawings, utility and "
            "medical-equipment logs, medical-gas records, fire and non-fire plans named for FMS.7.")
    h(doc, 1, "Disclaimer")
    disclaimer(doc)
    save_and_verify(doc, "HCO_FMS_7_v2_REWRITE_DRAFT.docx")


# ══════════════════════════════════════════════════════════════════════════════
if __name__ == "__main__":
    gen_fms1()
    gen_fms2()
    gen_fms3()
    gen_fms4()
    gen_fms5()
    gen_fms6()
    gen_fms7()
    print("\nAll 7 FMS drafts generated.")
