# -*- coding: utf-8 -*-
"""
Generate HCO_ROM_4_v2_REWRITE_DRAFT.docx
Plain-language rewrite per HCO_Policy_Rewrite_Rulebook.pdf
No stop-work: Governance=6, QM=7, Training=8, Distribution=9,
Abbreviations=10, Traceability=11, Records=12, References=13, Disclaimer (unnumbered)
CORE: b. Achievement: d, e. Asterisked: none.
"""
from docx import Document

OUT = "policies/build/rewrite_reference/HCO_ROM_4_v2_REWRITE_DRAFT.docx"

HN = "«Hospital Name»"
EM = "—"
doc = Document()


def h(level, text):
    styles = {0: "Title", 1: "Heading 1", 2: "Heading 2"}
    return doc.add_paragraph(text, style=styles[level])


def p(text):
    return doc.add_paragraph(text, style="Normal")


def ln(text):
    return doc.add_paragraph(text, style="List Number")


def lb(text):
    return doc.add_paragraph(text, style="List Bullet")


def tbl(rows, cols):
    t = doc.add_table(rows=rows, cols=cols)
    try:
        t.style = "Table Grid"
    except KeyError:
        pass
    return t


# ── Title ─────────────────────────────────────────────────────────────────────
h(0, "Policy on Day-to-Day Leadership of the Organisation")
p(HN)

# ── Document control ──────────────────────────────────────────────────────────
h(1, "Document control")

dc = tbl(6, 4)
for ri, (a, b, c, d) in enumerate([
    ("Document No.", "HCO/ROM/POL/04", "Version", "2.0"),
    ("Issue No.", "01", "Review due", "One year from implementation"),
    ("Date created", "________", "Date of implementation", "________"),
]):
    dc.cell(ri, 0).text = a
    dc.cell(ri, 1).text = b
    dc.cell(ri, 2).text = c
    dc.cell(ri, 3).text = d

for ri, (label, content) in enumerate([
    ("Prepared by",
     "Medical Superintendent Name: ________ Signature: ________"),
    ("Reviewed by",
     "Quality Coordinator Name: ________ Signature: ________"),
    ("Approved by",
     "Medical Superintendent Name: ________ Signature: ________"),
], start=3):
    dc.cell(ri, 0).text = label
    dc.cell(ri, 1).text = content
    dc.cell(ri, 1).merge(dc.cell(ri, 3))

p("A blank marked ________ must be completed before issue.")

# ── Statement of intent ───────────────────────────────────────────────────────
h(1, "Statement of intent")
p(f"{HN} is led by a qualified and experienced person who holds the "
  f"organisation's day-to-day accountability. That person's identity, "
  f"qualifications and legal compliance responsibilities are documented. Every "
  f"department and programme has an identified leader, and the organisation's "
  f"leader is formally reviewed for effectiveness every year.")

# ── 1. Purpose ────────────────────────────────────────────────────────────────
h(1, "1. Purpose")
p(f"This policy explains how {HN} appoints a person with the right "
  f"qualifications and experience to lead the organisation on a day-to-day "
  f"basis; keeps the hospital compliant with all applicable laws, regulations "
  f"and notifications; appoints leaders for each department and service; ensures "
  f"every programme and department has effective leadership; and reviews the "
  f"performance of the organisation's leader for effectiveness.")
p("This policy does not cover patient assessment, clinical care, medication "
  "management, patient rights, infection control, or quality and safety "
  "monitoring — those are covered in other hospital policies. The other "
  "Responsibilities of Management standards have their own policies too.")

# ── 2. Scope ──────────────────────────────────────────────────────────────────
h(1, "2. Scope")
p(f"This policy applies to those responsible for governance (as identified "
  f"under ROM.1), the Medical Superintendent as the person heading the "
  f"organisation, departmental leaders, the Quality Coordinator, and all staff "
  f"at {HN} who hold named leadership roles or are involved in the appointment "
  f"of department leaders.")

# ── 3. Policy standards ───────────────────────────────────────────────────────
h(1, "3. Policy standards")
p(f"{HN} appoints a qualified and experienced leader to head the organisation, "
  f"keeps the hospital legally compliant, ensures every department has an "
  f"identified leader, maintains effective leadership across all programmes and "
  f"departments, and reviews the organisation's leader annually.")
p("Staff follow the written guidance below, keep the records it requires, and "
  "raise concerns through the channels it describes.")

# ── 4. Non-negotiable rules ───────────────────────────────────────────────────
h(1, "4. Non-negotiable rules")
ln("Do not appoint a person to head the organisation without first verifying "
   "and recording their administrative qualification and experience "
   "— this check must happen before the appointment is confirmed, not after.")
ln("The hospital must maintain a current legislation register — the leader "
   "cannot plead ignorance of applicable requirements because the register "
   "was not kept up to date.")
ln("The legislation register must be actively maintained — the hospital must "
   "have a working mechanism for regularly picking up and incorporating "
   "amendments; a register that was created once and never reviewed since "
   "is not sufficient.")
ln("Every programme, service, site or department must have an identified leader "
   "on file — a gap must be escalated the same week it is found, not left open.")
ln("Each department leader must document the services their department provides "
   "— a verbal summary is not sufficient.")
ln("The organisation's leader must be formally reviewed for effectiveness at "
   "least once a year by those responsible for governance — a year without a "
   "review is a gap.")
ln("Staff who see any of these rules broken report it the same shift to the "
   "department in-charge or the Medical Superintendent.")

# ── 5. What we do ─────────────────────────────────────────────────────────────
h(1, "5. What we do")

h(2, "5.1 Appoint a qualified and experienced leader")
p(f"The person heading {HN} holds a qualification in hospital management or "
  f"administration and has administrative experience in a hospital. Those "
  f"responsible for governance verify both at the time of appointment and "
  f"record that check. When the role-holder changes, the verification is "
  f"repeated before the appointment is confirmed.")
p("The specific level of qualification and the minimum years of experience "
  "are for those responsible for governance to define. The Guidebook specifies "
  "that 'appropriate qualification' means qualification in hospital management "
  "or administration, and 'appropriate experience' means administrative "
  "experience in a hospital.")
p("Qualification certificates and a current CV are held on file.")

h(2, "5.2 Comply with applicable laws, regulations and notifications")
p(f"The leader of {HN} is responsible for compliance with all applicable "
  f"legislations, regulations and notifications, and takes active measures to "
  f"adhere to them. The hospital operates as a duly permitted legal entity "
  f"under all relevant registering authorities.")
p("The leader maintains a current applicable-legislation register specific to "
  "this hospital's legal form and the services it actually runs. The register "
  "is not a copy of every central Act — it is a working list of the "
  "requirements that apply to this hospital. A mechanism exists to regularly "
  "update the register when amendments are issued, and statutory documents are "
  "updated within the timelines the relevant laws and registration authorities "
  "require. Where the hospital conducts research, including clinical trials, "
  "that research is conducted in accordance with statutory norms.")
p("The form the update mechanism takes is for the hospital to determine "
  "— one possible approach is a tracker that logs each known requirement "
  "and flags pending updates.")

h(2, "5.3 Appoint department leaders")
p(f"The leader appoints, or actively participates in recruiting, the leaders "
  f"for each department and service that assist the day-to-day functioning of "
  f"{HN}. This applies at a minimum to the next level of the organisational "
  f"structure. Appointment files show the leader's involvement in each "
  f"decision.")
p("Each department leader documents the services their department provides and "
  "has an integrated approach to aligning those services with the rest of the "
  "hospital. Department leaders participate in the hospital's patient safety "
  "and quality improvement activities in line with institutional priorities.")
p("A department found without an identified leader is an open gap "
  "— this is escalated to the Medical Superintendent the same week.")

h(2, "5.4 Ensure effective leadership across all programmes and departments")
p(f"The leader ensures that every programme, service, site or department at "
  f"{HN} has effective leadership: a named person, an allocation of time for "
  f"the leadership duties, and a clear reporting line. The organisation "
  f"maintains a current leadership map recording these details.")
p("Department leaders must have relevant qualification and/or experience for "
  "the area they lead, and domain knowledge of that department. The "
  "organisation develops and uses metrics to assess the effectiveness of its "
  "leaders — what those metrics are is for the organisation to determine.")
p("When a leadership position becomes vacant, interim coverage is arranged and "
  "recorded until a substantive appointment is confirmed.")

h(2, "5.5 Review the organisation's leader for effectiveness")
p(f"The performance of the organisation's leader is reviewed for effectiveness "
  f"at least annually. The review is carried out by those responsible for "
  f"governance or the appointing authority.")
p("The form the review takes is for the organisation to determine "
  "— defined key result areas or a formal performance appraisal are both "
  "acceptable approaches. Where those responsible for governance and the person "
  "heading the organisation are the same individual, self-evaluation is an "
  "acceptable form of review.")
p("The review is dated, documented and filed. A year without a review is a gap.")

# ── 6. Governance and responsibility ──────────────────────────────────────────
h(1, "6. Governance and responsibility")
gt = tbl(6, 2)
for ri, (role, resp) in enumerate([
    ("Role", "Responsibility"),
    ("Medical Superintendent",
     "Accountable that this policy is followed and properly resourced."),
    ("Medical Superintendent",
     "Owns the legislation register, leadership map, and day-to-day records "
     "for this policy."),
    ("Those responsible for governance",
     "Verify the leader's qualifications at appointment; carry out or assign "
     "the annual performance review."),
    ("Quality Coordinator",
     "Audits this policy every quarter; keeps training records; supports "
     "governance packs and committee reviews this policy requires."),
    ("Departmental leaders",
     "Maintain documentation of their department's services; participate in "
     "quality improvement activities."),
]):
    gt.cell(ri, 0).text = role
    gt.cell(ri, 1).text = resp

# ── 7. Quality monitoring ─────────────────────────────────────────────────────
h(1, "7. Quality monitoring")
p("The Quality Coordinator audits this policy every quarter, checking a sample "
  "of records against each of the What-we-do steps.")
p("CORE objective elements show no critical gaps in the sample "
  "(ROM.4.b is the CORE element in this standard — see Level column, "
  "Section 11).")
p("If an audit finds a gap that is still open after 90 days, the hospital "
  "carries out a root-cause analysis and records the findings and corrective "
  "actions.")
p("This policy is reviewed every year, and sooner if the legislation register, "
  "the leadership structure, or the governance review process changes "
  "significantly.")

# ── 8. Training and staff acknowledgement ─────────────────────────────────────
h(1, "8. Training and staff acknowledgement")
p("Staff covered by this policy are trained when they join, and again every "
  "year after that. Training covers the What-we-do steps and the non-negotiable "
  "rules.")
p("I have read the Policy on Day-to-Day Leadership of the Organisation of "
  "«Hospital Name». I will follow the processes described.")
p("Name: ___________________________ Designation: ___________________________")
p("Department / floor: ____________________ Date: ____________")
p("Signature: ___________________________")
p("(One row per staff member. The Quality Coordinator holds signed "
  "acknowledgements with the induction record.)")

# ── 9. Distribution ───────────────────────────────────────────────────────────
h(1, "9. Distribution")
p("Medical Superintendent; those responsible for governance; Quality "
  "Coordinator; departmental leaders; staff covered by ROM.4")

# ── 10. Abbreviations ─────────────────────────────────────────────────────────
h(1, "10. Abbreviations")
at = tbl(6, 2)
for ri, (abbr, meaning) in enumerate([
    ("Abbreviation", "Meaning"),
    ("CORE", "Core objective element (NABH)"),
    ("HCO", "Hospital (Full Accreditation programme under NABH Hospitals "
             "6th Edition)"),
    ("NABH", "National Accreditation Board for Hospitals and Healthcare "
              "Providers"),
    ("OE", "Objective Element"),
    ("ROM", "Responsibilities of Management (NABH Hospitals 6th Edition "
             "chapter)"),
]):
    at.cell(ri, 0).text = abbr
    at.cell(ri, 1).text = meaning

# ── 11. Traceability ──────────────────────────────────────────────────────────
h(1, "11. Traceability to NABH HCO Full Accreditation 6th Edition ROM.4")
p("This table is an index. It is not how the policy is organised.")

tt = tbl(6, 5)
for ci, txt in enumerate(("OE", "Level", "Requirement",
                           "Where addressed", "Responsible")):
    tt.cell(0, ci).text = txt

for ri, (oe, level, req, where, who) in enumerate([
    ("ROM.4.a", "Commitment",
     "The person heading the organisation has requisite and appropriate "
     "administrative qualifications and experience.",
     "Section 3; 5.1", "Medical Superintendent"),
    ("ROM.4.b", "CORE",
     "The leader is responsible for and complies with the laid-down and "
     "applicable legislations, regulations and notifications.",
     "Section 3; 5.2", "Medical Superintendent"),
    ("ROM.4.c", "Commitment",
     "The leader appoints / participates in the recruitment of department "
     "leaders of the organisation who will assist in the day-to-day "
     "functioning of the organisation.",
     "Section 3; 5.3", "Medical Superintendent"),
    ("ROM.4.d", "Achievement",
     "The leader ensures that each organisational programme, service, site "
     "or department has effective leadership.",
     "Section 3; 5.4", "Medical Superintendent"),
    ("ROM.4.e", "Achievement",
     "The performance of the organisation's leader is reviewed for "
     "effectiveness.",
     "Section 3; 5.5", "Medical Superintendent"),
], start=1):
    tt.cell(ri, 0).text = oe
    tt.cell(ri, 1).text = level
    tt.cell(ri, 2).text = req
    tt.cell(ri, 3).text = where
    tt.cell(ri, 4).text = who

# ── 12. Required Records / Evidence Checklist ─────────────────────────────────
h(1, "12. Required Records / Evidence Checklist")
p("Records the hospital holds under this policy, listed by objective element.")

h(2, "ROM.4.a — Leader qualifications")
lb("Qualification certificates and CV of the person heading the organisation, "
   "on file.")
lb("Record of the governing body's qualification and experience check at "
   "appointment (or when the role-holder changes).")

h(2, "ROM.4.b — Statutory compliance")
lb("Current applicable-legislation register specific to this hospital's legal "
   "form and services.")
lb("Evidence of compliance against each item on the register.")
lb("Dated review or update record showing amendments were captured and "
   "statutory documents updated within required timelines.")

h(2, "ROM.4.c — Department leader appointment")
lb("Appointment file for each department leader showing the "
   "organisation-head's involvement in the decision.")
lb("Current list of department leaders with appointment dates.")
lb("Each department leader's documented service scope for their department.")
lb("Escalation record for any department found without an identified leader.")

h(2, "ROM.4.d — Effective departmental leadership")
lb("Leadership map naming a person, time allocation and reporting line for "
   "each programme, service, site or department.")
lb("Dated review of the leadership map.")
lb("Vacancy-coverage record for any leadership gap.")
lb("Effectiveness metrics the organisation uses for its leaders.")

h(2, "ROM.4.e — Leader performance review")
lb("Dated annual performance-review record for the organisation's leader.")
lb("Review criteria or key result areas used.")
lb("Filed outcome and any follow-up action.")

# ── 13. References ────────────────────────────────────────────────────────────
h(1, "13. References")
lb("National Accreditation Board for Hospitals and Healthcare Providers "
   "(NABH), Accreditation Standards for Hospitals, 6th Edition "
   "(January 2025) — Responsibilities of Management, standard ROM.4.")
lb("NABH Guidebook to Accreditation Standards for Hospitals, 6th Edition "
   "— ROM.4 interpretations.")
lb(f"Internal documents of {HN}: applicable-legislation register, "
   "department-leader appointment files, leadership map, leader "
   "performance-review records.")

# ── Disclaimer ────────────────────────────────────────────────────────────────
h(1, "Disclaimer")
p(f"This document is a template prepared for the guidance of {HN} and must be "
  f"reviewed, adapted and formally approved by {HN} before use. Every entry "
  f"marked ________ must be completed before the document is issued.")
p("The requirements in this document are accreditation requirements of the "
  "NABH Accreditation Standards for Hospitals, 6th Edition, not duties under "
  "a named Act of Parliament. This policy does not import the Consumer "
  "Protection Act, 2019, the Clinical Establishments Act, 2010, or the Mental "
  "Healthcare Act, 2017 as a checklist. Statutory duties that arise under other "
  "documents of «Hospital Name» remain those documents. "
  "«Hospital Name» is responsible for verifying any statutory duty "
  "that applies to it; this document does not constitute legal advice.")
p("«Hospital Name» remains responsible for verifying that it is "
  "current and consistent with the edition of the accreditation standard "
  "against which it is being assessed. The clinical and technical content "
  "reflects recognised national and international guidance current at the date "
  "of preparation.")
p("This document is not issued by, endorsed by, or affiliated with NABH, the "
  "World Health Organization, the National Centre for Disease Control, the Food "
  "Safety and Standards Authority of India, any Pollution Control Board, or "
  "any other body named in it. Wording is original; no text has been reproduced "
  "from the standards, rules or guidelines referenced.")

# ── Verify styles ─────────────────────────────────────────────────────────────
print("=== Paragraph style verification ===")
for i, para in enumerate(doc.paragraphs):
    print(f"{i:3d}  {para.style.name!r:32s}  {para.text[:55]!r}")

print(f"\nTotal paragraphs: {len(doc.paragraphs)}")

# ── Save ──────────────────────────────────────────────────────────────────────
doc.save(OUT)
print(f"\nSaved: {OUT}")
